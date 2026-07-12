"""Patch frozen groundtruth + eval literals to stay in sync with the russified
InSales (woocommerce-fork) seed.

Single source of truth = scripts/wc_relabel_map.py (FLAT_VALUE_MAP + atoms).
This walks every woo task's groundtruth_workspace/ and replaces English realia
STRING VALUES (categories, shipping zone/method titles, payment titles, tax names,
product-attribute names, shipping zones, customer full names, city/country) with
their RU value. NUMBERS / DATES / FORMULAS / identifiers / SKUs / emails /
status-slugs / column-header identifiers are NEVER touched.

Eval .py literals that hardcode a russified realia VALUE as a lookup key/label are
co-mapped via an explicit, audited edit table (EVAL_EDITS) — NOT a blind text sub.

Conservative by construction:
  * xlsx: header row (row 1) skipped (avoids translating column headers like
    'Brand'/'City'/'Status'/'Zone'); formula cells skipped; only whole-cell exact
    matches translated.
  * person-name mapping fires ONLY when BOTH first+last atoms are WC customer atoms
    (avoids corrupting foreign names e.g. canvas students 'Harry Taylor'); tasks
    with known mixed non-WC name pools are in NAME_MAP_DENY_TASKS.
  * json/csv: FIELD-AWARE (first_name->CUSTOMER_FIRST, city->CITY, ...) so the
    'Charlotte' first_name vs 'Charlotte' city collision is resolved correctly.
  * txt: only whole ' - '-delimited segments equal to a realia key are replaced.
  * docx: run-level whole-text exact match only (never substring -> no SKU/brand
    corruption).

Ambiguous strings are SKIPPED and logged (issues[]) rather than risk corrupting GT.
Does NOT git commit. Run:  python3 scripts/wc_patch_groundtruth.py
"""
import argparse
import csv
import glob
import importlib.util
import io
import json
import os
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TASKS_DIR = os.path.join(ROOT, "tasks", "finalpool")
MAP_PATH = os.path.join(ROOT, "scripts", "wc_relabel_map.py")

# ---- load the single-source-of-truth relabel map ----------------------------
_spec = importlib.util.spec_from_file_location("wc_relabel_map", MAP_PATH)
M = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(M)

FLAT = M.FLAT_VALUE_MAP            # all whole-value realia (incl. city/country)
CFIRST = M.CUSTOMER_FIRST
CLAST = M.CUSTOMER_LAST
CITY = M.CITY
COUNTRY = M.COUNTRY

# Non-person realia subset (excludes city/country so a generic category/zone/title
# field never accidentally turns a city into a person-name target, etc.)
NONPERSON = {}
for _name in ("CATEGORIES", "TAGS", "SHIPPING_ZONES", "SHIPPING_METHOD_TITLES",
              "PAYMENT_TITLES", "PAYMENT_METHOD_TITLES", "TAX_NAMES",
              "PRODUCT_ATTRIBUTES", "COUPON_DESC"):
    NONPERSON.update(getattr(M, _name))

# Tasks whose GT mixes non-WC person names (e.g. canvas students) with WC data.
# Person-name mapping is DISABLED for these (both-atom collisions corrupt foreign
# names). Their WC realia (categories) are handled in EVAL_EDITS / DB live.
NAME_MAP_DENY_TASKS = {"terminal-canvas-insales-excel-ppt-email"}

issues = []          # ambiguities skipped, for human fix-pass
gt_files_changed = []
eval_files_changed = []


def log_issue(task, where, msg):
    issues.append(f"[{task}] {where}: {msg}")


# ---------------------------------------------------------------------------
# value mappers
# ---------------------------------------------------------------------------
def map_full_name_strict(s):
    """Return RU 'First Last' ONLY if BOTH atoms are WC customer atoms; else None."""
    parts = s.split()
    if len(parts) != 2:
        return None
    f, l = parts
    rf, rl = CFIRST.get(f), CLAST.get(l)
    if rf is not None and rl is not None:
        return f"{rf} {rl}"
    return None


def map_cell_value(s, allow_names=True):
    """Whole-cell map for xlsx/docx data cells. Exact FLAT match first, then a
    strict (both-atom) full-name. Returns RU string or None (leave as-is)."""
    if s in FLAT:
        return FLAT[s]
    if allow_names:
        return map_full_name_strict(s)
    return None


def map_json_field(key, value):
    """Field-aware map for json/csv. Returns RU string or None."""
    if not isinstance(value, str):
        return None
    s = value.strip()
    if not s:
        return None
    kl = (key or "").lower()
    if kl in ("first_name", "firstname", "billing_first_name", "shipping_first_name"):
        return CFIRST.get(s)
    if kl in ("last_name", "lastname", "billing_last_name", "shipping_last_name"):
        return CLAST.get(s)
    if kl in ("city", "billing_city", "shipping_city"):
        return CITY.get(s)
    if kl in ("country", "billing_country", "shipping_country"):
        return COUNTRY.get(s)
    # category / zone / shipping / payment / tax / attribute name fields
    if kl in ("category", "category_name", "categories", "zone", "zone_name",
              "title", "method_title", "shipping_method", "payment_method_title",
              "payment_title", "tax_name", "attribute", "attribute_name"):
        return NONPERSON.get(s)
    # generic 'name' field: only non-person realia (avoid touching product names,
    # reviewer names, identifiers). Person full names handled by *_name fields.
    if kl == "name":
        return NONPERSON.get(s)
    return None


# ---------------------------------------------------------------------------
# per-format patchers
# ---------------------------------------------------------------------------
def patch_xlsx(path, task):
    import openpyxl
    try:
        wb = openpyxl.load_workbook(path, data_only=False)
    except Exception as e:
        log_issue(task, os.path.basename(path), f"xlsx unreadable: {e}")
        return 0
    allow_names = task not in NAME_MAP_DENY_TASKS
    n = 0
    for ws in wb.worksheets:
        # skip header row (row 1): column-header identifiers must not be translated
        for row in ws.iter_rows(min_row=2):
            for c in row:
                v = c.value
                if not isinstance(v, str):
                    continue
                s = v.strip()
                if not s or s.startswith("="):  # formula / empty
                    continue
                ru = map_cell_value(s, allow_names=allow_names)
                if ru and ru != v:
                    c.value = ru
                    n += 1
    if n:
        wb.save(path)
    if task in NAME_MAP_DENY_TASKS:
        log_issue(task, os.path.basename(path),
                  "person-name mapping DISABLED (mixed non-WC name pool); "
                  "realia categories handled via eval/DB")
    return n


def patch_docx(path, task):
    from docx import Document
    try:
        doc = Document(path)
    except Exception as e:
        log_issue(task, os.path.basename(path), f"docx unreadable: {e}")
        return 0
    allow_names = task not in NAME_MAP_DENY_TASKS
    n = 0

    def patch_runs(paras):
        nonlocal n
        for p in paras:
            for r in p.runs:
                t = r.text
                if not t:
                    continue
                s = t.strip()
                ru = map_cell_value(s, allow_names=allow_names)
                if ru and ru != t:
                    # preserve surrounding whitespace of the run
                    r.text = t.replace(s, ru) if s in t else ru
                    n += 1

    patch_runs(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                patch_runs(cell.paragraphs)
    if n:
        doc.save(path)
    return n


def _walk_json(o, key=None):
    """Yield (container, accessor, key, value) so caller can mutate in place."""
    if isinstance(o, dict):
        for k, v in o.items():
            if isinstance(v, (dict, list)):
                yield from _walk_json(v, k)
            else:
                yield (o, k, k, v)
    elif isinstance(o, list):
        for i, v in enumerate(o):
            if isinstance(v, (dict, list)):
                yield from _walk_json(v, key)
            else:
                yield (o, i, key, v)


def patch_json(path, task):
    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        log_issue(task, os.path.basename(path), f"json unreadable: {e}")
        return 0
    n = 0
    skipped_names = set()
    for container, accessor, key, value in list(_walk_json(data)):
        if not isinstance(value, str):
            continue
        ru = map_json_field(key, value)
        if ru and ru != value:
            container[accessor] = ru
            n += 1
            continue
        # flag a full person-name living in a non-WC name field (e.g. student_name)
        kl = (key or "").lower()
        if "name" in kl and kl not in (
            "first_name", "last_name", "category_name", "zone_name", "method_title",
            "tax_name", "attribute_name", "payment_method_title", "name",
        ):
            if map_full_name_strict(value.strip()):
                skipped_names.add(f"{kl}={value!r}")
    for s in sorted(skipped_names):
        log_issue(task, os.path.basename(path),
                  f"SKIPPED person-name in non-WC field {s} (not russified seed data)")
    if n:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    return n


def patch_csv(path, task):
    try:
        with open(path, newline="", encoding="utf-8") as f:
            rows = list(csv.reader(f))
    except Exception as e:
        log_issue(task, os.path.basename(path), f"csv unreadable: {e}")
        return 0
    if not rows:
        return 0
    header = rows[0]
    n = 0
    for r in rows[1:]:
        for ci, val in enumerate(r):
            col = header[ci] if ci < len(header) else ""
            ru = map_json_field(col, val)
            if ru and ru != val:
                r[ci] = ru
                n += 1
    if n:
        with open(path, "w", newline="", encoding="utf-8") as f:
            csv.writer(f).writerows(rows)
    return n


def patch_txt(path, task):
    try:
        with open(path, encoding="utf-8") as f:
            lines = f.readlines()
    except Exception as e:
        log_issue(task, os.path.basename(path), f"txt unreadable: {e}")
        return 0
    n = 0
    out = []
    for line in lines:
        nl = "\n" if line.endswith("\n") else ""
        body = line[:-1] if nl else line
        # split on ' - ' (the observed delimiter); replace only whole segments
        if " - " in body:
            segs = body.split(" - ")
            changed = False
            for i, seg in enumerate(segs):
                s = seg.strip()
                if s in NONPERSON:  # category/shipping realia only (no names/cities)
                    segs[i] = seg.replace(s, NONPERSON[s])
                    changed = True
            if changed:
                n += 1
            out.append(" - ".join(segs) + nl)
        else:
            s = body.strip()
            if s in NONPERSON:
                out.append(body.replace(s, NONPERSON[s]) + nl)
                n += 1
            else:
                out.append(line)
    if n:
        with open(path, "w", encoding="utf-8") as f:
            f.writelines(out)
    return n


PATCHERS = {
    ".xlsx": patch_xlsx, ".docx": patch_docx, ".json": patch_json,
    ".csv": patch_csv, ".txt": patch_txt,
}


# ---------------------------------------------------------------------------
# eval .py literal co-mapping — explicit, audited edits (old -> new, exact match)
# Only literals that are a russified realia VALUE used as a lookup key/label.
# Region labels (North America/Europe) and tiers (Gold/Silver/Bronze/VIP) are
# DERIVED, not seed values -> intentionally LEFT (noted in issues).
# ---------------------------------------------------------------------------
EVAL_EDITS = {
    "insales-customer-loyalty-word-gsheet": [
        # top spender full name + last-name substring fallback
        ('"william gonzalez" in full_text or "gonzalez" in full_text',
         '"василий гончаров" in full_text or "гончаров" in full_text'),
        ("Mentions 'william gonzalez' (top spender)",
         "Mentions 'василий гончаров' (top spender)"),
    ],
    "insales-inventory-gold-repricing": [
        # CATEGORY_MULTIPLIERS keys = product_categories.name (russified DB lookup)
        ("    'Electronics': 1.05,", "    'Электроника': 1.05,"),
        ("    'Cameras': 1.08,", "    'Камеры': 1.08,"),
        ("    'Audio': 0.97,", "    'Аудио': 0.97,"),
        ("    'TV & Home Theater': 1.10,", "    'ТВ и домашний кинотеатр': 1.10,"),
        ("    'Home Appliances': 1.02,", "    'Бытовая техника': 1.02,"),
        ("    'Watches': 0.95,", "    'Часы': 0.95,"),
        ("    'Headphones': 1.03,", "    'Наушники': 1.03,"),
        ("    'Speakers': 1.00,", "    'Колонки': 1.00,"),
    ],
    "insales-shipping-performance-gsheet-word-email": [
        # EXPECTED_ZONES names = shipping method titles (russified)
        ('{"name": "standard shipping", "orders": 64, "revenue": 25369.91},',
         '{"name": "стандартная доставка", "orders": 64, "revenue": 25369.91},'),
        ('{"name": "free shipping", "orders": 34, "revenue": 12734.76},',
         '{"name": "бесплатная доставка", "orders": 34, "revenue": 12734.76},'),
        ('{"name": "express shipping", "orders": 30, "revenue": 12673.21},',
         '{"name": "экспресс-доставка", "orders": 30, "revenue": 12673.21},'),
        # gsheet substring checks against russified agent output
        ('"standard" in all_lower, "Standard not found")',
         '"стандартная" in all_lower, "Standard not found")'),
        ('"express" in all_lower, "Express not found")',
         '"экспресс" in all_lower, "Express not found")'),
        # word-doc table + body substring checks
        ('"standard" in table_text or "express" in table_text,',
         '"стандартная" in table_text or "экспресс" in table_text,'),
        ('"standard" in full_text, "Standard shipping not mentioned")',
         '"стандартная" in full_text, "Standard shipping not mentioned")'),
        ('"express" in full_text, "Express shipping not mentioned")',
         '"экспресс" in full_text, "Express shipping not mentioned")'),
    ],
    "insales-vip-customer-gsheet-gcal-email": [
        # GOLD_CUSTOMERS display names (matching is by email, kept; names cosmetic
        # but russified for consistency with russified seed)
        ('{"name": "Scarlett Wright", "email": "scarlett.wright@x.dummyjson.com"},',
         '{"name": "Скарлетт Рыжов", "email": "scarlett.wright@x.dummyjson.com"},'),
        ('{"name": "Ethan Martinez", "email": "ethan.martinez@x.dummyjson.com"},',
         '{"name": "Артём Мартынов", "email": "ethan.martinez@x.dummyjson.com"},'),
        ('{"name": "Olivia Wilson", "email": "olivia.wilson@x.dummyjson.com"},',
         '{"name": "Оливия Виноградов", "email": "olivia.wilson@x.dummyjson.com"},'),
    ],
    "insales-yf-commodity-pricing-impact": [
        # category substring lookups against russified agent gsheet
        ('"watch" in str(r[0]).lower()]', '"час" in str(r[0]).lower()]'),
        ('"electronics" in str(r[0]).lower()]', '"электроник" in str(r[0]).lower()]'),
    ],
    "terminal-insales-yf-commodity-gsheet-word": [
        # EXPECTED_CATEGORIES fallback lists (matched vs russified agent cells)
        ('return cats if cats else ["Audio", "Cameras", "Electronics", "Home Appliances",',
         'return cats if cats else ["Аудио", "Камеры", "Электроника", "Бытовая техника",'),
        ('                                  "TV & Home Theater", "Watches"]',
         '                                  "ТВ и домашний кинотеатр", "Часы"]'),
        ('        return ["Audio", "Cameras", "Electronics", "Home Appliances",',
         '        return ["Аудио", "Камеры", "Электроника", "Бытовая техника",'),
        ('                "TV & Home Theater", "Watches"]',
         '                "ТВ и домашний кинотеатр", "Часы"]'),
        # SENSITIVITY_MAP keys (currently UNUSED dead code; russified for consistency)
        ('SENSITIVITY_MAP = {"Watches": 40, "Electronics": 15, "Audio": 10,',
         'SENSITIVITY_MAP = {"Часы": 40, "Электроника": 15, "Аудио": 10,'),
        ('                   "Cameras": 10, "TV & Home Theater": 5, "Home Appliances": 8}',
         '                   "Камеры": 10, "ТВ и домашний кинотеатр": 5, "Бытовая техника": 8}'),
    ],
    "insales-product-launch-dashboard": [
        # MARKET_DATA keys = category labels matched vs russified agent 'Market
        # Comparison' sheet (growth/size/trend VALUES are external benchmarks: KEEP)
        ('    "TV & Home Theater": {"growth": 12.5, "size": 48200, "trend": "Strong Growth"},',
         '    "ТВ и домашний кинотеатр": {"growth": 12.5, "size": 48200, "trend": "Strong Growth"},'),
        ('    "Electronics": {"growth": 8.3, "size": 125400, "trend": "Moderate Growth"},',
         '    "Электроника": {"growth": 8.3, "size": 125400, "trend": "Moderate Growth"},'),
        ('    "Audio": {"growth": 15.2, "size": 22800, "trend": "Strong Growth"},',
         '    "Аудио": {"growth": 15.2, "size": 22800, "trend": "Strong Growth"},'),
        ('    "Cameras": {"growth": 3.1, "size": 8900, "trend": "Slow Growth"},',
         '    "Камеры": {"growth": 3.1, "size": 8900, "trend": "Slow Growth"},'),
        ('    "Watches": {"growth": 6.7, "size": 15600, "trend": "Moderate Growth"},',
         '    "Часы": {"growth": 6.7, "size": 15600, "trend": "Moderate Growth"},'),
        ('    "Home Appliances": {"growth": -2.4, "size": 35100, "trend": "Declining"},',
         '    "Бытовая техника": {"growth": -2.4, "size": 35100, "trend": "Declining"},'),
        # top opportunity = Audio (russified substring + expected dict key)
        ('"audio" in top_cat.lower(),', '"ауди" in top_cat.lower(),'),
        ('expected.get("Audio", {}).get("avg_rating", 4.54)',
         'expected.get("Аудио", {}).get("avg_rating", 4.54)'),
        ('record("Top opportunity is Audio",', 'record("Top opportunity is Аудио",'),
        ('record(f"Audio Opportunity_Score ~{exp_audio_score}", ok,',
         'record(f"Аудио Opportunity_Score ~{exp_audio_score}", ok,'),
    ],
    "insales-product-pricing-excel-gform": [
        # EXPECTED_CATEGORIES (currently UNUSED dead code; russified for consistency)
        ('EXPECTED_CATEGORIES = ["Audio", "TV & Home Theater", "Speakers", "Electronics", "Cameras", "Headphones"]',
         'EXPECTED_CATEGORIES = ["Аудио", "ТВ и домашний кинотеатр", "Колонки", "Электроника", "Камеры", "Наушники"]'),
    ],
    "insales-sales-category-excel-email": [
        ('"headphones" in body_lower,', '"наушники" in body_lower,'),
        ('"tv" in body_lower and "theater" in body_lower or "tv & home theater" in body_lower,',
         '"тв" in body_lower and "кинотеатр" in body_lower or "тв и домашний кинотеатр" in body_lower,'),
    ],
    "terminal-canvas-insales-excel-ppt-email": [
        # WC category filter against russified wc.product_categories.name
        ("WHERE lower(name) IN ('electronics', 'cameras')",
         "WHERE lower(name) IN ('электроника', 'камеры')"),
        # slide-title fallback substring (agent title now russified)
        ('"purchase" in title.text.lower() or "electronics" in title.text.lower()',
         '"purchase" in title.text.lower() or "электроник" in title.text.lower()'),
    ],
    "insales-tax-compliance-review": [
        # tax_rates.country russified US->Россия (WHERE predicate, else 0 rows)
        ("WHERE country = 'US' AND state != '' AND class = 'standard'",
         "WHERE country = 'Россия' AND state != '' AND class = 'standard'"),
    ],
}


def apply_eval_edits(task):
    """Apply audited literal edits to a task's eval main.py. Returns files changed."""
    edits = EVAL_EDITS.get(task)
    if not edits:
        return []
    py = os.path.join(TASKS_DIR, task, "evaluation", "main.py")
    if not os.path.isfile(py):
        log_issue(task, "evaluation/main.py", "eval file missing; edits skipped")
        return []
    with open(py, encoding="utf-8") as f:
        txt = f.read()
    orig = txt
    for old, new in edits:
        if old not in txt:
            log_issue(task, "evaluation/main.py",
                      f"eval literal not found (manual review): {old[:80]!r}")
            continue
        if txt.count(old) > 1:
            log_issue(task, "evaluation/main.py",
                      f"eval literal AMBIGUOUS ({txt.count(old)}x), skipped: {old[:60]!r}")
            continue
        txt = txt.replace(old, new)
    if txt != orig:
        with open(py, "w", encoding="utf-8") as f:
            f.write(txt)
        return [py]
    return []


# ---------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args()

    with open("/tmp/woo_tasks.txt") as f:
        tasks = [l.strip() for l in f if l.strip()]

    tasks_scanned = 0
    per_task = {}
    for task in sorted(tasks):
        tasks_scanned += 1
        tdir = os.path.join(TASKS_DIR, task)
        gt = os.path.join(tdir, "groundtruth_workspace")
        changed_here = []
        if os.path.isdir(gt):
            for path in sorted(glob.glob(os.path.join(gt, "**", "*"), recursive=True)):
                if not os.path.isfile(path):
                    continue
                ext = os.path.splitext(path)[1].lower()
                fn = PATCHERS.get(ext)
                if not fn:
                    continue
                if args.dry_run:
                    continue
                n = fn(path, task)
                if n:
                    gt_files_changed.append(path)
                    changed_here.append(f"{os.path.relpath(path, tdir)} ({n} cells)")
        # eval literal co-map
        if not args.dry_run:
            ev = apply_eval_edits(task)
            for p in ev:
                eval_files_changed.append(p)
                changed_here.append(os.path.relpath(p, tdir))
        if changed_here:
            per_task[task] = changed_here

    print("=== PER-TASK CHANGES ===")
    for task in sorted(per_task):
        print(f"\n{task}:")
        for c in per_task[task]:
            print(f"   {c}")
    print("\n=== ISSUES / SKIPPED (human fix-pass) ===")
    for i in issues:
        print(f"   {i}")
    print("\n=== COUNTS ===")
    print(f"   tasks_scanned    : {tasks_scanned}")
    print(f"   gt_files_changed : {len(gt_files_changed)}")
    print(f"   eval_files_changed: {len(eval_files_changed)}")


if __name__ == "__main__":
    main()
