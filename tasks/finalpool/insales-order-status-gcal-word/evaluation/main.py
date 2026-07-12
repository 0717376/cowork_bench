"""Evaluation for insales-order-status-gcal-word."""
import argparse
import os
import sys
import psycopg2
import openpyxl


DB = {"host": os.environ.get("PGHOST", "localhost"), "port": 5432, "dbname": "cowork_gym", "user": "eigent", "password": "camel"}


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def get_order_status_data():
    # Hard-fail on DB error: no silent hardcoded fallback (stale numbers must not pass).
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("SELECT status, count(*), sum(total) FROM wc.orders GROUP BY status ORDER BY count(*) DESC")
    rows = cur.fetchall()
    cur.close()
    conn.close()
    total_orders = sum(r[1] for r in rows)
    total_rev = round(sum(float(r[2]) for r in rows), 2)
    completed = next((r for r in rows if r[0] == 'completed'), None)
    completed_count = completed[1] if completed else 0
    completion_rate = round(completed_count / total_orders * 100, 2) if total_orders > 0 else 0
    most_common = rows[0][0] if rows else "completed"
    refunded = next((r for r in rows if r[0] == 'refunded'), None)
    refund_count = refunded[1] if refunded else 0
    return {
        "statuses": rows,
        "total_orders": total_orders,
        "total_rev": total_rev,
        "completed_count": completed_count,
        "completed_total": round(float(completed[2]), 2) if completed else 0,
        "completion_rate": completion_rate,
        "most_common": most_common,
        "refund_count": refund_count,
    }


def check_gcal_event():
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute(
            "SELECT count(*) FROM gcal.events WHERE "
            "LOWER(summary) LIKE '%operations%' OR LOWER(summary) LIKE '%order%' "
            "OR LOWER(summary) LIKE '%операц%' OR LOWER(summary) LIKE '%заказ%' "
            "OR LOWER(summary) LIKE '%статус%' OR LOWER(summary) LIKE '%обзор%'"
        )
        cnt = cur.fetchone()[0]
        cur.close()
        conn.close()
        return cnt >= 1
    except Exception:
        return False


def check_email_sent():
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute(
            "SELECT count(*) FROM email.messages WHERE LOWER(to_addr::text) LIKE '%ops%' AND ("
            "LOWER(subject) LIKE '%order%' OR LOWER(subject) LIKE '%status%' "
            "OR LOWER(subject) LIKE '%report%' OR LOWER(subject) LIKE '%заказ%' "
            "OR LOWER(subject) LIKE '%статус%' OR LOWER(subject) LIKE '%отчёт%' "
            "OR LOWER(subject) LIKE '%отчет%')"
        )
        cnt = cur.fetchone()[0]
        cur.close()
        conn.close()
        return cnt >= 1
    except Exception:
        return False


def check_word_file(agent_workspace):
    word_path = os.path.join(agent_workspace, "Operations_Summary.docx")
    if not os.path.exists(word_path):
        return False, "Operations_Summary.docx not found"
    try:
        from docx import Document
        doc = Document(word_path)
    except Exception as e:
        # Real failure to open/parse the doc is a FAIL, not a silent pass.
        return False, f"Could not open Operations_Summary.docx: {e}"
    full_text = " ".join(p.text for p in doc.paragraphs)
    full_text += " " + " ".join(cell.text for t in doc.tables for row in t.rows for cell in row.cells)
    ft = full_text.lower()
    keywords = [
        "order status", "operations", "completed", "status",
        "статус заказ", "статус", "заказ", "операц", "завершён", "завершен", "отчёт", "отчет",
    ]
    if any(k in ft for k in keywords):
        # Also require a table actually exists (the deliverable is a status breakdown table).
        if not doc.tables:
            return False, "Word doc has no table (expected status breakdown table)"
        return True, ""
    return False, "Word doc does not contain expected content about order status"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    agent_file = os.path.join(args.agent_workspace, "Order_Status_Report.xlsx")
    gt_file = os.path.join(gt_dir, "Order_Status_Report.xlsx")

    if not os.path.exists(agent_file):
        print(f"FAIL: Agent output not found: {agent_file}")
        sys.exit(1)
    if not os.path.exists(gt_file):
        print(f"FAIL: Groundtruth not found: {gt_file}")
        sys.exit(1)

    all_errors = []
    critical_errors = []

    # Hard-fail (no silent fallback): a broken DB must not let stale numbers pass.
    try:
        order_data = get_order_status_data()
    except Exception as e:
        print(f"FAIL: Could not query live order data from wc.orders: {e}")
        sys.exit(1)

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    # Check Status Breakdown sheet
    print("  Checking Status Breakdown sheet...")
    a_rows = load_sheet_rows(agent_wb, "Status Breakdown")
    g_rows = load_sheet_rows(gt_wb, "Status Breakdown")
    if a_rows is None:
        all_errors.append("Sheet 'Status Breakdown' not found in agent output")
    else:
        data_rows = [r for r in a_rows[1:] if r and any(c is not None for c in r)]
        if len(data_rows) < 5:
            all_errors.append(f"Status Breakdown has {len(data_rows)} rows, expected >= 5")
        else:
            print(f"    PASS ({len(data_rows)} data rows)")

        # Build lookup by status
        a_lookup = {str(r[0]).strip().lower(): r for r in data_rows if r and r[0] is not None}

        # Check completed status (CRITICAL: live DB count + value)
        completed_row = a_lookup.get("completed")
        if completed_row is None:
            msg = "'completed' status not found in Status Breakdown"
            all_errors.append(msg)
            critical_errors.append(msg)
        else:
            errors = []
            if len(completed_row) > 1:
                if not num_close(completed_row[1], order_data["completed_count"], 0):
                    errors.append(f"completed.Order_Count: {completed_row[1]} vs {order_data['completed_count']}")
            if len(completed_row) > 2:
                if not num_close(completed_row[2], order_data["completed_total"], 1.0):
                    errors.append(f"completed.Total_Value: {completed_row[2]} vs {order_data['completed_total']} (tol=1.0)")
            if errors:
                all_errors.extend(errors)
                critical_errors.extend(errors)
                for e in errors:
                    print(f"    ERROR: {e}")
            else:
                print("    completed row PASS")

    # Check Summary sheet
    print("  Checking Summary sheet...")
    a_rows = load_sheet_rows(agent_wb, "Summary")
    if a_rows is None:
        all_errors.append("Sheet 'Summary' not found in agent output")
    else:
        a_data = {str(r[0]).strip().lower(): r[1] for r in a_rows[1:] if r and r[0] is not None}
        errors = []          # non-critical Summary errors
        crit = []            # CRITICAL Summary errors (live-DB semantic values)

        total_orders_val = a_data.get("total_orders")
        if total_orders_val is None:
            crit.append("Missing metric: Total_Orders")
        elif not num_close(total_orders_val, order_data["total_orders"], 0):
            crit.append(f"Total_Orders: {total_orders_val} vs {order_data['total_orders']}")

        total_rev_val = a_data.get("total_revenue")
        if total_rev_val is None:
            crit.append("Missing metric: Total_Revenue")
        elif not num_close(total_rev_val, order_data["total_rev"], 2.0):
            crit.append(f"Total_Revenue: {total_rev_val} vs {order_data['total_rev']} (tol=2.0)")

        completion_val = a_data.get("completion_rate_pct")
        if completion_val is None:
            crit.append("Missing metric: Completion_Rate_Pct")
        elif not num_close(completion_val, order_data["completion_rate"], 0.5):
            crit.append(f"Completion_Rate_Pct: {completion_val} vs {order_data['completion_rate']} (tol=0.5)")

        most_common_val = a_data.get("most_common_status")
        if most_common_val is None:
            crit.append("Missing metric: Most_Common_Status")
        elif str(most_common_val).strip().lower() != order_data["most_common"].lower():
            crit.append(f"Most_Common_Status: {most_common_val} vs {order_data['most_common']}")

        refund_val = a_data.get("refund_count")
        if refund_val is None:
            crit.append("Missing metric: Refund_Count")
        elif not num_close(refund_val, order_data["refund_count"], 0):
            crit.append(f"Refund_Count: {refund_val} vs {order_data['refund_count']}")

        if crit:
            all_errors.extend(crit)
            critical_errors.extend(crit)
        if errors:
            all_errors.extend(errors)
        for e in (crit + errors)[:5]:
            print(f"    ERROR: {e}")
        if not crit and not errors:
            print("    PASS")

    # Check GCal event
    print("  Checking GCal event...")
    if check_gcal_event():
        print("    PASS")
    else:
        all_errors.append("Expected calendar event with 'Operations' or 'Order' in title, not found")

    # Check email sent
    print("  Checking email to ops.team...")
    if check_email_sent():
        print("    PASS")
    else:
        all_errors.append("Email to ops.team@company.com with 'order' or 'status' subject not found")

    # Check Word document
    print("  Checking Word document...")
    ok, detail = check_word_file(args.agent_workspace)
    if ok:
        print("    PASS")
    else:
        all_errors.append(detail)

    # CRITICAL gate: any critical semantic failure => FAIL regardless of accuracy.
    if critical_errors:
        print(f"\n=== CRITICAL FAIL ({len(critical_errors)} critical errors) ===")
        for e in critical_errors[:10]:
            print(f"  CRITICAL: {e}")
        sys.exit(1)

    # Accuracy gate: ~7 weighted check groups; require >= 70% with no critical fail.
    # Check groups: Status Breakdown rows, completed row, Summary metrics,
    # GCal event, email sent, Word doc, (and the >=5-rows structural check).
    TOTAL_CHECKS = 7
    failed = len(all_errors)
    passed = max(0, TOTAL_CHECKS - failed)
    accuracy = passed / TOTAL_CHECKS * 100

    print(f"\n=== ACCURACY: {accuracy:.1f}% ({passed}/{TOTAL_CHECKS} check groups) ===")
    if all_errors:
        for e in all_errors[:10]:
            print(f"  {e}")

    if accuracy >= 70:
        print("=== RESULT: PASS ===")
        sys.exit(0)
    else:
        print(f"=== RESULT: FAIL (accuracy {accuracy:.1f}% < 70%) ===")
        sys.exit(1)


if __name__ == "__main__":
    main()
