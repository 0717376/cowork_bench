"""Evaluation for terminal-arxiv-canvas-gsheet-word-gcal."""
import argparse
import json
import os
import sys
from datetime import datetime, timedelta

import openpyxl
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILURES = []


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRITICAL]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        print(f"  [FAIL]{' [CRITICAL]' if critical else ''} {name}: {str(detail)[:200]}")
        if critical:
            CRITICAL_FAILURES.append(name)


def parse_dt(value):
    """Parse a gcal start_datetime stored as string into a naive datetime."""
    if value is None:
        return None
    s = str(value)
    # Strip timezone offset like -05:00 / +00:00 if present
    for fmt in ("%Y-%m-%dT%H:%M:%S%z", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S%z", "%Y-%m-%d %H:%M:%S"):
        try:
            dt = datetime.strptime(s, fmt)
            return dt.replace(tzinfo=None)
        except Exception:
            continue
    # Last resort: try fromisoformat
    try:
        dt = datetime.fromisoformat(s)
        return dt.replace(tzinfo=None)
    except Exception:
        return None


def num_close(a, b, tol=2.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except Exception:
        return False


# ---------- Word Document ----------
def check_word(ws):
    print("\n=== Checking Word Document ===")
    fpath = os.path.join(ws, "Curriculum_Modernization_Proposal.docx")
    if not os.path.isfile(fpath):
        check("Word document exists", False, f"Not found: {fpath}")
        return
    check("Word document exists", True)

    from docx import Document
    doc = Document(fpath)
    full_text = " ".join(p.text for p in doc.paragraphs).lower()

    check("Title mentions curriculum modernization",
          "curriculum modernization proposal" in full_text or "curriculum" in full_text and "modernization" in full_text)
    check("Has executive summary section", "executive summary" in full_text)
    check("Has current state analysis", "current state" in full_text or "current" in full_text and "analysis" in full_text)
    check("Has research landscape section", "research landscape" in full_text or "research" in full_text and "landscape" in full_text)
    check("Has identified gaps section", "gap" in full_text and "identified" in full_text or "gaps" in full_text)
    check("Has proposed updates section", "proposed" in full_text and ("update" in full_text or "recommendation" in full_text))
    check("Has implementation timeline", "implementation" in full_text or "timeline" in full_text or "phase" in full_text)
    check("Mentions applied analytics", "applied analytics" in full_text or "analytics" in full_text and "algorithms" in full_text)
    check("Mentions at least 2 paper titles",
          sum(1 for t in ["deep learning", "graph neural", "reinforcement learning", "analytics methods", "benchmarking"]
              if t in full_text) >= 2,
          "Expected references to research papers")
    check("At least 3 recommendations",
          full_text.count("recommendation") >= 3 or full_text.count("propose") >= 2 or
          (full_text.count("recommendation 1") >= 1 and full_text.count("recommendation 3") >= 1) or
          full_text.count("рекомендац") >= 3 or full_text.count("предлага") >= 2)

    # CRITICAL: proposal must cite >=2 specific RELEVANT paper titles AND no noise papers.
    relevant_markers = ["deep learning", "graph neural", "reinforcement learning",
                        "analytics methods", "benchmarking", "combinatorial optimization"]
    noise_markers = ["marine biodiversity", "paleolithic cave art", "arctic ecosystems", "uranium-thorium"]
    n_relevant = sum(1 for t in relevant_markers if t in full_text)
    has_noise = any(nt in full_text for nt in noise_markers)
    check("Proposal cites >=2 relevant paper titles and no noise papers",
          n_relevant >= 2 and not has_noise,
          f"relevant={n_relevant}, noise_present={has_noise}",
          critical=True)


# ---------- Google Sheet ----------
def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        # Check spreadsheet exists
        cur.execute("SELECT id, title FROM gsheet.spreadsheets WHERE title ILIKE '%curriculum%gap%'")
        rows = cur.fetchall()
        check("Spreadsheet 'Curriculum Gap Analysis' exists", len(rows) >= 1,
              f"Found {len(rows)} matching spreadsheets")
        if not rows:
            check("CRITICAL: Spreadsheet 'Curriculum Gap Analysis' with all three sheets",
                  False, "Spreadsheet not found", critical=True)
            cur.close()
            conn.close()
            return

        ss_id = rows[0][0]

        # Check sheets
        cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (ss_id,))
        sheets = cur.fetchall()
        sheet_names = [s[1].lower() for s in sheets]
        sheet_map = {s[1].lower(): s[0] for s in sheets}

        has_current = any("current" in n and "topic" in n for n in sheet_names)
        has_research = any("research" in n and "topic" in n for n in sheet_names)
        has_gap = any("gap" in n and "analysis" in n for n in sheet_names)

        check("Has Current_Topics sheet", has_current, f"Sheets: {sheet_names}")
        check("Has Research_Topics sheet", has_research, f"Sheets: {sheet_names}")
        check("Has Gap_Analysis sheet", has_gap, f"Sheets: {sheet_names}")

        # CRITICAL: core deliverable structure -- spreadsheet with all three sheets.
        check("CRITICAL: Spreadsheet 'Curriculum Gap Analysis' with all three sheets",
              has_current and has_research and has_gap,
              f"current={has_current}, research={has_research}, gap={has_gap}",
              critical=True)

        # Check Current_Topics content
        ct_id = None
        for n, sid in sheet_map.items():
            if "current" in n and "topic" in n:
                ct_id = sid
                break
        if ct_id:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, ct_id))
            row_count = cur.fetchone()[0]
            check("Current_Topics has data rows (>= 5)", row_count >= 5, f"Got {row_count}")

        # Check Research_Topics content
        rt_id = None
        for n, sid in sheet_map.items():
            if "research" in n and "topic" in n:
                rt_id = sid
                break
        if rt_id:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, rt_id))
            row_count = cur.fetchone()[0]
            check("Research_Topics has >= 4 paper rows", row_count >= 4, f"Got {row_count}")

            # CRITICAL: Research_Topics must reference the RELEVANT papers and NO noise papers.
            cur.execute("""
                SELECT LOWER(value) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, rt_id))
            rt_text = " ".join(r[0] for r in cur.fetchall() if r[0])
            relevant_markers = ["deep learning", "graph neural", "reinforcement learning",
                                "analytics methods", "benchmarking", "combinatorial optimization"]
            noise_markers = ["marine biodiversity", "paleolithic cave art", "arctic ecosystems", "uranium-thorium"]
            n_relevant = sum(1 for t in relevant_markers if t in rt_text)
            rt_has_noise = any(nt in rt_text for nt in noise_markers)
            check("CRITICAL: Research_Topics lists >=3 relevant papers and zero noise papers",
                  n_relevant >= 3 and not rt_has_noise,
                  f"relevant={n_relevant}, noise_present={rt_has_noise}",
                  critical=True)

        # Check Gap_Analysis content
        ga_id = None
        for n, sid in sheet_map.items():
            if "gap" in n and "analysis" in n:
                ga_id = sid
                break
        if ga_id:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, ga_id))
            row_count = cur.fetchone()[0]
            check("Gap_Analysis has data rows (>= 2)", row_count >= 2, f"Got {row_count}")

            # Check gap types
            cur.execute("""
                SELECT DISTINCT LOWER(value) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND col_index = 0 AND row_index > 0
            """, (ss_id, ga_id))
            gap_types = set(r[0] for r in cur.fetchall() if r[0])
            check("Gap_Analysis has 'Curriculum Gap' or 'Missing Research' types",
                  any("curriculum" in g or "gap" in g or "missing" in g for g in gap_types),
                  f"Found: {gap_types}")

        cur.close()
        conn.close()
    except Exception as e:
        check("Google Sheet DB check", False, str(e))


# ---------- Google Calendar ----------
def check_gcal(launch_time=None):
    print("\n=== Checking Google Calendar ===")
    try:
        # launch_time may carry a trailing weekday suffix (e.g. "2026-06-08 04:50:05 Monday");
        # strip to just the date+time before parsing so the +14d/+21d anchors track the real launch.
        lt_clean = " ".join((launch_time or "2026-03-07 10:00:00").split()[:2])
        lt = datetime.strptime(lt_clean, "%Y-%m-%d %H:%M:%S")
    except Exception:
        lt = datetime.strptime("2026-03-07 10:00:00", "%Y-%m-%d %H:%M:%S")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        # The required new event titles are kept in English in task.md, but accept RU variants too.
        # Check for Curriculum Review Meeting (~+14d at 10:00, 2h).
        cur.execute("""SELECT summary, start_datetime, description FROM gcal.events
                       WHERE summary ILIKE '%curriculum%review%meeting%'
                          OR (summary ILIKE '%пересмотр%' AND summary ILIKE '%учебн%')
                          OR summary ILIKE '%curriculum review%'""")
        rows = cur.fetchall()
        review_exists = len(rows) >= 1
        check("'Curriculum Review Meeting' event exists", review_exists, f"Found {len(rows)}")
        review_time_ok = False
        if rows:
            check("Review meeting has description", rows[0][2] is not None and len(str(rows[0][2])) > 10,
                  f"Description: {str(rows[0][2])[:100]}")
            dt = parse_dt(rows[0][1])
            if dt is not None:
                expected = lt + timedelta(days=14)
                day_close = abs((dt.date() - expected.date()).days) <= 1
                hour_close = abs(dt.hour - 10) <= 1
                review_time_ok = day_close and hour_close
            check("Review meeting scheduled ~+14d at 10:00", review_time_ok,
                  f"start={rows[0][1]}, expected ~{(lt + timedelta(days=14)).date()} 10:00")

        # Check for Faculty Workshop (~+21d at 13:00, 3h).
        cur.execute("""SELECT summary, start_datetime, description FROM gcal.events
                       WHERE summary ILIKE '%faculty%workshop%'
                          OR (summary ILIKE '%семинар%' AND summary ILIKE '%преподавател%')
                          OR summary ILIKE '%workshop%new%topic%'""")
        rows = cur.fetchall()
        workshop_exists = len(rows) >= 1
        check("'Faculty Workshop on New Topics' event exists", workshop_exists, f"Found {len(rows)}")
        workshop_time_ok = False
        if rows:
            check("Workshop has description", rows[0][2] is not None and len(str(rows[0][2])) > 10,
                  f"Description: {str(rows[0][2])[:100]}")
            dt = parse_dt(rows[0][1])
            if dt is not None:
                expected = lt + timedelta(days=21)
                day_close = abs((dt.date() - expected.date()).days) <= 1
                hour_close = abs(dt.hour - 13) <= 1
                workshop_time_ok = day_close and hour_close
            check("Workshop scheduled ~+21d at 13:00", workshop_time_ok,
                  f"start={rows[0][1]}, expected ~{(lt + timedelta(days=21)).date()} 13:00")

        # CRITICAL: both required events exist AND fall on the correct days/times from task.md.
        check("CRITICAL: Both gcal events exist with correct +14d/10:00 and +21d/13:00 scheduling",
              review_exists and workshop_exists and review_time_ok and workshop_time_ok,
              f"review_exists={review_exists}/{review_time_ok}, workshop_exists={workshop_exists}/{workshop_time_ok}",
              critical=True)

        cur.close()
        conn.close()
    except Exception as e:
        check("Google Calendar check", False, str(e))
        check("CRITICAL: Both gcal events exist with correct scheduling", False, str(e), critical=True)


# ---------- Terminal outputs ----------
def check_terminal_outputs(ws):
    print("\n=== Checking Terminal Outputs ===")

    # current_topics.json
    ct_ok = False
    ct_path = os.path.join(ws, "current_topics.json")
    if os.path.isfile(ct_path):
        check("current_topics.json exists", True)
        try:
            with open(ct_path) as f:
                data = json.load(f)
        except Exception:
            data = None
        is_list5 = isinstance(data, list) and len(data) >= 5
        check("current_topics.json is a list with >= 5 items",
              is_list5, f"Got {type(data).__name__ if data is not None else 'invalid'} with {len(data) if isinstance(data, list) else 'N/A'} items")
        if isinstance(data, list) and len(data) > 0:
            first = data[0]
            check("current_topics items have 'topic' key", "topic" in first, f"Keys: {list(first.keys())}")
            ct_ok = is_list5 and isinstance(first, dict) and "topic" in first
    else:
        check("current_topics.json exists", False)

    # topic_gaps.json
    tg_all_keys = False
    tg_path = os.path.join(ws, "topic_gaps.json")
    if os.path.isfile(tg_path):
        check("topic_gaps.json exists", True)
        try:
            with open(tg_path) as f:
                data = json.load(f)
        except Exception:
            data = None
        check("topic_gaps.json has expected keys",
              isinstance(data, dict) and ("covered_topics" in data or "gap_topics" in data or "missing_research_areas" in data),
              f"Keys: {list(data.keys()) if isinstance(data, dict) else 'not a dict'}")
        tg_all_keys = isinstance(data, dict) and all(
            k in data for k in ("covered_topics", "gap_topics", "missing_research_areas"))
    else:
        check("topic_gaps.json exists", False)

    # CRITICAL: the script-produced artifacts the whole analysis depends on.
    check("CRITICAL: current_topics.json (>=5 objects w/ topic) and topic_gaps.json (all three keys)",
          ct_ok and tg_all_keys,
          f"current_topics_ok={ct_ok}, topic_gaps_all_keys={tg_all_keys}",
          critical=True)


def check_reverse_validation(workspace):
    """Check that noise/irrelevant data is NOT present in outputs."""
    print("\n=== Reverse Validation ===")

    # Noise arxiv paper titles that should NOT appear in gsheet
    noise_titles = ["marine biodiversity", "paleolithic cave art", "arctic ecosystems", "uranium-thorium"]
    noise_ids = ["2025.20001", "2025.20002"]

    # Check gsheet does not contain noise papers
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        cur.execute("SELECT id FROM gsheet.spreadsheets WHERE title ILIKE '%curriculum%gap%'")
        rows = cur.fetchall()
        if rows:
            ss_id = rows[0][0]
            cur.execute("SELECT LOWER(value) FROM gsheet.cells WHERE spreadsheet_id = %s", (ss_id,))
            all_values = " ".join(r[0] for r in cur.fetchall() if r[0])

            no_noise = not any(nt in all_values for nt in noise_titles)
            check("No noise arxiv papers in gsheet (marine biodiversity, cave art)",
                  no_noise,
                  f"Found noise content in gsheet values")
        else:
            check("No noise arxiv papers in gsheet", True, "No spreadsheet to check")

        cur.close()
        conn.close()
    except Exception as e:
        check("Reverse validation (gsheet noise)", False, str(e))

    # Check Word document does not reference noise papers
    fpath = os.path.join(workspace, "Curriculum_Modernization_Proposal.docx")
    if os.path.isfile(fpath):
        from docx import Document
        doc = Document(fpath)
        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        no_noise_word = not any(nt in full_text for nt in noise_titles)
        check("No noise arxiv papers in Word document", no_noise_word,
              "Found marine biodiversity or cave art references in proposal")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gsheet()
    check_gcal(args.launch_time)
    check_terminal_outputs(args.agent_workspace)
    check_reverse_validation(args.agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\nOverall: {PASS_COUNT}/{total} ({accuracy:.1f}%)")
    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy,
              "critical_failures": CRITICAL_FAILURES}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if CRITICAL_FAILURES:
        print(f"\nCRITICAL CHECK(S) FAILED ({len(CRITICAL_FAILURES)}): {CRITICAL_FAILURES}")
        print("Result: FAIL (critical check failed, overriding accuracy gate)")
        sys.exit(1)

    sys.exit(0 if accuracy >= 70 else 1)


if __name__ == "__main__":
    main()
