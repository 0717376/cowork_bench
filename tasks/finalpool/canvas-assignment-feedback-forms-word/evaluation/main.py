"""Evaluation for canvas-assignment-feedback-forms-word (RU stack: canvas kept + forms swap).

Сценарий: аналитик курса «Основы финансов» (Основы финансов, Fall 2013,
course ID 16) собирает статистику по заданиям из Canvas, создаёт опрос обратной
связи (forms MCP, схема gform.*), документ Word Assignment_Analysis.docx и
отправляет итоговое письмо преподавателю.

Критические чеки (CRITICAL_CHECKS): любой их fail => задача FAIL, независимо от
общей accuracy. Это семантические чеки сути задачи (правильные числа из Canvas,
покрытие тем опроса, корректное письмо), а НЕ структура (наличие колонки/таблицы).

Порог: accuracy >= 70 И нет проваленных критических чеков.
"""
import argparse
import json
import os
import re
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []

# Эталонные данные из Canvas (course 16, Основы финансов).
# (Assignment_Name, Total_Submissions, Avg_Score, Late_Submissions, Late_Rate%)
EXPECTED_ASSIGNMENTS = [
    ("CMA 34878", 1470, 83.32, 36, 2.45),
    ("CMA 34879", 1352, 87.97, 43, 3.18),
    ("CMA 34880", 1252, 76.50, 61, 4.87),
    ("CMA 34881", 1224, 78.30, 59, 4.82),
    ("CMA 34882", 1193, 78.65, 75, 6.29),
    ("CMA 34883", 1196, 76.80, 119, 9.95),
    ("CMA 34884", 1160, 76.16, 88, 7.59),
    ("TMA 34873", 1859, 78.16, 157, 8.45),
    ("TMA 34874", 1661, 72.45, 193, 11.62),
    ("TMA 34875", 1402, 70.49, 376, 26.82),
    ("TMA 34876", 1313, 71.05, 268, 20.41),
    ("TMA 34877", 1158, 75.89, 148, 12.78),
    # Final Exam 34885 существует в курсе 16, но имеет 0 сдач (ungraded).
    # task.md Part 3 требует "включить все задания, найденные в курсе",
    # поэтому строка с 13-м заданием допустима. Avg/Late_Rate при 0 сдачах = 0.
    ("Final Exam 34885", 0, 0.0, 0, 0.0),
]

# Критические чеки — по строке name, как в check()
CRITICAL_CHECKS = {
    "Word: >=10 из 12 строк заданий совпадают с Canvas по числовым значениям",
    "Word: 12 или 13 строк, отсортированы по алфавиту (CMA до TMA)",
    "Опрос: >=5 вопросов И покрыты все 5 тем обратной связи",
    "Письмо instructor@financeou.example.com: верная тема, тело про 12/13/34879/34875",
    "Письмо содержит ссылку на опрос (forms URL / id)",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILED
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILED.append(name)


def _to_float(s):
    if s is None:
        return None
    m = re.search(r"-?\d+(?:[.,]\d+)?", str(s).replace(" ", ""))
    if not m:
        return None
    try:
        return float(m.group(0).replace(",", "."))
    except ValueError:
        return None


def _to_int(s):
    f = _to_float(s)
    return int(round(f)) if f is not None else None


def check_word(agent_workspace):
    """Проверка документа Word."""
    print("\n=== Проверка документа Word ===")
    crit_data = "Word: >=10 из 12 строк заданий совпадают с Canvas по числовым значениям"
    crit_sort = "Word: 12 или 13 строк, отсортированы по алфавиту (CMA до TMA)"
    try:
        from docx import Document
    except ImportError:
        check("python-docx доступен", False, "python-docx not installed")
        check(crit_data, False, "no docx lib")
        check(crit_sort, False, "no docx lib")
        return

    agent_file = os.path.join(agent_workspace, "Assignment_Analysis.docx")
    exists = os.path.isfile(agent_file)
    check("Файл Word существует", exists, f"Expected {agent_file}")
    if not exists:
        check(crit_data, False, "no file")
        check(crit_sort, False, "no file")
        return

    try:
        doc = Document(agent_file)
    except Exception as e:
        check("Файл Word читается", False, str(e))
        check(crit_data, False, "unreadable")
        check(crit_sort, False, "unreadable")
        return

    # Заголовок: принимаем EN-литерал, а также RU-альтернативы.
    full_text = "\n".join(p.text for p in doc.paragraphs).lower()
    title_ok = (
        ("foundations of finance" in full_text and "assignment analysis" in full_text)
        or ("основы финансов" in full_text and ("анализ задани" in full_text or "анализ зада" in full_text))
    )
    check("Заголовок документа присутствует", title_ok, f"Title not found in: {full_text[:200]}")

    check("В документе есть хотя бы одна таблица", len(doc.tables) >= 1, f"Found {len(doc.tables)} tables")
    if not doc.tables:
        check(crit_data, False, "no table")
        check(crit_sort, False, "no table")
        return

    tbl = doc.tables[0]
    # Заголовки колонок — английские идентификаторы (eval их грепает).
    headers = [c.text.strip() for c in tbl.rows[0].cells] if tbl.rows else []
    hl = [h.lower() for h in headers]
    check("Колонка 'Assignment_Name' присутствует",
          any("assignment_name" in h for h in hl), f"Headers: {headers}")
    check("Колонка 'Total_Submissions' присутствует",
          any("total_submissions" in h for h in hl), f"Headers: {headers}")
    check("Колонка 'Avg_Score' присутствует",
          any("avg_score" in h or "avg" in h for h in hl), f"Headers: {headers}")
    check("Колонка 'Late_Submissions' присутствует",
          any("late_submissions" in h for h in hl), f"Headers: {headers}")
    check("Колонка 'Late_Rate(%)' присутствует",
          any("late_rate" in h for h in hl), f"Headers: {headers}")

    # Сопоставление позиций колонок (на случай произвольного порядка — но строки сортируем по name).
    def col_idx(*subs):
        for i, h in enumerate(hl):
            if any(s in h for s in subs):
                return i
        return None

    i_name = col_idx("assignment_name")
    i_total = col_idx("total_submissions")
    i_avg = col_idx("avg_score", "avg")
    i_late = col_idx("late_submissions")
    i_rate = col_idx("late_rate")

    data_rows = list(tbl.rows)[1:]
    row_cells = [[c.text.strip() for c in r.cells] for r in data_rows if r.cells]

    # Структурные чеки (non-critical)
    check("Таблица содержит 12 строк данных", len(row_cells) >= 12, f"Found {len(row_cells)} rows")
    row_names = [r[i_name] if i_name is not None and i_name < len(r) else (r[0] if r else "") for r in row_cells]
    rl = [n.lower() for n in row_names]
    check("Строка CMA 34879 присутствует", any("cma 34879" in n for n in rl), f"Row names: {row_names[:5]}")
    check("Строка TMA 34875 присутствует", any("tma 34875" in n for n in rl), f"Row names: {row_names[:5]}")

    # CRITICAL: числовые значения совпадают с Canvas (>=10 из 12).
    exp_by_name = {name.lower(): vals for name, *vals in
                   [(a[0], a[1], a[2], a[3], a[4]) for a in EXPECTED_ASSIGNMENTS]}
    matched = 0
    if None not in (i_name, i_total, i_avg, i_late, i_rate):
        for r in row_cells:
            if i_name >= len(r):
                continue
            key = None
            for k in exp_by_name:
                if k in r[i_name].lower():
                    key = k
                    break
            if key is None:
                continue
            e_total, e_avg, e_late, e_rate = exp_by_name[key]
            a_total = _to_int(r[i_total]) if i_total < len(r) else None
            a_avg = _to_float(r[i_avg]) if i_avg < len(r) else None
            a_late = _to_int(r[i_late]) if i_late < len(r) else None
            a_rate = _to_float(r[i_rate]) if i_rate < len(r) else None
            ok = (
                a_total == e_total
                and a_avg is not None and abs(a_avg - e_avg) <= 0.01
                and a_late == e_late
                and a_rate is not None and abs(a_rate - e_rate) <= 0.05
            )
            if ok:
                matched += 1
    check(crit_data, matched >= 10, f"matched {matched}/12 rows numerically")

    # CRITICAL: 12 (только оцениваемые) или 13 (включая Final Exam 34885) строк
    # И отсортированы по Assignment_Name по возрастанию.
    names_sorted = (row_names == sorted(row_names, key=lambda s: s.lower()))
    check(crit_sort, len(row_cells) in (12, 13) and names_sorted,
          f"rows={len(row_cells)} sorted={names_sorted} names={row_names}")


def check_gform():
    """Проверка опроса обратной связи (forms MCP, схема gform.*)."""
    print("\n=== Проверка опроса (forms) ===")
    crit = "Опрос: >=5 вопросов И покрыты все 5 тем обратной связи"
    try:
        conn = psycopg2.connect(**DB)
    except Exception as e:
        check("Опрос создан", False, str(e))
        check(crit, False, "no db")
        return
    cur = conn.cursor()

    cur.execute("SELECT id, title FROM gform.forms")
    forms = cur.fetchall()
    check("Создан хотя бы один опрос", len(forms) >= 1, f"Found {len(forms)} forms")

    # Выбираем форму по релевантному заголовку (EN-литерал или RU-ключи).
    form_id = None
    for fid, title in forms:
        t = (title or "").lower()
        if any(k in t for k in ("assignment", "feedback", "finance", "обратн", "задани", "финанс")):
            form_id = fid
            break
    if form_id is None and forms:
        form_id = forms[0][0]
    check("Опрос с релевантным заголовком найден", form_id is not None, f"Forms: {forms}")

    if form_id is None:
        check(crit, False, "no form")
        cur.close()
        conn.close()
        return

    cur.execute("SELECT title FROM gform.questions WHERE form_id = %s ORDER BY position", (form_id,))
    questions = [(q[0] or "").lower() for q in cur.fetchall()]
    cur.close()
    conn.close()

    q_count = len(questions)
    check("Опрос содержит >=5 вопросов", q_count >= 5, f"Found {q_count} questions")

    all_q = " | ".join(questions)

    # 5 тем: нагрузка, ясность инструкций, время на выполнение, поддержка преподавателя,
    # общая удовлетворённость. Принимаем RU+EN ключи в исходном .lower() тексте.
    topic_workload = any(k in all_q for k in ("workload", "load", "нагрузк", "объём", "объем"))
    topic_clarity = any(k in all_q for k in ("clarity", "clear", "instruction", "ясност", "понятн", "формулиров", "инструкц"))
    topic_time = any(k in all_q for k in ("time", "duration", "hours", "врем", "час", "продолжительн"))
    topic_support = any(k in all_q for k in ("support", "instructor", "help", "поддержк", "помощ", "преподавател"))
    topic_satisfaction = any(k in all_q for k in ("satisf", "overall", "удовлетвор", "довольн", "общая оценк", "общую оценк"))

    check("Тема: нагрузка по заданиям", topic_workload, all_q[:200])
    check("Тема: ясность инструкций", topic_clarity, all_q[:200])
    check("Тема: требуемое время", topic_time, all_q[:200])
    check("Тема: поддержка преподавателя", topic_support, all_q[:200])
    check("Тема: общая удовлетворённость", topic_satisfaction, all_q[:200])

    topics_ok = all([topic_workload, topic_clarity, topic_time, topic_support, topic_satisfaction])
    check(crit, q_count >= 5 and topics_ok,
          f"count={q_count} workload={topic_workload} clarity={topic_clarity} "
          f"time={topic_time} support={topic_support} satisf={topic_satisfaction}")


def check_emails():
    """Проверка итогового письма преподавателю."""
    print("\n=== Проверка письма ===")
    crit_body = "Письмо instructor@financeou.example.com: верная тема, тело про 12/13/34879/34875"
    crit_link = "Письмо содержит ссылку на опрос (forms URL / id)"
    try:
        conn = psycopg2.connect(**DB)
    except Exception as e:
        check("Итоговое письмо отправлено", False, str(e))
        check(crit_body, False, "no db")
        check(crit_link, False, "no db")
        return
    cur = conn.cursor()
    cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages")
    all_emails = cur.fetchall()

    # Достаём id/url форм для проверки ссылки на опрос.
    try:
        cur.execute("SELECT id FROM gform.forms")
        form_ids = [str(r[0]) for r in cur.fetchall()]
    except Exception:
        form_ids = []
    cur.close()
    conn.close()

    def recipients_of(to_addr):
        if not to_addr:
            return []
        if isinstance(to_addr, list):
            return [str(r).strip().lower() for r in to_addr]
        try:
            parsed = json.loads(to_addr)
            if isinstance(parsed, list):
                return [str(r).strip().lower() for r in parsed]
            return [str(to_addr).strip().lower()]
        except (json.JSONDecodeError, TypeError):
            return [str(to_addr).strip().lower()]

    result = None
    for subj, from_addr, to_addr, body in all_emails:
        if "instructor@financeou.example.com" in recipients_of(to_addr):
            result = (subj, from_addr, to_addr, body)
            break

    check("Письмо отправлено на instructor@financeou.example.com", result is not None,
          f"Total emails found: {len(all_emails)}")
    if result is None:
        check(crit_body, False, "no email")
        check(crit_link, False, "no email")
        return

    subj, from_addr, to_addr, body = result
    sl = (subj or "").lower()
    subj_ok = "assignment" in sl and "analysis" in sl
    check("Тема письма содержит 'Assignment Analysis'", subj_ok, f"Subject: {subj}")
    from_ok = "analytics@university.example.com" in (from_addr or "").lower()
    check("Письмо от analytics@university.example.com", from_ok, f"From: {from_addr}")

    body_lower = (body or "").lower()
    # Число заданий: 12 (только оцениваемые) или 13 (включая Final Exam 34885).
    n_ok = any(k in body_lower for k in
               ("12", "twelve", "двенадцать", "13", "thirteen", "тринадцать"))
    check("Письмо упоминает 12 или 13 заданий", n_ok, "Expected mention of 12 or 13")
    a_ok = "34879" in body_lower
    check("Письмо упоминает CMA 34879 (наивысший средний балл)", a_ok, "Expected 34879")
    l_ok = "34875" in body_lower
    check("Письмо упоминает TMA 34875 (больше всего поздних сдач)", l_ok, "Expected 34875")

    # CRITICAL: тема + все три ключевых факта
    check(crit_body, subj_ok and n_ok and a_ok and l_ok,
          f"subj={subj_ok} n12={n_ok} 34879={a_ok} 34875={l_ok}; body={body[:200]}")

    # CRITICAL: ссылка на опрос (forms URL или id формы в теле письма).
    link_ok = ("form" in body_lower and ("http" in body_lower or "/forms" in body_lower)) or \
              any(fid and fid in (body or "") for fid in form_ids)
    check(crit_link, link_ok, f"form_ids={form_ids}; body={body[:200]}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("CANVAS ASSIGNMENT FEEDBACK FORMS WORD - EVALUATION")
    print("=" * 70)

    check_word(args.agent_workspace)
    check_gform()
    check_emails()

    total = PASS_COUNT + FAIL_COUNT
    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    if total == 0:
        print("FAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"  Accuracy: {accuracy:.1f}%")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_failed": CRITICAL_FAILED,
    }
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if CRITICAL_FAILED:
        print(f"  CRITICAL FAIL: {CRITICAL_FAILED}")
        print("  Overall: FAIL")
        sys.exit(1)

    if accuracy >= 70:
        print("  Overall: PASS")
        sys.exit(0)
    print("  Overall: FAIL (accuracy < 70%)")
    sys.exit(1)


if __name__ == "__main__":
    main()
