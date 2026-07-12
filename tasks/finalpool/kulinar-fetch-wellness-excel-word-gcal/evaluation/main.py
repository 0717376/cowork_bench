"""Evaluation script for kulinar-fetch-wellness-excel-word-gcal."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel"
}

# Известные блюда kulinar (русская кулинарная база, 50 рецептов).
# Используется критическим чеком, чтобы подтвердить, что агент реально
# обращался к базе рецептов, а не выдумывал названия.
KULINAR_RECIPES = {
    'салат оливье', 'винегрет', 'сельдь под шубой', 'салат мимоза',
    'крабовый салат', 'греческий салат', 'салат с курицей и грибами',
    'холодец', 'икра кабачковая', 'грибы маринованные', 'сало солёное',
    'селёдка с луком', 'борщ', 'щи из квашеной капусты', 'солянка мясная',
    'уха', 'окрошка', 'грибной суп', 'рассольник', 'куриный бульон с лапшой',
    'бефстроганов', 'пельмени домашние', 'голубцы', 'котлеты домашние',
    'жаркое в горшочках', 'курица в сметане', 'рыба запечённая по-русски',
    'цыплёнок табака', 'гречка с тушёнкой', 'плов узбекский',
    'картофельное пюре', 'гречневая каша', 'перловая каша',
    'картофель отварной с укропом', 'рис отварной',
    'пирожки с капустой жареные', 'пирожки с мясом печёные', 'блины тонкие',
    'кулебяка с капустой и яйцом', 'расстегаи с рыбой', 'медовик', 'наполеон',
    'сырники', 'пасха творожная', 'ватрушки с творогом', 'кисель ягодный',
    'морс клюквенный', 'компот из сухофруктов', 'сбитень', 'квас домашний',
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []

def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILED
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRITICAL]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL]{' [CRITICAL]' if critical else ''} {name}: {detail_str}")
        if critical:
            CRITICAL_FAILED.append(name)

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILED
    PASS_COUNT = 0
    FAIL_COUNT = 0
    CRITICAL_FAILED = []

    # Check Corporate_Wellness_Plan.xlsx
    excel_path = os.path.join(agent_workspace, "Corporate_Wellness_Plan.xlsx")
    check("Corporate_Wellness_Plan.xlsx exists", os.path.exists(excel_path))
    recipe_names_lower = []
    metrics = {}
    eval_row_count = 0
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Corporate_Wellness_Plan.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        # --- Сбор данных для критических чеков ---
        if "Recipe_Evaluation" in wb.sheetnames:
            ws_re = wb["Recipe_Evaluation"]
            re_headers = [str(c.value).strip() if c.value else "" for c in ws_re[1]]
            try:
                name_idx = re_headers.index("Recipe_Name")
            except ValueError:
                name_idx = 0
            for row in ws_re.iter_rows(min_row=2, values_only=True):
                if row and any(v is not None for v in row):
                    eval_row_count += 1
                    nm = row[name_idx] if name_idx < len(row) else None
                    if nm:
                        recipe_names_lower.append(str(nm).strip().lower())
        if "Program_Metrics" in wb.sheetnames:
            ws_pm = wb["Program_Metrics"]
            for row in ws_pm.iter_rows(min_row=2, values_only=True):
                if row and row[0]:
                    metrics[str(row[0]).strip()] = row[1] if len(row) > 1 else None

        if gt_wb:
            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Check row count
                    gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                    data_rows = list(ws.iter_rows(min_row=2, values_only=True))
                    min_rows = max(1, len(gt_rows) - 2)
                    check(f"{sheet_name} has >= {min_rows} data rows", len(data_rows) >= min_rows, f"got {len(data_rows)}")

                    # Структурная валидация (значения определяет агент,
                    # сравнение с одной произвольной GT-реализацией некорректно)
                    header_map = {h: i for i, h in enumerate(headers)}

                    def col_vals(col):
                        ci = header_map.get(col)
                        if ci is None:
                            return []
                        return [r[ci] for r in data_rows
                                if r and ci < len(r) and r[ci] is not None]

                    if sheet_name == "Recipe_Evaluation":
                        scores = [safe_float(v) for v in col_vals("wellness_score")]
                        check("Recipe_Evaluation Wellness_Score values in 1-10",
                              bool(scores) and all(s is not None and 0 < s <= 10 for s in scores),
                              f"scores: {scores[:10]}")
                        cals = [safe_float(v) for v in col_vals("estimated_calories")]
                        check("Recipe_Evaluation Estimated_Calories plausible (50-2000)",
                              bool(cals) and all(c is not None and 50 <= c <= 2000 for c in cals),
                              f"calories: {cals[:10]}")
                    elif sheet_name == "Weekly_Plan":
                        WEEKDAYS = {
                            "понедельник", "вторник", "среда", "четверг", "пятница",
                            "суббота", "воскресенье",
                            "monday", "tuesday", "wednesday", "thursday", "friday",
                            "saturday", "sunday",
                        }
                        days = [str(v).strip().lower() for v in col_vals("day")]
                        check("Weekly_Plan Day values are weekdays (RU or EN)",
                              bool(days) and all(d in WEEKDAYS for d in days),
                              f"days: {days[:7]}")
                        meals = [str(v).strip().lower()
                                 for col in ("breakfast", "lunch", "snack")
                                 for v in col_vals(col)]
                        real = sum(1 for m in meals if m in KULINAR_RECIPES)
                        check("Weekly_Plan meal slots are real kulinar recipes (>=80%)",
                              bool(meals) and real >= 0.8 * len(meals),
                              f"real {real}/{len(meals)}, sample: {meals[:5]}")
                    elif sheet_name == "Nutritional_Summary":
                        prot = [safe_float(v) for v in col_vals("protein_pct")]
                        carb = [safe_float(v) for v in col_vals("carb_pct")]
                        fat = [safe_float(v) for v in col_vals("fat_pct")]
                        sums_ok = (
                            prot and len(prot) == len(carb) == len(fat)
                            and all(p is not None and c is not None and f is not None
                                    and abs(p + c + f - 100) <= 10
                                    for p, c, f in zip(prot, carb, fat))
                        )
                        check("Nutritional_Summary macro pcts sum to ~100 per row",
                              sums_ok, f"prot={prot} carb={carb} fat={fat}")
                        avgs = [safe_float(v) for v in col_vals("avg_calories")]
                        check("Nutritional_Summary Avg_Calories plausible (100-1500)",
                              bool(avgs) and all(a is not None and 100 <= a <= 1500 for a in avgs),
                              f"avg_calories: {avgs}")
                    # Program_Metrics: значения покрыты CRITICAL-чеком согласованности ниже

    # Check Wellness_Program_Guide.docx
    docx_path = os.path.join(agent_workspace, "Wellness_Program_Guide.docx")
    check("Wellness_Program_Guide.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        check("Wellness_Program_Guide.docx has content", len(text) > 50, f"text length: {len(text)}")
        # Check headings match groundtruth (RU+EN pairs)
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        # Каждый обязательный раздел задан парой вариантов (RU и EN) —
        # агент может назвать раздел на любом из языков.
        REQUIRED_SECTIONS = [
            ["обзор программы", "program overview"],
            ["недельный план питания", "weekly meal plans", "weekly meal plan"],
            ["соответствие нормам питания", "nutritional guidelines compliance",
             "nutritional guidelines"],
            ["рекомендации по внедрению", "implementation recommendations"],
        ]
        all_headings_text = " | ".join(headings)
        for variants in REQUIRED_SECTIONS:
            found = any(
                v in h or h in v for h in headings for v in variants
            ) or any(v in all_headings_text for v in variants)
            check(
                f"Wellness_Program_Guide.docx has section \"{variants[0][:40]}\"",
                found, f"agent headings: {headings[:6]}", critical=True,
            )

    # Check Python script exists (terminal usage) + wellness_plan.json produced
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    check("Python analysis script exists", len(py_files) >= 1, f"found: {py_files}")
    plan_json_path = os.path.join(agent_workspace, "wellness_plan.json")
    # КРИТИЧЕСКИЙ: pipeline терминал+python отработал до конца
    check("Python pipeline produced output (any .py + wellness_plan.json)",
          len(py_files) >= 1 and os.path.exists(plan_json_path),
          f"py_files={py_files}, wellness_plan.json={os.path.exists(plan_json_path)}",
          critical=True)

    # --- КРИТИЧЕСКИЙ: реальные рецепты kulinar в Recipe_Evaluation ---
    real_recipe_hits = sum(1 for n in recipe_names_lower if n in KULINAR_RECIPES)
    check("Recipe_Evaluation has >=10 rows with real kulinar recipes",
          eval_row_count >= 10 and real_recipe_hits >= 8,
          f"rows={eval_row_count}, real_kulinar_matches={real_recipe_hits}, names={recipe_names_lower[:5]}",
          critical=True)

    # --- КРИТИЧЕСКИЙ: согласованность Program_Metrics ---
    total_eval = safe_float(metrics.get("Total_Recipes_Evaluated"))
    weekly_avg = safe_float(metrics.get("Weekly_Avg_Calories"))
    compliance = safe_float(metrics.get("Program_Compliance_Pct"))
    metrics_ok = (
        total_eval is not None and abs(total_eval - eval_row_count) <= 1
        and weekly_avg is not None and 1500 <= weekly_avg <= 2500
        and compliance is not None and 0 <= compliance <= 100
    )
    check("Program_Metrics consistent (total==rows, weekly_avg in 1500-2500 band, compliance 0-100)",
          metrics_ok,
          f"total={total_eval} rows={eval_row_count} weekly_avg={weekly_avg} compliance={compliance}",
          critical=True)

    # Database checks
    try:
        conn = get_conn()
        cur = conn.cursor()
        # Запуск велнес-программы: агент может назвать событие на EN ('Wellness
        # Program Launch') или на RU ('велнес' / 'велнес-программа'), 18.03.2026.
        cur.execute(
            "SELECT summary, start_datetime FROM gcal.events "
            "WHERE summary ILIKE '%wellness%' OR summary ILIKE '%велнес%'"
        )
        event_row = cur.fetchone()
        # КРИТИЧЕСКИЙ: событие запуска существует
        check("Calendar launch event exists (wellness/велнес summary)",
              event_row is not None, "no matching event found", critical=True)
        # Reverse verification: шумовые события (Планёрка/Обед) не удалены
        cur.execute(
            "SELECT COUNT(*) FROM gcal.events "
            "WHERE summary ILIKE '%планёрка%' OR summary ILIKE '%планерка%' "
            "OR summary ILIKE '%обед%' OR summary ILIKE '%standup%' OR summary ILIKE '%lunch%'"
        )
        noise_events = cur.fetchone()[0]
        check("Noise events exist (not deleted by agent)", noise_events >= 1,
              f"noise events: {noise_events}", critical=True)
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e))

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    if CRITICAL_FAILED:
        print(f"CRITICAL checks failed: {CRITICAL_FAILED}")
        print(f"Passed {PASS_COUNT}/{total} checks (accuracy {accuracy:.1f}%) but a CRITICAL check failed -> FAIL")
        sys.exit(1)
    success = accuracy >= 70
    return success, f"Passed {PASS_COUNT}/{total} checks (accuracy {accuracy:.1f}%)"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()