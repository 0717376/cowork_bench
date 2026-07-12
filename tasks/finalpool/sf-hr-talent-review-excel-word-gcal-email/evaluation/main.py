"""Evaluation for sf-hr-talent-review-excel-word-gcal-email.

Checks:
1. Excel file: Talent_Review_Report.xlsx with 3 sheets and correct data
2. Word file: Talent_Review_Summary.docx with correct structure and content
3. GCal: 7 department review meetings scheduled ~14-20 days from launch
4. Email: summary email sent to hr_leadership@company.com
"""
import argparse
import json
import os
import sys
from datetime import datetime, timedelta, timezone

import openpyxl
import psycopg2

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)

from docx import Document

DB = dict(
    host=os.environ.get("PGHOST", "localhost"), port=5432,
    dbname=os.environ.get("PGDATABASE", "cowork_gym"),
    user="eigent", password="camel"
)

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []

# Departments russified by the central ClickHouse relabel map
# (scripts/clickhouse_relabel_map.py): the agent reads russified sf_data,
# so its Excel/Word/GCal output uses Russian department names. Numeric
# expectations below are NOT realia and are unchanged after russification.
DEPARTMENTS = ["Инженерия", "Финансы", "Кадры", "Операции", "НИОКР", "Продажи", "Поддержка"]

# Expected department data (from actual ClickHouse queries; keyed by russified name).
DEPT_EXPECTED = {
    "Инженерия":  {"headcount": 7096, "avg_perf": 3.21, "avg_salary": 58991.61, "promo": 1578},
    "Финансы":    {"headcount": 7148, "avg_perf": 3.21, "avg_salary": 57878.19, "promo": 1588},
    "Кадры":      {"headcount": 7077, "avg_perf": 3.20, "avg_salary": 58920.45, "promo": 1588},
    "Операции":   {"headcount": 7120, "avg_perf": 3.18, "avg_salary": 57808.74, "promo": 1541},
    "НИОКР":      {"headcount": 7083, "avg_perf": 3.20, "avg_salary": 57905.93, "promo": 1519},
    "Продажи":    {"headcount": 7232, "avg_perf": 3.19, "avg_salary": 58864.79, "promo": 1596},
    "Поддержка":  {"headcount": 7244, "avg_perf": 3.20, "avg_salary": 58400.48, "promo": 1540},
}

INDUSTRY_BENCHMARKS = {
    "Инженерия": 62000, "Финансы": 60000, "Кадры": 55000, "Операции": 54000,
    "НИОКР": 63000, "Продажи": 61000, "Поддержка": 53000,
}

# Salary stats (russified key) used to value-verify Salary_Analysis math.
SALARY_EXPECTED = {
    "Инженерия":  {"avg": 58991.61, "min": 15360, "max": 695267, "range": 679907, "bpe": 89283.49, "gap": -3008.39},
    "Финансы":    {"avg": 57878.19, "min": 15760, "max": 638897, "range": 623137, "bpe": 86315.19, "gap": -2121.81},
    "Кадры":      {"avg": 58920.45, "min": 18307, "max": 692232, "range": 673925, "bpe": 98616.84, "gap": 3920.45},
    "Операции":   {"avg": 57808.74, "min": 17168, "max": 656505, "range": 639337, "bpe": 80430.93, "gap": 3808.74},
    "НИОКР":      {"avg": 57905.93, "min": 15128, "max": 680490, "range": 665362, "bpe": 101075.09, "gap": -5094.07},
    "Продажи":    {"avg": 58864.79, "min": 15885, "max": 652806, "range": 636921, "bpe": 88851.66, "gap": -2135.21},
    "Поддержка":  {"avg": 58400.48, "min": 15916, "max": 608157, "range": 592241, "bpe": 94525.33, "gap": 5400.48},
}

# Top promotion candidates per the PDF rule (rating>=4 AND years>=3, sorted
# perf desc then salary desc). Names russified by the central map
# (Luke Lewis -> Лука Лебедев, Leo Johnson -> Лев Иванов).
TOP_CANDIDATES = ["Лука Лебедев", "Лев Иванов"]

# Highest / lowest average-salary departments (russified) for the email check.
HIGHEST_SALARY_DEPT = "Инженерия"   # 58991.61
LOWEST_SALARY_DEPT = "Операции"     # 57808.74


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRIT]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        if critical:
            CRITICAL_FAILED.append(name)
        detail_str = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL]{' [CRIT]' if critical else ''} {name}{detail_str}")


def safe_float(val, default=None):
    try:
        return float(val)
    except (TypeError, ValueError):
        return default


def parse_recipients(to_addr):
    if to_addr is None:
        return []
    if isinstance(to_addr, list):
        return [str(r).strip().lower() for r in to_addr]
    to_str = str(to_addr).strip()
    try:
        parsed = json.loads(to_str)
        if isinstance(parsed, list):
            return [str(r).strip().lower() for r in parsed]
        return [to_str.lower()]
    except (json.JSONDecodeError, TypeError):
        return [to_str.lower()]


def check_excel(agent_workspace):
    print("\n=== Checking Excel: Talent_Review_Report.xlsx ===")
    xlsx_path = os.path.join(agent_workspace, "Talent_Review_Report.xlsx")
    check("Excel file exists", os.path.isfile(xlsx_path), f"Not found: {xlsx_path}")
    if not os.path.isfile(xlsx_path):
        return

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    sheet_names = [s.lower() for s in wb.sheetnames]

    # Sheet 1: Department_Summary
    has_dept_summary = any("department" in s and "summary" in s for s in sheet_names)
    check("Sheet 'Department_Summary' exists", has_dept_summary, f"Sheets: {wb.sheetnames}")

    if has_dept_summary:
        ws = None
        for name in wb.sheetnames:
            if "department" in name.lower() and "summary" in name.lower():
                ws = wb[name]
                break

        headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        rows = [r for r in rows if r[0] is not None]

        check("Department_Summary has >= 7 data rows", len(rows) >= 7, f"Found {len(rows)} rows")

        # Find column indices
        dept_col = next((i for i, h in enumerate(headers) if "department" in h), 0)
        hc_col = next((i for i, h in enumerate(headers) if "headcount" in h), 1)
        perf_col = next((i for i, h in enumerate(headers) if "perf" in h), 2)
        sal_col = next((i for i, h in enumerate(headers) if "avg" in h and "sal" in h), 3)
        promo_col = next((i for i, h in enumerate(headers) if "promo" in h), 5)

        depts_found = set()
        for row in rows:
            dept_name = str(row[dept_col]).strip() if row[dept_col] else ""
            matched_dept = None
            for d in DEPARTMENTS:
                if d.lower() == dept_name.lower():
                    matched_dept = d
                    break
            if not matched_dept:
                continue
            depts_found.add(matched_dept)
            exp = DEPT_EXPECTED[matched_dept]

            hc = safe_float(row[hc_col])
            if hc is not None:
                check(f"{matched_dept} Headcount", abs(hc - exp['headcount']) <= 5,
                      f"Expected ~{exp['headcount']}, got {hc}", critical=True)

            perf = safe_float(row[perf_col])
            if perf is not None:
                check(f"{matched_dept} Avg_Performance", abs(perf - exp['avg_perf']) <= 0.05,
                      f"Expected ~{exp['avg_perf']}, got {perf}")

            promo = safe_float(row[promo_col])
            if promo is not None:
                check(f"{matched_dept} Promo_Candidates", abs(promo - exp['promo']) <= 10,
                      f"Expected ~{exp['promo']}, got {promo}", critical=True)

        check("All 7 departments in Department_Summary", len(depts_found) == 7,
              f"Found {len(depts_found)}: {depts_found}", critical=True)

    # Sheet 2: Promotion_Candidates
    has_promo = any("promotion" in s or "promo" in s for s in sheet_names)
    check("Sheet 'Promotion_Candidates' exists", has_promo, f"Sheets: {wb.sheetnames}")
    if has_promo:
        ws2 = None
        for name in wb.sheetnames:
            if "promotion" in name.lower() or "promo" in name.lower():
                ws2 = wb[name]
                break
        rows2 = list(ws2.iter_rows(min_row=2, values_only=True))
        rows2 = [r for r in rows2 if r[0] is not None]
        check("Promotion_Candidates has 21 rows (3 per dept)", abs(len(rows2) - 21) <= 2,
              f"Found {len(rows2)} rows")

        # Check known russified top candidates (rating>=4 AND years>=3, sorted desc).
        all_text = " ".join(str(c) for row in rows2 for c in row if c)
        for cand in TOP_CANDIDATES:
            check(f"Promotion_Candidates contains '{cand}'", cand in all_text,
                  f"Missing top candidate {cand}", critical=True)

    # Sheet 3: Salary_Analysis
    has_salary = any("salary" in s for s in sheet_names)
    check("Sheet 'Salary_Analysis' exists", has_salary, f"Sheets: {wb.sheetnames}")
    if has_salary:
        ws3 = None
        for name in wb.sheetnames:
            if "salary" in name.lower():
                ws3 = wb[name]
                break
        headers3 = [str(c.value).strip().lower() if c.value else "" for c in ws3[1]]
        rows3 = list(ws3.iter_rows(min_row=2, values_only=True))
        rows3 = [r for r in rows3 if r[0] is not None]
        check("Salary_Analysis has >= 7 rows", len(rows3) >= 7, f"Found {len(rows3)} rows")

        # Check for benchmark columns
        has_benchmark = any("benchmark" in h for h in headers3)
        check("Salary_Analysis has benchmark column", has_benchmark, f"Headers: {headers3}")
        has_gap = any("gap" in h for h in headers3)
        check("Salary_Analysis has gap column", has_gap, f"Headers: {headers3}")

        # Locate columns by header keyword.
        s_dept = next((i for i, h in enumerate(headers3) if "department" in h), 0)
        s_avg = next((i for i, h in enumerate(headers3) if "avg" in h and "sal" in h), None)
        s_range = next((i for i, h in enumerate(headers3) if "range" in h), None)
        s_bench = next((i for i, h in enumerate(headers3) if "benchmark" in h), None)
        s_gap = next((i for i, h in enumerate(headers3) if "gap" in h), None)
        s_bpe = next((i for i, h in enumerate(headers3)
                      if ("budget" in h and ("per" in h or "employee" in h))), None)

        # Value-verify Salary_Analysis math with num_close (was previously unchecked).
        for row in rows3:
            dname = str(row[s_dept]).strip() if row[s_dept] else ""
            matched = next((d for d in DEPARTMENTS if d.lower() == dname.lower()), None)
            if not matched:
                continue
            exp = SALARY_EXPECTED[matched]
            if s_avg is not None and safe_float(row[s_avg]) is not None:
                check(f"{matched} Avg_Salary value", num_close(row[s_avg], exp['avg']),
                      f"Expected ~{exp['avg']}, got {row[s_avg]}")
            if s_range is not None and safe_float(row[s_range]) is not None:
                check(f"{matched} Salary_Range value", num_close(row[s_range], exp['range']),
                      f"Expected ~{exp['range']}, got {row[s_range]}")
            if s_bpe is not None and safe_float(row[s_bpe]) is not None:
                check(f"{matched} Budget_Per_Employee value", num_close(row[s_bpe], exp['bpe']),
                      f"Expected ~{exp['bpe']}, got {row[s_bpe]}")
            if s_gap is not None and safe_float(row[s_gap]) is not None:
                # Benchmark_Gap = Avg_Salary - Industry_Benchmark. Critical for Engineering & Support.
                is_crit = matched in ("Инженерия", "Поддержка")
                check(f"{matched} Benchmark_Gap value", num_close(row[s_gap], exp['gap']),
                      f"Expected ~{exp['gap']}, got {row[s_gap]}", critical=is_crit)

    wb.close()


def check_word(agent_workspace):
    print("\n=== Checking Word: Talent_Review_Summary.docx ===")
    docx_path = os.path.join(agent_workspace, "Talent_Review_Summary.docx")
    check("Word file exists", os.path.isfile(docx_path), f"Not found: {docx_path}")
    if not os.path.isfile(docx_path):
        return

    doc = Document(docx_path)
    full_text = " ".join(p.text for p in doc.paragraphs)
    full_text_lower = full_text.lower()

    # Check headings (accept both "Heading X" and "Title" styles)
    headings = [p.text for p in doc.paragraphs
                if p.style.name.startswith("Heading") or p.style.name == "Title"]
    heading_text = " ".join(headings).lower()
    # Headings: task keeps English heading templates, but accept RU equivalents too.
    check("Word has 'Talent Review' heading",
          any(k in heading_text for k in ("talent review", "обзор талантов", "talent")),
          f"Headings: {headings[:5]}")
    check("Word has 'Department' section",
          any(k in heading_text for k in ("department", "подразделен", "highlights")),
          f"Headings: {headings[:5]}")
    check("Word has 'Salary' section",
          any(k in heading_text for k in ("salary", "competitiveness", "зарплат", "конкурентоспособн")),
          f"Headings: {headings[:5]}")

    # Check content: total headcount = 50000 (RU prose -> accept RU keyword + number)
    check("Word mentions total headcount ~50000",
          ("50000" in full_text or "50,000" in full_text or "50 000" in full_text),
          "50000/50,000/50 000 not found in text")

    # Check content: mentions promotion candidates
    total_promo = sum(d["promo"] for d in DEPT_EXPECTED.values())
    check("Word mentions promotion candidates",
          ("promo" in full_text_lower or "повыш" in full_text_lower or "кандидат" in full_text_lower
           or str(total_promo) in full_text or "10950" in full_text or "10,950" in full_text
           or "10 950" in full_text),
          f"Expected mention of ~{total_promo} promo candidates")

    # Check content: mentions all (russified) departments
    dept_count = sum(1 for d in DEPARTMENTS if d.lower() in full_text_lower)
    check("Word mentions all 7 departments", dept_count == 7,
          f"Found {dept_count}/7 departments")

    # Check: has at least one table
    check("Word contains at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")

    # Check table has department data
    if doc.tables:
        table_text = " ".join(cell.text for row in doc.tables[0].rows for cell in row.cells).lower()
        check("Word table contains department names",
              "инженерия" in table_text and "продажи" in table_text,
              "Missing russified department names in table")


def check_gcal(launch_time_str):
    print("\n=== Checking Google Calendar ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT summary, start_datetime, end_datetime, description
        FROM gcal.events
        WHERE LOWER(summary) LIKE '%talent review%'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    check("At least 7 Talent Review events created", len(events) >= 7,
          f"Found {len(events)} events")

    if launch_time_str:
        try:
            launch_dt = datetime.strptime(launch_time_str, "%Y-%m-%d %H:%M:%S")
            if launch_dt.tzinfo is None:
                launch_dt = launch_dt.replace(tzinfo=timezone.utc)
        except Exception:
            launch_dt = datetime(2026, 3, 7, 10, 0, 0, tzinfo=timezone.utc)
    else:
        launch_dt = datetime(2026, 3, 7, 10, 0, 0, tzinfo=timezone.utc)

    target_start = launch_dt + timedelta(days=14)

    # Check department coverage
    depts_found = set()
    for summary, start_dt, end_dt, desc in events:
        for dept in DEPARTMENTS:
            if dept.lower() in summary.lower():
                depts_found.add(dept)
    check("Calendar events cover all 7 departments", len(depts_found) == 7,
          f"Found {len(depts_found)}: {depts_found}", critical=True)

    # Check event timing: first event should be ~14 days from launch
    if events:
        first_start = events[0][1]
        if first_start.tzinfo is None:
            first_start = first_start.replace(tzinfo=timezone.utc)
        diff_days = abs((first_start.date() - target_start.date()).days)
        check("First meeting ~14 days from launch", diff_days <= 3,
              f"First event at {first_start.date()}, expected ~{target_start.date()}")

    # Check event duration: should be ~90 minutes
    if events:
        for summary, start_dt, end_dt, desc in events[:1]:
            if start_dt and end_dt:
                duration_min = (end_dt - start_dt).total_seconds() / 60
                check("Meeting duration ~90 minutes", abs(duration_min - 90) <= 15,
                      f"Duration: {duration_min} min")

    # Check descriptions mention headcount
    if events:
        descs_with_count = sum(1 for _, _, _, desc in events if desc and any(c.isdigit() for c in str(desc)))
        check("Event descriptions include numeric data", descs_with_count >= 5,
              f"{descs_with_count}/7 events have numeric descriptions")


def check_email():
    print("\n=== Checking Email ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    # Find the talent review email
    target_email = None
    for subj, from_addr, to_addr, body in all_emails:
        recipients = parse_recipients(to_addr)
        if "hr_leadership@company.com" in recipients:
            subj_lower = (subj or "").lower()
            if "talent review" in subj_lower or "q1 2026" in subj_lower:
                target_email = (subj, from_addr, to_addr, body)
                break

    check("Summary email sent to hr_leadership@company.com", target_email is not None,
          "No matching email found", critical=True)

    if target_email:
        subj, from_addr, to_addr, body = target_email
        body_lower = (body or "").lower()

        check("Email from talent_review@company.com",
              "talent_review@company.com" in (from_addr or "").lower(),
              f"From: {from_addr}")

        check("Email subject contains 'Talent Review'",
              "talent review" in (subj or "").lower(),
              f"Subject: {subj}")

        # Check body mentions total headcount (RU prose -> RU keyword + number).
        check("Email body mentions headcount",
              ("50000" in body or "50,000" in body or "50 000" in body
               or "headcount" in body_lower or "численност" in body_lower or "штат" in body_lower),
              "No headcount mention")

        check("Email body mentions performance or rating",
              ("performance" in body_lower or "rating" in body_lower
               or "эффективн" in body_lower or "рейтинг" in body_lower),
              "No performance/rating mention")

        check("Email body mentions promotion",
              ("promo" in body_lower or "повыш" in body_lower or "кандидат" in body_lower),
              "No promotion mention")

        check("Email body mentions salary",
              ("salary" in body_lower or "compensation" in body_lower
               or "зарплат" in body_lower or "компенсац" in body_lower),
              "No salary/compensation mention")

        # task.md requires naming highest & lowest avg-salary departments (russified).
        check(f"Email names highest-avg-salary dept ({HIGHEST_SALARY_DEPT})",
              HIGHEST_SALARY_DEPT.lower() in body_lower,
              f"Missing '{HIGHEST_SALARY_DEPT}' in body", critical=True)
        check(f"Email names lowest-avg-salary dept ({LOWEST_SALARY_DEPT})",
              LOWEST_SALARY_DEPT.lower() in body_lower,
              f"Missing '{LOWEST_SALARY_DEPT}' in body", critical=True)

    # Reverse check: russified noise email ("График ремонта офиса") should still exist.
    noise_found = any(("ремонт" in (subj or "").lower() or "renovation" in (subj or "").lower())
                      for subj, _, _, _ in all_emails)
    check("Noise email not deleted", noise_found, "Renovation/ремонт noise email missing")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("SF HR TALENT REVIEW EXCEL WORD GCAL EMAIL - EVALUATION")
    print("=" * 70)

    check_excel(args.agent_workspace)
    check_word(args.agent_workspace)
    check_gcal(args.launch_time)
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0

    print(f"\n{'='*70}")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Accuracy: {accuracy:.1f}%")

    # CRITICAL gate: any critical failure => immediate FAIL regardless of accuracy.
    if CRITICAL_FAILED:
        print(f"  CRITICAL CHECKS FAILED: {CRITICAL_FAILED}")
        print("  Overall: FAIL (critical)")
        if args.res_log_file:
            with open(args.res_log_file, "w") as f:
                json.dump({"pass": PASS_COUNT, "fail": FAIL_COUNT, "accuracy": accuracy,
                           "critical_failed": CRITICAL_FAILED, "result": "FAIL"}, f)
        sys.exit(1)

    overall = accuracy >= 70
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"pass": PASS_COUNT, "fail": FAIL_COUNT, "accuracy": accuracy,
                       "result": "PASS" if overall else "FAIL"}, f)

    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
