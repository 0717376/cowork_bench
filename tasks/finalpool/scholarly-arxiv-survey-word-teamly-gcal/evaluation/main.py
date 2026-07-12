"""
Evaluation for scholarly-arxiv-survey-word-notion-gcal task.

Checks:
1. LLM_Reasoning_Survey.docx exists and has required sections (RU+EN headings)
2. Word doc mentions all 5 reasoning methods and papers (English names preserved)
3. Word doc has comparative analysis section/table
4. Teamly tracker "Reasoning Papers" page exists with required field labels
5. Teamly has at least 5 per-paper entry pages under the tracker
6. GCal has 5 reading group sessions in the week of March 16-20, 2026

CRITICAL_CHECKS (semantic): any failure => overall FAIL regardless of accuracy.
Pass threshold otherwise: accuracy >= 70%.
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2
from docx import Document

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "cowork_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Critical checks: any failure => overall FAIL regardless of accuracy.
# These reflect the task's substance (real survey with the methods/sections,
# a real Teamly tracker with >=5 entries, and the 5 GCal sessions).
CRITICAL_CHECKS = {
    "Mentions at least 4 of 5 reasoning methods",
    "Has at least 5 required sections",
    "Teamly 'Reasoning Papers' tracker page exists",
    "At least 5 paper entry pages under the Reasoning Papers tracker",
    "At least 5 reading group events in March 16-20",
}


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Word Document LLM_Reasoning_Survey.docx ===")

    docx_path = os.path.join(agent_workspace, "LLM_Reasoning_Survey.docx")
    if not os.path.exists(docx_path):
        record("LLM_Reasoning_Survey.docx exists", False, f"Not found at {docx_path}")
        return
    record("LLM_Reasoning_Survey.docx exists", True)

    try:
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    # Get all text
    all_text = "\n".join(p.text for p in doc.paragraphs).lower()

    # Check for table
    has_table = len(doc.tables) >= 1
    table_text = ""
    if has_table:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text.lower() + " "
    combined_text = all_text + " " + table_text

    # Check required sections
    headings = [p.text.strip() for p in doc.paragraphs
                if p.style.name.startswith("Heading") and p.text.strip()]
    headings_lower = [h.lower() for h in headings]

    # Required sections matched against RU+EN heading alternatives, because the
    # agent legitimately writes the survey prose in Russian.
    def heading_has(alts):
        return any(any(a in h for a in alts) for h in headings_lower)

    has_abstract = (heading_has(["abstract", "аннотац", "реферат"])
                    or "abstract" in all_text[:500]
                    or "аннотац" in all_text[:600])
    has_intro = heading_has(["introduction", "введение"])
    has_background = heading_has(["background", "предпосылк", "общие сведения", "фон"])
    has_taxonomy = heading_has(["taxonomy", "method", "таксоном", "метод"])
    has_comparative = heading_has(["comparative", "comparison", "сравнит", "сопоставит"])
    has_challenges = heading_has(["challenge", "open", "вызов", "проблем", "открыт"])
    has_conclusion = heading_has(["conclusion", "заключение", "выводы"])

    section_count = sum([has_abstract, has_intro, has_background, has_taxonomy,
                         has_comparative, has_challenges, has_conclusion])
    record("Has at least 5 required sections", section_count >= 5,
           f"Found {section_count}/7: abstract={has_abstract}, intro={has_intro}, "
           f"background={has_background}, taxonomy={has_taxonomy}, "
           f"comparative={has_comparative}, challenges={has_challenges}, conclusion={has_conclusion}")

    # Check for 5 reasoning methods
    has_cot = "chain-of-thought" in combined_text or "chain of thought" in combined_text
    has_tot = "tree of thought" in combined_text or "tree-of-thought" in combined_text
    has_sc = "self-consistency" in combined_text or "self consistency" in combined_text
    has_verify = ("step by step" in combined_text and "verif" in combined_text) or "process supervision" in combined_text
    has_auto = "auto-cot" in combined_text or "automatic chain" in combined_text or "automatic cot" in combined_text

    method_count = sum([has_cot, has_tot, has_sc, has_verify, has_auto])
    record("Mentions at least 4 of 5 reasoning methods", method_count >= 4,
           f"Found {method_count}/5: CoT={has_cot}, ToT={has_tot}, SC={has_sc}, "
           f"Verify={has_verify}, AutoCoT={has_auto}")

    # Check for paper titles/authors
    has_wei = "wei" in combined_text
    has_yao = "yao" in combined_text
    has_wang = "wang" in combined_text
    has_lightman = "lightman" in combined_text or "let's verify" in combined_text or "lets verify" in combined_text
    has_zhang = "zhang" in combined_text and "auto" in combined_text

    author_count = sum([has_wei, has_yao, has_wang, has_lightman, has_zhang])
    record("References at least 3 paper authors", author_count >= 3,
           f"Found {author_count}/5: Wei={has_wei}, Yao={has_yao}, Wang={has_wang}, "
           f"Lightman={has_lightman}, Zhang(Auto)={has_zhang}")

    # Check comparative table or analysis
    has_comparison_content = (
        has_table or
        ("comparative" in combined_text and any(m in combined_text for m in ["accuracy", "cost", "performance"]))
    )
    record("Has comparative analysis content (table or structured comparison)",
           has_comparison_content,
           "No table found and no comparative analysis keywords")


def check_teamly_tracker():
    """Teamly is a Confluence-analog (spaces/pages, no typed-property DB), so the
    'Reasoning Papers' tracker is modelled as a parent page plus one child page
    per paper. We count entry pages that actually belong to the tracker (its
    children, or — as a fallback — pages in the same space that carry the
    required field labels), not a loose keyword scan across all pages."""
    print("\n=== Check 2: Teamly 'Reasoning Papers' tracker ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    # Locate the tracker parent page by its preserved English title.
    cur.execute("""
        SELECT id, space_id, COALESCE(body, '')
        FROM teamly.pages
        WHERE title ILIKE '%%reasoning%%paper%%'
    """)
    tracker_rows = cur.fetchall()
    tracker = tracker_rows[0] if tracker_rows else None

    record("Teamly 'Reasoning Papers' tracker page exists", tracker is not None,
           "No page whose title matches 'Reasoning Papers'")

    if tracker is None:
        cur.close()
        conn.close()
        return

    tracker_id, tracker_space, tracker_body = tracker

    # Field labels are kept English per task.md; check the tracker (parent or
    # any of its child entry pages) carries the required schema labels.
    cur.execute("""
        SELECT title, COALESCE(body, '')
        FROM teamly.pages
        WHERE id = %s OR parent_id = %s OR space_id = %s
    """, (tracker_id, tracker_id, tracker_space))
    space_rows = cur.fetchall()
    schema_text = " ".join((t + " " + b).lower() for t, b in space_rows)

    has_method = "method" in schema_text or "метод" in schema_text
    has_year = "year" in schema_text or "год" in schema_text
    has_key = "key_contribution" in schema_text or "key contribution" in schema_text \
        or "contribution" in schema_text or "вклад" in schema_text
    prop_count = sum([has_method, has_year, has_key])
    record("Tracker carries key fields (Method, Year, Key_Contribution)",
           prop_count >= 2,
           f"Method={has_method}, Year={has_year}, Key_Contribution={has_key}")

    # Count per-paper ENTRY pages belonging to the tracker. Primary signal:
    # child pages (parent_id == tracker). Each entry should reference a reasoning
    # method/paper. We do NOT count the parent or the RU noise pages.
    cur.execute("""
        SELECT title, COALESCE(body, '')
        FROM teamly.pages
        WHERE parent_id = %s
    """, (tracker_id,))
    child_rows = cur.fetchall()

    method_keywords = ["chain", "tree", "self-consistency", "self consistency",
                       "verify", "step by step", "process supervision",
                       "auto-cot", "automatic", "thought"]

    def is_entry(title, body):
        text = (title + " " + body).lower()
        return any(kw in text for kw in method_keywords)

    entry_pages = sum(1 for t, b in child_rows if is_entry(t, b))

    # Fallback for agents that put entries flat in the tracker's space rather
    # than as children: count same-space non-noise pages that look like entries
    # and are not the parent itself.
    if entry_pages < 5:
        cur.execute("""
            SELECT id, title, COALESCE(body, '')
            FROM teamly.pages
            WHERE space_id = %s AND id <> %s
        """, (tracker_space, tracker_id))
        flat_entries = sum(1 for _id, t, b in cur.fetchall() if is_entry(t, b))
        entry_pages = max(entry_pages, flat_entries)

    record("At least 5 paper entry pages under the Reasoning Papers tracker",
           entry_pages >= 5,
           f"Found {entry_pages} entry pages")

    cur.close()
    conn.close()


def check_gcal():
    print("\n=== Check 3: Google Calendar Reading Group Sessions ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    # Check for reading group events in the week of March 16-20, 2026
    cur.execute("""
        SELECT summary, start_datetime, end_datetime
        FROM gcal.events
        WHERE start_datetime >= '2026-03-16' AND start_datetime < '2026-03-21'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()

    reading_events = [
        e for e in events
        if "reading group" in (e[0] or "").lower()
    ]

    record("At least 5 reading group events in March 16-20", len(reading_events) >= 5,
           f"Found {len(reading_events)} reading group events in target week")

    if reading_events:
        # Check duration (should be ~1 hour)
        summary, start_dt, end_dt = reading_events[0]
        if start_dt and end_dt:
            duration_hours = (end_dt - start_dt).total_seconds() / 3600
            record("Reading sessions are ~1 hour", 0.5 <= duration_hours <= 1.5,
                   f"Duration: {duration_hours:.1f} hours")

        # Check they are on different days
        dates = set(e[1].date() for e in reading_events if e[1])
        record("Sessions on different days (at least 4 distinct dates)", len(dates) >= 4,
               f"Found {len(dates)} distinct dates: {sorted(dates)}")

        # Check summaries contain topic descriptors
        summaries_text = " ".join(e[0].lower() for e in reading_events)
        topic_keywords = ["chain", "tree", "self", "verif", "step", "auto", "thought", "consistency"]
        topic_matches = sum(1 for kw in topic_keywords if kw in summaries_text)
        record("Session titles contain paper topic descriptors", topic_matches >= 3,
               f"Matched {topic_matches} topic keywords in summaries: {[e[0] for e in reading_events]}")

    cur.close()
    conn.close()


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_teamly_tracker()
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    if critical_failed:
        print(f"  CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"    - {n}")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_failed": critical_failed,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if critical_failed:
        print("FAIL (critical check failed)")
        sys.exit(1)
    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
