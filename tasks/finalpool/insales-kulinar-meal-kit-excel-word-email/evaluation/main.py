"""Evaluation script for insales-kulinar-meal-kit-excel-word-email (InSales + Kulinar).

Hardened, swap-aware evaluation.

Context after the insales->insales / recipe-source->kulinar swaps:
  * Store products are read live from wc.* (InSales). Category NAMES are
    russified centrally; product names/SKUs stay English. Total_Products is a
    LIVE count recomputed here from wc.products, never a frozen literal.
  * Recipes come from the kulinar (Кулинар) MCP: recipe names are RUSSIAN
    (Cyrillic), categories are RUSSIAN (салат/суп/горячее/...), difficulty is an
    INTEGER 1-4. The agent legitimately writes these Russian strings, so we do
    NOT compare Recipe_Collection / Kit_Proposals rows against any English
    literal. Instead we validate structure: Cyrillic recipe names, valid kulinar
    categories, integer difficulty, the pricing FORMULA, and row counts.

CRITICAL (semantic) checks abort with sys.exit(1) before the accuracy gate.
PASS requires: no critical failure AND accuracy >= 70%.
"""
import argparse
import json
import os
import sys

import openpyxl

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": os.environ.get("PGUSER", "eigent"),
    "password": os.environ.get("PGPASSWORD", "camel"),
}

# Valid kulinar (Кулинар) recipe categories.
VALID_RECIPE_CATEGORIES = {
    "выпечка", "гарнир", "горячее", "десерт",
    "закуска", "напиток", "салат", "суп",
}

TARGET_RECIPIENT = "product-team@company.com"

# docx required sections matched RU+EN (agent may write either language).
REQUIRED_SECTIONS = [
    ["product analysis", "анализ товаров", "анализ продуктов"],
    ["recipe selection", "подбор рецептов", "выбор рецептов"],
    ["kit design", "дизайн наборов", "проектирование наборов"],
    ["financial projections", "финансовые прогнозы", "финансовый прогноз"],
]

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILURES = []


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRITICAL]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL]{' [CRITICAL]' if critical else ''} {name}{detail_str}")
        if critical:
            CRITICAL_FAILURES.append(name)


def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default


def has_cyrillic(s):
    return any("Ѐ" <= ch <= "ӿ" for ch in str(s))


def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)


def live_total_products():
    """Live count of store products read from wc.products (InSales).

    Returns None if the DB is unavailable, so the check degrades gracefully.
    """
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM wc.products")
        n = cur.fetchone()[0]
        conn.close()
        return int(n)
    except Exception as e:
        print(f"  WARNING: could not read live wc.products count: {e}")
        return None


def col_map(ws):
    return {str(c.value).strip().lower(): i for i, c in enumerate(ws[1]) if c.value}


def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    # ── Meal_Kit_Report.xlsx ─────────────────────────────────────────────
    excel_path = os.path.join(agent_workspace, "Meal_Kit_Report.xlsx")
    check("Meal_Kit_Report.xlsx exists", os.path.exists(excel_path), critical=True)

    summary_values = {}
    kit_rows_count = 0

    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        required_sheets = ["Product_Catalog", "Recipe_Collection", "Kit_Proposals", "Summary"]
        sheet_lookup = {s.lower(): s for s in wb.sheetnames}
        for sn in required_sheets:
            check(f"{sn} sheet exists", sn.lower() in sheet_lookup,
                  f"sheets: {wb.sheetnames}", critical=True)

        # Required English headers per sheet.
        required_headers = {
            "Product_Catalog": ["product_name", "price", "stock", "category"],
            "Recipe_Collection": ["recipe_name", "category", "difficulty", "ingredient_count"],
            "Kit_Proposals": ["kit_name", "recipe_name", "estimated_cost", "margin_pct", "recommended_price"],
            "Summary": ["metric", "value"],
        }
        for sn, hdrs in required_headers.items():
            if sn.lower() in sheet_lookup:
                ws = wb[sheet_lookup[sn.lower()]]
                cmap = col_map(ws)
                for h in hdrs:
                    check(f"{sn} has {h} column", h in cmap, f"headers: {list(cmap)}")

        # ── Product_Catalog: structural ──
        if "product_catalog" in sheet_lookup:
            ws = wb[sheet_lookup["product_catalog"]]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Product_Catalog has data rows", len(rows) >= 1, f"got {len(rows)}")

        # ── Recipe_Collection: kulinar structure (RU names, valid cats, int difficulty) ──
        if "recipe_collection" in sheet_lookup:
            ws = wb[sheet_lookup["recipe_collection"]]
            cmap = col_map(ws)
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Recipe_Collection has >= 3 recipes", len(rows) >= 3, f"got {len(rows)}",
                  critical=True)
            ni = cmap.get("recipe_name")
            ci = cmap.get("category")
            di = cmap.get("difficulty")
            cyr_ok = bad_cat = bad_diff = 0
            cats_seen = set()
            for r in rows:
                if ni is not None and ni < len(r) and has_cyrillic(r[ni]):
                    cyr_ok += 1
                if ci is not None and ci < len(r) and r[ci] is not None:
                    cat = str(r[ci]).strip().lower()
                    cats_seen.add(cat)
                    if cat not in VALID_RECIPE_CATEGORIES:
                        bad_cat += 1
                if di is not None and di < len(r):
                    dv = safe_float(r[di])
                    if dv is None or not (1 <= dv <= 4) or abs(dv - round(dv)) > 1e-6:
                        bad_diff += 1
            check("Recipe_Collection recipe names are Russian (Cyrillic)",
                  rows and cyr_ok >= max(1, len(rows) - 1), f"cyrillic rows: {cyr_ok}/{len(rows)}",
                  critical=True)
            check("Recipe_Collection categories are valid kulinar categories",
                  bad_cat == 0, f"unexpected categories: {cats_seen - VALID_RECIPE_CATEGORIES}",
                  critical=True)
            check("Recipe_Collection difficulty is integer 1-4",
                  bad_diff == 0, f"{bad_diff} rows out of range/non-integer")
            check("Recipe_Collection covers >= 2 distinct categories",
                  len(cats_seen & VALID_RECIPE_CATEGORIES) >= 2, f"categories: {cats_seen}")

        # ── Kit_Proposals: >=5 rows AND pricing formula consistency ──
        if "kit_proposals" in sheet_lookup:
            ws = wb[sheet_lookup["kit_proposals"]]
            cmap = col_map(ws)
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            rows = [r for r in rows if r and any(v is not None for v in r)]
            kit_rows_count = len(rows)
            check("Kit_Proposals has >= 5 proposals", kit_rows_count >= 5,
                  f"got {kit_rows_count}", critical=True)
            ei = cmap.get("estimated_cost")
            mi = cmap.get("margin_pct")
            ri = cmap.get("recommended_price")
            bad_price = checked = 0
            if ei is not None and mi is not None and ri is not None:
                for r in rows:
                    cost = safe_float(r[ei]) if ei < len(r) else None
                    margin = safe_float(r[mi]) if mi < len(r) else None
                    rec = safe_float(r[ri]) if ri < len(r) else None
                    if cost is None or margin is None or rec is None:
                        continue
                    if margin >= 100:
                        continue
                    expected = cost / (1 - margin / 100.0)
                    checked += 1
                    if abs(rec - expected) > max(1.0, expected * 0.05):
                        bad_price += 1
            check("Kit_Proposals Recommended_Price == Estimated_Cost/(1-Margin) for every row",
                  checked >= 5 and bad_price == 0,
                  f"checked={checked}, inconsistent={bad_price}", critical=True)

        # ── Summary ──
        if "summary" in sheet_lookup:
            ws = wb[sheet_lookup["summary"]]
            for r in ws.iter_rows(min_row=2, values_only=True):
                if r and r[0]:
                    summary_values[str(r[0]).strip().lower()] = r[1] if len(r) > 1 else None
            for metric in ["total_products", "total_recipes", "proposed_kits",
                           "avg_kit_cost", "avg_margin_pct"]:
                check(f"Summary has {metric}", metric in summary_values,
                      f"metrics: {list(summary_values)}")

            # CRITICAL: Total_Products == live wc.products count.
            live_tp = live_total_products()
            tp = safe_float(summary_values.get("total_products"))
            if live_tp is not None:
                check("Summary Total_Products == live wc.products count",
                      tp is not None and abs(tp - live_tp) <= 1,
                      f"reported {tp}, live {live_tp}", critical=True)
            else:
                check("Summary Total_Products is a positive number",
                      tp is not None and tp > 0, f"got {tp}")

            # Proposed_Kits must match the Kit_Proposals row count.
            pk = safe_float(summary_values.get("proposed_kits"))
            check("Summary Proposed_Kits == Kit_Proposals row count",
                  pk is not None and kit_rows_count > 0 and abs(pk - kit_rows_count) <= 0,
                  f"reported {pk}, rows {kit_rows_count}", critical=True)

    # ── Meal_Kit_Proposal.docx ───────────────────────────────────────────
    docx_path = os.path.join(agent_workspace, "Meal_Kit_Proposal.docx")
    check("Meal_Kit_Proposal.docx exists", os.path.exists(docx_path), critical=True)
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join(p.text for p in doc.paragraphs)
        check("Meal_Kit_Proposal.docx has content", len(text) > 50, f"len {len(text)}")
        headings = [p.text.strip().lower() for p in doc.paragraphs
                    if p.style.name.startswith("Heading")]
        all_text_lower = text.lower()
        missing = []
        for variants in REQUIRED_SECTIONS:
            found = any(any(v in h for h in headings) for v in variants) or \
                    any(v in all_text_lower for v in variants)
            if not found:
                missing.append(variants[0])
        check("Meal_Kit_Proposal.docx has all 4 required sections (RU/EN)",
              len(missing) == 0, f"missing: {missing}", critical=True)

    # ── Python analysis script ──
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    check("Python analysis script exists", len(py_files) >= 1, f"found: {py_files}")

    # ── Database / email checks ──
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute(
            "SELECT subject, to_addr FROM email.messages "
            "WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1) "
            "AND subject ILIKE '%meal%'"
        )
        email_row = cur.fetchone()
        check("Email with subject containing 'meal' sent", email_row is not None,
              "no matching email in Sent", critical=True)
        if email_row:
            to_raw = email_row[1]
            # to_addr may be a json/text array or a plain string.
            to_text = json.dumps(to_raw) if not isinstance(to_raw, str) else to_raw
            check("Email sent to exactly product-team@company.com",
                  TARGET_RECIPIENT in (to_text or ""),
                  f"to_addr: {to_text}", critical=True)

        # Reverse verification: injected noise emails must NOT be in Sent.
        cur.execute(
            "SELECT COUNT(*) FROM email.messages "
            "WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1) "
            "AND message_id LIKE '%noise-%'"
        )
        noise_sent = cur.fetchone()[0]
        check("No injected noise emails in Sent folder", noise_sent == 0,
              f"found {noise_sent} noise emails in Sent")
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e))

    # ── Gate ─────────────────────────────────────────────────────────────
    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100.0) if total else 0.0
    critical_ok = len(CRITICAL_FAILURES) == 0
    success = critical_ok and accuracy >= 70.0

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}  Failed: {FAIL_COUNT}  Accuracy: {accuracy:.1f}%")
    print(f"  Critical failures: {CRITICAL_FAILURES}")

    if res_log_file:
        try:
            with open(res_log_file, "w") as f:
                json.dump({
                    "passed": PASS_COUNT, "failed": FAIL_COUNT, "accuracy": accuracy,
                    "success": success, "critical_failures": CRITICAL_FAILURES,
                }, f, indent=2)
        except Exception:
            pass

    if not critical_ok:
        print(f"=== RESULT: FAIL (critical checks failed: {CRITICAL_FAILURES}) ===")
        sys.exit(1)

    return success, f"Passed {PASS_COUNT}/{total} checks, accuracy {accuracy:.1f}%"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
