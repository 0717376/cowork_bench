"""
Evaluation script for sf-support-priority-review-word task (ClickHouse, sf_data schema).

Expected values are computed LIVE from the source tables
sf_data."SUPPORT_CENTER__PUBLIC__TICKETS" and
sf_data."SUPPORT_CENTER__PUBLIC__SLA_POLICIES" -- nothing is hardcoded.

Critical (semantic) checks (any fail => FAIL regardless of accuracy):
1. Word doc contains the avg RESPONSE_TIME_HOURS per priority (High/Medium/Low), within tolerance.
2. Word doc contains the SLA target response hours per priority.
3. Google Sheet 'By Priority' has the required header columns and exactly one
   row per priority whose Ticket_Count matches COUNT(*) GROUP BY PRIORITY.
4. Email from support-analytics to support-director with subject
   'Priority Handling Performance Report' and a non-trivial body.
5. Word doc contains the per-priority ticket counts (core analytical deliverable).

Structural (non-critical) checks: file exists, sheet/title present, ISO date,
RU/EN keyword presence, etc.
"""

import argparse
import json
import os
import re
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "cowork_gym",
    "user": "eigent",
    "password": "camel",
}

# PRIORITY enum values stay English (eval greps these; not in russification map).
PRIORITIES = ["High", "Medium", "Low"]

# Accept RU + EN phrasing where the agent legitimately writes Russian prose.
RESPONSE_TIME_TERMS = ["response time", "время отклика", "время ответа", "врем"]

CHECKS = []  # list of (name, passed, critical)


def add(name, passed, critical=False):
    CHECKS.append((name, bool(passed), critical))


def num_present_in_text(value, text, tol=0.5):
    """True if a number within tol of `value` appears in `text`.

    Scans every numeric token (int or decimal) in the text and checks
    closeness, so it tolerates rounding/formatting differences.
    """
    try:
        target = float(value)
    except (TypeError, ValueError):
        return False
    # Pre-pass: integers written with comma/space thousands separators
    # (e.g. "6,466" / "6 466" -> 6466). Done first so they aren't mis-read
    # as European decimals by the token scan below.
    for m in re.finditer(r"\d{1,3}(?:[ ,]\d{3})+", text):
        try:
            if abs(int(m.group(0).replace(",", "").replace(" ", "")) - target) <= tol:
                return True
        except ValueError:
            continue
    for m in re.finditer(r"\d+(?:[.,]\d+)?", text):
        tok = m.group(0).replace(",", ".")
        try:
            if abs(float(tok) - target) <= tol:
                return True
        except ValueError:
            continue
    return False


def fetch_source_stats():
    """Return (counts, avg_resp, avg_sat, sla_target) keyed by lowercase priority."""
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute(
        'SELECT "PRIORITY", COUNT(*), AVG("RESPONSE_TIME_HOURS"), '
        'AVG("CUSTOMER_SATISFACTION") '
        'FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS" GROUP BY "PRIORITY"'
    )
    counts, avg_resp, avg_sat = {}, {}, {}
    for prio, cnt, resp, sat in cur.fetchall():
        k = str(prio).strip().lower()
        counts[k] = int(cnt)
        avg_resp[k] = float(resp) if resp is not None else None
        avg_sat[k] = float(sat) if sat is not None else None
    cur.execute(
        'SELECT "PRIORITY", "RESPONSE_TARGET_HOURS" '
        'FROM sf_data."SUPPORT_CENTER__PUBLIC__SLA_POLICIES"'
    )
    sla_target = {}
    for prio, tgt in cur.fetchall():
        sla_target[str(prio).strip().lower()] = float(tgt) if tgt is not None else None
    cur.close()
    conn.close()
    return counts, avg_resp, avg_sat, sla_target


def check_word(agent_workspace, counts, avg_resp, sla_target):
    print("\n=== Checking Word Output ===")
    docx_path = os.path.join(agent_workspace, "Priority_Handling_Report.docx")
    exists = os.path.isfile(docx_path)
    add("Word file exists", exists)
    if not exists:
        # core semantic deliverables fail too
        add("Word: avg response hours per priority present", False, critical=True)
        add("Word: SLA target hours per priority present", False, critical=True)
        add("Word: per-priority ticket counts present", False, critical=True)
        return

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        add("Word file readable", False)
        add("Word: avg response hours per priority present", False, critical=True)
        add("Word: SLA target hours per priority present", False, critical=True)
        add("Word: per-priority ticket counts present", False, critical=True)
        print(f"  docx read error: {e}")
        return

    text = ""
    for para in doc.paragraphs:
        text += para.text + " "
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
    low = text.lower()

    # --- Structural (non-critical) ---
    add("Word mentions response-time term (RU/EN)",
        any(t in low for t in RESPONSE_TIME_TERMS))
    for p in PRIORITIES:
        add(f"Word contains priority level '{p}'", p.lower() in low)
    add("Word title 'priority' + 'analysis' present",
        "priority" in low and "analysis" in low)
    add("Word contains date 2026-03-06", "2026-03-06" in low)

    # --- CRITICAL semantic checks (computed live) ---
    # 5. per-priority ticket counts
    counts_ok = all(
        counts.get(p.lower()) is not None
        and num_present_in_text(counts[p.lower()], text, tol=0.0)
        for p in PRIORITIES
    )
    add("Word: per-priority ticket counts present", counts_ok, critical=True)

    # 1. avg response hours per priority
    resp_ok = all(
        avg_resp.get(p.lower()) is not None
        and num_present_in_text(round(avg_resp[p.lower()], 1), text, tol=0.5)
        for p in PRIORITIES
    )
    add("Word: avg response hours per priority present", resp_ok, critical=True)

    # 2. SLA target hours per priority
    sla_ok = all(
        sla_target.get(p.lower()) is not None
        and num_present_in_text(sla_target[p.lower()], text, tol=0.5)
        for p in PRIORITIES
    )
    add("Word: SLA target hours per priority present", sla_ok, critical=True)


def check_gsheet(counts):
    print("\n=== Checking Google Sheet ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT id, title FROM gsheet.spreadsheets")
    spreadsheets = cur.fetchall()

    matching = [s for s in spreadsheets if s[1] and "priority" in s[1].lower()]
    add("Spreadsheet with 'priority' in title exists", len(matching) > 0)

    # Locate a sheet named 'By Priority' (case-insensitive) and read its cells.
    target_sheet_id = None
    for ss_id, _title in matching:
        cur.execute(
            "SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (ss_id,)
        )
        for sh_id, sh_title in cur.fetchall():
            if sh_title and sh_title.strip().lower() == "by priority":
                target_sheet_id = sh_id
                break
        if target_sheet_id:
            break

    sheet_found = target_sheet_id is not None
    add("Sheet 'By Priority' exists", sheet_found)

    if not sheet_found:
        add("Sheet 'By Priority': per-priority Ticket_Count matches source",
            False, critical=True)
        cur.close()
        conn.close()
        return

    cur.execute(
        "SELECT row_index, col_index, value FROM gsheet.cells WHERE sheet_id = %s",
        (target_sheet_id,),
    )
    cells = cur.fetchall()
    cur.close()
    conn.close()

    # Build a grid: grid[row][col] = value
    grid = {}
    max_row = 0
    for r, c, v in cells:
        grid.setdefault(r, {})[c] = v
        max_row = max(max_row, r)

    # Header row = smallest row index present.
    header_rows = sorted(grid.keys())
    header = {}
    header_row_idx = None
    if header_rows:
        header_row_idx = header_rows[0]
        header = {c: str(v).strip() for c, v in grid[header_row_idx].items() if v is not None}
    header_vals = [h.lower() for h in header.values()]
    required_cols = ["priority", "ticket_count", "avg_response_hours",
                     "avg_satisfaction", "sla_target"]
    cols_ok = all(rc in header_vals for rc in required_cols)
    add("Sheet 'By Priority' has required header columns", cols_ok)

    # Map column name -> col index
    name_to_col = {h.lower(): c for c, h in header.items()}
    prio_col = name_to_col.get("priority")
    count_col = name_to_col.get("ticket_count")

    # Read data rows (everything below the header row).
    data_rows = [r for r in header_rows if header_row_idx is not None and r > header_row_idx]
    found_priorities = {}
    for r in data_rows:
        prio_val = grid[r].get(prio_col) if prio_col is not None else None
        cnt_val = grid[r].get(count_col) if count_col is not None else None
        if prio_val is None:
            continue
        found_priorities[str(prio_val).strip().lower()] = cnt_val

    # exactly 3 data rows, one per priority
    add("Sheet 'By Priority' has exactly 3 priority data rows",
        len(found_priorities) == 3)

    # CRITICAL: each priority's Ticket_Count matches COUNT(*) from source.
    count_match = True
    for p in PRIORITIES:
        k = p.lower()
        src = counts.get(k)
        got = found_priorities.get(k)
        try:
            if src is None or got is None or int(float(got)) != int(src):
                count_match = False
        except (TypeError, ValueError):
            count_match = False
    add("Sheet 'By Priority': per-priority Ticket_Count matches source",
        count_match, critical=True)


def check_email():
    print("\n=== Checking Email ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT subject, to_addr, from_addr, body_text FROM email.messages")
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    def to_lower_join(addr):
        if isinstance(addr, list):
            return " ".join(str(r).lower() for r in addr)
        if isinstance(addr, str):
            try:
                parsed = json.loads(addr)
                if isinstance(parsed, list):
                    return " ".join(str(r).lower() for r in parsed)
            except (json.JSONDecodeError, TypeError):
                pass
            return addr.lower()
        return str(addr).lower()

    director_found = False
    semantic_ok = False
    for subject, to_addr, from_addr, body_text in all_emails:
        to_str = to_lower_join(to_addr)
        from_str = to_lower_join(from_addr)
        subj = (subject or "").lower()
        body = (body_text or "").lower()
        if "support-director@company.com" not in to_str:
            continue
        director_found = True
        # structural
        add("Email subject contains 'priority' or 'handling'",
            "priority" in subj or "handling" in subj)
        # CRITICAL: correct from, exact subject, non-trivial body mentioning response time
        subject_ok = "priority handling performance report" in subj
        from_ok = "support-analytics@company.com" in from_str
        body_ok = len(body.strip()) >= 40 and any(t in body for t in RESPONSE_TIME_TERMS)
        semantic_ok = subject_ok and from_ok and body_ok
        break

    add("Email sent to support-director@company.com", director_found)
    add("Email from support-analytics with exact subject + response-time body",
        semantic_ok, critical=True)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    # Compute expected values live from the source warehouse.
    try:
        counts, avg_resp, avg_sat, sla_target = fetch_source_stats()
    except Exception as e:
        print(f"FATAL: could not read source sf_data tables: {e}")
        sys.exit(1)

    check_word(args.agent_workspace, counts, avg_resp, sla_target)
    check_gsheet(counts)
    check_email()

    total = len(CHECKS)
    passed = sum(1 for _, p, _ in CHECKS if p)
    accuracy = (passed / total * 100.0) if total else 0.0
    critical_fail = [n for n, p, c in CHECKS if c and not p]

    print("\n=== CHECKS ===")
    for name, p, c in CHECKS:
        tag = "CRIT" if c else "    "
        print(f"  [{'PASS' if p else 'FAIL'}] {tag} {name}")
    print(f"\nAccuracy: {passed}/{total} = {accuracy:.1f}%")

    all_ok = (not critical_fail) and accuracy >= 70

    if args.res_log_file:
        result = {
            "passed": passed,
            "failed": total - passed,
            "accuracy": accuracy,
            "critical_failures": critical_fail,
            "success": all_ok,
        }
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if critical_fail:
        print(f"=== RESULT: FAIL (critical checks failed: {critical_fail}) ===")
        sys.exit(1)
    if accuracy >= 70:
        print("=== RESULT: PASS ===")
        sys.exit(0)
    print(f"=== RESULT: FAIL (accuracy {accuracy:.1f}% < 70%) ===")
    sys.exit(1)


if __name__ == "__main__":
    main()
