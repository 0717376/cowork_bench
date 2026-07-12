"""Evaluation script for pw-sf-benchmark-excel-word-email."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Critical (semantic) checks: any failure => overall FAIL regardless of accuracy.
# These reflect the core data-warehouse deliverable, gap-classification correctness,
# correct email routing, the Word narrative, and the N/A business rule.
CRITICAL_CHECKS = {
    "Internal Metrics: Avg_Salary value matches",
    "Internal Metrics: Avg_Order_Value value matches",
    "Internal Metrics: Avg_Satisfaction value matches",
    "Gap Analysis: Avg_Order_Value classification matches",
    "Gap Analysis: Avg_Salary classification matches",
    "Gap Analysis: Avg_Satisfaction classification matches",
    "Gap Analysis: Revenue_Per_Employee classification matches",
    "Customer_Retention_Rate marked N/A in Internal Metrics",
    "SLA_Compliance_Rate marked N/A in Internal Metrics",
    "Email to hr_director routes HR figures (salary/satisfaction)",
    "Email to sales_vp routes Sales figures (order value/revenue)",
    "Word doc surfaces internal salary value and key classifications",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")


def safe_float(val, default=None):
    try:
        if val is None:
            return default
        s = str(val).replace(',', '').replace('%', '').replace('$', '').strip()
        if s.upper() == 'N/A':
            return default
        return float(s)
    except (ValueError, TypeError):
        return default


def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)


def values_close(a, b, tolerance=0.15):
    """Check if two numeric values are within tolerance (relative)."""
    if a is None or b is None:
        return False
    if b == 0:
        return abs(a) < 1.0
    return abs(a - b) / max(abs(b), 1e-6) <= tolerance


# Discriminating keyword aliases per canonical GT metric key. The guide maps
# Avg_Satisfaction to the dashboard name "Employee Satisfaction" (and similar
# human-readable labels), so the agent's row label legitimately differs from the
# GT's snake_case token. Match on the discriminating keyword instead of requiring
# the exact `avg`-prefixed token as a substring. EN + RU aliases supported.
METRIC_ALIASES = {
    'Avg_Salary': ['salary', 'зарплат', 'оклад'],
    'Avg_Order_Value': ['order', 'заказ', 'стоимост'],
    'Avg_Satisfaction': ['satisfaction', 'удовлетвор'],
    'Revenue_Per_Employee': ['revenue', 'выручк', 'доход'],
    'Customer_Retention_Rate': ['retention', 'удержан', 'отток'],
    'SLA_Compliance_Rate': ['sla', 'соблюдение sla', 'уровень обслуживан'],
}


def find_agent_entry(metric_name, agent_map):
    """Return the agent value/label for a canonical GT metric key.

    Resolution order (correctness checks downstream stay strict on the returned
    value): (1) exact-token substring match against the snake_case GT key,
    (2) any discriminating alias keyword as a substring of the agent label,
    (3) exact key lookup.
    """
    gt_token = metric_name.lower().replace('_', '')
    for k, v in agent_map.items():
        if gt_token in k.lower().replace('_', '').replace(' ', ''):
            return v
    for alias in METRIC_ALIASES.get(metric_name, []):
        for k, v in agent_map.items():
            if alias in k.lower():
                return v
    return agent_map.get(metric_name)


def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT, FAILED_NAMES
    PASS_COUNT = 0
    FAIL_COUNT = 0
    FAILED_NAMES = []

    # --- Excel Evaluation ---
    excel_path = os.path.join(agent_workspace, "Benchmark_Analysis.xlsx")
    check("Benchmark_Analysis.xlsx exists", os.path.exists(excel_path))

    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Benchmark_Analysis.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        # Sheet 1: Internal Metrics
        check("Internal Metrics sheet exists", "Internal Metrics" in wb.sheetnames,
              f"sheets: {wb.sheetnames}")
        if "Internal Metrics" in wb.sheetnames:
            ws = wb["Internal Metrics"]
            headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Internal Metrics has >= 4 rows", len(data_rows) >= 4, f"got {len(data_rows)}")

            for expected_col in ['metric', 'internal_value', 'source_note']:
                check(f"Internal Metrics has '{expected_col}' column",
                      expected_col in headers, f"headers: {headers}")

            # Verify key metric values against groundtruth
            if gt_wb and "Internal Metrics" in gt_wb.sheetnames:
                gt_ws = gt_wb["Internal Metrics"]
                gt_data = {str(r[0]).strip(): r[1] for r in gt_ws.iter_rows(min_row=2, values_only=True) if r[0]}
                agent_data = {str(r[0]).strip(): r[1] for r in data_rows if r and r[0]}

                for metric_name in ['Avg_Salary', 'Avg_Order_Value', 'Avg_Satisfaction']:
                    gt_val = safe_float(gt_data.get(metric_name))
                    # Resolve the agent row via token/alias-aware lookup, then
                    # compare the value strictly.
                    agent_val = safe_float(find_agent_entry(metric_name, agent_data))
                    check(f"Internal Metrics: {metric_name} value matches",
                          values_close(agent_val, gt_val),
                          f"agent={agent_val}, gt={gt_val}")

                # Business rule: retention & SLA must be N/A (not directly calculable)
                for na_metric in ['Customer_Retention_Rate', 'SLA_Compliance_Rate']:
                    raw = find_agent_entry(na_metric, agent_data)
                    is_na = raw is not None and 'n/a' in str(raw).strip().lower()
                    check(f"{na_metric} marked N/A in Internal Metrics", is_na,
                          f"value='{raw}'")

        # Sheet 2: Gap Analysis
        check("Gap Analysis sheet exists", "Gap Analysis" in wb.sheetnames,
              f"sheets: {wb.sheetnames}")
        if "Gap Analysis" in wb.sheetnames:
            ws = wb["Gap Analysis"]
            headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Gap Analysis has >= 4 rows", len(data_rows) >= 4, f"got {len(data_rows)}")

            for expected_col in ['metric', 'internal_value', 'industry_avg', 'classification', 'priority']:
                check(f"Gap Analysis has '{expected_col}' column",
                      expected_col in headers, f"headers: {headers}")

            # Check classifications
            if gt_wb and "Gap Analysis" in gt_wb.sheetnames:
                gt_ws = gt_wb["Gap Analysis"]
                gt_class = {}
                for r in gt_ws.iter_rows(min_row=2, values_only=True):
                    if r and r[0]:
                        gt_class[str(r[0]).strip()] = str(r[7]).strip() if r[7] else ""

                agent_class = {}
                class_col_idx = None
                for i, h in enumerate(headers):
                    if 'classification' in h:
                        class_col_idx = i
                        break
                if class_col_idx is not None:
                    for r in data_rows:
                        if r and r[0]:
                            agent_class[str(r[0]).strip()] = str(r[class_col_idx]).strip() if r[class_col_idx] else ""

                for metric_name in ['Avg_Order_Value', 'Avg_Salary', 'Avg_Satisfaction', 'Revenue_Per_Employee']:
                    gt_c = gt_class.get(metric_name, "")
                    agent_c = find_agent_entry(metric_name, agent_class)
                    if agent_c is None:
                        agent_c = ""
                    check(f"Gap Analysis: {metric_name} classification matches",
                          gt_c.lower() == str(agent_c).lower(),
                          f"agent='{agent_c}', gt='{gt_c}'")

        # Sheet 3: Action Plan
        check("Action Plan sheet exists", "Action Plan" in wb.sheetnames,
              f"sheets: {wb.sheetnames}")
        if "Action Plan" in wb.sheetnames:
            ws = wb["Action Plan"]
            headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Action Plan has >= 2 rows", len(data_rows) >= 2, f"got {len(data_rows)}")

            for expected_col in ['metric', 'classification', 'recommended_action']:
                check(f"Action Plan has '{expected_col}' column",
                      expected_col in headers, f"headers: {headers}")

            # Priority-ascending ordering: the two Critical metrics must precede
            # the two Moderate metrics. Derive priority from the Classification
            # column (Critical=1, Moderate=2) since Action Plan need not carry a
            # numeric Priority column.
            class_idx = None
            for i, h in enumerate(headers):
                if 'classification' in h:
                    class_idx = i
                    break
            if class_idx is not None:
                order = []
                for r in data_rows:
                    if r and r[0] and class_idx < len(r) and r[class_idx]:
                        c = str(r[class_idx]).strip().lower()
                        if 'critical' in c or 'критич' in c:
                            order.append(1)
                        elif 'moderate' in c or 'умерен' in c:
                            order.append(2)
                        elif 'leading' in c or 'лидир' in c or 'on track' in c or 'в норме' in c:
                            order.append(0)
                        else:
                            order.append(99)
                check("Action Plan sorted by priority (ascending)",
                      len(order) >= 2 and order == sorted(order),
                      f"order={order}")

    # --- Word Document Evaluation ---
    word_path = os.path.join(agent_workspace, "Benchmark_Report.docx")
    check("Benchmark_Report.docx exists", os.path.exists(word_path))

    if os.path.exists(word_path):
        from docx import Document
        doc = Document(word_path)
        full_text = "\n".join(
            [p.text for p in doc.paragraphs]
            + [c.text for t in doc.tables for r in t.rows for c in r.cells])
        full_lower = full_text.lower()

        # Agent may write Russian section headings; accept EN or RU equivalents.
        required_sections = [
            ("Executive Summary", "Краткое резюме", "Резюме для руководства"),
            ("Methodology", "Методология"),
            ("Internal Performance Overview", "Обзор внутренних показателей", "Внутренние показатели"),
            ("Benchmark Comparison", "Сравнение с эталоном", "Сравнительный анализ"),
            ("Gap Analysis", "Анализ отклонений", "Анализ разрывов"),
            ("Strategic Recommendations", "Стратегические рекомендации"),
            ("Implementation Timeline", "План внедрения", "Сроки внедрения"),
        ]
        for variants in required_sections:
            found = any(v.lower() in full_lower for v in variants)
            check(f"Word doc has '{variants[0]}' section",
                  found, "section not found")

        # Key values mentioned (clickhouse-preserved numeric AVGs, language-neutral)
        salary_found = "58396" in full_text or "58,396" in full_text or "58 396" in full_text
        order_found = "152.45" in full_text or "152,45" in full_text or "152" in full_text
        check("Word doc mentions avg salary value", salary_found, "avg salary not found")
        check("Word doc mentions avg order value", order_found, "avg order value not found")
        crit_found = "critical" in full_lower or "критич" in full_lower
        mod_found = "moderate" in full_lower or "умерен" in full_lower
        check("Word doc mentions 'Critical' classification", crit_found)
        check("Word doc mentions 'Moderate' classification", mod_found)

        # CRITICAL narrative deliverable: surfaces internal salary number AND both
        # the Critical and Moderate classifications (EN or RU).
        check("Word doc surfaces internal salary value and key classifications",
              salary_found and order_found and crit_found and mod_found,
              "missing salary/order value or Critical/Moderate classifications")

    # --- Email Evaluation ---
    try:
        conn = get_conn()
        cur = conn.cursor()

        # Check for operations head email
        cur.execute("""SELECT subject, to_addr, body_text FROM email.messages
            WHERE to_addr::text ILIKE %s""", ('%operations_head%',))
        ops_emails = cur.fetchall()
        check("Email to operations_head@company.com sent", len(ops_emails) >= 1,
              f"found {len(ops_emails)}")

        # Check for HR director email
        cur.execute("""SELECT subject, to_addr, body_text FROM email.messages
            WHERE to_addr::text ILIKE %s""", ('%hr_director%',))
        hr_emails = cur.fetchall()
        check("Email to hr_director@company.com sent", len(hr_emails) >= 1,
              f"found {len(hr_emails)}")

        # HR email body must route HR-relevant metrics: salary AND satisfaction.
        hr_body = " ".join((row[2] or "") for row in hr_emails).lower()
        hr_salary = ("58396" in hr_body or "58,396" in hr_body or "58 396" in hr_body
                     or "salary" in hr_body or "зарплат" in hr_body or "оклад" in hr_body)
        hr_satis = ("6.55" in hr_body or "6,55" in hr_body
                    or "satisfaction" in hr_body or "удовлетвор" in hr_body)
        check("Email to hr_director routes HR figures (salary/satisfaction)",
              len(hr_emails) >= 1 and hr_salary and hr_satis,
              f"salary={hr_salary}, satisfaction={hr_satis}")

        # Check for sales VP email
        cur.execute("""SELECT subject, to_addr, body_text FROM email.messages
            WHERE to_addr::text ILIKE %s""", ('%sales_vp%',))
        sales_emails = cur.fetchall()
        check("Email to sales_vp@company.com sent", len(sales_emails) >= 1,
              f"found {len(sales_emails)}")

        # Sales email body must route sales-relevant metrics: order value AND revenue.
        sales_body = " ".join((row[2] or "") for row in sales_emails).lower()
        sales_order = ("152.45" in sales_body or "152,45" in sales_body or "152" in sales_body
                       or "order value" in sales_body or "стоимост" in sales_body
                       or "order" in sales_body or "заказ" in sales_body)
        sales_rev = ("revenue" in sales_body or "выручк" in sales_body or "доход" in sales_body
                     or "60.98" in sales_body or "60,98" in sales_body)
        check("Email to sales_vp routes Sales figures (order value/revenue)",
              len(sales_emails) >= 1 and sales_order and sales_rev,
              f"order={sales_order}, revenue={sales_rev}")

        # Check email subjects
        cur.execute("""SELECT subject FROM email.messages
            WHERE subject ILIKE %s""", ('%benchmark%',))
        benchmark_emails = cur.fetchall()
        check("At least 3 benchmark-related emails sent", len(benchmark_emails) >= 3,
              f"found {len(benchmark_emails)}")

        conn.close()
    except Exception as e:
        check("Email verification", False, str(e))

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    print(f"\n=== Results: {PASS_COUNT}/{total} passed ({accuracy:.1f}%) ===")
    if critical_failed:
        print(f"CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")

    if res_log_file:
        try:
            with open(res_log_file, "w") as f:
                json.dump({
                    "total_passed": PASS_COUNT,
                    "total_checks": total,
                    "accuracy": accuracy,
                    "critical_failed": critical_failed,
                }, f, indent=2)
        except Exception:
            pass

    success = (not critical_failed) and accuracy >= 70
    return success, f"Passed {PASS_COUNT}/{total} checks ({accuracy:.1f}%)"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    print("PASS" if success else "FAIL")
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
