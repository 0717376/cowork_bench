"""Evaluation for sf-support-resolution-word-gcal (ClickHouse / sf_data)."""
import argparse
import json
import os
import re
import sys

import psycopg2


def normalize_ru_numbers(text):
    """RU number normalization for substring/regex checks: collapse digit-group
    separators (space/NBSP/NNBSP/dot/comma before a 3-digit group) and turn
    decimal commas into dots ("31 588" -> "31588", "4 586,91" -> "4586.91")."""
    t = str(text or "")
    t = re.sub(r"(?<=\d)[ \xa0\u202f\u2009.,](?=\d{3}\b)", "", t)
    return re.sub(r"(?<=\d),(?=\d)", ".", t)


DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Semantic checks. Any failure here => FAIL regardless of accuracy.
CRITICAL_CHECKS = {
    "Priority High ticket count",
    "Priority High avg hours",
    "Priority Medium ticket count",
    "Priority Medium avg hours",
    "Priority Low ticket count",
    "Priority Low avg hours",
    "Top category present with correct count and avg hours",
    "Summary states overall resolved total",
    "Summary states overall average resolution time",
    "Calendar event on 2026-03-12 from 14:00 to 15:00",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def get_priority_data(cur):
    cur.execute("""
        SELECT "PRIORITY", COUNT(*),
               ROUND(AVG(EXTRACT(EPOCH FROM ("RESOLVED_AT" - "CREATED_AT"))/3600)::numeric, 2)
        FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"
        WHERE "STATUS" = 'Resolved' AND "RESOLVED_AT" IS NOT NULL
        GROUP BY "PRIORITY"
        ORDER BY "PRIORITY"
    """)
    return cur.fetchall()


def get_category_data(cur):
    cur.execute("""
        SELECT "ISSUE_TYPE", COUNT(*),
               ROUND(AVG(EXTRACT(EPOCH FROM ("RESOLVED_AT" - "CREATED_AT"))/3600)::numeric, 2)
        FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"
        WHERE "STATUS" = 'Resolved' AND "RESOLVED_AT" IS NOT NULL
        GROUP BY "ISSUE_TYPE"
        ORDER BY 3 DESC
    """)
    return cur.fetchall()


def get_overall(cur):
    cur.execute("""
        SELECT COUNT(*),
               ROUND(AVG(EXTRACT(EPOCH FROM ("RESOLVED_AT" - "CREATED_AT"))/3600)::numeric, 2)
        FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"
        WHERE "STATUS" = 'Resolved' AND "RESOLVED_AT" IS NOT NULL
    """)
    return cur.fetchone()


def cell_has_int(cells, target, tol):
    for cell in cells:
        try:
            val = int(cell.replace(",", "").replace(" ", ""))
            if abs(val - target) <= tol:
                return True
        except (ValueError, AttributeError):
            continue
    return False


def cell_has_float(cells, target, tol):
    for cell in cells:
        try:
            val = float(cell.replace(",", "").replace(" ", ""))
            if num_close(val, float(target), tol):
                return True
        except (ValueError, AttributeError):
            continue
    return False


def check_word_doc(agent_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Resolution_Analysis.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Heading: accept EN ("resolution"+"analysis") or RU equivalents.
    full_text = " ".join(p.text for p in doc.paragraphs)
    low = full_text.lower()
    has_heading = (("resolution" in low and "analysis" in low) or
                   ("реш" in low and "анализ" in low) or
                   ("тикет" in low and "анализ" in low))
    check("Document has resolution analysis heading", has_heading)

    check("Document has at least 2 tables", len(doc.tables) >= 2,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 2:
        return False

    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    priority_data = get_priority_data(cur)
    category_data = get_category_data(cur)
    overall_count, overall_avg = get_overall(cur)
    cur.close()
    conn.close()

    # ---- Priority table (table[0]) ----
    table1 = doc.tables[0]
    rows = []
    for row in table1.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    check("Priority table has 3 rows", len(rows) == 3, f"Got {len(rows)} rows")

    for priority, count, avg_hours in priority_data:
        matched = None
        for r in rows:
            if r and r[0].lower() == priority.lower():
                matched = r
                break
        if matched:
            check(f"Priority {priority} ticket count",
                  cell_has_int(matched[1:], count, 5), f"Expected ~{count}")
            check(f"Priority {priority} avg hours",
                  cell_has_float(matched[1:], float(avg_hours), 0.5),
                  f"Expected ~{float(avg_hours)}")
        else:
            check(f"Priority {priority} ticket count", False, "row not found")
            check(f"Priority {priority} avg hours", False, "row not found")

    # ---- Category table (table[1]) ----
    # Find the table that contains category (ISSUE_TYPE) rows. Default to table[1].
    cat_table = doc.tables[1]
    cat_rows = []
    for row in cat_table.rows[1:]:
        cat_rows.append([cell.text.strip() for cell in row.cells])

    # Top category = highest avg resolution time (category_data is ORDER BY avg DESC).
    if category_data:
        top_cat, top_count, top_avg = category_data[0]
        matched = None
        for r in cat_rows:
            if r and r[0].strip().lower() == str(top_cat).strip().lower():
                matched = r
                break
        ok = bool(matched) and cell_has_int(matched[1:], top_count, 5) and \
            cell_has_float(matched[1:], float(top_avg), 0.5)
        check("Top category present with correct count and avg hours", ok,
              f"Top category '{top_cat}': expected count ~{top_count}, avg ~{float(top_avg)}; row={matched}")

    # ---- Summary paragraph: verify overall total AND overall avg numbers ----
    nums = re.findall(r"\d[\d,\.]*", normalize_ru_numbers(full_text).replace(",", ""))
    int_vals = []
    float_vals = []
    for n in nums:
        try:
            if "." in n:
                float_vals.append(float(n))
            else:
                int_vals.append(int(n))
        except ValueError:
            continue
    has_total = any(abs(v - overall_count) <= 5 for v in int_vals + float_vals)
    has_avg = any(num_close(v, float(overall_avg), 0.5) for v in float_vals + int_vals)
    check("Summary states overall resolved total", has_total,
          f"Expected overall total ~{overall_count}")
    check("Summary states overall average resolution time", has_avg,
          f"Expected overall avg ~{float(overall_avg)}")

    return True


def check_gcal():
    """Check Google Calendar event: date 2026-03-12 AND 14:00-15:00 window."""
    print("\n=== Checking Google Calendar ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, description, start_datetime, end_datetime
        FROM gcal.events
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    check("At least 1 calendar event created", len(events) >= 1,
          f"Found {len(events)}")

    found_title = False
    found_window = False
    for summary, description, start_dt, end_dt in events:
        s = str(summary or "").lower()
        title_ok = "resolution" in s or "review" in s or "реш" in s or "разбор" in s
        if title_ok:
            found_title = True
        sstr = str(start_dt or "")
        estr = str(end_dt or "")
        date_ok = "2026-03-12" in sstr
        start_time_ok = ("14:00" in sstr) or ("T14:" in sstr) or (" 14:" in sstr) or sstr.endswith("14:00:00")
        end_time_ok = ("15:00" in estr) or ("T15:" in estr) or (" 15:" in estr) or estr.endswith("15:00:00")
        if title_ok and date_ok and start_time_ok and end_time_ok:
            found_window = True

    check("Event title mentions resolution/review", found_title,
          f"Events: {[e[0] for e in events]}")
    check("Calendar event on 2026-03-12 from 14:00 to 15:00", found_window,
          f"Events (start,end): {[(str(e[2]), str(e[3])) for e in events]}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("SF SUPPORT RESOLUTION WORD GCAL - EVALUATION (ClickHouse)")
    print("=" * 70)

    check_word_doc(args.agent_workspace)
    check_gcal()

    total = (PASS_COUNT + FAIL_COUNT) or 1
    accuracy = PASS_COUNT / total * 100
    print(f"\n=== SUMMARY ===")
    print(f"  Total checks - Passed: {PASS_COUNT}, Failed: {FAIL_COUNT} ({accuracy:.1f}%)")

    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    success = (not critical_failed) and accuracy >= 70

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": success}, f, indent=2)

    if critical_failed:
        print(f"  CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"    - {n}")
        print("\n=== RESULT: FAIL (critical check failed) ===")
        sys.exit(1)

    if accuracy >= 70:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)
    else:
        print("\n=== RESULT: FAIL (accuracy below 70%) ===")
        sys.exit(1)


if __name__ == "__main__":
    main()
