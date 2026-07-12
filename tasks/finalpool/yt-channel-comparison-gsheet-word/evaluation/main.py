"""
Evaluation for yt-channel-comparison-gsheet-word task.

Checks:
1. Channel_Analysis_Report.docx exists with 4 headings
2. Word doc mentions both channels with key metrics
3. GSheet "Channel Comparison Analysis" exists with Comparison sheet
4. Comparison sheet has 2 data rows for Fireship and Veritasium
5. Email sent to media@company.com
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "cowork_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0

# CRITICAL semantic checks. Any failure here -> hard FAIL (sys.exit(1)),
# regardless of overall accuracy. A correct RU agent passes all of these;
# a non-doer (no real computation / missing deliverables) fails.
CRITICAL_FAILED = []

# Non-volatile anchors computed from the READ-ONLY youtube.videos data
# (preprocess never touches youtube.*). Used to prove real aggregation,
# not just empty sheet/doc creation.
FIRESHIP_TOTAL_VIEWS = 121886817
VERITASIUM_TOTAL_VIEWS = 365002206


def record(name, passed, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    tag = "CRITICAL " if critical else ""
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {tag}{name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {tag}{name}{msg}")
        if critical:
            CRITICAL_FAILED.append(name)


def check_word(agent_workspace):
    print("\n=== Check 1: Channel_Analysis_Report.docx ===")
    docx_path = os.path.join(agent_workspace, "Channel_Analysis_Report.docx")
    if not os.path.exists(docx_path):
        record("Channel_Analysis_Report.docx exists", False,
               f"Not found at {docx_path}", critical=True)
        return
    record("Channel_Analysis_Report.docx exists", True, critical=True)

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e), critical=True)
        return
    record("Word doc readable", True)

    # Check headings (all four required sections are CRITICAL structure).
    headings = [p.text.strip() for p in doc.paragraphs
                if p.style.name.startswith('Heading')]
    heading_text = " ".join(headings).lower()
    has_exec = "executive" in heading_text or "summary" in heading_text
    has_metrics = "metric" in heading_text or "comparison" in heading_text
    has_findings = "finding" in heading_text or "key" in heading_text
    has_recs = "recommendation" in heading_text
    all_headings = has_exec and has_metrics and has_findings and has_recs
    record("Has Executive Summary heading", has_exec, f"Headings: {headings}")
    record("Has Channel Metrics Comparison heading", has_metrics, f"Headings: {headings}")
    record("Has Key Findings heading", has_findings, f"Headings: {headings}")
    record("Has Recommendations heading", has_recs, f"Headings: {headings}")
    record("All 4 required report sections present", all_headings,
           f"Headings: {headings}", critical=True)

    # Check content
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    record("Doc mentions Fireship", "fireship" in full_text, "Fireship not found in document")
    record("Doc mentions Veritasium", "veritasium" in full_text, "Veritasium not found")


def check_gsheet():
    print("\n=== Check 2: GSheet 'Channel Comparison Analysis' ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, title FROM gsheet.spreadsheets
        WHERE title ILIKE %s
    """, ("%Channel Comparison%",))
    sheets = cur.fetchall()
    record("GSheet 'Channel Comparison Analysis' exists", len(sheets) >= 1,
           f"Found: {[s[1] for s in sheets]}")

    if sheets:
        ss_id = sheets[0][0]
        cur.execute("""
            SELECT title FROM gsheet.sheets WHERE spreadsheet_id = %s
        """, (ss_id,))
        sheet_titles = [r[0] for r in cur.fetchall()]
        sheet_lower = [t.lower() for t in sheet_titles]
        record("Comparison sheet exists in GSheet",
               any("comparison" in t for t in sheet_lower),
               f"Sheets: {sheet_titles}")
        record("Monthly_Fireship sheet exists in GSheet",
               any("monthly" in t or "fireship" in t for t in sheet_lower),
               f"Sheets: {sheet_titles}")

        # CRITICAL: prove the Comparison sheet actually contains the
        # correctly-aggregated Total_Views for BOTH channels. These large
        # exact numbers come from real aggregation over the READ-ONLY
        # youtube.videos data, so they cannot be guessed or preseeded.
        # Tolerate thousands separators / decimals by stripping non-digits.
        cur.execute("""
            SELECT value, formatted_value FROM gsheet.cells
            WHERE spreadsheet_id = %s
        """, (ss_id,))
        cell_digits = set()
        for value, fmt in cur.fetchall():
            for raw in (value, fmt):
                if raw is None:
                    continue
                digits = "".join(ch for ch in str(raw) if ch.isdigit())
                if digits:
                    cell_digits.add(digits)
        has_fireship_views = str(FIRESHIP_TOTAL_VIEWS) in cell_digits
        has_veritasium_views = str(VERITASIUM_TOTAL_VIEWS) in cell_digits
        record("Comparison sheet has correct Fireship Total_Views",
               has_fireship_views,
               f"Expected {FIRESHIP_TOTAL_VIEWS} in sheet cells", critical=True)
        record("Comparison sheet has correct Veritasium Total_Views",
               has_veritasium_views,
               f"Expected {VERITASIUM_TOTAL_VIEWS} in sheet cells", critical=True)

    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 3: Email sent ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT m.to_addr, m.subject FROM email.messages m
        JOIN email.sent_log sl ON sl.message_id = m.id
        WHERE m.to_addr::text ILIKE %s
        ORDER BY sl.sent_at DESC LIMIT 5
    """, ("%media%",))
    emails = cur.fetchall()

    if not emails:
        cur.execute("""
            SELECT to_addr, subject FROM email.messages
            WHERE to_addr::text ILIKE %s
            ORDER BY date DESC LIMIT 5
        """, ("%media%",))
        emails = cur.fetchall()

    cur.close()
    conn.close()

    record("Email sent to media@company.com", len(emails) >= 1,
           f"Found: {emails}", critical=True)
    if emails:
        subject = str(emails[0][1]).lower() if emails[0][1] else ""
        record("Email subject mentions 'Channel Comparison'",
               "channel" in subject or "comparison" in subject or "youtube" in subject,
               f"Subject: {emails[0][1]}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gsheet()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_failed": CRITICAL_FAILED,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    # CRITICAL gate: any failed critical check -> hard FAIL before accuracy.
    if CRITICAL_FAILED:
        print(f"\nFAIL: {len(CRITICAL_FAILED)} CRITICAL check(s) failed: {CRITICAL_FAILED}")
        sys.exit(1)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
