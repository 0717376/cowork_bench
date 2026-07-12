"""
Generate groundtruth files:
  - Q4_2025_Sales_Report.xlsx (two sheets)
  - Executive_Summary.docx

NOTE: After the snowflake->clickhouse fork, the warehouse returns RUSSIAN
REGION/SEGMENT values (db/zzz_clickhouse_after_init.sql). The agent therefore
writes Russian region/segment strings into the xlsx, so the groundtruth cells
must be russified through the SAME central map (scripts/clickhouse_relabel_map.py
REGIONS/SEGMENTS). ALL NUMBERS ARE FROZEN; sheet and column names stay English.
"""
import os
import sys
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt

# Import the central English->Russian relabel map (single source of truth).
_SCRIPTS = os.path.abspath(os.path.join(
    os.path.dirname(__file__), "..", "..", "..", "..", "scripts"))
if _SCRIPTS not in sys.path:
    sys.path.insert(0, _SCRIPTS)
from clickhouse_relabel_map import REGIONS, SEGMENTS as SEGMENT_MAP  # noqa: E402

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def R(region):
    """English region -> Russian via central map (frozen numbers)."""
    return REGIONS[region]


def S(segment):
    """English segment -> Russian via central map."""
    return SEGMENT_MAP[segment]


# ── Data (English keys; russified on emit via central map) ──────────────────

TARGETS = {
    "Asia Pacific": 65000,
    "Europe": 60000,
    "Latin America": 55000,
    "Middle East": 50000,
    "North America": 55000,
}

ACTUALS = {
    "Asia Pacific":  {"revenue": 70510.11, "orders": 386, "customers": 246},
    "Europe":        {"revenue": 54490.62, "orders": 364, "customers": 246},
    "Latin America": {"revenue": 57100.57, "orders": 335, "customers": 231},
    "Middle East":   {"revenue": 57505.34, "orders": 341, "customers": 233},
    "North America": {"revenue": 51818.56, "orders": 358, "customers": 242},
}

SEGMENTS = {
    "Asia Pacific": [
        ("Enterprise", 25754.31, 108),
        ("Consumer", 16491.17, 93),
        ("SMB", 14248.70, 90),
        ("Government", 14015.93, 95),
    ],
    "Europe": [
        ("Enterprise", 16347.91, 98),
        ("Consumer", 16321.67, 93),
        ("SMB", 14196.64, 86),
        ("Government", 7624.40, 87),
    ],
    "Latin America": [
        ("Consumer", 20924.91, 109),
        ("Enterprise", 16106.22, 75),
        ("Government", 11244.69, 76),
        ("SMB", 8824.75, 75),
    ],
    "Middle East": [
        ("Government", 18668.73, 89),
        ("SMB", 16965.25, 92),
        ("Consumer", 11197.00, 80),
        ("Enterprise", 10674.36, 80),
    ],
    "North America": [
        ("Government", 16309.09, 92),
        ("Enterprise", 14532.33, 92),
        ("Consumer", 13054.90, 98),
        ("SMB", 7922.24, 76),
    ],
}


def create_excel():
    wb = openpyxl.Workbook()

    # ── Sheet 1: Regional Performance ──
    ws1 = wb.active
    ws1.title = "Regional Performance"
    headers = ["Region", "Target", "Actual", "Variance", "Variance_Pct", "Order_Count", "Customer_Count"]
    ws1.append(headers)

    # Style header
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
    for col_idx, h in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    # Sort by the RUSSIAN region label A-Z (that is what the agent sorts on).
    for region in sorted(ACTUALS.keys(), key=lambda r: R(r)):
        target = TARGETS[region]
        actual = round(ACTUALS[region]["revenue"], 2)
        variance = round(actual - target, 2)
        variance_pct = round(variance / target * 100, 1)
        order_count = ACTUALS[region]["orders"]
        customer_count = ACTUALS[region]["customers"]
        ws1.append([R(region), target, actual, variance, variance_pct, order_count, customer_count])

    # Auto-width
    for col in ws1.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws1.column_dimensions[col[0].column_letter].width = max_len + 4

    # ── Sheet 2: Segment Breakdown ──
    ws2 = wb.create_sheet("Segment Breakdown")
    seg_headers = ["Region", "Segment", "Revenue", "Orders"]
    ws2.append(seg_headers)

    for col_idx, h in enumerate(seg_headers, 1):
        cell = ws2.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    # Region A-Z by Russian label; segments already revenue-desc within region.
    for region in sorted(SEGMENTS.keys(), key=lambda r: R(r)):
        for segment, revenue, orders in SEGMENTS[region]:
            ws2.append([R(region), S(segment), round(revenue, 2), orders])

    for col in ws2.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws2.column_dimensions[col[0].column_letter].width = max_len + 4

    xlsx_path = os.path.join(OUTPUT_DIR, "Q4_2025_Sales_Report.xlsx")
    wb.save(xlsx_path)
    print(f"Created: {xlsx_path}")


def create_docx():
    doc = Document()

    # Title
    title = doc.add_heading("Сверка продаж за 4 квартал 2025 — Сводка для руководства", level=1)

    total_actual = sum(v["revenue"] for v in ACTUALS.values())
    total_target = sum(TARGETS.values())

    beat_regions = [r for r in sorted(ACTUALS.keys(), key=lambda r: R(r)) if ACTUALS[r]["revenue"] > TARGETS[r]]
    missed_regions = [r for r in sorted(ACTUALS.keys(), key=lambda r: R(r)) if ACTUALS[r]["revenue"] <= TARGETS[r]]

    beat_str = ", ".join(R(r) for r in beat_regions)
    missed_str = ", ".join(R(r) for r in missed_regions)

    summary_text = (
        f"В 4 квартале 2025 года компания получила суммарную выручку ${total_actual:,.2f} при совокупном "
        f"плане ${total_target:,.2f} по всем пяти регионам. Три региона превысили план: "
        f"{beat_str}, продемонстрировав уверенные результаты за счёт сильных корпоративного и "
        f"клиентского сегментов. Однако два региона не достигли целевых показателей: {missed_str}. "
        f"Регион {R('Europe')} недобрал к плану ${TARGETS['Europe']:,} примерно ${5509:,}, а {R('North America')} "
        f"не дотянул до плана ${TARGETS['North America']:,} примерно ${3181:,}. В целом компания превысила "
        f"совокупный план на ${total_actual - total_target:,.2f}, что говорит о хороших итогах квартала "
        f"несмотря на региональные расхождения. Руководству следует разобраться с отставанием регионов "
        f"{R('Europe')} и {R('North America')} и определить корректирующие действия на 1 квартал 2026 года."
    )

    para = doc.add_paragraph(summary_text)
    for run in para.runs:
        run.font.size = Pt(11)

    docx_path = os.path.join(OUTPUT_DIR, "Executive_Summary.docx")
    doc.save(docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    create_excel()
    create_docx()
