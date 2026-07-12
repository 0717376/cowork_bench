"""
Evaluation for insales-bestsellers-word-email task.

Checks:
1. Word document Bestsellers_Report.docx with correct top 10 products
2. Email sent to sales-leads@store.com ( DB check)
"""
import argparse
import sys
import os
from pathlib import Path

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []

# Names of SEMANTIC checks whose failure must force a FAIL regardless of accuracy.
CRITICAL_CHECKS = {
    "Has a table",
    "Table has 10 data rows",
    "Top product name matches",
    "Top product Total_Sold matches",
    "All 10 rows present & values correct & descending",
    "Summary has total/avg-price/#1 figures",
    "Email body has exact top name + units + total",
    "Email subject references bestsellers/top",
}


def _norm_sep(s):
    # Strip thousands separators (comma, regular/NBSP/narrow-NBSP/thin spaces)
    # so a 4-digit figure like "1,914" or "1 914" still matches "1914".
    return (s or "").replace(",", "").replace(" ", "").replace(" ", "") \
        .replace(" ", "").replace(" ", "")


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILS.append(name)
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def get_expected_top10():
    """Query PostgreSQL for top 10 products by total_sales."""
    import psycopg2

    conn = psycopg2.connect(
        host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym",
        user="eigent", password="camel"
    )
    cur = conn.cursor()
    cur.execute("""
        SELECT name, total_sales::int as total_sold, price::float as price
        FROM wc.products
        ORDER BY total_sales::int DESC
        LIMIT 10
    """)
    products = cur.fetchall()
    conn.close()
    return products


def check_word(workspace, expected):
    """Check Bestsellers_Report.docx for correctness."""
    from docx import Document

    print("\n=== Checking Word Document ===")
    doc_path = Path(workspace) / "Bestsellers_Report.docx"

    if not doc_path.exists():
        check("Word file exists", False, f"Not found: {doc_path}")
        # Fail closed: the table CRITICAL deliverables cannot exist without the file.
        check("Has a table", False, "Word file missing")
        check("Table has 10 data rows", False, "Word file missing")
        return
    check("Word file exists", True)

    doc = Document(str(doc_path))

    # Check for title heading
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    full_text = "\n".join(p.text for p in doc.paragraphs).lower()

    HEADING_TOKENS = ["best", "bestsell", "бестселлер", "лучш", "продава"]
    check("Has heading with bestsellers/top (RU+EN)",
          any(any(t in h.lower() for t in HEADING_TOKENS) for h in headings),
          f"Headings: {[h[:50] for h in headings]}")

    # Check tables
    if len(doc.tables) == 0:
        check("Has a table", False, "No tables found")
        # Fail closed: without a table there are no 10 data rows (CRITICAL).
        check("Table has 10 data rows", False, "No tables found")
        return
    check("Has a table", True)

    table = doc.tables[0]
    rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    check("Table has 10 data rows", len(rows) == 10, f"Got {len(rows)} rows")

    # CRITICAL: top product name (English-preserved, first 30 chars) and its units
    if len(rows) >= 1:
        top_name_in_doc = rows[0][1] if len(rows[0]) > 1 else ""
        exp_top_name = expected[0][0]
        check("Top product name matches",
              top_name_in_doc[:30].lower() == exp_top_name[:30].lower(),
              f"Expected '{exp_top_name[:50]}', got '{top_name_in_doc[:50]}'")
        try:
            top_sold_in_doc = int(rows[0][2]) if len(rows[0]) >= 3 else None
        except ValueError:
            top_sold_in_doc = None
        check("Top product Total_Sold matches",
              top_sold_in_doc == expected[0][1],
              f"Expected {expected[0][1]}, got {top_sold_in_doc}")

    # Per-rank Total_Sold values (non-critical structural granularity)
    parsed_sold = []
    for i, (exp_name, exp_sold, exp_price) in enumerate(expected):
        if i >= len(rows):
            break
        row = rows[i]
        if len(row) >= 3:
            try:
                actual_sold = int(row[2])
                parsed_sold.append(actual_sold)
                check(f"Rank {i+1} Total_Sold",
                      actual_sold == exp_sold,
                      f"Expected {exp_sold}, got {actual_sold}")
            except ValueError:
                parsed_sold.append(None)
                check(f"Rank {i+1} Total_Sold", False, f"Cannot parse: {row[2]}")

    # CRITICAL: all 10 rows present, each equals expected, and sequence is non-increasing
    exp_sold_seq = [p[1] for p in expected]
    all_rows = len(rows) == 10 and len(parsed_sold) == 10
    values_correct = all_rows and parsed_sold == exp_sold_seq
    descending = all_rows and all(
        parsed_sold[i] is not None and parsed_sold[i + 1] is not None
        and parsed_sold[i] >= parsed_sold[i + 1]
        for i in range(len(parsed_sold) - 1)
    )
    check("All 10 rows present & values correct & descending",
          values_correct and descending,
          f"rows={len(rows)} parsed={parsed_sold} expected={exp_sold_seq}")

    # CRITICAL: Summary must contain all three figures
    total_sold_all = sum(p[1] for p in expected)
    avg_price = round(sum(p[2] for p in expected) / len(expected), 2)
    avg_price_str = f"{avg_price:.2f}"
    top_name_lc = expected[0][0][:30].lower()
    has_total = str(total_sold_all) in _norm_sep(full_text)
    # Accept the 2-decimal string or a comma-decimal RU variant
    has_avg = avg_price_str in full_text or avg_price_str.replace(".", ",") in full_text
    has_top1 = top_name_lc in full_text and str(expected[0][1]) in full_text
    check("Summary has total/avg-price/#1 figures",
          has_total and has_avg and has_top1,
          f"total={has_total}({total_sold_all}) avg={has_avg}({avg_price_str}) "
          f"top1_name+units={has_top1}")


def check_email(expected):
    """Check email was sent ( DB check)."""
    import psycopg2

    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(
            host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym",
            user="eigent", password="camel"
        )
        cur = conn.cursor()
    except Exception as e:
        check("DB connection", False, str(e))
        return

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    target = "sales-leads@store.com"
    found = None
    for subj, from_addr, to_addr, body in all_emails:
        to_str = str(to_addr or "").lower()
        if target in to_str:
            found = (subj, from_addr, to_addr, body)
            break

    check(f"Email sent to {target}", found is not None,
          f"Found {len(all_emails)} total emails")

    if found:
        subj, _, _, body = found
        body_lower = (body or "").lower()
        subj_lower = (subj or "").lower()

        SUBJ_TOKENS = ["bestseller", "top", "best", "бестселлер", "топ"]
        check("Email subject references bestsellers/top",
              bool(subj_lower.strip()) and any(t in subj_lower for t in SUBJ_TOKENS),
              f"Subject: {(subj or '')[:100]}")

        # CRITICAL: body must contain exact top product name, its exact units,
        # and the total-top-10 units figure.
        total_sold_all = sum(p[1] for p in expected)
        top_name_lc = expected[0][0][:30].lower()
        has_name = top_name_lc in body_lower
        body_norm = _norm_sep(body or "")
        has_units = str(expected[0][1]) in body_norm
        has_total = str(total_sold_all) in body_norm
        check("Email body has exact top name + units + total",
              has_name and has_units and has_total,
              f"name={has_name}({top_name_lc}) units={has_units}({expected[0][1]}) "
              f"total={has_total}({total_sold_all})")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("Fetching expected data...")
    expected = get_expected_top10()
    print(f"  Top 10 products loaded")

    check_word(args.agent_workspace, expected)
    check_email(expected)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0

    print(f"\n=== SUMMARY ===")
    print(f"  Total checks - Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    print(f"  Accuracy: {accuracy:.1f}%")

    if CRITICAL_FAILS:
        print(f"  CRITICAL failures: {CRITICAL_FAILS}")
        print("  Overall: FAIL (critical check failed)")
        print("\nSome critical checks failed.")
        sys.exit(1)

    all_ok = accuracy >= 70
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if all_ok:
        print("\nPass all tests!")
        sys.exit(0)
    else:
        print("\nAccuracy below threshold.")
        sys.exit(1)
