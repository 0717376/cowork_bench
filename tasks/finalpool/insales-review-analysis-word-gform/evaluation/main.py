"""Evaluation для insales-review-analysis-word-gform (RU-стек: insales/forms).

Проверяет два артефакта:
  - Word Review_Analysis.docx: заголовок отчёта, 3 таблицы (категории, лучшие,
    худшие товары) и абзац-резюме. Содержательные значения сверяются с
    groundtruth_workspace/Review_Data.xlsx (листы By Category / Best Products /
    Worst Products), который уже централизованно русифицирован (RU-категории).
  - Forms (gform.*): форма обратной связи об улучшении товаров с тремя вопросами —
    выбор категории, оценка опыта покупок (шкала 1-5) и текстовое предложение.
    RU forms-mcp (local_servers/forms-mcp) умеет ровно два типа вопросов:
    add_text_question (question_type='textQuestion') и
    add_multiple_choice_question (question_type='choiceQuestion'), поэтому
    LINEAR_SCALE недостижим — шкалу 1-5 проверяем как choice-вопрос про опыт.

CRITICAL_CHECKS: любой их fail => задача FAIL, даже если accuracy >= 70%.
Структурные чеки (файл есть, таблица есть, 8 строк) — мягкие. Порог: accuracy
>= 70% И нет критических провалов.
"""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []

# Критические (содержательные) чеки — по имени, переданному в check().
CRITICAL_CHECKS = {
    "Таблица категорий воспроизводит все 8 категорий с верным Avg_Rating и Review_Count",
    "Лучшие товары: >=4 из 5 строк — валидные члены tie-frontier (название+категория+рейтинг)",
    "Худшие товары: >=4 из 5 строк валидны + оба однозначно худших товара присутствуют",
    "Forms: форма обратной связи с 3 вопросами (категория / опыт-шкала / текст-предложение)",
    "Резюме: общее число отзывов, число категорий (8), общий средний рейтинг, наивысшая+наинизшая категория",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILED
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILED.append(name)


def num_close(a, b, tol=0.5):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def _to_float(cell):
    try:
        return float(str(cell).replace(",", ".").strip())
    except (ValueError, AttributeError):
        return None


def _short(name):
    """Короткий устойчивый ключ названия товара (первые слова до 40 симв., lower)."""
    return (name or "").strip().lower()[:40]


def _table_rows(table):
    rows = []
    for row in table.rows[1:]:
        rows.append([c.text.strip() for c in row.cells])
    return rows


def _find_products_table(doc, gt_products):
    """Возвращает rows таблицы товаров, лучше всего совпадающей с эталоном.

    Таблицы товаров имеют 4 колонки; ищем ту, что максимально покрывает
    эталонные названия (Best/Worst)."""
    best_rows, best_score = [], -1
    for t in doc.tables:
        rows = _table_rows(t)
        if not rows or len(rows[0]) < 4:
            continue
        joined = " ".join(" ".join(r).lower() for r in rows)
        score = sum(1 for p in gt_products if _short(p[0]) and _short(p[0]) in joined)
        if score > best_score:
            best_score, best_rows = score, rows
    return best_rows


def _count_products_present(rows, gt_products):
    """Сколько эталонных товаров (name+category+avg) представлено в строках таблицы."""
    found = 0
    for name, cat, _reviews, avg in gt_products:
        sn = _short(name)
        for r in rows:
            joined = " ".join(r).lower()
            name_ok = sn and sn in joined
            cat_ok = (cat or "").lower() in joined
            avg_ok = any(num_close(_to_float(c), avg, 0.1) for c in r if _to_float(c) is not None)
            if name_ok and cat_ok and avg_ok:
                found += 1
                break
    return found


def _row_matches(row, member):
    """Строка таблицы соответствует кандидату (name+category+avg)."""
    name, cat, _reviews, avg = member
    joined = " ".join(row).lower()
    sn = _short(name)
    name_ok = sn and sn in joined
    cat_ok = (cat or "").lower() in joined
    avg_ok = any(num_close(_to_float(c), avg, 0.1) for c in row if _to_float(c) is not None)
    return name_ok and cat_ok and avg_ok


def _count_valid_rows(rows, members):
    """(валидных строк, всего строк): строка валидна, если совпадает с
    каким-либо членом tie-frontier по названию+категории+рейтингу."""
    data = [r for r in rows if any((c or "").strip() for c in r)]
    valid = sum(1 for r in data if any(_row_matches(r, m) for m in members))
    return valid, len(data)


def _product_stats_from_db():
    """Пер-(товар, категория) статистика отзывов из БД (>=3 отзывов)."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT p.name, cat->>'name', COUNT(*), ROUND(AVG(r.rating)::numeric, 2)
        FROM wc.product_reviews r
        JOIN wc.products p ON p.id = r.product_id
        CROSS JOIN LATERAL jsonb_array_elements(p.categories) cat
        GROUP BY p.id, p.name, cat->>'name'
        HAVING COUNT(*) >= 3
    """)
    stats = [(n, c, int(cnt), float(avg)) for n, c, cnt, avg in cur.fetchall()]
    cur.close()
    conn.close()
    return stats


def _tie_frontier(stats, worst=False):
    """(members, required): члены tie-frontier топ/боттом-5 с учётом ничьих.

    members — все (name, cat, n, avg), допустимые в корректном топ/боттом-5
    (avg за/на границе 5-го различного товара); required — названия товаров,
    которые ОБЯЗАНЫ присутствовать (avg строго за границей ничьей)."""
    by_product = {}
    for name, _cat, _n, avg in stats:
        by_product[name] = avg
    avgs = sorted(by_product.values(), reverse=not worst)
    cutoff = avgs[min(4, len(avgs) - 1)] if avgs else None
    if cutoff is None:
        return [], set()
    eps = 1e-9
    if worst:
        members = [s for s in stats if s[3] <= cutoff + eps]
        required = {nm for nm, a in by_product.items() if a < cutoff - eps}
    else:
        members = [s for s in stats if s[3] >= cutoff - eps]
        required = {nm for nm, a in by_product.items() if a > cutoff + eps}
    return members, required


def check_word_doc(agent_workspace, groundtruth_workspace):
    """Проверка структуры и содержимого Word-документа."""
    print("\n=== Проверка Word-документа ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx установлен", False, "pip install python-docx")
        return

    doc_path = os.path.join(agent_workspace, "Review_Analysis.docx")
    check("Файл Word существует", os.path.isfile(doc_path), f"Ожидался {doc_path}")
    if not os.path.isfile(doc_path):
        return

    doc = Document(doc_path)

    # Заголовок отчёта (RU+EN)
    has_heading = False
    for p in doc.paragraphs:
        t = p.text.lower()
        if ("анализ" in t and "отзыв" in t) or ("review" in t and "analysis" in t):
            has_heading = True
            break
    check("Документ имеет заголовок анализа отзывов", has_heading)

    # Таблицы
    check("Документ имеет минимум 3 таблицы", len(doc.tables) >= 3,
          f"Найдено таблиц: {len(doc.tables)}")

    # Groundtruth
    gt_file = os.path.join(groundtruth_workspace, "Review_Data.xlsx")
    if not os.path.isfile(gt_file):
        check("Groundtruth файл существует", False)
        return
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    gt_cats = list(gt_wb["By Category"].iter_rows(min_row=2, values_only=True))
    gt_best = list(gt_wb["Best Products"].iter_rows(min_row=2, values_only=True))
    gt_worst = list(gt_wb["Worst Products"].iter_rows(min_row=2, values_only=True))

    # --- Таблица категорий (table 0) ---
    cat_rows = _table_rows(doc.tables[0]) if doc.tables else []
    check("Таблица категорий имеет 8 строк", len(cat_rows) == 8, f"Получено {len(cat_rows)} строк")

    # Построчная сверка: категория найдена + Avg_Rating (tol 0.05) + Review_Count точно.
    cat_full_ok = (len(gt_cats) == 8)
    for gt_row in gt_cats:
        cat_name, review_count, avg_rating, _pos, _neg = gt_row
        matched = None
        for r in cat_rows:
            if r and cat_name.lower() in r[0].lower():
                matched = r
                break
        if not matched:
            check(f"Категория «{cat_name}» найдена в таблице", False)
            cat_full_ok = False
            continue
        nums = [_to_float(c) for c in matched[1:]]
        nums = [n for n in nums if n is not None]
        avg_ok = any(num_close(n, avg_rating, 0.05) for n in nums)
        cnt_ok = any(int(n) == int(review_count) for n in nums)
        check(f"Категория «{cat_name}»: Avg_Rating ~{avg_rating}", avg_ok, f"Числа строки: {nums}")
        check(f"Категория «{cat_name}»: Review_Count = {review_count}", cnt_ok, f"Числа строки: {nums}")
        if not (avg_ok and cnt_ok):
            cat_full_ok = False
    # CRITICAL
    check("Таблица категорий воспроизводит все 8 категорий с верным Avg_Rating и Review_Count",
          cat_full_ok and len(cat_rows) == 8)

    # --- Лучшие / худшие товары ---
    try:
        stats = _product_stats_from_db()
        best_members, _ = _tie_frontier(stats, worst=False)
        worst_members, worst_required = _tie_frontier(stats, worst=True)
    except Exception as e:
        print(f"  WARNING: tie-frontier из БД недоступен ({e}), fallback на GT")
        best_members, worst_members = gt_best, gt_worst
        worst_required = set()

    best_rows = _find_products_table(doc, best_members)
    worst_rows = _find_products_table(doc, worst_members)
    n_best, t_best = _count_valid_rows(best_rows, best_members)
    n_worst, t_worst = _count_valid_rows(worst_rows, worst_members)
    check("Лучшие товары: >=4 из 5 строк — валидные члены tie-frontier (название+категория+рейтинг)",
          n_best >= 4 and t_best - n_best <= 1,
          f"Валидно {n_best}/{t_best}")
    required_ok = all(
        any(_row_matches(r, m) for r in worst_rows
            for m in worst_members if m[0] == req)
        for req in worst_required
    )
    check("Худшие товары: >=4 из 5 строк валидны + оба однозначно худших товара присутствуют",
          n_worst >= 4 and t_worst - n_worst <= 1 and required_ok,
          f"Валидно {n_worst}/{t_worst}, required_ok={required_ok}")

    # Разделы лучших/худших (RU+EN)
    has_top = False
    has_worst_sec = False
    for p in doc.paragraphs:
        t = p.text.lower()
        if ("лучш" in t and ("рейтинг" in t or "товар" in t)) or ("top" in t and "rated" in t):
            has_top = True
        if "низк" in t or "худш" in t or "lowest" in t or "worst" in t:
            has_worst_sec = True
    check("Есть раздел лучших товаров", has_top)
    check("Есть раздел товаров с низким рейтингом", has_worst_sec)

    # --- Резюме ---
    full_text = " ".join(p.text for p in doc.paragraphs)
    low = full_text.lower()
    # Общее число отзывов: сумма по категориям (474) двойным счётом учитывает
    # мультикатегорийные товары; принимаем и число уникальных отзывов из БД (396).
    total_reviews = sum(int(r[1]) for r in gt_cats)
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM wc.product_reviews")
        distinct_reviews = cur.fetchone()[0]
        cur.close()
        conn.close()
    except Exception:
        distinct_reviews = None
    n_categories = len(gt_cats)                            # 8
    # общий средний рейтинг (взвешенный)
    overall_avg = round(sum(int(r[1]) * float(r[2]) for r in gt_cats) / total_reviews, 2)
    # наивысшая/наинизшая категория (gt отсортирован по убыванию avg)
    highest_cat = gt_cats[0][0]
    lowest_cat = gt_cats[-1][0]

    has_summary_kw = any(k in low for k in ("итог", "всего", "общ", "категор", "total", "overall", "categor"))
    check("Документ содержит текст резюме", has_summary_kw)

    total_ok = (str(total_reviews) in full_text) or (
        distinct_reviews is not None and str(distinct_reviews) in full_text)
    ncat_ok = (str(n_categories) in full_text)
    # средний рейтинг: ищем число близкое к overall_avg среди чисел текста
    text_floats = []
    for tok in low.replace(",", ".").replace("(", " ").replace(")", " ").split():
        f = _to_float(tok.strip(".:;"))
        if f is not None:
            text_floats.append(f)
    avg_ok = any(num_close(f, overall_avg, 0.1) for f in text_floats)

    def cat_named(cat):
        """Категория названа: полное имя ИЛИ её самый длинный токен (>=4 симв.)."""
        c = (cat or "").lower()
        if c and c in low:
            return True
        toks = [t for t in c.replace("&", " ").split() if len(t) >= 4]
        toks.sort(key=len, reverse=True)
        return bool(toks) and toks[0] in low

    hi_ok = cat_named(highest_cat)
    lo_ok = cat_named(lowest_cat)
    check("Резюме упоминает общее число отзывов", total_ok, f"Ожидалось {total_reviews}")
    check("Резюме упоминает число категорий (8)", ncat_ok)
    check("Резюме упоминает общий средний рейтинг", avg_ok, f"Ожидалось ~{overall_avg}")
    check("Резюме называет наивысшую и наинизшую категории", hi_ok and lo_ok,
          f"highest={highest_cat}, lowest={lowest_cat}")
    # CRITICAL
    check("Резюме: общее число отзывов, число категорий (8), общий средний рейтинг, наивысшая+наинизшая категория",
          has_summary_kw and total_ok and ncat_ok and avg_ok and hi_ok and lo_ok)


def check_gform():
    """Проверка формы обратной связи (gform.*)."""
    print("\n=== Проверка формы (gform) ===")
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("SELECT id, title FROM gform.forms")
        forms = cur.fetchall()
        cur.execute("SELECT form_id, title, question_type FROM gform.questions")
        questions = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        check("Forms: форма обратной связи с 3 вопросами (категория / опыт-шкала / текст-предложение)", False, str(e))
        return

    check("Создана минимум 1 форма", len(forms) >= 1, f"Найдено {len(forms)}")

    def title_ok(title):
        t = (title or "").lower()
        ru = any(k in t for k in ("обратн", "улучш", "отзыв", "опрос", "анкет"))
        en = any(k in t for k in ("feedback", "improvement"))
        return ru or en

    found_feedback = any(title_ok(f[1]) for f in forms)
    check("Заголовок формы про обратную связь/улучшение", found_feedback,
          f"Формы: {[f[1] for f in forms]}")

    check("Создано минимум 3 вопроса", len(questions) >= 3, f"Найдено {len(questions)}")

    # Содержательные вопросы по заголовкам (RU+EN в ОРИГИНАЛЬНОМ .lower())
    q_titles = [(q[1] or "").lower() for q in questions]
    q_types = [(q[2] or "") for q in questions]

    def any_title(*keys):
        return any(any(k in qt for k in keys) for qt in q_titles)

    has_category_q = any_title("категор", "category")
    has_rating_q = any_title("опыт", "оцен", "покуп", "experience", "rate", "rating", "shopping")
    has_suggestion_q = any_title("улучш", "предлож", "improve", "suggest")
    check("Есть вопрос про категорию", has_category_q, f"Заголовки: {q_titles}")
    check("Есть вопрос про опыт покупок (оценка)", has_rating_q, f"Заголовки: {q_titles}")
    check("Есть вопрос с предложением улучшений", has_suggestion_q, f"Заголовки: {q_titles}")

    # RU forms-mcp пишет ТОЛЬКО два типа: 'choiceQuestion' и 'textQuestion'.
    # LINEAR_SCALE недостижим -> шкала 1-5 реализуется как choice. Проверяем
    # достижимую схему: есть choice-вопрос(ы) и есть текстовый вопрос.
    has_choice = any(qt in ("choiceQuestion", "RADIO", "MULTIPLE_CHOICE", "CHECKBOX") for qt in q_types)
    has_text = any(qt in ("textQuestion", "TEXT", "PARAGRAPH", "SHORT_ANSWER") for qt in q_types)
    check("Есть вопрос(ы) с выбором (choiceQuestion)", has_choice, f"Типы: {q_types}")
    check("Есть текстовый вопрос (textQuestion)", has_text, f"Типы: {q_types}")

    # CRITICAL: достижимая схема — форма про обратную связь, >=3 вопроса,
    # покрыты категория/опыт/предложение, есть choice и text.
    check("Forms: форма обратной связи с 3 вопросами (категория / опыт-шкала / текст-предложение)",
          found_feedback and len(questions) >= 3
          and has_category_q and has_rating_q and has_suggestion_q
          and has_choice and has_text,
          f"types={q_types}, titles={q_titles}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("WC REVIEW ANALYSIS WORD GFORM (RU: insales/forms) - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace, gt_dir)
    check_gform()

    total = PASS_COUNT + FAIL_COUNT
    pct = 100.0 * PASS_COUNT / total if total else 0.0

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Accuracy: {pct:.1f}%")

    if CRITICAL_FAILED:
        print(f"  CRITICAL FAIL: {CRITICAL_FAILED}")
        print("  Overall: FAIL")
        if args.res_log_file:
            with open(args.res_log_file, "w") as f:
                json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": False}, f, indent=2)
        sys.exit(1)

    overall = pct >= 70.0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": overall}, f, indent=2)
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
