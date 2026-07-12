"""Evaluation script for pw-insales-price-optimization-excel-word."""
import os
import argparse, json, os, sys
import openpyxl


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []

# Centrally-russified wc category labels (db/zzz_wc_after_init.sql + wc_patch_groundtruth.py)
RU_CATEGORY_LABELS = {
    "электроника", "аудио", "камеры",
    "тв и домашний кинотеатр", "бытовая техника",
}

def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        tag = "FAIL-CRITICAL" if critical else "FAIL"
        print(f"  [{tag}] {name}: {detail_str}")
        if critical:
            CRITICAL_FAILED.append(name)

def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(',', '').replace('%', '').replace('$', '').strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    
    excel_path = os.path.join(agent_workspace, "Price_Optimization_Report.xlsx")
    check("Price_Optimization_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Price_Optimization_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        check("Price_Comparison sheet exists", "Price_Comparison" in wb.sheetnames)
        if "Price_Comparison" in wb.sheetnames:
            ws = wb["Price_Comparison"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Price_Comparison has >= 10 rows", len(data_rows) >= 10, f"got {len(data_rows)}")

            # Check headers
            headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
            hidx = {h: i for i, h in enumerate(headers)}
            for expected_col in ['Product_Name', 'Our_Price', 'Competitor_Price', 'Price_Difference', 'Difference_Pct', 'Recommendation']:
                check(f"Price_Comparison has {expected_col} column",
                      expected_col.lower() in headers, f"headers: {headers[:8]}")

            # --- CRITICAL: row-internal price math + recommendation rule ---
            need = {'our_price', 'competitor_price', 'price_difference', 'difference_pct', 'recommendation'}
            if need.issubset(hidx):
                diff_ok = 0
                pct_ok = 0
                rec_ok = 0
                considered = 0
                over = comp = under = 0
                for row in data_rows:
                    op = safe_float(row[hidx['our_price']])
                    cp = safe_float(row[hidx['competitor_price']])
                    pd_written = safe_float(row[hidx['price_difference']])
                    pct_written = safe_float(row[hidx['difference_pct']])
                    rec = str(row[hidx['recommendation']] or "").strip().lower()
                    if op is None or cp is None or cp == 0:
                        continue
                    considered += 1
                    exp_diff = round(op - cp, 2)
                    exp_pct = round((op - cp) / cp * 100, 1)
                    if pd_written is not None and abs(pd_written - exp_diff) <= 0.05:
                        diff_ok += 1
                    if pct_written is not None and abs(pct_written - exp_pct) <= 0.2:
                        pct_ok += 1
                    # Recommendation rule (task.md / analysis_guide.md): Reduce if >15% above,
                    # Maintain if within +-15%, Consider increase if >15% below.
                    # The [-15, 0) band is accepted with either label.
                    if exp_pct > 15:
                        rec_match = rec == "reduce price"
                    elif exp_pct < -15:
                        rec_match = rec == "consider increase"
                    elif exp_pct < 0:
                        rec_match = rec in ("maintain", "consider increase")
                    else:
                        rec_match = rec == "maintain"
                    if rec_match:
                        rec_ok += 1
                    # Exec-summary buckets (analysis_guide.md): +-15% band
                    if exp_pct > 15:
                        over += 1
                    elif exp_pct >= -15:
                        comp += 1
                    else:
                        under += 1
                thr = max(8, int(considered * 0.8)) if considered >= 10 else considered
                check("Price_Comparison: Price_Difference values correct (>=8/10)",
                      diff_ok >= thr, f"{diff_ok}/{considered} correct", critical=True)
                check("Price_Comparison: Difference_Pct values correct (>=8/10)",
                      pct_ok >= thr, f"{pct_ok}/{considered} correct", critical=True)
                check("Price_Comparison: Recommendation matches 15% rule (>=8/10)",
                      rec_ok >= thr, f"{rec_ok}/{considered} correct", critical=True)
                # stash expected tallies for exec-summary cross-check
                wb._expected_tallies = (considered, over, comp, under)

        check("Category_Summary sheet exists", "Category_Summary" in wb.sheetnames)
        if "Category_Summary" in wb.sheetnames:
            ws = wb["Category_Summary"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Category_Summary has >= 5 rows", len(data_rows) >= 5, f"got {len(data_rows)}")

            # Check headers
            headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
            chidx = {h: i for i, h in enumerate(headers)}
            for expected_col in ['Category', 'Product_Count', 'Avg_Our_Price', 'Avg_Stock']:
                check(f"Category_Summary has {expected_col} column",
                      expected_col.lower() in headers, f"headers: {headers[:8]}")

            # --- CRITICAL: RU category labels + plausible numbers ---
            if 'category' in chidx:
                cat_idx = chidx['category']
                ru_matches = 0
                plausible = 0
                ncat = 0
                for row in data_rows:
                    if row[cat_idx] is None:
                        continue
                    ncat += 1
                    cat = str(row[cat_idx]).strip().lower()
                    if cat in RU_CATEGORY_LABELS:
                        ru_matches += 1
                    aop = safe_float(row[chidx['avg_our_price']]) if 'avg_our_price' in chidx else None
                    ast = safe_float(row[chidx['avg_stock']]) if 'avg_stock' in chidx else None
                    if (aop is None or aop > 0) and (ast is None or ast >= 0):
                        plausible += 1
                check("Category_Summary: >=4 rows with russified category labels",
                      ru_matches >= 4, f"{ru_matches} ru labels in {ncat} rows", critical=True)
                check("Category_Summary: Avg_Our_Price/Avg_Stock plausible",
                      ncat > 0 and plausible == ncat, f"{plausible}/{ncat} plausible")

        check("Executive_Summary sheet exists", "Executive_Summary" in wb.sheetnames)
        if "Executive_Summary" in wb.sheetnames:
            ws = wb["Executive_Summary"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Executive_Summary has >= 5 rows", len(data_rows) >= 5, f"got {len(data_rows)}")

            # Check headers
            headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
            for expected_col in ['Metric', 'Value']:
                check(f"Executive_Summary has {expected_col} column",
                      expected_col.lower() in headers, f"headers: {headers[:8]}")

            # --- CRITICAL: exec-summary tallies are self-consistent and match comparison sheet ---
            metrics = {}
            for row in data_rows:
                if row and row[0] is not None:
                    metrics[str(row[0]).strip()] = safe_float(row[1]) if len(row) > 1 else None
            total = metrics.get('Total_Products_Compared')
            mover = metrics.get('Products_Overpriced')
            mcomp = metrics.get('Products_Competitive')
            munder = metrics.get('Products_Underpriced')
            if None not in (total, mover, mcomp, munder):
                check("Executive_Summary: over+competitive+underpriced == total",
                      abs((mover + mcomp + munder) - total) < 0.5,
                      f"{mover}+{mcomp}+{munder} vs {total}", critical=True)
                exp = getattr(wb, "_expected_tallies", None)
                if exp is not None:
                    e_total, e_over, e_comp, e_under = exp
                    matches = (abs((total or 0) - e_total) < 0.5
                               and abs((mover or 0) - e_over) < 0.5
                               and abs((mcomp or 0) - e_comp) < 0.5
                               and abs((munder or 0) - e_under) < 0.5)
                    check("Executive_Summary: tallies match Price_Comparison buckets",
                          matches, f"got {(total,mover,mcomp,munder)} expected {exp}", critical=True)
            else:
                check("Executive_Summary: required metrics present",
                      False, f"metrics={list(metrics)}", critical=True)

        word_path = os.path.join(agent_workspace, "Pricing_Strategy.docx")
        check("Pricing strategy Word exists", os.path.exists(word_path))
        if os.path.exists(word_path):
            from docx import Document
            doc = Document(word_path)
            text = " ".join(p.text for p in doc.paragraphs).lower()
            check("Word mentions pricing (RU/EN)",
                  any(k in text for k in ("pric", "цен", "стоимост")), text[:120])
            check("Word mentions recommendation (RU/EN)",
                  any(k in text for k in ("recommend", "рекоменд")), text[:120])
            # three required sections (RU or EN) + a heading present
            sections = {
                "market position": ("market position", "позиц", "положение на рынке"),
                "product-level recommendations": ("product-level", "product level", "по товар", "рекоменд"),
                "implementation timeline": ("implementation", "timeline", "график", "внедрен", "сроки"),
            }
            sec_found = sum(1 for keys in sections.values() if any(k in text for k in keys))
            check("Word has >=3 required sections (RU/EN)",
                  sec_found >= 3, f"found {sec_found}/3", critical=True)
            check("Word heading present",
                  len(doc.paragraphs) > 0 and any(p.text.strip() for p in doc.paragraphs[:3]),
                  "no heading text")
        check("price_optimizer.py exists", os.path.exists(os.path.join(agent_workspace, "price_optimizer.py")))

    if CRITICAL_FAILED:
        print(f"CRITICAL checks failed: {CRITICAL_FAILED}")
        sys.exit(1)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0
    return accuracy >= 70, f"Passed {PASS_COUNT}/{total} checks ({accuracy:.1f}%)"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
