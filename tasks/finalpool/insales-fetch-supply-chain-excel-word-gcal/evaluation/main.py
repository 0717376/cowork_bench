"""Evaluation script for insales-fetch-supply-chain-excel-word-gcal."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Semantic critical checks: any failure here => overall FAIL regardless of
# accuracy. These verify the load-bearing deliverables demanded by task.md
# (gcal event date/time/description, docx title + a computed value, core
# Excel values) so a non-doing agent cannot pass via the accuracy gate.
CRITICAL_CHECKS = {
    "CRITICAL: gcal event on 19 Mar 2026, 9:00-10:30 UTC",
    "CRITICAL: gcal event description enumerates critical stockout risks",
    "CRITICAL: docx title 'Supply Chain Optimization Report'",
    "CRITICAL: Inventory_Status daily_sales_rate/reorder computed correctly",
    "CRITICAL: Summary core counts match groundtruth",
}

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    # Check Supply_Chain_Optimization.xlsx
    excel_path = os.path.join(agent_workspace, "Supply_Chain_Optimization.xlsx")
    check("Supply_Chain_Optimization.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Supply_Chain_Optimization.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Check row count
                    gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                    data_rows = list(ws.iter_rows(min_row=2, values_only=True))
                    min_rows = max(1, len(gt_rows) - 2)
                    check(f"{sheet_name} has >= {min_rows} data rows", len(data_rows) >= min_rows, f"got {len(data_rows)}")

                    # Cell value comparison against groundtruth
                    header_map = {h: i for i, h in enumerate(headers)}
                    gt_header_map = {h: i for i, h in enumerate(gt_headers)}
                    for ri in range(min(3, len(gt_rows), len(data_rows))):
                        gt_row = gt_rows[ri]
                        agent_row = data_rows[ri]
                        for ci, gt_h in enumerate(gt_headers):
                            if not gt_h or ci >= len(gt_row):
                                continue
                            gv = gt_row[ci]
                            agent_ci = header_map.get(gt_h)
                            if agent_ci is None or agent_ci >= len(agent_row):
                                continue
                            av = agent_row[agent_ci]
                            gf = safe_float(gv)
                            af = safe_float(av)
                            if gf is not None and af is not None:
                                tol = max(0.5, abs(gf) * 0.15)
                                check(f"{sheet_name} R{ri+2} {gt_h} ~{gf:.1f}",
                                      abs(gf - af) <= tol, f"got {af}")
                            elif gv is not None and av is not None:
                                gs = str(gv).strip().lower()
                                avs = str(av).strip().lower()
                                if gs:
                                    check(f"{sheet_name} R{ri+2} {gt_h} text",
                                          gs == avs or gs in avs or avs in gs,
                                          f"expected {gs[:50]}, got {avs[:50]}")

        # CRITICAL: verify computed values (daily_sales_rate = Total/90 and
        # reorder_point = daily_sales_rate * lead_time * 1.5) are correct, not
        # placeholder/empty. Compare agent's Inventory_Status against GT by
        # Product_Name so row ordering differences don't break the check.
        inv_ok = False
        inv_detail = "Inventory_Status sheet missing or unreadable"
        if gt_wb and "Inventory_Status" in wb.sheetnames and "Inventory_Status" in gt_wb.sheetnames:
            def index_inv(ws):
                hdr = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                hm = {h: i for i, h in enumerate(hdr)}
                out = {}
                for row in ws.iter_rows(min_row=2, values_only=True):
                    if not row or row[hm.get("product_name", 0)] is None:
                        continue
                    name = str(row[hm["product_name"]]).strip().lower()
                    out[name] = {
                        "daily": safe_float(row[hm["daily_sales_rate"]]) if "daily_sales_rate" in hm else None,
                        "reorder": safe_float(row[hm["reorder_point"]]) if "reorder_point" in hm else None,
                    }
                return out
            try:
                agent_inv = index_inv(wb["Inventory_Status"])
                gt_inv = index_inv(gt_wb["Inventory_Status"])

                # GT product names are stored truncated to ~125 chars ending in
                # "..." while agents read full names from the InSales API. Match
                # on a normalized prefix so a full agent name resolves to its
                # truncated GT counterpart instead of failing as "missing".
                def lookup(name):
                    if name in agent_inv:
                        return agent_inv[name]
                    base = name[:-3].strip() if name.endswith("...") else name
                    if len(base) >= 10:
                        for an, av in agent_inv.items():
                            if an.startswith(base) or base.startswith(an):
                                return av
                    return None

                matched = 0
                checked = 0
                bad = []
                for name, gvals in gt_inv.items():
                    avals = lookup(name)
                    if not avals:
                        bad.append(f"{name}: missing")
                        continue
                    row_ok = True
                    for key in ("daily", "reorder"):
                        gv, av = gvals[key], avals[key]
                        if gv is None:
                            continue
                        checked += 1
                        tol = max(0.5, abs(gv) * 0.15)
                        if av is None or abs(gv - av) > tol:
                            row_ok = False
                            bad.append(f"{name}.{key}: exp {gv}, got {av}")
                    if row_ok:
                        matched += 1
                # Require the majority of GT products to have correct computed values.
                inv_ok = checked > 0 and matched >= max(1, (len(gt_inv) + 1) // 2)
                inv_detail = f"matched {matched}/{len(gt_inv)} products; issues: {bad[:4]}"
            except Exception as e:
                inv_detail = f"error: {e}"
        check("CRITICAL: Inventory_Status daily_sales_rate/reorder computed correctly",
              inv_ok, inv_detail)

        # CRITICAL: core Summary counts must match groundtruth (proves the
        # classification pipeline actually ran, not just a stub workbook).
        sum_ok = False
        sum_detail = "Summary sheet missing or unreadable"
        if gt_wb and "Summary" in wb.sheetnames and "Summary" in gt_wb.sheetnames:
            def index_summary(ws):
                out = {}
                for row in ws.iter_rows(min_row=2, values_only=True):
                    if not row or row[0] is None:
                        continue
                    out[str(row[0]).strip().lower()] = safe_float(row[1]) if len(row) > 1 else None
                return out
            try:
                agent_sum = index_summary(wb["Summary"])
                gt_sum = index_summary(gt_wb["Summary"])
                core_metrics = ["total_products", "critical_products",
                                "warning_products", "healthy_products"]
                ok_count = 0
                present = 0
                bad = []
                for m in core_metrics:
                    gv = gt_sum.get(m)
                    av = agent_sum.get(m)
                    if gv is None:
                        continue
                    present += 1
                    tol = max(0.5, abs(gv) * 0.15)
                    if av is not None and abs(gv - av) <= tol:
                        ok_count += 1
                    else:
                        bad.append(f"{m}: exp {gv}, got {av}")
                sum_ok = present > 0 and ok_count >= present  # all core counts must match
                sum_detail = f"matched {ok_count}/{present} core metrics; issues: {bad[:4]}"
            except Exception as e:
                sum_detail = f"error: {e}"
        check("CRITICAL: Summary core counts match groundtruth", sum_ok, sum_detail)

    # Check Supply_Chain_Report.docx
    docx_path = os.path.join(agent_workspace, "Supply_Chain_Report.docx")
    check("Supply_Chain_Report.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        check("Supply_Chain_Report.docx has content", len(text) > 50, f"text length: {len(text)}")
        # CRITICAL: title must be the exact report title required by task.md.
        title_texts = [p.text.strip().lower() for p in doc.paragraphs
                       if p.style.name == "Title" or p.style.name.startswith("Title")]
        title_found = any("supply chain optimization report" in t for t in title_texts) \
            or "supply chain optimization report" in text.lower()
        check("CRITICAL: docx title 'Supply Chain Optimization Report'", title_found,
              f"titles: {title_texts[:3]}")
        # Check headings match groundtruth
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        gt_doc_path = os.path.join(groundtruth_workspace, "Supply_Chain_Report.docx")
        if os.path.exists(gt_doc_path):
            gt_doc = Document(gt_doc_path)
            gt_headings = [p.text.strip().lower() for p in gt_doc.paragraphs if p.style.name.startswith("Heading")]
            for gh in gt_headings:
                if gh:
                    found = any(gh in h or h in gh for h in headings)
                    check(f"Supply_Chain_Report.docx has heading \"{gh[:40]}\"", found, f"agent headings: {headings[:5]}")
        else:
            check("Supply_Chain_Report.docx has headings", len(headings) >= 2, f"found {len(headings)} headings")

    # Check Python script exists (terminal usage)
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    check("Python analysis script exists", len(py_files) >= 1, f"found: {py_files}")

    # Database checks
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("""SELECT summary, description, start_datetime, end_datetime
                       FROM gcal.events WHERE summary ILIKE '%supply%'""")
        event_row = cur.fetchone()
        check("Calendar event with correct summary", event_row is not None, "no matching event found")

        # CRITICAL: the event must fall on 19 Mar 2026, 9:00-10:30 UTC.
        # Normalize to UTC so an agent storing tz-aware datetimes still passes.
        dt_ok = False
        dt_detail = "no matching 'supply' event found"
        if event_row is not None:
            from datetime import timezone
            start_dt, end_dt = event_row[2], event_row[3]

            def to_utc_naive(dt):
                if dt is None:
                    return None
                if dt.tzinfo is not None:
                    dt = dt.astimezone(timezone.utc).replace(tzinfo=None)
                return dt
            s = to_utc_naive(start_dt)
            e = to_utc_naive(end_dt)
            if s is not None:
                date_ok = (s.year, s.month, s.day) == (2026, 3, 19)
                start_ok = (s.hour, s.minute) == (9, 0)
                end_ok = e is not None and (e.year, e.month, e.day) == (2026, 3, 19) \
                    and (e.hour, e.minute) == (10, 30)
                dt_ok = date_ok and start_ok and end_ok
                dt_detail = f"start={start_dt}, end={end_dt} (UTC start={s}, end={e})"
        check("CRITICAL: gcal event on 19 Mar 2026, 9:00-10:30 UTC", dt_ok, dt_detail)

        # CRITICAL: description must enumerate critical stockout risks, i.e. name
        # at least 2 of the 3 critical products from the inventory analysis.
        desc_ok = False
        desc_detail = "no event / empty description"
        if event_row is not None:
            desc = (event_row[1] or "").lower()
            critical_products = ["jbl glide", "boxtudio", "type-c earphone",
                                 "30kg digital scale", "silencer panels"]
            hits = sum(1 for p in critical_products if p in desc)
            mentions_risk = any(k in desc for k in
                                ["stockout", "stock-out", "critical", "shortage", "depletion"])
            desc_ok = len(desc.strip()) > 20 and hits >= 2 and mentions_risk
            desc_detail = (f"desc len={len(desc)}, product hits={hits}, "
                           f"risk_kw={mentions_risk}, desc[:120]={desc[:120]!r}")
        check("CRITICAL: gcal event description enumerates critical stockout risks",
              desc_ok, desc_detail)

        # Reverse verification: noise events should not match task keyword
        cur.execute("SELECT COUNT(*) FROM gcal.events WHERE summary ILIKE '%standup%' OR summary ILIKE '%lunch%'")
        noise_events = cur.fetchone()[0]
        check("Noise events exist (not deleted by agent)", noise_events >= 1, f"noise events: {noise_events}")
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e))

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== Results: {PASS_COUNT}/{total} passed ({accuracy:.1f}%) ===")
    if critical_failed:
        print(f"CRITICAL FAILURES ({len(critical_failed)}):")
        for n in critical_failed:
            print(f"  - {n}")

    # Convention: any critical failure => FAIL (sys.exit(1)) before the
    # accuracy gate; otherwise PASS requires accuracy >= 70%.
    if critical_failed:
        return False, f"FAIL (critical check failed): {PASS_COUNT}/{total} passed ({accuracy:.1f}%)"
    if accuracy >= 70:
        return True, f"PASS: {PASS_COUNT}/{total} passed ({accuracy:.1f}%)"
    return False, f"FAIL (accuracy {accuracy:.1f}% < 70%): {PASS_COUNT}/{total} passed"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()