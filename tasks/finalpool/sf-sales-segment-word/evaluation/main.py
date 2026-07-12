"""Evaluation for sf-sales-segment-word (ClickHouse / RU)."""
import argparse
import os
import re
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = False

# Q4 2025 plans (RUB) from Q4_Targets.pdf, keyed by the RUSSIAN segment names
# that the central ClickHouse seed produces in sf_data."...CUSTOMERS"."SEGMENT".
TARGETS = {
    "Частные клиенты": 210000,        # Consumer
    "Корпоративный": 200000,          # Enterprise
    "Государственный": 170000,        # Government
    "Малый и средний бизнес": 180000, # SMB
}

EXPECTED_SEGMENTS = set(TARGETS.keys())


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILED
    tag = "CRITICAL " if critical else ""
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {tag}{name}")
    else:
        FAIL_COUNT += 1
        if critical:
            CRITICAL_FAILED = True
        detail_str = f": {detail[:200]}" if detail else ""
        print(f"  [FAIL] {tag}{name}{detail_str}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def parse_num(cell):
    """Extract a float from a table cell that may contain currency/thousands sep."""
    if cell is None:
        return None
    s = str(cell).strip().replace(" ", "").replace(" ", "")
    s = s.replace("RUB", "").replace("руб", "").replace("₽", "").replace("$", "")
    s = s.replace(",", "")
    m = re.search(r"-?\d+(?:\.\d+)?", s)
    if not m:
        return None
    try:
        return float(m.group(0))
    except ValueError:
        return None


def fetch_db_actuals():
    """SUM(TOTAL_AMOUNT) and COUNT(ORDER_ID) per RU segment for Q4 2025."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT c."SEGMENT", COUNT(o."ORDER_ID"), ROUND(SUM(o."TOTAL_AMOUNT")::numeric, 2)
        FROM sf_data."SALES_DW__PUBLIC__ORDERS" o
        JOIN sf_data."SALES_DW__PUBLIC__CUSTOMERS" c ON o."CUSTOMER_ID" = c."CUSTOMER_ID"
        WHERE o."ORDER_DATE" >= '2025-10-01' AND o."ORDER_DATE" <= '2025-12-31'
        GROUP BY c."SEGMENT"
    """)
    out = {}
    for seg, cnt, actual in cur.fetchall():
        out[seg] = {"count": int(cnt), "actual": float(actual)}
    conn.close()
    return out


def check_word_doc(agent_workspace, db):
    """Check the Word document structure and content (semantic, critical)."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Q4_Segment_Report.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}", critical=True)
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Heading: accept RU or EN variant.
    has_heading = False
    for p in doc.paragraphs:
        low = p.text.lower()
        if "q4 2025" in low and ("segment" in low or "сегмент" in low):
            has_heading = True
            break
    check("Document has Q4 2025 segment heading (RU/EN)", has_heading)

    check("Document has at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 1:
        return False

    table = doc.tables[0]
    rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    check("Table has 4 segment rows", len(rows) == 4, f"Got {len(rows)} rows")

    # --- CRITICAL: segment set ---
    table_segments = set()
    for r in rows:
        if r:
            table_segments.add(r[0].strip())
    seg_set_ok = table_segments == EXPECTED_SEGMENTS
    check("Table contains exactly the 4 RU segments (each once)",
          seg_set_ok and len(rows) == 4,
          f"Got {sorted(table_segments)}", critical=True)

    # Per-segment semantic checks (use DB actuals + PDF targets).
    for segment, target in TARGETS.items():
        info = db.get(segment)
        if info is None:
            check(f"DB has data for {segment}", False, "no DB rows", critical=True)
            continue
        actual = info["actual"]
        count = info["count"]
        variance = round(actual - target, 2)
        achievement = round(actual / target * 100, 1)

        matched = None
        for r in rows:
            if r and str_match(r[0], segment):
                matched = r
                break
        if not matched:
            check(f"Segment {segment} found in table", False, critical=True)
            continue

        nums = [parse_num(c) for c in matched[1:]]
        nums = [n for n in nums if n is not None]

        # CRITICAL: Q4_Actual present.
        check(f"Segment {segment} Q4_Actual matches DB (~{actual})",
              any(num_close(n, actual, 50.0) for n in nums),
              f"row nums={nums}", critical=True)
        # CRITICAL: Variance present (= actual - target, negative).
        check(f"Segment {segment} Variance == {variance}",
              any(num_close(n, variance, 50.0) for n in nums),
              f"row nums={nums}", critical=True)
        # CRITICAL: Achievement_Pct present.
        check(f"Segment {segment} Achievement_Pct == {achievement}",
              any(num_close(n, achievement, 0.5) for n in nums),
              f"row nums={nums}", critical=True)
        # NON-critical: Order_Count present.
        check(f"Segment {segment} Order_Count == {count}",
              any(num_close(n, count, 0.5) for n in nums),
              f"row nums={nums}")
        # NON-critical: target present.
        check(f"Segment {segment} Q4_Target == {target}",
              any(num_close(n, target, 0.5) for n in nums),
              f"row nums={nums}")

    # --- Summary paragraph (CRITICAL on substance) ---
    full_text = " ".join(p.text for p in doc.paragraphs)
    flat = full_text.replace(" ", "").replace(" ", "").replace(",", "").lower()

    total_target = sum(TARGETS.values())                       # 760000
    total_actual = round(sum(v["actual"] for v in db.values()), 2)  # ~393653.23

    has_total_target = "760000" in flat
    # total_actual may appear as 393653.23 or 393653.2 etc.
    has_total_actual = ("393653" in flat)
    # 0 segments met target.
    n_meeting = sum(1 for seg, t in TARGETS.items()
                    if db.get(seg, {}).get("actual", 0) >= t)
    # accept "0 сегментов" / "0 segments" / "ни один (из ... ) сегментов"
    low_text = full_text.lower()
    has_zero = bool(
        re.search(r"(?:0|ноль|ни\s+од(?:ин|ного|на|но)(?:\s+из\s+\S+)?)[^.]{0,40}сегмент", low_text)
        or re.search(r"0\s+segment", low_text)
    )

    check("Summary states total target 760000", has_total_target,
          f"text={full_text[:200]}", critical=True)
    check("Summary states total actual ~393653.23", has_total_actual,
          f"text={full_text[:200]}", critical=True)
    check(f"Summary states {n_meeting} segments meeting target", has_zero and n_meeting == 0,
          f"n_meeting={n_meeting}; text={full_text[:200]}", critical=True)

    return True


def check_excel_groundtruth(groundtruth_workspace, db):
    """Cross-check groundtruth Excel for numeric accuracy (non-critical)."""
    print("\n=== Cross-checking with Groundtruth Data ===")

    gt_file = os.path.join(groundtruth_workspace, "Q4_Segment_Data.xlsx")
    if not os.path.isfile(gt_file):
        check("Groundtruth file exists", False)
        return

    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)
    gt_rows = list(gt_wb["Segment Comparison"].iter_rows(min_row=2, values_only=True))

    for segment, info in db.items():
        actual = info["actual"]
        gt_match = None
        for r in gt_rows:
            if str_match(r[0], segment):
                gt_match = r
                break
        check(f"Groundtruth {segment} actual matches DB",
              gt_match is not None and num_close(gt_match[2], actual, 1.0),
              f"GT={gt_match[2] if gt_match else None}, DB={actual}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("SF SALES SEGMENT WORD - EVALUATION (ClickHouse / RU)")
    print("=" * 70)

    db = fetch_db_actuals()

    check_word_doc(args.agent_workspace, db)
    check_excel_groundtruth(gt_dir, db)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Accuracy: {accuracy:.1f}%")
    print(f"  Critical failed: {CRITICAL_FAILED}")

    if CRITICAL_FAILED:
        print("  Overall: FAIL (critical check failed)")
        sys.exit(1)

    overall = accuracy >= 70.0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
