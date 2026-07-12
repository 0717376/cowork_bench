"""Evaluation for moex-earnings-report-word-teamly task.

The agent pulls annual income-statement data for 3 MOEX tickers from the
moex-finance MCP, builds Earnings_Report.docx (heading, one section+table per
company with columns Period/Revenue_B/Net_Income_B/Gross_Profit_B/
Operating_Income_B sorted ascending, plus a RUB summary paragraph), and creates
a teamly knowledge-base page "Earnings Analysis Key Insights".

Critical checks (see CRITICAL_CHECKS): any failure => overall FAIL regardless
of accuracy. Otherwise pass threshold: accuracy >= 70%.

Values in moex.financial_statements are raw RUB; /1e9 -> billions of RUB.
Tickers used: SBER.ME, GAZP.ME, LKOH.ME.
Highest 2025 revenue and highest 2025 net income both -> GAZP.ME.
"""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB = dict(
    host=os.environ.get("PGHOST", "localhost"),
    port=5432,
    dbname=os.environ.get("PGDATABASE", "cowork_gym"),
    user="eigent",
    password="camel",
)

SYMBOLS = ["SBER.ME", "GAZP.ME", "LKOH.ME"]
EXPECTED_COLS = ["Period", "Revenue_B", "Net_Income_B", "Gross_Profit_B", "Operating_Income_B"]

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Semantic critical checks: any failure => overall FAIL regardless of accuracy.
CRITICAL_CHECKS = {
    "SBER.ME latest revenue + net income present",
    "GAZP.ME latest revenue + net income present",
    "LKOH.ME latest revenue + net income present",
    "Teamly Insights page exists with highest-revenue ticker (GAZP.ME)",
    "Teamly Insights page names highest-net-income ticker (GAZP.ME)",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        d = (str(detail)[:300]) if detail else ""
        print(f"  [FAIL] {name}: {d}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def parse_float(cell):
    try:
        return float(str(cell).replace(",", "").replace("$", "").replace(" ", "").strip())
    except (ValueError, AttributeError):
        return None


def value_in_tables(doc, target, tol):
    for table in doc.tables:
        for row in table.rows[1:]:
            for cell in row.cells:
                v = parse_float(cell.text.strip())
                if v is not None and num_close(v, target, tol):
                    return True
    return False


def load_gt(groundtruth_workspace):
    gt_file = os.path.join(groundtruth_workspace, "Earnings_Data.xlsx")
    if not os.path.isfile(gt_file):
        return None
    wb = openpyxl.load_workbook(gt_file, data_only=True)
    gt = {}
    for sym in SYMBOLS:
        if sym not in wb.sheetnames:
            continue
        rows = list(wb[sym].iter_rows(min_row=2, values_only=True))
        if rows:
            gt[sym] = rows  # each row: (Period, Revenue_B, Net_Income_B, Gross_Profit_B, Operating_Income_B)
    return gt


def check_word_doc(agent_workspace, gt):
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return

    doc_path = os.path.join(agent_workspace, "Earnings_Report.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}")
    if not os.path.isfile(doc_path):
        return
    doc = Document(doc_path)

    # Heading (structural, non-critical)
    has_heading = any(
        "earnings" in p.text.lower() and ("analysis" in p.text.lower() or "report" in p.text.lower())
        for p in doc.paragraphs
    )
    check("Document has earnings analysis heading", has_heading)

    # At least 3 tables (structural)
    check("Document has at least 3 tables", len(doc.tables) >= 3, f"Found {len(doc.tables)} tables")

    # Each table header contains the full required schema (structural but verifies layout)
    full_schema_tables = 0
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        header_join = " ".join(header)
        if all(col.lower() in header_join for col in EXPECTED_COLS):
            full_schema_tables += 1
    check("At least 3 tables have full column schema (Period/Revenue_B/Net_Income_B/Gross_Profit_B/Operating_Income_B)",
          full_schema_tables >= 3, f"Tables with full schema: {full_schema_tables}")

    # Each company section heading "<TICKER> Annual Earnings" (structural)
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    for sym in SYMBOLS:
        check(f"Document has section for {sym}", sym.lower() in full_text)

    # Per-company latest-year (2025) revenue + net income within tight tolerance (CRITICAL).
    if gt:
        for sym in SYMBOLS:
            rows = gt.get(sym)
            if not rows:
                check(f"{sym} latest revenue + net income present", False, "no groundtruth rows")
                continue
            latest = rows[-1]
            latest_rev = latest[1]
            latest_ni = latest[2]
            rev_ok = value_in_tables(doc, latest_rev, 0.5)
            ni_ok = value_in_tables(doc, latest_ni, 0.5)
            check(f"{sym} latest revenue + net income present", rev_ok and ni_ok,
                  f"rev({latest_rev})={rev_ok}, ni({latest_ni})={ni_ok}")

            # Both 2024 & 2025 period rows present, ascending sort (non-critical correctness).
            periods = [r[0] for r in rows]
            found_periods = []
            for table in doc.tables:
                cells_text = [c.text.strip() for row in table.rows for c in row.cells]
                joined = " ".join(cells_text)
                if value_in_tables(doc, latest_rev, 0.5) and all(
                    str(p)[:10] in joined for p in periods
                ):
                    # at least one table contains all periods for this company
                    found_periods = periods
                    break
            check(f"{sym} has 2024 and 2025 period rows", len(found_periods) >= 2,
                  f"expected periods {periods}")
    else:
        check("Groundtruth Earnings_Data.xlsx loaded", False)

    # Summary paragraph mentions revenue/выручка (structural)
    has_summary = any(
        ("revenue" in p.text.lower() or "выручк" in p.text.lower() or "прибыл" in p.text.lower())
        for p in doc.paragraphs if len(p.text) > 50
    )
    check("Document has summary paragraph", has_summary)


def check_teamly(gt):
    print("\n=== Checking Teamly Knowledge Base ===")
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("SELECT title, COALESCE(body, '') FROM teamly.pages")
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        check("Teamly Insights page exists with highest-revenue ticker (GAZP.ME)", False, str(e))
        check("Teamly Insights page names highest-net-income ticker (GAZP.ME)", False, str(e))
        return

    # Find the deliverable page by title.
    insights = [
        (t, b) for t, b in rows
        if t and ("earnings" in t.lower() or "insights" in t.lower())
    ]
    check("At least 1 teamly Insights page created", len(insights) >= 1, f"Total pages: {len(rows)}")

    # Compute highest 2025 revenue / net income tickers from groundtruth.
    top_rev_sym = None
    top_ni_sym = None
    if gt:
        best_rev = best_ni = -1e30
        for sym, grows in gt.items():
            latest = grows[-1]
            if latest[1] is not None and float(latest[1]) > best_rev:
                best_rev = float(latest[1]); top_rev_sym = sym
            if latest[2] is not None and float(latest[2]) > best_ni:
                best_ni = float(latest[2]); top_ni_sym = sym

    page_text = " ".join((str(t) + " " + str(b)) for t, b in insights).lower()

    # CRITICAL: page exists AND body names the highest-revenue ticker.
    rev_ok = bool(insights) and top_rev_sym is not None and (
        top_rev_sym.lower() in page_text or top_rev_sym.split(".")[0].lower() in page_text
    )
    check(f"Teamly Insights page exists with highest-revenue ticker ({top_rev_sym})", rev_ok,
          f"top_rev={top_rev_sym}, page_text_present={bool(insights)}")

    # CRITICAL: body names the highest-net-income ticker.
    ni_ok = bool(insights) and top_ni_sym is not None and (
        top_ni_sym.lower() in page_text or top_ni_sym.split(".")[0].lower() in page_text
    )
    check(f"Teamly Insights page names highest-net-income ticker ({top_ni_sym})", ni_ok,
          f"top_ni={top_ni_sym}")

    # Non-critical: page mentions revenue / выручка.
    has_rev_word = "revenue" in page_text or "выручк" in page_text
    check("Teamly Insights page mentions revenue", has_rev_word)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("MOEX EARNINGS REPORT WORD TEAMLY - EVALUATION")
    print("=" * 70)

    gt = load_gt(gt_dir)
    check_word_doc(args.agent_workspace, gt)
    check_teamly(gt)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}, Accuracy: {accuracy:.1f}%")
    if critical_failed:
        print(f"  CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"    - {n}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({
                "total_passed": PASS_COUNT,
                "total_checks": total,
                "accuracy": accuracy,
                "critical_failed": critical_failed,
                "success": (not critical_failed) and accuracy >= 70,
            }, f, indent=2)

    if critical_failed:
        print("Overall: FAIL (critical check failed)")
        sys.exit(1)
    if accuracy >= 70:
        print("Overall: PASS")
        sys.exit(0)
    print("Overall: FAIL")
    sys.exit(1)


if __name__ == "__main__":
    main()
