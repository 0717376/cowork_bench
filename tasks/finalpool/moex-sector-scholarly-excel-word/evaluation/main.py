"""Evaluation script for moex-sector-scholarly-excel-word.

Критические проверки (CRITICAL_CHECKS): любой провал => общий FAIL независимо
от accuracy. Иначе PASS требует accuracy >= 70%.

Замечание по волатильным данным: цены MOEX (RUB) читаются «честно» из live-данных
и не сверяются с захардкоженным groundtruth. Вместо равенства конкретным ценам
проверяются структура, диапазоны-санити (положительные цены, Total == Avg*Count)
и принадлежность enum-значений допустимым множествам.
"""
import os
import argparse, json, os, sys
import openpyxl

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

VALIDATION_VALUES = {"confirmed", "partial", "inconclusive"}
OUTLOOK_VALUES = {"bullish", "neutral", "bearish"}

# Заголовки разделов docx: пары RU+EN, принимаем любой вариант.
REQUIRED_HEADINGS = [
    ["Cross-Disciplinary Sector Analysis", "Кросс-дисциплинарный секторный анализ", "Секторный анализ"],
    ["Financial Performance Review", "Обзор финансовой динамики", "Финансовая динамика"],
    ["Academic Research Insights", "Выводы академических исследований", "Академические исследования"],
    ["Theory vs Practice Comparison", "Сравнение теории и практики", "Теория и практика"],
    ["Investment Implications", "Инвестиционные выводы", "Инвестиционные импликации"],
]

# Критические проверки: любой провал => FAIL независимо от accuracy.
CRITICAL_CHECKS = {
    "Sector_Analysis_Report.xlsx exists",
    "Sector_Performance sheet present with required headers",
    "Sector_Performance has >= 4 sector rows",
    "Sector_Performance Total_Market_Value == Avg_Price * Stock_Count (all rows)",
    "Sector_Performance positive Avg_Price (all rows)",
    "Research_Mapping has >= 2 valid rows (Paper_Title/Key_Finding/Applicable_Sector + Validation_Status enum)",
    "Investment_Thesis Outlook enum + non-empty Supporting_Evidence/Risk_Factor (all rows)",
    "Sector_Research_Brief.docx has heading \"Cross-Disciplinary Sector Analysis\"",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)


def header_index(ws):
    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
    return {h: i for i, h in enumerate(headers) if h}, headers


def check_excel(agent_workspace):
    excel_path = os.path.join(agent_workspace, "Sector_Analysis_Report.xlsx")
    exists = os.path.exists(excel_path)
    check("Sector_Analysis_Report.xlsx exists", exists)
    if not exists:
        return
    wb = openpyxl.load_workbook(excel_path)

    # --- Sector_Performance ---
    sp_ok = "Sector_Performance" in wb.sheetnames
    required_sp = ["sector", "stock_count", "avg_price", "total_market_value", "volatility_score"]
    sp_headers_ok = False
    if sp_ok:
        ws = wb["Sector_Performance"]
        hmap, headers = header_index(ws)
        sp_headers_ok = all(h in hmap for h in required_sp)
        check("Sector_Performance sheet present with required headers", sp_headers_ok, f"headers: {headers[:10]}")
        rows = [r for r in ws.iter_rows(min_row=2, values_only=True) if any(c is not None and str(c).strip() for c in r)]
        check("Sector_Performance has >= 4 sector rows", len(rows) >= 4, f"got {len(rows)}")

        prod_ok = True
        pos_ok = True
        for r in rows:
            avg = safe_float(r[hmap["avg_price"]]) if hmap["avg_price"] < len(r) else None
            cnt = safe_float(r[hmap["stock_count"]]) if hmap["stock_count"] < len(r) else None
            tmv = safe_float(r[hmap["total_market_value"]]) if hmap["total_market_value"] < len(r) else None
            if avg is None or avg <= 0:
                pos_ok = False
            if avg is None or cnt is None or tmv is None:
                prod_ok = False
            else:
                tol = max(0.5, abs(avg * cnt) * 0.02)
                if abs(tmv - avg * cnt) > tol:
                    prod_ok = False
        check("Sector_Performance Total_Market_Value == Avg_Price * Stock_Count (all rows)", prod_ok)
        check("Sector_Performance positive Avg_Price (all rows)", pos_ok)
        # volatility presence sanity (non-critical)
        vol_ok = all(safe_float(r[hmap["volatility_score"]]) is not None
                     for r in rows if hmap["volatility_score"] < len(r))
        check("Sector_Performance Volatility_Score numeric (all rows)", vol_ok)
    else:
        check("Sector_Performance sheet present with required headers", False, "sheet missing")
        check("Sector_Performance has >= 4 sector rows", False)
        check("Sector_Performance Total_Market_Value == Avg_Price * Stock_Count (all rows)", False)
        check("Sector_Performance positive Avg_Price (all rows)", False)

    # --- Research_Mapping ---
    rm_ok = "Research_Mapping" in wb.sheetnames
    check("Research_Mapping sheet exists", rm_ok)
    if rm_ok:
        ws = wb["Research_Mapping"]
        hmap, headers = header_index(ws)
        req = ["paper_title", "key_finding", "applicable_sector", "validation_status"]
        check("Research_Mapping has required headers", all(h in hmap for h in req), f"headers: {headers[:10]}")
        rows = [r for r in ws.iter_rows(min_row=2, values_only=True) if any(c is not None and str(c).strip() for c in r)]
        valid = 0
        for r in rows:
            try:
                title = str(r[hmap["paper_title"]] or "").strip()
                finding = str(r[hmap["key_finding"]] or "").strip()
                sector = str(r[hmap["applicable_sector"]] or "").strip()
                vs = str(r[hmap["validation_status"]] or "").strip().lower()
            except (IndexError, KeyError):
                continue
            if title and finding and sector and vs in VALIDATION_VALUES:
                valid += 1
        check("Research_Mapping has >= 2 valid rows (Paper_Title/Key_Finding/Applicable_Sector + Validation_Status enum)",
              valid >= 2, f"valid rows: {valid}")
    else:
        check("Research_Mapping has required headers", False)
        check("Research_Mapping has >= 2 valid rows (Paper_Title/Key_Finding/Applicable_Sector + Validation_Status enum)", False)

    # --- Investment_Thesis ---
    it_ok = "Investment_Thesis" in wb.sheetnames
    check("Investment_Thesis sheet exists", it_ok)
    if it_ok:
        ws = wb["Investment_Thesis"]
        hmap, headers = header_index(ws)
        req = ["sector", "outlook", "supporting_evidence", "risk_factor"]
        check("Investment_Thesis has required headers", all(h in hmap for h in req), f"headers: {headers[:10]}")
        rows = [r for r in ws.iter_rows(min_row=2, values_only=True) if any(c is not None and str(c).strip() for c in r)]
        all_ok = len(rows) >= 1
        for r in rows:
            try:
                outlook = str(r[hmap["outlook"]] or "").strip().lower()
                ev = str(r[hmap["supporting_evidence"]] or "").strip()
                risk = str(r[hmap["risk_factor"]] or "").strip()
            except (IndexError, KeyError):
                all_ok = False
                continue
            if outlook not in OUTLOOK_VALUES or not ev or not risk:
                all_ok = False
        check("Investment_Thesis Outlook enum + non-empty Supporting_Evidence/Risk_Factor (all rows)",
              all_ok, f"rows: {len(rows)}")
    else:
        check("Investment_Thesis has required headers", False)
        check("Investment_Thesis Outlook enum + non-empty Supporting_Evidence/Risk_Factor (all rows)", False)


def check_docx(agent_workspace):
    docx_path = os.path.join(agent_workspace, "Sector_Research_Brief.docx")
    exists = os.path.exists(docx_path)
    check("Sector_Research_Brief.docx exists", exists)
    if not exists:
        check("Sector_Research_Brief.docx has heading \"Cross-Disciplinary Sector Analysis\"", False)
        return
    from docx import Document
    doc = Document(docx_path)
    text = " ".join([p.text for p in doc.paragraphs])
    check("Sector_Research_Brief.docx has content", len(text) > 50, f"text length: {len(text)}")
    # Все заголовки + полный текст (RU+EN). Заголовки могут быть оформлены не как
    # Heading-стиль (Word MCP), поэтому сопоставляем по тексту документа тоже.
    headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
    hay = (text + " " + " ".join(headings)).lower()
    for variants in REQUIRED_HEADINGS:
        found = any(v.lower() in hay for v in variants)
        primary = variants[0]
        name = f"Sector_Research_Brief.docx has heading \"{primary}\""
        check(name, found, f"variants: {variants}")


def check_script(agent_workspace):
    try:
        py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    except OSError:
        py_files = []
    check("Python analysis script exists", len(py_files) >= 1, f"found: {py_files}")


def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT, FAILED_NAMES
    PASS_COUNT = 0
    FAIL_COUNT = 0
    FAILED_NAMES = []

    check_excel(agent_workspace)
    check_docx(agent_workspace)
    check_script(agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total else 0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== Results: {PASS_COUNT}/{total} passed ({accuracy:.1f}%) ===")
    if critical_failed:
        print(f"CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")

    if res_log_file:
        try:
            with open(res_log_file, "w") as f:
                json.dump({
                    "total_passed": PASS_COUNT, "total_checks": total,
                    "accuracy": accuracy, "critical_failed": critical_failed,
                }, f, indent=2)
        except Exception:
            pass

    success = (not critical_failed) and accuracy >= 70
    return success, f"Passed {PASS_COUNT}/{total} checks ({accuracy:.1f}%)"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
