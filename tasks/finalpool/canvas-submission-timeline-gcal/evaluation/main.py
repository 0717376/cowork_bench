"""
Оценка задачи canvas-submission-timeline-gcal.

Критические проверки (CRITICAL_CHECKS): провал любой => общий FAIL независимо
от accuracy. В остальном PASS требует accuracy >= 70%.

Данные о курсах, заданиях и сдачах поступают с ЖИВОГО сервера Canvas (внешний
server-side seed, схема canvas.*), а НЕ из preprocess этой задачи. Поэтому
ожидаемые числа (всего сдач, просроченных сдач, число заданий со сроком сдачи)
НЕ зашиты в eval, а пересчитываются ВЖИВУЮ из БД Canvas тем же способом, что
использует агент: поле canvas.submissions.late. Сравнение значений в документе
ведётся с допуском (проценты округляются по-разному, ',' vs '.').

Проверки:
1. Word-документ Late_Submission_Report.docx: заголовок, маркер 'late submission',
   все 6 кодов курсов Fall 2013, итоговый раздел, и — критично — числа всего/late
   по каждому курсу совпадают с пересчитанными вживую из Canvas (с допуском).
2. Google Calendar: события в формате '[Course_Code]: [Assignment_Name] Due',
   >= 30 событий, покрытие >= 4 кодов; число событий согласуется с числом
   заданий Fall 2013, у которых задан срок сдачи (с допуском).
"""

import argparse
import json
import os
import re
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

FALL_2013_CODES = ["AAA-2013J", "BBB-2013J", "DDD-2013J", "EEE-2013J", "FFF-2013J", "GGG-2013J"]

# Заголовок документа: принимаем английский (грепается eval) ИЛИ русский эквивалент.
TITLE_EN = "fall 2013 late submission analysis"
TITLE_RU = "анализ просроченных сдач"

# Ключевые слова итогового/сводного раздела (RU + EN).
OVERALL_KEYWORDS = ["overall", "summary", "total", "итог", "итого", "сводк", "всего по"]

# Маркер «Due» в заголовке события (EN, грепается) или RU-варианты.
DUE_MARKERS = ["due", "срок", "дедлайн"]

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Критические (семантические) проверки: провал любой => общий FAIL.
CRITICAL_CHECKS = {
    "Word doc has title 'Fall 2013 Late Submission Analysis' (RU equiv accepted)",
    "Word doc contains all 6 Fall-2013 course codes",
    "Per-course late & total counts in doc match live Canvas recompute (tolerant)",
    "Word doc has an overall/summary section across all courses",
    "Calendar events use summary format '[Code]: [Name] Due' (RU 'Срок' accepted)",
    "Calendar event count consistent with Fall-2013 assignments having a due date",
}


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        msg = f": {str(detail)[:300]}" if detail else ""
        tag = " (CRITICAL)" if critical or name in CRITICAL_CHECKS else ""
        print(f"  [FAIL]{tag} {name}{msg}")


# ---------------------------------------------------------------------------
# Живой пересчёт из Canvas (server-side seed). Числа НЕ хардкодятся.
# ---------------------------------------------------------------------------
def canvas_live_stats():
    """Возвращает (per_course, due_assignments_count).

    per_course: {code_lower: {'total': int, 'late': int}} — всего и просроченных
                сдач по каждому курсу Fall 2013.
    due_assignments_count: число заданий Fall 2013 с заданным сроком сдачи
                (или None, если поле срока недоступно — тогда проверка не блокирует).
    """
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    per_course = {}
    cur.execute("""
        SELECT LOWER(c.course_code),
               COUNT(s.id) AS total,
               SUM(CASE WHEN s.late THEN 1 ELSE 0 END) AS late
        FROM canvas.courses c
        JOIN canvas.assignments a ON a.course_id = c.id
        JOIN canvas.submissions s ON s.assignment_id = a.id
        WHERE c.course_code LIKE %s
        GROUP BY LOWER(c.course_code)
    """, ("%2013J",))
    for code, total, late in cur.fetchall():
        per_course[code] = {"total": int(total or 0), "late": int(late or 0)}

    # Число заданий Fall 2013 со сроком сдачи. Имя столбца срока в стандартной
    # схеме Canvas — due_at; делаем устойчиво к отсутствию столбца.
    due_count = None
    for col in ("due_at", "due_date"):
        try:
            cur.execute(f"""
                SELECT COUNT(*)
                FROM canvas.assignments a
                JOIN canvas.courses c ON c.id = a.course_id
                WHERE c.course_code LIKE %s AND a.{col} IS NOT NULL
            """, ("%2013J",))
            due_count = int(cur.fetchone()[0] or 0)
            break
        except Exception:
            conn.rollback()
            continue

    cur.close()
    conn.close()
    return per_course, due_count


def _nums_in(text):
    """Все целые числа в фрагменте текста (для tolerant-сравнения).

    Сначала убираем разделители групп разрядов между цифрами (пробел, NBSP
    U+00A0, узкий неразрывный пробел U+202F, тонкий пробел U+2009, '.', ','),
    чтобы русская локаль «1 633» / «14 375» парсилась как одно число, а не {1,633}.
    """
    t = re.sub(r"(?<=\d)[\s   .,](?=\d{3}(?:\D|$))", "", text or "")
    return set(int(x) for x in re.findall(r"\d+", t))


def check_word(agent_workspace, per_course):
    print("\n=== Проверка Word-документа ===")
    docx_path = os.path.join(agent_workspace, "Late_Submission_Report.docx")
    exists = os.path.isfile(docx_path)
    check("Word file exists", exists, f"Expected {docx_path}")
    if not exists:
        # Зависимые критические проверки тоже валим явно.
        check("Word doc has title 'Fall 2013 Late Submission Analysis' (RU equiv accepted)", False, "no docx")
        check("Word doc contains all 6 Fall-2013 course codes", False, "no docx")
        check("Per-course late & total counts in doc match live Canvas recompute (tolerant)", False, "no docx")
        check("Word doc has an overall/summary section across all courses", False, "no docx")
        return

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        check("Word file readable", False, str(e))
        check("Word doc has title 'Fall 2013 Late Submission Analysis' (RU equiv accepted)", False, "unreadable")
        check("Word doc contains all 6 Fall-2013 course codes", False, "unreadable")
        check("Per-course late & total counts in doc match live Canvas recompute (tolerant)", False, "unreadable")
        check("Word doc has an overall/summary section across all courses", False, "unreadable")
        return

    # Текст по абзацам (для посекционного сопоставления) и общий текст.
    para_texts = [p.text for p in doc.paragraphs]
    table_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text += " " + cell.text
    all_text = " ".join(para_texts) + " " + table_text
    lower = all_text.lower()

    # --- Заголовок (EN грепается eval, RU принимается как альтернатива) [CRITICAL] ---
    check("Word doc has title 'Fall 2013 Late Submission Analysis' (RU equiv accepted)",
          TITLE_EN in lower or TITLE_RU in lower,
          f"Sample: {lower[:200]}")

    # --- Маркер 'late submission' (структурный, NON-critical) ---
    check("Word contains 'late submission'",
          "late submission" in lower or "просроч" in lower or "опоздан" in lower,
          f"Sample: {lower[:200]}")

    # --- Все 6 кодов курсов [CRITICAL] ---
    missing = [c for c in FALL_2013_CODES if c.lower() not in lower]
    check("Word doc contains all 6 Fall-2013 course codes",
          not missing,
          f"Missing: {missing}")

    # --- Процент просрочки присутствует (структурный) ---
    check("Word doc reports a late-submission percentage ('%')",
          "%" in all_text or "процент" in lower,
          "No '%' / 'процент' found")

    # --- Итоговый/сводный раздел [CRITICAL] ---
    check("Word doc has an overall/summary section across all courses",
          any(k in lower for k in OVERALL_KEYWORDS),
          f"Keywords {OVERALL_KEYWORDS} not found")

    # --- Числа всего/late совпадают с живым пересчётом из Canvas [CRITICAL] ---
    # Для каждого курса находим в документе фрагмент, относящийся к коду курса,
    # и проверяем, что в нём встречаются ожидаемые total и late (с допуском).
    # Допуск: точное совпадение ИЛИ в пределах +-1 (округление/расхождение выборки).
    blocks = para_texts + [c.text for t in doc.tables for r in t.rows for c in r.cells]

    def value_present(target, texts):
        for t in texts:
            for n in _nums_in(t):
                if abs(n - target) <= 1:
                    return True
        return False

    def section_for(cl):
        """Контекст секции курса: блок с кодом + последующие блоки до следующего
        кода курса (но не более 8 блоков). Покрывает оба макета: код и числа в
        одном блоке (ячейка таблицы) ИЛИ код-заголовок и числа в соседних абзацах.
        """
        other = [c.lower() for c in FALL_2013_CODES if c.lower() != cl]
        ctx = []
        for i, b in enumerate(blocks):
            if cl in b.lower():
                ctx.append(b)
                for nb in blocks[i + 1:i + 9]:
                    nl = nb.lower()
                    if any(o in nl for o in other):
                        break
                    ctx.append(nb)
        return ctx

    matched_courses = 0
    detail_bits = []
    if not per_course:
        # Не смогли пересчитать (БД недоступна/пусто) — не блокируем критически,
        # но фиксируем как провал значения, чтобы не считать пройденным «честно».
        check("Per-course late & total counts in doc match live Canvas recompute (tolerant)",
              False, "Live Canvas recompute returned no Fall-2013 data")
    else:
        for code in FALL_2013_CODES:
            cl = code.lower()
            stats = per_course.get(cl)
            if not stats:
                detail_bits.append(f"{code}:no-canvas-data")
                continue
            ctx = section_for(cl)
            if not ctx:
                detail_bits.append(f"{code}:no-section")
                continue
            total_ok = value_present(stats["total"], ctx)
            late_ok = value_present(stats["late"], ctx)
            if total_ok and late_ok:
                matched_courses += 1
            else:
                detail_bits.append(
                    f"{code}:total={stats['total']}({'ok' if total_ok else 'MISS'}),"
                    f"late={stats['late']}({'ok' if late_ok else 'MISS'})")
        # Критично: значения сходятся как минимум для большинства курсов (>=5/6),
        # чтобы стерпеть единичные расхождения в форматировании секции.
        check("Per-course late & total counts in doc match live Canvas recompute (tolerant)",
              matched_courses >= 5,
              f"matched {matched_courses}/6; issues: {detail_bits}")


def check_calendar(due_count):
    print("\n=== Проверка Google Calendar ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("SELECT summary, start_datetime FROM gcal.events")
        events = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        check("Calendar reachable", False, str(e))
        check("Calendar events use summary format '[Code]: [Name] Due' (RU 'Срок' accepted)", False, str(e))
        check("Calendar event count consistent with Fall-2013 assignments having a due date", False, str(e))
        return

    print(f"[check_calendar] Найдено событий: {len(events)}")

    # События с кодом курса в summary.
    course_events = []
    for summary, _start in events:
        if not summary:
            continue
        s_lower = summary.lower()
        if any(code.lower() in s_lower for code in FALL_2013_CODES):
            course_events.append(summary)

    # --- Структурные (NON-critical) ---
    check("At least 30 calendar events with course codes",
          len(course_events) >= 30,
          f"Found {len(course_events)} events with course codes")

    codes_found = set()
    for summary in course_events:
        for code in FALL_2013_CODES:
            if code.lower() in summary.lower():
                codes_found.add(code)
    check("Events cover at least 4 different course codes",
          len(codes_found) >= 4,
          f"Found codes: {codes_found}")

    # --- Формат заголовка '[Code]: [Name] Due' [CRITICAL] ---
    # Требуем: код курса, двоеточие после кода, и Due-маркер (EN/RU) в summary.
    well_formed = 0
    for summary in course_events:
        s_lower = summary.lower()
        code_with_colon = any(
            re.search(re.escape(code.lower()) + r"\s*:", s_lower)
            for code in FALL_2013_CODES
        )
        has_due = any(m in s_lower for m in DUE_MARKERS)
        if code_with_colon and has_due:
            well_formed += 1
    # «Bulk»: не менее 80% событий с кодом курса соответствуют формату.
    fmt_ok = course_events and well_formed >= max(30, int(0.8 * len(course_events)))
    check("Calendar events use summary format '[Code]: [Name] Due' (RU 'Срок' accepted)",
          bool(fmt_ok),
          f"well_formed={well_formed}/{len(course_events)}")

    # --- Число событий согласуется с числом заданий со сроком сдачи [CRITICAL] ---
    if due_count is None:
        # Не смогли определить столбец срока — не блокируем жёстко по точному числу,
        # но требуем, чтобы число событий было правдоподобным (>=30).
        check("Calendar event count consistent with Fall-2013 assignments having a due date",
              len(course_events) >= 30,
              "due-date column unavailable; fell back to >=30 sanity check")
    else:
        # Допуск: события должны покрывать почти все задания со сроком и не сильно
        # их превышать. |events - due_count| <= max(3, 15% от due_count).
        tol = max(3, int(0.15 * due_count))
        ok = abs(len(course_events) - due_count) <= tol and len(course_events) >= max(1, due_count - tol)
        check("Calendar event count consistent with Fall-2013 assignments having a due date",
              bool(ok),
              f"events={len(course_events)}, due_assignments={due_count}, tol={tol}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    # Живой пересчёт из Canvas (server-side seed); при сбое БД — пустые значения,
    # соответствующие проверки тогда честно провалятся.
    try:
        per_course, due_count = canvas_live_stats()
        print(f"[canvas_live_stats] per_course={per_course}, due_count={due_count}")
    except Exception as e:
        print(f"[canvas_live_stats] ОШИБКА живого пересчёта: {e}")
        per_course, due_count = {}, None

    check_word(args.agent_workspace, per_course)
    check_calendar(due_count)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total else 0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== Результат: {PASS_COUNT}/{total} пройдено ({accuracy:.1f}%) ===")

    if args.res_log_file:
        result = {
            "passed": PASS_COUNT,
            "failed": FAIL_COUNT,
            "accuracy": accuracy,
            "critical_failed": critical_failed,
            "success": (not critical_failed) and accuracy >= 70,
        }
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if critical_failed:
        print(f"КРИТИЧЕСКИЕ ПРОВАЛЫ: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")
        sys.exit(1)

    if accuracy >= 70:
        print("Все условия выполнены (нет критических провалов, accuracy >= 70%).")
        sys.exit(0)
    else:
        print(f"accuracy {accuracy:.1f}% < 70%")
        sys.exit(1)


if __name__ == "__main__":
    main()
