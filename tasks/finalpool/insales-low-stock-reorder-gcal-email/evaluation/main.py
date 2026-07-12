"""Evaluation for insales-low-stock-reorder-gcal-email (InSales / wc.* schema)."""
import argparse
import os
import sys
import psycopg2
import openpyxl


DB = {"host": os.environ.get("PGHOST", "localhost"), "port": 5432, "dbname": "cowork_gym", "user": "eigent", "password": "camel"}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRITICAL]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL]{' [CRITICAL]' if critical else ''} {name}: {detail_str}")
        if critical:
            CRITICAL_FAILED.append(name)


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def get_low_stock_counts():
    """Live counts from wc.products: stock_quantity<=5 bucketed 0 / 1-2 / 3-5."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("SELECT stock_quantity, count(*) FROM wc.products WHERE stock_quantity IS NOT NULL AND stock_quantity <= 5 GROUP BY stock_quantity ORDER BY stock_quantity")
    rows = cur.fetchall()
    cur.close()
    conn.close()
    urgent = sum(r[1] for r in rows if r[0] == 0)
    high = sum(r[1] for r in rows if r[0] in [1, 2])
    normal = sum(r[1] for r in rows if 3 <= r[0] <= 5)
    return {"total": urgent + high + normal, "urgent": urgent, "high": high, "normal": normal}


def bucket_for_stock(stock_val):
    try:
        s = int(float(str(stock_val)))
    except (TypeError, ValueError):
        return None
    if s == 0:
        return "urgent"
    if s in (1, 2):
        return "high"
    if 3 <= s <= 5:
        return "normal"
    return None  # stock>5 or negative => should not be present


def check_gcal_events():
    """At least 2 stock/reorder/replenishment events, EN or RU titles."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute(
        "SELECT count(*) FROM gcal.events WHERE "
        "LOWER(summary) LIKE '%stock%' OR LOWER(summary) LIKE '%reorder%' "
        "OR LOWER(summary) LIKE '%replenishment%' "
        "OR LOWER(summary) LIKE '%запас%' OR LOWER(summary) LIKE '%заказ%' "
        "OR LOWER(summary) LIKE '%пополнен%'"
    )
    cnt = cur.fetchone()[0]
    cur.close()
    conn.close()
    return cnt >= 2


def check_email_to_procurement(urgent_names):
    """Email to procurement@company.com, low-stock/reorder subject, body lists a zero-stock product."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute(
        "SELECT subject, to_addr::text, body_text FROM email.messages "
        "WHERE LOWER(to_addr::text) LIKE '%procurement%'"
    )
    rows = cur.fetchall()
    cur.close()
    conn.close()
    detail = []
    for subject, to_addr, body in rows:
        subj_l = (subject or "").lower()
        if not ("low stock" in subj_l or "reorder" in subj_l):
            detail.append(f"subject={subject!r}")
            continue
        body_l = (body or "").lower()
        if urgent_names:
            listed = any(
                (str(p).lower()[:25] in body_l) or
                any(tok for tok in str(p).lower().split() if len(tok) > 4 and tok in body_l)
                for p in urgent_names
            )
        else:
            listed = len(body_l.strip()) > 20
        if listed:
            return True, ""
        detail.append("body lists no zero-stock product")
    return False, "; ".join(detail[:4]) or "no matching email"


def check_word_file(agent_workspace):
    word_path = os.path.join(agent_workspace, "Reorder_Report.docx")
    if not os.path.exists(word_path):
        return False, "Reorder_Report.docx not found"
    try:
        from docx import Document
    except Exception as e:
        # python-docx unavailable: fall back to file existence only (cannot inspect content).
        return True, f"(python-docx unavailable: {e}; existence-only)"
    try:
        doc = Document(word_path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        full_text += " ".join(cell.text for t in doc.tables for row in t.rows for cell in row.cells)
        if "low stock" in full_text.lower() or "reorder" in full_text.lower():
            return True, ""
        return False, "Word doc does not mention 'Low Stock' or 'Reorder'"
    except Exception as e:
        # Malformed document is a real failure, not a silent pass.
        return False, f"Could not parse Reorder_Report.docx: {e}"


def main():
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILED
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    agent_file = os.path.join(args.agent_workspace, "Low_Stock_Reorder.xlsx")
    check("Low_Stock_Reorder.xlsx exists", os.path.exists(agent_file))
    if not os.path.exists(agent_file):
        print(f"FAIL: Agent output not found: {agent_file}")
        sys.exit(1)

    # Live source-of-truth counts. If the seed/DB is broken, that is a hard failure,
    # not a silently-masked hardcoded fallback.
    try:
        counts = get_low_stock_counts()
    except Exception as e:
        print(f"FAIL: Could not query wc.products for low-stock counts: {e}")
        sys.exit(1)

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)

    # ---------- Reorder List ----------
    print("  Checking Reorder List sheet...")
    urgent_names = []
    a_rows = load_sheet_rows(agent_wb, "Reorder List")
    if a_rows is None:
        check("Reorder List sheet exists", False, "sheet not found")
    else:
        check("Reorder List sheet exists", True)
        data_rows = [r for r in a_rows[1:] if r and any(c is not None for c in r)]

        # CRITICAL: exact row count (== catches both omission and over-inclusion).
        check("CRITICAL: Reorder List row count == live low-stock count",
              len(data_rows) == counts["total"],
              f"{len(data_rows)} rows vs expected {counts['total']}", critical=True)

        # CRITICAL: every row's Reorder_Priority matches its Current_Stock bucket (all three buckets).
        bad = []
        for row in data_rows:
            if not row or len(row) < 4:
                bad.append("malformed row")
                continue
            stock_val = row[1]
            priority_val = str(row[3]).strip().lower() if row[3] is not None else ""
            exp = bucket_for_stock(stock_val)
            # collect zero-stock product names for the email body check
            if exp == "urgent" and row[0] is not None:
                urgent_names.append(str(row[0]))
            if exp is None:
                bad.append(f"stock={stock_val!r} outside 0-5 buckets")
            elif priority_val != exp:
                bad.append(f"stock={stock_val} -> '{priority_val}' (expected {exp})")
        check("CRITICAL: Reorder_Priority matches Current_Stock bucket (0->Urgent,1-2->High,3-5->Normal)",
              len(bad) == 0, "; ".join(bad[:5]), critical=True)

    # ---------- Summary ----------
    print("  Checking Summary sheet...")
    a_rows = load_sheet_rows(agent_wb, "Summary")
    if a_rows is None:
        check("Summary sheet exists", False, "sheet not found")
        check("CRITICAL: Summary metrics equal live counts", False, "no Summary sheet", critical=True)
    else:
        check("Summary sheet exists", True)
        a_data = {str(r[0]).strip().lower(): r[1] for r in a_rows[1:] if r and r[0] is not None}
        errors = []
        for metric, exp_val in [
            ("total_low_stock_items", counts["total"]),
            ("urgent_count", counts["urgent"]),
            ("high_priority_count", counts["high"]),
            ("normal_priority_count", counts["normal"]),
        ]:
            val = a_data.get(metric)
            if val is None:
                errors.append(f"missing {metric}")
            elif not num_close(val, exp_val, 0):
                errors.append(f"{metric}: {val} vs {exp_val}")
        check("CRITICAL: Summary metrics equal live counts (tol=0)",
              len(errors) == 0, "; ".join(errors[:5]), critical=True)

    # ---------- GCal ----------
    print("  Checking GCal events...")
    try:
        check("At least 2 stock/reorder calendar events (EN or RU)", check_gcal_events())
    except Exception as e:
        check("At least 2 stock/reorder calendar events (EN or RU)", False, str(e))

    # ---------- Email ----------
    print("  Checking email to procurement...")
    try:
        ok, detail = check_email_to_procurement(urgent_names)
    except Exception as e:
        ok, detail = False, str(e)
    check("CRITICAL: email -> procurement@company.com (low stock/reorder subject) lists zero-stock products",
          ok, detail, critical=True)

    # ---------- Word ----------
    print("  Checking Word document...")
    ok, detail = check_word_file(args.agent_workspace)
    check("Reorder_Report.docx mentions Low Stock / Reorder", ok, detail)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0

    # Critical gate: any critical failure => hard FAIL regardless of accuracy.
    if CRITICAL_FAILED:
        print(f"\nCRITICAL checks failed: {CRITICAL_FAILED}")
        print(f"Passed {PASS_COUNT}/{total} checks ({accuracy:.1f}%) but CRITICAL gate failed")
        print("=== RESULT: FAIL ===")
        sys.exit(1)

    if accuracy >= 70:
        print(f"\nPassed {PASS_COUNT}/{total} checks ({accuracy:.1f}%)")
        print("=== RESULT: PASS ===")
        sys.exit(0)
    else:
        print(f"\nPassed {PASS_COUNT}/{total} checks ({accuracy:.1f}%) — below 70% threshold")
        print("=== RESULT: FAIL ===")
        sys.exit(1)


if __name__ == "__main__":
    main()
