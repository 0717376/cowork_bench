"""Evaluation for insales-tax-compliance-review."""
import argparse
import os
import sys

import openpyxl
import psycopg2

DB = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0

# Official rates from mock website (state base rates only)
OFFICIAL_RATES = {
    "CA": 7.25,
    "FL": 6.50,
    "NY": 8.875,
    "TX": 6.25,
    "WA": 6.50,
}

COMPLIANCE_THRESHOLD = 0.25


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=0.01):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def safe_str(v):
    return str(v).strip() if v is not None else ""


def get_wc_state_rates():
    """Get state-level tax rates from the InSales store database (wc.* schema).

    Excludes district-level rows (city != '') so each state keeps only its
    base state-level rate. If a state still has multiple base rows, pick the
    lowest rate deterministically (DB-order independent). The russified store
    tax names (НДС (Москва)/(Уфа)/...) join to US state codes by state code
    only, per the fork convention.
    """
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT state, name, rate::numeric, city
        FROM wc.tax_rates
        WHERE country = 'Россия' AND state != '' AND class = 'standard'
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()
    # Keep only the state-level (non-district) row per state.
    state_rates = {}
    for state, name, rate, city in rows:
        # Skip district-level rows (those carry a city, e.g. Краснодар).
        if city is not None and str(city).strip() != "":
            continue
        rate = float(rate)
        if state not in state_rates:
            state_rates[state] = (name, rate)
        else:
            # Deterministic tie-break: pick the lower base rate.
            if rate < state_rates[state][1]:
                state_rates[state] = (name, rate)
    return state_rates


def compute_expected():
    """Compute expected comparison data."""
    wc_rates = get_wc_state_rates()
    results = []
    for state in sorted(wc_rates.keys()):
        if state in OFFICIAL_RATES:
            name, wc_rate = wc_rates[state]
            official = OFFICIAL_RATES[state]
            diff = round(abs(wc_rate - official), 4)
            compliant = "Yes" if diff <= COMPLIANCE_THRESHOLD else "No"
            results.append((state, name, wc_rate, official, diff, compliant))
    return results


def check_excel(agent_workspace):
    """Check Tax_Compliance.xlsx."""
    print("\n=== Checking Tax_Compliance.xlsx ===")

    excel_path = os.path.join(agent_workspace, "Tax_Compliance.xlsx")
    if not os.path.isfile(excel_path):
        record("Excel file exists", False, f"Not found: {excel_path}")
        return False
    record("Excel file exists", True)

    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
    except Exception as e:
        record("Excel readable", False, str(e))
        return False

    expected = compute_expected()
    all_ok = True

    # Check Rate Comparison sheet
    rc_sheet = None
    for name in wb.sheetnames:
        if "rate" in name.lower() and "compar" in name.lower():
            rc_sheet = wb[name]
            break
    if rc_sheet is None:
        record("Sheet 'Rate Comparison' exists", False, f"Sheets: {wb.sheetnames}")
        all_ok = False
    else:
        record("Sheet 'Rate Comparison' exists", True)

        # Check headers
        headers = [safe_str(rc_sheet.cell(1, c).value).lower() for c in range(1, 8)]
        record("Has State column", any("state" in h for h in headers))
        record("Has WC_Rate column", any("wc" in h and "rate" in h for h in headers))
        record("Has Official_Rate column", any("official" in h for h in headers))
        record("Has Difference column", any("diff" in h for h in headers))
        record("Has Compliant column", any("compli" in h for h in headers))

        rows = list(rc_sheet.iter_rows(min_row=2, values_only=True))
        record("Rate Comparison has correct row count",
               len(rows) >= len(expected),
               f"Expected {len(expected)}, got {len(rows)}")

        for exp_state, exp_name, exp_wc, exp_off, exp_diff, exp_compl in expected:
            found = False
            for r in rows:
                if r and r[0] and safe_str(r[0]).upper() == exp_state:
                    found = True
                    ok_wc = num_close(r[2], exp_wc, 0.01)
                    record(f"{exp_state} WC_Rate={exp_wc}", ok_wc,
                           f"Got {r[2]}")
                    if not ok_wc:
                        all_ok = False

                    ok_off = num_close(r[3], exp_off, 0.01)
                    record(f"{exp_state} Official_Rate={exp_off}", ok_off,
                           f"Got {r[3]}")
                    if not ok_off:
                        all_ok = False

                    ok_diff = num_close(r[4], exp_diff, 0.01)
                    record(f"{exp_state} Difference={exp_diff}", ok_diff,
                           f"Got {r[4]}")
                    if not ok_diff:
                        all_ok = False

                    ok_compl = safe_str(r[5]).lower() == exp_compl.lower()
                    record(f"{exp_state} Compliant={exp_compl}", ok_compl,
                           f"Got {r[5]}")
                    if not ok_compl:
                        all_ok = False
                    break
            if not found:
                record(f"State {exp_state} found in Rate Comparison", False)
                all_ok = False

    # Check Discrepancies sheet
    disc_sheet = None
    for name in wb.sheetnames:
        if "discrep" in name.lower():
            disc_sheet = wb[name]
            break
    if disc_sheet is None:
        record("Sheet 'Discrepancies' exists", False, f"Sheets: {wb.sheetnames}")
        all_ok = False
    else:
        record("Sheet 'Discrepancies' exists", True)
        disc_rows = list(disc_sheet.iter_rows(min_row=2, values_only=True))
        non_compliant = [e for e in expected if e[5] == "No"]
        record("Discrepancies has correct row count",
               len(disc_rows) >= len(non_compliant),
               f"Expected {len(non_compliant)}, got {len(disc_rows)}")

        for exp_state, _, _, _, _, _ in non_compliant:
            found = any(r and r[0] and safe_str(r[0]).upper() == exp_state for r in disc_rows)
            record(f"Discrepancy for {exp_state} listed", found)
            if not found:
                all_ok = False

    # Check Summary sheet
    sum_sheet = None
    for name in wb.sheetnames:
        if "summ" in name.lower():
            sum_sheet = wb[name]
            break
    if sum_sheet is None:
        record("Sheet 'Summary' exists", False, f"Sheets: {wb.sheetnames}")
        all_ok = False
    else:
        record("Sheet 'Summary' exists", True)
        summary = {}
        for row in sum_sheet.iter_rows(min_row=2, values_only=True):
            if row and row[0]:
                summary[safe_str(row[0]).lower().replace(" ", "_")] = row[1]

        total_reviewed = len(expected)
        compliant_count = sum(1 for e in expected if e[5] == "Yes")
        non_compliant_count = sum(1 for e in expected if e[5] == "No")
        max_diff = max(e[4] for e in expected)

        for key, val in summary.items():
            if "total" in key and "rate" in key or "reviewed" in key:
                ok = num_close(val, total_reviewed, 1)
                record(f"Summary Total_Rates_Reviewed={total_reviewed}", ok, f"Got {val}")
                if not ok:
                    all_ok = False
            elif "compliant_count" in key or ("compliant" in key and "non" not in key and "count" in key):
                ok = num_close(val, compliant_count, 1)
                record(f"Summary Compliant_Count={compliant_count}", ok, f"Got {val}")
                if not ok:
                    all_ok = False
            elif "non_compliant" in key or "noncompliant" in key:
                ok = num_close(val, non_compliant_count, 1)
                record(f"Summary Non_Compliant_Count={non_compliant_count}", ok, f"Got {val}")
                if not ok:
                    all_ok = False
            elif "max" in key and "diff" in key:
                ok = num_close(val, max_diff, 0.01)
                record(f"Summary Max_Difference={max_diff}", ok, f"Got {val}")
                if not ok:
                    all_ok = False

    return all_ok


def check_word(agent_workspace):
    """Check Tax_Compliance_Report.docx."""
    print("\n=== Checking Tax_Compliance_Report.docx ===")
    from docx import Document

    docx_path = os.path.join(agent_workspace, "Tax_Compliance_Report.docx")
    if not os.path.isfile(docx_path):
        record("Word file exists", False, f"Not found: {docx_path}")
        return False
    record("Word file exists", True)

    try:
        doc = Document(docx_path)
    except Exception as e:
        record("Word readable", False, str(e))
        return False

    full_text = " ".join(p.text.lower() for p in doc.paragraphs)
    all_ok = True

    def w(name, cond):
        nonlocal all_ok
        record(name, cond)
        if not cond:
            all_ok = False

    w("Mentions 'compliance' / 'соответств'",
      "complian" in full_text or "соответств" in full_text)
    w("Mentions 'florida' / 'флорида' / 'FL'",
      "florida" in full_text or "флорида" in full_text or " fl " in full_text)
    w("Mentions 'new york' / 'нью-йорк' / 'NY'",
      "new york" in full_text or "нью-йорк" in full_text
      or " ny " in full_text)
    w("Mentions 'discrepan' / 'расхожден' / 'non-compliant' / 'несоответств'",
      "discrepan" in full_text or "расхожден" in full_text
      or "non-compliant" in full_text or "non compliant" in full_text
      or "несоответств" in full_text or "не соответств" in full_text)
    w("Mentions 'recommend' / 'рекоменд'",
      "recommend" in full_text or "рекоменд" in full_text)
    w("Has at least 3 paragraphs", len(doc.paragraphs) >= 3)

    return all_ok


def _load_excel(agent_workspace):
    excel_path = os.path.join(agent_workspace, "Tax_Compliance.xlsx")
    if not os.path.isfile(excel_path):
        return None
    try:
        return openpyxl.load_workbook(excel_path, data_only=True)
    except Exception:
        return None


def _find_sheet(wb, *needles):
    for name in wb.sheetnames:
        low = name.lower()
        if all(n in low for n in needles):
            return wb[name]
    return None


def critical_checks(agent_workspace):
    """SEMANTIC critical checks. Any failure => hard FAIL (sys.exit(1)).

    These reflect the substance of the deliverable: the two non-compliant
    states (FL, NY) with their exact rates/diff/flag, the three compliant
    states, the Summary totals, the Discrepancies sheet content, and that the
    Word report names both discrepancies and a recommendations section.
    """
    print("\n=== CRITICAL CHECKS ===")
    results = []

    def crit(name, passed, detail=""):
        results.append((name, bool(passed)))
        if passed:
            print(f"  [CRIT PASS] {name}")
        else:
            msg = f": {detail[:200]}" if detail else ""
            print(f"  [CRIT FAIL] {name}{msg}")

    expected = compute_expected()  # from live wc.tax_rates (state-code join)
    exp_by_state = {e[0]: e for e in expected}

    wb = _load_excel(agent_workspace)
    if wb is None:
        crit("Excel workbook loads", False, "Tax_Compliance.xlsx missing/unreadable")
        return results

    # --- Rate Comparison rows by state ---
    rc = _find_sheet(wb, "rate", "compar")
    rc_by_state = {}
    if rc is not None:
        for r in rc.iter_rows(min_row=2, values_only=True):
            if r and r[0] and safe_str(r[0]):
                rc_by_state[safe_str(r[0]).upper()] = r

    def row_ok(state):
        e = exp_by_state.get(state)
        r = rc_by_state.get(state)
        if not e or not r:
            return False, f"missing row for {state}"
        _, _, exp_wc, exp_off, exp_diff, exp_compl = e
        ok = (num_close(r[2], exp_wc, 0.01) and num_close(r[3], exp_off, 0.01)
              and num_close(r[4], exp_diff, 0.01)
              and safe_str(r[5]).lower() == exp_compl.lower())
        return ok, f"got {r[2:6]}, expected wc={exp_wc} off={exp_off} diff={exp_diff} compl={exp_compl}"

    # FL non-compliant
    ok, d = row_ok("FL")
    crit("FL non-compliant (WC=6.0, Official=6.5, Diff=0.5, Compliant=No)", ok, d)
    # NY non-compliant
    ok, d = row_ok("NY")
    crit("NY non-compliant (WC=8.0, Official=8.875, Diff=0.875, Compliant=No)", ok, d)
    # Three compliant states
    comp_ok = True
    comp_det = []
    for st in ("CA", "TX", "WA"):
        ok, d = row_ok(st)
        if not ok:
            comp_ok = False
            comp_det.append(d)
    crit("CA/TX/WA compliant (Difference≈0, Compliant=Yes)", comp_ok,
         "; ".join(comp_det))

    # --- Summary totals ---
    sm = _find_sheet(wb, "summ")
    summary = {}
    if sm is not None:
        for row in sm.iter_rows(min_row=2, values_only=True):
            if row and row[0]:
                summary[safe_str(row[0]).lower().replace(" ", "_")] = row[1]

    def sval(*keys):
        """First summary value whose key contains all given substrings."""
        for k, v in summary.items():
            if all(p in k for p in keys):
                return v
        return None

    total = len(expected)
    comp_count = sum(1 for e in expected if e[5] == "Yes")
    noncomp_count = sum(1 for e in expected if e[5] == "No")
    max_diff = max(e[4] for e in expected)
    # "non" rows also match "compliant"; resolve Non/Compliant by excluding "non".
    comp_val = None
    for k, v in summary.items():
        if "compliant" in k and "count" in k and "non" not in k:
            comp_val = v
            break
    noncomp_val = None
    for k, v in summary.items():
        if "non" in k and "compliant" in k:
            noncomp_val = v
            break
    sum_ok = (num_close(sval("total"), total, 1)
              and num_close(comp_val, comp_count, 1)
              and num_close(noncomp_val, noncomp_count, 1)
              and num_close(sval("max"), max_diff, 0.01))
    crit(f"Summary totals (Total={total}, Compliant={comp_count}, "
         f"Non-Compliant={noncomp_count}, Max_Diff={max_diff})", sum_ok,
         f"summary={summary}")

    # --- Discrepancies sheet contains exactly FL and NY, no compliant rows ---
    disc = _find_sheet(wb, "discrep")
    disc_states = set()
    if disc is not None:
        for r in disc.iter_rows(min_row=2, values_only=True):
            if r and r[0] and safe_str(r[0]):
                disc_states.add(safe_str(r[0]).upper())
    expected_disc = {e[0] for e in expected if e[5] == "No"}
    crit(f"Discrepancies sheet = exactly {sorted(expected_disc)}",
         disc_states == expected_disc, f"got {sorted(disc_states)}")

    # --- Word report names both discrepancies + recommendations (RU or EN) ---
    word_crit = False
    word_detail = ""
    try:
        from docx import Document
        docx_path = os.path.join(agent_workspace, "Tax_Compliance_Report.docx")
        if os.path.isfile(docx_path):
            doc = Document(docx_path)
            ft = " ".join(p.text.lower() for p in doc.paragraphs)
            has_fl = "florida" in ft or "флорида" in ft or " fl " in ft
            has_ny = "new york" in ft or "нью-йорк" in ft or " ny " in ft
            has_rec = "recommend" in ft or "рекоменд" in ft
            word_crit = has_fl and has_ny and has_rec
            word_detail = f"FL={has_fl} NY={has_ny} rec={has_rec}"
        else:
            word_detail = "Tax_Compliance_Report.docx missing"
    except Exception as e:
        word_detail = str(e)
    crit("Word report names FL+NY discrepancies + recommendations (RU/EN)",
         word_crit, word_detail)

    return results


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    excel_ok = check_excel(args.agent_workspace)
    word_ok = check_word(args.agent_workspace)

    crit_results = critical_checks(args.agent_workspace)
    crit_failed = [n for n, ok in crit_results if not ok]

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (100.0 * PASS_COUNT / total) if total else 0.0

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Accuracy: {accuracy:.1f}%")
    print(f"  Excel OK: {excel_ok} | Word OK: {word_ok}")

    if crit_failed:
        print(f"  CRITICAL FAILURES ({len(crit_failed)}): {crit_failed}")
        print(f"  Overall: FAIL (critical check failed)")
        sys.exit(1)

    overall = (accuracy >= 70.0) and excel_ok and word_ok
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
