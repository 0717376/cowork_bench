"""
Evaluation for yt-transcript-word-gform task (RU, forms fork).

Gate: accuracy >= 70% AND no CRITICAL check failed => PASS.
Any CRITICAL failure => immediate FAIL (sys.exit(1)) regardless of accuracy.

The forms MCP (RU fork) writes to the same Postgres schema gform.* as the
original google_forms server, so DB-level checks are unchanged.

Checks:
1. Afrobeat_Playlist.docx exists with 3 headings (Introduction, Song List, Summary)
   and substantive Afrobeat content from the transcript.
2. Form "Afrobeat Mix Feedback" exists with exactly 3 questions of the correct
   types and option sets (Q1 rating choice, Q2 favorite text, Q3 recommend choice).
3. Email sent to music@company.com with the Afrobeat analysis subject.

CRITICAL checks (each gates the whole task):
- docx with all 3 required headings + Afrobeat content (the playlist artifact)
- form with exactly 3 correctly-typed questions and required option sets
- email to music@company.com with the analysis subject
A correct RU agent passes all three; a non-doer (no doc / no form / no email)
fails a CRITICAL check and the task FAILs regardless of accuracy.
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "cowork_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []


def record(name, passed, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    tag = "[CRIT]" if critical else ""
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS]{tag} {name}")
    else:
        FAIL_COUNT += 1
        if critical:
            CRITICAL_FAILED.append(name)
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL]{tag} {name}{msg}")


# State the doc/form/email CRITICAL gates depend on.
DOC_OK = False
FORM_OK = False
EMAIL_OK = False


def check_word(agent_workspace):
    global DOC_OK
    print("\n=== Check 1: Afrobeat_Playlist.docx ===")
    docx_path = os.path.join(agent_workspace, "Afrobeat_Playlist.docx")
    if not os.path.exists(docx_path):
        record("Afrobeat_Playlist.docx exists", False, f"Not found at {docx_path}")
        record("CRITICAL: playlist doc with 3 headings + Afrobeat content",
               False, "doc missing", critical=True)
        return
    record("Afrobeat_Playlist.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        record("CRITICAL: playlist doc with 3 headings + Afrobeat content",
               False, f"unreadable: {e}", critical=True)
        return
    record("Word doc readable", True)

    headings = [p.text.strip() for p in doc.paragraphs
                if p.style.name.startswith('Heading')]
    heading_text = " ".join(headings).lower()
    has_intro = "introduction" in heading_text or "intro" in heading_text
    has_songs = "song" in heading_text or "playlist" in heading_text or "track" in heading_text
    has_summary = "summary" in heading_text or "conclusion" in heading_text
    record("Has Introduction heading", has_intro, f"Headings: {headings}")
    record("Has Song List heading", has_songs, f"Headings: {headings}")
    record("Has Summary heading", has_summary, f"Headings: {headings}")

    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    has_afrobeat = "afrobeat" in full_text or "afrobeats" in full_text
    record("Doc mentions Afrobeat content", has_afrobeat,
           "Afrobeat not mentioned in document")
    has_paragraphs = len([p for p in doc.paragraphs if len(p.text.strip()) > 30]) >= 3
    record("Doc has at least 3 paragraphs of content", has_paragraphs,
           "Less than 3 substantial paragraphs")

    # CRITICAL: the playlist artifact must be real and complete.
    DOC_OK = has_intro and has_songs and has_summary and has_afrobeat and has_paragraphs
    record("CRITICAL: playlist doc with 3 headings + Afrobeat content",
           DOC_OK,
           f"intro={has_intro} songs={has_songs} summary={has_summary} "
           f"afrobeat={has_afrobeat} paragraphs={has_paragraphs}",
           critical=True)


def _is_choice(t):
    t = (t or "").upper()
    return t in ("RADIO", "MULTIPLE_CHOICE", "CHOICE", "CHOICEQUESTION", "CHECKBOX")


def _is_text(t):
    t = (t or "").upper()
    return t in ("TEXT", "SHORT_ANSWER", "PARAGRAPH", "TEXTQUESTION")


def _options(config):
    if not config:
        return []
    if isinstance(config, str):
        try:
            config = json.loads(config)
        except Exception:
            return []
    opts = config.get("options", []) or []
    # The forms MCP stores each option as {"value": "..."}; unwrap to the raw value.
    return [(o.get("value") if isinstance(o, dict) else o) for o in opts]


def check_gform(cur):
    global FORM_OK
    print("\n=== Check 2: Form 'Afrobeat Mix Feedback' (forms RU fork, gform schema) ===")

    cur.execute("SELECT id, title FROM gform.forms WHERE title ILIKE %s",
                ("%Afrobeat Mix Feedback%",))
    forms = cur.fetchall()
    form_exists = len(forms) >= 1
    record("Form 'Afrobeat Mix Feedback' exists", form_exists,
           f"Forms found: {[f[1] for f in forms]}")
    if not form_exists:
        record("CRITICAL: form with 3 correctly-typed questions", False,
               "form missing", critical=True)
        return

    form_id = forms[0][0]
    cur.execute("SELECT COUNT(*) FROM gform.questions WHERE form_id = %s", (form_id,))
    q_count = cur.fetchone()[0]
    three_questions = q_count == 3
    record("Form has exactly 3 questions", three_questions, f"Found {q_count} questions")

    cur.execute("""
        SELECT title, question_type, config
        FROM gform.questions WHERE form_id = %s
        ORDER BY position
    """, (form_id,))
    questions = cur.fetchall()

    q1_ok = q2_ok = q3_ok = False

    if len(questions) >= 1:
        q1_title = (questions[0][0] or "").lower()
        q1_type = questions[0][1]
        q1_is_choice = _is_choice(q1_type)
        q1_title_ok = "rate" in q1_title or "rating" in q1_title
        record("Q1 is 'How would you rate this mix?' (multiple choice)",
               q1_title_ok and q1_is_choice,
               f"Q1 title: {questions[0][0]}, type: {q1_type}")
        opts1 = [str(o).lower() for o in _options(questions[0][2])]
        has_excellent = any("excellent" in o for o in opts1)
        has_good = any(o.strip() == "good" or "good" in o for o in opts1)
        has_average = any("average" in o for o in opts1)
        has_poor = any("poor" in o for o in opts1)
        # Exactly the four required options.
        four_opts = len(opts1) == 4 and has_excellent and has_good and has_average and has_poor
        record("Q1 has exactly 4 options: Excellent/Good/Average/Poor", four_opts,
               f"Options: {_options(questions[0][2])}")
        q1_ok = q1_title_ok and q1_is_choice and four_opts

    if len(questions) >= 2:
        q2_title = (questions[1][0] or "").lower()
        q2_type = questions[1][1]
        q2_is_text = _is_text(q2_type)
        q2_title_ok = "favorite" in q2_title or "favourite" in q2_title or "song" in q2_title
        record("Q2 is 'Which song was your favorite?' (text)",
               q2_title_ok and q2_is_text,
               f"Q2: {questions[1][0]}, type: {q2_type}")
        q2_ok = q2_title_ok and q2_is_text

    if len(questions) >= 3:
        q3_title = (questions[2][0] or "").lower()
        q3_type = questions[2][1]
        q3_is_choice = _is_choice(q3_type)
        q3_title_ok = "recommend" in q3_title
        opts3 = [str(o).lower() for o in _options(questions[2][2])]
        has_yes = any(o.strip() == "yes" or "yes" in o for o in opts3)
        has_no = any(o.strip() == "no" for o in opts3)
        has_maybe = any("maybe" in o for o in opts3)
        three_opts = len(opts3) == 3 and has_yes and has_no and has_maybe
        record("Q3 is 'Would you recommend?' with exactly Yes/No/Maybe options",
               q3_title_ok and q3_is_choice and three_opts,
               f"Q3: {questions[2][0]}, type: {q3_type}, options: {_options(questions[2][2])}")
        q3_ok = q3_title_ok and q3_is_choice and three_opts

    # CRITICAL: the form must be a real survey with all three correctly-built questions.
    FORM_OK = three_questions and q1_ok and q2_ok and q3_ok
    record("CRITICAL: form with 3 correctly-typed questions",
           FORM_OK,
           f"three_questions={three_questions} q1={q1_ok} q2={q2_ok} q3={q3_ok}",
           critical=True)


def check_email(cur):
    global EMAIL_OK
    print("\n=== Check 3: Email sent ===")
    cur.execute("""
        SELECT m.to_addr, m.subject FROM email.messages m
        JOIN email.sent_log sl ON sl.message_id = m.id
        WHERE m.to_addr::text ILIKE %s
        ORDER BY sl.sent_at DESC LIMIT 5
    """, ("%music@company.com%",))
    emails = cur.fetchall()

    if not emails:
        cur.execute("""
            SELECT to_addr, subject FROM email.messages
            WHERE to_addr::text ILIKE %s
            ORDER BY date DESC LIMIT 5
        """, ("%music@company.com%",))
        emails = cur.fetchall()

    sent = len(emails) >= 1
    record("Email sent to music@company.com", sent, f"Found: {emails}")

    subject_ok = False
    if emails:
        subject = str(emails[0][1]).lower() if emails[0][1] else ""
        subject_ok = "afrobeat" in subject or "analysis" in subject
        record("Email subject mentions 'Afrobeat' or 'Analysis'",
               subject_ok, f"Subject: {emails[0][1]}")

    # CRITICAL: notification email to the right recipient with the right subject.
    EMAIL_OK = sent and subject_ok
    record("CRITICAL: email to music@company.com with analysis subject",
           EMAIL_OK, f"sent={sent} subject_ok={subject_ok}", critical=True)


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)

    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        check_gform(cur)
        check_email(cur)
        cur.close()
        conn.close()
    except Exception as e:
        record("DB checks", False, str(e), critical=True)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_failed": CRITICAL_FAILED,
        "success": not CRITICAL_FAILED and accuracy >= 70,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if CRITICAL_FAILED:
        print("CRITICAL checks failed:")
        for c in CRITICAL_FAILED:
            print(f"  - {c}")
        print("=> FAIL (critical)")
        sys.exit(1)

    if accuracy >= 70:
        print("=> PASS (accuracy >= 70% and no critical failure)")
        sys.exit(0)
    print(f"=> FAIL (accuracy {accuracy:.1f}% < 70%)")
    sys.exit(1)


if __name__ == "__main__":
    main()
