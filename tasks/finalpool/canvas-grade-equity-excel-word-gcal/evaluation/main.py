"""Evaluation для canvas-grade-equity-excel-word-gcal.

Архетип keep_foreign_ru_seed: MCP не меняются (canvas/gcal/excel/word/pdf).
Идентификаторы (имена файлов, листов, колонок, названия курсов,
статусы Equity_Status, суффикс заголовка 'Grade Equity Review') остаются
на английском — eval грепает их. Прозаический текст агента — на русском,
поэтому проверки разделов Word приняты в двуязычном виде (EN+RU).

CRITICAL_CHECKS (любой провал => FAIL всей задачи независимо от accuracy):
- точный набор из 6 сопоставленных курсов в Course Comparison;
- Биохимия и биоинформатика: Score_Difference ≈ -12.7 И Equity_Status == 'Action Required';
- Summary: Total_Courses_Compared==6, Action_Required==1, Acceptable==5, Concerning==0;
- хотя бы 1 событие GCal с названием курса Action Required и 'Grade Equity Review',
  длительность 45 мин (40-50), на неделе 2026-03-16..2026-03-20, отличное от шумовых;
- Overall_Avg_2013 ≈ 74.4 и Overall_Avg_2014 ≈ 72.2.
"""
import argparse
import os
import sys
from datetime import datetime

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "cowork_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []

# Known course data from actual DB (read-only canvas, нерусифицируемо)
COURSE_NAMES = [
    "Прикладная аналитика и алгоритмы",
    "Биохимия и биоинформатика",
    "Проектирование на основе данных",
    "Экологическая экономика и этика",
    "Основы финансов",
    "Глобальное управление и геополитика",
]

ACTION_REQUIRED_COURSES = ["Биохимия и биоинформатика"]

# Чеки, провал которых = содержательное невыполнение задачи.
CRITICAL_CHECKS = {
    "CRITICAL: Course Comparison содержит ровно 6 сопоставленных курсов",
    "CRITICAL: Биохимия и биоинформатика Score_Difference ≈ -12.7",
    "CRITICAL: Биохимия и биоинформатика Equity_Status == 'Action Required'",
    "CRITICAL: Summary Total_Courses_Compared == 6",
    "CRITICAL: Summary Courses_Action_Required == 1",
    "CRITICAL: Summary Courses_Acceptable == 5",
    "CRITICAL: Summary Courses_Concerning == 0",
    "CRITICAL: GCal — встреча Action Required курса 'Grade Equity Review', 45 мин, 16-20 марта 2026",
    "CRITICAL: Summary Overall_Avg_2013 ≈ 74.4",
    "CRITICAL: Summary Overall_Avg_2014 ≈ 72.2",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        marker = " [CRITICAL]" if name in CRITICAL_CHECKS else ""
        print(f"  [FAIL]{marker} {name}{msg}")
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILS.append(name)


def check_excel(agent_workspace, groundtruth_workspace):
    print("\n=== Checking Excel ===")
    xlsx_path = os.path.join(agent_workspace, "Grade_Equity_Analysis.xlsx")
    if not os.path.isfile(xlsx_path):
        check("Grade_Equity_Analysis.xlsx exists", False, f"Not found: {xlsx_path}")
        return
    check("Grade_Equity_Analysis.xlsx exists", True)

    try:
        from openpyxl import load_workbook

        wb = load_workbook(xlsx_path)

        # Check sheet names
        check("Has 'Course Comparison' sheet", "Course Comparison" in wb.sheetnames,
              f"Sheets: {wb.sheetnames}")
        check("Has 'Grade Distribution' sheet", "Grade Distribution" in wb.sheetnames,
              f"Sheets: {wb.sheetnames}")
        check("Has 'Summary' sheet", "Summary" in wb.sheetnames,
              f"Sheets: {wb.sheetnames}")

        # Load groundtruth for comparison
        gt_path = os.path.join(groundtruth_workspace, "Grade_Equity_Analysis.xlsx")
        gt_wb = load_workbook(gt_path)

        # Check Course Comparison sheet
        if "Course Comparison" in wb.sheetnames:
            ws = wb["Course Comparison"]
            gt_ws = gt_wb["Course Comparison"]

            rows = list(ws.iter_rows(min_row=2, values_only=True))
            gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))

            check(f"Course Comparison has {len(gt_rows)} data rows",
                  len(rows) == len(gt_rows),
                  f"Found {len(rows)} rows, expected {len(gt_rows)}")

            agent_names = [str(r[0]).strip() if r[0] else "" for r in rows]

            # CRITICAL: точный набор из 6 курсов (равенство множеств, без лишних/недостающих)
            check("CRITICAL: Course Comparison содержит ровно 6 сопоставленных курсов",
                  set(n for n in agent_names if n) == set(COURSE_NAMES),
                  f"Found: {sorted(agent_names)}")

            for name in COURSE_NAMES:
                check(f"Course '{name}' in Course Comparison",
                      name in agent_names,
                      f"Found: {agent_names}")

            # Check numeric values with tolerance
            gt_dict = {str(r[0]).strip(): r for r in gt_rows}
            for row in rows:
                name = str(row[0]).strip() if row[0] else ""
                if name in gt_dict:
                    gt_row = gt_dict[name]
                    for col_idx, col_name in [(1, "Fall_2013_Mean"), (2, "Fall_2014_Mean"),
                                               (3, "Score_Difference")]:
                        if row[col_idx] is not None and gt_row[col_idx] is not None:
                            diff = abs(float(row[col_idx]) - float(gt_row[col_idx]))
                            cname = f"{name} {col_name} within tolerance"
                            # повысить до CRITICAL для ключевого курса Score_Difference
                            if name == "Биохимия и биоинформатика" and col_idx == 3:
                                cname = "CRITICAL: Биохимия и биоинформатика Score_Difference ≈ -12.7"
                            check(cname, diff <= 1.0,
                                  f"Agent={row[col_idx]}, GT={gt_row[col_idx]}, diff={diff:.2f}")

                    # Check equity status
                    if row[7] is not None:
                        cname = f"{name} Equity_Status matches"
                        if name == "Биохимия и биоинформатика":
                            cname = "CRITICAL: Биохимия и биоинформатика Equity_Status == 'Action Required'"
                        check(cname,
                              str(row[7]).strip() == str(gt_row[7]).strip(),
                              f"Agent='{row[7]}', GT='{gt_row[7]}'")

        # Check Grade Distribution sheet
        if "Grade Distribution" in wb.sheetnames:
            ws = wb["Grade Distribution"]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Grade Distribution has 12 data rows (6 courses x 2 years)",
                  len(rows) == 12,
                  f"Found {len(rows)} rows")

        # Check Summary sheet
        if "Summary" in wb.sheetnames:
            ws = wb["Summary"]
            gt_ws = gt_wb["Summary"]
            rows = {str(r[0]).strip(): r[1] for r in ws.iter_rows(min_row=2, values_only=True) if r[0]}
            gt_rows = {str(r[0]).strip(): r[1] for r in gt_ws.iter_rows(min_row=2, values_only=True) if r[0]}

            check("Summary has Total_Courses_Compared",
                  "Total_Courses_Compared" in rows,
                  f"Keys: {list(rows.keys())}")

            def summary_int(key, expected, critical_name):
                if key in rows:
                    try:
                        ok = int(rows[key]) == expected
                    except (TypeError, ValueError):
                        ok = False
                    check(critical_name, ok, f"Got {rows.get(key)}")
                else:
                    check(critical_name, False, f"Missing {key}")

            summary_int("Total_Courses_Compared", 6,
                        "CRITICAL: Summary Total_Courses_Compared == 6")
            summary_int("Courses_Action_Required", 1,
                        "CRITICAL: Summary Courses_Action_Required == 1")
            summary_int("Courses_Acceptable", 5,
                        "CRITICAL: Summary Courses_Acceptable == 5")
            summary_int("Courses_Concerning", 0,
                        "CRITICAL: Summary Courses_Concerning == 0")

            for key, cname in [
                ("Overall_Avg_2013", "CRITICAL: Summary Overall_Avg_2013 ≈ 74.4"),
                ("Overall_Avg_2014", "CRITICAL: Summary Overall_Avg_2014 ≈ 72.2"),
            ]:
                if key in rows and key in gt_rows:
                    try:
                        diff = abs(float(rows[key]) - float(gt_rows[key]))
                        ok = diff <= 1.0
                    except (TypeError, ValueError):
                        ok = False
                        diff = None
                    check(cname, ok, f"Agent={rows[key]}, GT={gt_rows[key]}")
                else:
                    check(cname, False, f"Missing {key}")

    except ImportError:
        check("openpyxl available", False, "openpyxl not installed")
    except Exception as e:
        check("Excel parsing", False, str(e))


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    docx_path = os.path.join(agent_workspace, "Equity_Report.docx")
    if not os.path.isfile(docx_path):
        check("Equity_Report.docx exists", False, f"Not found: {docx_path}")
        return
    check("Equity_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += " " + cell.text.lower()

        check("Word doc has meaningful content (>= 200 chars)",
              len(all_text.strip()) >= 200,
              f"Content length: {len(all_text)}")

        # Required sections — двуязычное соответствие (EN заголовки + RU синонимы).
        section_alts = [
            ("executive summary", ["executive summary", "краткое резюме", "аннотация", "резюме"]),
            ("methodology", ["methodology", "методология", "методика"]),
            ("course-by-course", ["course-by-course", "по курсам", "покурсов", "анализ курсов"]),
            ("recommendation", ["recommendation", "рекомендац"]),
            ("appendix", ["appendix", "приложение"]),
        ]
        for label, alts in section_alts:
            check(f"Word doc contains '{label}' section",
                  any(a in all_text for a in alts),
                  f"Not found (EN/RU): {alts}")

        # Check course names mentioned
        found_courses = sum(1 for name in COURSE_NAMES if name.lower() in all_text)
        check("Word doc mentions at least 4 course names",
              found_courses >= 4,
              f"Found {found_courses} of {len(COURSE_NAMES)} course names")

        # Check for action required course (RU prose; либо EN-идентификатор курса)
        check("Word doc mentions 'Биохимия и биоинформатика'",
              ("биохими" in all_text or "биоинформат" in all_text
               or "biochemistry" in all_text),
              "Action required course not mentioned")

        # Check for equity-related content (EN identifiers + RU prose)
        check("Word doc discusses equity/grading",
              any(k in all_text for k in
                  ["equity", "grade", "score", "оцен", "балл", "справедлив"]),
              f"Sample: {all_text[:300]}")

    except ImportError:
        check("Word doc has content", os.path.getsize(docx_path) > 1000,
              f"Size: {os.path.getsize(docx_path)}")
    except Exception as e:
        check("Word doc readable", False, str(e))


def _in_target_week(dt_val):
    """True если дата события приходится на 2026-03-16..2026-03-20."""
    if not dt_val:
        return False
    try:
        dt = dt_val if isinstance(dt_val, datetime) else datetime.fromisoformat(str(dt_val))
        return datetime(2026, 3, 16) <= dt < datetime(2026, 3, 21)
    except (ValueError, TypeError):
        s = str(dt_val)
        return any(f"2026-03-{d}" in s for d in ("16", "17", "18", "19", "20"))


def check_calendar():
    print("\n=== Checking Google Calendar ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        cur.execute("""
            SELECT summary, description, start_datetime, end_datetime
            FROM gcal.events
        """)
        all_events = cur.fetchall()

        # События, связанные с проверкой справедливости оценок
        equity_events = [
            e for e in all_events
            if e[0] and (
                "equity" in e[0].lower()
                or ("grade" in e[0].lower() and "review" in e[0].lower())
                or "biochemistry" in e[0].lower()
                or "bioinformatics" in e[0].lower()
            )
        ]
        check("At least 1 grade equity review meeting scheduled",
              len(equity_events) >= 1,
              f"Found {len(equity_events)} matching events")

        # CRITICAL: встреча по Action Required курсу, в целевую неделю,
        # с длительностью 45 минут, отличная от шумовых событий.
        valid = []
        for e in equity_events:
            summary = (e[0] or "")
            slow = summary.lower()
            mentions_course = ("biochemistry" in slow or "bioinformatics" in slow)
            mentions_review = ("grade equity review" in slow or
                               ("equity" in slow and "review" in slow))
            in_week = _in_target_week(e[2])
            dur_ok = False
            if e[2] and e[3]:
                try:
                    start = e[2] if isinstance(e[2], datetime) else datetime.fromisoformat(str(e[2]))
                    end = e[3] if isinstance(e[3], datetime) else datetime.fromisoformat(str(e[3]))
                    dur = (end - start).total_seconds() / 60
                    dur_ok = 40 <= dur <= 50
                except (ValueError, TypeError):
                    dur_ok = False
            if mentions_course and mentions_review and in_week and dur_ok:
                valid.append(e)

        check("CRITICAL: GCal — встреча Action Required курса 'Grade Equity Review', 45 мин, 16-20 марта 2026",
              len(valid) >= 1,
              f"Valid={len(valid)}; equity_events dates={[str(e[2]) for e in equity_events]}")

        # Доп. (некритичные) структурные проверки
        in_week_any = [e for e in equity_events if _in_target_week(e[2])]
        check("Meeting(s) scheduled in week of March 16-20, 2026",
              len(in_week_any) >= 1,
              f"{len(in_week_any)} of {len(equity_events)} in target week. "
              f"Dates: {[str(e[2]) for e in equity_events]}")

        event_texts = " ".join(
            (str(e[0]) + " " + str(e[1] or "")).lower() for e in equity_events
        )
        check("Meeting mentions 'Biochemistry' or relevant course",
              "biochemistry" in event_texts or "bioinformatics" in event_texts,
              f"Event text: {event_texts[:300]}")

        cur.close()
        conn.close()
    except Exception as e:
        check("Calendar check", False, str(e))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    gt = args.groundtruth_workspace or os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "groundtruth_workspace"
    )

    check_excel(args.agent_workspace, gt)
    check_word(args.agent_workspace)
    check_calendar()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    print(f"\n=== Results: {PASS_COUNT}/{total} passed ({accuracy:.1f}%) ===")
    if CRITICAL_FAILS:
        print(f"Critical fails ({len(CRITICAL_FAILS)}): {CRITICAL_FAILS}")
        print("FAIL: провалены критичные проверки.")
        sys.exit(1)
    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    print(f"FAIL: accuracy {accuracy:.1f}% < 70%")
    sys.exit(1)


if __name__ == "__main__":
    main()
