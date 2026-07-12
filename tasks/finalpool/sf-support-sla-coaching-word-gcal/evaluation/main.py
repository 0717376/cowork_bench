"""Evaluation for sf-support-sla-coaching-word-gcal (ClickHouse / russified).

The SUPPORT_CENTER ticket data lives in the ClickHouse-identity sf_data schema and
is russified centrally by db/zzz_clickhouse_after_init.sql. PRIORITY values
(High/Medium/Low) are intentionally NOT mapped and stay English, as do the column
names, the document title, section names, the event titles and the email address.
The agent legitimately writes Russian analysis prose and a Russian email; subject
matching is RU+EN tolerant.

CRITICAL gate: the semantic anchors (real per-priority ticket counts queried from
the DWH; the three required sections; all four quarterly GCal events on the exact
dates) must pass or the task FAILs regardless of accuracy. Structural niceties stay
non-critical. Threshold: accuracy >= 70 AND no critical fail => PASS.
"""
import argparse
import json
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "cowork_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRITICAL]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL]{' [CRITICAL]' if critical else ''} {name}: {d}")
        if critical:
            CRITICAL_FAILS.append(name)


def check_word_doc(agent_ws):
    print("\n=== Checking Word Document ===")
    docx_path = os.path.join(agent_ws, "SLA_Coaching_Plan.docx")
    check("SLA_Coaching_Plan.docx exists", os.path.exists(docx_path),
          f"Expected {docx_path}", critical=True)
    if not os.path.exists(docx_path):
        return

    from docx import Document
    doc = Document(docx_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    low = full_text.lower()

    # Title (lenient, non-critical).
    check("Document has title 'SLA Performance Coaching Plan'",
          "sla performance coaching" in low)

    # CRITICAL: all three priority labels present (stay English).
    missing_pri = [p for p in ["High", "Medium", "Low"] if p not in full_text]
    check("Document contains priority labels High / Medium / Low",
          not missing_pri, f"missing={missing_pri}", critical=True)

    # CRITICAL: real per-priority ticket counts queried from the DWH.
    missing_counts = [c for c in ["6466", "15774", "9348"] if c not in full_text]
    check("Document contains ticket counts 6466 / 15774 / 9348",
          not missing_counts, f"missing={missing_counts}", critical=True)

    # CRITICAL: the three required section headings exist (not just any 'coaching'
    # substring from the title). RU prose allowed, but the section names stay English.
    has_summary = "sla performance summary" in low
    has_analysis = "priority analysis" in low
    has_reco = "coaching recommendations" in low
    check("Document has the three sections (SLA Performance Summary / Priority "
          "Analysis / Coaching Recommendations)",
          has_summary and has_analysis and has_reco,
          f"summary={has_summary} analysis={has_analysis} reco={has_reco}",
          critical=True)

    # The required column headers appear in the table (non-critical structural).
    cols = ["Priority", "Total_Tickets", "Avg_Response_Hours", "SLA_Hours", "Compliance_Rate"]
    missing_cols = [c for c in cols if c not in full_text]
    check("Summary table has all 5 column headers", not missing_cols,
          f"missing={missing_cols}")


def check_gcal():
    print("\n=== Checking Google Calendar events ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        quarters = [
            ("Q1", "2026-03-15"),
            ("Q2", "2026-06-15"),
            ("Q3", "2026-09-15"),
            ("Q4", "2026-12-15"),
        ]
        all_dates_ok = True
        all_time_ok = True
        for q, date in quarters:
            cur.execute("""
                SELECT start_datetime::time, end_datetime::time
                FROM gcal.events
                WHERE LOWER(summary) LIKE %s
                  AND start_datetime::date = %s
            """, (f"%{q.lower()}%", date))
            rows = cur.fetchall()
            found = len(rows) > 0
            check(f"{q} SLA Coaching Review event on {date}", found,
                  f"no event for {q} on {date}")
            if not found:
                all_dates_ok = False
                all_time_ok = False
                continue
            # 10:00-11:00 window (verify the time component, not just the date).
            time_ok = any(
                str(s).startswith("10:00") and str(e).startswith("11:00")
                for s, e in rows
            )
            if not time_ok:
                all_time_ok = False
            check(f"{q} event scheduled 10:00-11:00", time_ok,
                  f"times={[ (str(s), str(e)) for s, e in rows ]}")

        # CRITICAL: all four quarterly events exist on the exact dates.
        check("All four quarterly coaching events exist on correct 2026 dates",
              all_dates_ok, critical=True)
        # Time window is a hardening check (non-critical).
        check("All four events scheduled in the 10:00-11:00 window", all_time_ok)
    finally:
        cur.close()
        conn.close()


def check_email():
    print("\n=== Checking Email ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT subject, body_text FROM email.messages
            WHERE to_addr::text ILIKE '%support_manager@company.com%'
              AND folder_id != 0
        """)
        msgs = cur.fetchall()
        cur.execute("""
            SELECT subject, body_text FROM email.drafts
            WHERE to_addr::text ILIKE '%support_manager@company.com%'
        """)
        try:
            msgs += cur.fetchall()
        except Exception:
            pass
    finally:
        cur.close()
        conn.close()

    check("Email sent to support_manager@company.com", len(msgs) >= 1,
          f"found {len(msgs)} message(s)", critical=True)
    if not msgs:
        return

    # Subject references the coaching plan: RU+EN tolerant grep.
    subj_terms = ["coaching", "sla", "коучинг", "план"]
    subj_ok = any(
        any(t in str(subj or "").lower() for t in subj_terms)
        for subj, _ in msgs
    )
    check("Email subject references the coaching plan (RU/EN)", subj_ok,
          f"subjects={[s for s, _ in msgs][:5]}", critical=True)

    # Body summarizes SLA findings (RU+EN) and mentions the scheduled events.
    body_terms = ["sla", "коучинг", "приоритет", "тикет", "priority", "ticket"]
    body_ok = any(
        any(t in str(body or "").lower() for t in body_terms)
        for _, body in msgs
    )
    check("Email body summarizes SLA findings (RU/EN)", body_ok)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    agent_ws = args.agent_workspace or task_root

    print("=" * 70)
    print("SF SUPPORT SLA COACHING WORD GCAL - EVALUATION (ClickHouse / RU)")
    print("=" * 70)

    check_word_doc(agent_ws)
    try:
        check_gcal()
    except Exception as e:
        check("GCal check ran", False, str(e), critical=True)
    try:
        check_email()
    except Exception as e:
        check("Email check ran", False, str(e), critical=True)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100.0) if total else 0.0
    critical_ok = len(CRITICAL_FAILS) == 0
    all_ok = critical_ok and accuracy >= 70.0

    print("\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}, Accuracy: {accuracy:.1f}%")
    if CRITICAL_FAILS:
        print(f"  CRITICAL FAILURES: {CRITICAL_FAILS}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT,
                       "accuracy": accuracy, "critical_fails": CRITICAL_FAILS,
                       "success": all_ok}, f, indent=2)

    if not critical_ok:
        print("FAIL: critical check(s) failed.")
        sys.exit(1)
    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
