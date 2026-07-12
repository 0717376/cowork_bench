"""Evaluation script for pw-sf-hr-skills-gap-excel-word.

Критические проверки (CRITICAL_CHECKS): любой провал => общий FAIL независимо от
accuracy. Иначе PASS требует accuracy >= 70%.
"""
import os
import argparse, json, os, sys
import openpyxl
from docx import Document as DocxDocument

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Критические проверки: любой провал => общий FAIL независимо от accuracy.
CRITICAL_CHECKS = {
    "Department_Overview semantic (7 depts, avg_salary +/-50, salary_gap +/-100)",
    "Summary semantic core (Total_Departments=7, Below_Benchmark=7, Largest_Gap=R&D, Avg_Salary_Gap)",
    "Skills_Matrix semantic (17 rows, sampled Gap_Pct + Role/Skill literals)",
    "Word report: >=3 headings and gap/recommendation/benchmark concept (RU+EN)",
    "Terminal artifacts: skills_analysis.py, raw_skills_data.json, processed_analysis.json exist",
}

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(',', '').replace('%', '').replace('$', '').strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)


def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT, FAILED_NAMES
    PASS_COUNT = 0
    FAIL_COUNT = 0
    FAILED_NAMES = []

    # ---- Excel ----
    excel_path = os.path.join(agent_workspace, "Skills_Gap_Analysis.xlsx")
    check("Excel file exists", os.path.exists(excel_path))

    # Track per-section semantic correctness for critical gating.
    dept_overview_ok = False
    summary_ok = False
    skills_matrix_ok = False

    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)

        gt_path = os.path.join(groundtruth_workspace, "Skills_Gap_Analysis.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path)

        # ---- Sheet 1: Department_Overview ----
        check("Department_Overview sheet exists", "Department_Overview" in wb.sheetnames)
        if "Department_Overview" in wb.sheetnames:
            ws = wb["Department_Overview"]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            rows = [r for r in rows if r and r[0] is not None]
            check("Department_Overview has 7 departments", len(rows) == 7, f"got {len(rows)}")

            gt_ws = gt_wb["Department_Overview"]
            gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))

            dept_overview_ok = (len(rows) == 7)
            for gt_row in gt_rows:
                dept = gt_row[0]
                agent_row = None
                for r in rows:
                    if r[0] and str(r[0]).strip().lower() == str(dept).strip().lower():
                        agent_row = r
                        break
                if agent_row is None:
                    dept_overview_ok = False
                    check(f"{dept} present in Department_Overview", False, "department missing")
                    continue
                gt_sal = safe_float(gt_row[2])
                ag_sal = safe_float(agent_row[2])
                sal_ok = gt_sal is not None and ag_sal is not None and abs(gt_sal - ag_sal) <= 50
                check(f"{dept} avg salary", sal_ok, f"expected ~{gt_sal}, got {ag_sal}")
                gt_gap = safe_float(gt_row[6])
                ag_gap = safe_float(agent_row[6])
                gap_ok = gt_gap is not None and ag_gap is not None and abs(gt_gap - ag_gap) <= 100
                check(f"{dept} salary gap", gap_ok, f"expected ~{gt_gap}, got {ag_gap}")
                if not (sal_ok and gap_ok):
                    dept_overview_ok = False

        # ---- Sheet 2: Skills_Matrix ----
        check("Skills_Matrix sheet exists", "Skills_Matrix" in wb.sheetnames)
        if "Skills_Matrix" in wb.sheetnames:
            ws = wb["Skills_Matrix"]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            rows = [r for r in rows if r and any(c is not None for c in r)]
            n_ok = len(rows) == 17
            check("Skills_Matrix has 17 skills", n_ok, f"got {len(rows)}")

            # Build header map from agent sheet.
            headers = [str(c.value).strip() if c.value else "" for c in ws[1]]
            hmap = {h.lower(): i for i, h in enumerate(headers)}
            role_i = hmap.get("role")
            skill_i = hmap.get("skill")
            gap_i = hmap.get("gap_pct")

            # Sampled value validation: Data Engineer / Python => Gap_Pct ~ -52.8.
            sample_ok = False
            literals_ok = False
            if role_i is not None and skill_i is not None and gap_i is not None:
                literals_ok = True  # required English literals present somewhere
                found_de_python = False
                # Verify a few English role/skill literals are present.
                role_skill_pairs = set()
                for r in rows:
                    role = str(r[role_i]).strip() if r[role_i] else ""
                    skill = str(r[skill_i]).strip() if r[skill_i] else ""
                    role_skill_pairs.add((role, skill))
                    if role == "Data Engineer" and skill == "Python":
                        gp = safe_float(r[gap_i])
                        if gp is not None and abs(gp - (-52.8)) <= 1.0:
                            sample_ok = True
                        found_de_python = True
                required_pairs = {
                    ("Data Engineer", "Python"),
                    ("Software Developer", "Java"),
                    ("Sales Executive", "Negotiation"),
                    ("Support Specialist", "ITIL"),
                }
                literals_ok = required_pairs.issubset(role_skill_pairs)
                if not found_de_python:
                    sample_ok = False
            check("Skills_Matrix Role/Skill English literals present", literals_ok,
                  "missing expected Role/Skill pairs")
            check("Skills_Matrix sampled Gap_Pct (Data Engineer/Python ~ -52.8)", sample_ok)
            skills_matrix_ok = n_ok and sample_ok and literals_ok

        # ---- Sheet 3: Summary ----
        check("Summary sheet exists", "Summary" in wb.sheetnames)
        if "Summary" in wb.sheetnames:
            ws = wb["Summary"]
            srows = {str(r[0]).strip(): r[1] for r in ws.iter_rows(min_row=2, values_only=True) if r[0]}
            total_dep = safe_float(srows.get("Total_Departments"))
            below = safe_float(srows.get("Departments_Below_Benchmark"))
            largest = str(srows.get("Largest_Gap_Department")).strip() if srows.get("Largest_Gap_Department") else ""
            avg_gap = safe_float(srows.get("Average_Salary_Gap"))

            check("Total_Departments = 7", total_dep == 7, f"got {srows.get('Total_Departments')}")
            check("Departments_Below_Benchmark = 7", below == 7, f"got {srows.get('Departments_Below_Benchmark')}")
            check("Largest_Gap_Department = R&D", largest == "R&D", f"got {largest}")
            avg_ok = avg_gap is not None and abs(avg_gap - (-45747.12)) <= 200
            check("Average_Salary_Gap ~ -45747.12", avg_ok, f"got {avg_gap}")
            summary_ok = (total_dep == 7 and below == 7 and largest == "R&D" and avg_ok)

    # ---- Word document ----
    word_path = os.path.join(agent_workspace, "Skills_Gap_Report.docx")
    word_exists = os.path.exists(word_path)
    check("Word report exists", word_exists)
    word_ok = False
    if word_exists:
        from docx import Document
        doc = Document(word_path)
        full_text = " ".join([p.text for p in doc.paragraphs]).lower()
        # Broadened RU+EN matching (агент пишет прозу на русском).
        gap_terms = ["skills gap", "навык", "пробел", "дефицит", "gap"]
        rec_terms = ["recommend", "рекоменд"]
        bench_terms = ["benchmark", "бенчмарк", "эталон", "ориентир"]
        has_gap = any(t in full_text for t in gap_terms)
        has_rec = any(t in full_text for t in rec_terms)
        has_bench = any(t in full_text for t in bench_terms)
        check("Word contains gap concept (RU+EN)", has_gap)
        check("Word contains recommendation concept (RU+EN)", has_rec)
        check("Word contains benchmark concept (RU+EN)", has_bench)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        has_headings = len(headings) >= 3
        check("Word has at least 3 headings", has_headings, f"got {len(headings)}")
        word_ok = has_gap and has_rec and has_bench and has_headings

    # ---- Terminal artifacts ----
    script_path = os.path.join(agent_workspace, "skills_analysis.py")
    json_path = os.path.join(agent_workspace, "raw_skills_data.json")
    processed_path = os.path.join(agent_workspace, "processed_analysis.json")
    script_ex = os.path.exists(script_path)
    json_ex = os.path.exists(json_path)
    proc_ex = os.path.exists(processed_path)
    check("Python script exists", script_ex)
    check("Raw JSON data exists", json_ex)
    check("Processed JSON exists", proc_ex)
    artifacts_ok = script_ex and json_ex and proc_ex

    # ---- Register critical (semantic) gate results ----
    check("Department_Overview semantic (7 depts, avg_salary +/-50, salary_gap +/-100)",
          dept_overview_ok)
    check("Summary semantic core (Total_Departments=7, Below_Benchmark=7, Largest_Gap=R&D, Avg_Salary_Gap)",
          summary_ok)
    check("Skills_Matrix semantic (17 rows, sampled Gap_Pct + Role/Skill literals)",
          skills_matrix_ok)
    check("Word report: >=3 headings and gap/recommendation/benchmark concept (RU+EN)",
          word_ok)
    check("Terminal artifacts: skills_analysis.py, raw_skills_data.json, processed_analysis.json exist",
          artifacts_ok)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total else 0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== Results: {PASS_COUNT}/{total} passed ({accuracy:.1f}%) ===")
    if critical_failed:
        print(f"CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")

    if res_log_file:
        try:
            with open(res_log_file, "w") as f:
                json.dump({
                    "total_passed": PASS_COUNT, "total_checks": total,
                    "accuracy": accuracy, "critical_failed": critical_failed,
                }, f, indent=2)
        except Exception:
            pass

    success = (not critical_failed) and accuracy >= 70
    return success, f"Passed {PASS_COUNT}/{total} checks ({accuracy:.1f}%)"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
