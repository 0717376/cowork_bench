"""
Evaluation for academic-conf-volga-rus-kazan-forms-teamly-excel-word-email.

Проверки:
- Excel Submissions.xlsx: оба листа, 10/7 строк, статусы, тематики, средние.
- Teamly: страница «Программа конференции ДРВБ-2026» создана, упомянуты все 7 принятых.
- Word: Conference_Program.docx создан, упоминает даты, площадку, ключевые имена.
- Emails: 10 писем отправлены каждому автору на правильный email с правильным статусом.

Критические чеки (CRITICAL_CHECKS): любой их fail = задача FAIL даже при высокой
общей accuracy. Это:
  - C1: Excel: 10 строк всех авторов с верным статусом acc/rej.
  - C2: Teamly: страница «Программа конференции ДРВБ-2026» создана, упомянуты все 7 принятых.
  - C3: Emails: 10 писем — по одному каждому из 10 авторов на их email.
"""
from argparse import ArgumentParser
import json
import os
import sys
from pathlib import Path

import psycopg2

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

SUBMISSIONS = [
    # (ФИО, email, accept?)
    ("Хабибуллин М.Х.", "habibullin@kfu.ru", True),
    ("Дроздова О.В.", "drozdova@iling.spb.ru", True),
    ("Петров А.С.", "petrov@hist.msu.ru", True),
    ("Шарифуллина Е.Р.", "sharifullina@kfu.ru", True),
    ("Иванов К.М.", "ivanov@nov-arch.ru", True),
    ("Ласкина Н.Ю.", "laskina@inion.ru", True),
    ("Соколов Д.А.", "sokolov@kfu.ru", True),
    ("Краснов В.П.", "krasnov@example.ru", False),
    ("Михайлов С.К.", "mikhailov@yandex.ru", False),
    ("Орлов Р.Е.", "orlov@rambler.ru", False),
]

ACCEPTED_LAST_NAMES = {fio.split()[0] for fio, _e, ok in SUBMISSIONS if ok}
ALL_LAST_NAMES = {fio.split()[0] for fio, _e, _ok in SUBMISSIONS}
EMAILS_TO = [e for _fio, e, _ok in SUBMISSIONS]

# Критические чеки (по строке name, как при record())
CRITICAL_CHECKS = {
    "Excel All_Submissions: 10 строк со всеми ФИО и верными статусами",
    "Teamly: страница «Программа конференции ДРВБ-2026» содержит всех 7 принятых",
    "Emails: 10 писем отправлены — по одному каждому автору на правильный email",
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILED
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILED.append(name)


def norm(s):
    return (s or "").strip().lower()


# ---------------------------------------------------------------------------
# Excel checks
# ---------------------------------------------------------------------------
def check_excel(workspace: Path):
    print("\n=== Check 1: Excel Submissions.xlsx ===")
    xlsx_path = workspace / "Submissions.xlsx"
    if not xlsx_path.exists():
        record("Submissions.xlsx exists", False, str(xlsx_path))
        for nm in (
            "Excel readable",
            "Excel has All_Submissions sheet",
            "Excel has Accepted_Program sheet",
            "Excel All_Submissions: 10 строк со всеми ФИО и верными статусами",
            "Accepted_Program: 7 строк всех принятых",
        ):
            record(nm, False, "no file")
        return
    record("Submissions.xlsx exists", True)

    if openpyxl is None:
        record("Excel readable", False, "openpyxl не установлен")
        return

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    except Exception as e:
        record("Excel readable", False, str(e))
        return
    record("Excel readable", True)

    sheets = wb.sheetnames
    record("Excel has All_Submissions sheet", "All_Submissions" in sheets, str(sheets))
    record("Excel has Accepted_Program sheet", "Accepted_Program" in sheets, str(sheets))

    # ---- All_Submissions ----
    if "All_Submissions" in sheets:
        ws = wb["All_Submissions"]
        rows = list(ws.iter_rows(values_only=True))
        record("All_Submissions: есть заголовок и >= 10 строк данных",
               len(rows) >= 11, f"rows={len(rows)}")

        # Собираем словарь last_name → row
        header = [str(c or "").strip() for c in rows[0]] if rows else []

        def find_col(name_substr):
            for i, h in enumerate(header):
                if name_substr.lower() in h.lower():
                    return i
            return None

        col_fio = find_col("ФИО") or find_col("автор")
        col_status = find_col("статус")
        col_topic = find_col("тематик")
        col_score = find_col("средн")

        found_authors = {}  # last_name → row tuple
        for r in rows[1:]:
            if not r:
                continue
            fio_cell = str(r[col_fio] or "") if col_fio is not None else ""
            if not fio_cell:
                continue
            ln = fio_cell.strip().split()[0]
            found_authors[ln] = r

        record("All_Submissions: упомянуты все 10 фамилий",
               ALL_LAST_NAMES.issubset(set(found_authors.keys())),
               f"missing={ALL_LAST_NAMES - set(found_authors.keys())}")

        # CRITICAL: 10 строк со статусами accepted/rejected правильно
        correct_status = 0
        wrong = []
        for fio, _e, ok in SUBMISSIONS:
            ln = fio.split()[0]
            r = found_authors.get(ln)
            if r is None:
                wrong.append(f"{ln}: нет строки")
                continue
            status_raw = ""
            if col_status is not None:
                status_raw = str(r[col_status] or "")
            sn = norm(status_raw)
            expected_accept = ok
            if expected_accept:
                ok_match = ("прин" in sn or "accept" in sn or "✓" in sn)
            else:
                ok_match = ("отклон" in sn or "reject" in sn or "✗" in sn or "не прин" in sn)
            if ok_match:
                correct_status += 1
            else:
                wrong.append(f"{ln}: status='{status_raw}' (ожидался {'accept' if ok else 'reject'})")
        record("Excel All_Submissions: 10 строк со всеми ФИО и верными статусами",
               correct_status == 10, f"correct={correct_status}/10; wrong={wrong[:5]}")

        # Тематики и средние оценки — мягкие чеки
        thematic_ok = 0
        if col_topic is not None:
            for r in rows[1:]:
                if r and r[col_topic]:
                    s = norm(str(r[col_topic]))
                    if any(k in s for k in ("археолог", "источник", "истор", "иное")):
                        thematic_ok += 1
        record("All_Submissions: тематики заполнены (>=8)", thematic_ok >= 8, f"good={thematic_ok}")

        scores_ok = 0
        if col_score is not None:
            for r in rows[1:]:
                if r and r[col_score] is not None:
                    try:
                        v = float(str(r[col_score]).replace(",", "."))
                        if 0 <= v <= 5:
                            scores_ok += 1
                    except (TypeError, ValueError):
                        pass
        record("All_Submissions: средние оценки в диапазоне 0..5 (>=8)",
               scores_ok >= 8, f"good={scores_ok}")

    # ---- Accepted_Program ----
    if "Accepted_Program" in sheets:
        ws = wb["Accepted_Program"]
        rows = list(ws.iter_rows(values_only=True))
        record("Accepted_Program: есть заголовок и >= 7 строк данных",
               len(rows) >= 8, f"rows={len(rows)}")

        header = [str(c or "").strip() for c in rows[0]] if rows else []
        all_text = " ".join(str(c or "") for r in rows[1:] for c in r)
        accepted_found = sum(1 for ln in ACCEPTED_LAST_NAMES if ln in all_text)
        record("Accepted_Program: 7 строк всех принятых",
               accepted_found >= 7, f"found {accepted_found}/7")

        # Секции должны быть упомянуты
        nl = all_text.lower()
        sections = sum(1 for s in ("археолог", "источник", "истор") if s in nl)
        record("Accepted_Program: упомянуты все 3 секции (археология/источниковедение/история)",
               sections == 3, f"sections={sections}")

        # Даты конференции
        date_hits = sum(1 for d in ("2026-03-17", "2026-03-18", "2026-03-19",
                                     "17.03", "18.03", "19.03",
                                     "17 марта", "18 марта", "19 марта")
                        if d.lower() in nl)
        record("Accepted_Program: упомянуты даты 17/18/19 марта",
               date_hits >= 3, f"hits={date_hits}")


# ---------------------------------------------------------------------------
# Teamly checks
# ---------------------------------------------------------------------------
def check_teamly():
    print("\n=== Check 2: Teamly «Программа конференции ДРВБ-2026» ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
    except Exception as e:
        record("Teamly: подключение к БД", False, str(e))
        for nm in ("Teamly: страница «Программа конференции ДРВБ-2026» создана",
                   "Teamly: страница «Программа конференции ДРВБ-2026» содержит всех 7 принятых",
                   "Teamly: упомянуты все 3 секции",
                   "Teamly: упомянуты даты 17-19 марта"):
            record(nm, False, "no db")
        return
    with conn.cursor() as cur:
        cur.execute("""
            SELECT title, body
            FROM teamly.pages
            WHERE title ILIKE '%Программа%конференции%ДРВБ%'
               OR title ILIKE '%Программа%ДРВБ%2026%'
        """)
        rows = cur.fetchall()
    conn.close()

    record("Teamly: страница «Программа конференции ДРВБ-2026» создана",
           len(rows) >= 1, f"found={len(rows)}")
    if not rows:
        for nm in ("Teamly: страница «Программа конференции ДРВБ-2026» содержит всех 7 принятых",
                   "Teamly: упомянуты все 3 секции",
                   "Teamly: упомянуты даты 17-19 марта"):
            record(nm, False, "no page")
        return

    body = "\n\n".join(b or "" for _t, b in rows)
    accepted_in_body = sum(1 for ln in ACCEPTED_LAST_NAMES if ln in body)
    record("Teamly: страница «Программа конференции ДРВБ-2026» содержит всех 7 принятых",
           accepted_in_body >= 7, f"found {accepted_in_body}/7")

    nl = body.lower()
    sections = sum(1 for s in ("археолог", "источник", "истор") if s in nl)
    record("Teamly: упомянуты все 3 секции", sections == 3, f"sections={sections}")

    date_hits = sum(1 for d in ("2026-03-17", "2026-03-18", "2026-03-19",
                                 "17.03", "18.03", "19.03",
                                 "17 марта", "18 марта", "19 марта")
                    if d.lower() in nl)
    record("Teamly: упомянуты даты 17-19 марта", date_hits >= 3, f"hits={date_hits}")


# ---------------------------------------------------------------------------
# Word checks
# ---------------------------------------------------------------------------
def check_word(workspace: Path):
    print("\n=== Check 3: Word Conference_Program.docx ===")
    docx_path = workspace / "Conference_Program.docx"
    if not docx_path.exists():
        record("Conference_Program.docx exists", False, str(docx_path))
        for nm in ("Word readable",
                   "Word: упомянуты даты 17-19 марта 2026",
                   "Word: упомянута площадка КФУ или Казанский федеральный университет",
                   "Word: упомянуто >= 5 принятых авторов",
                   "Word: упомянуты все 3 секции"):
            record(nm, False, "no file")
        return
    record("Conference_Program.docx exists", True)

    if DocxDocument is None:
        record("Word readable", False, "python-docx не установлен")
        return

    try:
        doc = DocxDocument(docx_path)
    except Exception as e:
        record("Word readable", False, str(e))
        return
    record("Word readable", True)

    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    nl = text.lower()

    date_hits = sum(1 for d in ("2026-03-17", "2026-03-18", "2026-03-19",
                                 "17.03", "18.03", "19.03",
                                 "17 марта", "18 марта", "19 марта") if d.lower() in nl)
    record("Word: упомянуты даты 17-19 марта 2026", date_hits >= 2, f"hits={date_hits}")

    venue_ok = ("кфу" in nl
                or "казанский федеральный" in nl
                or "казань" in nl)
    record("Word: упомянута площадка КФУ или Казанский федеральный университет",
           venue_ok, f"venue check")

    authors_in_word = sum(1 for ln in ACCEPTED_LAST_NAMES if ln in text)
    record("Word: упомянуто >= 5 принятых авторов",
           authors_in_word >= 5, f"found {authors_in_word}/7")

    sections = sum(1 for s in ("археолог", "источник", "истор") if s in nl)
    record("Word: упомянуты все 3 секции", sections == 3, f"sections={sections}")


# ---------------------------------------------------------------------------
# Email checks
# ---------------------------------------------------------------------------
def check_emails():
    print("\n=== Check 4: Emails — 10 уведомлений ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
    except Exception as e:
        record("Emails: 10 писем отправлены — по одному каждому автору на правильный email",
               False, str(e))
        return

    sent = {}  # to_email → subject+body
    with conn.cursor() as cur:
        # email.messages: from_addr / to_addr / subject / body_text. Письма агента
        # имеют from_addr, отличающийся от seeded inbox (там from = moscow_team@uni.ru
        # и т.п.), но проще — берём все письма, отправленные на наш список адресов.
        try:
            cur.execute("""
                SELECT to_addr::text, subject, body_text
                FROM email.messages
                WHERE to_addr::text ILIKE ANY (ARRAY[%s,%s,%s,%s,%s,%s,%s,%s,%s,%s])
            """, tuple(f"%{e}%" for e in EMAILS_TO))
            rows = cur.fetchall()
        except Exception as e:
            rows = []
            record("Emails: подключение и выборка", False, str(e))
            return
    conn.close()

    for to_field, subj, body in rows:
        # to_field — JSONB или массив; ищем нашу почту как substring
        for e in EMAILS_TO:
            if e in (to_field or ""):
                sent.setdefault(e, []).append((subj or "", body or ""))

    delivered = sum(1 for e in EMAILS_TO if e in sent)
    record("Emails: 10 писем отправлены — по одному каждому автору на правильный email",
           delivered >= 10, f"delivered={delivered}/10")

    # Soft: правильный смысл (accept/reject) в письме
    correct_meaning = 0
    for fio, e, ok in SUBMISSIONS:
        msgs = sent.get(e, [])
        if not msgs:
            continue
        text = " ".join(s + " " + b for s, b in msgs).lower()
        if ok:
            if ("прин" in text and "доклад" in text) or "accept" in text:
                correct_meaning += 1
        else:
            if "отклон" in text or "reject" in text or "не прин" in text:
                correct_meaning += 1
    record("Emails: тон/смысл совпадают со статусом (>=8 из 10)",
           correct_meaning >= 8, f"correct={correct_meaning}/10")

    # Soft: упоминание ДРВБ-2026 хотя бы в 7 письмах
    drvb_hits = 0
    for e, msgs in sent.items():
        text = " ".join(s + " " + b for s, b in msgs).lower()
        if "дрвб" in text or "drvb" in text or "древняя русь" in text or "волжск" in text:
            drvb_hits += 1
    record("Emails: упоминают конференцию ДРВБ-2026 (>=7)",
           drvb_hits >= 7, f"hits={drvb_hits}")


# ---------------------------------------------------------------------------
def main():
    parser = ArgumentParser()
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    args = parser.parse_args()

    if not args.agent_workspace:
        print("No agent_workspace provided.")
        sys.exit(1)
    workspace = Path(args.agent_workspace)

    check_excel(workspace)
    check_teamly()
    check_word(workspace)
    check_emails()

    total = PASS_COUNT + FAIL_COUNT
    pct = 100.0 * PASS_COUNT / total if total else 0
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({pct:.1f}%)")
    if CRITICAL_FAILED:
        print(f"CRITICAL FAIL: {CRITICAL_FAILED}")
        print("FAIL")
        sys.exit(1)
    if pct < 80.0:
        print("FAIL (accuracy < 80%)")
        sys.exit(1)
    print("PASS")
    sys.exit(0)


if __name__ == "__main__":
    main()
