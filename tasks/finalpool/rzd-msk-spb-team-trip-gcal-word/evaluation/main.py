"""Evaluation for rzd-msk-spb-team-trip-gcal-word (hard version).

Команда из 3 человек, бюджет 40 000 ₽. Эконом 5 500 ₽ × 3 × 2 = 33 000 ≤ 40 000 → fits.
Эконом+ 7 500 ₽ × 3 × 2 = 45 000 > 40 000 → не fits. Единственный валидный класс — Эконом.

Проверки (20):
  Word (10):  docx exists + readable + 3 headings + trains + cities + times +
              цены: 5500 за билет, 16500 на команду в одну сторону, 33000 итого +
              ≥2 таблицы
  trip_notes (3): файл + поезда+дата + 3 человека и итого 33000
  Calendar (2):   ≥2 события на 2026-03-10 + направление Москва ↔ СПб
  Email (5):      sent + оба поезда + subject + 3 человека/трое + итого 33000
"""
import json
import os
import sys
import unicodedata
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

# Expected answer: outbound 752А (06:50→10:50), return 759А (19:25→23:25).
# Class: Эконом, единственный укладывающийся в бюджет на команду из 3.
TRAIN_OUTBOUND = "752"
TRAIN_RETURN   = "759"
DEPART_OUT, ARRIVE_OUT = "06:50", "10:50"
DEPART_RET, ARRIVE_RET = "19:25", "23:25"
PRICE_PER_TICKET = 5500            # Эконом, одно направление, один пассажир
TEAM_SIZE        = 3
PER_LEG_TEAM     = 16500           # 5500 * 3
TOTAL_TEAM       = 33000           # 5500 * 3 * 2
BUDGET           = 40000

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def normalize(s: str) -> str:
    """Map cyrillic А/О/Е/Р/С/У/К/Х to latin lookalikes so "752А" == "752A"."""
    s = s.lower()
    s = unicodedata.normalize("NFKD", s)
    table = str.maketrans({"а": "a", "о": "o", "е": "e", "р": "p",
                           "с": "c", "у": "y", "к": "k", "х": "x"})
    return s.translate(table)


def contains_price(text: str, value: int) -> bool:
    """Match price like 33000, 33 000, 33,000."""
    patterns = [str(value), f"{value // 1000} {value % 1000:03d}",
                f"{value // 1000},{value % 1000:03d}"]
    return any(p in text for p in patterns)


# ---------------------------------------------------------------------------
# Word doc
# ---------------------------------------------------------------------------

def check_word_doc(agent_workspace):
    print("\n=== Word: Travel_Plan.docx ===")
    import glob
    candidates = glob.glob(os.path.join(agent_workspace, "*.docx"))
    if not candidates:
        record("Travel_Plan.docx существует", False, f"docx not found in {agent_workspace}")
        for _ in range(9):
            record("(skipped — no docx)", False)
        return
    doc_path = next((c for c in candidates if any(k in c.lower() for k in
                     ("travel", "plan", "поездк"))), candidates[0])
    record("Travel_Plan.docx существует", True)

    try:
        import docx
        doc = docx.Document(doc_path)
    except Exception as e:
        record("Документ читается", False, str(e))
        for _ in range(8):
            record("(skipped — docx unreadable)", False)
        return
    record("Документ читается", True)

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    heading_text = " | ".join(headings).lower()
    required = ["поездка туда", "поездка обратно", "сводк"]
    record("Заголовки 'Поездка туда / Поездка обратно / Сводка'",
           all(h in heading_text for h in required) and len(headings) >= 3,
           f"headings: {headings}")

    full_text = " ".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text += " " + cell.text
    norm = normalize(full_text)
    lower = full_text.lower()

    has_out = TRAIN_OUTBOUND in norm and "a" in norm.split(TRAIN_OUTBOUND, 1)[1][:2]
    has_ret = TRAIN_RETURN in norm and "a" in norm.split(TRAIN_RETURN, 1)[1][:2]
    record(f"Номера поездов {TRAIN_OUTBOUND}А и {TRAIN_RETURN}А",
           has_out and has_ret, f"out={has_out}, ret={has_ret}")

    has_msk = "москв" in lower
    has_spb = "петербург" in lower or "спб" in lower
    record("Города: Москва и Санкт-Петербург", has_msk and has_spb,
           f"msk={has_msk}, spb={has_spb}")

    has_dep = DEPART_OUT in full_text and DEPART_RET in full_text
    has_arr = ARRIVE_OUT in full_text and ARRIVE_RET in full_text
    record("Времена отправления и прибытия", has_dep and has_arr,
           f"dep={has_dep}, arr={has_arr}")

    record(f"Цена одного билета {PRICE_PER_TICKET}₽",
           contains_price(full_text, PRICE_PER_TICKET))
    record(f"Стоимость на команду за одно направление {PER_LEG_TEAM}₽ "
           f"({PRICE_PER_TICKET}×{TEAM_SIZE})",
           contains_price(full_text, PER_LEG_TEAM),
           f"text sample: {full_text[:150]}")
    record(f"Итого на команду {TOTAL_TEAM}₽",
           contains_price(full_text, TOTAL_TEAM),
           f"text sample: {full_text[:150]}")

    record("Не менее 2 таблиц в документе", len(doc.tables) >= 2,
           f"tables={len(doc.tables)}")


# ---------------------------------------------------------------------------
# trip_notes.md
# ---------------------------------------------------------------------------

def check_trip_notes(agent_workspace):
    print("\n=== trip_notes.md ===")
    path = os.path.join(agent_workspace, "trip_notes.md")
    if not os.path.isfile(path):
        import glob
        candidates = [p for p in glob.glob(os.path.join(agent_workspace, "*.md"))
                      if any(kw in os.path.basename(p).lower()
                             for kw in ("trip", "поездк", "notes"))]
        if not candidates:
            record("trip_notes.md существует в workspace", False,
                   f"file not found in {agent_workspace}")
            for _ in range(2):
                record("(skipped — no notes file)", False)
            return
        path = candidates[0]
    record("trip_notes.md существует в workspace", True)

    text = open(path, encoding="utf-8", errors="replace").read()
    norm = normalize(text)
    has_trains = TRAIN_OUTBOUND in norm and TRAIN_RETURN in norm
    has_date = ("10.03" in text or "10 марта" in text.lower()
                or "2026-03-10" in text)
    record("trip_notes.md упоминает оба поезда и дату",
           has_trains and has_date, f"trains={has_trains}, date={has_date}")

    has_team = "3" in text and any(w in text.lower() for w in
                                    ["человек", "пассажир", "участник",
                                     "трое", "команд"])
    has_total = contains_price(text, TOTAL_TEAM)
    record(f"trip_notes.md содержит 3 человека и итог {TOTAL_TEAM}₽",
           has_team and has_total, f"team={has_team}, total={has_total}")


# ---------------------------------------------------------------------------
# Calendar
# ---------------------------------------------------------------------------

def check_gcal():
    print("\n=== Calendar: 2 события на 2026-03-10 ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, description, start_datetime, end_datetime
          FROM gcal.events
         WHERE start_datetime >= '2026-03-10'
           AND start_datetime <  '2026-03-11'
           AND summary NOT ILIKE '%встреча с клиентом%'
           AND summary NOT ILIKE '%client meeting%'
         ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    record("Не менее 2 событий поездки в календаре", len(events) >= 2,
           f"events: {[e[0] for e in events]}")

    direction_ok = False
    for summary, descr, _s, _e in events:
        combined = f"{summary or ''} {descr or ''}".lower()
        if "москв" in combined and ("петербург" in combined or "спб" in combined):
            direction_ok = True
            break
    record("Событие упоминает направление Москва ↔ СПб",
           direction_ok, f"events: {[e[0] for e in events]}")


# ---------------------------------------------------------------------------
# Email
# ---------------------------------------------------------------------------

def check_email():
    print("\n=== Email: travel@consulting.ru ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute(
        "SELECT m.subject, m.body_text "
        "  FROM email.messages m "
        "  JOIN email.sent_log s ON s.message_id = m.id "
        " WHERE m.to_addr::text ILIKE %s "
        "   AND m.from_addr NOT ILIKE %s ",
        ("%travel@consulting.ru%", "%travel@consulting.ru%"),
    )
    rows = cur.fetchall()
    if not rows:
        cur.execute(
            "SELECT subject, body_text FROM email.messages "
            " WHERE to_addr::text ILIKE %s "
            "   AND from_addr NOT ILIKE %s ",
            ("%travel@consulting.ru%", "%travel@consulting.ru%"),
        )
        rows = cur.fetchall()
    cur.close()
    conn.close()

    record("Письмо отправлено на travel@consulting.ru", len(rows) >= 1,
           f"matched rows: {len(rows)}")
    if not rows:
        for _ in range(4):
            record("(skipped — no email found)", False)
        return

    body = " ".join((b or "") for _s, b in rows)
    subj = " ".join((s or "") for s, _b in rows).lower()
    body_norm = normalize(body)
    body_lower = body.lower()

    has_both_trains = TRAIN_OUTBOUND in body_norm and TRAIN_RETURN in body_norm
    record(f"В теле письма оба поезда ({TRAIN_OUTBOUND}А и {TRAIN_RETURN}А)",
           has_both_trains, f"body sample: {body[:180]!r}")

    subj_ok = any(root in subj for root in ["поездк", "travel", "деловая", "командировк"])
    record("Subject содержит признак деловой поездки", subj_ok,
           f"subject: {subj[:120]!r}")

    has_team = "3" in body and any(w in body_lower for w in
                                    ["человек", "пассажир", "участник",
                                     "трое", "команд"])
    record("В письме упомянута команда (3 человека)", has_team,
           f"body sample: {body[:180]!r}")

    has_total = contains_price(body, TOTAL_TEAM)
    record(f"В письме итоговая стоимость {TOTAL_TEAM}₽", has_total,
           f"body sample: {body[:180]!r}")


# ---------------------------------------------------------------------------

def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_trip_notes(args.agent_workspace)
    check_gcal()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\nИтого: {PASS_COUNT}/{total} проверок пройдено ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
