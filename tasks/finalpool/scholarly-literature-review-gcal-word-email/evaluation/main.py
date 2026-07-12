"""Evaluation for scholarly-literature-review-gcal-word-email.

Строгая проверка. Критические проверки (CRITICAL_CHECKS): любой провал ->
итоговый FAIL независимо от accuracy. Иначе PASS требует accuracy >= 70%.

Проверки:
  Word (Literature_Review.docx):
    - файл существует и читается
    - заголовок содержит Literature/Review (или "Обзор литературы")
    - присутствуют данные о цитировании (годы/авторы)
    - [CRITICAL] >= 6 из отобранных по теме статей присутствуют (по названию)
    - [CRITICAL] off-topic статьи-дистракторы исключены (ImageNet 2012, Survey 2015)
    - [CRITICAL] реальные числа цитирований из scholarly присутствуют (не выдуманы)
    - раздел Summary/Conclusion (RU: Заключение/Итоги/Выводы)
    - >= 7 заголовков (титул + 6 статей + summary)
  Calendar:
    - [CRITICAL] событие "Literature Review Presentation" на целевую дату (launch+14),
      начало 14:00 UTC, описание про LLM/large language models
  Email:
    - [CRITICAL] письмо на research-team@university.example.com с темой Literature Review
      (RU: "Обзор литературы"), в теле — статьи для чтения и дата презентации
"""
import argparse
import json
import os
import sys
from datetime import datetime, timedelta

import psycopg2

DB = {"host": os.environ.get("PGHOST", "localhost"), "port": 5432, "dbname": "cowork_gym", "user": "eigent", "password": "camel"}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# In-scope seeded paper titles (must be selected). Each tuple: (key fragment, citation count).
IN_SCOPE = [
    ("attention is all you need", 86000),
    ("bert", 62000),
    ("language models are few-shot learners", 42000),
    ("training language models to follow instructions", 18000),
    ("llama", 15000),
    ("scaling laws for neural language models", 8500),
]
# Off-topic distractors that MUST be excluded (outside 2017-2023 LLM scope).
DISTRACTORS = ["imagenet classification with deep convolutional", "a survey of deep learning"]
# Actual numeric citation values seeded for in-scope papers.
CITATION_VALUES = ["86000", "62000", "42000", "18000", "15000", "8500"]

# Critical checks: any failure => overall FAIL regardless of accuracy.
CRITICAL_CHECKS = {
    "Doc contains >= 6 in-scope seeded paper titles",
    "Off-topic distractor papers are excluded from doc",
    "Doc contains real per-paper citation counts (data read, not fabricated)",
    "GCal event on target date with 14:00 UTC start and LLM description",
    "Email to research-team with subject and substantive body",
}

RU_MONTHS = {1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая", 6: "июня",
             7: "июля", 8: "августа", 9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"}
EN_MONTHS = {1: "January", 2: "February", 3: "March", 4: "April", 5: "May", 6: "June",
             7: "July", 8: "August", 9: "September", 10: "October", 11: "November", 12: "December"}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        d = (str(detail)[:300] + "...") if len(str(detail)) > 300 else str(detail)
        print(f"  [FAIL] {name}: {d}")


def compute_target_date(launch_time):
    """Презентация через 14 дней после запуска. Fallback: 2026-03-21."""
    if launch_time:
        # tolerate weekday tail, e.g. "2026-03-07 10:00:00 Saturday"
        lt = " ".join(str(launch_time).split("+")[0].strip().split()[:2])
        for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S.%f",
                    "%Y-%m-%d", "%Y-%m-%dT%H:%M:%SZ"):
            try:
                base = datetime.strptime(lt, fmt)
                return (base + timedelta(days=14)).date()
            except (ValueError, AttributeError):
                continue
    return datetime(2026, 3, 7).date() + timedelta(days=14)  # -> 2026-03-21


def check_word(agent_ws):
    print("\n=== Check 1: Literature_Review.docx ===")
    path = os.path.join(agent_ws, "Literature_Review.docx")
    check("File Literature_Review.docx exists", os.path.isfile(path))
    if not os.path.isfile(path):
        # Mark dependent critical checks failed.
        check("Doc contains >= 6 in-scope seeded paper titles", False, "no docx")
        check("Off-topic distractor papers are excluded from doc", False, "no docx")
        check("Doc contains real per-paper citation counts (data read, not fabricated)", False, "no docx")
        return

    try:
        from docx import Document
        doc = Document(path)
    except Exception as e:
        check("Word doc is readable", False, str(e))
        check("Doc contains >= 6 in-scope seeded paper titles", False, "unreadable")
        check("Off-topic distractor papers are excluded from doc", False, "unreadable")
        check("Doc contains real per-paper citation counts (data read, not fabricated)", False, "unreadable")
        return

    full_text = " ".join(p.text for p in doc.paragraphs)
    full_text_lower = full_text.lower()
    digits_only = "".join(ch for ch in full_text if ch.isdigit())

    check("Doc contains 'Literature'/'Review' or 'Обзор литературы' in title/heading",
          ("literature" in full_text_lower and "review" in full_text_lower)
          or ("обзор" in full_text_lower and "литератур" in full_text_lower))

    citation_indicators = ["citation", "citations", "cite", "цитир", "year:", "authors:",
                           "год", "автор", "2017", "2019", "2020", "2022", "2023"]
    has_citations = sum(1 for kw in citation_indicators if kw in full_text_lower)
    check("Doc contains citation data (years/authors)", has_citations >= 3,
          f"Citation indicators found: {has_citations}")

    # CRITICAL: at least 6 distinct in-scope seeded paper titles present.
    found_in_scope = [frag for frag, _ in IN_SCOPE if frag in full_text_lower]
    check("Doc contains >= 6 in-scope seeded paper titles", len(found_in_scope) >= 6,
          f"Found {len(found_in_scope)}/6: {found_in_scope}")

    # CRITICAL: off-topic distractors excluded.
    present_distractors = [d for d in DISTRACTORS if d in full_text_lower]
    check("Off-topic distractor papers are excluded from doc", len(present_distractors) == 0,
          f"Distractors present: {present_distractors}")

    # CRITICAL: real numeric citation counts appear (>=4 of the seeded values).
    found_vals = [v for v in CITATION_VALUES if v in digits_only]
    check("Doc contains real per-paper citation counts (data read, not fabricated)",
          len(found_vals) >= 4, f"Found citation values: {found_vals}")

    check("Doc has Summary/Conclusion section (RU: Заключение/Итоги/Выводы)",
          any(k in full_text_lower for k in ["summary", "conclusion", "заключени", "итог", "вывод"]))

    heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
    check("Doc has at least 7 headings (title + 6+ paper sections + summary)",
          heading_count >= 7, f"Found {heading_count} headings")


def check_gcal(target_date):
    print("\n=== Check 2: Google Calendar Event ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, summary, start_datetime, description FROM gcal.events
        WHERE summary ILIKE '%literature%' AND summary ILIKE '%review%'
        LIMIT 10
    """)
    events = cur.fetchall()
    found = len(events) > 0
    check("GCal event with 'Literature Review' in title found", found,
          "No matching events in gcal.events")

    target_str = target_date.isoformat()
    if events:
        event = events[0]
        evt_start = str(event[2] or "")
        desc = str(event[3] or "").lower()
        date_ok = target_str in evt_start
        time_ok = "14:00" in evt_start or "T14" in evt_start
        desc_ok = any(k in desc for k in ["llm", "large language", "больших языков",
                                          "большие языков", "языковых модел", "literature", "литератур"])
        check("GCal event on target date with 14:00 UTC start and LLM description",
              date_ok and time_ok and desc_ok,
              f"start={evt_start} (target {target_str}), desc={str(event[3])[:120]}")
    else:
        check("GCal event on target date with 14:00 UTC start and LLM description",
              False, "no event")

    cur.close()
    conn.close()


def check_email(target_date):
    print("\n=== Check 3: Email ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, to_addr, body_text FROM email.messages
        WHERE to_addr::text ILIKE '%research-team%'
           OR subject ILIKE '%literature%review%'
           OR subject ILIKE '%обзор%литератур%'
        LIMIT 20
    """)
    rows = cur.fetchall()
    check("Email to research-team / Literature Review subject found", len(rows) > 0,
          "No matching email found")

    # Build date strings (EN + RU) for the target presentation date.
    en_date = f"{EN_MONTHS[target_date.month]} {target_date.day}, {target_date.year}"
    en_date_alt = f"{EN_MONTHS[target_date.month]} {target_date.day} {target_date.year}"
    ru_date = f"{target_date.day} {RU_MONTHS[target_date.month]} {target_date.year}"
    iso_date = target_date.isoformat()

    if rows:
        # Prefer the row actually addressed to research-team.
        target_rows = [r for r in rows if "research-team" in str(r[1] or "").lower()] or rows
        to_addrs = [str(r[1] or "") for r in rows]
        subj_ok = any(("literature" in str(r[0] or "").lower() and "review" in str(r[0] or "").lower())
                      or ("обзор" in str(r[0] or "").lower() and "литератур" in str(r[0] or "").lower())
                      for r in target_rows)
        to_ok = any("research-team" in addr.lower() for addr in to_addrs)

        bodies = [str(r[2] or "").lower() for r in target_rows]
        paper_kw = ["attention", "bert", "gpt", "llama", "vaswani", "devlin", "brown", "touvron",
                    "scaling laws", "few-shot", "instruction"]
        body_papers_ok = any(any(k in b for k in paper_kw) for b in bodies)
        d, m, y = target_date.day, target_date.month, target_date.year
        date_variants = [en_date.lower(), en_date_alt.lower(), ru_date.lower(), iso_date,
                         f"{d:02d}.{m:02d}.{y}", f"{d}.{m}.{y}",
                         f"{d:02d}/{m:02d}/{y}", f"{d}/{m}/{y}",
                         f"{d:02d}.{m:02d}.{y % 100:02d}"]
        body_date_ok = any(any(dv in b for dv in date_variants) for b in bodies)

        check("Email to research-team with subject and substantive body",
              to_ok and subj_ok and body_papers_ok and body_date_ok,
              f"to_ok={to_ok} subj_ok={subj_ok} papers_ok={body_papers_ok} date_ok={body_date_ok}; "
              f"expected date one of {date_variants}; body[:200]={bodies[0][:200] if bodies else ''}")
    else:
        check("Email to research-team with subject and substantive body", False, "no email")

    cur.close()
    conn.close()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=== Evaluation: scholarly-literature-review-gcal-word-email ===")

    target_date = compute_target_date(args.launch_time)
    print(f"Target presentation date (launch+14): {target_date.isoformat()}")

    check_word(args.agent_workspace)
    check_gcal(target_date)
    check_email(target_date)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total else 0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== SUMMARY: {PASS_COUNT}/{total} passed ({accuracy:.1f}%) ===")
    if critical_failed:
        print(f"CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"pass": PASS_COUNT, "fail": FAIL_COUNT, "accuracy": accuracy,
                       "critical_failed": critical_failed}, f)

    success = (not critical_failed) and accuracy >= 70
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
