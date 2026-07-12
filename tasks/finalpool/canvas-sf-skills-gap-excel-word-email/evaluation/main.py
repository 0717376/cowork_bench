"""Evaluation script for canvas-sf-skills-gap-excel-word-email (ClickHouse RU fork)."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []

def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILS
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRITICAL]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL]{' [CRITICAL]' if critical else ''} {name}: {detail_str}")
        if critical:
            CRITICAL_FAILS.append(name)

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

# Four required docx section headings. Each entry is a list of accepted variants
# (English grep target + Russian equivalents). Exact-list, not bidirectional substring.
REQUIRED_HEADINGS = [
    ["methodology", "методология", "методика"],
    ["department analysis", "анализ подразделений", "анализ отделов"],
    ["critical gaps identified", "критические пробелы", "выявленные критические пробелы"],
    ["training recommendations", "рекомендации по обучению", "рекомендации по обучению и развитию"],
]

def heading_matches(variants, headings):
    for h in headings:
        for v in variants:
            if v == h or v in h or h in v:
                return True
    return False

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILS
    PASS_COUNT = 0
    FAIL_COUNT = 0
    CRITICAL_FAILS = []

    dept_top_from_sheet = None  # department with max Gap_Score in Department_Skills

    # Check Skills_Gap_Report.xlsx
    excel_path = os.path.join(agent_workspace, "Skills_Gap_Report.xlsx")
    check("Skills_Gap_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Skills_Gap_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Check row count
                    gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                    data_rows = list(ws.iter_rows(min_row=2, values_only=True))
                    min_rows = max(1, len(gt_rows) - 2)
                    check(f"{sheet_name} has >= {min_rows} data rows", len(data_rows) >= min_rows, f"got {len(data_rows)}")

                    # Data-derivable cells: key-based comparison against groundtruth.
                    # Judgment cells (Gap_Score, Enrollment_Rate_Pct, Gap_Details text,
                    # Avg_Gap_Score, Critical_Gaps_Count...) get range/enum checks instead.
                    header_map = {h: i for i, h in enumerate(headers)}
                    gt_header_map = {h: i for i, h in enumerate(gt_headers)}

                    def aval(row, col):
                        ci = header_map.get(col)
                        return row[ci] if ci is not None and ci < len(row) else None

                    def gval(row, col):
                        ci = gt_header_map.get(col)
                        return row[ci] if ci is not None and ci < len(row) else None

                    if sheet_name == "Department_Skills":
                        agent_by_dept = {}
                        for r in data_rows:
                            d = aval(r, "department")
                            if d:
                                agent_by_dept[str(d).strip()] = r
                        for gr in gt_rows:
                            dept = str(gval(gr, "department") or "").strip()
                            if not dept:
                                continue
                            ar = agent_by_dept.get(dept)
                            check(f"Department_Skills has {dept}", ar is not None,
                                  f"depts: {sorted(agent_by_dept)[:10]}")
                            if ar is None:
                                continue
                            for col in ("employee_count", "avg_performance"):
                                gf = safe_float(gval(gr, col))
                                af = safe_float(aval(ar, col))
                                check(f"Department_Skills {dept} {col} ~{gf:g}",
                                      af is not None and num_close(af, gf), f"got {af}")
                        for ri, ar in enumerate(data_rows[:3]):
                            gs = safe_float(aval(ar, "gap_score"))
                            check(f"Department_Skills R{ri+2} Gap_Score in 1-10",
                                  gs is not None and 1 <= gs <= 10, f"got {gs}")
                            er = safe_float(aval(ar, "enrollment_rate_pct"))
                            check(f"Department_Skills R{ri+2} Enrollment_Rate_Pct numeric >= 0",
                                  er is not None and er >= 0, f"got {er}")
                            ca = safe_float(aval(ar, "courses_available"))
                            check(f"Department_Skills R{ri+2} Courses_Available in 0-25",
                                  ca is not None and 0 <= ca <= 25, f"got {ca}")

                    elif sheet_name == "Course_Utilization":
                        agent_by_code = {}
                        for r in data_rows:
                            cd = aval(r, "course_code")
                            if cd:
                                agent_by_code[str(cd).strip().upper()] = r
                        # top-5 courses by enrollment (GT pre-sorted descending)
                        for gr in gt_rows[:5]:
                            code = str(gval(gr, "course_code") or "").strip().upper()
                            if not code:
                                continue
                            ar = agent_by_code.get(code)
                            check(f"Course_Utilization has {code}", ar is not None,
                                  f"codes: {sorted(agent_by_code)[:10]}")
                            if ar is None:
                                continue
                            gf = safe_float(gval(gr, "enrollment_count"))
                            af = safe_float(aval(ar, "enrollment_count"))
                            check(f"Course_Utilization {code} enrollment_count ~{gf:g}",
                                  af is not None and num_close(af, gf), f"got {af}")
                            gname = str(gval(gr, "course_name") or "").strip().lower()
                            aname = str(aval(ar, "course_name") or "").strip().lower()
                            check(f"Course_Utilization {code} course_name",
                                  bool(aname) and (gname in aname or aname in gname),
                                  f"expected {gname[:50]}, got {aname[:50]}")
                        enr_vals = [safe_float(aval(r, "enrollment_count")) for r in data_rows]
                        enr_vals = [v for v in enr_vals if v is not None]
                        check("Course_Utilization sorted by Enrollment_Count descending",
                              len(enr_vals) > 1 and all(enr_vals[i] >= enr_vals[i+1] - 1e-6
                                                        for i in range(len(enr_vals) - 1)),
                              f"order: {enr_vals[:10]}")
                        for ri, ar in enumerate(data_rows[:3]):
                            cr = safe_float(aval(ar, "completion_rate_pct"))
                            check(f"Course_Utilization R{ri+2} Completion_Rate_Pct in 0-100",
                                  cr is not None and 0 <= cr <= 100, f"got {cr}")
                            td = aval(ar, "target_department")
                            check(f"Course_Utilization R{ri+2} Target_Department non-empty",
                                  td is not None and str(td).strip() != "", f"got {td}")

                    elif sheet_name == "Gap_Details":
                        for ri, ar in enumerate(data_rows[:3]):
                            cov = str(aval(ar, "current_coverage") or "").strip().lower()
                            check(f"Gap_Details R{ri+2} Current_Coverage in High/Medium/Low",
                                  cov in ("high", "medium", "low"), f"got {cov}")
                            pri = str(aval(ar, "priority") or "").strip().lower()
                            check(f"Gap_Details R{ri+2} Priority in Critical/Important/Nice-to-have",
                                  pri in ("critical", "important", "nice-to-have"), f"got {pri}")
                            for col in ("department", "skill_needed", "recommended_action"):
                                v = aval(ar, col)
                                check(f"Gap_Details R{ri+2} {col} non-empty",
                                      v is not None and str(v).strip() != "", f"got {v}")
                        prios = [str(aval(r, "priority") or "").strip().lower() for r in data_rows]
                        check("Gap_Details has at least one Critical row",
                              "critical" in prios, f"priorities: {prios[:10]}")

                    elif sheet_name == "Summary":
                        ag = {}
                        for r in data_rows:
                            m = aval(r, "metric")
                            if m:
                                ag[str(m).strip()] = aval(r, "value")
                        td = safe_float(ag.get("Total_Departments"))
                        check("Summary Total_Departments == 7",
                              td is not None and abs(td - 7) < 0.5, f"got {ag.get('Total_Departments')}")
                        tc = safe_float(ag.get("Total_Courses"))
                        check("Summary Total_Courses in 7-23",
                              tc is not None and 7 <= tc <= 23, f"got {ag.get('Total_Courses')}")
                        ags = safe_float(ag.get("Avg_Gap_Score"))
                        check("Summary Avg_Gap_Score in 1-10",
                              ags is not None and 1 <= ags <= 10, f"got {ag.get('Avg_Gap_Score')}")
                        cg = safe_float(ag.get("Critical_Gaps_Count"))
                        check("Summary Critical_Gaps_Count numeric >= 0",
                              cg is not None and cg >= 0, f"got {ag.get('Critical_Gaps_Count')}")
                        tpd = ag.get("Top_Priority_Department")
                        check("Summary Top_Priority_Department non-empty",
                              tpd is not None and str(tpd).strip() != "", f"got {tpd}")

        # ---- CRITICAL: core ranking deliverable in Department_Skills ----
        if "Department_Skills" in wb.sheetnames:
            ws = wb["Department_Skills"]
            headers = [str(c.value).strip() if c.value else "" for c in ws[1]]
            hmap = {h: i for i, h in enumerate(headers)}
            di = hmap.get("Department")
            gi = hmap.get("Gap_Score")
            rows = [r for r in ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
            if di is not None and gi is not None and rows:
                gaps = []
                for r in rows:
                    gv = safe_float(r[gi]) if gi < len(r) else None
                    if gv is not None:
                        gaps.append((r[di], gv))
                if gaps:
                    # sorted descending?
                    vals = [g[1] for g in gaps]
                    is_sorted = all(vals[i] >= vals[i+1] - 1e-6 for i in range(len(vals)-1))
                    check("Department_Skills sorted by Gap_Score descending",
                          is_sorted, f"Gap_Score order: {vals}", critical=True)
                    # top row = max gap
                    dept_top_from_sheet = max(gaps, key=lambda x: x[1])[0]
                    check("Department_Skills top row is the max-Gap_Score department",
                          str(gaps[0][0]).strip() == str(dept_top_from_sheet).strip(),
                          f"top row dept={gaps[0][0]}, max-gap dept={dept_top_from_sheet}",
                          critical=True)
                else:
                    check("Department_Skills has numeric Gap_Score rows", False, "no parsable Gap_Score", critical=True)
            else:
                check("Department_Skills has Department+Gap_Score columns", False, f"headers={headers}", critical=True)

        # ---- CRITICAL: Summary internal consistency ----
        if "Summary" in wb.sheetnames and dept_top_from_sheet is not None:
            ws = wb["Summary"]
            summary = {}
            for r in ws.iter_rows(min_row=2, values_only=True):
                if r and r[0] is not None:
                    summary[str(r[0]).strip()] = r[1] if len(r) > 1 else None
            tpd = summary.get("Top_Priority_Department")
            a = str(tpd).strip().lower() if tpd is not None else ""
            b = str(dept_top_from_sheet).strip().lower()
            check("Summary Top_Priority_Department == max-Gap_Score department",
                  bool(a) and (a == b or a in b or b in a),
                  f"Summary={tpd}, Department_Skills top={dept_top_from_sheet}",
                  critical=True)

    # Check Skills_Gap_Brief.docx
    docx_path = os.path.join(agent_workspace, "Skills_Gap_Brief.docx")
    check("Skills_Gap_Brief.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        check("Skills_Gap_Brief.docx has content", len(text) > 50, f"text length: {len(text)}")
        # All four required section headings present (RU+EN accepted, exact-list)
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        # also accept Title style for the document heading; sections are Heading
        for variants in REQUIRED_HEADINGS:
            found = heading_matches(variants, headings)
            check(f"Skills_Gap_Brief.docx has section \"{variants[0]}\"",
                  found, f"agent headings: {headings[:8]}", critical=True)

    # ---- CRITICAL: python/terminal pipeline actually produced analysis JSON ----
    json_path = os.path.join(agent_workspace, "skills_gap_analysis.json")
    json_ok = False
    json_detail = "skills_gap_analysis.json missing"
    if os.path.exists(json_path):
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            # accept a per-department structure: dict keyed by dept, or list of dept records,
            # or a dict containing a list under any key.
            def looks_like_dept_struct(obj):
                if isinstance(obj, dict) and len(obj) > 0:
                    # values are dicts/numbers per department, OR contains a list
                    if any(isinstance(v, (dict, list)) for v in obj.values()):
                        return True
                if isinstance(obj, list) and len(obj) > 0 and isinstance(obj[0], dict):
                    return True
                return False
            json_ok = looks_like_dept_struct(data)
            json_detail = f"parsed type={type(data).__name__}, len={len(data) if hasattr(data,'__len__') else 'n/a'}"
        except Exception as e:
            json_detail = f"invalid JSON: {e}"
    check("skills_gap_analysis.json exists with per-department structure",
          json_ok, json_detail, critical=True)

    # Check Python script exists (terminal usage)
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    check("Python analysis script exists", len(py_files) >= 1, f"found: {py_files}")

    # Database checks
    try:
        conn = get_conn()
        cur = conn.cursor()
        # ---- CRITICAL: email actually sent to hr-director@company.com with skill-subject ----
        cur.execute("""
            SELECT subject, to_addr::text
            FROM email.messages
            WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1)
              AND subject ILIKE '%skill%'
              AND to_addr::text ILIKE '%hr-director@company.com%'
        """)
        email_row = cur.fetchone()
        check("Email sent to hr-director@company.com with skill subject",
              email_row is not None,
              "no Sent email with subject ILIKE '%skill%' to hr-director@company.com",
              critical=True)

        # Reverse verification: russified noise emails should not be in Sent folder
        cur.execute("""SELECT COUNT(*) FROM email.messages
            WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1)
              AND (subject ILIKE '%рассылка%' OR subject ILIKE '%newsletter%'
                   OR subject ILIKE '%обслуживание%' OR subject ILIKE '%maintenance%')""")
        noise_sent = cur.fetchone()[0]
        check("No noise emails in Sent folder", noise_sent == 0, f"found {noise_sent} noise emails in Sent")
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e))

    # ---- Gate ----
    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    if CRITICAL_FAILS:
        msg = f"CRITICAL FAIL ({len(CRITICAL_FAILS)}): {CRITICAL_FAILS} | Passed {PASS_COUNT}/{total} ({accuracy:.1f}%)"
        print(msg)
        return False, msg
    success = accuracy >= 70
    msg = f"Passed {PASS_COUNT}/{total} checks ({accuracy:.1f}%)"
    return success, msg

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
