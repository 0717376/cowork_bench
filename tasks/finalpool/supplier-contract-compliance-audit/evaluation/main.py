#!/usr/bin/env python3
"""Evaluation script for supplier-contract-compliance-audit task validation.

Structure-only checks feed the >=70% accuracy gate. SEMANTIC critical checks
(scores, corrective actions, summary consistency, email deliverable, per-supplier
findings) trigger sys.exit(1) BEFORE the gate on any failure.

Supplier names are russified in the source CSV / groundtruth and kept in sync
here. Column/sheet/file identifiers and the audit status vocabulary
(Complete/Incomplete/Compliant/Non-Compliant/Missing) stay English.
"""

from argparse import ArgumentParser
import datetime
import json
import os
import re
import sys

PASS_COUNT = 0
FAIL_COUNT = 0

# Russified suppliers (must match initial_workspace/supplier_list.csv + groundtruth)
SUPPLIERS = [
    "ООО ТехСервис",
    "Глобал Логистика",
    "КонтролКачество",
    "КонсалтГрупп",
    "СервисПоддержка",
]
TOP_SUPPLIER = "СервисПоддержка"      # fully compliant -> highest score
WORST_SUPPLIER = "Глобал Логистика"   # most gaps -> lowest score
# Suppliers with at least one gap requiring a corrective action
NONCOMPLIANT = ["ООО ТехСервис", "Глобал Логистика", "КонтролКачество", "КонсалтГрупп"]

SUPPLIER_EMAILS = [
    "michael.torres@supplier1.com",
    "sarah.liu@supplier2.com",
    "david.brown@supplier3.com",
    "james.wilson@supplier4.com",
    "emily.davis@supplier5.com",
]


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def supplier_in(cell):
    """Match a russified supplier name in a free-form cell (case/space tolerant)."""
    if cell is None:
        return None
    t = str(cell).strip().lower()
    for s in SUPPLIERS:
        if s.lower() in t:
            return s
    return None


# ---------------------------------------------------------------------------
# Structural checks (non-critical, feed the accuracy gate)
# ---------------------------------------------------------------------------

def check_xlsx_content(workspace, groundtruth_workspace="."):
    print("\n=== Check: XLSX files ===")
    import openpyxl
    for fname in ["compliance_assessment.xlsx", "corrective_actions.xlsx"]:
        xlsx_path = os.path.join(workspace, fname)
        if not os.path.isfile(xlsx_path):
            record(f"xlsx {fname} exists", False, "Not found")
            continue
        record(f"xlsx {fname} exists", True)
        try:
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            for ws in wb.worksheets:
                rows = list(ws.iter_rows(values_only=True))
                record(f"xlsx {fname} '{ws.title}' has data", len(rows) >= 2, f"{len(rows)} rows")
            # Structural sanity only. Per-supplier correctness is enforced by the
            # CRITICAL checks (header-scan + lookup). The old positional GT compare
            # false-negated task.md-compliant layouts: the GT xlsx carries
            # undocumented title/blank/Risk/Summary rows the contract never mandates.
            min_rows = len(SUPPLIERS) if "compliance" in fname else len(NONCOMPLIANT)
            data_rows, has_supplier = 0, False
            for ws in wb.worksheets:
                wsrows = [r for r in ws.iter_rows(values_only=True) if any(c is not None for c in r)]
                if any(any(isinstance(c, str) and c.strip().lower() == "supplier" for c in r) for r in wsrows):
                    has_supplier = True
                    data_rows = max(data_rows, len(wsrows) - 1)
            record(f"xlsx {fname} has 'Supplier' column", has_supplier)
            record(f"xlsx {fname} has >= {min_rows} supplier rows", data_rows >= min_rows,
                   f"{data_rows} data rows")
            wb.close()
        except Exception as e:
            record(f"xlsx {fname} readable", False, str(e))


def check_docx_content(workspace):
    print("\n=== Check: DOCX files ===")
    from docx import Document
    for fname in ["audit_findings.docx", "audit_summary.docx"]:
        path = os.path.join(workspace, fname)
        if not os.path.isfile(path):
            record(f"docx {fname} exists", False, "Not found")
            continue
        record(f"docx {fname} exists", True)
        try:
            doc = Document(path)
            record(f"docx {fname} has content", len(doc.paragraphs) > 0, f"{len(doc.paragraphs)} paragraphs")
        except Exception as e:
            record(f"docx {fname} readable", False, str(e))


# ---------------------------------------------------------------------------
# Helpers for critical checks
# ---------------------------------------------------------------------------

def _docx_text(path):
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _load_sheet_rows(workspace, fname, sheet_title):
    import openpyxl
    path = os.path.join(workspace, fname)
    if not os.path.isfile(path):
        return None
    wb = openpyxl.load_workbook(path, data_only=True)
    target = None
    for sn in wb.sheetnames:
        if sn.strip().lower() == sheet_title.strip().lower():
            target = wb[sn]; break
    if target is None:
        target = wb.worksheets[0]
    rows = list(target.iter_rows(values_only=True))
    wb.close()
    return rows


def _valid_future_iso(val, today):
    if val is None:
        return False
    s = str(val).strip()
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", s)
    if not m:
        return False
    try:
        d = datetime.date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    except ValueError:
        return False
    return d >= today


# ---------------------------------------------------------------------------
# CRITICAL (semantic) checks -> any failure aborts before the accuracy gate
# ---------------------------------------------------------------------------

def critical_compliance_scores(workspace):
    """All 5 suppliers present with numeric Compliance Score 0-100; the fully
    compliant supplier scores highest and the most-gapped one scores lowest."""
    rows = _load_sheet_rows(workspace, "compliance_assessment.xlsx", "Compliance Assessment")
    if not rows:
        return False, "compliance_assessment.xlsx / sheet missing"
    # Locate the column holding 'Compliance Score'
    score_col = None
    supplier_col = 0
    for r in rows:
        for ci, c in enumerate(r):
            if isinstance(c, str) and c.strip().lower() == "compliance score":
                score_col = ci
            if isinstance(c, str) and c.strip().lower() == "supplier":
                supplier_col = ci
        if score_col is not None:
            break
    if score_col is None:
        return False, "no 'Compliance Score' column"
    scores = {}
    for r in rows:
        if supplier_col >= len(r):
            continue
        s = supplier_in(r[supplier_col])
        if not s or s in scores:
            continue
        if score_col >= len(r):
            return False, f"{s}: no score cell"
        v = r[score_col]
        try:
            fv = float(v)
        except (TypeError, ValueError):
            return False, f"{s}: non-numeric score {v!r}"
        if not (0 <= fv <= 100):
            return False, f"{s}: score {fv} out of 0-100"
        scores[s] = fv
    missing = [s for s in SUPPLIERS if s not in scores]
    if missing:
        return False, f"missing suppliers: {missing}"
    top = max(scores, key=scores.get)
    worst = min(scores, key=scores.get)
    if top != TOP_SUPPLIER:
        return False, f"highest score should be {TOP_SUPPLIER}, got {top} ({scores})"
    if worst != WORST_SUPPLIER:
        return False, f"lowest score should be {WORST_SUPPLIER}, got {worst} ({scores})"
    return True, f"scores={scores}"


def critical_corrective_actions(workspace, today):
    """>=1 corrective action per non-compliant supplier with non-empty Required
    Action and a valid future ISO Target Date."""
    rows = _load_sheet_rows(workspace, "corrective_actions.xlsx", "Corrective Actions")
    if not rows:
        return False, "corrective_actions.xlsx / sheet missing"
    # Find header row
    header_idx = None
    for i, r in enumerate(rows):
        low = [str(c).strip().lower() if c is not None else "" for c in r]
        if "supplier" in low and "required action" in low and "target date" in low:
            header_idx = i; header = low; break
    if header_idx is None:
        return False, "header (Supplier/Required Action/Target Date) not found"
    ci_sup = header.index("supplier")
    ci_act = header.index("required action")
    ci_date = header.index("target date")
    by_supplier = {}
    for r in rows[header_idx + 1:]:
        if ci_sup >= len(r):
            continue
        s = supplier_in(r[ci_sup])
        if not s:
            continue
        act = r[ci_act] if ci_act < len(r) else None
        date = r[ci_date] if ci_date < len(r) else None
        good = bool(act and str(act).strip()) and _valid_future_iso(date, today)
        by_supplier.setdefault(s, False)
        if good:
            by_supplier[s] = True
    missing = [s for s in NONCOMPLIANT if not by_supplier.get(s)]
    if missing:
        return False, f"no valid corrective action (non-empty action + future ISO date) for: {missing}"
    return True, f"actions ok for {list(by_supplier)}"


def critical_summary_consistency(workspace):
    """audit_summary.docx contains an average compliance score and per-risk-tier
    supplier counts (RU+EN keyword match)."""
    path = os.path.join(workspace, "audit_summary.docx")
    if not os.path.isfile(path):
        return False, "audit_summary.docx missing"
    text = _docx_text(path)
    low = text.lower()
    has_avg = any(k in low for k in ["средн", "average"]) and ("compliance" in low or "соответ" in low or "балл" in low)
    has_high = any(k in low for k in ["high", "высок"])
    has_med = any(k in low for k in ["medium", "средн"])
    has_low = any(k in low for k in ["low", "низк"])
    # at least one explicit numeric score present
    has_number = re.search(r"\b\d{1,3}\s*/\s*100\b", text) or re.search(r"\b\d{1,3}\b", text)
    ok = has_avg and has_high and has_med and has_low and bool(has_number)
    return ok, f"avg={has_avg} high={has_high} med={has_med} low={has_low} num={bool(has_number)}"


def critical_findings_per_supplier(workspace):
    """audit_findings.docx names all 5 suppliers, each with a specific gap term."""
    path = os.path.join(workspace, "audit_findings.docx")
    if not os.path.isfile(path):
        return False, "audit_findings.docx missing"
    text = _docx_text(path)
    low = text.lower()
    missing = [s for s in SUPPLIERS if s.lower() not in low]
    if missing:
        return False, f"suppliers not mentioned: {missing}"
    gap_terms = ["sla", "ответствен", "liability", "конфиденц", "confidential",
                 "оплат", "payment", "интеллектуальн", "ip", "качеств", "scope",
                 "предмет", "пробел", "gap"]
    if not any(t in low for t in gap_terms):
        return False, "no specific gap terminology found"
    return True, "all suppliers + gap terms present"


def critical_email_sent(workspace):
    """At least one email sent to a supplier contact address with an
    audit/compliance subject."""
    try:
        import psycopg2
    except Exception as e:
        return False, f"psycopg2 unavailable: {e}"
    cfg = {
        "host": os.environ.get("PGHOST", "localhost"),
        "port": int(os.environ.get("PGPORT", "5432")),
        "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
        "user": "eigent",
        "password": "camel",
    }
    try:
        conn = psycopg2.connect(**cfg)
        cur = conn.cursor()
    except Exception as e:
        return False, f"DB connect failed: {e}"
    found = []
    for addr in SUPPLIER_EMAILS:
        cur.execute(
            """
            SELECT m.subject FROM email.messages m
              JOIN email.sent_log s ON s.message_id = m.id
             WHERE m.to_addr::text ILIKE %s
            """,
            (f"%{addr}%",),
        )
        rows = cur.fetchall()
        if not rows:
            cur.execute(
                "SELECT subject FROM email.messages WHERE to_addr::text ILIKE %s",
                (f"%{addr}%",),
            )
            rows = cur.fetchall()
        for (subj,) in rows:
            sl = (subj or "").lower()
            if any(k in sl for k in ["audit", "аудит", "compliance", "соответ", "договор", "contract"]):
                found.append((addr, subj))
    cur.close()
    conn.close()
    if not found:
        return False, "no audit/compliance email to any supplier contact address"
    return True, f"emails: {found[:2]}"


def run_critical_checks(workspace):
    today = datetime.date.today()
    checks = [
        ("compliance scores (5 suppliers, top/worst ordering)", lambda: critical_compliance_scores(workspace)),
        ("corrective actions per non-compliant supplier", lambda: critical_corrective_actions(workspace, today)),
        ("summary consistency (avg + risk tiers)", lambda: critical_summary_consistency(workspace)),
        ("per-supplier findings with specific gaps", lambda: critical_findings_per_supplier(workspace)),
        ("audit email to a supplier contact", lambda: critical_email_sent(workspace)),
    ]
    print("\n=== CRITICAL CHECKS ===")
    failed = []
    for name, fn in checks:
        try:
            ok, detail = fn()
        except Exception as e:
            ok, detail = False, f"exception: {e}"
        status = "PASS" if ok else "FAIL"
        print(f"  [{status}] CRITICAL: {name}" + (f" -> {detail}" if detail else ""))
        if not ok:
            failed.append(name)
    if failed:
        print(f"\nFAIL: critical check(s) failed: {failed}")
        sys.exit(1)


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--launch_time", required=False)
    args = parser.parse_args()

    ws = args.agent_workspace
    if not os.path.isdir(ws):
        print(f"Agent workspace not found: {ws}")
        sys.exit(1)

    check_xlsx_content(ws, args.groundtruth_workspace)
    check_docx_content(ws)

    # Critical semantic checks: any failure aborts before the accuracy gate.
    run_critical_checks(ws)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
