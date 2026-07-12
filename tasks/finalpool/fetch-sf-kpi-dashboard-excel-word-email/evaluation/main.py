"""Evaluation script for fetch-sf-kpi-dashboard-excel-word-email."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []

def critical(name, condition, detail=""):
    """Semantic gate. Any failure => hard FAIL via sys.exit(1) before accuracy gate."""
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [CRITICAL PASS] {name}")
    else:
        FAIL_COUNT += 1
        CRITICAL_FAILS.append(name)
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [CRITICAL FAIL] {name}: {detail_str}")

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def sheet_rows(wb, sheet_name):
    """Return list of dicts keyed by header (lowercased, stripped)."""
    if sheet_name not in wb.sheetnames:
        return []
    ws = wb[sheet_name]
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    headers = [str(h).strip().lower() if h is not None else "" for h in rows[0]]
    out = []
    for r in rows[1:]:
        if all(c is None for c in r):
            continue
        out.append({headers[i]: r[i] for i in range(len(headers)) if i < len(r)})
    return out

def find_row(rows, key_col, key_val):
    """Find a row where key_col value (case/space-insensitive) equals key_val."""
    kv = str(key_val).strip().lower()
    for row in rows:
        cell = row.get(key_col)
        if cell is not None and str(cell).strip().lower() == kv:
            return row
    return None

def run_critical_checks(agent_workspace):
    """Semantic-substance gate: correct reconciliation values from source + key rules."""
    excel_path = os.path.join(agent_workspace, "KPI_Dashboard_Report.xlsx")
    if not os.path.exists(excel_path):
        critical("KPI_Dashboard_Report.xlsx exists (critical)", False, "missing workbook")
        return
    try:
        wb = openpyxl.load_workbook(excel_path)
    except Exception as e:
        critical("KPI_Dashboard_Report.xlsx readable (critical)", False, str(e))
        return

    # --- KPI_Scorecard: Total_Revenue reconciliation (SALES_DW vs target 2.5M) ---
    # Source: SUM(SALES_DW.ORDERS.TOTAL_AMOUNT) = 3,048,998.33; achievement = 3.049M/2.5M = 122.0% => Met.
    sc = sheet_rows(wb, "KPI_Scorecard")
    tr = find_row(sc, "kpi_name", "Total_Revenue")
    if tr is None:
        critical("KPI_Scorecard has Total_Revenue row", False, f"kpi_names: {[r.get('kpi_name') for r in sc][:10]}")
    else:
        act = safe_float(tr.get("actual"))
        ach = safe_float(tr.get("achievement_pct"))
        status = str(tr.get("status") or "").strip().lower()
        critical("Total_Revenue Actual ~3.05M", act is not None and num_close(act, 3048998.33, rel_tol=0.05, abs_tol=1), f"actual={act}")
        critical("Total_Revenue Achievement_Pct ~122", ach is not None and abs(ach - 122.0) <= 3, f"ach={ach}")
        critical("Total_Revenue Status == Met", status == "met", f"status={status}")

    # --- Executive_Summary: the central deliverable verdict ---
    # From source-derived scorecard: Met=3 (Total_Revenue, Avg_Order_Value, Ticket_Resolution_Rate),
    # Near=0, Missed=3 (Support_Response_Time, Customer_Satisfaction, Employee_Satisfaction).
    # Top_Risk_Area = worst achievement = Support_Response_Time (26.6%, avg 15.03h vs 4h target).
    es = sheet_rows(wb, "Executive_Summary")
    def es_val(metric):
        row = find_row(es, "metric", metric)
        return row.get("value") if row else None
    critical("Executive_Summary KPIs_Missed == 3", safe_float(es_val("KPIs_Missed")) == 3, f"got {es_val('KPIs_Missed')}")
    critical("Executive_Summary KPIs_Near == 0", safe_float(es_val("KPIs_Near")) == 0, f"got {es_val('KPIs_Near')}")
    critical("Executive_Summary KPIs_Met == 3", safe_float(es_val("KPIs_Met")) == 3, f"got {es_val('KPIs_Met')}")
    tra = str(es_val("Top_Risk_Area") or "").strip().lower()
    critical("Executive_Summary Top_Risk_Area == Support_Response_Time", "support_response_time" in tra, f"got {tra}")

    # --- Revenue_Detail: russified region join survived + worst region identified ---
    # Source: ORDERS JOIN CUSTOMERS, Северная Америка actual 606,318.35 vs target 800,000 => variance -24.2%.
    rd = sheet_rows(wb, "Revenue_Detail")
    na = find_row(rd, "region", "Северная Америка")
    if na is None:
        critical("Revenue_Detail has region 'Северная Америка'", False,
                 f"regions: {[r.get('region') for r in rd][:10]}")
    else:
        var = safe_float(na.get("variance_pct"))
        critical("Revenue_Detail 'Северная Америка' Variance_Pct ~ -24.2",
                 var is not None and abs(var - (-24.2)) <= 3, f"variance={var}")

    # --- Support_Detail: SLA classification rule applied correctly ---
    # Source: AVG(TICKETS.RESPONSE_TIME_HOURS) by PRIORITY vs SLA_POLICIES targets.
    # High 6.23h > 4h => No; Medium 12.28h > 8h => No. (No 'Critical' priority tickets exist in seed.)
    sd = sheet_rows(wb, "Support_Detail")
    high_p = find_row(sd, "priority", "High")
    med_p = find_row(sd, "priority", "Medium")
    critical("Support_Detail 'High' Met_SLA == No",
             high_p is not None and str(high_p.get("met_sla") or "").strip().lower() == "no",
             f"got {high_p.get('met_sla') if high_p else None}")
    critical("Support_Detail 'Medium' Met_SLA == No",
             med_p is not None and str(med_p.get("met_sla") or "").strip().lower() == "no",
             f"got {med_p.get('met_sla') if med_p else None}")

def run_critical_email_check():
    """Required communication deliverable: KPI email to executive-team with substantive body."""
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("""SELECT subject, to_addr, body_text FROM email.messages
                       WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1)
                         AND subject ILIKE '%kpi%'""")
        row = cur.fetchone()
        conn.close()
    except Exception as e:
        critical("KPI email deliverable", False, str(e))
        return
    if row is None:
        critical("KPI email in Sent folder", False, "no KPI email found")
        return
    subject, to_addr, body = row
    to_s = str(to_addr or "")
    critical("KPI email recipient == executive-team@company.com",
             "executive-team@company.com" in to_s, f"to_addr={to_s[:120]}")
    body_l = (str(body or "")).lower()
    # Overall achievement = mean of per-KPI achievement_pct = 100.5 (source-derived).
    critical("KPI email body mentions overall achievement (~100.5)",
             ("100.5" in body_l) or ("100,5" in body_l) or ("100" in body_l and ("общ" in body_l or "overall" in body_l or "достижен" in body_l)),
             f"body[:120]={body_l[:120]}")
    # Missed KPIs (achievement < 90%): Support_Response_Time, Customer_Satisfaction, Employee_Satisfaction.
    critical("KPI email body lists at least one missed KPI",
             any(k in body_l for k in ["support_response_time", "customer_satisfaction", "employee_satisfaction"]),
             f"body[:160]={body_l[:160]}")

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILS
    PASS_COUNT = 0
    FAIL_COUNT = 0
    CRITICAL_FAILS = []

    # === CRITICAL semantic checks (hard gate before accuracy) ===
    print("--- CRITICAL CHECKS ---")
    run_critical_checks(agent_workspace)
    run_critical_email_check()
    if CRITICAL_FAILS:
        print(f"CRITICAL FAILURE: {len(CRITICAL_FAILS)} critical check(s) failed: {CRITICAL_FAILS}")
        sys.exit(1)
    print("--- STRUCTURAL CHECKS ---")

    # Check KPI_Dashboard_Report.xlsx
    excel_path = os.path.join(agent_workspace, "KPI_Dashboard_Report.xlsx")
    check("KPI_Dashboard_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "KPI_Dashboard_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Check row count
                    gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                    data_rows = list(ws.iter_rows(min_row=2, values_only=True))
                    min_rows = max(1, len(gt_rows) - 2)
                    check(f"{sheet_name} has >= {min_rows} data rows", len(data_rows) >= min_rows, f"got {len(data_rows)}")

                    # Cell value comparison against groundtruth
                    header_map = {h: i for i, h in enumerate(headers)}
                    gt_header_map = {h: i for i, h in enumerate(gt_headers)}
                    for ri in range(min(3, len(gt_rows), len(data_rows))):
                        gt_row = gt_rows[ri]
                        agent_row = data_rows[ri]
                        for ci, gt_h in enumerate(gt_headers):
                            if not gt_h or ci >= len(gt_row):
                                continue
                            gv = gt_row[ci]
                            agent_ci = header_map.get(gt_h)
                            if agent_ci is None or agent_ci >= len(agent_row):
                                continue
                            av = agent_row[agent_ci]
                            gf = safe_float(gv)
                            af = safe_float(av)
                            if gf is not None and af is not None:
                                tol = max(0.5, abs(gf) * 0.15)
                                check(f"{sheet_name} R{ri+2} {gt_h} ~{gf:.1f}",
                                      abs(gf - af) <= tol, f"got {af}")
                            elif gv is not None and av is not None:
                                gs = str(gv).strip().lower()
                                avs = str(av).strip().lower()
                                if gs:
                                    check(f"{sheet_name} R{ri+2} {gt_h} text",
                                          gs == avs or gs in avs or avs in gs,
                                          f"expected {gs[:50]}, got {avs[:50]}")

    # Check KPI_Review.docx
    docx_path = os.path.join(agent_workspace, "KPI_Review.docx")
    check("KPI_Review.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        check("KPI_Review.docx has content", len(text) > 50, f"text length: {len(text)}")
        # Check headings: accept RU or EN variant for each required section.
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        required_headings = [
            ["executive summary", "сводка для руководства"],
            ["revenue performance", "эффективность выручки"],
            ["support center kpis", "kpi центра поддержки"],
            ["action items for missed targets", "меры по непрошедшим целям"],
        ]
        for variants in required_headings:
            found = any(any(v in h or h in v for v in variants) for h in headings)
            check(f"KPI_Review.docx has heading \"{variants[0]}\"", found, f"agent headings: {headings[:5]}")

    # Check Python script exists (terminal usage)
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    check("Python analysis script exists", len(py_files) >= 1, f"found: {py_files}")

    # Database checks
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("SELECT subject, to_addr FROM email.messages WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1) AND subject ILIKE '%kpi%'")
        email_row = cur.fetchone()
        check("Email with correct subject sent", email_row is not None, "no matching email found")
        if email_row:
            check("Email has recipient", email_row[1] is not None, f"to_addr: {email_row[1]}")
        # Reverse verification: noise emails should not be in Sent folder
        cur.execute("SELECT COUNT(*) FROM email.messages WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1) AND subject ILIKE '%newsletter%'")
        noise_sent = cur.fetchone()[0]
        check("No noise emails in Sent folder", noise_sent == 0, f"found {noise_sent} noise emails in Sent")
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e))

    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()