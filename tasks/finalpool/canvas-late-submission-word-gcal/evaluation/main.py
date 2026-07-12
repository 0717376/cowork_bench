"""Evaluation for canvas-late-submission-word-gcal.

Критические проверки (CRITICAL_CHECKS): провал любой => общий FAIL независимо
от accuracy. В остальном PASS требует accuracy >= 70%.

Данные о курсах и поздних сдачах поступают с ЖИВОГО сервера Canvas (внешний
глобальный seed, не из preprocess этой задачи) и читаются "честно": в eval НЕ
зашиты волатильные числовые проценты. Проверяются только стабильные, server-seeded
названия курсов и структура (наличие столбцов с числовыми значениями).
"""
import argparse
import os
import re
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

# Top 3 курса по доле поздних сдач (стабильные, server-seeded названия Canvas).
TOP3_COURSES = [
    "Креативные вычисления и культура (Весна 2014)",
    "Креативные вычисления и культура (Осень 2014)",
    "Проектирование на основе данных (Весна 2013)",
]
# Подстроки для устойчивого поиска названий top-3 в тексте (русифицированные
# названия Canvas, в нижнем регистре; сезон/год сохраняют различимость).
TOP3_SUBSTR = [
    "креативные вычисления и культура (весна 2014)",
    "креативные вычисления и культура (осень 2014)",
    "проектирование на основе данных (весна 2013)",
]

# Обязательные столбцы таблицы (language-neutral identifiers).
REQUIRED_COLUMNS = ["Course_Name", "Total_Submissions", "Late_Submissions", "Late_Rate_Pct"]

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Критические (семантические) проверки: провал любой => общий FAIL.
CRITICAL_CHECKS = {
    "Word doc table has all four required columns",
    "Word doc names all three top late-rate courses",
    "At least 3 'Late Submission Review' events on 2026-04-01 at 10:00/11:00/14:00 (1h each)",
    "Email to registrar@university.edu, subject correct, names a top course",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    docx_path = os.path.join(agent_workspace, "Late_Submission_Report.docx")
    if not os.path.isfile(docx_path):
        check("Late_Submission_Report.docx exists", False, f"Not found: {docx_path}")
        # Без файла остальные проверки бессмысленны.
        check("Word doc table has all four required columns", False, "no docx")
        check("Word doc names all three top late-rate courses", False, "no docx")
        return
    check("Late_Submission_Report.docx exists", True)

    try:
        from docx import Document
    except ImportError:
        check("Word doc has content", os.path.getsize(docx_path) > 1000,
              f"Size: {os.path.getsize(docx_path)}")
        check("Word doc table has all four required columns", False, "python-docx unavailable")
        check("Word doc names all three top late-rate courses", False, "python-docx unavailable")
        return

    try:
        doc = Document(docx_path)
    except Exception as e:
        check("Word doc readable", False, str(e))
        check("Word doc table has all four required columns", False, "unreadable")
        check("Word doc names all three top late-rate courses", False, "unreadable")
        return

    all_text = " ".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += " " + cell.text
    lower = all_text.lower()

    check("Word doc has meaningful content (>= 100 chars)",
          len(all_text.strip()) >= 100,
          f"Content length: {len(all_text)}")
    check("Word doc contains 'late submission'",
          "late submission" in lower or "late_submission" in lower,
          f"Sample: {lower[:200]}")

    # --- Структурная проверка таблицы: 4 обязательных столбца + числовые значения ---
    header_ok = False
    numeric_late_rate = False
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        header_norm = [h.lower() for h in header]
        if all(col.lower() in header_norm for col in REQUIRED_COLUMNS):
            header_ok = True
            # Проверяем, что в столбце Late_Rate_Pct есть числовые значения.
            try:
                idx = header_norm.index("late_rate_pct")
                for row in table.rows[1:]:
                    val = row.cells[idx].text.strip()
                    if re.search(r"\d+(\.\d+)?", val):
                        numeric_late_rate = True
                        break
            except (ValueError, IndexError):
                pass
            break
    check("Word doc table has all four required columns",
          header_ok,
          f"Required: {REQUIRED_COLUMNS}")
    check("Word doc table Late_Rate_Pct has numeric values",
          numeric_late_rate,
          "No numeric value found in Late_Rate_Pct column")

    # --- Названы все три top-3 курса (стабильные английские названия Canvas) ---
    named = [s for s in TOP3_SUBSTR if s in lower]
    check("Word doc names all three top late-rate courses",
          len(named) >= 3,
          f"Found {len(named)}/3: {named}")

    # --- Секция рекомендаций (RU + EN ключевые слова) ---
    rec_keywords = ["recommend", "policy", "intervention",
                    "рекоменд", "политик", "вмешательств", "меры", "мер"]
    check("Word doc has recommendations section",
          any(k in lower for k in rec_keywords),
          f"Sample: {lower[:200]}")


def check_calendar():
    print("\n=== Checking Google Calendar ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT summary, description, start_datetime, end_datetime
            FROM gcal.events
            WHERE LOWER(summary) LIKE '%%late submission review%%'
               OR LOWER(summary) LIKE '%%late submission%%'
        """)
        events = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        check("Calendar check", False, str(e))
        check("At least 3 'Late Submission Review' events on 2026-04-01 at 10:00/11:00/14:00 (1h each)",
              False, str(e))
        return

    check("At least 3 'Late Submission Review' events scheduled",
          len(events) >= 3,
          f"Found {len(events)} events")

    # События на 2026-04-01 в нужные часы.
    target_hours = {10, 11, 14}
    on_date = []
    for summ, desc, start, end in events:
        if start is None:
            continue
        if str(start)[:10] != "2026-04-01":
            continue
        hour = getattr(start, "hour", None)
        if hour is None:
            m = re.search(r"\b(\d{1,2}):", str(start))
            hour = int(m.group(1)) if m else None
        # Длительность ~1 час.
        dur_ok = False
        if end is not None:
            try:
                dur = (end - start).total_seconds()
                dur_ok = 0 < dur <= 3 * 3600
            except Exception:
                dur_ok = True  # не блокируем, если тип не datetime
        else:
            dur_ok = True
        on_date.append((hour, dur_ok, summ, desc))

    hours_present = {h for (h, _d, _s, _de) in on_date if h is not None}
    durations_ok = all(d for (_h, d, _s, _de) in on_date) if on_date else False

    crit_time_ok = (
        len(on_date) >= 3
        and target_hours.issubset(hours_present)
        and durations_ok
    )

    # Каждое событие должно упоминать один из top-3 курсов в summary/description.
    courses_covered = 0
    for h, _d, summ, desc in on_date:
        text = (str(summ or "") + " " + str(desc or "")).lower()
        if any(s in text for s in TOP3_SUBSTR):
            courses_covered += 1
    crit_course_ok = courses_covered >= 3

    check("At least 3 'Late Submission Review' events on 2026-04-01 at 10:00/11:00/14:00 (1h each)",
          crit_time_ok and crit_course_ok,
          f"on_date={len(on_date)}, hours={sorted(hours_present)}, "
          f"durations_ok={durations_ok}, courses_covered={courses_covered}")

    check("Events scheduled on 2026-04-01",
          len(on_date) >= 3,
          f"{len(on_date)} events on 2026-04-01 out of {len(events)} total")


def check_email():
    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT id, subject, to_addr, body_text
            FROM email.messages
            WHERE to_addr::text ILIKE '%%registrar@university.edu%%'
        """)
        emails = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        check("Email check", False, str(e))
        check("Email to registrar@university.edu, subject correct, names a top course",
              False, str(e))
        return

    check("Email sent to registrar@university.edu", len(emails) >= 1,
          "No matching email found")

    if not emails:
        check("Email to registrar@university.edu, subject correct, names a top course",
              False, "no email")
        return

    # Берём письмо с правильной темой, если есть.
    chosen = None
    for em in emails:
        subj = str(em[1] or "").lower()
        if "late submission analysis report" in subj:
            chosen = em
            break
    if chosen is None:
        chosen = emails[0]

    subject = str(chosen[1] or "").lower()
    body = str(chosen[3] or "")
    body_lower = body.lower()

    subject_ok = "late submission analysis report" in subject
    check("Email subject is 'Late Submission Analysis Report'",
          subject_ok,
          f"Subject: {chosen[1]}")

    check("Email body has content", len(body) > 30,
          f"Body length: {len(body)}")

    names_course = any(s in body_lower for s in TOP3_SUBSTR)

    check("Email to registrar@university.edu, subject correct, names a top course",
          subject_ok and names_course,
          f"subject_ok={subject_ok}, names_course={names_course}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_calendar()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total else 0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== Результат: {PASS_COUNT}/{total} пройдено ({accuracy:.1f}%) ===")
    if critical_failed:
        print(f"КРИТИЧЕСКИЕ ПРОВАЛЫ: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")
        sys.exit(1)

    if accuracy >= 70:
        print("Все условия выполнены (нет критических провалов, accuracy >= 70%).")
        sys.exit(0)
    else:
        print(f"accuracy {accuracy:.1f}% < 70%")
        sys.exit(1)


if __name__ == "__main__":
    main()
