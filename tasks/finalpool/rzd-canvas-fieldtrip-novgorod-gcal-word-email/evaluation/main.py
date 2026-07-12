"""Evaluation for rzd-canvas-fieldtrip-novgorod-gcal-word-email.

Строгий evaluator. Проверки (12):
  Word (7):
    1.  Field_Trip_Notice.docx существует в workspace
    2.  Документ читается
    3.  >= 4 параграфов со стилем Heading
    4.  Содержит хотя бы один из номеров поездов "туда" (818А или 820А)
    5.  Содержит хотя бы один из номеров поездов "обратно" (819А или 821А)
    6.  Содержит города "Москва" и "Великий Новгород"
    7.  Содержит сумму туда-обратно: 5000₽ (либо 5 000)
  Calendar (2):
    8.  >= 2 событий поездки (не "обычное занятие")
    9.  Хотя бы одно событие на 12.03.2026, второе на 15.03.2026
  Canvas (1):
   10.  >= 1 announcement в курсе 9991 с упоминанием Новгорода/экскурсии
  Email (2):
   11.  >= 1 письмо отправлено на students@university.ru
   12.  Тело письма упоминает оба номера поездов и стоимость
"""
import json
import os
import sys
import unicodedata
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

TRAINS_OUTBOUND = ("818", "820")          # 818А / 820А
TRAINS_RETURN   = ("819", "821")          # 819А / 821А
PRICE_TOTAL_RUB = 5000
COURSE_ID       = 9991

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def normalize(s: str) -> str:
    """Lowercase + collapse common cyrillic/latin lookalikes (А/A, С/C, etc.)
    so '818А' and '818A' compare equal."""
    s = s.lower()
    s = unicodedata.normalize("NFKD", s)
    table = str.maketrans({"а": "a", "о": "o", "е": "e", "р": "p",
                           "с": "c", "у": "y", "к": "k", "х": "x"})
    return s.translate(table)


def contains_price(text: str, value: int) -> bool:
    """Match 5000 / 5 000 / 5,000 / 5000₽."""
    patterns = [str(value), f"{value // 1000} {value % 1000:03d}",
                f"{value // 1000},{value % 1000:03d}"]
    return any(p in text for p in patterns)


# ---------------------------------------------------------------------------
# Word doc checks
# ---------------------------------------------------------------------------

def check_word_doc(agent_workspace):
    print("\n=== Word: Field_Trip_Notice.docx ===")
    import glob
    candidates = glob.glob(os.path.join(agent_workspace, "*.docx"))
    if not candidates:
        record("Field_Trip_Notice.docx существует", False, f"docx not found in {agent_workspace}")
        return
    doc_path = next((c for c in candidates if any(
        kw in os.path.basename(c).lower()
        for kw in ("field", "trip", "notice", "новгород", "уведомлен", "экскурс")
    )), candidates[0])
    record("Field_Trip_Notice.docx существует", True)

    try:
        import docx
        doc = docx.Document(doc_path)
    except Exception as e:
        record("Документ читается", False, str(e))
        return
    record("Документ читается", True)

    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    record("Не менее 4 заголовков (Heading)", len(headings) >= 4,
           f"headings: {[p.text for p in headings]}")

    full_text = " ".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text += " " + cell.text
    norm = normalize(full_text)

    has_out = any(tn in norm for tn in TRAINS_OUTBOUND)
    record(f"Номер поезда туда ({' или '.join(TRAINS_OUTBOUND)}А)",
           has_out, f"text sample: {full_text[:200]!r}")

    has_ret = any(tn in norm for tn in TRAINS_RETURN)
    record(f"Номер поезда обратно ({' или '.join(TRAINS_RETURN)}А)",
           has_ret, f"text sample: {full_text[:200]!r}")

    lower = full_text.lower()
    has_msk = "москв" in lower
    has_nvg = "новгород" in lower
    record("Города: Москва и Великий Новгород", has_msk and has_nvg,
           f"msk={has_msk}, nvg={has_nvg}")

    has_total = contains_price(full_text, PRICE_TOTAL_RUB)
    record(f"Сумма туда-обратно {PRICE_TOTAL_RUB}₽", has_total,
           f"text sample: {full_text[:200]!r}")


# ---------------------------------------------------------------------------
# Calendar checks
# ---------------------------------------------------------------------------

def check_gcal():
    print("\n=== Calendar: события поездки 12 и 15 марта ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime, end_datetime
          FROM gcal.events
         WHERE (
                  (start_datetime >= '2026-03-12' AND start_datetime < '2026-03-13')
               OR (start_datetime >= '2026-03-15' AND start_datetime < '2026-03-16')
              )
           AND summary NOT ILIKE '%обычное занятие%'
           AND summary NOT ILIKE '%regular class%'
         ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    record("Не менее 2 событий поездки в календаре", len(events) >= 2,
           f"events: {[(e[0], str(e[1])) for e in events]}")

    has_outbound = any(e[1].day == 12 for e in events)
    has_return   = any(e[1].day == 15 for e in events)
    record("Есть событие на 12.03 и на 15.03",
           has_outbound and has_return,
           f"out={has_outbound}, ret={has_return}")


# ---------------------------------------------------------------------------
# Canvas announcement checks
# ---------------------------------------------------------------------------

def check_canvas_announcement():
    print("\n=== Canvas: объявление в курсе 'Изучение культурного наследия' ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        cur.execute(
            """
            SELECT COUNT(*) FROM canvas.announcements
             WHERE course_id = %s
               AND (
                       title   ILIKE '%%новгород%%'
                    OR title   ILIKE '%%экскурс%%'
                    OR title   ILIKE '%%поездк%%'
                    OR title   ILIKE '%%trip%%'
                    OR title   ILIKE '%%field%%'
                    OR message ILIKE '%%новгород%%'
                    OR message ILIKE '%%экскурс%%'
                    OR message ILIKE '%%поездк%%'
                   )
            """,
            (COURSE_ID,),
        )
        cnt = cur.fetchone()[0]
        record("Объявление об экскурсии создано в курсе 9991",
               cnt >= 1, f"matching announcements: {cnt}")
    except Exception as e:
        record("Canvas announcement check", False, str(e))
    cur.close()
    conn.close()


# ---------------------------------------------------------------------------
# Email checks
# ---------------------------------------------------------------------------

def check_email():
    print("\n=== Email: students@university.ru ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute(
        """
        SELECT m.subject, m.body_text
          FROM email.messages m
          JOIN email.sent_log s ON s.message_id = m.id
         WHERE m.to_addr::text ILIKE %s
           AND m.from_addr  NOT ILIKE %s
        """,
        ("%students@university.ru%", "%students@university.ru%"),
    )
    rows = cur.fetchall()
    if not rows:
        cur.execute(
            """
            SELECT subject, body_text FROM email.messages
             WHERE to_addr::text ILIKE %s
               AND from_addr  NOT ILIKE %s
            """,
            ("%students@university.ru%", "%students@university.ru%"),
        )
        rows = cur.fetchall()
    cur.close()
    conn.close()

    record("Письмо отправлено на students@university.ru", len(rows) >= 1,
           f"matched rows: {len(rows)}")
    if not rows:
        record("(skipped — нет письма)", False)
        return

    body_combined = " ".join((b or "") for _s, b in rows)
    body_norm = normalize(body_combined)

    has_out = any(tn in body_norm for tn in TRAINS_OUTBOUND)
    has_ret = any(tn in body_norm for tn in TRAINS_RETURN)
    has_price = contains_price(body_combined, PRICE_TOTAL_RUB)
    record("В теле письма оба поезда и сумма",
           has_out and has_ret and has_price,
           f"out={has_out}, ret={has_ret}, price={has_price}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_gcal()
    check_canvas_announcement()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\nИтого: {PASS_COUNT}/{total} проверок пройдено ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
