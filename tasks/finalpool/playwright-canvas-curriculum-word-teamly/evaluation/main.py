"""
Evaluation script for playwright-canvas-curriculum-word task (teamly variant).

Checks:
1. Word document Accreditation_Compliance_Report.docx with course evaluations
2. Teamly knowledge-base page "Course Compliance Tracker" with all 22 courses
3. Department notification emails for departments with non-compliant courses

Compliance is RECOMPUTED LIVE from Canvas (canvas.assignments / canvas.quizzes)
against the standards (min 8 assignments, min 3 quizzes) — nothing is hardcoded,
because the Canvas LMS data is live/volatile.

CRITICAL_CHECKS (semantic): any failure => overall FAIL regardless of accuracy.
Pass threshold otherwise: accuracy >= 70%.
"""

import argparse
import os
import re
import sys

import psycopg2

try:
    from docx import Document
except ImportError:
    Document = None

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

# Accreditation standards published on standards.html (the agent reads these honestly).
MIN_ASSIGNMENTS = 8
MIN_QUIZZES = 3
FOLLOW_UP_DATE = "2026-04-15"

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Critical checks: any failure => overall FAIL regardless of accuracy.
CRITICAL_CHECKS = {
    "Word: compliant/non-compliant summary counts match live recompute",
    "Teamly tracker page exists with all 22 courses",
    "Teamly tracker: Overall_Status correct for every course",
    "Teamly tracker: Follow_Up_Date present for non-compliant / empty for compliant",
    "Emails: one review email per non-compliant department (correct from/to/subject)",
}


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


# ── Live compliance recompute from Canvas ────────────────────────────────────

def short_name(course_name):
    """Course name without the '(Term Year)' suffix, lowercased and trimmed."""
    return course_name.split("(")[0].strip().lower()


def department_of(course_name):
    """Department key = subject-area prefix of the course name (before '&' / '(').

    standards.html: 'Each department is identified by the subject area prefix in
    the course name'. Used for grouping courses into departments, e.g.
    'Прикладная аналитика и алгоритмы' -> 'applied analytics',
    'Проектирование на основе данных' -> 'data-driven design'.
    """
    prefix = course_name.split("(")[0].split("&")[0].strip()
    return prefix.lower()


def dept_match_token(dept_key):
    """A distinctive token to look for in free text / email subjects when the
    agent may abbreviate the department name (e.g. 'Applied Analytics' or just
    'Applied'). We require the first word of the prefix, which is unique across
    all departments (Applied / Biochemistry / Creative / Data-Driven /
    Environmental / Foundations / Global ...)."""
    return dept_key.split()[0]


def get_live_courses():
    """Recompute per-course assignment/quiz counts and compliance live from Canvas."""
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT c.name,
               (SELECT COUNT(*) FROM canvas.assignments a WHERE a.course_id = c.id) AS asgn,
               (SELECT COUNT(*) FROM canvas.quizzes q WHERE q.course_id = c.id) AS quiz
        FROM canvas.courses c
        WHERE c.id <> 9991
        ORDER BY c.name
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()
    courses = []
    for name, asgn, quiz in rows:
        a_ok = asgn >= MIN_ASSIGNMENTS
        q_ok = quiz >= MIN_QUIZZES
        compliant = a_ok and q_ok
        courses.append({
            "name": name,
            "short": short_name(name),
            "dept": department_of(name),
            "assignments": asgn,
            "quizzes": quiz,
            "a_ok": a_ok,
            "q_ok": q_ok,
            "compliant": compliant,
        })
    return courses


# ── Word document ────────────────────────────────────────────────────────────

def check_word(agent_workspace, courses):
    """Check Word compliance report against live-recomputed totals."""
    print("\n=== Checking Word Document ===")

    doc_path = os.path.join(agent_workspace, "Accreditation_Compliance_Report.docx")
    if not os.path.isfile(doc_path):
        record("Accreditation_Compliance_Report.docx exists", False, doc_path)
        record("Word: compliant/non-compliant summary counts match live recompute",
               False, "Document missing")
        return

    record("Accreditation_Compliance_Report.docx exists", True)

    if Document is None:
        record("python-docx available", False, "Cannot import docx")
        return

    try:
        doc = Document(doc_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        record("Word: compliant/non-compliant summary counts match live recompute",
               False, str(e))
        return

    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Table cells too: answers laid out in a docx table are legitimate.
    full_text += "\n" + "\n".join(
        c.text for t in doc.tables for r in t.rows for c in r.cells)
    low = full_text.lower()

    record(
        "Doc mentions accreditation/compliance",
        any(kw in low for kw in
            ["accreditation", "compliance", "аккредитац", "соответств"]),
    )

    # Course-name coverage (course names are seeded English substrings).
    course_shorts = sorted({c["short"] for c in courses})
    mentioned = sum(1 for s in course_shorts if s in low)
    record(
        "Doc covers most departments/courses",
        mentioned >= max(5, int(0.7 * len(course_shorts))),
        f"Found {mentioned}/{len(course_shorts)} distinct course names",
    )

    # Compliance status indicators (RU + EN).
    has_status = (
        ("compliant" in low and "non-compliant" in low)
        or ("pass" in low and "fail" in low)
        or ("соответств" in low and "не соответств" in low)
    )
    record("Doc has compliance status indicators", has_status)

    # Summary section keywords (RU + EN).
    has_summary = any(kw in low for kw in
                      ["summary", "total", "rate", "overall",
                       "итог", "всего", "процент"])
    record("Doc has summary section", has_summary)

    # CRITICAL: the summary compliant/non-compliant counts must equal live totals.
    compliant_expected = sum(1 for c in courses if c["compliant"])
    non_compliant_expected = len(courses) - compliant_expected
    numbers = set(re.findall(r"\d+", full_text))
    counts_match = (str(compliant_expected) in numbers
                    and str(non_compliant_expected) in numbers)
    record(
        "Word: compliant/non-compliant summary counts match live recompute",
        counts_match,
        f"Expected compliant={compliant_expected}, "
        f"non-compliant={non_compliant_expected}; numbers in doc: "
        f"{sorted(numbers, key=lambda x: int(x))[:30]}",
    )


# ── Teamly tracker page ──────────────────────────────────────────────────────

def _date_present(chunk):
    """Whether the follow-up date appears in the given text chunk (RU + EN forms)."""
    return (FOLLOW_UP_DATE in chunk
            or "15 апреля 2026" in chunk
            or "april 15, 2026" in chunk
            or "15.04.2026" in chunk
            or "04/15/2026" in chunk)


def check_teamly(courses):
    """Check the Teamly 'Course Compliance Tracker' page."""
    print("\n=== Checking Teamly Tracker Page ===")

    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT id, title, COALESCE(body, '')
            FROM teamly.pages
            WHERE title ILIKE '%%course compliance tracker%%'
               OR title ILIKE '%%compliance tracker%%'
               OR (title ILIKE '%%course%%' AND title ILIKE '%%compliance%%')
        """)
        pages = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Teamly DB accessible", False, str(e))
        record("Teamly tracker page exists with all 22 courses", False, str(e))
        record("Teamly tracker: Overall_Status correct for every course", False, str(e))
        record("Teamly tracker: Follow_Up_Date present for non-compliant / empty for compliant",
               False, str(e))
        return

    if not pages:
        record("Teamly tracker page exists with all 22 courses", False,
               "No 'Course Compliance Tracker' page found")
        record("Teamly tracker: Overall_Status correct for every course", False,
               "No page")
        record("Teamly tracker: Follow_Up_Date present for non-compliant / empty for compliant",
               False, "No page")
        return

    body = "\n".join(str(b) for _, _, b in pages)
    body_low = body.lower()

    # Column header presence (preserved English property names).
    for col in ["Course_Name", "Department", "Assignment_Count", "Quiz_Count",
                "Assignments_Compliant", "Quizzes_Compliant",
                "Overall_Status", "Follow_Up_Date"]:
        record(f"Tracker mentions column {col}", col.lower() in body_low,
               f"'{col}' not found in page body")

    # Locate each course's row by the FULL course name (with the term suffix),
    # which disambiguates same-department courses that differ only by term.
    def row_index(course):
        full = course["name"].lower()
        idx = body_low.find(full)
        if idx >= 0:
            return idx
        # Fallback: anchor on department prefix co-located with the term, e.g.
        # 'biochemistry ... fall 2014' across a table row.
        term = ""
        if "(" in course["name"] and ")" in course["name"]:
            term = course["name"].split("(", 1)[1].split(")", 1)[0].strip().lower()
        if not term:
            return -1
        start = 0
        while True:
            di = body_low.find(course["short"], start)
            if di < 0:
                return -1
            if term in body_low[di: di + 400]:
                return di
            start = di + 1

    # Map each course to its row start, then bound each row's chunk by the next
    # row's start so date/status lookups do not bleed into adjacent rows.
    course_idx = {c["name"]: row_index(c) for c in courses}
    sorted_starts = sorted(i for i in course_idx.values() if i >= 0)

    def row_chunk(course):
        idx = course_idx[course["name"]]
        if idx < 0:
            return None
        end = len(body_low)
        for s in sorted_starts:
            if s > idx:
                end = s
                break
        return body_low[idx: min(end, idx + 600)]

    # CRITICAL: all 22 course rows present (located individually).
    missing = [c["name"] for c in courses if course_idx[c["name"]] < 0]
    record(
        "Teamly tracker page exists with all 22 courses",
        len(missing) == 0 and len(courses) == 22,
        f"{len(courses)} live courses; missing rows: {missing[:6]}",
    )

    # CRITICAL: Overall_Status correct for every course.
    def status_ok(course):
        chunk = row_chunk(course)
        if chunk is None:
            return False
        if course["compliant"]:
            has_compliant = ("compliant" in chunk or "соответствует" in chunk)
            has_non = ("non-compliant" in chunk or "non compliant" in chunk
                       or "не соответствует" in chunk)
            return has_compliant and not has_non
        else:
            return ("non-compliant" in chunk or "non compliant" in chunk
                    or "не соответствует" in chunk)

    status_correct = sum(1 for c in courses if status_ok(c))
    record(
        "Teamly tracker: Overall_Status correct for every course",
        status_correct == len(courses),
        f"{status_correct}/{len(courses)} courses have correct Overall_Status",
    )

    # CRITICAL: Follow_Up_Date present for non-compliant, empty for compliant.
    def followup_ok(course):
        chunk = row_chunk(course)
        if chunk is None:
            return False
        present = _date_present(chunk)
        if course["compliant"]:
            return not present   # compliant -> empty follow-up
        return present           # non-compliant -> 2026-04-15

    followup_correct = sum(1 for c in courses if followup_ok(c))
    record(
        "Teamly tracker: Follow_Up_Date present for non-compliant / empty for compliant",
        followup_correct == len(courses),
        f"{followup_correct}/{len(courses)} courses have correct Follow_Up_Date",
    )


# ── Emails ───────────────────────────────────────────────────────────────────

def check_emails(courses):
    """Check department notification emails against the live non-compliant depts."""
    print("\n=== Checking Emails ===")

    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute(
            "SELECT subject, from_addr, to_addr, body_text FROM email.messages"
        )
        emails = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Email DB accessible", False, str(e))
        record("Emails: one review email per non-compliant department (correct from/to/subject)",
               False, str(e))
        return

    # Departments with >=1 non-compliant course (live).
    non_compliant_depts = sorted({c["dept"] for c in courses if not c["compliant"]})

    # Candidate review emails: an accreditation/review subject.
    review_emails = []
    for subject, from_addr, to_addr, body_text in emails:
        subj_low = (subject or "").lower()
        if any(kw in subj_low for kw in
               ["accreditation review", "accreditation", "compliance", "review",
                "аккредитац"]):
            review_emails.append({
                "subject": subject or "",
                "from": (from_addr or "").lower(),
                "to": str(to_addr or "").lower(),
                "body": body_text or "",
            })

    record(
        "Accreditation review emails sent",
        len(review_emails) >= len(non_compliant_depts),
        f"Found {len(review_emails)} review emails for "
        f"{len(non_compliant_depts)} non-compliant departments",
    )

    # Departments mentioned anywhere in review emails (RU agent may write RU body).
    combined = " ".join((e["subject"] + " " + e["body"]).lower() for e in review_emails)
    depts_mentioned = sum(1 for d in non_compliant_depts
                          if dept_match_token(d) in combined)
    record(
        "Non-compliant departments mentioned in emails",
        depts_mentioned >= max(1, int(0.7 * len(non_compliant_depts))),
        f"{depts_mentioned}/{len(non_compliant_depts)} departments mentioned",
    )

    # Body mentions compliance-failure details (RU + EN).
    has_details = any(kw in combined for kw in [
        "assignment", "quiz", "fail", "non-compliant", "below", "minimum",
        "задани", "тест", "квиз", "не соответств", "ниже", "минимум",
    ])
    record("Email body mentions compliance failures", has_details)

    # CRITICAL: exactly one correctly-addressed review email per non-compliant dept,
    # with from=accreditation@university.edu, to=department-review@university.edu,
    # and subject 'Accreditation Review: <Department> Courses'.
    all_depts_ok = True
    details = []
    for dept in non_compliant_depts:
        token = dept_match_token(dept)
        matches = []
        for e in review_emails:
            subj_low = e["subject"].lower()
            from_ok = "accreditation@university.edu" in e["from"]
            to_ok = "department-review@university.edu" in e["to"]
            subj_ok = ("accreditation review" in subj_low
                       and token in subj_low
                       and "course" in subj_low)
            if from_ok and to_ok and subj_ok:
                matches.append(e)
        if len(matches) < 1:
            all_depts_ok = False
            details.append(f"{dept}: {len(matches)} matching emails (expected 1)")

    # No review email for fully-compliant departments.
    all_depts = {c["dept"] for c in courses}
    compliant_only_depts = sorted(all_depts - set(non_compliant_depts))
    for dept in compliant_only_depts:
        token = dept_match_token(dept)
        # Skip ambiguous tokens that also prefix a non-compliant dept.
        if any(token == dept_match_token(nd) for nd in non_compliant_depts):
            continue
        for e in review_emails:
            subj_low = e["subject"].lower()
            if ("accreditation review" in subj_low and token in subj_low
                    and "course" in subj_low
                    and "accreditation@university.edu" in e["from"]
                    and "department-review@university.edu" in e["to"]):
                all_depts_ok = False
                details.append(f"{dept}: unexpected review email (dept is fully compliant)")

    record(
        "Emails: one review email per non-compliant department (correct from/to/subject)",
        all_depts_ok,
        "; ".join(details) if details else "ok",
    )


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    try:
        courses = get_live_courses()
    except Exception as e:
        print(f"[eval] FATAL: could not recompute compliance from Canvas: {e}")
        sys.exit(1)

    compliant = sum(1 for c in courses if c["compliant"])
    print(f"[eval] Live recompute: {len(courses)} courses, "
          f"{compliant} compliant, {len(courses) - compliant} non-compliant.")

    check_word(args.agent_workspace, courses)
    check_teamly(courses)
    check_emails(courses)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}, Accuracy: {accuracy:.1f}%")

    # Critical checks: any failure => overall FAIL regardless of accuracy.
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    if critical_failed:
        print("  CRITICAL checks failed:")
        for n in critical_failed:
            print(f"    - {n}")
        print("  Overall: FAIL (critical)")
        sys.exit(1)

    overall = accuracy >= 70
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
