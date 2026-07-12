#!/usr/bin/env python3
"""Evaluation for financial-audit-cross-system-reconciliation.

The reconciliation is computed against local fixtures (authoritative_ledger.csv,
the InSales/store extract) vs q1_manual_tracking.xlsx (Manual Tracking), per the
PDF matching rules. Deterministic groundtruth:
  total orders = 15, matched = 11, discrepancies = 4
  discrepancy orders = WC-005 (amount), WC-008 (missing), WC-011 (date), WC-013 (amount)
  total amount (Manual Tracking) = 3263.98

CRITICAL semantic checks (any fail => sys.exit(1) before the accuracy gate):
  - Reconciliation per-order classification matches groundtruth (matched/discrepancy counts)
  - Summary total amount matches groundtruth within tolerance
  - Audit email sent to finance@company.com with the fixed subject literal
  - Calendar review-meeting event created
"""
import os
import argparse
import sys

import openpyxl

# ---------------------------------------------------------------------------
# Deterministic groundtruth (derived from the seeded local fixtures).
# ---------------------------------------------------------------------------
GT_TOTAL_ORDERS = 15
GT_MATCHED = 11
GT_DISCREPANCIES = 4
GT_DISCREPANCY_ORDERS = {"WC-005", "WC-008", "WC-011", "WC-013"}
GT_TOTAL_AMOUNT = 3263.98
GT_ALL_ORDERS = [f"WC-{i:03d}" for i in range(1, 16)]

DB = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

EMAIL_SUBJECT_LITERAL = "financial audit reconciliation"  # lowercase grep target
EMAIL_RECIPIENT = "finance@company.com"

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []


def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    try:
        return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)
    except (TypeError, ValueError):
        return False


def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").replace("₽", "").strip())
    except (ValueError, TypeError):
        return default


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRITICAL]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        if critical:
            CRITICAL_FAILED.append(name)
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL]{' [CRITICAL]' if critical else ''} {name}: {detail_str}")


def get_conn():
    import psycopg2
    return psycopg2.connect(**DB)


def load_reconciliation(ws):
    """Return {order_id: status_lower} from the Reconciliation sheet."""
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return {}
    hdr = [str(c).strip().lower() if c is not None else "" for c in rows[0]]
    try:
        oi = hdr.index("order id")
    except ValueError:
        oi = 0
    # match-status column: any header containing 'match' or 'status'
    si = None
    for i, h in enumerate(hdr):
        if "match" in h or "status" in h:
            si = i
            break
    if si is None:
        si = 1 if len(hdr) > 1 else 0
    out = {}
    for r in rows[1:]:
        if oi >= len(r) or r[oi] is None:
            continue
        oid = str(r[oi]).strip()
        st = str(r[si]).strip().lower() if si < len(r) and r[si] is not None else ""
        out[oid] = st
    return out


def status_is_discrepancy(s):
    s = (s or "").lower()
    return ("discrep" in s) or ("mismatch" in s) or ("расхожд" in s) or ("exception" in s)


def status_is_matched(s):
    s = (s or "").lower()
    return ("match" in s and "mismatch" not in s) or ("сверен" in s) or ("совпад" in s) or s == "ok"


def load_summary(ws):
    """Return {metric_lower: value} from the Summary sheet (Metric/Value rows)."""
    rows = list(ws.iter_rows(values_only=True))
    out = {}
    for r in rows:
        if not r or r[0] is None:
            continue
        key = str(r[0]).strip().lower()
        val = r[1] if len(r) > 1 else None
        out[key] = val
    return out


def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    if not agent_workspace or not os.path.isdir(agent_workspace):
        return False, f"Agent workspace not found: {agent_workspace}"

    # -------------------------------------------------------------------
    # 1. Audit reconciliation spreadsheet
    # -------------------------------------------------------------------
    excel_path = os.path.join(agent_workspace, "Audit_Reconciliation_Report.xlsx")
    check("Audit_Reconciliation_Report.xlsx exists", os.path.exists(excel_path),
          critical=True)

    recon = {}
    summary = {}
    if os.path.exists(excel_path):
        try:
            wb = openpyxl.load_workbook(excel_path, data_only=True)
        except Exception as e:
            wb = None
            check("Audit spreadsheet readable", False, str(e), critical=True)
        if wb is not None:
            # Locate sheets (tolerant of exact names)
            recon_sheet = None
            summary_sheet = None
            for name in wb.sheetnames:
                nl = name.strip().lower()
                if "reconcil" in nl or "сверк" in nl:
                    recon_sheet = name
                if "summary" in nl or "итог" in nl or "сводк" in nl:
                    summary_sheet = name
            if recon_sheet is None and wb.sheetnames:
                recon_sheet = wb.sheetnames[0]
            check("Reconciliation sheet present", recon_sheet is not None)
            check("Summary sheet present", summary_sheet is not None)

            if recon_sheet:
                recon = load_reconciliation(wb[recon_sheet])
                check("Reconciliation lists all 15 orders",
                      sum(1 for o in GT_ALL_ORDERS if o in recon) >= 15,
                      f"found {len(recon)} orders")
            if summary_sheet:
                summary = load_summary(wb[summary_sheet])

    # Per-order classification accuracy (non-critical granular checks)
    correct = 0
    for oid in GT_ALL_ORDERS:
        st = recon.get(oid, "")
        expected_disc = oid in GT_DISCREPANCY_ORDERS
        if expected_disc:
            ok = status_is_discrepancy(st)
        else:
            ok = status_is_matched(st)
        if ok:
            correct += 1
        check(f"{oid} classified {'discrepancy' if expected_disc else 'matched'}",
              ok, f"got status={st!r}")

    # CRITICAL: counts match groundtruth
    agent_disc = sum(1 for oid, st in recon.items()
                     if oid in GT_ALL_ORDERS and status_is_discrepancy(st))
    agent_matched = sum(1 for oid, st in recon.items()
                        if oid in GT_ALL_ORDERS and status_is_matched(st))
    check("CRITICAL matched count == 11", agent_matched == GT_MATCHED,
          f"got {agent_matched}", critical=True)
    check("CRITICAL discrepancy count == 4", agent_disc == GT_DISCREPANCIES,
          f"got {agent_disc}", critical=True)
    check("CRITICAL discrepancy orders exactly {WC-005,WC-008,WC-011,WC-013}",
          {oid for oid, st in recon.items()
           if oid in GT_ALL_ORDERS and status_is_discrepancy(st)} == GT_DISCREPANCY_ORDERS,
          f"got {sorted(o for o, s in recon.items() if status_is_discrepancy(s))}",
          critical=True)

    # Summary metrics
    def summ(*keys):
        for k in keys:
            for sk, sv in summary.items():
                if k in sk:
                    return safe_float(sv)
        return None

    total_orders = summ("total order")
    matched_v = summ("matched", "сверен")
    disc_v = summ("discrep", "расхожд")
    total_amount = summ("total amount", "сумм")

    check("Summary Total Orders == 15", total_orders is not None and num_close(total_orders, GT_TOTAL_ORDERS),
          f"got {total_orders}")
    check("Summary Matched == 11", matched_v is not None and num_close(matched_v, GT_MATCHED),
          f"got {matched_v}")
    check("Summary Discrepancies == 4", disc_v is not None and num_close(disc_v, GT_DISCREPANCIES),
          f"got {disc_v}")
    check("CRITICAL Summary Total Amount ~= 3263.98",
          total_amount is not None and num_close(total_amount, GT_TOTAL_AMOUNT),
          f"got {total_amount}", critical=True)

    # -------------------------------------------------------------------
    # 2. Audit Word report
    # -------------------------------------------------------------------
    docx_path = os.path.join(agent_workspace, "Audit_Report.docx")
    check("Audit_Report.docx exists", os.path.exists(docx_path), critical=True)
    if os.path.exists(docx_path):
        try:
            from docx import Document
            doc = Document(docx_path)
            text = " ".join(p.text for p in doc.paragraphs)
            check("Audit_Report.docx has content", len(text.strip()) > 50,
                  f"len={len(text)}")
            headings = [p.text.strip().lower() for p in doc.paragraphs
                        if p.style.name.startswith("Heading")]
            all_lower = text.lower()
            # Executive summary (RU+EN)
            exec_ok = any(("executive summary" in h or "резюме" in h or "сводка" in h)
                          for h in headings) or ("executive summary" in all_lower or "резюме" in all_lower)
            check("Audit_Report.docx has Executive Summary / Резюме section", exec_ok,
                  f"headings={headings[:6]}")
            # Discrepancies (RU+EN)
            disc_ok = any(("discrep" in h or "расхожд" in h or "exception" in h)
                          for h in headings) or ("discrep" in all_lower or "расхожд" in all_lower)
            check("Audit_Report.docx has Discrepancies / Расхождения section", disc_ok,
                  f"headings={headings[:6]}")
        except Exception as e:
            check("Audit_Report.docx readable", False, str(e))

    # -------------------------------------------------------------------
    # 3. Email + 4. Calendar (database-backed)
    # -------------------------------------------------------------------
    try:
        conn = get_conn()
        cur = conn.cursor()

        # Email: audit report sent to finance@company.com with the fixed subject.
        cur.execute(
            "SELECT subject, to_addr::text, body_text FROM email.messages "
            "WHERE LOWER(subject) LIKE %s",
            (f"%{EMAIL_SUBJECT_LITERAL}%",),
        )
        rows = cur.fetchall()
        check("CRITICAL Audit email with subject 'Financial Audit Reconciliation' sent",
              len(rows) >= 1, f"found {len(rows)} matching emails", critical=True)
        recipient_ok = any(EMAIL_RECIPIENT in (str(r[1] or "").lower()) for r in rows)
        check("CRITICAL Email sent to finance@company.com", recipient_ok,
              f"to_addrs={[r[1] for r in rows]}", critical=True)

        # Calendar: review meeting event created (RU+EN keywords).
        cur.execute(
            "SELECT count(*) FROM gcal.events WHERE "
            "LOWER(summary) LIKE '%сверк%' OR LOWER(summary) LIKE '%аудит%' "
            "OR LOWER(summary) LIKE '%audit%' OR LOWER(summary) LIKE '%review%' "
            "OR LOWER(summary) LIKE '%reconcil%'"
        )
        ev_cnt = cur.fetchone()[0]
        check("CRITICAL Audit review meeting event created", ev_cnt >= 1,
              f"found {ev_cnt} events", critical=True)

        cur.close()
        conn.close()
    except Exception as e:
        check("Database checks (email + calendar)", False, str(e), critical=True)

    # -------------------------------------------------------------------
    # Aggregate
    # -------------------------------------------------------------------
    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100.0) if total else 0.0
    critical_ok = len(CRITICAL_FAILED) == 0
    success = critical_ok and accuracy >= 70.0
    msg = f"Passed {PASS_COUNT}/{total} checks (accuracy {accuracy:.1f}%)"
    if not critical_ok:
        msg += f" | CRITICAL FAILED: {CRITICAL_FAILED}"
    return success, msg


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--launch_time", required=False, default="2026-05-30 10:00:00")
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace,
        args.groundtruth_workspace,
        args.launch_time,
        args.res_log_file,
    )
    print(message)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
