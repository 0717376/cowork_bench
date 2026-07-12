"""
Evaluation for arxiv-latex-analysis-teamly-word task.

Checks:
1. Paper_Analysis_Report.docx exists and has at least 4 sections
2. Word doc contains all 3 paper title keywords
3. Word doc mentions Scaling Laws and RLHF/InstructGPT and OPT
4. Word doc has a Comparative Analysis section
5. Teamly has 3 pages about the papers
6. Email sent to research_lead@university.edu

Critical checks (see CRITICAL_CHECKS): any failure there => overall FAIL
regardless of accuracy. Pass threshold otherwise: accuracy >= 70%.
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2
from docx import Document

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "cowork_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Critical checks: any failure => overall FAIL regardless of accuracy.
# The core deliverable is a Word report covering all 3 papers plus the email.
CRITICAL_CHECKS = {
    "Paper_Analysis_Report.docx exists",
    "Mentions Scaling Laws paper",
    "Mentions RLHF/InstructGPT paper",
    "Mentions OPT paper",
    "Email sent to research_lead@university.edu",
}


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Word Document Paper_Analysis_Report.docx ===")

    docx_path = os.path.join(agent_workspace, "Paper_Analysis_Report.docx")
    if not os.path.exists(docx_path):
        record("Paper_Analysis_Report.docx exists", False, f"Not found at {docx_path}")
        return
    record("Paper_Analysis_Report.docx exists", True)

    try:
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    # Get all text
    all_text = "\n".join(p.text for p in doc.paragraphs).lower()

    # Count headings/sections
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading") or
                (p.text.strip() and len(p.text.strip()) < 100 and
                 any(kw in p.text.lower() for kw in ["scaling", "instruct", "opt:", "analysis", "paper"]))]

    record("Word doc has at least 4 sections", len(headings) >= 4,
           f"Found {len(headings)} section-like headings")

    # Check paper keywords
    has_scaling = "scaling laws" in all_text or "scaling" in all_text
    has_rlhf = "rlhf" in all_text or "instructgpt" in all_text or "follow instructions" in all_text or "human feedback" in all_text
    has_opt = "opt" in all_text and ("open pre-trained" in all_text or "open-source" in all_text or "175b" in all_text)

    record("Mentions Scaling Laws paper", has_scaling, "No scaling laws content found")
    record("Mentions RLHF/InstructGPT paper", has_rlhf, "No RLHF/InstructGPT content found")
    record("Mentions OPT paper", has_opt, "No OPT content found")

    has_comparative = "comparative" in all_text or "comparison" in all_text or "connect" in all_text or "relate" in all_text
    record("Has comparative analysis section", has_comparative,
           "No comparative analysis content found")


def check_teamly():
    print("\n=== Check 2: Teamly Pages ===")

    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("SELECT title, COALESCE(body, '') FROM teamly.pages")
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Teamly pages for all 3 papers", False, f"Query failed: {e}")
        record("At least 3 paper pages created in Teamly", False, f"Query failed: {e}")
        return

    all_text = " ".join((str(t) + " " + str(b)).lower() for t, b in rows)
    has_scaling = "scaling" in all_text
    has_instruct = "instruct" in all_text or "human feedback" in all_text or "rlhf" in all_text
    has_opt = "opt" in all_text and ("pre-trained" in all_text or "transformer" in all_text)

    paper_pages_found = sum([has_scaling, has_instruct, has_opt])
    record("Teamly pages for all 3 papers", paper_pages_found >= 3,
           f"scaling={has_scaling}, instruct={has_instruct}, opt={has_opt}. Total pages: {len(rows)}")

    # Count agent-created paper pages by title markers (seed pages don't match).
    paper_markers = ["scaling", "instruct", "follow instructions", "opt:", "paper:"]
    paper_pages = [t for t, _ in rows
                   if any(m in (t or "").lower() for m in paper_markers)]
    record("At least 3 paper pages created in Teamly", len(paper_pages) >= 3,
           f"Found {len(paper_pages)} paper pages: {paper_pages[:5]}")


def check_email():
    print("\n=== Check 3: Email to research_lead@university.edu ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    messages = cur.fetchall()
    cur.close()
    conn.close()

    matching = None
    for subject, from_addr, to_addr, body_text in messages:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                to_str = " ".join(str(r).lower() for r in parsed) if isinstance(parsed, list) else str(to_addr).lower()
            except Exception:
                to_str = str(to_addr).lower()
        if "research_lead@university.edu" in to_str:
            matching = (subject, from_addr, to_addr, body_text)
            break

    record("Email sent to research_lead@university.edu", matching is not None,
           f"Messages found: {len(messages)}")

    if matching:
        subject, _, _, body_text = matching
        subject_lower = (subject or "").lower()
        body_lower = (body_text or "").lower()
        has_paper_ref = (
            "paper" in subject_lower or "analysis" in subject_lower or "report" in subject_lower or
            "scaling" in body_lower or "rlhf" in body_lower or "opt" in body_lower or
            "статья" in body_lower or "статей" in body_lower or "отчёт" in body_lower or "отчет" in body_lower
        )
        record("Email mentions paper analysis", has_paper_ref,
               f"Subject: {subject}, body preview: {body_text[:100] if body_text else ''}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_teamly()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    if critical_failed:
        print(f"CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_failed": critical_failed,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if critical_failed:
        print("FAIL (critical check failed)")
        sys.exit(1)
    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
