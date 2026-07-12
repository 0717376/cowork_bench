"""Evaluation for rzd-msk-spb-trip-notion-gcal-word.

Строгий evaluator — без substring-ловушек. Все артефакты локальные.

Проверки (16):
  Word (9):
    1-9. см. check_word_doc — заголовки, поезда, города, времена, цены, таблицы
  trip_notes.md (2):
   10. Файл существует в workspace
   11. Содержит оба номера поездов + дату
  Calendar (2):
   12. >= 2 новых события на 2026-03-10 (помимо "Встреча с клиентом")
   13. Хотя бы одно событие упоминает направление Москва ↔ СПб
  Email (3):
   14-16. отправлено на travel@consulting.ru с обоими поездами в теле и датой
"""
import json
import os
import re
import sys
import unicodedata
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

# Expected answer (Сапсан на 2026-03-10): outbound 752А (06:50→10:50), return 759А (19:25→23:25)
TRAIN_OUTBOUND = "752"
TRAIN_RETURN   = "759"
DEPART_OUT, ARRIVE_OUT = "06:50", "10:50"
DEPART_RET, ARRIVE_RET = "19:25", "23:25"
PRICE_ONE_WAY = 5500
PRICE_TOTAL   = 11000

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def normalize(s: str) -> str:
    """Lowercase + collapse common cyrillic/latin lookalikes (А/A, С/C, etc.)
    so '752А' and '752A' compare equal."""
    s = s.lower()
    s = unicodedata.normalize("NFKD", s)
    table = str.maketrans({"а": "a", "о": "o", "е": "e", "р": "p",
                           "с": "c", "у": "y", "к": "k", "х": "x"})
    return s.translate(table)


def contains_price(text: str, value: int) -> bool:
    """Match price like 5500, 5 500, 5,500, 5500₽."""
    patterns = [str(value), f"{value // 1000} {value % 1000:03d}",
                f"{value // 1000},{value % 1000:03d}"]
    return any(p in text for p in patterns)


# ---------------------------------------------------------------------------
# Word doc checks
# ---------------------------------------------------------------------------

def check_word_doc(agent_workspace):
    print("\n=== Word: Travel_Plan.docx ===")
    import glob
    candidates = glob.glob(os.path.join(agent_workspace, "*.docx"))
    if not candidates:
        record("Travel_Plan.docx существует", False, f"docx not found in {agent_workspace}")
        return
    # Prefer one named Travel_Plan; else take the first.
    doc_path = next((c for c in candidates if "travel" in c.lower() or "plan" in c.lower()
                     or "поездк" in c.lower()), candidates[0])
    record("Travel_Plan.docx существует", True)

    try:
        import docx
        doc = docx.Document(doc_path)
    except Exception as e:
        record("Документ читается", False, str(e))
        return
    record("Документ читается", True)

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    heading_text = " | ".join(headings).lower()
    required_headings = ["поездка туда", "поездка обратно", "сводк"]
    have_all = all(h in heading_text for h in required_headings)
    record("Заголовки 'Поездка туда / Поездка обратно / Сводка'",
           have_all and len(headings) >= 3,
           f"headings: {headings}")

    full_text = " ".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text += " " + cell.text
    norm = normalize(full_text)

    has_out = TRAIN_OUTBOUND in norm and "a" in norm.split(TRAIN_OUTBOUND, 1)[1][:2]
    has_ret = TRAIN_RETURN in norm and "a" in norm.split(TRAIN_RETURN, 1)[1][:2]
    record(f"Номера поездов {TRAIN_OUTBOUND}А и {TRAIN_RETURN}А",
           has_out and has_ret, f"out={has_out}, ret={has_ret}")

    lower = full_text.lower()
    has_msk = "москв" in lower
    has_spb = "петербург" in lower or "спб" in lower
    record("Города: Москва и Санкт-Петербург", has_msk and has_spb,
           f"msk={has_msk}, spb={has_spb}")

    has_dep_times = DEPART_OUT in full_text and DEPART_RET in full_text
    record(f"Время отправления {DEPART_OUT} и {DEPART_RET}", has_dep_times,
           f"text snippet: {full_text[:160]}")

    has_arr_times = ARRIVE_OUT in full_text and ARRIVE_RET in full_text
    record(f"Время прибытия {ARRIVE_OUT} и {ARRIVE_RET}", has_arr_times,
           f"text snippet: {full_text[:160]}")

    has_fare = contains_price(full_text, PRICE_ONE_WAY)
    has_total = contains_price(full_text, PRICE_TOTAL)
    record(f"Цены {PRICE_ONE_WAY}₽ и итого {PRICE_TOTAL}₽",
           has_fare and has_total, f"one_way={has_fare}, total={has_total}")

    record("Не менее 2 таблиц в документе", len(doc.tables) >= 2,
           f"tables={len(doc.tables)}")


# ---------------------------------------------------------------------------
# trip_notes.md check (replaces Notion knowledge-base page)
# ---------------------------------------------------------------------------

def check_trip_notes(agent_workspace):
    print("\n=== trip_notes.md: сводка по поездке ===")
    path = os.path.join(agent_workspace, "trip_notes.md")
    if not os.path.isfile(path):
        # Fallback: any .md file with trip/поездк in the name.
        import glob
        candidates = [p for p in glob.glob(os.path.join(agent_workspace, "*.md"))
                      if any(kw in os.path.basename(p).lower()
                             for kw in ("trip", "поездк", "notes"))]
        if not candidates:
            record("trip_notes.md существует в workspace", False,
                   f"file not found in {agent_workspace}")
            record("(skipped — no notes file)", False)
            return
        path = candidates[0]
    record("trip_notes.md существует в workspace", True)

    text = open(path, encoding="utf-8", errors="replace").read()
    norm = normalize(text)
    has_both_trains = TRAIN_OUTBOUND in norm and TRAIN_RETURN in norm
    has_date = ("10.03" in text or "10 марта" in text.lower()
                or "2026-03-10" in text)
    record("trip_notes.md упоминает оба поезда и дату поездки",
           has_both_trains and has_date,
           f"trains={has_both_trains}, date={has_date}")


# ---------------------------------------------------------------------------
# Calendar checks
# ---------------------------------------------------------------------------

def check_gcal():
    print("\n=== Calendar: 2 события на 2026-03-10 ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, description, start_datetime, end_datetime
          FROM gcal.events
         WHERE start_datetime >= '2026-03-10'
           AND start_datetime <  '2026-03-11'
           AND summary NOT ILIKE '%встреча с клиентом%'
           AND summary NOT ILIKE '%client meeting%'
         ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    record("Не менее 2 событий поездки в календаре", len(events) >= 2,
           f"events: {[e[0] for e in events]}")

    direction_ok = False
    for summary, descr, _s, _e in events:
        combined = f"{summary or ''} {descr or ''}".lower()
        if ("москв" in combined and ("петербург" in combined or "спб" in combined)):
            direction_ok = True
            break
    record("Событие упоминает направление Москва ↔ Санкт-Петербург",
           direction_ok, f"events: {[e[0] for e in events]}")


# ---------------------------------------------------------------------------
# Email checks
# ---------------------------------------------------------------------------

def check_email():
    print("\n=== Email: travel@consulting.ru ===")
    # email.messages is the source of truth (used by all sent_email tools).
    # email.sent_log is just a pointer table (id, message_id, sent_at), so
    # joining via that is the right way if we want "actually sent" filter.
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute(
        "SELECT m.subject, m.body_text "
        "  FROM email.messages m "
        "  JOIN email.sent_log s ON s.message_id = m.id "
        " WHERE m.to_addr::text ILIKE %s "
        "   AND m.from_addr NOT ILIKE %s ",
        ("%travel@consulting.ru%", "%travel@consulting.ru%"),
    )
    rows = cur.fetchall()
    if not rows:
        # Fallback: some email backends INSERT into messages without sent_log.
        cur.execute(
            "SELECT subject, body_text FROM email.messages "
            " WHERE to_addr::text ILIKE %s "
            "   AND from_addr NOT ILIKE %s ",
            ("%travel@consulting.ru%", "%travel@consulting.ru%"),
        )
        rows = cur.fetchall()
    cur.close()
    conn.close()

    record("Письмо отправлено на travel@consulting.ru", len(rows) >= 1,
           f"matched rows: {len(rows)}")
    if not rows:
        for _ in range(3):
            record("(skipped because no email found)", False)
        return

    body_combined = " ".join((b or "") for _s, b in rows)
    subj_combined = " ".join((s or "") for s, _b in rows).lower()
    body_norm = normalize(body_combined)

    has_both_trains = TRAIN_OUTBOUND in body_norm and TRAIN_RETURN in body_norm
    record(f"В теле письма упомянуты оба поезда ({TRAIN_OUTBOUND}А и {TRAIN_RETURN}А)",
           has_both_trains, f"body sample: {body_combined[:200]!r}")

    subj_ok = any(root in subj_combined for root in
                  ["поездк", "travel", "деловая", "командировк"])
    record("Subject содержит признак деловой поездки",
           subj_ok, f"subject: {subj_combined[:140]!r}")

    date_ok = ("10.03" in body_combined or "10 марта" in body_combined.lower()
               or "2026-03-10" in body_combined)
    record("В теле письма указана дата поездки", date_ok,
           f"body sample: {body_combined[:200]!r}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_trip_notes(args.agent_workspace)
    check_gcal()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\nИтого: {PASS_COUNT}/{total} проверок пройдено ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
