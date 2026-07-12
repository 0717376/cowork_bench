"""Evaluation for sf-hr-attrition-forecast-excel-word-gcal (ClickHouse / RU).

Checks:
1. Attrition_Risk_Analysis.xlsx with 3 sheets matching groundtruth
2. Intervention_Report.docx with required sections (RU+EN keywords)
3. 7 department review meetings in gcal.events (RU dept names)

Departments are russified centrally by db/zzz_clickhouse_after_init.sql; the
groundtruth artifacts use the same Russian names. Risk_Level (Low/Medium/High)
and Top_Factor values stay English in both data and groundtruth.

CRITICAL_CHECKS: a semantic subset that must ALL pass, else FAIL regardless of
overall accuracy. Otherwise PASS requires accuracy >= 70 and no critical fail.
"""
import json
import os
import sys
from argparse import ArgumentParser

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []

# Russian department names (must match clickhouse map + groundtruth)
DEPARTMENTS = ["Инженерия", "Финансы", "Кадры", "Операции", "НИОКР", "Продажи", "Поддержка"]


def record(name, passed, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    tag = "CRITICAL " if critical else ""
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {tag}{name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {tag}{name}{msg}")
        if critical:
            CRITICAL_FAILS.append(name)


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def sheet_dicts(wb, name):
    """Get rows as list of dicts from a sheet, matching by case-insensitive name."""
    for sn in wb.sheetnames:
        if sn.strip().lower() == name.strip().lower():
            ws = wb[sn]
            rows = list(ws.iter_rows(values_only=True))
            if len(rows) < 2:
                return []
            hdrs = [str(h).strip() if h else "" for h in rows[0]]
            return [
                {hdrs[i]: row[i] for i in range(len(hdrs))}
                for row in rows[1:]
                if not all(v is None for v in row)
            ]
    return None


def get_groundtruth(gt_path):
    """Load groundtruth from Excel file."""
    gt_xlsx = os.path.join(gt_path, "Attrition_Risk_Analysis.xlsx")
    if not os.path.isfile(gt_xlsx):
        return None
    wb = openpyxl.load_workbook(gt_xlsx, data_only=True)
    gt = {
        "risk_assessment": sheet_dicts(wb, "Risk Assessment"),
        "dept_summary": sheet_dicts(wb, "Department Summary"),
        "top_factors": sheet_dicts(wb, "Top Risk Factors"),
    }
    wb.close()
    return gt


def check_excel(agent_ws, gt):
    print("\n=== Checking Excel: Attrition_Risk_Analysis.xlsx ===")
    xlsx_path = os.path.join(agent_ws, "Attrition_Risk_Analysis.xlsx")
    if not os.path.isfile(xlsx_path):
        record("Excel file exists", False, xlsx_path, critical=True)
        return
    record("Excel file exists", True)

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    except Exception as e:
        record("Excel readable", False, str(e), critical=True)
        return
    record("Excel readable", True)

    # --- Sheet 1: Risk Assessment ---
    agent_ra = sheet_dicts(wb, "Risk Assessment")
    gt_ra = gt["risk_assessment"]
    if agent_ra is None:
        record("Sheet Risk Assessment exists", False, str(wb.sheetnames), critical=True)
    else:
        record("Sheet Risk Assessment exists", True)
        record("Risk Assessment row count", len(agent_ra) == len(gt_ra),
               f"Expected {len(gt_ra)}, got {len(agent_ra)}", critical=True)

        for gt_row in gt_ra:
            dept = gt_row.get("Department", "")
            risk = gt_row.get("Risk_Level", "")
            match = next(
                (r for r in agent_ra
                 if str_match(r.get("Department"), dept) and str_match(r.get("Risk_Level"), risk)),
                None
            )
            if not match:
                record(f"Row {dept}/{risk} exists", False, "Missing")
                continue

            # CRITICAL for High-risk rows: exact employee count + tight growth
            is_high = str(risk).strip().lower() == "high"
            record(f"{dept}/{risk} Employee_Count",
                   num_close(match.get("Employee_Count"), gt_row.get("Employee_Count"), 0),
                   f"Got {match.get('Employee_Count')} vs {gt_row.get('Employee_Count')}",
                   critical=is_high)
            record(f"{dept}/{risk} Avg_Satisfaction",
                   num_close(match.get("Avg_Satisfaction"), gt_row.get("Avg_Satisfaction"), 0.2),
                   f"Got {match.get('Avg_Satisfaction')} vs {gt_row.get('Avg_Satisfaction')}")
            record(f"{dept}/{risk} Avg_Performance",
                   num_close(match.get("Avg_Performance"), gt_row.get("Avg_Performance"), 0.2),
                   f"Got {match.get('Avg_Performance')} vs {gt_row.get('Avg_Performance')}")
            record(f"{dept}/{risk} Avg_Salary_Growth_Pct",
                   num_close(match.get("Avg_Salary_Growth_Pct"), gt_row.get("Avg_Salary_Growth_Pct"), 0.5),
                   f"Got {match.get('Avg_Salary_Growth_Pct')} vs {gt_row.get('Avg_Salary_Growth_Pct')}",
                   critical=is_high)

    # --- Sheet 2: Department Summary ---
    agent_ds = sheet_dicts(wb, "Department Summary")
    gt_ds = gt["dept_summary"]
    if agent_ds is None:
        record("Sheet Department Summary exists", False, str(wb.sheetnames), critical=True)
    else:
        record("Sheet Department Summary exists", True)
        record("Department Summary row count", len(agent_ds) == len(gt_ds),
               f"Expected {len(gt_ds)}, got {len(agent_ds)}", critical=True)

        for gt_row in gt_ds:
            dept = gt_row.get("Department", "")
            match = next(
                (r for r in agent_ds if str_match(r.get("Department"), dept)),
                None
            )
            if not match:
                record(f"Dept Summary {dept} exists", False, "Missing", critical=True)
                continue

            record(f"{dept} Total_Employees",
                   num_close(match.get("Total_Employees"), gt_row.get("Total_Employees"), 0),
                   f"Got {match.get('Total_Employees')} vs {gt_row.get('Total_Employees')}")
            # CRITICAL: exact High_Risk_Count and exact budget prove the scoring +
            # classification + budget formula were applied correctly.
            record(f"{dept} High_Risk_Count",
                   num_close(match.get("High_Risk_Count"), gt_row.get("High_Risk_Count"), 0),
                   f"Got {match.get('High_Risk_Count')} vs {gt_row.get('High_Risk_Count')}",
                   critical=True)
            record(f"{dept} High_Risk_Pct",
                   num_close(match.get("High_Risk_Pct"), gt_row.get("High_Risk_Pct"), 0.2),
                   f"Got {match.get('High_Risk_Pct')} vs {gt_row.get('High_Risk_Pct')}")
            record(f"{dept} Total_Intervention_Budget",
                   num_close(match.get("Total_Intervention_Budget"),
                             gt_row.get("Total_Intervention_Budget"), 1),
                   f"Got {match.get('Total_Intervention_Budget')} vs {gt_row.get('Total_Intervention_Budget')}",
                   critical=True)

    # --- Sheet 3: Top Risk Factors ---
    agent_tf = sheet_dicts(wb, "Top Risk Factors")
    gt_tf = gt["top_factors"]
    if agent_tf is None:
        record("Sheet Top Risk Factors exists", False, str(wb.sheetnames), critical=True)
    else:
        record("Sheet Top Risk Factors exists", True)
        record("Top Risk Factors row count", len(agent_tf) == len(gt_tf),
               f"Expected {len(gt_tf)}, got {len(agent_tf)}")

        for gt_row in gt_tf:
            dept = gt_row.get("Department", "")
            match = next(
                (r for r in agent_tf if str_match(r.get("Department"), dept)),
                None
            )
            if not match:
                record(f"Top Factor {dept} exists", False, "Missing", critical=True)
                continue
            # CRITICAL: the single highest-contributing factor per department.
            record(f"{dept} Top_Factor",
                   str_match(match.get("Top_Factor"), gt_row.get("Top_Factor")),
                   f"Got '{match.get('Top_Factor')}' vs '{gt_row.get('Top_Factor')}'",
                   critical=True)
            record(f"{dept} Factor_Score",
                   num_close(match.get("Factor_Score"), gt_row.get("Factor_Score"), 0.3),
                   f"Got {match.get('Factor_Score')} vs {gt_row.get('Factor_Score')}")

    wb.close()


def check_word(agent_ws, gt):
    print("\n=== Checking Word: Intervention_Report.docx ===")
    docx_path = os.path.join(agent_ws, "Intervention_Report.docx")
    if not os.path.isfile(docx_path):
        record("Word file exists", False, docx_path, critical=True)
        return
    record("Word file exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        record("Word file readable", False, str(e), critical=True)
        return
    record("Word file readable", True)

    # Collect all paragraph + table text
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            all_text += "\n" + "\t".join(c.text for c in row.cells)
    low = all_text.lower()

    # Required sections — accept RU or EN keyword for each.
    required_sections = [
        ("executive summary", ["executive summary", "краткое резюме"]),
        ("department analysis", ["department", "подразделени"]),
        ("budget", ["budget", "бюджет"]),
        ("recommended action", ["recommended action", "рекомендуем", "рекомендации", "меры"]),
        ("timeline", ["timeline", "срок", "график"]),
    ]
    for label, kws in required_sections:
        record(f"Word contains '{label}' section",
               any(k in low for k in kws),
               "Not found in document text")

    # CRITICAL: all 7 (Russian) departments mentioned.
    for dept in DEPARTMENTS:
        record(f"Word mentions {dept}",
               dept.lower() in low,
               "Department not found in document",
               critical=True)

    # Risk levels discussed (RU or EN)
    record("Word mentions risk levels",
           any(k in low for k in ["high risk", "high-risk", "высок", "риск"]),
           "No risk level discussion found")

    # CRITICAL: total intervention budget figure present, with a currency marker.
    total_budget = sum(int(r.get("Total_Intervention_Budget") or 0) for r in gt["dept_summary"])
    budget_variants = {
        str(total_budget),
        f"{total_budget:,}",
        f"{total_budget:,}".replace(",", " "),
        f"{total_budget:,}".replace(",", "."),
        f"{total_budget:,}".replace(",", " "),
    }
    has_total = any(v in all_text for v in budget_variants)
    has_currency = any(m in all_text for m in ["₽", "руб", "RUB", "$"])
    record("Word states total intervention budget with currency",
           has_total and has_currency,
           f"total={total_budget} present={has_total} currency={has_currency}",
           critical=True)


def check_gcal():
    print("\n=== Checking Calendar Events ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, description, start_datetime, end_datetime
        FROM gcal.events
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    # A review meeting: summary mentions оттока/attrition/обзор/review AND a dept name.
    REVIEW_KW = ["оттока", "обзор риска", "attrition", "risk review", "review", "обзор"]
    review_events = []
    for summary, desc, start_dt, end_dt in events:
        s_lower = (summary or "").lower()
        if any(k in s_lower for k in REVIEW_KW):
            for dept in DEPARTMENTS:
                if dept.lower() in s_lower:
                    review_events.append((summary, desc, start_dt, end_dt, dept))
                    break

    record("At least 7 department review events",
           len(review_events) >= 7,
           f"Found {len(review_events)} review events out of {len(events)} total",
           critical=True)

    found_depts = {dept for _, _, _, _, dept in review_events}
    for dept in DEPARTMENTS:
        record(f"Calendar event for {dept}",
               dept in found_depts,
               f"Found events for: {found_depts}",
               critical=True)

    # In the correct week
    in_range_count = 0
    for summary, desc, start_dt, end_dt, dept in review_events:
        if start_dt:
            date_str = start_dt.strftime("%Y-%m-%d")
            in_range = "2026-03-16" <= date_str <= "2026-03-20"
            if in_range:
                in_range_count += 1
            record(f"{dept} meeting in Mar 16-20 week",
                   in_range,
                   f"Scheduled on {date_str}")

    # 1-hour duration
    for summary, desc, start_dt, end_dt, dept in review_events:
        if start_dt and end_dt:
            duration_min = (end_dt - start_dt).total_seconds() / 60
            record(f"{dept} meeting is 1 hour",
                   55 <= duration_min <= 65,
                   f"Duration: {duration_min} minutes")

    # CRITICAL: all 7 dept meetings fall within the target week.
    record("All 7 dept meetings within Mar 16-20 week",
           in_range_count >= 7 and len(found_depts) >= 7,
           f"in_range={in_range_count} depts={len(found_depts)}",
           critical=True)

    # CRITICAL: no time conflicts among review meetings.
    intervals = sorted(
        (start_dt, end_dt) for _, _, start_dt, end_dt, _ in review_events
        if start_dt and end_dt
    )
    conflict = any(intervals[i][1] > intervals[i + 1][0] for i in range(len(intervals) - 1))
    record("No time conflicts among review meetings",
           not conflict and len(intervals) >= 7,
           f"intervals={len(intervals)} conflict={conflict}",
           critical=True)


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", default=".")
    parser.add_argument("--groundtruth_workspace", default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    gt = get_groundtruth(args.groundtruth_workspace)
    if gt is None:
        print("ERROR: Could not load groundtruth Excel.")
        sys.exit(1)

    check_excel(args.agent_workspace, gt)
    check_word(args.agent_workspace, gt)
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\n=== SUMMARY: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%) ===")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_fails": CRITICAL_FAILS,
    }
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    # Critical gate: any critical failure => FAIL regardless of accuracy.
    if CRITICAL_FAILS:
        print(f"\nFAIL: {len(CRITICAL_FAILS)} critical check(s) failed: {CRITICAL_FAILS}")
        sys.exit(1)

    # Accuracy gate.
    if accuracy >= 70:
        print("\nPASS")
        sys.exit(0)
    print(f"\nFAIL: accuracy {accuracy:.1f}% < 70%")
    sys.exit(1)


if __name__ == "__main__":
    main()
