"""Evaluation for playwright-moex-earnings-report-word-email-gcal.

The agent reads a mock earnings calendar site (RU prose, but English column
identifiers and MOEX tickers SBER.ME/GAZP.ME/LKOH.ME/MGNT.ME/MTSS.ME), pulls
live financials from the moex-finance MCP (moex.* schema), then produces:
  - Earnings_Preview_Report.docx (per-stock sections + market commentary)
  - Earnings_Data_Appendix.xlsx ("Estimates vs Actuals" + "Market Summary")
  - an email research@company.com -> equity-team@company.com
  - 5 GCal "Earnings Watch: <TICKER>" events (Apr 20-24, 2026).

Current_Price / Trailing_EPS are read LIVE from moex.stock_info (not hardcoded).
Consensus EPS values are static mock-site data authored in files/mock_pages.tar.gz.
EPS_Gap = Consensus_EPS - Trailing_EPS. Largest gap -> LKOH.ME (core finding).

Critical checks (see CRITICAL_CHECKS): any failure => overall FAIL regardless of
accuracy. Otherwise pass threshold: accuracy >= 70%.
"""
import argparse
import json
import os
import re
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

TICKERS = ["SBER.ME", "GAZP.ME", "LKOH.ME", "MGNT.ME", "MTSS.ME"]

# Static mock-site consensus data (authored in files/mock_pages.tar.gz index.html).
MOCK_CONSENSUS_EPS = {
    "SBER.ME": 54.0,
    "GAZP.ME": 95.0,
    "LKOH.ME": 1180.0,
    "MGNT.ME": 520.0,
    "MTSS.ME": 34.0,
}
# Ticker with the largest EPS gap (consensus - trailing) -> core semantic finding.
LARGEST_GAP_TICKER = "LKOH.ME"

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Semantic critical checks: any failure => overall FAIL regardless of accuracy.
CRITICAL_CHECKS = {
    "Excel Estimates vs Actuals: all 5 MOEX tickers with live Current_Price + Trailing_EPS",
    "Excel EPS_Gap equals Consensus-Trailing for >=4 of 5 stocks",
    "Word report names all 5 MOEX tickers with per-stock trailing EPS",
    "Email research@company.com -> equity-team@company.com names largest-gap ticker LKOH.ME",
    "5 GCal Earnings Watch events covering all 5 MOEX tickers",
}


def normalize_ru_numbers(text):
    """RU number normalization for substring/regex checks: collapse digit-group
    separators (space/NBSP/NNBSP/dot/comma before a 3-digit group) and turn
    decimal commas into dots ("31 588" -> "31588", "4 586,91" -> "4586.91")."""
    t = str(text or "")
    t = re.sub(r"(?<=\d)[ \xa0\u202f\u2009.,](?=\d{3}\b)", "", t)
    return re.sub(r"(?<=\d),(?=\d)", ".", t)


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        d = (str(detail)[:300]) if detail else ""
        print(f"  [FAIL] {name}: {d}")


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def parse_float(s):
    try:
        return float(re.sub(r"[^0-9.\-]", "", str(s)))
    except (ValueError, TypeError):
        return None


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def load_moex_live():
    """Read live currentPrice / trailingEps / sector matching the moex-finance MCP.

    trailingEps / sector come from moex.stock_info JSON, but currentPrice is the
    latest moex.stock_prices.close per symbol -- exactly what the MCP's
    get_stock_info returns (pg_adapter overrides the static JSON price with the
    stock_prices close, the single source of truth). Reading the stale JSON price
    instead would diverge ~5-7% from the only price tool the agent is given.
    """
    live = {}
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute(
        "SELECT symbol, data->>'trailingEps', data->>'sector' "
        "FROM moex.stock_info WHERE symbol = ANY(%s)",
        (TICKERS,),
    )
    for sym, teps, sector in cur.fetchall():
        live[sym] = {
            "price": None,
            "teps": parse_float(teps),
            "sector": sector,
        }
    cur.execute(
        "SELECT DISTINCT ON (symbol) symbol, close FROM moex.stock_prices "
        "WHERE symbol = ANY(%s) ORDER BY symbol, date DESC",
        (TICKERS,),
    )
    for sym, close in cur.fetchall():
        live.setdefault(sym, {"teps": None, "sector": None})
        live[sym]["price"] = parse_float(close)
    cur.close()
    conn.close()
    return live


def check_excel(agent_ws, live):
    import openpyxl

    agent_file = os.path.join(agent_ws, "Earnings_Data_Appendix.xlsx")
    print("Checking Excel file...")
    if not os.path.exists(agent_file):
        check("Earnings_Data_Appendix.xlsx exists", False, "file not found")
        check("Excel Estimates vs Actuals: all 5 MOEX tickers with live Current_Price + Trailing_EPS", False, "no file")
        check("Excel EPS_Gap equals Consensus-Trailing for >=4 of 5 stocks", False, "no file")
        check("Market Summary sheet has >=5 rows", False, "no file")
        return
    check("Earnings_Data_Appendix.xlsx exists", True)

    wb = openpyxl.load_workbook(agent_file, data_only=True)

    a_rows = load_sheet_rows(wb, "Estimates vs Actuals")
    check("Sheet 'Estimates vs Actuals' present", a_rows is not None)
    a_lookup = {}
    if a_rows:
        for row in a_rows[1:]:
            if row and row[0]:
                a_lookup[str(row[0]).strip().upper()] = row

    # Critical: all 5 tickers with live price + trailing EPS.
    ok_live = True
    detail_live = []
    for t in TICKERS:
        r = a_lookup.get(t.upper())
        if r is None:
            ok_live = False
            detail_live.append(f"{t} missing")
            continue
        lv = live.get(t, {})
        # col indices: 0 Ticker,1 Company,2 Sector,3 Current_Price,4 Trailing_EPS,5 Consensus_EPS,6 EPS_Gap
        price = parse_float(r[3]) if len(r) > 3 else None
        teps = parse_float(r[4]) if len(r) > 4 else None
        ptol = max(5.0, (lv.get("price") or 0) * 0.02)
        etol = max(0.5, abs(lv.get("teps") or 0) * 0.02)
        if not num_close(price, lv.get("price"), ptol):
            ok_live = False
            detail_live.append(f"{t} price {price} vs {lv.get('price')}")
        if not num_close(teps, lv.get("teps"), etol):
            ok_live = False
            detail_live.append(f"{t} teps {teps} vs {lv.get('teps')}")
    check("Excel Estimates vs Actuals: all 5 MOEX tickers with live Current_Price + Trailing_EPS",
          ok_live, "; ".join(detail_live))

    # Critical: EPS_Gap = Consensus_EPS - Trailing_EPS for >=4 of 5.
    gap_ok = 0
    gap_detail = []
    for t in TICKERS:
        r = a_lookup.get(t.upper())
        if r is None or len(r) < 7:
            gap_detail.append(f"{t} no row/gap")
            continue
        teps = parse_float(r[4])
        ceps = parse_float(r[5])
        gap = parse_float(r[6])
        if teps is None or ceps is None or gap is None:
            gap_detail.append(f"{t} non-numeric")
            continue
        if num_close(gap, ceps - teps, 0.5):
            gap_ok += 1
        else:
            gap_detail.append(f"{t} gap {gap} vs {round(ceps - teps, 2)}")
    check("Excel EPS_Gap equals Consensus-Trailing for >=4 of 5 stocks",
          gap_ok >= 4, f"{gap_ok}/5 ok; " + "; ".join(gap_detail))

    # Non-critical: Consensus_EPS matches static mock-site values.
    cons_ok = 0
    for t in TICKERS:
        r = a_lookup.get(t.upper())
        if r and len(r) > 5 and num_close(parse_float(r[5]), MOCK_CONSENSUS_EPS[t], 0.5):
            cons_ok += 1
    check("Excel Consensus_EPS matches mock site for >=4 of 5", cons_ok >= 4, f"{cons_ok}/5")

    # Non-critical: Market Summary structural.
    a_rows2 = load_sheet_rows(wb, "Market Summary")
    n2 = (len(a_rows2) - 1) if a_rows2 else 0
    check("Market Summary sheet has >=5 rows", n2 >= 5, f"{n2} rows")


def check_word(agent_ws, live):
    print("Checking Word document...")
    doc_path = os.path.join(agent_ws, "Earnings_Preview_Report.docx")
    if not os.path.exists(doc_path):
        check("Earnings_Preview_Report.docx exists", False, "file not found")
        check("Word report names all 5 MOEX tickers with per-stock trailing EPS", False, "no file")
        return
    check("Earnings_Preview_Report.docx exists", True)

    from docx import Document

    doc = Document(doc_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    low = full_text.lower()

    # Non-critical: mentions earnings (English token kept in report).
    check("Word doc mentions 'earnings'", "earnings" in low)

    # Critical: all 5 tickers present AND each has a nearby trailing-EPS number.
    missing = [t for t in TICKERS if t not in full_text]
    eps_ok = 0
    eps_detail = []
    for t in TICKERS:
        if t not in full_text:
            eps_detail.append(f"{t} absent")
            continue
        teps = (live.get(t) or {}).get("teps")
        # Scan every occurrence of the ticker; pass if any nearby window has a
        # number close to the live trailing EPS (the per-stock section).
        found = False
        start = 0
        while True:
            idx = full_text.find(t, start)
            if idx < 0:
                break
            start = idx + 1
            window = normalize_ru_numbers(full_text[idx: idx + 600])
            # Accept both EN (1180.5) and RU (1 180,5) number formats.
            nums = []
            for m in re.findall(r"\d+(?:[  ]\d{3})*(?:[.,]\d+)?", window):
                nums.append(parse_float(m))
                nums.append(parse_float(m.replace(" ", "").replace(" ", "").replace(",", ".")))
            if teps is not None and any(
                n is not None and abs(n - teps) <= max(1.0, abs(teps) * 0.05) for n in nums
            ):
                found = True
                break
        if found:
            eps_ok += 1
        else:
            eps_detail.append(f"{t} no near-EPS({teps})")
    check("Word report names all 5 MOEX tickers with per-stock trailing EPS",
          not missing and eps_ok >= 4, f"missing={missing}; eps_ok={eps_ok}/5; " + "; ".join(eps_detail))


def check_email():
    print("Checking email...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT subject, body_text, from_addr
            FROM email.messages
            WHERE to_addr::text ILIKE '%equity-team@company.com%'
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        check("Email sent to equity-team@company.com", False, str(e))
        check("Email research@company.com -> equity-team@company.com names largest-gap ticker LKOH.ME", False, str(e))
        return

    earnings_rows = [r for r in rows if r[0] and "earnings" in r[0].lower()]
    check("Email to equity-team@company.com about earnings", len(earnings_rows) > 0, f"{len(rows)} to addr")

    # Critical: from research@company.com AND body names largest-gap ticker.
    ok = False
    detail = []
    for subj, body, frm in rows:
        body = body or ""
        frm = frm or ""
        if "research@company.com" in frm.lower() and LARGEST_GAP_TICKER in body:
            ok = True
            break
        detail.append(f"from={frm} hasLKOH={LARGEST_GAP_TICKER in body}")
    check("Email research@company.com -> equity-team@company.com names largest-gap ticker LKOH.ME",
          ok, "; ".join(detail[:5]))


def check_gcal():
    print("Checking GCal events...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT summary, start_datetime::date
            FROM gcal.events
            WHERE summary ILIKE '%earnings%watch%'
            ORDER BY start_datetime
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        check("5 GCal Earnings Watch events covering all 5 MOEX tickers", False, str(e))
        return

    check("GCal has 5 'Earnings Watch' events", len(rows) >= 5, f"found {len(rows)}")
    summaries = " ".join(r[0] for r in rows).upper()
    missing = [t for t in TICKERS if t not in summaries]
    check("5 GCal Earnings Watch events covering all 5 MOEX tickers",
          len(rows) >= 5 and not missing, f"missing={missing}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    agent_ws = args.agent_workspace or task_root

    try:
        live = load_moex_live()
    except Exception as e:
        print(f"FATAL: cannot read moex.stock_info: {e}")
        sys.exit(1)

    check_excel(agent_ws, live)
    check_word(agent_ws, live)
    check_email()
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print("\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}, Accuracy: {accuracy:.1f}%")
    if critical_failed:
        print(f"  CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"    - {n}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({
                "total_passed": PASS_COUNT,
                "total_checks": total,
                "accuracy": accuracy,
                "critical_failed": critical_failed,
                "success": (not critical_failed) and accuracy >= 70,
            }, f, indent=2)

    if critical_failed:
        print("Overall: FAIL (critical check failed)")
        sys.exit(1)
    if accuracy >= 70:
        print("Overall: PASS")
        sys.exit(0)
    print("Overall: FAIL")
    sys.exit(1)


if __name__ == "__main__":
    main()
