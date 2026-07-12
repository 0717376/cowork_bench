"""
Evaluation for yt-veritasium-science-quiz-word-gcal task.

CRITICAL (содержательные) проверки гейтят результат: если хоть одна провалена,
итог = FAIL независимо от accuracy. Структурные проверки (наличие файла/листа)
обязательны, но не являются единственным условием прохождения.

Гейт прохождения: нет провалов CRITICAL И accuracy >= 70%.

Замечание: список видео берётся вживую с YouTube (Veritasium), поэтому точные
названия/счётчики не захардкожены — проверки устойчивы к живым данным, но при этом
содержательны (структура колонок, статусы, число строк, согласованность вопросов,
реальные события календаря, осмысленное письмо).
"""
import os
import sys
import json
from argparse import ArgumentParser

import psycopg2

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []

# CRITICAL — содержательные проверки. Любой их провал = FAIL всей задачи.
CRITICAL_CHECKS = {
    "Word: документ содержит >= 4 заголовка уровня 1 (по одному на видео)",
    "GCal: создано >= 3 события 'Science Study Session' в окне 14.03–14.04.2026",
    "GSheet: лист Videos содержит обязательные колонки и >= 6 строк данных",
    "GSheet: лист Videos — все статусы Study_Status = 'Pending'",
    "Email: письмо отправлено на studygroup@university.edu с осмысленной темой и телом",
}

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        marker = " [CRITICAL]" if name in CRITICAL_CHECKS else ""
        print(f"  [FAIL]{marker} {name}{msg}")
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILS.append(name)


def check_word_doc(agent_workspace):
    print("\n=== Проверка 1-3: Veritasium_Science_Quiz.docx ===")
    docx_path = os.path.join(agent_workspace, "Veritasium_Science_Quiz.docx")
    if not os.path.exists(docx_path):
        record("Veritasium_Science_Quiz.docx существует", False, f"Не найден: {docx_path}")
        record("Word: документ содержит >= 4 заголовка уровня 1 (по одному на видео)",
               False, "Файл отсутствует")
        record("Word: текст содержит научные термины", False, "Файл отсутствует")
        return
    record("Veritasium_Science_Quiz.docx существует", True)

    try:
        import docx
        doc = docx.Document(docx_path)
    except ImportError:
        size = os.path.getsize(docx_path)
        # Без библиотеки docx критичную проверку структуры выполнить нельзя -> CRITICAL fail.
        record("Word: документ содержит >= 4 заголовка уровня 1 (по одному на видео)",
               False, "Библиотека python-docx недоступна, структуру проверить нельзя")
        record("Word: текст содержит научные термины", size > 1000, f"Size: {size}")
        return
    except Exception as e:
        record("Word: документ содержит >= 4 заголовка уровня 1 (по одному на видео)",
               False, str(e))
        record("Word: текст содержит научные термины", False, str(e))
        return

    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    record("Word: документ содержит >= 4 заголовка уровня 1 (по одному на видео)",
           len(headings) >= 4, f"Найдено заголовков Heading 1: {len(headings)}")

    all_text = " ".join(p.text.lower() for p in doc.paragraphs)
    science_keywords = ["physics", "math", "science", "engineering", "biology",
                        "chemistry", "quantum", "gravity", "speed", "light",
                        "energy", "wave", "probability", "paradox", "experiment",
                        "психолог", "физик", "биолог", "энерги", "квант", "наук"]
    found = [kw for kw in science_keywords if kw in all_text]
    record("Word: текст содержит научные термины", len(found) >= 1,
           f"Найдены термины: {found[:5]}")

    # Документ должен содержать осмысленные вопросы (по три на видео).
    qmarks = all_text.count("?")
    record("Word: документ содержит вопросы на понимание (>= 8 знаков '?')",
           qmarks >= 8, f"Найдено знаков вопроса: {qmarks}")


def check_gcal():
    print("\n=== Проверка 4: события Science Study Session в календаре ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        with conn.cursor() as cur:
            # Скобки вокруг OR обязательны: иначе AND с датой относится только ко второй ветке.
            cur.execute("""
                SELECT COUNT(*) FROM gcal.events
                WHERE (summary ILIKE '%Science Study%' OR summary ILIKE '%Study Session%')
                  AND start_datetime >= '2026-03-14'
                  AND start_datetime < '2026-04-15'
            """)
            count = cur.fetchone()[0]
        conn.close()
        record("GCal: создано >= 3 события 'Science Study Session' в окне 14.03–14.04.2026",
               count >= 3, f"Найдено событий: {count}")
    except Exception as e:
        record("GCal: создано >= 3 события 'Science Study Session' в окне 14.03–14.04.2026",
               False, str(e))


def check_gsheet():
    print("\n=== Проверка 5: таблица Veritasium Study Tracker ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        with conn.cursor() as cur:
            cur.execute("""
                SELECT s.id, s.title FROM gsheet.spreadsheets s
                WHERE s.title ILIKE '%Veritasium%' OR s.title ILIKE '%Study Tracker%'
            """)
            spreadsheets = cur.fetchall()
            if not spreadsheets:
                record("GSheet: таблица 'Veritasium Study Tracker' существует", False,
                       "Подходящая таблица не найдена")
                record("GSheet: лист Videos содержит обязательные колонки и >= 6 строк данных",
                       False, "Таблица отсутствует")
                record("GSheet: лист Videos — все статусы Study_Status = 'Pending'",
                       False, "Таблица отсутствует")
                conn.close()
                return
            record("GSheet: таблица 'Veritasium Study Tracker' существует", True,
                   f"Найдено: {[r[1] for r in spreadsheets]}")

            ss_id = spreadsheets[0][0]
            cur.execute("""
                SELECT sh.id, sh.title FROM gsheet.sheets sh
                WHERE sh.spreadsheet_id = %s AND sh.title ILIKE '%%video%%'
            """, (ss_id,))
            sheets = cur.fetchall()
            if not sheets:
                record("GSheet: лист Videos содержит обязательные колонки и >= 6 строк данных",
                       False, "Лист Videos не найден")
                record("GSheet: лист Videos — все статусы Study_Status = 'Pending'",
                       False, "Лист Videos не найден")
                conn.close()
                return
            sh_id = sheets[0][0]

            # Загрузить все ячейки листа Videos в карту (row, col) -> value.
            cur.execute("""
                SELECT row_index, col_index, value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s
            """, (ss_id, sh_id))
            cells = cur.fetchall()
        conn.close()

        grid = {}
        max_row = 0
        for r, c, v in cells:
            grid[(r, c)] = (v or "")
            if r > max_row:
                max_row = r

        # Заголовки в строке 0 (или 1, если 1-индексация).
        header_row = 0 if any(k[0] == 0 for k in grid) else 1
        headers = {}
        for (r, c), v in grid.items():
            if r == header_row:
                headers[v.strip().lower()] = c
        required = ["video_id", "title", "topic", "duration_min", "difficulty", "study_status"]
        missing = [h for h in required if h not in headers]

        data_rows = sorted({r for (r, c) in grid if r > header_row})
        enough_rows = len(data_rows) >= 6
        record("GSheet: лист Videos содержит обязательные колонки и >= 6 строк данных",
               not missing and enough_rows,
               f"Отсутствуют колонки: {missing}; строк данных: {len(data_rows)}")

        # Все статусы Study_Status = 'Pending'.
        if "study_status" in headers:
            sc = headers["study_status"]
            statuses = [grid.get((r, sc), "").strip() for r in data_rows]
            non_pending = [s for s in statuses if s.lower() != "pending"]
            record("GSheet: лист Videos — все статусы Study_Status = 'Pending'",
                   len(statuses) > 0 and not non_pending,
                   f"Статусы != Pending: {non_pending[:5]}")
        else:
            record("GSheet: лист Videos — все статусы Study_Status = 'Pending'",
                   False, "Колонка Study_Status отсутствует")
    except Exception as e:
        record("GSheet: лист Videos содержит обязательные колонки и >= 6 строк данных",
               False, str(e))
        record("GSheet: лист Videos — все статусы Study_Status = 'Pending'", False, str(e))


def check_email():
    print("\n=== Проверка 6: письмо на studygroup@university.edu ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        rows = []
        with conn.cursor() as cur:
            cur.execute("""
                SELECT subject, body_text FROM email.messages
                WHERE to_addr::text ILIKE '%studygroup@university.edu%'
            """)
            rows = cur.fetchall()
            if not rows:
                try:
                    cur.execute("""
                        SELECT subject, body_text FROM email.sent_log
                        WHERE to_addr ILIKE '%studygroup@university.edu%'
                    """)
                    rows = cur.fetchall()
                except Exception:
                    pass
        conn.close()

        if not rows:
            record("Email: письмо отправлено на studygroup@university.edu с осмысленной темой и телом",
                   False, "Письмо не найдено")
            return

        ok = False
        for subject, body in rows:
            subject = (subject or "")
            body = (body or "")
            subj_l = subject.lower()
            body_l = body.lower()
            # Тема должна ссылаться на викторину/расписание/учебные сессии.
            subj_ok = any(k in subj_l for k in
                          ["quiz", "викторин", "study", "session", "сесси", "veritasium", "расписан"])
            # Тело должно упоминать сессии и быть содержательным.
            body_ok = (len(body.strip()) >= 60 and
                       any(k in body_l for k in ["session", "сесси", "study", "учебн"]))
            if subj_ok and body_ok:
                ok = True
                break
        record("Email: письмо отправлено на studygroup@university.edu с осмысленной темой и телом",
               ok, f"Писем найдено: {len(rows)}, релевантное: {ok}")
    except Exception as e:
        record("Email: письмо отправлено на studygroup@university.edu с осмысленной темой и телом",
               False, str(e))


def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    print(f"Running evaluation for yt-veritasium-science-quiz-word-gcal")
    print(f"Agent workspace: {agent_workspace}")

    check_word_doc(agent_workspace)
    check_gcal()
    check_gsheet()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    print(f"\n{'='*40}")
    print(f"Результат: {PASS_COUNT}/{total} пройдено ({accuracy:.1f}%)")
    if CRITICAL_FAILS:
        print(f"Провалены CRITICAL ({len(CRITICAL_FAILS)}): {CRITICAL_FAILS}")

    all_passed = (not CRITICAL_FAILS) and (accuracy >= 70)
    summary = f"Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}, accuracy={accuracy:.1f}%"

    if res_log_file:
        with open(res_log_file, "w") as f:
            json.dump({
                "passed": PASS_COUNT,
                "failed": FAIL_COUNT,
                "accuracy": accuracy,
                "critical_fails": CRITICAL_FAILS,
                "all_passed": all_passed,
            }, f)

    return all_passed, summary


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()
    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("FAIL: ни одной проверки не выполнено")
        sys.exit(1)
    if CRITICAL_FAILS:
        print(f"FAIL: провалены критичные проверки ({len(CRITICAL_FAILS)})")
        sys.exit(1)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
