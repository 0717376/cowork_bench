"""Evaluation for sf-salary-market-benchmark (ClickHouse / RU fork)."""
import argparse
import os
import sys

import openpyxl


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    agent_file = os.path.join(args.agent_workspace, "Salary_Benchmark.xlsx")
    gt_file = os.path.join(gt_dir, "Salary_Benchmark.xlsx")

    if not os.path.exists(agent_file):
        print(f"FAIL: Agent output not found: {agent_file}")
        sys.exit(1)
    if not os.path.exists(gt_file):
        print(f"FAIL: Groundtruth not found: {gt_file}")
        sys.exit(1)

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    all_errors = []        # non-critical errors -> aggregate into accuracy
    critical_errors = []   # any one of these => immediate FAIL
    total_checks = 0
    passed_checks = 0

    def check(ok, label, critical=False):
        nonlocal total_checks, passed_checks
        total_checks += 1
        if ok:
            passed_checks += 1
        else:
            if critical:
                critical_errors.append(label)
            else:
                all_errors.append(label)
        return ok

    # ---------- Department Analysis ----------
    print("  Checking Department Analysis...")
    a_rows = load_sheet_rows(agent_wb, "Department Analysis")
    g_rows = load_sheet_rows(gt_wb, "Department Analysis")
    if a_rows is None:
        critical_errors.append("Sheet 'Department Analysis' not found in agent output")
    elif g_rows is None:
        all_errors.append("Sheet 'Department Analysis' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {str(r[0]).strip().lower(): r for r in a_data if r and r[0] is not None}
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                # missing department row is a CRITICAL failure (key-language desync guard)
                check(False, f"DeptAnalysis missing row: {g_row[0]}", critical=True)
                continue
            # Avg_Salary [1], Market_Benchmark [2] : soft numeric (non-critical)
            if len(a_row) > 1 and len(g_row) > 1:
                check(num_close(a_row[1], g_row[1], 500),
                      f"{key}.Avg_Salary: {a_row[1]} vs {g_row[1]}")
            if len(a_row) > 2 and len(g_row) > 2:
                check(num_close(a_row[2], g_row[2], 500),
                      f"{key}.Market_Benchmark: {a_row[2]} vs {g_row[2]}")
            # Pay_Ratio [3] : CRITICAL (core deliverable)
            if len(a_row) > 3 and len(g_row) > 3:
                check(num_close(a_row[3], g_row[3], 0.05),
                      f"{key}.Pay_Ratio: {a_row[3]} vs {g_row[3]}", critical=True)
            # Above_Market_Pct [6] / Below_Market_Pct [7] : CRITICAL (previously unchecked)
            if len(a_row) > 6 and len(g_row) > 6:
                check(num_close(a_row[6], g_row[6], 3.0),
                      f"{key}.Above_Market_Pct: {a_row[6]} vs {g_row[6]}", critical=True)
            if len(a_row) > 7 and len(g_row) > 7:
                check(num_close(a_row[7], g_row[7], 3.0),
                      f"{key}.Below_Market_Pct: {a_row[7]} vs {g_row[7]}", critical=True)

    # ---------- Role Details ----------
    print("  Checking Role Details...")
    a_rows = load_sheet_rows(agent_wb, "Role Details")
    g_rows = load_sheet_rows(gt_wb, "Role Details")
    if a_rows is None:
        critical_errors.append("Sheet 'Role Details' not found in agent output")
    elif g_rows is None:
        all_errors.append("Sheet 'Role Details' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {}
        for r in a_data:
            if r and r[0] is not None and r[1] is not None:
                k = f"{str(r[0]).strip().lower()}|{str(r[1]).strip().lower()}"
                a_lookup[k] = r
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = f"{str(g_row[0]).strip().lower()}|{str(g_row[1]).strip().lower()}"
            a_row = a_lookup.get(key)
            if a_row is None:
                check(False, f"RoleDetails missing row: {g_row[0]}|{g_row[1]}", critical=True)
                continue
            # Avg_Salary [3], Market_Benchmark [4] : soft numeric (non-critical)
            if len(a_row) > 3 and len(g_row) > 3:
                check(num_close(a_row[3], g_row[3], 500),
                      f"{key}.Avg_Salary: {a_row[3]} vs {g_row[3]}")
            if len(a_row) > 4 and len(g_row) > 4:
                check(num_close(a_row[4], g_row[4], 500),
                      f"{key}.Market_Benchmark: {a_row[4]} vs {g_row[4]}")
            # Pay_Ratio [5] : non-critical numeric (status below is the critical classification)
            if len(a_row) > 5 and len(g_row) > 5:
                check(num_close(a_row[5], g_row[5], 0.05),
                      f"{key}.Pay_Ratio: {a_row[5]} vs {g_row[5]}")
            # Status [7] : CRITICAL (central business-logic classification, English literals)
            if len(a_row) > 7 and len(g_row) > 7:
                check(str_match(a_row[7], g_row[7]),
                      f"{key}.Status: {a_row[7]} vs {g_row[7]}", critical=True)

    # ---------- Summary ----------
    print("  Checking Summary...")
    a_rows = load_sheet_rows(agent_wb, "Summary")
    g_rows = load_sheet_rows(gt_wb, "Summary")
    if a_rows is None:
        critical_errors.append("Sheet 'Summary' not found in agent output")
    elif g_rows is None:
        all_errors.append("Sheet 'Summary' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {str(r[0]).strip().lower(): r for r in a_data if r and r[0] is not None}
        # which summary metrics are CRITICAL
        crit_counts = ("total_employees", "total_roles_analyzed",
                       "roles_above_market", "roles_below_market")
        crit_roles = ("most_above_market_role", "most_below_market_role")
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                # missing critical metric rows fail critically
                check(False, f"Summary missing row: {g_row[0]}",
                      critical=(key in crit_counts or key in crit_roles))
                continue
            if len(a_row) > 1 and len(g_row) > 1:
                if key in crit_counts:
                    check(num_close(a_row[1], g_row[1], 2),
                          f"{key}: {a_row[1]} vs {g_row[1]}", critical=True)
                elif key in ("overall_pay_ratio", "most_above_market_ratio", "most_below_market_ratio"):
                    check(num_close(a_row[1], g_row[1], 0.05),
                          f"{key}: {a_row[1]} vs {g_row[1]}")
                elif key in ("overall_avg_salary", "overall_market_avg"):
                    check(num_close(a_row[1], g_row[1], 500),
                          f"{key}: {a_row[1]} vs {g_row[1]}")
                elif "role" in key:
                    # Most_Above/Below_Market_Role: exact "Role (Dept)" identification
                    check(str_match(a_row[1], g_row[1]),
                          f"{key}: {a_row[1]} vs {g_row[1]}",
                          critical=(key in crit_roles))

    # ---------- Word document ----------
    print("  Checking Compensation_Report.docx...")
    word_file = os.path.join(args.agent_workspace, "Compensation_Report.docx")
    if not os.path.exists(word_file):
        critical_errors.append("Compensation_Report.docx not found")
        print("    FAIL: file not found")
    else:
        try:
            from docx import Document
            doc = Document(word_file)
            # ORIGINAL lowercased text (NOT normalized) for RU keyword checks
            text = " ".join(p.text for p in doc.paragraphs).lower()

            def has_any(keywords):
                return any(k in text for k in keywords)

            # executive summary section (RU or EN)
            check(has_any(["executive summary", "итоговое резюме", "резюме", "summary"]),
                  "Word doc missing executive summary / итоговое резюме section",
                  critical=True)
            # recommendations section (RU or EN)
            check(has_any(["recommend", "рекомендац"]),
                  "Word doc missing recommendations / рекомендации section",
                  critical=True)
            # non-trivial length
            check(len(doc.paragraphs) >= 5, "Word doc too short (< 5 paragraphs)")
            # at least one concrete pay-ratio / threshold figure mentioned
            check(has_any(["0.95", "1.05", "1.25", "0.90", "1.30",
                           "0,95", "1,05", "1,25", "0,90", "1,30",
                           "pay ratio", "коэффициент опла"]),
                  "Word doc missing a concrete pay-ratio / threshold figure")
        except Exception as e:
            critical_errors.append(f"Word doc read error: {e}")
            print(f"    ERROR: {e}")

    # ---------- Verdict ----------
    # CRITICAL gate first
    if critical_errors:
        print(f"\n=== CRITICAL FAILURE ({len(critical_errors)} critical checks failed) ===")
        for e in critical_errors[:15]:
            print(f"  CRITICAL: {e}")
        print("=== RESULT: FAIL ===")
        sys.exit(1)

    accuracy = (passed_checks / total_checks * 100.0) if total_checks else 0.0
    print(f"\nNon-critical errors: {len(all_errors)}")
    for e in all_errors[:10]:
        print(f"  {e}")
    print(f"Accuracy: {passed_checks}/{total_checks} = {accuracy:.1f}%")

    if accuracy >= 70.0:
        print("=== RESULT: PASS ===")
        sys.exit(0)
    else:
        print("=== RESULT: FAIL (accuracy < 70) ===")
        sys.exit(1)


if __name__ == "__main__":
    main()
