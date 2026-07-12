"""Evaluation for sf-regional-sales-forecast (ClickHouse, russified)."""
import argparse
import os
import sys

import openpyxl
import psycopg2

# Russian region names from the central ClickHouse map
# (db/zzz_clickhouse_after_init.sql). These are the shared join keys across the
# warehouse, the mock economic API, the methodology doc and the groundtruth.
RU_ASIA = "азиатско-тихоокеанский регион"
RU_EUROPE = "европа"
RU_LATAM = "латинская америка"
RU_MIDEAST = "ближний восток"
RU_NA = "северная америка"


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    agent_file = os.path.join(args.agent_workspace, "Sales_Forecast.xlsx")
    gt_file = os.path.join(gt_dir, "Sales_Forecast.xlsx")

    if not os.path.exists(agent_file):
        print(f"FAIL: Agent output not found: {agent_file}")
        sys.exit(1)
    if not os.path.exists(gt_file):
        print(f"FAIL: Groundtruth not found: {gt_file}")
        sys.exit(1)

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    all_errors = []
    critical_failures = []

    # ------------------------------------------------------------------
    # CRITICAL CHECKS (semantic): verify the warehouse<->API region join and
    # the core forecast formula, keyed on the Russian region names from the
    # central ClickHouse map. Any critical failure => sys.exit(1) below,
    # BEFORE the accuracy gate. Structural/loose checks remain non-critical.
    # ------------------------------------------------------------------
    def forecast_lookup(wb):
        rows = load_sheet_rows(wb, "Forecast")
        if not rows or len(rows) < 2:
            return {}
        out = {}
        for r in rows[1:]:
            if r and r[0] is not None:
                out[str(r[0]).strip().lower()] = r
        return out

    def summary_lookup(wb):
        rows = load_sheet_rows(wb, "Summary")
        if not rows or len(rows) < 2:
            return {}
        out = {}
        for r in rows[1:]:
            if r and r[0] is not None:
                out[str(r[0]).strip().lower()] = r
        return out

    a_fc = forecast_lookup(agent_wb)
    a_sm = summary_lookup(agent_wb)

    # Expected per-region Next_Quarter_Forecast / Economic_Adjustment / GDP
    # keyed on the Russian region names (must match the russified warehouse+API).
    EXP_REGIONS = {
        RU_ASIA: {"q": 93992.01, "adj": 1.038, "gdp": 3.8},
        RU_EUROPE: {"q": 73711.86, "adj": 1.015, "gdp": 1.5},
        RU_LATAM: {"q": 76586.28, "adj": 1.012, "gdp": 1.2},
        RU_MIDEAST: {"q": 75943.92, "adj": 1.025, "gdp": 2.5},
        RU_NA: {"q": 82404.75, "adj": 1.021, "gdp": 2.1},
    }

    # CRITICAL 1+2: per-region forecast & economic adjustment for ALL 5 regions,
    # keyed on Russian region names -> proves correct join + formula.
    for region, exp in EXP_REGIONS.items():
        row = a_fc.get(region)
        if row is None:
            critical_failures.append(
                f"CRITICAL: Forecast region missing or wrong key: '{region}'")
            continue
        # Next_Quarter_Forecast (col idx 5), tight tolerance
        if len(row) <= 5 or not num_close(row[5], exp["q"], 50):
            critical_failures.append(
                f"CRITICAL: {region}.Next_Quarter_Forecast="
                f"{row[5] if len(row) > 5 else None} vs {exp['q']} (+/-50)")
        # Economic_Adjustment (col idx 3) == 1+GDP/100
        if len(row) <= 3 or not num_close(row[3], exp["adj"], 0.001):
            critical_failures.append(
                f"CRITICAL: {region}.Economic_Adjustment="
                f"{row[3] if len(row) > 3 else None} vs {exp['adj']} (+/-0.001)")

    # CRITICAL 3: Summary totals + region resolution to Russian names.
    def sm_val(key):
        r = a_sm.get(key)
        return r[1] if r and len(r) > 1 else None

    if not num_close(sm_val("total_q1_2026_forecast"), 402638.85, 100):
        critical_failures.append(
            f"CRITICAL: Total_Q1_2026_Forecast={sm_val('total_q1_2026_forecast')} "
            f"vs 402638.85 (+/-100)")
    if not str_match(sm_val("highest_forecast_region"), "Азиатско-Тихоокеанский регион"):
        critical_failures.append(
            f"CRITICAL: Highest_Forecast_Region={sm_val('highest_forecast_region')} "
            f"vs 'Азиатско-Тихоокеанский регион'")
    if not str_match(sm_val("highest_gdp_growth_region"), "Азиатско-Тихоокеанский регион"):
        critical_failures.append(
            f"CRITICAL: Highest_GDP_Growth_Region={sm_val('highest_gdp_growth_region')} "
            f"vs 'Азиатско-Тихоокеанский регион'")
    if not str_match(sm_val("lowest_forecast_region"), "Европа"):
        critical_failures.append(
            f"CRITICAL: Lowest_Forecast_Region={sm_val('lowest_forecast_region')} "
            f"vs 'Европа'")

    # CRITICAL 4: QoQ growth derived from correct Q4 actual vs Q1 forecast.
    if not num_close(sm_val("qoq_growth_pct"), 2.3, 0.3):
        critical_failures.append(
            f"CRITICAL: QoQ_Growth_Pct={sm_val('qoq_growth_pct')} vs 2.3 (+/-0.3)")

    # Check Historical sheet
    print("  Checking Historical...")
    a_rows = load_sheet_rows(agent_wb, "Historical")
    g_rows = load_sheet_rows(gt_wb, "Historical")
    if a_rows is None:
        all_errors.append("Sheet 'Historical' not found")
    elif g_rows is None:
        all_errors.append("Sheet 'Historical' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {}
        for r in a_data:
            if r and r[0] is not None and r[1] is not None:
                k = f"{str(r[0]).strip().lower()}|{str(r[1]).strip()}"
                a_lookup[k] = r
        errors = []
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = f"{str(g_row[0]).strip().lower()}|{str(g_row[1]).strip()}"
            a_row = a_lookup.get(key)
            if a_row is None:
                errors.append(f"Missing: {g_row[0]}|{g_row[1]}")
                continue
            # Revenue
            if len(a_row) > 3 and len(g_row) > 3:
                if not num_close(a_row[3], g_row[3], 100):
                    errors.append(f"{key}.Revenue: {a_row[3]} vs {g_row[3]}")
        if errors:
            all_errors.extend(errors)
            print(f"    ERRORS: {len(errors)}")
            for e in errors[:5]:
                print(f"      {e}")
        else:
            print("    PASS")

    # Check Forecast sheet
    print("  Checking Forecast...")
    a_rows = load_sheet_rows(agent_wb, "Forecast")
    g_rows = load_sheet_rows(gt_wb, "Forecast")
    if a_rows is None:
        all_errors.append("Sheet 'Forecast' not found")
    elif g_rows is None:
        all_errors.append("Sheet 'Forecast' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {str(r[0]).strip().lower(): r for r in a_data if r and r[0] is not None}
        errors = []
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                errors.append(f"Missing region: {g_row[0]}")
                continue
            # Trailing_3M_Avg
            if len(a_row) > 1 and len(g_row) > 1:
                if not num_close(a_row[1], g_row[1], 200):
                    errors.append(f"{key}.Trailing_Avg: {a_row[1]} vs {g_row[1]}")
            # GDP_Growth_Pct
            if len(a_row) > 2 and len(g_row) > 2:
                if not num_close(a_row[2], g_row[2], 0.1):
                    errors.append(f"{key}.GDP: {a_row[2]} vs {g_row[2]}")
            # Next_Quarter_Forecast
            if len(a_row) > 5 and len(g_row) > 5:
                if not num_close(a_row[5], g_row[5], 1000):
                    errors.append(f"{key}.Q_Forecast: {a_row[5]} vs {g_row[5]}")
        if errors:
            all_errors.extend(errors)
            print(f"    ERRORS: {len(errors)}")
            for e in errors[:5]:
                print(f"      {e}")
        else:
            print("    PASS")

    # Check Summary sheet
    print("  Checking Summary...")
    a_rows = load_sheet_rows(agent_wb, "Summary")
    g_rows = load_sheet_rows(gt_wb, "Summary")
    if a_rows is None:
        all_errors.append("Sheet 'Summary' not found")
    elif g_rows is None:
        all_errors.append("Sheet 'Summary' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {str(r[0]).strip().lower(): r for r in a_data if r and r[0] is not None}
        errors = []
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                errors.append(f"Missing: {g_row[0]}")
                continue
            if len(a_row) > 1 and len(g_row) > 1:
                if "region" in key:
                    if not str_match(a_row[1], g_row[1]):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
                elif "count" in key:
                    if not num_close(a_row[1], g_row[1], 1):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
                elif "pct" in key:
                    if not num_close(a_row[1], g_row[1], 1.0):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
                else:
                    if not num_close(a_row[1], g_row[1], 2000):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
        if errors:
            all_errors.extend(errors)
            print(f"    ERRORS: {len(errors)}")
            for e in errors[:5]:
                print(f"      {e}")
        else:
            print("    PASS")

    # Check Word document
    print("  Checking Forecast_Report.docx...")
    word_file = os.path.join(args.agent_workspace, "Forecast_Report.docx")
    if not os.path.exists(word_file):
        all_errors.append("Forecast_Report.docx not found")
        critical_failures.append("CRITICAL: Forecast_Report.docx not found")
        print("    FAIL: file not found")
    else:
        try:
            from docx import Document
            doc = Document(word_file)
            text = " ".join(p.text for p in doc.paragraphs).lower()
            has_method = ("методолог" in text or "формул" in text
                          or "methodology" in text or "formula" in text)
            has_forecast = ("прогноз" in text or "forecast" in text)
            checks = [
                (has_method, "Missing methodology section"),
                (has_forecast, "Missing forecast content"),
                (len(doc.paragraphs) >= 5, "Document too short"),
            ]
            for cond, msg in checks:
                if not cond:
                    all_errors.append(msg)
            print("    PASS" if all(c for c, _ in checks) else "    ERRORS found")

            # CRITICAL 5: substantive report -- methodology/forecast prose (ru+en)
            # AND at least 2 concrete region forecast numbers, proving it is not
            # a stub.
            digits = "".join(ch for ch in text if ch.isdigit())
            number_hits = sum(
                1 for tok in (
                    "93992", "82404", "73711", "76586", "75943",
                    "402638", "393653",
                )
                if tok in digits
            )
            if not (has_method and has_forecast):
                critical_failures.append(
                    "CRITICAL: Forecast_Report.docx missing methodology/forecast prose")
            if number_hits < 2:
                critical_failures.append(
                    f"CRITICAL: Forecast_Report.docx lacks concrete region forecast "
                    f"numbers (found {number_hits}, need >=2)")
        except Exception as e:
            all_errors.append(f"Word doc error: {e}")
            critical_failures.append(f"CRITICAL: Word doc error: {e}")
            print(f"    ERROR: {e}")

    # Check Google Sheet
    print("  Checking Google Sheet...")
    try:
        db_config = {
            "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
            "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
            "user": "eigent", "password": "camel",
        }
        conn = psycopg2.connect(**db_config)
        cur = conn.cursor()
        cur.execute("SELECT id, title FROM gsheet.spreadsheets WHERE title LIKE '%Forecast%' OR title LIKE '%forecast%'")
        sheets = cur.fetchall()
        if len(sheets) < 1:
            all_errors.append("No Google Sheet with 'Forecast' in title found")
            print("    FAIL: no forecast spreadsheet")
        else:
            ss_id = sheets[0][0]
            cur.execute("SELECT COUNT(*) FROM gsheet.cells WHERE spreadsheet_id = %s", (ss_id,))
            cell_count = cur.fetchone()[0]
            if cell_count < 10:
                all_errors.append(f"Google Sheet has only {cell_count} cells")
                print(f"    FAIL: only {cell_count} cells")
            else:
                print("    PASS")
        cur.close()
        conn.close()
    except Exception as e:
        all_errors.append(f"GSheet check error: {e}")
        print(f"    ERROR: {e}")

    # ------------------------------------------------------------------
    # CRITICAL GATE: any critical (semantic) failure => immediate FAIL,
    # independent of accuracy.
    # ------------------------------------------------------------------
    if critical_failures:
        print(f"\n=== CRITICAL FAILURES ({len(critical_failures)}) ===")
        for e in critical_failures:
            print(f"  {e}")
        print("\n=== RESULT: FAIL (critical) ===")
        sys.exit(1)

    # ------------------------------------------------------------------
    # ACCURACY GATE: structural / loose checks. PASS requires accuracy >= 70.
    # Sections: Historical, Forecast, Summary, Word(method), Word(forecast),
    # Word(length), GSheet.
    # ------------------------------------------------------------------
    def errored(*subs):
        return any(any(s in e for s in subs) for e in all_errors)

    section_ok = [
        not errored(".Revenue", "Sheet 'Historical'", "Missing: "),          # Historical
        not errored(".Trailing_Avg", ".GDP", ".Q_Forecast",
                    "Sheet 'Forecast'", "Missing region"),                    # Forecast
        not errored("total_", "qoq_", "highest_", "lowest_", "regions_count",
                    "Sheet 'Summary'"),                                       # Summary
        not errored("Missing methodology"),                                   # Word method
        not errored("Missing forecast content", "not found", "Word doc error"),  # Word forecast
        not errored("Document too short"),                                    # Word length
        not errored("Google Sheet", "forecast spreadsheet", "GSheet check"),  # GSheet
    ]
    passed = sum(1 for ok in section_ok if ok)
    total = len(section_ok)
    accuracy = 100.0 * passed / total

    print(f"\n=== Section accuracy: {passed}/{total} = {accuracy:.1f}% ===")
    if all_errors:
        print(f"--- {len(all_errors)} non-critical errors ---")
        for e in all_errors[:10]:
            print(f"  {e}")

    if accuracy >= 70:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)
    else:
        print("\n=== RESULT: FAIL (accuracy < 70) ===")
        sys.exit(1)


if __name__ == "__main__":
    main()
