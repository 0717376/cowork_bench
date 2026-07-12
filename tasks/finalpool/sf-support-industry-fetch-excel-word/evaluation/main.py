"""Evaluation script for sf-support-industry-fetch-excel-word (ClickHouse fork).

The support center metrics come from the global sf_data seed (SUPPORT_CENTER, russified
centrally) and the task-local mock benchmark JSON (numeric values 5.2/87.5/4.1/3500).
Identifiers (file/sheet/column names, PRIORITY/STATUS realia) stay English by design;
the agent may write Russian prose in recommendations / docx body, so free-text checks
accept RU+EN alternatives.
"""
import os
import argparse, json, os, sys
import openpyxl

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRIT]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL]{' [CRIT]' if critical else ''} {name}: {detail_str}")
        if critical:
            CRITICAL_FAILS.append(name)


def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    try:
        return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)
    except (ValueError, TypeError):
        return False


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel"
}


def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default


def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)


def find_row(rows, header_map, key_col, key_val):
    """Return the first data row whose key_col cell equals key_val (case-insensitive)."""
    ci = header_map.get(key_col.lower())
    if ci is None:
        return None
    for r in rows:
        if ci < len(r) and r[ci] is not None and str(r[ci]).strip().lower() == key_val.lower():
            return r
    return None


def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILS
    PASS_COUNT = 0
    FAIL_COUNT = 0
    CRITICAL_FAILS = []

    # ------------------------------------------------------------------ Excel
    excel_path = os.path.join(agent_workspace, "Support_Benchmark_Report.xlsx")
    check("Support_Benchmark_Report.xlsx exists", os.path.exists(excel_path))

    our_perf_map = {}      # priority -> row dict for critical checks
    industry_map = {}      # metric -> row dict for critical checks

    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Support_Benchmark_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name not in wb.sheetnames:
                    continue
                ws = wb[sheet_name]
                gt_ws = gt_wb[sheet_name]
                gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                for h in gt_headers:
                    if h:
                        check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                data_rows = list(ws.iter_rows(min_row=2, values_only=True))
                min_rows = max(1, len(gt_rows) - 2)
                check(f"{sheet_name} has >= {min_rows} data rows", len(data_rows) >= min_rows, f"got {len(data_rows)}")

                header_map = {h: i for i, h in enumerate(headers)}

                # Structural cell comparison against groundtruth (non-critical).
                # Numeric cells -> tolerance compare; text cells -> we DO NOT compare
                # recommendation prose here (agent may legitimately write Russian);
                # only numeric / identifier text are spot-checked structurally.
                for ri in range(min(3, len(gt_rows), len(data_rows))):
                    gt_row = gt_rows[ri]
                    agent_row = data_rows[ri]
                    for ci, gt_h in enumerate(gt_headers):
                        if not gt_h or ci >= len(gt_row):
                            continue
                        gv = gt_row[ci]
                        agent_ci = header_map.get(gt_h)
                        if agent_ci is None or agent_ci >= len(agent_row):
                            continue
                        av = agent_row[agent_ci]
                        gf = safe_float(gv)
                        af = safe_float(av)
                        if gf is not None and af is not None:
                            tol = max(0.5, abs(gf) * 0.15)
                            check(f"{sheet_name} R{ri+2} {gt_h} ~{gf:.1f}",
                                  abs(gf - af) <= tol, f"got {af}")

                # Cache maps for critical checks.
                if sheet_name == "Our_Performance":
                    pc = header_map.get("priority")
                    if pc is not None:
                        for r in data_rows:
                            if pc < len(r) and r[pc] is not None:
                                our_perf_map[str(r[pc]).strip().lower()] = (r, header_map)
                if sheet_name == "Industry_Comparison":
                    mc = header_map.get("metric")
                    if mc is not None:
                        for r in data_rows:
                            if mc < len(r) and r[mc] is not None:
                                industry_map[str(r[mc]).strip().lower()] = (r, header_map)

    # ---------------------------------------------- CRITICAL: Industry_Comparison
    # Industry_Avg must match fetched API benchmarks exactly (proves real fetch),
    # and Status must reflect the correct comparison vs the live DWH values.
    #
    # Industry_Avg is the fetched API benchmark (frozen constants below). The
    # expected Status is DERIVED FROM THE LIVE WAREHOUSE (not hardcoded), so the
    # check always reflects the real SUPPORT_CENTER ticket data: ours-vs-benchmark
    # with the correct comparison direction (lower-is-better for response time,
    # higher-is-better for resolution rate / CSAT / tickets-per-agent). This keeps
    # the check data-faithful even if the seed's STATUS / CSAT distribution changes.
    industry_avg_api = {
        "avg_response_time": 5.2,
        "resolution_rate": 87.5,
        "customer_satisfaction": 4.1,
        "tickets_per_agent": 3500,
    }
    # lower_is_better=True -> "above" (beats benchmark) when our value is LOWER.
    metric_lower_is_better = {
        "avg_response_time": True,
        "resolution_rate": False,
        "customer_satisfaction": False,
        "tickets_per_agent": False,
    }

    def derive_our_values():
        """Compute our actual metrics from the live SUPPORT_CENTER warehouse."""
        conn = get_conn()
        cur = conn.cursor()
        T = 'sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"'
        cur.execute(f'SELECT AVG("RESPONSE_TIME_HOURS") FROM {T}')
        avg_resp = cur.fetchone()[0]
        cur.execute(
            f'SELECT 100.0*COUNT(*) FILTER (WHERE "STATUS"=\'Resolved\')'
            f'/NULLIF(COUNT(*),0) FROM {T}')
        res_rate = cur.fetchone()[0]
        cur.execute(f'SELECT AVG("CUSTOMER_SATISFACTION") FROM {T}')
        csat = cur.fetchone()[0]
        cur.execute(f'SELECT COUNT(*) FROM {T}')
        total_tickets = cur.fetchone()[0]
        cur.execute('SELECT COUNT(*) FROM sf_data."SUPPORT_CENTER__PUBLIC__AGENTS"')
        agents = cur.fetchone()[0]
        cur.close()
        conn.close()
        return {
            "avg_response_time": safe_float(avg_resp),
            "resolution_rate": safe_float(res_rate),
            "customer_satisfaction": safe_float(csat),
            "tickets_per_agent": (safe_float(total_tickets) / safe_float(agents))
                if agents else None,
        }

    our_values = derive_our_values()

    expected_industry = {}
    for m, ind_avg in industry_avg_api.items():
        ours = our_values.get(m)
        if ours is None:
            status = "above" if not metric_lower_is_better[m] else "below"
        elif metric_lower_is_better[m]:
            status = "above" if ours < ind_avg else "below"
        else:
            status = "above" if ours > ind_avg else "below"
        expected_industry[m] = (ind_avg, status)
    for metric, (ind_avg, status) in expected_industry.items():
        entry = industry_map.get(metric)
        if entry is None:
            check(f"Industry_Comparison metric {metric} present", False,
                  f"have: {list(industry_map.keys())}", critical=True)
            continue
        row, hm = entry
        ia_ci = hm.get("industry_avg")
        st_ci = hm.get("status")
        ia = safe_float(row[ia_ci]) if (ia_ci is not None and ia_ci < len(row)) else None
        st = str(row[st_ci]).strip().lower() if (st_ci is not None and st_ci < len(row) and row[st_ci] is not None) else ""
        check(f"Industry_Comparison {metric} Industry_Avg == {ind_avg} (fetched API)",
              ia is not None and num_close(ia, ind_avg, rel_tol=0.001, abs_tol=0.05),
              f"got {ia}", critical=True)
        check(f"Industry_Comparison {metric} Status == {status}",
              status in st, f"got {st}", critical=True)

    # ---------------------------------------------- CRITICAL: Our_Performance
    # Medium is the largest tier (~15774, ~49.9%) and Critical == 0, confirming the
    # agent aggregated the real SUPPORT_CENTER ticket data instead of inventing numbers.
    med = our_perf_map.get("medium")
    crit = our_perf_map.get("critical")
    if med is not None:
        row, hm = med
        tc_ci = hm.get("ticket_count")
        pct_ci = hm.get("pct_of_total")
        tc = safe_float(row[tc_ci]) if (tc_ci is not None and tc_ci < len(row)) else None
        pct = safe_float(row[pct_ci]) if (pct_ci is not None and pct_ci < len(row)) else None
        check("Our_Performance Medium Ticket_Count ~15774",
              tc is not None and num_close(tc, 15774, rel_tol=0.1, abs_tol=50),
              f"got {tc}", critical=True)
        check("Our_Performance Medium Pct_of_Total ~49.9",
              pct is not None and num_close(pct, 49.9, rel_tol=0.1, abs_tol=2.0),
              f"got {pct}", critical=True)
    else:
        check("Our_Performance Medium row present", False,
              f"have: {list(our_perf_map.keys())}", critical=True)
    if crit is not None:
        row, hm = crit
        tc_ci = hm.get("ticket_count")
        tc = safe_float(row[tc_ci]) if (tc_ci is not None and tc_ci < len(row)) else None
        check("Our_Performance Critical Ticket_Count == 0",
              tc is not None and abs(tc) < 1, f"got {tc}", critical=True)
    else:
        check("Our_Performance Critical row present", False,
              f"have: {list(our_perf_map.keys())}", critical=True)

    # ---------------------------------------------- CRITICAL: pipeline artifact
    # benchmark_comparison.json produced by benchmark_analyzer.py, derived from both
    # support_metrics.json and industry_benchmarks.json (terminal+fetch+warehouse).
    cmp_path = os.path.join(agent_workspace, "benchmark_comparison.json")
    analyzer_path = os.path.join(agent_workspace, "benchmark_analyzer.py")
    check("benchmark_analyzer.py exists", os.path.exists(analyzer_path), critical=True)
    cmp_ok = False
    if os.path.exists(cmp_path):
        try:
            with open(cmp_path) as f:
                cmp_blob = f.read()
            cmp_low = cmp_blob.lower()
            # Must reference the comparison metrics derived from the two inputs.
            hits = sum(1 for k in ("response", "resolution", "satisfaction", "agent")
                       if k in cmp_low)
            # Must carry at least one fetched benchmark value.
            has_bench = any(v in cmp_blob for v in ("5.2", "87.5", "4.1", "3500"))
            cmp_ok = hits >= 3 and has_bench
        except Exception as e:
            cmp_ok = False
    check("benchmark_comparison.json has the four comparison metrics + fetched benchmarks",
          cmp_ok, "missing/incomplete benchmark_comparison.json", critical=True)
    # Inputs the analyzer was told to create.
    check("support_metrics.json exists",
          os.path.exists(os.path.join(agent_workspace, "support_metrics.json")))
    check("industry_benchmarks.json exists",
          os.path.exists(os.path.join(agent_workspace, "industry_benchmarks.json")))

    # ------------------------------------------------------------------ Word
    docx_path = os.path.join(agent_workspace, "Benchmark_Analysis.docx")
    check("Benchmark_Analysis.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        text_low = text.lower()
        check("Benchmark_Analysis.docx has content", len(text) > 50, f"text length: {len(text)}")
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        headings_blob = " || ".join(headings)
        # Required sections — accept RU or EN heading wording.
        required_sections = [
            ("Performance Overview", ["performance overview", "обзор производительности", "обзор показателей"]),
            ("Industry Comparison", ["industry comparison", "сравнение с отраслью", "отраслевое сравнение"]),
            ("Improvement Recommendations", ["improvement recommendations", "рекомендации по улучшению", "рекомендации"]),
        ]
        sections_present = True
        for label, alts in required_sections:
            found = any(a in headings_blob or a in text_low for a in alts)
            if not found:
                sections_present = False
            check(f"Benchmark_Analysis.docx has section \"{label}\" (RU/EN)", found,
                  f"headings: {headings[:6]}")
        # CRITICAL: all three sections present AND at least one concrete data point in body.
        body_low = text_low
        data_points = [
            "8.5", "8,5",                # our avg response time
            "5.2", "5,2",                # benchmark response time
            "82",                        # resolution rate
            "87.5", "87,5",              # benchmark resolution rate
            "3.8", "3,8",                # our CSAT
            "4.1", "4,1",                # benchmark CSAT
        ]
        has_data_point = any(dp in body_low for dp in data_points)
        check("Benchmark_Analysis.docx has all 3 sections AND a concrete data point",
              sections_present and has_data_point,
              f"sections={sections_present} data_point={has_data_point}", critical=True)

    # ------------------------------------------------ Python script (terminal usage)
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    check("Python analysis script exists", len(py_files) >= 1, f"found: {py_files}")

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}, Accuracy: {accuracy:.1f}%")
    if CRITICAL_FAILS:
        print(f"  CRITICAL FAILURES: {CRITICAL_FAILS}")
        return False, f"Critical check failed: {CRITICAL_FAILS}"

    overall = accuracy >= 70
    return overall, f"Passed {PASS_COUNT}/{total} checks (accuracy {accuracy:.1f}%)"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
