"""
Evaluation for yt-veritasium-kulinar-wellness-excel-gcal task.

CRITICAL checks (gate, any failure => FAIL regardless of accuracy):
  C1. Wellness_Plan.xlsx exists with Videos (>=4), Recipes (>=8), Weekly_Meal_Plan (==7) sheets.
  C2. Recipes/Weekly_Meal_Plan use real Kulinar (RU) dish names (agent actually queried the recipe KB).
  C3. Wellness_Guide.docx exists with >= 3 headings.
  C4. GCal has >= 3 "Wellness Check-in" events on Wednesdays in April 2026.

Then accuracy >= 70 gate over all (critical + non-critical) checks.
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []


def record(name, passed, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    tag = "CRITICAL " if critical else ""
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {tag}{name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {tag}{name}{msg}")
        if critical:
            CRITICAL_FAILED.append(name)


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


# --- Real Kulinar (RU) dish names: agent must have queried the kulinar recipe KB ---
def load_kulinar_names():
    fallback = {
        'Бефстроганов', 'Блины тонкие', 'Борщ', 'Ватрушки с творогом', 'Винегрет',
        'Голубцы', 'Греческий салат', 'Гречка с тушёнкой', 'Гречневая каша',
        'Грибной суп', 'Грибы маринованные', 'Жаркое в горшочках', 'Икра кабачковая',
        'Картофель отварной с укропом', 'Картофельное пюре', 'Квас домашний',
        'Кисель ягодный', 'Компот из сухофруктов', 'Котлеты домашние', 'Крабовый салат',
        'Кулебяка с капустой и яйцом', 'Куриный бульон с лапшой', 'Курица в сметане',
        'Медовик', 'Морс клюквенный', 'Наполеон', 'Окрошка', 'Пасха творожная',
        'Пельмени домашние', 'Перловая каша', 'Пирожки с капустой жареные',
        'Пирожки с мясом печёные', 'Плов узбекский', 'Рассольник', 'Расстегаи с рыбой',
        'Рис отварной', 'Рыба запечённая по-русски', 'Салат Мимоза', 'Салат Оливье',
        'Салат с курицей и грибами', 'Сало солёное', 'Сбитень', 'Сельдь под шубой',
        'Селёдка с луком', 'Солянка мясная', 'Сырники', 'Уха', 'Холодец',
        'Цыплёнок табака', 'Щи из квашеной капусты',
    }
    candidates = [
        "/app/local_servers/kulinar-mcp/src/data/all_recipes.json",
        os.path.join(os.path.dirname(__file__),
                     "../../../../local_servers/kulinar-mcp/src/data/all_recipes.json"),
    ]
    for p in candidates:
        try:
            with open(p, encoding="utf-8") as f:
                data = json.load(f)
            names = {r["name"] for r in data if r.get("name")}
            if names:
                return names
        except Exception:
            continue
    return fallback


KULINAR_NAMES = load_kulinar_names()
KULINAR_NAMES_LOWER = {n.strip().lower() for n in KULINAR_NAMES}


def check_excel(agent_workspace, groundtruth_workspace="."):
    print("\n=== Check 1-3: Wellness_Plan.xlsx ===")
    xlsx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".xlsx") and ("wellness" in fname.lower() or "plan" in fname.lower()):
            xlsx_path = os.path.join(agent_workspace, fname)
            break
    if not xlsx_path:
        for fname in os.listdir(agent_workspace):
            if fname.lower().endswith(".xlsx"):
                xlsx_path = os.path.join(agent_workspace, fname)
                break

    record("Wellness_Plan.xlsx exists", xlsx_path is not None,
           f"No matching xlsx in {agent_workspace}", critical=True)

    if not xlsx_path:
        for chk in ["Videos sheet >= 4 rows", "Recipes sheet >= 8 rows", "Weekly_Meal_Plan sheet = 7 rows"]:
            record(chk, False, "xlsx not found", critical=True)
        record("Recipes use real Kulinar dish names", False, "xlsx not found", critical=True)
        return

    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_path)

        # Videos sheet
        videos_sheet = None
        for name in wb.sheetnames:
            if "video" in name.lower():
                videos_sheet = wb[name]
                break
        if not videos_sheet and wb.sheetnames:
            videos_sheet = wb[wb.sheetnames[0]]

        if videos_sheet:
            data_rows = [r for r in videos_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Videos sheet has >= 4 data rows", len(data_rows) >= 4,
                   f"Found {len(data_rows)} rows", critical=True)
        else:
            record("Videos sheet has >= 4 data rows", False, "No Videos sheet", critical=True)

        # Recipes sheet
        recipe_sheet = None
        for name in wb.sheetnames:
            if "recipe" in name.lower():
                recipe_sheet = wb[name]
                break
        record("Recipes sheet exists", recipe_sheet is not None, f"Sheets: {wb.sheetnames}")

        recipe_rows = []
        if recipe_sheet:
            recipe_rows = [r for r in recipe_sheet.iter_rows(min_row=2, values_only=True)
                           if any(c is not None for c in r)]
            record("Recipes sheet has >= 8 data rows", len(recipe_rows) >= 8,
                   f"Found {len(recipe_rows)} rows", critical=True)
        else:
            record("Recipes sheet has >= 8 data rows", False, "Sheet not found", critical=True)

        # CRITICAL: Recipes must be real Kulinar (RU) dish names (not placeholder/Chinese/invented).
        recipe_names = [str(r[0]).strip() for r in recipe_rows if r and r[0] is not None]
        matched = [n for n in recipe_names if n.lower() in KULINAR_NAMES_LOWER]
        record("Recipes use real Kulinar dish names (>= 5 matches)",
               len(matched) >= 5,
               f"Matched {len(matched)}/{len(recipe_names)} against Kulinar KB; sample={recipe_names[:5]}",
               critical=True)

        # Weekly_Meal_Plan sheet
        meal_sheet = None
        for name in wb.sheetnames:
            if "meal" in name.lower() or "weekly" in name.lower():
                meal_sheet = wb[name]
                break
        record("Weekly_Meal_Plan sheet exists", meal_sheet is not None, f"Sheets: {wb.sheetnames}")

        if meal_sheet:
            meal_rows = [r for r in meal_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Weekly_Meal_Plan sheet has exactly 7 rows", len(meal_rows) == 7,
                   f"Found {len(meal_rows)} rows (expected 7)", critical=True)
            # meal plan recipe names should also be Kulinar dishes
            mp_names = []
            for r in meal_rows:
                for ci in (1, 2, 3):  # Breakfast, Lunch, Dinner columns
                    if ci < len(r) and r[ci] is not None:
                        mp_names.append(str(r[ci]).strip())
            mp_matched = [n for n in mp_names if n.lower() in KULINAR_NAMES_LOWER]
            record("Weekly_Meal_Plan uses real Kulinar dishes (>= 10 cells)",
                   len(mp_matched) >= 10,
                   f"Matched {len(mp_matched)}/{len(mp_names)} meal-plan cells")
        else:
            record("Weekly_Meal_Plan sheet has exactly 7 rows", False, "Sheet not found", critical=True)

        # --- Groundtruth XLSX value comparison (non-critical, shape/row-count + sampled rows) ---
        gt_path = os.path.join(groundtruth_workspace, "Wellness_Plan.xlsx")
        if os.path.isfile(gt_path):
            gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
            for gt_sname in gt_wb.sheetnames:
                gt_ws = gt_wb[gt_sname]
                a_ws = None
                for asn in wb.sheetnames:
                    if asn.strip().lower() == gt_sname.strip().lower():
                        a_ws = wb[asn]
                        break
                if a_ws is None:
                    record(f"GT sheet '{gt_sname}' exists in agent xlsx", False, f"Available: {wb.sheetnames}")
                    continue
                gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                a_rows = [r for r in a_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                record(f"GT '{gt_sname}' row count", len(a_rows) == len(gt_rows),
                       f"Expected {len(gt_rows)}, got {len(a_rows)}")
            gt_wb.close()

    except Exception as e:
        record("Wellness_Plan.xlsx readable", False, str(e), critical=True)


def check_word(agent_workspace):
    print("\n=== Check 4-6: Wellness_Guide.docx ===")
    docx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".docx") and ("wellness" in fname.lower() or "guide" in fname.lower()):
            docx_path = os.path.join(agent_workspace, fname)
            break
    if not docx_path:
        for fname in os.listdir(agent_workspace):
            if fname.lower().endswith(".docx"):
                docx_path = os.path.join(agent_workspace, fname)
                break

    record("Wellness_Guide.docx exists", docx_path is not None,
           f"No wellness/guide docx in {agent_workspace}", critical=True)

    if not docx_path:
        record("Word doc has >= 3 headings", False, "docx not found", critical=True)
        record("Word doc contains wellness keywords", False, "docx not found")
        return

    try:
        from docx import Document
        doc = Document(docx_path)
        headings = [p for p in doc.paragraphs if p.style.name.lower().startswith("heading")]
        record("Word doc has >= 3 headings", len(headings) >= 3,
               f"Found {len(headings)} headings", critical=True)

        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        # English section headings stay English; prose is Russian. Accept either language.
        keywords = [
            # English (kept identifiers / section headings)
            "wellness", "recipe", "meal", "science", "veritasium",
            # Russian prose terms
            "здоров", "питани", "рецепт", "велнес", "наук", "рацион", "блюд",
        ]
        found = [k for k in keywords if k in full_text]
        record("Word doc contains health/wellness/recipe keywords (RU or EN)",
               len(found) >= 3, f"Found keywords: {found}")
    except Exception as e:
        record("Word doc has >= 3 headings", False, str(e), critical=True)
        record("Word doc contains wellness keywords", False, str(e))


def check_gcal():
    print("\n=== Check 7: GCal Wellness Check-in Events in April 2026 ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime FROM gcal.events
        WHERE start_datetime >= '2026-04-01' AND start_datetime < '2026-05-01'
        AND summary ILIKE '%wellness%'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    # Wednesday-only wellness events in April 2026 (the 4 required check-ins land on Wednesdays).
    wed_events = [e for e in events if e[1] and e[1].weekday() == 2]
    record("GCal has >= 3 Wellness Check-in events on Wednesdays in April 2026",
           len(wed_events) >= 3,
           f"Found {len(events)} wellness events ({len(wed_events)} on Wednesdays)",
           critical=True)


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_excel(args.agent_workspace, args.groundtruth_workspace)
    check_word(args.agent_workspace)
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_failed": CRITICAL_FAILED,
    }
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    # CRITICAL gate: any critical failure => hard FAIL regardless of accuracy.
    if CRITICAL_FAILED:
        print(f"\nFAIL: {len(CRITICAL_FAILED)} CRITICAL check(s) failed: {CRITICAL_FAILED}")
        sys.exit(1)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
