"""Генератор эталона (groundtruth) для terminal-yf-insales-excel-word-notion (RU).

Все финансовые значения берутся из живой схемы moex.* (GLDRUB_TOM, OZON.ME),
товарные агрегаты — из wc.products. Логика рекомендаций rescale'нута под
актуальный масштаб золота (GLDRUB_TOM, изменение ~27% — НИЖЕ порога 50%).

Используется и эталонными файлами, и evaluation/main.py (compute_finance/
compute_products), чтобы пороги и значения были выведены из БД, а не зашиты.
"""
import os
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"),
          port=int(os.environ.get("PGPORT", "5432")),
          dbname=os.environ.get("PGDATABASE", "cowork_gym"),
          user="eigent", password="camel")

COMMODITY = "GLDRUB_TOM"   # золото (RUB/грамм)
CONSUMER = "OZON.ME"       # Ozon, потребительский циклический сектор
GOLD_SIGNIFICANT_PCT = 50.0
STANDARD_MARGIN = 40
REDUCED_MARGIN = 35


def get_conn():
    return psycopg2.connect(**DB)


def compute_finance():
    """Возвращает dict symbol -> {avg_price, price_change_pct, volatility, name}."""
    conn = get_conn()
    cur = conn.cursor()
    out = {}
    for sym in (COMMODITY, CONSUMER):
        cur.execute("""
            SELECT close, date FROM moex.stock_prices
            WHERE symbol=%s ORDER BY date
        """, (sym,))
        rows = cur.fetchall()
        closes = [float(r[0]) for r in rows]
        n = len(closes)
        avg = sum(closes) / n
        var = sum((c - avg) ** 2 for c in closes) / n  # population std
        vol = var ** 0.5
        pct = (closes[-1] - closes[0]) / closes[0] * 100.0
        cur.execute("SELECT data->>'shortName' FROM moex.stock_info WHERE symbol=%s", (sym,))
        r = cur.fetchone()
        name = r[0] if r and r[0] else sym
        out[sym] = {
            "avg_price": round(avg, 2),
            "price_change_pct": round(pct, 2),
            "volatility": round(vol, 2),
            "name": name,
        }
    cur.close()
    conn.close()
    return out


def compute_products():
    """Агрегаты по категориям из wc.products (первичная категория товара).

    Возвращает список dict: category, avg_price, avg_cost_estimate,
    margin_pct, total_sales — отсортировано по category.
    """
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("""
        SELECT (categories->0->>'name') AS cat,
               AVG(regular_price::numeric) AS avg_price,
               SUM(total_sales) AS sales
        FROM wc.products
        WHERE status='publish' AND categories->0->>'name' IS NOT NULL
        GROUP BY cat ORDER BY cat
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()
    cats = []
    for cat, avg_price, sales in rows:
        ap = round(float(avg_price), 2)
        cats.append({
            "category": cat,
            "avg_price": ap,
            "avg_cost_estimate": round(ap * 0.6, 2),
            "margin_pct": STANDARD_MARGIN,
            "total_sales": int(sales or 0),
        })
    return cats


def gold_significant():
    """True, если изменение золота превышает порог значимости (50%)."""
    fin = compute_finance()
    return fin[COMMODITY]["price_change_pct"] > GOLD_SIGNIFICANT_PCT


def target_margin_and_action():
    if gold_significant():
        return REDUCED_MARGIN, "Monitor costs and consider price adjustment"
    return STANDARD_MARGIN, "Maintain current pricing"


# --------------------------------------------------------------------------- #
# Генерация файлов эталона
# --------------------------------------------------------------------------- #
def build_files(out_dir):
    import openpyxl
    from docx import Document

    fin = compute_finance()
    prods = compute_products()
    g = fin[COMMODITY]
    c = fin[CONSUMER]
    tgt, action = target_margin_and_action()

    os.makedirs(out_dir, exist_ok=True)

    # --- Excel ---
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Stock_Trends"
    ws.append(["symbol", "name", "avg_price", "price_change_pct", "volatility"])
    ws.append([COMMODITY, g["name"], g["avg_price"], g["price_change_pct"], g["volatility"]])
    ws.append([CONSUMER, c["name"], c["avg_price"], c["price_change_pct"], c["volatility"]])

    ws2 = wb.create_sheet("Product_Margins")
    ws2.append(["category", "avg_price", "avg_cost_estimate", "margin_pct", "total_sales"])
    for p in prods:
        ws2.append([p["category"], p["avg_price"], p["avg_cost_estimate"],
                    p["margin_pct"], p["total_sales"]])

    ws3 = wb.create_sheet("Correlation_Analysis")
    ws3.append(["factor_pair", "correlation_description", "implication"])
    ws3.append([
        "Gold vs Product Costs",
        f"Цена золота (GLDRUB_TOM) изменилась на {g['price_change_pct']}% за период, "
        f"что указывает на умеренное давление сырьевых издержек.",
        "Входные затраты под контролем; маржу можно сохранить на текущем уровне.",
    ])
    ws3.append([
        "OZON.ME vs Consumer Spending",
        f"Акции OZON.ME изменились на {c['price_change_pct']}% за период, "
        f"что говорит об умеренной уверенности потребительского спроса.",
        "Потребительский спрос стабилен; текущие цены выглядят устойчивыми.",
    ])

    ws4 = wb.create_sheet("Strategic_Recommendations")
    ws4.append(["category", "current_margin", "target_margin", "action"])
    for p in prods:
        ws4.append([p["category"], STANDARD_MARGIN, tgt, action])

    xlsx_path = os.path.join(out_dir, "Commodity_Impact_Analysis.xlsx")
    wb.save(xlsx_path)

    # --- Word ---
    doc = Document()
    doc.add_heading("Commodity Impact and Pricing Strategy Memo", 0)
    doc.add_heading("Market Overview", level=1)
    doc.add_paragraph(
        f"Золото (GLDRUB_TOM) изменилось примерно на {g['price_change_pct']}% за период "
        f"анализа: средняя цена закрытия {g['avg_price']} RUB/грамм, волатильность "
        f"(стандартное отклонение) {g['volatility']}. Это умеренный рост, который не "
        f"превышает порог значимости в 50% и указывает на ограниченное давление "
        f"сырьевых издержек."
    )
    doc.add_paragraph(
        f"Акции Ozon (OZON.ME), используемые как индикатор уверенности потребительского "
        f"циклического спроса, изменились примерно на {c['price_change_pct']}% за тот же "
        f"период: средняя цена {c['avg_price']}, волатильность {c['volatility']}. Это "
        f"говорит о том, что потребительский спрос остаётся относительно здоровым."
    )
    doc.add_heading("Product Margin Analysis", level=1)
    doc.add_paragraph(
        "Маржа по всем категориям товаров сейчас оценивается в 40% (себестоимость — 60% "
        "от розничной цены). Каталог охватывает категории: "
        + ", ".join(p["category"] for p in prods) + ". "
        "Наибольший объём продаж приходится на категорию «Электроника»."
    )
    doc.add_heading("Strategic Recommendations", level=1)
    if gold_significant():
        doc.add_paragraph(
            f"Поскольку цены на золото выросли более чем на 50% (изменение "
            f"{g['price_change_pct']}%), рекомендуется снизить целевую маржу до 35% по "
            f"всем категориям для компенсации роста издержек."
        )
    else:
        doc.add_paragraph(
            f"Поскольку цены на золото выросли умеренно ({g['price_change_pct']}%, ниже "
            f"порога значимости 50%), рекомендуется сохранить целевую маржу на уровне 40% "
            f"по всем категориям (Maintain current pricing). Положительная динамика "
            f"OZON.ME подтверждает устойчивость потребительского спроса."
        )
    docx_path = os.path.join(out_dir, "Pricing_Strategy_Memo.docx")
    doc.save(docx_path)

    # --- correlation_analysis.py (эталонный скрипт) ---
    script = os.path.join(out_dir, "correlation_analysis.py")
    with open(script, "w", encoding="utf-8") as f:
        f.write(
            "# Эталонный correlation_analysis.py\n"
            f"GOLD_CHANGE_PCT = {g['price_change_pct']}\n"
            f"OZON_CHANGE_PCT = {c['price_change_pct']}\n"
            "if GOLD_CHANGE_PCT > 50:\n"
            "    print('Market conditions suggest CONTRACTING margins (gold inflation).')\n"
            "else:\n"
            "    print('Market conditions suggest MAINTAINING margins (limited cost pressure).')\n"
        )

    return xlsx_path, docx_path, script


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    gt = os.path.join(os.path.dirname(here), "groundtruth_workspace")
    paths = build_files(gt)
    print("Сгенерирован эталон:")
    for p in paths:
        print(" ", p)
    print("Финансы:", compute_finance())
    print("Gold significant (>50%)?", gold_significant())
    print("Target margin/action:", target_margin_and_action())
