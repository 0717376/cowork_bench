"""Evaluation for playwright-insales-market-trends-gsheet-word-email.

Что проверяем:
1. Google Sheet "Market Trend Analysis 2026" с 3 листами (Industry Trends,
   Internal Performance, Gap Analysis); 6 категорий присутствуют как значения.
2. Market_Insights_Report.docx: упоминает категории, рыночные/трендовые разделы,
   рекомендации, зарождающиеся тренды с конкретными примерами.
3. Письмо на product-team@company.com с темой про market trends; тело письма
   ссылается на возможность из gap-анализа и сильнейшую внутреннюю категорию.
4. Market_Trend_Analysis.xlsx — зеркало Google Sheet (3 листа).

Категории — это realia, общие для mock-панели, PDF-гайда и каталога InSales
(wc.*), и центрально руссифицированы. Поэтому здесь они проверяются на РУССКОМ:
Аудио / Камеры / Электроника / Бытовая техника / ТВ и домашний кинотеатр / Часы.

CRITICAL-чеки отражают СУТЬ задачи (корректные значения с панели + вычисление
gap + порог Market_Position + вывод из живого каталога). Любой их провал =>
итог FAIL независимо от accuracy.
"""
import argparse
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

# Эталонные значения с mock-панели (русские категории <-> числа дашборда).
# Industry_Sentiment и Growth_Rate_Pct — это то, что агент ОБЯЗАН извлечь
# через playwright; market size — для справки.
DASHBOARD = {
    "аудио":                    {"growth": 8.2,  "market": 12400, "sentiment": 9.2},
    "камеры":                   {"growth": -2.1, "market": 6800,  "sentiment": 8.5},
    "электроника":              {"growth": 12.5, "market": 45000, "sentiment": 9.6},
    "бытовая техника":          {"growth": 5.4,  "market": 28000, "sentiment": 8.8},
    "тв и домашний кинотеатр":  {"growth": 3.8,  "market": 18500, "sentiment": 9.0},
    "часы":                     {"growth": 15.3, "market": 9200,  "sentiment": 9.5},
}

# Эталонные внутренние показатели из ЖИВОГО каталога InSales (wc.*),
# руссифицированные. Avg_Rating / Total_Revenue / Product_Count привязывают
# Internal Performance к реальным данным каталога (а не к правдоподобным
# выдуманным строкам). Значения совпадают с groundtruth_workspace.
#
# Канонический вывод из ТЕКУЩЕГО сида wc.products (db/init.sql.gz), где каждый
# продукт относится к одной из 6 главных категорий через jsonb categories[].name:
#   count    = COUNT(*) по продуктам категории
#   revenue  = SUM(price * total_sales) по продуктам категории
#   rating   = SUM(average_rating * rating_count) / SUM(rating_count)  (взвеш.)
# Воспроизводимо агентом через woo_products_list (price/total_sales/
# average_rating/rating_count/categories).
CATALOG = {
    "аудио":                    {"count": 15, "revenue": 134825.65, "rating": 4.48},
    "камеры":                   {"count": 10, "revenue": 23323.13,  "rating": 4.69},
    "электроника":              {"count": 30, "revenue": 192516.44, "rating": 4.57},
    "бытовая техника":          {"count": 8,  "revenue": 10897.97,  "rating": 4.72},
    "тв и домашний кинотеатр":  {"count": 13, "revenue": 449220.62, "rating": 4.78},
    "часы":                     {"count": 6,  "revenue": 11469.72,  "rating": 4.43},
}

# Русские названия категорий, ожидаемые в ячейках/тексте.
RU_CATEGORIES = ["аудио", "камеры", "электроника",
                 "бытовая техника", "тв и домашний кинотеатр", "часы"]
# Допускаем и английские эквиваленты в свободном тексте Word (агент мог
# смешать), но НЕ требуем их.
EN_CATEGORIES = ["audio", "cameras", "electronics",
                 "home appliances", "tv & home theater", "watches"]

# Чеки, провал которых = содержательное невыполнение задачи.
CRITICAL_CHECKS = {
    "GapAnalysis: Industry_Sentiment и Gap соответствуют данным панели",
    "GapAnalysis: Market_Position соответствует правилу порогов",
    "IndustryTrends: темпы роста совпадают с панелью",
    "InternalPerformance: данные выведены из живого каталога InSales",
    "Email: тело ссылается на gap-возможность и сильную категорию",
}

CRITICAL_FAILS = []


def norm(s):
    if s is None:
        return ""
    return str(s).strip().lower().replace("ё", "е")


def to_float(v):
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().replace(",", "").replace("%", "").replace("$", "")
    try:
        return float(s)
    except ValueError:
        return None


def num_close(a, b, tol=1.0):
    fa, fb = to_float(a), to_float(b)
    if fa is None or fb is None:
        return False
    return abs(fa - fb) <= tol


def cat_key(v):
    """Свести значение ячейки к ключу категории (RU или EN -> RU-ключ)."""
    n = norm(v)
    for i, en in enumerate(EN_CATEGORIES):
        if n == en:
            return RU_CATEGORIES[i]
    return n if n in DASHBOARD else None


def expected_position(gap):
    if gap is None:
        return None
    if gap >= 0:
        return "strong"
    if -1 <= gap < 0:
        return "needs attention"
    return "critical"


def read_sheet_rows(cur, ss_id, name_substr):
    """Вернуть (header_lower->col_index, list-of-rowdicts) для листа,
    название которого содержит name_substr. rowdict: col_index -> value."""
    cur.execute("""
        SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s
    """, (ss_id,))
    target = None
    for sid, title in cur.fetchall():
        if name_substr in norm(title):
            target = sid
            break
    if target is None:
        return None, None
    cur.execute("""
        SELECT row_index, col_index, value FROM gsheet.cells
        WHERE spreadsheet_id = %s AND sheet_id = %s
        ORDER BY row_index, col_index
    """, (ss_id, target))
    grid = {}
    for ri, ci, val in cur.fetchall():
        grid.setdefault(ri, {})[ci] = val
    if not grid:
        return {}, []
    header_ri = min(grid.keys())
    header = {norm(v): ci for ci, v in grid[header_ri].items()}
    rows = [grid[ri] for ri in sorted(grid) if ri != header_ri]
    return header, rows


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    agent_ws = args.agent_workspace or task_root

    all_errors = []

    def fail(name, detail=""):
        msg = f"{name}" + (f": {detail}" if detail else "")
        marker = " [CRITICAL]" if name in CRITICAL_CHECKS else ""
        all_errors.append(marker + " " + msg if marker else msg)
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILS.append(name)

    # ---------------- Google Sheet ----------------
    print("Checking Google Sheet...")
    ss_id = None
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        cur.execute("""
            SELECT id, title FROM gsheet.spreadsheets
            WHERE title ILIKE '%market%trend%'
        """)
        rows = cur.fetchall()
        if not rows:
            fail("Google Sheet 'Market Trend Analysis 2026' not found")
        else:
            ss_id = rows[0][0]
            cur.execute("""
                SELECT title FROM gsheet.sheets WHERE spreadsheet_id = %s
                ORDER BY "index"
            """, (ss_id,))
            sheet_names = [norm(r[0]) for r in cur.fetchall()]
            if not any("industry" in s for s in sheet_names):
                fail("Sheet 'Industry Trends' not found in GSheet")
            if not any("internal" in s for s in sheet_names):
                fail("Sheet 'Internal Performance' not found in GSheet")
            if not any("gap" in s for s in sheet_names):
                fail("Sheet 'Gap Analysis' not found in GSheet")

            # 6 категорий присутствуют как значения (RU; EN допустимо).
            cur.execute("""
                SELECT DISTINCT LOWER(value) FROM gsheet.cells c
                JOIN gsheet.sheets s ON c.sheet_id = s.id
                WHERE s.spreadsheet_id = %s
            """, (ss_id,))
            vals = {norm(r[0]) for r in cur.fetchall()}
            found = sum(1 for i, ru in enumerate(RU_CATEGORIES)
                        if ru in vals or EN_CATEGORIES[i] in vals)
            if found < 6:
                fail(f"GSheet: found {found}/6 categories in cells")

            # ---- Industry Trends: темпы роста (CRITICAL) ----
            ihdr, irows = read_sheet_rows(cur, ss_id, "industry")
            chk = "IndustryTrends: темпы роста совпадают с панелью"
            if not ihdr:
                fail(chk, "лист Industry Trends не читается")
            else:
                gci = ihdr.get("growth_rate_pct")
                cci = (ihdr.get("category") if "category" in ihdr
                       else next((ci for k, ci in ihdr.items() if "categor" in k), None))
                if gci is None or cci is None:
                    fail(chk, "нет колонок Category/Growth_Rate_Pct")
                else:
                    matched = 0
                    for row in irows:
                        k = cat_key(row.get(cci))
                        if k in DASHBOARD and num_close(row.get(gci),
                                                        DASHBOARD[k]["growth"], 0.2):
                            matched += 1
                    if matched < 4:
                        fail(chk, f"совпало {matched}/6 темпов роста")

            # ---- Internal Performance: живой каталог (CRITICAL) ----
            phdr, prows = read_sheet_rows(cur, ss_id, "internal")
            chk = "InternalPerformance: данные выведены из живого каталога InSales"
            if not phdr:
                fail(chk, "лист Internal Performance не читается")
            else:
                cci = next((ci for k, ci in phdr.items() if "categor" in k), None)
                pcci = next((ci for k, ci in phdr.items()
                             if "product_count" in k or k == "product_count"), None)
                rev_ci = next((ci for k, ci in phdr.items()
                               if "total_revenue" in k or "revenue" in k), None)
                rat_ci = next((ci for k, ci in phdr.items()
                               if "avg_rating" in k or "rating" in k), None)
                if (cci is None or pcci is None or rev_ci is None
                        or rat_ci is None):
                    fail(chk, "нет колонок Category/Product_Count/"
                              "Total_Revenue/Avg_Rating")
                else:
                    # Привязываем к реальному каталогу: Avg_Rating + Total_Revenue
                    # (+ Product_Count) должны совпасть с эталонными значениями
                    # InSales хотя бы для 4 категорий. Правдоподобных выдуманных
                    # строк недостаточно.
                    good = 0
                    for row in prows:
                        k = cat_key(row.get(cci))
                        if k not in CATALOG:
                            continue
                        gt = CATALOG[k]
                        rating_ok = num_close(row.get(rat_ci), gt["rating"], 0.1)
                        rev_ok = num_close(row.get(rev_ci), gt["revenue"], 1.0)
                        count_ok = num_close(row.get(pcci), gt["count"], 0)
                        if rating_ok and rev_ok and count_ok:
                            good += 1
                    if good < 4:
                        fail(chk, f"только {good}/6 категорий совпали с живым "
                                  "каталогом InSales (Avg_Rating+Total_Revenue+"
                                  "Product_Count)")

            # ---- Gap Analysis: сентимент+gap и Market_Position (CRITICAL) ----
            ghdr, grows = read_sheet_rows(cur, ss_id, "gap")
            chk_g = "GapAnalysis: Industry_Sentiment и Gap соответствуют данным панели"
            chk_p = "GapAnalysis: Market_Position соответствует правилу порогов"
            if not ghdr:
                fail(chk_g, "лист Gap Analysis не читается")
                fail(chk_p, "лист Gap Analysis не читается")
            else:
                cci = next((ci for k, ci in ghdr.items() if "categor" in k), None)
                r10 = next((ci for k, ci in ghdr.items()
                            if "internal_rating" in k or "10pt" in k or "10" in k), None)
                sci = next((ci for k, ci in ghdr.items() if "sentiment" in k), None)
                gci = ghdr.get("gap") or next(
                    (ci for k, ci in ghdr.items() if k == "gap"), None)
                mci = next((ci for k, ci in ghdr.items()
                            if "position" in k or "market_position" in k), None)
                if None in (cci, sci, gci, mci, r10):
                    fail(chk_g, "не хватает колонок Gap Analysis")
                    fail(chk_p, "не хватает колонок Gap Analysis")
                else:
                    sent_ok, pos_ok = 0, 0
                    for row in grows:
                        k = cat_key(row.get(cci))
                        if k not in DASHBOARD:
                            continue
                        exp_sent = DASHBOARD[k]["sentiment"]
                        sent_match = num_close(row.get(sci), exp_sent, 0.15)
                        ir = to_float(row.get(r10))
                        gap = to_float(row.get(gci))
                        gap_match = (ir is not None and gap is not None
                                     and num_close(gap, ir - exp_sent, 0.15))
                        if sent_match and gap_match:
                            sent_ok += 1
                        if gap is not None:
                            exp_pos = expected_position(gap)
                            if exp_pos and exp_pos in norm(row.get(mci)):
                                pos_ok += 1
                    if sent_ok < 4:
                        fail(chk_g, f"совпало {sent_ok}/6 строк "
                                    "(Industry_Sentiment + Gap)")
                    if pos_ok < 4:
                        fail(chk_p, f"корректный Market_Position у {pos_ok}/6 строк")

        cur.close()
        conn.close()
    except Exception as e:
        all_errors.append(f"Error checking GSheet: {e}")

    # ---------------- Word document ----------------
    print("Checking Word document...")
    doc_path = os.path.join(agent_ws, "Market_Insights_Report.docx")
    if not os.path.exists(doc_path):
        all_errors.append("Market_Insights_Report.docx not found")
    else:
        try:
            from docx import Document
            doc = Document(doc_path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += "\n" + cell.text
            full_lower = norm(full_text)

            found_cats = sum(1 for i in range(len(RU_CATEGORIES))
                             if RU_CATEGORIES[i] in full_lower
                             or EN_CATEGORIES[i] in full_lower)
            if found_cats < 4:
                all_errors.append(f"Word doc mentions only {found_cats}/6 categories")

            if not any(w in full_lower for w in
                       ("trend", "market", "тренд", "рынок", "рыноч")):
                all_errors.append("Word doc missing market/trend references")

            if not any(w in full_lower for w in
                       ("recommendation", "strategic", "рекомендац", "стратег")):
                all_errors.append("Word doc missing recommendations section")

            if not any(w in full_lower for w in ("emerging", "зарожда", "зарождающ")):
                all_errors.append("Word doc missing emerging trends section")

            # Конкретный зарождающийся тренд, привязанный к рекомендации (RU+EN).
            concrete = ("smartwatch", "smart watch", "health", "smart speaker",
                        "ai speaker", "ai-powered", "wearable", "8k",
                        "умные час", "умных час", "здоров", "умные колонк",
                        "умных колонок", "колонк", "носим", "ии", "8к",
                        "беспроводн", "умный дом", "умного дома")
            if not any(c in full_lower for c in concrete):
                all_errors.append("Word doc emerging section lacks a concrete "
                                  "emerging trend tied to a recommendation")
        except Exception as e:
            all_errors.append(f"Error reading Word doc: {e}")

    # ---------------- Email ----------------
    print("Checking email...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT subject, COALESCE(body_text, '') FROM email.messages
            WHERE to_addr::text ILIKE '%product-team@company.com%'
            AND subject ILIKE '%market%trend%'
        """)
        msgs = cur.fetchall()
        if not msgs:
            all_errors.append("No email sent to product-team@company.com "
                              "about market trends")
        else:
            # Тело письма должно ссылаться на возможность из gap-анализа
            # И на сильнейшую внутреннюю категорию (CRITICAL).
            chk = "Email: тело ссылается на gap-возможность и сильную категорию"
            best = ""
            for subj, body in msgs:
                if norm(body) and len(norm(body)) > len(norm(best)):
                    best = body
            bl = norm(best)
            # Сильнейшая категория по выручке/позиции: Электроника / ТВ /
            # Камеры / Бытовая техника попадают в "Strong"; требуем упоминание
            # хотя бы одной gap-сильной категории И слова про возможность/инвест.
            cat_mentioned = any(ru in bl or EN_CATEGORIES[i] in bl
                                for i, ru in enumerate(RU_CATEGORIES))
            opp_word = any(w in bl for w in
                           ("возможност", "инвест", "opportunity", "invest",
                            "gap", "разрыв", "приоритет", "рекоменд",
                            "recommend", "сильн", "strong"))
            if not bl:
                fail(chk, "пустое тело письма")
            elif not (cat_mentioned and opp_word):
                fail(chk, "в теле нет связки категория + возможность/инвестиция")
        cur.close()
        conn.close()
    except Exception as e:
        all_errors.append(f"Error checking email: {e}")

    # ---------------- XLSX mirror ----------------
    print("Checking XLSX content...")
    xlsx_path = os.path.join(agent_ws, "Market_Trend_Analysis.xlsx")
    if not os.path.exists(xlsx_path):
        all_errors.append("Market_Trend_Analysis.xlsx not found")
    else:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            n_sheets = len(wb.worksheets)
            if n_sheets < 3:
                all_errors.append(f"XLSX has {n_sheets} sheets (need 3 mirror sheets)")
            for ws in wb.worksheets:
                rs = list(ws.iter_rows(values_only=True))
                if len(rs) < 2:
                    all_errors.append(
                        f"XLSX sheet '{ws.title}' has only {len(rs)} rows (need >= 2)")
            wb.close()
        except Exception as e:
            all_errors.append(f"Error reading XLSX: {e}")

    # ---------------- Result ----------------
    total_checks = 13  # ориентировочное число содержательных проверок
    n_err = len(all_errors)
    # accuracy как доля непровалённых проверок (грубая оценка)
    accuracy = max(0.0, 100.0 * (1 - n_err / float(total_checks)))

    print("\n=== Errors ===")
    for e in all_errors[:20]:
        print(f"  {e}")
    print(f"\naccuracy ~ {accuracy:.1f}  (errors={n_err})")

    if CRITICAL_FAILS:
        print(f"\n=== RESULT: FAIL — CRITICAL checks failed: {CRITICAL_FAILS} ===")
        sys.exit(1)

    if accuracy >= 70:
        print(f"\n=== RESULT: PASS (accuracy {accuracy:.1f}, no critical fail) ===")
        sys.exit(0)
    print(f"\n=== RESULT: FAIL (accuracy {accuracy:.1f} < 70) ===")
    sys.exit(1)


if __name__ == "__main__":
    main()
