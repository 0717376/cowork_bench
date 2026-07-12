"""
Evaluation for canvas-assignment-word-notion (russified, teamly swap).

The agent reads all assignments for course_id=7 (Креативные вычисления и культура
Fall 2014) honestly from the live Canvas MCP, builds Assignment_Guide.docx
(English heading + course code + intro line, a 4-column table sorted by due
date, then 'Total Assignments: N' and 'Total Points: X' lines), and creates a
teamly knowledge-base page titled "CCC-2014J Assignment Overview".

Critical checks (see CRITICAL_CHECKS): any failure => overall FAIL regardless
of accuracy. Otherwise pass threshold: accuracy >= 70%.

RU note: the agent legitimately writes Russian prose around the English literal
markers. RU keyword checks search .lower() ORIGINAL text (never normalize()).
"""

import argparse
import json
import os
import re
import sys

import psycopg2

DB = dict(
    host=os.environ.get("PGHOST", "localhost"),
    port=5432,
    dbname=os.environ.get("PGDATABASE", "cowork_gym"),
    user="eigent",
    password="camel",
)

# Expected assignments for course_id=7 (Креативные вычисления и культура Fall 2014)
EXPECTED_ASSIGNMENT_COUNT = 10
EXPECTED_TOTAL_POINTS = 300.0
EXPECTED_COURSE_NAME = "Креативные вычисления и культура (Осень 2014)"
EXPECTED_ASSIGNMENT_NAMES = [
    "CMA 24295", "CMA 24296", "CMA 24297", "CMA 24298",
    "TMA 24291", "TMA 24292", "TMA 24293", "TMA 24294",
    "Final Exam 24299", "Final Exam 40088",
]
FINAL_EXAM_NAMES = ["Final Exam 24299", "Final Exam 40088"]

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Semantic critical checks: any failure => overall FAIL regardless of accuracy.
CRITICAL_CHECKS = {
    "Table has 10 rows and >=8 expected assignment names",
    "Total Points line equals 300.0",
    "Total Assignments line equals 10",
    "Exactly 2 no-due-date rows and both Final Exams present with 100 points",
    "Teamly 'CCC-2014J Assignment Overview' page mentions course, count 10, points 300",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        d = (str(detail)[:300]) if detail else ""
        print(f"  [FAIL] {name}: {d}")


def find_value_near_keyword(text_lower, keywords, value_regex):
    """Find `value_regex` on a line that also contains one of `keywords`.

    Guards against substring false-positives (e.g. '1300' matching '300'):
    the number must appear adjacent to the keyword on the same line.
    Returns the matched value string or None.
    """
    for line in text_lower.splitlines():
        if any(k in line for k in keywords):
            m = re.search(value_regex, line)
            if m:
                return m.group(0)
    return None


def check_word_doc(agent_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Assignment_Guide.docx ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Assignment_Guide.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Not found: {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Heading stays English (literal marker, not russified). Structural.
    has_heading = False
    for p in doc.paragraphs:
        if "creative computing" in p.text.lower() and "assignment guide" in p.text.lower():
            has_heading = True
            break
    check("Document has correct heading", has_heading)

    # Course code stays English. Structural.
    full_text = "\n".join(p.text for p in doc.paragraphs)
    full_text_lower = full_text.lower()
    check("Document mentions CCC-2014J", "CCC-2014J" in full_text)

    # Table present. Structural.
    check("Document has at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 1:
        return False

    table = doc.tables[0]
    data_rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        data_rows.append(cells)

    # ---- CRITICAL: 10 rows + >=8 expected names (core Canvas data correctness) ----
    row_texts = " ".join(str(cell) for row in data_rows for cell in row)
    found_names = sum(1 for nm in EXPECTED_ASSIGNMENT_NAMES if nm in row_texts)
    check("Table has 10 rows and >=8 expected assignment names",
          len(data_rows) == EXPECTED_ASSIGNMENT_COUNT and found_names >= 8,
          f"rows={len(data_rows)}, names={found_names}/10")

    # Dates present (structural).
    has_dates = sum(1 for r in data_rows if len(r) > 1 and ("2014" in r[1] or "2015" in r[1]))
    check("Table has assignments with dates", has_dates >= 8, f"Found {has_dates}")

    # ---- CRITICAL: exactly 2 no-due-date rows AND both Final Exams w/ 100 points ----
    no_date_rows = [r for r in data_rows
                    if len(r) > 1 and "no due date" in r[1].lower()]
    finals_ok = True
    finals_detail = []
    for fname in FINAL_EXAM_NAMES:
        frow = next((r for r in data_rows if any(fname in c for c in r)), None)
        if frow is None:
            finals_ok = False
            finals_detail.append(f"{fname}: missing")
            continue
        joined = " ".join(frow)
        has_100 = bool(re.search(r"\b100(?:\.0)?\b", joined))
        has_no_date = "no due date" in joined.lower()
        if not (has_100 and has_no_date):
            finals_ok = False
        finals_detail.append(f"{fname}: 100={has_100}, nodate={has_no_date}")
    check("Exactly 2 no-due-date rows and both Final Exams present with 100 points",
          len(no_date_rows) == 2 and finals_ok,
          f"no_date_rows={len(no_date_rows)}; {finals_detail}")

    # ---- CRITICAL: Total Assignments line == 10 (keyword RU/EN + adjacent value) ----
    ta_val = find_value_near_keyword(
        full_text_lower, ["total assignments", "всего заданий", "количество заданий"],
        r"\b\d+\b")
    check("Total Assignments line equals 10", ta_val is not None and ta_val == "10",
          f"matched={ta_val}")

    # ---- CRITICAL: Total Points line == 300.0 (keyword RU/EN + adjacent value) ----
    tp_val = find_value_near_keyword(
        full_text_lower, ["total points", "итого баллов", "сумма баллов", "всего баллов"],
        r"\b\d+(?:\.\d+)?\b")
    tp_ok = tp_val is not None and abs(float(tp_val) - EXPECTED_TOTAL_POINTS) < 0.05
    check("Total Points line equals 300.0", tp_ok, f"matched={tp_val}")

    # Intro sentence present (structural; RU or EN). Broaden keywords.
    intro_ok = (
        "this document lists all assignments" in full_text_lower
        or ("задани" in full_text_lower and ("срок" in full_text_lower or "балл" in full_text_lower))
    )
    check("Document has intro/description line", intro_ok)

    return True


def check_teamly():
    """Check the teamly knowledge-base page — BLOCKING (critical)."""
    print("\n=== Checking Teamly Knowledge Base ===")
    crit = "Teamly 'CCC-2014J Assignment Overview' page mentions course, count 10, points 300"
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("SELECT title, COALESCE(body, '') FROM teamly.pages")
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        check(crit, False, f"db error: {e}")
        return

    # Locate the deliverable page by title (English literal marker).
    overview = [
        (t, b) for t, b in rows
        if t and ("ccc-2014j" in t.lower() or "assignment overview" in t.lower()
                  or "обзор заданий" in t.lower())
    ]
    check("Teamly overview page created", len(overview) >= 1,
          f"Total pages: {len(rows)}")

    page_text = " ".join((str(t) + " " + str(b)) for t, b in overview)
    page_lower = page_text.lower()

    mentions_course = (
        "creative computing" in page_lower or "ccc-2014j" in page_lower
    )
    mentions_count = bool(re.search(r"\b10\b", page_text))
    mentions_points = bool(re.search(r"\b300(?:\.0)?\b", page_text))

    # CRITICAL: page exists AND mentions course name/code, count 10, points 300.
    check(crit,
          bool(overview) and mentions_course and mentions_count and mentions_points,
          f"course={mentions_course}, count10={mentions_count}, points300={mentions_points}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("CANVAS ASSIGNMENT WORD TEAMLY - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace)
    check_teamly()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}, Accuracy: {accuracy:.1f}%")
    if critical_failed:
        print(f"  CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"    - {n}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({
                "total_passed": PASS_COUNT,
                "total_checks": total,
                "accuracy": accuracy,
                "critical_failed": critical_failed,
                "success": (not critical_failed) and accuracy >= 70,
            }, f, indent=2)

    if critical_failed:
        print("Overall: FAIL (critical check failed)")
        sys.exit(1)
    if accuracy >= 70:
        print("Overall: PASS")
        sys.exit(0)
    print("Overall: FAIL")
    sys.exit(1)


if __name__ == "__main__":
    main()
