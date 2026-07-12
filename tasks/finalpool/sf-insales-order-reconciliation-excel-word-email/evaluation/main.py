"""Evaluation script for sf-insales-order-reconciliation-excel-word-email."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []

# RU<->EN region equivalence. ClickHouse central map (db/zzz_clickhouse_after_init.sql)
# russifies sf_data CUSTOMERS.REGION, so a russified agent writes RU region names into
# DW_Summary, while the frozen English groundtruth Excel still has English region literals.
# Accept both forms when matching region text.
REGION_EQUIV = {
    "asia pacific": "азиатско-тихоокеанский регион",
    "europe": "европа",
    "latin america": "латинская америка",
    "middle east": "ближний восток",
    "north america": "северная америка",
}
REGION_EQUIV_REV = {v: k for k, v in REGION_EQUIV.items()}

def region_match(gs, avs):
    """True if groundtruth region string gs matches agent string avs in either RU or EN."""
    gs = gs.strip().lower()
    avs = avs.strip().lower()
    candidates = {gs}
    if gs in REGION_EQUIV:
        candidates.add(REGION_EQUIV[gs])
    if gs in REGION_EQUIV_REV:
        candidates.add(REGION_EQUIV_REV[gs])
    for c in candidates:
        if c == avs or c in avs or avs in c:
            return True
    return False

def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRITICAL]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        if critical:
            CRITICAL_FAILED.append(name)
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL]{' [CRITICAL]' if critical else ''} {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    # Check Order_Reconciliation_Report.xlsx
    excel_path = os.path.join(agent_workspace, "Order_Reconciliation_Report.xlsx")
    check("Order_Reconciliation_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Order_Reconciliation_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Check row count
                    gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                    data_rows = list(ws.iter_rows(min_row=2, values_only=True))
                    min_rows = max(1, len(gt_rows) - 2)
                    check(f"{sheet_name} has >= {min_rows} data rows", len(data_rows) >= min_rows, f"got {len(data_rows)}")

                    # Cell value comparison against groundtruth
                    header_map = {h: i for i, h in enumerate(headers)}
                    gt_header_map = {h: i for i, h in enumerate(gt_headers)}
                    for ri in range(min(3, len(gt_rows), len(data_rows))):
                        gt_row = gt_rows[ri]
                        agent_row = data_rows[ri]
                        for ci, gt_h in enumerate(gt_headers):
                            if not gt_h or ci >= len(gt_row):
                                continue
                            gv = gt_row[ci]
                            agent_ci = header_map.get(gt_h)
                            if agent_ci is None or agent_ci >= len(agent_row):
                                continue
                            av = agent_row[agent_ci]
                            gf = safe_float(gv)
                            af = safe_float(av)
                            if gf is not None and af is not None:
                                tol = max(0.5, abs(gf) * 0.15)
                                check(f"{sheet_name} R{ri+2} {gt_h} ~{gf:.1f}",
                                      abs(gf - af) <= tol, f"got {af}")
                            elif gv is not None and av is not None:
                                gs = str(gv).strip().lower()
                                avs = str(av).strip().lower()
                                if gs:
                                    if sheet_name == "DW_Summary" and gt_h == "region":
                                        matched = region_match(gs, avs)
                                    else:
                                        matched = gs == avs or gs in avs or avs in gs
                                    check(f"{sheet_name} R{ri+2} {gt_h} text",
                                          matched,
                                          f"expected {gs[:50]}, got {avs[:50]}")

    # Check Reconciliation_Audit.docx
    docx_path = os.path.join(agent_workspace, "Reconciliation_Audit.docx")
    check("Reconciliation_Audit.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        check("Reconciliation_Audit.docx has content", len(text) > 50, f"text length: {len(text)}")
        # Check headings match groundtruth
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        gt_doc_path = os.path.join(groundtruth_workspace, "Reconciliation_Audit.docx")
        # RU+EN alternatives: task.md keeps section names in English, but a russified
        # agent may write Russian headings. Accept either form.
        HEADING_ALT = {
            "data warehouse analysis": ["анализ хранилища данных", "хранилище данных"],
            "online store analysis": ["анализ интернет-магазина", "интернет-магазин", "магазин"],
            "reconciliation findings": ["результаты сверки", "выводы сверки", "находки сверки"],
        }
        if os.path.exists(gt_doc_path):
            gt_doc = Document(gt_doc_path)
            gt_headings = [p.text.strip().lower() for p in gt_doc.paragraphs if p.style.name.startswith("Heading")]
            for gh in gt_headings:
                if gh:
                    variants = [gh] + HEADING_ALT.get(gh, [])
                    found = any(v in h or h in v for h in headings for v in variants)
                    check(f"Reconciliation_Audit.docx has heading \"{gh[:40]}\"", found, f"agent headings: {headings[:5]}")
        else:
            check("Reconciliation_Audit.docx has headings", len(headings) >= 2, f"found {len(headings)} headings")

    # Check Python script exists (terminal usage)
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    check("Python analysis script exists", len(py_files) >= 1, f"found: {py_files}")

    # Database checks
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("SELECT subject, to_addr FROM email.messages WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1) AND subject ILIKE '%reconciliation%'")
        email_row = cur.fetchone()
        check("Email with correct subject sent", email_row is not None, "no matching email found", critical=True)
        if email_row:
            # recipient must include finance@company.com
            to_str = str(email_row[1]).lower() if email_row[1] is not None else ""
            check("Email sent to finance@company.com",
                  "finance@company.com" in to_str, f"to_addr: {email_row[1]}", critical=True)
        # Reverse verification: noise emails (russified subjects injected in preprocess:
        # 'Еженедельная рассылка', 'Обслуживание сервера') should not be in Sent folder
        cur.execute("SELECT COUNT(*) FROM email.messages WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1) AND (subject ILIKE '%рассылка%' OR subject ILIKE '%обслуживание сервера%')")
        noise_sent = cur.fetchone()[0]
        check("No noise emails in Sent folder", noise_sent == 0, f"found {noise_sent} noise emails in Sent", critical=True)
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e), critical=True)

    # ------------------------------------------------------------------
    # CRITICAL semantic checks: core reconciliation deliverables.
    # Any failure here => sys.exit(1) regardless of accuracy.
    # ------------------------------------------------------------------
    def sheet_to_rows(ws):
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            return [], []
        hdr = [str(c).strip().lower() if c is not None else "" for c in rows[0]]
        return hdr, rows[1:]

    gt_xlsx = os.path.join(groundtruth_workspace, "Order_Reconciliation_Report.xlsx")
    if os.path.exists(excel_path) and os.path.exists(gt_xlsx):
        a_wb = openpyxl.load_workbook(excel_path)
        g_wb = openpyxl.load_workbook(gt_xlsx)

        # DW_Summary: per-region Total_Revenue + Total_Orders (RU+EN region keys).
        if "DW_Summary" in a_wb.sheetnames and "DW_Summary" in g_wb.sheetnames:
            ah, ar = sheet_to_rows(a_wb["DW_Summary"])
            gh, gr = sheet_to_rows(g_wb["DW_Summary"])
            def col(hdr, name):
                return hdr.index(name) if name in hdr else None
            ai_reg, ai_rev, ai_ord = col(ah, "region"), col(ah, "total_revenue"), col(ah, "total_orders")
            gi_reg, gi_rev, gi_ord = col(gh, "region"), col(gh, "total_revenue"), col(gh, "total_orders")
            dw_ok = all(x is not None for x in (ai_reg, ai_rev, ai_ord, gi_reg, gi_rev, gi_ord))
            if dw_ok:
                matched_regions = 0
                for grow in gr:
                    g_region = str(grow[gi_reg]).strip().lower() if grow[gi_reg] is not None else ""
                    g_rev = safe_float(grow[gi_rev]); g_ord = safe_float(grow[gi_ord])
                    arow = None
                    for cand in ar:
                        if cand[ai_reg] is not None and region_match(g_region, str(cand[ai_reg])):
                            arow = cand
                            break
                    if arow is None:
                        dw_ok = False
                        break
                    a_rev = safe_float(arow[ai_rev]); a_ord = safe_float(arow[ai_ord])
                    if not (a_rev is not None and num_close(a_rev, g_rev) and
                            a_ord is not None and num_close(a_ord, g_ord)):
                        dw_ok = False
                        break
                    matched_regions += 1
                dw_ok = dw_ok and matched_regions == len([g for g in gr if g[gi_reg]])
            check("CRITICAL DW_Summary per-region revenue/orders correct (RU+EN regions)",
                  dw_ok, "region revenue/order aggregation mismatch", critical=True)

        # Store_Summary: per-status Order_Count + Total_Amount (English status slugs).
        if "Store_Summary" in a_wb.sheetnames and "Store_Summary" in g_wb.sheetnames:
            ah, ar = sheet_to_rows(a_wb["Store_Summary"])
            gh, gr = sheet_to_rows(g_wb["Store_Summary"])
            ai_st, ai_cnt, ai_amt = (ah.index("status") if "status" in ah else None,
                                     ah.index("order_count") if "order_count" in ah else None,
                                     ah.index("total_amount") if "total_amount" in ah else None)
            gi_st, gi_cnt, gi_amt = (gh.index("status") if "status" in gh else None,
                                     gh.index("order_count") if "order_count" in gh else None,
                                     gh.index("total_amount") if "total_amount" in gh else None)
            st_ok = all(x is not None for x in (ai_st, ai_cnt, ai_amt, gi_st, gi_cnt, gi_amt))
            if st_ok:
                a_map = {str(r[ai_st]).strip().lower(): r for r in ar if r[ai_st] is not None}
                for grow in gr:
                    g_st = str(grow[gi_st]).strip().lower() if grow[gi_st] is not None else ""
                    arow = a_map.get(g_st)
                    if arow is None:
                        st_ok = False
                        break
                    a_cnt = safe_float(arow[ai_cnt]); a_amt = safe_float(arow[ai_amt])
                    g_cnt = safe_float(grow[gi_cnt]); g_amt = safe_float(grow[gi_amt])
                    if not (a_cnt is not None and num_close(a_cnt, g_cnt) and
                            a_amt is not None and num_close(a_amt, g_amt)):
                        st_ok = False
                        break
            check("CRITICAL Store_Summary per-status count/amount correct",
                  st_ok, "status aggregation mismatch", critical=True)

        # Cross_Reference: DW_Value / Store_Value / Difference for the 3 metrics.
        if "Cross_Reference" in a_wb.sheetnames and "Cross_Reference" in g_wb.sheetnames:
            ah, ar = sheet_to_rows(a_wb["Cross_Reference"])
            gh, gr = sheet_to_rows(g_wb["Cross_Reference"])
            def cidx(hdr, name): return hdr.index(name) if name in hdr else None
            ai_m, ai_dw, ai_sv, ai_df = (cidx(ah, "metric"), cidx(ah, "dw_value"),
                                         cidx(ah, "store_value"), cidx(ah, "difference"))
            gi_m, gi_dw, gi_sv, gi_df = (cidx(gh, "metric"), cidx(gh, "dw_value"),
                                         cidx(gh, "store_value"), cidx(gh, "difference"))
            xr_ok = all(x is not None for x in (ai_m, ai_dw, ai_sv, ai_df, gi_m, gi_dw, gi_sv, gi_df))
            if xr_ok:
                a_map = {str(r[ai_m]).strip().lower(): r for r in ar if r[ai_m] is not None}
                for grow in gr:
                    g_m = str(grow[gi_m]).strip().lower() if grow[gi_m] is not None else ""
                    arow = a_map.get(g_m)
                    if arow is None:
                        xr_ok = False
                        break
                    for ai_c, gi_c in ((ai_dw, gi_dw), (ai_sv, gi_sv), (ai_df, gi_df)):
                        av = safe_float(arow[ai_c]); gvv = safe_float(grow[gi_c])
                        if gvv is None:
                            continue
                        if av is None or not num_close(av, gvv):
                            xr_ok = False
                            break
                    if not xr_ok:
                        break
            check("CRITICAL Cross_Reference reconciliation values correct",
                  xr_ok, "cross-reference metric mismatch", critical=True)

    # Pipeline artifacts: reconciler.py + docx existence prove the multi-tool run.
    check("CRITICAL reconciler.py / analysis script present in workspace",
          len(py_files) >= 1, f"found: {py_files}", critical=True)
    check("CRITICAL Reconciliation_Audit.docx exists",
          os.path.exists(docx_path), "missing audit docx", critical=True)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100.0) if total else 0.0
    critical_ok = len(CRITICAL_FAILED) == 0
    success = critical_ok and accuracy >= 70.0
    msg = f"Passed {PASS_COUNT}/{total} checks (accuracy {accuracy:.1f}%)"
    if not critical_ok:
        msg += f" | CRITICAL FAILED: {CRITICAL_FAILED}"
    return success, msg

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()