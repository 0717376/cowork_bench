"""
Evaluation for sf-sales-trend-word-pdf task.

Critical checks (CRITICAL_CHECKS): любой провал критического чека => общий FAIL
независимо от accuracy. Иначе PASS требует accuracy >= 70%.

Идентификаторы (имена файлов Sales_Trend_Analysis.docx/.pdf, формат месяца
YYYY-MM, английские имена столбцов/схемы/БД sf_data, SALES_DW, MONTH_KEY,
TOTAL_REVENUE и т.п.) — это литералы вывода и ДОЛЖНЫ оставаться английскими/
числовыми. Не переводить. Значения выручки читаются ЖИВЫМИ из БД sf_data
(ClickHouse-аналог на PG-схеме), а не захардкожены.

Заголовок/разделы/метки документа агент пишет по-русски, поэтому подстроковые
проверки расширены до RU+EN.
"""

import argparse
import json
import os
import re
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Критические семантические чеки: отражают суть отчёта (корректные значения из
# живой БД, ключевые выводы). Любой провал => общий FAIL независимо от accuracy.
CRITICAL_CHECKS = {
    "Таблица: >=11/12 месяцев с корректной выручкой (живая БД, допуск 1.0)",
    "Сводка: Total Annual Revenue совпадает с живой суммой",
    "Сводка: Best Month = месяц с макс. выручкой, указан с суммой",
    "Сводка: Worst Month = месяц с мин. выручкой, указан с суммой",
    "Динамика: H1 и H2 присутствуют, корректны, направление (H2>H1) согласовано",
    "PDF содержит помесячные ключи месяцев (>=11/12) и числа выручки",
}


def check(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        detail_str = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def contains_any(text_lower, variants):
    """variants: list of substrings; True if any present in text_lower."""
    return any(v in text_lower for v in variants)


def extract_numbers(text):
    """Return list of floats parsed from text, handling $1,234.56 style."""
    nums = []
    for m in re.findall(r"-?\$?\s?\d[\d,]*\.?\d*", text):
        cleaned = m.replace("$", "").replace(",", "").strip()
        if cleaned in ("", "-", "."):
            continue
        try:
            nums.append(float(cleaned))
        except ValueError:
            pass
    return nums


def find_amount(text, target, tol=1.0):
    """True if any number in text is within tol of target."""
    return any(num_close(n, target, tol) for n in extract_numbers(text))


def get_expected_data():
    """Fetch expected monthly revenue data from the live database."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT "MONTH_KEY", "TOTAL_REVENUE", "ORDER_COUNT", "UNIQUE_CUSTOMERS", "AVG_ORDER_VALUE"
        FROM sf_data."SALES_DW__ANALYTICS__MONTHLY_REVENUE"
        WHERE "MONTH_KEY" >= '2025-01-01' AND "MONTH_KEY" <= '2025-12-31'
        ORDER BY "MONTH_KEY"
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()
    return rows


def compute_summary(expected):
    """Derive total/best/worst/H1/H2 from live rows."""
    months = [(r[0].strftime("%Y-%m"), float(r[1])) for r in expected]
    total = sum(rev for _, rev in months)
    best = max(months, key=lambda x: x[1])
    worst = min(months, key=lambda x: x[1])
    h1 = sum(rev for m, rev in months if m <= "2025-06")
    h2 = sum(rev for m, rev in months if m >= "2025-07")
    return dict(total=total, best=best, worst=worst, h1=h1, h2=h2, months=months)


def check_word_doc(agent_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Sales_Trend_Analysis.docx ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Sales_Trend_Analysis.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Not found: {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)
    expected = get_expected_data()
    summ = compute_summary(expected)

    full_text = " ".join(p.text for p in doc.paragraphs)
    full_lower = full_text.lower()

    # Heading: RU+EN
    has_heading = ("2025" in full_text) and contains_any(
        full_lower, ["sales trend", "анализ динамики продаж", "динамика продаж", "анализ продаж"]
    )
    check("Document has Sales Trend 2025 heading (RU/EN)", has_heading,
          "Expected 'Анализ динамики продаж'/'Sales Trend' + 2025")

    # Table presence
    check("Document has at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 1:
        return False

    table = doc.tables[0]
    data_rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        data_rows.append(cells)

    check("Table has 12 month rows", len(data_rows) == 12, f"Got {len(data_rows)} rows")

    # CRITICAL: month revenue values within tight tolerance, >=11/12
    matched_months = 0
    for exp_row in expected:
        month_str = exp_row[0].strftime("%Y-%m")
        exp_revenue = float(exp_row[1])
        for dr in data_rows:
            if month_str in dr[0]:
                rev_text = dr[1].replace("$", "").replace(",", "").strip()
                if num_close(rev_text, exp_revenue, 1.0):
                    matched_months += 1
                break
    check("Таблица: >=11/12 месяцев с корректной выручкой (живая БД, допуск 1.0)",
          matched_months >= 11, f"Matched {matched_months}/12")

    # CRITICAL: Total Annual Revenue actual number present
    check("Сводка: Total Annual Revenue совпадает с живой суммой",
          find_amount(full_text, summ["total"], 1.0),
          f"Expected total ~{summ['total']:.2f}")
    # NON-critical: keyword presence for summary label (RU/EN)
    check("Document mentions total/annual revenue label (RU/EN)",
          contains_any(full_lower, ["total annual", "annual revenue", "total revenue",
                                    "итого", "годовая выручка", "общая выручка", "всего выручк"]),
          "Expected total-revenue label in summary")

    # CRITICAL: Best Month = max month, stated with amount
    best_m, best_v = summ["best"]
    check("Сводка: Best Month = месяц с макс. выручкой, указан с суммой",
          (best_m in full_text) and find_amount(full_text, best_v, 1.0),
          f"Expected best month {best_m} with ~{best_v:.2f}")

    # CRITICAL: Worst Month = min month, stated with amount
    worst_m, worst_v = summ["worst"]
    check("Сводка: Worst Month = месяц с мин. выручкой, указан с суммой",
          (worst_m in full_text) and find_amount(full_text, worst_v, 1.0),
          f"Expected worst month {worst_m} with ~{worst_v:.2f}")

    # NON-critical: best/worst labels (RU/EN)
    check("Document has best/worst month labels (RU/EN)",
          contains_any(full_lower, ["best month", "worst month",
                                    "лучший месяц", "худший месяц",
                                    "наибольш", "наименьш", "пиков"]),
          "Expected best/worst month labels")

    # CRITICAL: H1/H2 values present, correct, direction consistent
    h1_ok = find_amount(full_text, summ["h1"], 1.0)
    h2_ok = find_amount(full_text, summ["h2"], 1.0)
    increased = summ["h2"] > summ["h1"]
    # expect direction word for increase (RU/EN) since H2>H1 in live data
    direction_ok = contains_any(full_lower, ["increas", "grew", "growth", "higher", "up ",
                                             "вырос", "увеличил", "рост", "выше", "прирост"]) if increased \
        else contains_any(full_lower, ["decreas", "declin", "lower", "drop", "down",
                                       "снизил", "уменьшил", "падени", "ниже", "сокращ"])
    check("Динамика: H1 и H2 присутствуют, корректны, направление (H2>H1) согласовано",
          h1_ok and h2_ok and direction_ok,
          f"H1~{summ['h1']:.2f} ({h1_ok}), H2~{summ['h2']:.2f} ({h2_ok}), dir_ok={direction_ok}")

    # NON-critical: H1/H2 trend label (RU/EN)
    check("Document has H1/H2 trend analysis label (RU/EN)",
          contains_any(full_lower, ["h1", "h2", "first half", "second half",
                                    "полугодие", "первое полугод", "второе полугод",
                                    "анализ динамики", "trend analysis"]),
          "Expected H1/H2 / полугодие label")

    return True


def check_pdf(agent_workspace):
    """Check the PDF file exists and carries the numeric monthly data."""
    print("\n=== Checking Sales_Trend_Analysis.pdf ===")

    pdf_path = os.path.join(agent_workspace, "Sales_Trend_Analysis.pdf")
    check("PDF file exists", os.path.isfile(pdf_path), f"Not found: {pdf_path}")
    if not os.path.isfile(pdf_path):
        return False

    size = os.path.getsize(pdf_path)
    check("PDF file size reasonable", size > 1024, f"Size: {size} bytes")

    text = None
    try:
        import PyPDF2
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = "".join((page.extract_text() or "") for page in reader.pages)
    except ImportError:
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as p:
                text = "".join((page.extract_text() or "") for page in p.pages)
        except ImportError:
            # Без PDF-ридера нельзя проверить содержимое числовых данных ->
            # это CRITICAL чек, поэтому помечаем как провал, а не пропуск.
            check("PDF содержит помесячные ключи месяцев (>=11/12) и числа выручки",
                  False, "No PDF reader (PyPDF2/pdfplumber) available")
            return False

    text_lower = text.lower()
    check("PDF contains Sales Trend title (RU/EN)",
          contains_any(text_lower, ["sales trend", "анализ динамики продаж", "динамика продаж"]),
          "Expected RU/EN title")

    expected = get_expected_data()
    # CRITICAL: PDF carries the monthly month-keys AND the revenue figures
    months_found = sum(1 for r in expected if r[0].strftime("%Y-%m") in text)
    rev_found = 0
    for r in expected:
        if find_amount(text, float(r[1]), 1.0):
            rev_found += 1
    check("PDF содержит помесячные ключи месяцев (>=11/12) и числа выручки",
          months_found >= 11 and rev_found >= 11,
          f"month-keys {months_found}/12, revenue figures {rev_found}/12")

    return True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("SF SALES TREND WORD PDF - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace)
    check_pdf(args.agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}/{total} ({accuracy:.1f}%)")
    print(f"  Failed: {FAIL_COUNT}")
    if critical_failed:
        print(f"  CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"    - {n}")

    success = (not critical_failed) and accuracy >= 70
    print(f"  Overall: {'PASS' if success else 'FAIL'}")

    if args.res_log_file:
        result = {"passed": PASS_COUNT, "failed": FAIL_COUNT,
                  "accuracy": accuracy, "critical_failed": critical_failed,
                  "success": success}
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
