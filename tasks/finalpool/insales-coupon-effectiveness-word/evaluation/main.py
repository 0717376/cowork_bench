"""
Evaluation script for insales-coupon-effectiveness-word task.

Checks (values dynamically computed from the InSales wc.* DB):
1. Word document Coupon_Effectiveness_Report.docx: title (RU/EN) + used coupon codes.
2. Excel Coupon_Analysis.xlsx: structure AND numeric values verified against DB.
3. Email sent with correct subject/recipient and a non-empty findings body.

Critical checks (CRITICAL_CHECKS): any failure => overall FAIL regardless of
accuracy. Otherwise PASS requires accuracy >= 70%.
"""

import argparse
import json
import os
import re
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "cowork_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# SEMANTIC checks whose failure forces an overall FAIL regardless of accuracy.
CRITICAL_CHECKS = {
    "Excel top coupon value matches DB",
    "Excel Summary metrics match DB",
    "Word mentions >=3 used coupon codes + title present",
    "Email subject/recipient/non-empty body",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        d = detail[:300] if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def _num(s):
    """Strip currency symbols/spaces; accept comma decimals -> float or None."""
    if s is None:
        return None
    t = str(s).strip().replace(" ", "")
    t = re.sub(r"[^\d,.\-]", "", t)
    if t.count(",") == 1 and t.count(".") == 0:
        t = t.replace(",", ".")
    else:
        t = t.replace(",", "")
    try:
        return float(t)
    except ValueError:
        return None


def get_expected_coupon_data():
    """Query actual coupon usage from DB."""
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT
            cl->>'code' as coupon_code,
            COUNT(*) as usage_count,
            SUM((cl->>'discount')::numeric) as total_discount,
            AVG((cl->>'discount')::numeric) as avg_discount
        FROM wc.orders,
             jsonb_array_elements(coupon_lines) cl
        WHERE coupon_lines IS NOT NULL
          AND coupon_lines::text <> 'null'
          AND coupon_lines::text <> '[]'
        GROUP BY cl->>'code'
        ORDER BY total_discount DESC
    """)
    coupon_stats = cur.fetchall()

    cur.execute("""
        SELECT COUNT(*)
        FROM wc.orders
        WHERE coupon_lines IS NOT NULL
          AND coupon_lines::text <> 'null'
          AND coupon_lines::text <> '[]'
    """)
    total_orders_with_coupons = cur.fetchone()[0]

    cur.execute("SELECT code FROM wc.coupons")
    defined_codes = set(row[0] for row in cur.fetchall())

    used_codes = set(row[0] for row in coupon_stats)

    cur.close()
    conn.close()

    return coupon_stats, total_orders_with_coupons, defined_codes, used_codes


def check_word(workspace, used_codes):
    """Check Word document content."""
    from docx import Document

    print("\n=== Checking Word Document ===")
    docx_path = os.path.join(workspace, "Coupon_Effectiveness_Report.docx")
    if not os.path.exists(docx_path):
        check("Word file exists", False, f"Not found: {docx_path}")
        check("Word mentions >=3 used coupon codes + title present", False, "no file")
        return
    check("Word file exists", True)

    doc = Document(docx_path)
    all_text = " ".join(p.text for p in doc.paragraphs)
    # Table cells too: answers laid out in a docx table are legitimate.
    all_text += " " + " ".join(
        c.text for t in doc.tables for r in t.rows for c in r.cells)
    low = all_text.lower()

    # Title: accept English markers OR Russian equivalents.
    title_en = ("coupon" in low and "effectiveness" in low and "analysis" in low)
    title_ru = ("купон" in low and "эффективн" in low and "анализ" in low)
    # Codes mentioned (codes stay English in outputs).
    codes_found = sum(1 for code in used_codes if code.lower() in low)
    need = min(3, len(used_codes))

    check("Word mentions >=3 used coupon codes + title present",
          (title_en or title_ru) and codes_found >= need,
          f"title_en={title_en} title_ru={title_ru} codes_found={codes_found}/{need}")

    # Non-critical granularity.
    check("Word title markers present (RU/EN)", title_en or title_ru,
          f"text head: {all_text[:80]!r}")
    check("Word lists used coupon codes", codes_found >= need,
          f"found {codes_found} of needed {need}")


def check_excel(workspace, coupon_stats, total_orders_with_coupons,
                defined_codes, used_codes):
    """Check Excel file structure AND numeric values vs DB."""
    from openpyxl import load_workbook

    print("\n=== Checking Excel ===")
    xlsx_path = os.path.join(workspace, "Coupon_Analysis.xlsx")
    if not os.path.exists(xlsx_path):
        check("Excel file exists", False, f"Not found: {xlsx_path}")
        check("Excel top coupon value matches DB", False, "no file")
        check("Excel Summary metrics match DB", False, "no file")
        return
    check("Excel file exists", True)

    wb = load_workbook(xlsx_path)
    sheet_names_lower = [s.lower() for s in wb.sheetnames]

    # ---- Coupon Performance ----
    perf_ok = "coupon performance" in sheet_names_lower
    check("Has 'Coupon Performance' sheet", perf_ok, f"Found: {wb.sheetnames}")

    top_match = False
    if perf_ok:
        ws = wb[wb.sheetnames[sheet_names_lower.index("coupon performance")]]
        headers = [str(c.value).lower().replace(" ", "_") if c.value else "" for c in ws[1]]
        for rh in ["coupon_code", "usage_count", "total_discount", "avg_discount"]:
            check(f"Performance header '{rh}'",
                  any(rh in h or rh.replace("_", "") in h.replace("_", "") for h in headers),
                  f"headers={headers}")

        # Build {code -> (usage, total, avg)} from the sheet.
        # Locate columns by header.
        def col_idx(*keys):
            for i, h in enumerate(headers):
                hn = h.replace("_", "")
                if any(k.replace("_", "") in hn for k in keys):
                    return i
            return None

        c_code = col_idx("coupon_code", "code")
        c_use = col_idx("usage_count", "usage")
        c_tot = col_idx("total_discount")
        sheet_map = {}
        data_rows = 0
        for row in ws.iter_rows(min_row=2):
            if c_code is None or row[c_code].value is None:
                continue
            data_rows += 1
            code = str(row[c_code].value).strip()
            use = _num(row[c_use].value) if c_use is not None else None
            tot = _num(row[c_tot].value) if c_tot is not None else None
            sheet_map[code] = (use, tot)

        check("Performance has all used-coupon rows",
              data_rows >= len(coupon_stats),
              f"rows={data_rows}, expected>={len(coupon_stats)}")

        # CRITICAL: top coupon by total_discount must match DB (usage + total).
        if coupon_stats:
            top_code, top_use, top_tot, _ = coupon_stats[0]
            s_use, s_tot = sheet_map.get(top_code, (None, None))
            top_match = (
                s_use is not None and abs(s_use - float(top_use)) < 0.5
                and s_tot is not None and abs(s_tot - float(top_tot)) <= 0.05
            )
            check("Excel top coupon value matches DB", top_match,
                  f"{top_code}: sheet=({s_use},{s_tot}) db=({top_use},{top_tot})")
        else:
            check("Excel top coupon value matches DB", False, "no coupon_stats")
    else:
        check("Excel top coupon value matches DB", False, "no Performance sheet")

    # ---- Summary ----
    summ_ok = "summary" in sheet_names_lower
    check("Has 'Summary' sheet", summ_ok, f"Found: {wb.sheetnames}")

    if summ_ok:
        ws = wb[wb.sheetnames[sheet_names_lower.index("summary")]]
        metrics = {}
        for row in ws.iter_rows(min_row=1):
            cells = [c.value for c in row]
            if len(cells) >= 2 and cells[0] is not None:
                key = str(cells[0]).strip().lower().replace(" ", "_")
                metrics[key] = cells[1]

        exp_total_used = len(used_codes)
        # Defined coupons that never appear in any order (set difference, not a
        # count subtraction): used codes and defined codes overlap only partially,
        # so a plain count subtraction can go negative. The semantic metric is
        # "how many catalogue coupons were never used".
        exp_never = len(defined_codes - used_codes)
        exp_orders = total_orders_with_coupons
        exp_overall = sum(float(r[2]) for r in coupon_stats)

        m_used = _num(metrics.get("total_coupons_used"))
        m_orders = _num(metrics.get("total_orders_with_coupons"))
        m_never = _num(metrics.get("coupons_never_used"))
        m_overall = _num(metrics.get("overall_discount_total"))

        ok_used = m_used is not None and abs(m_used - exp_total_used) < 0.5
        ok_orders = m_orders is not None and abs(m_orders - exp_orders) < 0.5
        ok_never = m_never is not None and abs(m_never - exp_never) < 0.5
        ok_overall = m_overall is not None and abs(m_overall - exp_overall) <= 0.05

        # Non-critical per-metric granularity.
        check("Summary Total_Coupons_Used", ok_used, f"sheet={m_used} db={exp_total_used}")
        check("Summary Total_Orders_With_Coupons", ok_orders, f"sheet={m_orders} db={exp_orders}")
        check("Summary Coupons_Never_Used", ok_never, f"sheet={m_never} db={exp_never}")
        check("Summary Overall_Discount_Total", ok_overall,
              f"sheet={m_overall} db={round(exp_overall, 2)}")

        # CRITICAL: all four metrics correct.
        check("Excel Summary metrics match DB",
              ok_used and ok_orders and ok_never and ok_overall,
              f"used={ok_used} orders={ok_orders} never={ok_never} overall={ok_overall}")
    else:
        check("Excel Summary metrics match DB", False, "no Summary sheet")


def check_email():
    """Check email subject, recipient and non-empty body."""
    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
    except Exception as e:
        check("DB connection (email)", False, str(e))
        check("Email subject/recipient/non-empty body", False, "no db")
        return

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
        WHERE LOWER(subject) LIKE '%coupon effectiveness report%'
    """)
    emails = cur.fetchall()
    cur.close()
    conn.close()

    if not emails:
        check("Email subject/recipient/non-empty body", False,
              "No email with subject 'Coupon Effectiveness Report'")
        return

    subj, from_addr, to_addr, body = emails[0]
    to_str = str(to_addr or "").lower()
    from_str = str(from_addr or "").lower()
    body = body or ""

    recipient_ok = "marketing-director@company.com" in to_str
    sender_ok = "marketing@store.com" in from_str
    body_ok = len(body.strip()) >= 30

    check("Email recipient is marketing-director@company.com", recipient_ok, f"to={to_addr}")
    check("Email sender is marketing@store.com", sender_ok, f"from={from_addr}")
    check("Email body non-empty (findings)", body_ok, f"len={len(body.strip())}")

    check("Email subject/recipient/non-empty body",
          recipient_ok and sender_ok and body_ok,
          f"recipient={recipient_ok} sender={sender_ok} body={body_ok}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("Fetching expected data from InSales DB...")
    coupon_stats, total_orders_with_coupons, defined_codes, used_codes = get_expected_coupon_data()
    print(f"  used_codes={len(used_codes)} orders_with_coupons={total_orders_with_coupons}")

    check_word(args.agent_workspace, used_codes)
    check_excel(args.agent_workspace, coupon_stats, total_orders_with_coupons,
                defined_codes, used_codes)
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total else 0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== Results: {PASS_COUNT}/{total} passed ({accuracy:.1f}%) ===")
    if critical_failed:
        print(f"CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")

    success = (not critical_failed) and accuracy >= 70

    if args.res_log_file:
        try:
            with open(args.res_log_file, "w") as f:
                json.dump({
                    "total_passed": PASS_COUNT, "total_checks": total,
                    "accuracy": accuracy, "critical_failed": critical_failed,
                    "success": success,
                }, f, indent=2)
        except Exception:
            pass

    print(f"  Overall: {'PASS' if success else 'FAIL'}")
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
