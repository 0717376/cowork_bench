"""Evaluation for sf-sales-customer-segment-gsheet-word (ClickHouse / RU fork).

Data source: ClickHouse logical DWH SALES_DW (PG schema sf_data). The SEGMENT
data values are russified centrally by db/zzz_clickhouse_after_init.sql:
  Consumer   -> Частные клиенты
  Enterprise -> Корпоративный
  Government -> Государственный
  SMB        -> Малый и средний бизнес
The agent reads RU segment labels from the DB and writes them to the GSheet,
Word doc and email, so all label/prose checks below use the RU labels.

Critical checks (see CRITICAL_CHECKS): any failure => overall FAIL regardless
of accuracy. Otherwise pass threshold: accuracy >= 70%.

Checks:
1. Google Sheet "Customer Segment Analysis" with 4 segment rows
   - Contains all 4 RU segment labels
   - Per-segment Total_Revenue matches DB (num_close), rows ordered desc
2. Word doc Customer_Segment_Report.docx with:
   - Heading "Customer Segment Analysis Report"
   - Table with all 4 RU segment labels + their counts/revenue
   - Recommendations section, top-revenue segment (Частные клиенты) named
3. Email to sales-strategy@company.example.com from analytics@...
   Subject: "Customer Segment Analysis Report", body names top segment
"""
import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# RU segment labels (russified centrally) with actual DB metrics, ordered by
# Total_Revenue descending. Частные клиенты (ex-Consumer) is the top segment.
SEGMENTS = [
    {"name": "Частные клиенты", "customers": 532, "orders": 5423, "revenue": 839609.20},
    {"name": "Корпоративный", "customers": 513, "orders": 5058, "revenue": 793741.69},
    {"name": "Государственный", "customers": 474, "orders": 4679, "revenue": 712686.66},
    {"name": "Малый и средний бизнес", "customers": 481, "orders": 4840, "revenue": 702960.78},
]
TOP_SEGMENT = SEGMENTS[0]["name"]  # Частные клиенты (highest revenue)

# Critical checks: any failure => overall FAIL regardless of accuracy.
CRITICAL_CHECKS = {
    "GSheet contains all 4 RU segments",
    "GSheet top revenue (Частные клиенты ~839609) present",
    "GSheet rows ordered by Total_Revenue descending",
    "Word table contains all 4 RU segments",
    "Word doc names Частные клиенты as top revenue segment",
    "Email sent to sales-strategy@company.example.com",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        detail_str = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=500.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, title FROM gsheet.spreadsheets
        WHERE LOWER(title) LIKE '%customer%' AND LOWER(title) LIKE '%segment%'
    """)
    sheets = cur.fetchall()
    check("Customer Segment Analysis spreadsheet exists", len(sheets) >= 1,
          f"Found {len(sheets)} matching spreadsheets")

    if not sheets:
        cur.close()
        conn.close()
        return

    ss_id = sheets[0][0]
    # Pull cells with their row/column position to reconstruct ordering.
    cur.execute("""
        SELECT c.row_index, c.col_index, c.value FROM gsheet.cells c
        JOIN gsheet.sheets s ON c.spreadsheet_id = s.spreadsheet_id AND c.sheet_id = s.id
        WHERE c.spreadsheet_id = %s
        ORDER BY c.row_index, c.col_index
    """, (ss_id,))
    rows = cur.fetchall()
    cur.close()
    conn.close()

    all_values = " ".join(str(r[2]) for r in rows if r[2] is not None)

    # 1) All 4 RU segment labels present (CRITICAL).
    seg_present = {seg["name"]: (seg["name"] in all_values) for seg in SEGMENTS}
    check("GSheet contains all 4 RU segments", all(seg_present.values()),
          f"Present: {seg_present}")
    for seg in SEGMENTS:
        check(f"GSheet contains segment '{seg['name']}'", seg_present[seg["name"]],
              "Segment label not found in cells")

    # 2) Top revenue value present precisely (CRITICAL) — tightened from '839'.
    top_rev = SEGMENTS[0]["revenue"]
    rev_ok = any(num_close(str(r[2]).replace(",", "").replace(" ", ""), top_rev, 500.0)
                 for r in rows if r[2] is not None and _looks_numeric(r[2]))
    rev_ok = rev_ok or "839609" in all_values or "839,609" in all_values
    check("GSheet top revenue (Частные клиенты ~839609) present", rev_ok,
          "Consumer/Частные клиенты revenue not found")

    # 3) Per-segment Total_Revenue present within tolerance for all 4 segments.
    for seg in SEGMENTS:
        seg_rev_ok = any(
            num_close(str(r[2]).replace(",", "").replace(" ", ""), seg["revenue"], 500.0)
            for r in rows if r[2] is not None and _looks_numeric(r[2])
        )
        check(f"GSheet revenue for '{seg['name']}' (~{seg['revenue']:.0f}) present",
              seg_rev_ok, "Revenue value not found")

    # 4) Rows ordered by Total_Revenue descending (CRITICAL).
    order_ok = _gsheet_order_descending(rows)
    check("GSheet rows ordered by Total_Revenue descending", order_ok,
          "Segment rows not in descending revenue order")


def _looks_numeric(v):
    s = str(v).replace(",", "").replace(" ", "").replace("%", "").strip()
    try:
        float(s)
        return True
    except ValueError:
        return False


def _gsheet_order_descending(rows):
    """Reconstruct rows, find the segment-label column, verify the data rows
    appear in DB descending-revenue order (Частные клиенты, Корпоративный,
    Государственный, Малый и средний бизнес)."""
    # Group values by row.
    by_row = {}
    for r_idx, c_idx, val in rows:
        if val is None:
            continue
        by_row.setdefault(r_idx, {})[c_idx] = str(val)
    seg_order_expected = [seg["name"] for seg in SEGMENTS]
    # For each row, detect which (if any) segment label it contains.
    seq = []
    for r_idx in sorted(by_row):
        cells = " ".join(by_row[r_idx].values())
        for name in seg_order_expected:
            if name in cells:
                seq.append(name)
                break
    # Keep only first occurrence of each segment, in row order.
    seen = []
    for name in seq:
        if name not in seen:
            seen.append(name)
    return seen == seg_order_expected


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    doc_path = os.path.join(agent_workspace, "Customer_Segment_Report.docx")
    check("Customer_Segment_Report.docx exists", os.path.isfile(doc_path),
          f"Expected at {doc_path}")
    if not os.path.isfile(doc_path):
        return

    try:
        from docx import Document
        doc = Document(doc_path)
    except Exception as e:
        check("Word doc readable", False, str(e))
        return

    full_text = " ".join(p.text for p in doc.paragraphs)
    full_lower = full_text.lower()
    headings = [p.text.lower() for p in doc.paragraphs
                if p.style.name.startswith("Heading")]

    check("Word doc has 'Customer Segment Analysis Report' heading",
          any("customer segment analysis" in h for h in headings) or
          "customer segment analysis report" in full_lower,
          "Heading not found")

    check("Word doc has a table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")

    table_text = ""
    if doc.tables:
        table_text = " ".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        # All 4 RU segments in the table (CRITICAL).
        seg_in_table = {seg["name"]: (seg["name"] in table_text) for seg in SEGMENTS}
        check("Word table contains all 4 RU segments", all(seg_in_table.values()),
              f"Present: {seg_in_table}")
        for seg in SEGMENTS:
            check(f"Word table contains segment '{seg['name']}'",
                  seg_in_table[seg["name"]], "Not found in table")
        # Each segment's revenue in the table (within tolerance).
        for seg in SEGMENTS:
            rev_ok = any(
                num_close(_strip_num(cell.text), seg["revenue"], 500.0)
                for table in doc.tables for row in table.rows for cell in row.cells
                if _looks_numeric(cell.text)
            )
            check(f"Word table revenue for '{seg['name']}' (~{seg['revenue']:.0f})",
                  rev_ok, "Revenue value not found in table")

    check("Word doc has Recommendations section",
          "recommendation" in full_lower or "recommend" in full_lower or
          "рекоменд" in full_lower,
          "Recommendations not found")

    # Top revenue segment named in prose (RU keyword, original-case .lower()).
    check("Word doc names Частные клиенты as top revenue segment",
          TOP_SEGMENT.lower() in full_lower or TOP_SEGMENT.lower() in table_text.lower(),
          "Частные клиенты not mentioned as top segment")


def _strip_num(v):
    return str(v).replace(",", "").replace(" ", "").replace("%", "").strip()


def check_emails():
    print("\n=== Checking Emails ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    conn.close()

    def parse_recipients(to_addr):
        if to_addr is None:
            return []
        if isinstance(to_addr, list):
            return [str(r).strip().lower() for r in to_addr]
        to_str = str(to_addr).strip()
        try:
            parsed = json.loads(to_str)
            if isinstance(parsed, list):
                return [str(r).strip().lower() for r in parsed]
            return [to_str.lower()]
        except (json.JSONDecodeError, TypeError):
            return [to_str.lower()]

    target = "sales-strategy@company.example.com"
    found = None
    for subj, from_addr, to_addr, body in all_emails:
        recipients = parse_recipients(to_addr)
        if target in recipients:
            found = (subj, from_addr, to_addr, body)
            break

    check("Email sent to sales-strategy@company.example.com", found is not None)
    if found:
        subj, from_addr, to_addr, body = found
        check("Email from analytics@company.example.com",
              "analytics@company.example.com" in (from_addr or "").lower(),
              f"From: {from_addr}")
        check("Subject is 'Customer Segment Analysis Report'",
              "customer segment" in (subj or "").lower() and "analysis" in (subj or "").lower(),
              f"Subject: {subj}")
        # Body names the top-revenue segment (RU label, original-case).
        check("Email body names top-revenue segment (Частные клиенты)",
              TOP_SEGMENT.lower() in (body or "").lower(),
              "Top segment not named in body")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("SF SALES CUSTOMER SEGMENT GSHEET WORD - EVALUATION (ClickHouse/RU)")
    print("=" * 70)

    check_gsheet()
    check_word(args.agent_workspace)
    check_emails()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}/{total} ({accuracy:.1f}%)")
    print(f"  Failed: {FAIL_COUNT}")

    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    if critical_failed:
        print(f"  CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"    - {n}")

    if critical_failed:
        print("  Overall: FAIL (critical check failed)")
        sys.exit(1)
    if accuracy >= 70:
        print("  Overall: PASS")
        sys.exit(0)
    print("  Overall: FAIL (accuracy below threshold)")
    sys.exit(1)


if __name__ == "__main__":
    main()
