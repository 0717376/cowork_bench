"""
Evaluation for scholarly-nlp-survey-word task.
Checks Word doc and Excel spreadsheet for correct content.
"""
import argparse
import json
import os
import sys

import openpyxl
from docx import Document

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []  # имена провалившихся обязательных (CRITICAL) проверок

# CRITICAL — содержательные проверки. Если хоть одна FAIL, итог всей задачи = FAIL
# независимо от accuracy. Структурные чеки (наличие листа/колонки/файла) — НЕ critical.
CRITICAL_CHECKS = {
    "Оба файла существуют и читаемы",
    "Excel: все 6 статей в листе Papers и ровно 6 строк данных",
    "Excel: шумовые статьи (point cloud / soil microbiome) отсутствуют",
    "Summary: Total_Papers = 6",
    "Summary: Avg_Citations корректно (~45833)",
    "Summary: Most_Cited_Paper = 'Attention Is All You Need'",
    "Word: упомянуты все 6 ключевых статей",
    "Excel: цитирования совпадают с источником (данные не выдуманы)",
}

# The 6 NLP papers from scholarly.scholar_papers (IDs 169-174)
EXPECTED_NLP_PAPERS = [
    {"id": 169, "title": "Attention Is All You Need",
     "year": 2017, "venue": "NeurIPS 2017", "citations": 120000},
    {"id": 170, "title": "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding",
     "year": 2019, "venue": "NAACL 2019", "citations": 95000},
    {"id": 171, "title": "RoBERTa: A Robustly Optimized BERT Pretraining Approach",
     "year": 2019, "venue": "arXiv preprint", "citations": 18000},
    {"id": 172, "title": "Exploring the Limits of Transfer Learning with a Unified Text-to-Text Transformer",
     "year": 2020, "venue": "JMLR 2020", "citations": 22000},
    {"id": 173, "title": "GPT-4 Technical Report",
     "year": 2023, "venue": "arXiv preprint", "citations": 8000},
    {"id": 174, "title": "Training language models to follow instructions with human feedback",
     "year": 2022, "venue": "NeurIPS 2022", "citations": 12000},
]

# Short keywords to detect each paper in text
PAPER_KEYWORDS = [
    "bert",
    "gpt-4",
    "attention is all you need",
    "roberta",
    "text-to-text",
    "instructgpt",  # or "follow instructions"
]

PAPER_KEYWORD_ALTS = [
    ["bert"],
    ["gpt-4", "gpt4"],
    ["attention is all you need"],
    ["roberta"],
    ["text-to-text", "t5", "transfer learning"],
    ["instructgpt", "follow instructions", "human feedback"],
]

NOISE_KEYWORDS = ["point cloud", "soil microbiome"]


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        marker = " [CRITICAL]" if name in CRITICAL_CHECKS else ""
        print(f"  [FAIL]{marker} {name}{msg}")
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILS.append(name)


def num_close(a, b, tol=50):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    doc_path = os.path.join(agent_workspace, "NLP_Transformer_Survey.docx")

    if not os.path.isfile(doc_path):
        record("Word file exists", False, f"Not found: {doc_path}")
        return False
    record("Word file exists", True)

    try:
        doc = Document(doc_path)
    except Exception as e:
        record("Word file readable", False, str(e))
        return False
    record("Word file readable", True)

    # Extract all text
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    full_lower = full_text.lower()

    # Check title
    has_title = "transformer" in full_lower and "literature" in full_lower and "survey" in full_lower
    record("Word doc has survey title", has_title,
           f"Looking for 'Transformer-Based NLP: A Literature Survey'")

    # Check date
    has_date = "2026-03-06" in full_text or "march 6, 2026" in full_lower or "2026/03/06" in full_text or "march 2026" in full_lower
    record("Word doc has date", has_date, "Looking for 2026-03-06 or similar")

    # Check each paper is mentioned (структурный детальный чек, не critical)
    all_papers_mentioned = True
    for i, alt_list in enumerate(PAPER_KEYWORD_ALTS):
        found = any(kw in full_lower for kw in alt_list)
        if not found:
            all_papers_mentioned = False
        paper_name = EXPECTED_NLP_PAPERS[i]["title"][:50]
        record(f"Word mentions: {paper_name}...", found,
               f"Keywords checked: {alt_list}")

    # CRITICAL: все 6 статей упомянуты в Word
    record("Word: упомянуты все 6 ключевых статей", all_papers_mentioned,
           "Не все ключевые статьи присутствуют в обзоре")

    # Check noise papers are NOT mentioned
    for noise in NOISE_KEYWORDS:
        absent = noise not in full_lower
        record(f"Word does NOT mention: {noise}", absent,
               "Noise paper should not appear in NLP survey")

    return True


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower().replace(" ", "_") == sheet_name.strip().lower().replace(" ", "_"):
            ws = wb[name]
            return [[cell.value for cell in row] for row in ws.iter_rows()]
        if name.strip().lower().replace("_", " ") == sheet_name.strip().lower().replace("_", " "):
            ws = wb[name]
            return [[cell.value for cell in row] for row in ws.iter_rows()]
    return None


def find_col(header, names):
    if not header:
        return None
    for i, cell in enumerate(header):
        if cell is None:
            continue
        c = str(cell).strip().lower().replace(" ", "_")
        for n in names:
            if n.lower().replace(" ", "_") == c:
                return i
    return None


def check_excel(agent_workspace):
    print("\n=== Checking Excel Spreadsheet ===")
    excel_path = os.path.join(agent_workspace, "NLP_Paper_Catalog.xlsx")

    if not os.path.isfile(excel_path):
        record("Excel file exists", False, f"Not found: {excel_path}")
        return False
    record("Excel file exists", True)

    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
    except Exception as e:
        record("Excel file readable", False, str(e))
        return False
    record("Excel file readable", True)

    # Check Papers sheet
    papers_rows = load_sheet_rows(wb, "Papers")
    if papers_rows is None:
        record("Sheet 'Papers' exists", False, f"Available: {wb.sheetnames}")
        return False
    record("Sheet 'Papers' exists", True)

    header = papers_rows[0] if papers_rows else []
    data_rows = papers_rows[1:] if len(papers_rows) > 1 else []

    rows_ok = len(data_rows) == 6
    record("Papers sheet has 6 data rows", rows_ok,
           f"Found {len(data_rows)} rows, expected 6")

    title_col = find_col(header, ["Title", "title", "Paper_Title"])
    year_col = find_col(header, ["Year", "year", "Pub_Year", "pub_year"])
    cite_col = find_col(header, ["Citations", "citations", "Citation_Count", "citation_count"])
    venue_col = find_col(header, ["Venue", "venue"])

    all_titles_present = False
    noise_absent = True
    if title_col is not None:
        found_titles = set()
        for row in data_rows:
            if title_col < len(row) and row[title_col]:
                found_titles.add(str(row[title_col]).strip().lower())

        all_titles_present = True
        for paper in EXPECTED_NLP_PAPERS:
            t_lower = paper["title"].lower()
            found = any(t_lower in t or t in t_lower for t in found_titles)
            if not found:
                all_titles_present = False
            record(f"Papers sheet has: {paper['title'][:50]}...", found)

        # Шумовые статьи не должны попасть в каталог
        joined = " ".join(found_titles)
        for noise in NOISE_KEYWORDS:
            if noise in joined:
                noise_absent = False
    else:
        record("Title column found", False, f"Header: {header}")

    # CRITICAL: все 6 статей присутствуют в Papers и ровно 6 строк данных
    record("Excel: все 6 статей в листе Papers и ровно 6 строк данных",
           all_titles_present and rows_ok,
           f"all_titles_present={all_titles_present}, rows={len(data_rows)}")

    # CRITICAL: шумовые статьи отсутствуют
    record("Excel: шумовые статьи (point cloud / soil microbiome) отсутствуют",
           noise_absent, "Шумовая статья попала в каталог")

    # Check citations (CRITICAL: данные взяты из источника, а не выдуманы)
    citations_ok = True
    if cite_col is not None and title_col is not None:
        for row in data_rows:
            if title_col < len(row) and row[title_col]:
                t = str(row[title_col]).strip().lower()
                for paper in EXPECTED_NLP_PAPERS:
                    if paper["title"].lower() in t or t in paper["title"].lower():
                        actual = row[cite_col] if cite_col < len(row) else None
                        ok = num_close(actual, paper["citations"], tol=5000)
                        if not ok:
                            citations_ok = False
                        record(f"Citations for {paper['title'][:40]}...", ok,
                               f"Got {actual}, expected {paper['citations']}")
                        break
    else:
        citations_ok = False
    record("Excel: цитирования совпадают с источником (данные не выдуманы)",
           citations_ok, "Значения цитирований не совпадают с scholarly")

    # Check Summary sheet
    summary_rows = load_sheet_rows(wb, "Summary")
    if summary_rows is None:
        record("Sheet 'Summary' exists", False, f"Available: {wb.sheetnames}")
        return False
    record("Sheet 'Summary' exists", True)

    metrics = {}
    for row in summary_rows:
        if row and row[0] is not None:
            key = str(row[0]).strip().lower().replace(" ", "_")
            val = row[1] if len(row) > 1 else None
            metrics[key] = val

    # Total Papers
    total_key = None
    for k in metrics:
        if "total" in k and "paper" in k:
            total_key = k
            break
    if total_key:
        ok = num_close(metrics[total_key], 6, tol=0)
        record("Summary: Total_Papers = 6", ok, f"Got {metrics[total_key]}")
    else:
        record("Summary: Total_Papers = 6", False, f"Keys: {list(metrics.keys())}")

    # Avg Citations
    avg_key = None
    for k in metrics:
        if "avg" in k and "cit" in k:
            avg_key = k
            break
    expected_avg = (120000 + 95000 + 18000 + 22000 + 8000 + 12000) / 6  # ~45833
    if avg_key:
        ok = num_close(metrics[avg_key], expected_avg, tol=5000)
        record("Summary: Avg_Citations корректно (~45833)", ok,
               f"Got {metrics[avg_key]}, expected ~{expected_avg:.1f}")
    else:
        record("Summary: Avg_Citations корректно (~45833)", False,
               f"Keys: {list(metrics.keys())}")

    # Most Cited Paper
    most_key = None
    for k in metrics:
        if "most" in k and "cit" in k:
            most_key = k
            break
    if most_key:
        val = str(metrics[most_key]).lower() if metrics[most_key] else ""
        ok = "attention" in val or "all you need" in val
        record("Summary: Most_Cited_Paper = 'Attention Is All You Need'", ok,
               f"Got '{metrics[most_key]}'")
    else:
        record("Summary: Most_Cited_Paper = 'Attention Is All You Need'", False,
               f"Keys: {list(metrics.keys())}")

    return True


def check_files_present(agent_workspace):
    """CRITICAL: оба ключевых файла существуют и читаемы."""
    print("\n=== Checking Deliverables Present ===")
    docx_path = os.path.join(agent_workspace, "NLP_Transformer_Survey.docx")
    xlsx_path = os.path.join(agent_workspace, "NLP_Paper_Catalog.xlsx")
    ok = True
    try:
        if os.path.isfile(docx_path):
            Document(docx_path)
        else:
            ok = False
        if os.path.isfile(xlsx_path):
            openpyxl.load_workbook(xlsx_path, data_only=True)
        else:
            ok = False
    except Exception:
        ok = False
    record("Оба файла существуют и читаемы", ok,
           "Нужны NLP_Transformer_Survey.docx и NLP_Paper_Catalog.xlsx")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--launch_time", required=False)
    args = parser.parse_args()

    check_files_present(args.agent_workspace)
    check_word(args.agent_workspace)
    check_excel(args.agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\n=== Results: {PASS_COUNT}/{total} passed ({accuracy:.1f}%) ===")
    if CRITICAL_FAILS:
        print(f"Critical fails ({len(CRITICAL_FAILS)}): {CRITICAL_FAILS}")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_fails": CRITICAL_FAILS,
    }
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if CRITICAL_FAILS:
        print(f"FAIL: критичные чеки провалены ({len(CRITICAL_FAILS)})")
        sys.exit(1)
    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print(f"FAIL: accuracy {accuracy:.1f}% < 70%")
        sys.exit(1)


if __name__ == "__main__":
    main()
