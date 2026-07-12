"""Evaluation for sf-hr-attrition-word-email (ClickHouse / russified HR_ANALYTICS).

The live HR_ANALYTICS DB values (DEPARTMENT_NAME / DEPARTMENT / ROLE / LOCATION /
EMPLOYEE_NAME) are russified centrally by db/zzz_clickhouse_after_init.sql, so the
agent legitimately writes Russian department names and employee names. The
groundtruth Attrition_Data.xlsx is a binary file that still holds the ENGLISH
realia. To keep seed<->eval<->groundtruth in sync without hand-editing the binary,
this evaluation matches both the ENGLISH groundtruth form AND its russified
equivalent (RU+EN tolerant). Numeric aggregates (counts/salary/exp/percentage) are
unchanged by russification and are checked directly.
"""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []

# EN -> RU department map mirroring db/zzz_clickhouse_after_init.sql (HR_ANALYTICS).
DEPT_EN2RU = {
    "engineering": "инженерия",
    "finance": "финансы",
    "hr": "кадры",
    "operations": "операции",
    "r&d": "ниокр",
    "sales": "продажи",
    "support": "поддержка",
}

# EN -> RU employee-name map for the top-risk groundtruth names (from the central map).
NAME_EN2RU = {
    "vikram iyer": "виктор игнатьев",
    "olivia williams": "оливия васильев",
    "aisha thomas": "аиша тимофеев",
    "daniel wilson": "даниил виноградов",
    "karen gupta": "карина гущин",
    "emily iyer": "эмилия игнатьев",
    "nina smith": "нина смирнов",
    "karen singh": "карина семёнов",
    "emily kumar": "эмилия кузнецов",
    "arun kumar": "арун кузнецов",
}


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRITICAL]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL]{' [CRITICAL]' if critical else ''} {name}: {d}")
        if critical:
            CRITICAL_FAILS.append(name)


def dept_variants(dept):
    """Return lowercase EN + RU forms of a department name for substring matching."""
    en = str(dept).strip().lower()
    out = {en}
    if en in DEPT_EN2RU:
        out.add(DEPT_EN2RU[en])
    return out


def name_variants(name):
    en = str(name).strip().lower()
    out = {en}
    if en in NAME_EN2RU:
        out.add(NAME_EN2RU[en])
    return out


def cell_has_any(haystack, variants):
    h = haystack.lower()
    return any(v in h for v in variants if v)


def num_in_text(text, value, tol=1.0):
    """True if a number within tol of `value` appears anywhere in `text`.

    Normalizes Russian numeric formatting before tokenizing:
      - spaces / non-breaking spaces used as thousands separators are stripped
        between digits ("1 680" / "1 680" -> "1680");
      - a comma used as a decimal separator between digits becomes a dot
        ("3,4" -> "3.4").
    """
    import re
    cleaned = re.sub(r"(?<=\d)[   ](?=\d)", "", text)
    # strip comma thousands separators ("1,680" -> "1680") before decimal-comma rule
    cleaned = re.sub(r"(?<=\d),(?=\d{3}(?:\D|$))", "", cleaned)
    cleaned = re.sub(r"(?<=\d),(?=\d)", ".", cleaned)
    for tok in re.findall(r"\d+(?:\.\d+)?", cleaned):
        try:
            if abs(float(tok) - float(value)) <= tol:
                return True
        except ValueError:
            continue
    return False


def check_word_doc(agent_workspace, groundtruth_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return None

    doc_path = os.path.join(agent_workspace, "Attrition_Risk.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}", critical=True)
    if not os.path.isfile(doc_path):
        return None

    doc = Document(doc_path)

    # Heading
    has_heading = any("attrition" in p.text.lower() and "risk" in p.text.lower() for p in doc.paragraphs)
    check("Document has attrition risk heading", has_heading)

    check("Document has at least 2 tables", len(doc.tables) >= 2, f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 2:
        return None

    # Groundtruth
    gt_file = os.path.join(groundtruth_workspace, "Attrition_Data.xlsx")
    if not os.path.isfile(gt_file):
        check("Groundtruth file exists", False)
        return None
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    # ---- Department table ----
    gt_dept = list(gt_wb["By Department"].iter_rows(min_row=2, values_only=True))
    table1 = doc.tables[0]
    dept_rows = [[c.text.strip() for c in row.cells] for row in table1.rows[1:]]
    check("Department table has 7 rows", len(dept_rows) == 7, f"Got {len(dept_rows)} rows")

    # Per-department count correctness (RU+EN tolerant on name; numeric on count).
    dept_count_all_ok = True
    matched_counts = []  # (gt_count, doc_count) in doc order to check sort later
    for gt_row in gt_dept:
        dept, count, avg_sal, avg_exp = gt_row
        variants = dept_variants(dept)
        matched = None
        for r in dept_rows:
            if r and cell_has_any(r[0], variants):
                matched = r
                break
        if not matched:
            check(f"Dept {dept} found in table", False, f"variants={variants}")
            dept_count_all_ok = False
            continue
        # Find the Flight_Risk_Count column: a cell whose int is within tol of count
        # AND clearly not the salary/experience column (count is a small integer).
        found_count = False
        for cell in matched[1:]:
            try:
                val = int(str(cell).replace(",", "").split(".")[0])
            except (ValueError, AttributeError):
                continue
            if abs(val - count) <= 2 and val < 10000:  # counts are O(100s), salaries O(10k+)
                found_count = True
                break
        check(f"Dept {dept} flight risk count ~{count}", found_count)
        if not found_count:
            dept_count_all_ok = False

    # CRITICAL: all 7 departments present (RU+EN) with correct flight-risk counts.
    check("By-Department: all 7 depts present with correct flight-risk counts (RU+EN)",
          dept_count_all_ok and len(dept_rows) == 7, critical=True)

    # CRITICAL: top department is R&D / НИОКР (highest flight risk), first data row.
    if dept_rows:
        top_ok = cell_has_any(dept_rows[0][0], dept_variants("R&D"))
        check("By-Department sorted: top row is R&D / НИОКР", top_ok,
              f"top row={dept_rows[0][0] if dept_rows else None}", critical=True)

    # ---- Top risk employees table ----
    gt_risk = list(gt_wb["Top Risk Employees"].iter_rows(min_row=2, values_only=True))
    table2 = doc.tables[1]
    risk_rows = [[c.text.strip() for c in row.cells] for row in table2.rows[1:]]
    check("Top risk table has 10 rows", len(risk_rows) == 10, f"Got {len(risk_rows)} rows")

    # Ordering among ties (all perf=5, satisfaction=1) is non-deterministic, so we
    # only require that SOME groundtruth top-risk name (EN or RU) appears, and that
    # all rows are R&D / НИОКР (the highest-risk dept dominates the top-10).
    if gt_risk and risk_rows:
        joined = [" ".join(r).lower() for r in risk_rows]
        any_known = False
        for gt in gt_risk:
            variants = name_variants(gt[1])
            if any(any(v in j for v in variants) for j in joined):
                any_known = True
                break
        check("A known top-risk employee (RU/EN) appears in table", any_known)
        all_rnd = all(cell_has_any(j, dept_variants("R&D")) for j in joined)
        check("All top-10 rows are R&D / НИОКР dept", all_rnd)

    # ---- Summary ----
    gt_summary = list(gt_wb["Summary"].iter_rows(min_row=2, values_only=True))
    summary_dict = {r[0]: r[1] for r in gt_summary}
    full_text = " ".join(p.text for p in doc.paragraphs)

    total_emp = summary_dict.get("Total Employees")
    check("Summary mentions total employees", num_in_text(full_text, total_emp, tol=0),
          f"Expected {total_emp}")

    flight_count = summary_dict.get("Flight Risk Count")  # 1680
    has_flight = num_in_text(full_text, flight_count, tol=2)
    check("Summary mentions flight risk count (~1680)", has_flight,
          f"Expected {flight_count}", critical=True)

    pct = summary_dict.get("Flight Risk Percentage")  # 3.4
    has_pct = num_in_text(full_text, pct, tol=0.1)
    check("Summary mentions flight risk percentage (~3.4%)", has_pct,
          f"Expected {pct}", critical=True)

    # Highest risk department named in summary (R&D / НИОКР)
    has_top_dept = cell_has_any(full_text, dept_variants("R&D"))
    check("Summary names highest-risk department (R&D / НИОКР)", has_top_dept, critical=True)

    return summary_dict


def check_email(summary_dict):
    """Check email sent to hr-director with the required body content."""
    print("\n=== Checking Email ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
        WHERE folder_id != 0
        ORDER BY date DESC
    """)
    emails = cur.fetchall()
    cur.execute("SELECT * FROM email.sent_log ORDER BY id DESC LIMIT 10")
    sent = cur.fetchall()
    cur.close()
    conn.close()

    all_msgs = list(emails) + [
        (s[2] if len(s) > 2 else None, s[3] if len(s) > 3 else None,
         s[4] if len(s) > 4 else None, s[5] if len(s) > 5 else None)
        for s in sent
    ]

    check("At least 1 email sent", len(all_msgs) >= 1,
          f"Found {len(emails)} messages, {len(sent)} sent_log")
    if not all_msgs:
        check("Email sent to hr-director with required content", False, "no emails", critical=True)
        return

    # Identify the relevant email: to hr-director.
    hr_msgs = [m for m in all_msgs if "hr-director" in str(m[2] or "").lower()]
    check("Email sent to hr-director", len(hr_msgs) >= 1,
          f"recipients={[m[2] for m in all_msgs[:5]]}")

    candidates = hr_msgs if hr_msgs else all_msgs

    # Subject: prefer the exact required subject substring; accept attrition/risk fallback.
    subj_ok = any("attrition risk analysis report" in str(m[0] or "").lower() for m in candidates)
    subj_soft = any(("attrition" in str(m[0] or "").lower()) or ("risk" in str(m[0] or "").lower())
                    for m in candidates)
    check("Email subject is 'Attrition Risk Analysis Report'", subj_ok,
          f"subjects={[m[0] for m in candidates[:5]]}")
    if not subj_ok:
        check("Email subject mentions attrition/risk", subj_soft)

    # Body content: total flight count (1680), percentage (3.4), top dept (R&D/НИОКР).
    flight_count = summary_dict.get("Flight Risk Count") if summary_dict else 1680
    pct = summary_dict.get("Flight Risk Percentage") if summary_dict else 3.4

    body_has_count = any(num_in_text(str(m[3] or ""), flight_count, tol=2) for m in candidates)
    body_has_pct = any(num_in_text(str(m[3] or ""), pct, tol=0.1) for m in candidates)
    body_has_topdept = any(cell_has_any(str(m[3] or ""), dept_variants("R&D")) for m in candidates)

    check("Email body mentions total flight risk count (~1680)", body_has_count, critical=True)
    check("Email body mentions flight risk percentage (~3.4%)", body_has_pct)
    check("Email body mentions top department (R&D / НИОКР)", body_has_topdept)

    # CRITICAL composite: correct recipient + correct subject + count present in body.
    composite = (len(hr_msgs) >= 1) and subj_ok and body_has_count
    check("Email to hr-director with subject 'Attrition Risk Analysis Report' + count in body",
          composite, critical=True)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("SF HR ATTRITION WORD EMAIL - EVALUATION (ClickHouse / RU)")
    print("=" * 70)

    summary_dict = check_word_doc(args.agent_workspace, gt_dir)
    check_email(summary_dict)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100.0) if total else 0.0
    critical_ok = len(CRITICAL_FAILS) == 0
    all_ok = critical_ok and accuracy >= 70.0

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}, Accuracy: {accuracy:.1f}%")
    if CRITICAL_FAILS:
        print(f"  CRITICAL FAILURES: {CRITICAL_FAILS}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT,
                       "accuracy": accuracy, "critical_fails": CRITICAL_FAILS,
                       "success": all_ok}, f, indent=2)

    if not critical_ok:
        print("FAIL: critical check(s) failed.")
        sys.exit(1)
    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
