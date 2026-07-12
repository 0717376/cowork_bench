"""Evaluation for sf-sales-region-analysis-gsheet-email (ClickHouse DWH)."""
import argparse
import json
import os
import sys

import psycopg2

DB = {"host": os.environ.get("PGHOST", "localhost"), "port": 5432, "dbname": "cowork_gym", "user": "eigent", "password": "camel"}

# Region values are russified CENTRALLY by db/zzz_clickhouse_after_init.sql.
# The agent legitimately reads RUSSIAN region names from the DWH and writes them
# everywhere, while groundtruth keeps ENGLISH keys (central-map policy: do NOT
# hand-translate REGION literals in eval/groundtruth). To keep seed<->eval<->gt
# in sync we canonicalize every region string (EN or RU) to its English key.
REGION_ALIASES = {
    "europe": ["europe", "европа"],
    "asia pacific": ["asia pacific", "азиатско-тихоокеанский регион", "азиатско тихоокеанский регион"],
    "north america": ["north america", "северная америка"],
    "middle east": ["middle east", "ближний восток"],
    "latin america": ["latin america", "латинская америка"],
}
# variant (lowercase) -> canonical english key
_VARIANT_TO_CANON = {v.strip().lower(): canon for canon, variants in REGION_ALIASES.items() for v in variants}


def canon_region(s):
    """Map a region string (English or Russian) to its canonical English key."""
    if s is None:
        return None
    return _VARIANT_TO_CANON.get(str(s).strip().lower(), str(s).strip().lower())


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def region_match(a, b):
    """True if a and b refer to the same region (cross-language)."""
    if a is None or b is None:
        return a is None and b is None
    return canon_region(a) == canon_region(b)


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def check_excel(agent_workspace, gt_dir):
    """Returns (errors, critical_errors)."""
    errors, critical = [], []
    try:
        import openpyxl
    except ImportError:
        errors.append("openpyxl not installed")
        return errors, critical

    agent_file = os.path.join(agent_workspace, "Regional_Sales_Report.xlsx")
    gt_file = os.path.join(gt_dir, "Regional_Sales_Report.xlsx")

    if not os.path.exists(agent_file):
        errors.append("Regional_Sales_Report.xlsx not found in agent workspace")
        return errors, critical
    if not os.path.exists(gt_file):
        errors.append("Groundtruth Regional_Sales_Report.xlsx not found")
        return errors, critical

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    # Check Regional Performance sheet
    a_rows = load_sheet_rows(agent_wb, "Regional Performance")
    g_rows = load_sheet_rows(gt_wb, "Regional Performance")
    if a_rows is None:
        errors.append("Sheet 'Regional Performance' not found in agent output")
    else:
        a_data = [r for r in (a_rows[1:] if len(a_rows) > 1 else []) if r and r[0] is not None]
        g_data = [r for r in (g_rows[1:] if g_rows and len(g_rows) > 1 else []) if r and r[0] is not None]

        if len(a_data) < 5:
            errors.append(f"Regional Performance: expected 5 data rows, got {len(a_data)}")
        else:
            # canonicalize agent region keys (agent writes Russian names)
            a_lookup = {canon_region(r[0]): r for r in a_data if r[0]}
            for g_row in g_data:
                key = canon_region(g_row[0])  # groundtruth English -> canon
                a_row = a_lookup.get(key)
                if a_row is None:
                    msg = f"Missing region row: {g_row[0]}"
                    errors.append(msg)
                    critical.append(msg)
                    continue
                # Total_Revenue col 2 is the core per-region figure -> CRITICAL
                if len(a_row) > 2 and not num_close(a_row[2], g_row[2], 1.0):
                    msg = f"{g_row[0]} Total_Revenue: got {a_row[2]}, expected {g_row[2]} (tol=1.0)"
                    errors.append(msg)
                    critical.append(msg)
                # Order_Count col 1
                if len(a_row) > 1 and not num_close(a_row[1], g_row[1], 5):
                    errors.append(f"{g_row[0]} Order_Count: got {a_row[1]}, expected {g_row[1]} (tol=5)")
                # Avg_Order_Value col 3
                if len(a_row) > 3 and not num_close(a_row[3], g_row[3], 1.0):
                    errors.append(f"{g_row[0]} Avg_Order_Value: got {a_row[3]}, expected {g_row[3]} (tol=1.0)")
                # Revenue_Share_Pct col 4
                if len(a_row) > 4 and not num_close(a_row[4], g_row[4], 0.5):
                    errors.append(f"{g_row[0]} Revenue_Share_Pct: got {a_row[4]}, expected {g_row[4]} (tol=0.5)")

    # Check Summary sheet
    a_sum = load_sheet_rows(agent_wb, "Summary")
    g_sum = load_sheet_rows(gt_wb, "Summary")
    if a_sum is None:
        errors.append("Sheet 'Summary' not found in agent output")
    else:
        a_sum_data = {str(r[0]).strip().lower(): r[1] for r in (a_sum[1:] if len(a_sum) > 1 else []) if r and r[0]}
        g_sum_data = {str(r[0]).strip().lower(): r[1] for r in (g_sum[1:] if g_sum and len(g_sum) > 1 else []) if r and r[0]}

        gt_total_revenue = g_sum_data.get("total_revenue")
        gt_total_orders = g_sum_data.get("total_orders")
        gt_top = g_sum_data.get("top_region")
        gt_bottom = g_sum_data.get("bottom_region")

        # Total_Revenue -> CRITICAL (core aggregate derived from the DWH)
        tr = a_sum_data.get("total_revenue")
        if tr is None:
            msg = "Summary missing Total_Revenue"
            errors.append(msg); critical.append(msg)
        elif not num_close(tr, gt_total_revenue, 1.0):
            msg = f"Total_Revenue: got {tr}, expected {gt_total_revenue} (tol=1.0)"
            errors.append(msg); critical.append(msg)

        # Total_Orders -> CRITICAL
        to = a_sum_data.get("total_orders")
        if to is None:
            msg = "Summary missing Total_Orders"
            errors.append(msg); critical.append(msg)
        elif not num_close(to, gt_total_orders, 10):
            msg = f"Total_Orders: got {to}, expected {gt_total_orders} (tol=10)"
            errors.append(msg); critical.append(msg)

        # Top_Region -> CRITICAL (cross-language match)
        top = a_sum_data.get("top_region")
        if top is None:
            msg = "Summary missing Top_Region"
            errors.append(msg); critical.append(msg)
        elif not region_match(top, gt_top):
            msg = f"Top_Region: got '{top}', expected '{gt_top}' (EN or RU)"
            errors.append(msg); critical.append(msg)

        # Bottom_Region -> CRITICAL (cross-language match)
        bot = a_sum_data.get("bottom_region")
        if bot is None:
            msg = "Summary missing Bottom_Region"
            errors.append(msg); critical.append(msg)
        elif not region_match(bot, gt_bottom):
            msg = f"Bottom_Region: got '{bot}', expected '{gt_bottom}' (EN or RU)"
            errors.append(msg); critical.append(msg)

    return errors, critical


def _gt_region_figures(gt_dir):
    """canon region key -> (order_count, total_revenue) from groundtruth."""
    figs = {}
    try:
        import openpyxl
    except ImportError:
        return figs
    gt_file = os.path.join(gt_dir, "Regional_Sales_Report.xlsx")
    if not os.path.exists(gt_file):
        return figs
    wb = openpyxl.load_workbook(gt_file, data_only=True)
    rows = load_sheet_rows(wb, "Regional Performance") or []
    for r in rows[1:]:
        if r and r[0] is not None:
            figs[canon_region(r[0])] = (r[1], r[2])
    return figs


# manager email -> canonical region key
MANAGER_REGION = {
    "europe.manager": "europe",
    "apac.manager": "asia pacific",
    "na.manager": "north america",
    "me.manager": "middle east",
    "latam.manager": "latin america",
}


def check_gsheet():
    """Returns (errors, critical_errors)."""
    errors, critical = [], []
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("""
            SELECT s.id, s.title
            FROM gsheet.spreadsheets s
            WHERE LOWER(s.title) LIKE '%regional%' OR LOWER(s.title) LIKE '%sales regional%'
        """)
        sheets = cur.fetchall()
        if not sheets:
            msg = "No Google Sheet named 'Sales Regional Performance' found"
            errors.append(msg); critical.append(msg)
        else:
            ss_id = sheets[0][0]
            cur.execute("""
                SELECT COUNT(DISTINCT c.row_index)
                FROM gsheet.cells c
                WHERE c.spreadsheet_id = %s AND c.row_index > 0
            """, (ss_id,))
            row_count = cur.fetchone()[0]
            if row_count < 5:
                msg = f"Sales Regional Performance sheet has only {row_count} data rows, expected at least 5"
                errors.append(msg); critical.append(msg)
        cur.close()
        conn.close()
    except Exception as e:
        errors.append(f"GSheet DB check error: {e}")
    return errors, critical


def check_emails(gt_dir):
    """Returns (errors, critical_errors). Verifies each manager got a mail whose
    body contains that region's revenue AND order-count figures."""
    errors, critical = [], []
    figs = _gt_region_figures(gt_dir)
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("""
            SELECT subject, to_addr, body_text FROM email.messages
            WHERE LOWER(subject) LIKE '%regional%' OR LOWER(subject) LIKE '%sales performance%'
        """)
        emails = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        errors.append(f"Email DB check error: {e}")
        return errors, critical

    if len(emails) < 5:
        msg = f"Expected at least 5 regional emails, found {len(emails)}"
        errors.append(msg); critical.append(msg)

    def to_str(v):
        if isinstance(v, list):
            return " ".join(str(x).lower() for x in v)
        if isinstance(v, str):
            s = v.strip()
            if s.startswith("[") and s.endswith("]"):
                try:
                    p = json.loads(s)
                    if isinstance(p, list):
                        return " ".join(str(x).lower() for x in p)
                except Exception:
                    pass
            return s.lower()
        return str(v).lower()

    # collect bodies per manager
    mgr_bodies = {m: [] for m in MANAGER_REGION}
    for subj, to_addr, body in emails:
        addrs = to_str(to_addr)
        for mgr in MANAGER_REGION:
            if mgr in addrs:
                mgr_bodies[mgr].append((body or ""))

    for mgr, region_key in MANAGER_REGION.items():
        bodies = mgr_bodies.get(mgr, [])
        if not bodies:
            msg = f"No email found for {mgr}@company.com"
            errors.append(msg); critical.append(msg)
            continue
        # body must contain that region's revenue AND order-count figures
        oc, tr = figs.get(region_key, (None, None))
        # Strip thousands separators, spaces AND the decimal point so the
        # integer-part digits of the revenue are contiguous regardless of how
        # the agent formatted the (exact 2-decimal) figure (e.g. "642,644.81").
        import math
        joined = " ".join(b.lower() for b in bodies).replace(",", "").replace(" ", "").replace(".", "")
        rev_ok = True
        ord_ok = True
        if tr is not None:
            # Accept if the integer part (floor) of the revenue appears in the
            # body. Using floor (not round) avoids spurious failures when the
            # exact 2dp figure rounds UP (e.g. 642644.81 -> 642645 would not be
            # a substring of "64264481").
            rev_ok = str(int(math.floor(float(tr)))) in joined
        if oc is not None:
            ord_ok = str(int(oc)) in joined
        if not rev_ok:
            errors.append(f"{mgr} email body missing revenue figure (~{int(math.floor(float(tr)))})")
        if not ord_ok:
            errors.append(f"{mgr} email body missing order-count figure ({int(oc)})")
    return errors, critical


def check_word(agent_workspace):
    """Returns (errors, critical_errors)."""
    errors, critical = [], []
    docx_path = os.path.join(agent_workspace, "Regional_Summary.docx")
    if not os.path.exists(docx_path):
        errors.append("Regional_Summary.docx not found")
        return errors, critical
    try:
        from docx import Document
        doc = Document(docx_path)
        # include table cell text too (the breakdown lives in a table)
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        text = " ".join(parts).lower()
        if len(text.strip()) < 30:
            errors.append("Regional_Summary.docx has too little content")
        # accept EN or RU for "regional"/"europe": the agent writes Russian region names
        if not any(k in text for k in ["regional", "регион"]):
            errors.append("Regional_Summary.docx missing keyword: regional/регион")
        if not any(k in text for k in ["europe", "европа"]):
            errors.append("Regional_Summary.docx missing keyword: europe/европа")
    except ImportError:
        if os.path.getsize(docx_path) < 100:
            errors.append("Regional_Summary.docx too small")
    except Exception as e:
        errors.append(f"Error reading Regional_Summary.docx: {e}")
    return errors, critical


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    all_errors = []
    critical_errors = []
    checks_total = 0
    checks_passed = 0

    def run(name, errs_crit):
        nonlocal checks_total, checks_passed
        errs, crit = errs_crit
        checks_total += 1
        print(f"\n=== Checking {name} ===")
        if errs:
            for e in errs:
                print(f"  [FAIL] {e}")
            all_errors.extend(errs)
        else:
            checks_passed += 1
            print(f"  [PASS] {name} check passed")
        critical_errors.extend(crit)

    run("Excel", check_excel(args.agent_workspace, gt_dir))
    run("Google Sheet", check_gsheet())
    run("Emails", check_emails(gt_dir))
    run("Word Document", check_word(args.agent_workspace))

    accuracy = (checks_passed / checks_total * 100.0) if checks_total else 0.0

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({
                "errors": all_errors,
                "critical_errors": critical_errors,
                "accuracy": accuracy,
                "success": (not critical_errors) and accuracy >= 70,
            }, f, indent=2)

    # CRITICAL gate: any critical failure => FAIL regardless of accuracy.
    if critical_errors:
        print(f"\n=== CRITICAL FAILURES ({len(critical_errors)}) ===")
        for e in critical_errors:
            print(f"  [CRITICAL] {e}")
        print("\n=== RESULT: FAIL (critical check failed) ===")
        sys.exit(1)

    print(f"\n=== Accuracy: {accuracy:.1f}% ({checks_passed}/{checks_total} checks) ===")
    if accuracy >= 70:
        print("=== RESULT: PASS ===")
        sys.exit(0)
    else:
        print(f"=== RESULT: FAIL ({len(all_errors)} errors, accuracy {accuracy:.1f}% < 70) ===")
        sys.exit(1)


if __name__ == "__main__":
    main()
