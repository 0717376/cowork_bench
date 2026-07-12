"""Evaluation для yf-stock-comparison-word (RU / moex-finance).

Агент сравнивает три акции Московской биржи (SBER.ME, GAZP.ME, LKOH.ME) по
данным MCP `moex-finance` (схема moex.*). По реальному сиду доступно только
~65 торговых дней в окне 2026-02-25 .. 2026-05-26 (годовой ряд отсутствует),
поэтому период анализа фиксирован этим окном.

Все эталонные значения берутся НАПРЯМУЮ из БД (moex.stock_prices), а не из
устаревшего groundtruth_workspace — так eval остаётся честным относительно сида.

CRITICAL_CHECKS (семантические): любой их провал => общий FAIL независимо от
accuracy. Структурные проверки (наличие листов/таблиц, заголовки) — не
критические.
"""
import argparse
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym", user="eigent", password="camel")

# Тикеры MOEX и окно доступных данных
SYMBOLS = ["SBER.ME", "GAZP.ME", "LKOH.ME"]
DATE_START = "2026-02-25"
DATE_END = "2026-05-26"

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = False

# Семантические критические проверки (значения из БД, ключевые правила, ядро задачи).
CRITICAL_CHECKS = {
    "SBER.ME Start_Price (из БД)",
    "SBER.ME End_Price (из БД)",
    "SBER.ME Return_Pct (пересчёт из БД)",
    "GAZP.ME Start_Price (из БД)",
    "GAZP.ME End_Price (из БД)",
    "GAZP.ME Return_Pct (пересчёт из БД)",
    "LKOH.ME Start_Price (из БД)",
    "LKOH.ME End_Price (из БД)",
    "LKOH.ME Return_Pct (пересчёт из БД)",
    "Daily Prices: число строк == число торговых дней (±1)",
    "Daily Prices: значения close первого/последнего дня совпадают с БД",
    "Заключение называет лучшую акцию (из БД)",
    "Заключение называет худшую акцию (из БД)",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILED
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        tag = " (CRITICAL)" if name in CRITICAL_CHECKS else ""
        detail_str = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL]{tag} {name}{detail_str}")
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILED = True


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def sym_match(written, expected):
    """Сопоставление тикеров: принимаем как 'SBER', так и 'SBER.ME'."""
    if written is None:
        return False
    w = str(written).strip().lower()
    base = expected.lower().replace(".me", "")
    return w == expected.lower() or w == base or base in w


def get_expected_data():
    """Эталонные данные напрямую из moex.stock_prices."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    summary = {}
    series = {}  # symbol -> list[(date_str, close)]
    for sym in SYMBOLS:
        cur.execute(
            """
            SELECT date, close FROM moex.stock_prices
            WHERE symbol = %s AND date >= %s AND date <= %s
            ORDER BY date
            """,
            (sym, DATE_START, DATE_END),
        )
        rows = cur.fetchall()
        if rows:
            start_p = round(float(rows[0][1]), 2)
            end_p = round(float(rows[-1][1]), 2)
            ret = round(((end_p - start_p) / start_p) * 100, 2)
            summary[sym] = {"start": start_p, "end": end_p, "return": ret, "count": len(rows)}
            series[sym] = [(r[0].strftime("%Y-%m-%d"), round(float(r[1]), 2)) for r in rows]

    # Число торговых дней по эталонному тикеру SBER.ME
    cur.execute(
        """
        SELECT COUNT(DISTINCT date) FROM moex.stock_prices
        WHERE symbol = %s AND date >= %s AND date <= %s
        """,
        (SYMBOLS[0], DATE_START, DATE_END),
    )
    total_days = cur.fetchone()[0]

    conn.close()
    return summary, series, total_days


def check_excel(agent_workspace, summary, series, total_days):
    """Проверка вывода Excel (значения из БД)."""
    print("\n=== Проверка вывода Excel ===")

    agent_file = os.path.join(agent_workspace, "Stock_Comparison.xlsx")
    check("Excel file exists", os.path.isfile(agent_file), f"Ожидался {agent_file}")
    if not os.path.isfile(agent_file):
        return False

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)

    def get_sheet(wb, name):
        for s in wb.sheetnames:
            if s.strip().lower() == name.strip().lower():
                return wb[s]
        return None

    # ---------------- Daily Prices ----------------
    agent_dp = get_sheet(agent_wb, "Daily Prices")
    check("Sheet 'Daily Prices' exists", agent_dp is not None, f"Найдено: {agent_wb.sheetnames}")

    if agent_dp:
        header = [c.value for c in next(agent_dp.iter_rows(min_row=1, max_row=1))]
        header_lc = [str(h).strip().lower() if h is not None else "" for h in header]
        check("Daily Prices: столбец Date присутствует", "date" in header_lc, f"Заголовок: {header}")
        for sym in SYMBOLS:
            col = f"{sym.replace('.ME', '')}_Close".lower()
            check(f"Daily Prices: столбец {sym.replace('.ME', '')}_Close присутствует",
                  col in header_lc, f"Заголовок: {header}")

        dp_rows = list(agent_dp.iter_rows(min_row=2, values_only=True))
        # отбрасываем полностью пустые строки
        dp_rows = [r for r in dp_rows if r and any(c is not None and str(c).strip() != "" for c in r)]
        check("Daily Prices: число строк == число торговых дней (±1)",
              abs(len(dp_rows) - total_days) <= 1,
              f"Ожидалось ~{total_days}, получено {len(dp_rows)}")

        # CRITICAL: значения close первого и последнего дня совпадают с БД
        # Строим карту date->row из листа агента.
        date_idx = header_lc.index("date") if "date" in header_lc else 0
        col_idx = {}
        for sym in SYMBOLS:
            col = f"{sym.replace('.ME', '')}_Close".lower()
            if col in header_lc:
                col_idx[sym] = header_lc.index(col)
        row_by_date = {}
        for r in dp_rows:
            d = r[date_idx]
            d = str(d).strip()[:10] if d is not None else ""
            row_by_date[d] = r

        all_ok = True
        details = []
        for sym in SYMBOLS:
            if sym not in series or sym not in col_idx:
                all_ok = False
                details.append(f"{sym}: нет столбца/данных")
                continue
            for label, (d, expected_close) in (("first", series[sym][0]), ("last", series[sym][-1])):
                row = row_by_date.get(d)
                if row is None:
                    all_ok = False
                    details.append(f"{sym} {label} {d}: строка не найдена")
                    continue
                got = row[col_idx[sym]]
                if not num_close(got, expected_close, 1.0):
                    all_ok = False
                    details.append(f"{sym} {label} {d}: ожид {expected_close}, получ {got}")
        check("Daily Prices: значения close первого/последнего дня совпадают с БД",
              all_ok, "; ".join(details))

    # ---------------- Summary ----------------
    agent_sum = get_sheet(agent_wb, "Summary")
    check("Sheet 'Summary' exists", agent_sum is not None, f"Найдено: {agent_wb.sheetnames}")

    if agent_sum:
        agent_rows = list(agent_sum.iter_rows(min_row=2, values_only=True))
        agent_rows = [r for r in agent_rows if r and r[0] is not None and str(r[0]).strip() != ""]
        check("Summary: 3 строки акций", len(agent_rows) == 3, f"Получено {len(agent_rows)}")

        for sym in SYMBOLS:
            exp = summary.get(sym)
            matched = None
            for ar in agent_rows:
                if sym_match(ar[0], sym):
                    matched = ar
                    break
            if exp is None:
                check(f"{sym}: есть данные в БД", False, "нет данных в moex.stock_prices")
                continue
            if matched is None:
                check(f"{sym} найден в Summary", False)
                # критические проверки значений тоже зафиксируем как fail
                check(f"{sym} Start_Price (из БД)", False, "строка не найдена")
                check(f"{sym} End_Price (из БД)", False, "строка не найдена")
                check(f"{sym} Return_Pct (пересчёт из БД)", False, "строка не найдена")
                continue
            check(f"{sym} Start_Price (из БД)",
                  num_close(matched[1], exp["start"], 1.0),
                  f"Ожид {exp['start']}, получ {matched[1]}")
            check(f"{sym} End_Price (из БД)",
                  num_close(matched[2], exp["end"], 1.0),
                  f"Ожид {exp['end']}, получ {matched[2]}")
            check(f"{sym} Return_Pct (пересчёт из БД)",
                  num_close(matched[3], exp["return"], 0.5),
                  f"Ожид {exp['return']}, получ {matched[3]}")

    return True


def check_word_doc(agent_workspace, summary):
    """Проверка документа Word."""
    print("\n=== Проверка документа Word ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx установлен", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Stock_Analysis.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Ожидался {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Заголовок: принимаем RU и EN формулировки
    has_heading = False
    for p in doc.paragraphs:
        t = p.text.lower()
        ru_ok = ("акц" in t) and ("анализ" in t or "динамик" in t)
        en_ok = ("stock" in t) and ("performance" in t or "analysis" in t)
        if ru_ok or en_ok:
            has_heading = True
            break
    check("Документ содержит заголовок анализа акций", has_heading)

    # Таблица
    check("Документ содержит хотя бы одну таблицу", len(doc.tables) >= 1,
          f"Найдено {len(doc.tables)} таблиц")

    if doc.tables:
        table = doc.tables[0]
        check("В таблице заголовок + 3 строки данных", len(table.rows) >= 4,
              f"Получено {len(table.rows)} строк")
        all_text = " ".join(cell.text for row in table.rows for cell in row.cells).lower()
        for sym in SYMBOLS:
            base = sym.replace(".ME", "").lower()
            check(f"Таблица упоминает {sym}", base in all_text)

    full_text = " ".join(p.text for p in doc.paragraphs).lower()

    # Заключение про лучшую/худшую (RU+EN ключевые слова)
    ru_kw = any(k in full_text for k in ["лучш", "худш", "наибольш", "наименьш", "максим", "минимальн"])
    en_kw = any(k in full_text for k in ["best", "highest", "worst", "lowest"])
    check("Документ содержит заключение про лучшую/худшую акцию", ru_kw or en_kw)

    if summary:
        best = max(summary.items(), key=lambda x: x[1]["return"])[0]
        worst = min(summary.items(), key=lambda x: x[1]["return"])[0]
        check("Заключение называет лучшую акцию (из БД)",
              best.replace(".ME", "").lower() in full_text, f"ожид {best}")
        check("Заключение называет худшую акцию (из БД)",
              worst.replace(".ME", "").lower() in full_text, f"ожид {worst}")

    return True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("MOEX STOCK COMPARISON WORD - EVALUATION")
    print("=" * 70)

    summary, series, total_days = get_expected_data()

    check_excel(args.agent_workspace, summary, series, total_days)
    check_word_doc(args.agent_workspace, summary)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    print(f"\n=== ИТОГ ===")
    print(f"  Пройдено: {PASS_COUNT}")
    print(f"  Провалено: {FAIL_COUNT}")
    print(f"  Точность: {accuracy:.1f}%")

    if CRITICAL_FAILED:
        print("=== RESULT: FAIL (провалена критическая проверка) ===")
        sys.exit(1)

    if accuracy >= 70:
        print("=== RESULT: PASS ===")
        sys.exit(0)
    print("=== RESULT: FAIL (accuracy < 70%) ===")
    sys.exit(1)


if __name__ == "__main__":
    main()
