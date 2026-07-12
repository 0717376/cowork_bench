"""
Оценка задачи team-retreat (rzd + teamly + word + gcal + email).

Проверки:
  Word (Team_Retreat_Itinerary.docx): 4 раздела, корректный пункт назначения и
  даты, валидные номера поездов туда/обратно для обеих групп.
  Teamly: страница в пространстве TRIPS с планом поездки.
  Calendar: ровно 2 события ('Team Retreat Departs' / 'Team Retreat Returns').
  Email: письмо на hr@company.com с темой 'Team Retreat Travel Confirmed'.

Критические проверки (CRITICAL_CHECKS): любой провал => общий FAIL независимо
от accuracy. Структурные проверки (наличие файла, наличие события) — не критичны.

Номера поездов в схеме rzd (Великий Новгород):
  Москва ↔ НВГ туда: 818А / 820А; обратно: 819А / 821А
  СПб   ↔ НВГ туда: 822А / 824А; обратно: 823А / 825А
Агент легитимно ВЫБИРАЕТ один поезд на маршрут, поэтому проверяем,
что присутствует хотя бы один валидный код из соответствующего набора.
"""
import json
import os
import sys
import unicodedata
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

# Наборы валидных номеров поездов (сравниваем после normalize: '818А'->'818a')
MOW_OUTBOUND = ("818", "820")   # Москва -> Великий Новгород, 12.03
SPB_OUTBOUND = ("822", "824")   # Санкт-Петербург -> Великий Новгород, 12.03
MOW_RETURN   = ("819", "821")   # Великий Новгород -> Москва, 15.03
SPB_RETURN   = ("823", "825")   # Великий Новгород -> Санкт-Петербург, 15.03

PARTICIPANTS = [
    ("анна", "соколова"),
    ("борис", "иванов"),
    ("светлана", "орлова"),
    ("дмитрий", "захаров"),
    ("елена", "лебедева"),
]

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Критические (семантические) проверки.
CRITICAL_CHECKS = {
    "Word: раздел Outbound содержит валидный код поезда для обеих групп",
    "Word: раздел Return содержит валидный код поезда для обеих групп",
    "Word: Retreat Overview указывает Великий Новгород и обе даты",
    "Calendar: ровно 2 события Departs/Returns, в Departs все 5 участников",
    "Email на hr@company.com: тема и тело с участниками и кодами туда",
}


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def normalize(s: str) -> str:
    """Lowercase + схлопывание кириллица/латиница (А/A, С/C, ...),
    чтобы '818А' и '818A' сравнивались как равные."""
    s = s.lower()
    s = unicodedata.normalize("NFKD", s)
    table = str.maketrans({"а": "a", "о": "o", "е": "e", "р": "p",
                           "с": "c", "у": "y", "к": "k", "х": "x"})
    return s.translate(table)


def count_participants(text_lower: str) -> int:
    """Сколько участников упомянуто (по фамилии ИЛИ имени)."""
    n = 0
    for first, last in PARTICIPANTS:
        if last in text_lower or first in text_lower:
            n += 1
    return n


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def check_word(agent_workspace):
    print("\n=== Word: Team_Retreat_Itinerary.docx ===")

    docx_path = os.path.join(agent_workspace, "Team_Retreat_Itinerary.docx")
    if not os.path.exists(docx_path):
        record("Team_Retreat_Itinerary.docx существует", False, f"Не найден: {docx_path}")
        return
    record("Team_Retreat_Itinerary.docx существует", True)

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        record("Файл Word читается", False, str(e))
        return
    record("Файл Word читается", True)

    full_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    lower = full_text.lower()
    norm = normalize(full_text)

    # Структура: 4 раздела (RU+EN допускаются)
    has_overview = "retreat overview" in lower or "overview" in lower or "обзор" in lower
    has_outbound = "outbound" in lower or "туда" in lower or "поездка туда" in lower
    has_return = "return" in lower or "обратно" in lower
    has_notes = "schedule" in lower or "notes" in lower or "заметк" in lower or "расписан" in lower
    record("Документ: раздел Retreat Overview", has_overview, "нет 'Retreat Overview/Обзор'")
    record("Документ: раздел Outbound Journey", has_outbound, "нет 'Outbound/туда'")
    record("Документ: раздел Return Journey", has_return, "нет 'Return/обратно'")
    record("Документ: раздел Schedule Notes", has_notes, "нет 'Schedule Notes/заметки'")

    # Участники
    n_part = count_participants(lower)
    record("Документ упоминает участников (>=4 из 5)", n_part >= 4,
           f"найдено {n_part}/5")

    # --- CRITICAL: валидные коды поездов туда (обе группы) ---
    out_mow = any(t in norm for t in MOW_OUTBOUND)
    out_spb = any(t in norm for t in SPB_OUTBOUND)
    record("Word: раздел Outbound содержит валидный код поезда для обеих групп",
           out_mow and out_spb,
           f"MOW({'/'.join(MOW_OUTBOUND)})={out_mow}, SPB({'/'.join(SPB_OUTBOUND)})={out_spb}")

    # --- CRITICAL: валидные коды поездов обратно (обе группы) ---
    ret_mow = any(t in norm for t in MOW_RETURN)
    ret_spb = any(t in norm for t in SPB_RETURN)
    record("Word: раздел Return содержит валидный код поезда для обеих групп",
           ret_mow and ret_spb,
           f"MOW({'/'.join(MOW_RETURN)})={ret_mow}, SPB({'/'.join(SPB_RETURN)})={ret_spb}")

    # --- CRITICAL: Overview — пункт назначения + обе даты ---
    has_dest = "новгород" in lower or "novgorod" in lower
    date_out = ("2026-03-12" in lower or "12 март" in lower or "12.03" in lower
                or "march 12" in lower)
    date_ret = ("2026-03-15" in lower or "15 март" in lower or "15.03" in lower
                or "march 15" in lower)
    record("Word: Retreat Overview указывает Великий Новгород и обе даты",
           has_dest and date_out and date_ret,
           f"dest={has_dest}, out={date_out}, ret={date_ret}")


# ---------------------------------------------------------------------------
# Teamly
# ---------------------------------------------------------------------------

def check_teamly():
    print("\n=== Teamly: страница плана поездки в TRIPS ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT p.id, p.title, p.body FROM teamly.pages p
            JOIN teamly.spaces s ON s.id = p.space_id
            WHERE s.key = 'TRIPS'
              AND p.id > 3
              AND (p.title ILIKE '%ретрит%' OR p.title ILIKE '%retreat%'
                   OR p.title ILIKE '%новгород%'
                   OR p.body ILIKE '%818%' OR p.body ILIKE '%820%'
                   OR p.body ILIKE '%822%' OR p.body ILIKE '%824%')
            ORDER BY p.id DESC
        """)
        pages = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Создана страница в Teamly/TRIPS", False, f"DB error: {e}")
        return

    record("Создана страница в Teamly/TRIPS", len(pages) >= 1,
           f"найдено {len(pages)}: {[p[1] for p in pages]}")

    if pages:
        raw = str(pages[0][2] or "")
        body_lower = raw.lower()
        body_norm = normalize(raw)
        has_dest = "новгород" in body_lower or "novgorod" in body_norm
        has_train = any(t in body_norm for t in MOW_OUTBOUND + SPB_OUTBOUND)
        record("Страница Teamly содержит пункт назначения и коды поездов",
               has_dest and has_train,
               f"dest={has_dest}, train={has_train}")


# ---------------------------------------------------------------------------
# Calendar
# ---------------------------------------------------------------------------

def check_gcal():
    print("\n=== Calendar: события Departs / Returns ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime, end_datetime, description
          FROM gcal.events
         WHERE summary ILIKE '%retreat%' OR summary ILIKE '%ретрит%'
         ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    depart_events = [e for e in events if "depart" in (e[0] or "").lower()]
    return_events = [e for e in events if "return" in (e[0] or "").lower()]

    record("Есть событие 'Team Retreat Departs'", len(depart_events) >= 1,
           f"события: {[e[0] for e in events]}")
    record("Есть событие 'Team Retreat Returns'", len(return_events) >= 1,
           f"события: {[e[0] for e in events]}")

    # CRITICAL: ровно 2 события + в Departs упомянуты все 5 участников
    exactly_two = len(events) == 2
    departs_ok = False
    departs_names = 0
    if depart_events:
        desc = (depart_events[0][3] or "").lower()
        departs_names = count_participants(desc)
        departs_ok = departs_names >= 5
    record("Calendar: ровно 2 события Departs/Returns, в Departs все 5 участников",
           exactly_two and len(depart_events) >= 1 and len(return_events) >= 1 and departs_ok,
           f"всего={len(events)}, departs_names={departs_names}/5")


# ---------------------------------------------------------------------------
# Email
# ---------------------------------------------------------------------------

def check_email():
    print("\n=== Email: hr@company.com ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT to_addr, subject, body_text FROM email.messages
         WHERE to_addr::text ILIKE '%hr@company.com%'
    """)
    messages = cur.fetchall()
    cur.close()
    conn.close()

    record("Письмо на hr@company.com отправлено", len(messages) >= 1,
           f"найдено {len(messages)} писем")

    # CRITICAL: тема + участники + код туда из каждого маршрута
    subj_ok = any("travel confirmed" in (m[1] or "").lower()
                  or "retreat" in (m[1] or "").lower()
                  or "ретрит" in (m[1] or "").lower() for m in messages)
    body_combined = " ".join((m[2] or "") for m in messages)
    body_lower = body_combined.lower()
    body_norm = normalize(body_combined)
    n_names = count_participants(body_lower)
    out_mow = any(t in body_norm for t in MOW_OUTBOUND)
    out_spb = any(t in body_norm for t in SPB_OUTBOUND)
    record("Email на hr@company.com: тема и тело с участниками и кодами туда",
           bool(messages) and subj_ok and n_names >= 5 and out_mow and out_spb,
           f"subj={subj_ok}, names={n_names}/5, mow={out_mow}, spb={out_spb}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_teamly()
    check_gcal()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: проверки не выполнялись.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nИтого: {PASS_COUNT}/{total} проверок пройдено ({accuracy:.1f}%)")

    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    if critical_failed:
        print(f"  КРИТИЧЕСКИЕ ПРОВАЛЫ: {len(critical_failed)}")
        for n in critical_failed:
            print(f"    - {n}")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_failed": critical_failed,
    }
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if critical_failed:
        print("Overall: FAIL (провалена критическая проверка)")
        sys.exit(1)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
