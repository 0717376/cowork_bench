"""Evaluation for moex-portfolio-analysis-excel-word-email.

Checks:
1. Portfolio_Analysis.xlsx exists with Holdings sheet (5 rows) and Summary sheet
2. Investment_Memo.docx exists with required sections (RU/EN headings accepted)
3. Email sent to investment-committee@fund.example.com with correct subject + body

CRITICAL checks (any fail => overall FAIL regardless of accuracy %):
  - Holdings: каждый из 5 тикеров MOEX присутствует и его Current_Value ~= qty*close (tol 5%)
  - Summary: Total_Portfolio_Value ~= сумма всех current values (tol 5%)
  - Summary: Best_Performer / Worst_Performer совпадают с реальными max/min 30d доходностями
  - Summary: разбивка по секторам со стоимостью (Energy объединяет GAZP+LKOH)
  - Email: тело письма содержит итоговую стоимость портфеля (~total) и лучший/худший тикер

Prices и доходности читаются честно из moex.stock_prices; значения не хардкодятся.
Иначе PASS, если accuracy >= 70%.
"""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB = {"host": os.environ.get("PGHOST", "localhost"), "port": 5432, "dbname": "cowork_gym", "user": "eigent", "password": "camel"}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []

SYMBOLS = ["SBER.ME", "GAZP.ME", "LKOH.ME", "MGNT.ME", "MTSS.ME"]
HOLDINGS = {
    "SBER.ME": 100, "GAZP.ME": 200, "LKOH.ME": 10, "MGNT.ME": 15, "MTSS.ME": 150,
}
# Sector grouping (from moex.stock_info seed): Energy combines GAZP + LKOH.
SECTORS = {
    "SBER.ME": "Financial Services",
    "GAZP.ME": "Energy",
    "LKOH.ME": "Energy",
    "MGNT.ME": "Consumer Defensive",
    "MTSS.ME": "Communication Services",
}

CRITICAL_CHECKS = {
    "Holdings: все 5 тикеров с корректной Current_Value (~qty*close)",
    "Summary: Total_Portfolio_Value ~= сумма всех текущих стоимостей",
    "Summary: Best_Performer соответствует реальному max 30d доходности",
    "Summary: Worst_Performer соответствует реальному min 30d доходности",
    "Summary: разбивка по секторам со стоимостью (Energy = GAZP+LKOH)",
    "Email: тело содержит итоговую стоимость портфеля и лучший/худший тикер",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300] + "...") if len(detail) > 300 else detail
        marker = " [CRITICAL]" if name in CRITICAL_CHECKS else ""
        print(f"  [FAIL]{marker} {name}: {d}")
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILS.append(name)


def num_close(a, b, tol_pct=5.0):
    """Check if two values are within tol_pct percent of each other."""
    try:
        a, b = float(a), float(b)
    except (TypeError, ValueError):
        return False
    if abs(b) < 1e-6:
        return abs(a) < 0.01
    return abs(a - b) / abs(b) * 100 <= tol_pct


def get_latest_prices():
    """Latest close per symbol from moex.stock_prices."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    prices = {}
    for sym in SYMBOLS:
        cur.execute(
            "SELECT close FROM moex.stock_prices WHERE symbol=%s ORDER BY date DESC LIMIT 1",
            (sym,),
        )
        row = cur.fetchone()
        if row is not None:
            prices[sym] = float(row[0])
    cur.close()
    conn.close()
    return prices


def get_30d_returns():
    """Compute 30-day return % per symbol: (last_close / close_30d_ago - 1) * 100.
    'close 30 days ago' = the close on/just-before (max_date - 30 calendar days).
    """
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    returns = {}
    for sym in SYMBOLS:
        cur.execute("SELECT MAX(date) FROM moex.stock_prices WHERE symbol=%s", (sym,))
        max_date = cur.fetchone()[0]
        if max_date is None:
            continue
        cur.execute(
            "SELECT close FROM moex.stock_prices WHERE symbol=%s ORDER BY date DESC LIMIT 1",
            (sym,),
        )
        last_close = float(cur.fetchone()[0])
        cur.execute(
            "SELECT close FROM moex.stock_prices WHERE symbol=%s AND date <= (%s::date - INTERVAL '30 days') "
            "ORDER BY date DESC LIMIT 1",
            (sym, max_date),
        )
        r = cur.fetchone()
        if r is None:
            # fallback: oldest available close
            cur.execute(
                "SELECT close FROM moex.stock_prices WHERE symbol=%s ORDER BY date ASC LIMIT 1",
                (sym,),
            )
            r = cur.fetchone()
        old_close = float(r[0])
        if abs(old_close) > 1e-9:
            returns[sym] = (last_close / old_close - 1.0) * 100.0
    cur.close()
    conn.close()
    return returns


def check_excel(agent_ws, prices, returns):
    print("\n=== Check 1: Portfolio_Analysis.xlsx ===")
    path = os.path.join(agent_ws, "Portfolio_Analysis.xlsx")
    check("File Portfolio_Analysis.xlsx exists", os.path.isfile(path))
    if not os.path.isfile(path):
        return

    try:
        wb = openpyxl.load_workbook(path, data_only=True)
    except Exception as e:
        check("Excel is readable", False, str(e))
        return

    # ---------- Holdings sheet ----------
    holdings_ws = None
    for sname in wb.sheetnames:
        if "holdings" in sname.lower():
            holdings_ws = wb[sname]
            break
    check("Sheet 'Holdings' exists", holdings_ws is not None, f"Sheets: {wb.sheetnames}")

    holdings_value_ok = True
    if holdings_ws is not None:
        rows = list(holdings_ws.iter_rows(min_row=2, values_only=True))
        non_empty = [r for r in rows if any(c is not None for c in r)]
        check("Holdings has 5 data rows", len(non_empty) >= 5, f"Got {len(non_empty)}")

        all_text = " ".join(str(c) for row in non_empty for c in row if c is not None)
        for sym in SYMBOLS:
            check(f"Holdings contains symbol {sym}", sym in all_text)

        # CRITICAL: each symbol's current value ~= qty * latest close
        for sym in SYMBOLS:
            expected_val = round(HOLDINGS[sym] * prices.get(sym, 0), 2)
            found_val = False
            for row in non_empty:
                row_text = " ".join(str(c) for c in row if c is not None)
                if sym in row_text:
                    for c in row:
                        if num_close(c, expected_val, tol_pct=5.0):
                            found_val = True
                            break
                if found_val:
                    break
            if not found_val:
                holdings_value_ok = False
                print(f"    (Holdings {sym}: ожидалось ~{expected_val}, не найдено)")
        check("Holdings: все 5 тикеров с корректной Current_Value (~qty*close)",
              holdings_value_ok, "См. строки выше")

    # ---------- Summary sheet ----------
    summary_ws = None
    for sname in wb.sheetnames:
        if "summary" in sname.lower():
            summary_ws = wb[sname]
            break
    check("Sheet 'Summary' exists", summary_ws is not None, f"Sheets: {wb.sheetnames}")

    if summary_ws is not None:
        srows = list(summary_ws.iter_rows(values_only=True))
        all_text = " ".join(str(c) for row in srows for c in row if c is not None)
        all_lower = all_text.lower()

        # CRITICAL: total portfolio value
        total_value = sum(HOLDINGS[s] * prices.get(s, 0) for s in SYMBOLS)
        found_total = False
        for row in srows:
            for c in row:
                if num_close(c, total_value, tol_pct=5.0):
                    found_total = True
                    break
            if found_total:
                break
        check("Summary: Total_Portfolio_Value ~= сумма всех текущих стоимостей",
              found_total, f"Ожидалось ~{total_value:.2f}; контент: {all_text[:200]}")

        # CRITICAL: best / worst performer correctness (computed from price history)
        if returns:
            best_sym = max(returns, key=returns.get)
            worst_sym = min(returns, key=returns.get)
        else:
            best_sym = worst_sym = None

        # Locate Best/Worst rows and verify the named symbol matches the computed one.
        def symbol_in_labeled_row(labels):
            """Return set of MOEX symbols mentioned in any row whose text contains a label."""
            for row in srows:
                row_text = " ".join(str(c) for c in row if c is not None)
                rl = row_text.lower()
                if any(lbl in rl for lbl in labels):
                    found = {s for s in SYMBOLS if s in row_text}
                    if found:
                        return found
            return set()

        best_named = symbol_in_labeled_row(["best", "лучш"])
        worst_named = symbol_in_labeled_row(["worst", "худш"])
        check("Summary: Best_Performer соответствует реальному max 30d доходности",
              best_sym is not None and best_sym in best_named,
              f"Ожидался {best_sym} (доходность {returns.get(best_sym):.2f}%); найдено {best_named}" if best_sym else "Нет данных доходности")
        check("Summary: Worst_Performer соответствует реальному min 30d доходности",
              worst_sym is not None and worst_sym in worst_named,
              f"Ожидался {worst_sym} (доходность {returns.get(worst_sym):.2f}%); найдено {worst_named}" if worst_sym else "Нет данных доходности")

        # CRITICAL: per-sector breakdown with values (Energy = GAZP + LKOH)
        sector_values = {}
        for sym in SYMBOLS:
            sector_values.setdefault(SECTORS[sym], 0.0)
            sector_values[SECTORS[sym]] += HOLDINGS[sym] * prices.get(sym, 0)
        sector_ok = True
        sector_detail = []
        for sector, sval in sector_values.items():
            # a row that mentions the sector name and carries a value close to sval
            row_match = False
            for row in srows:
                row_text = " ".join(str(c) for c in row if c is not None)
                if sector.lower() in row_text.lower():
                    for c in row:
                        if num_close(c, sval, tol_pct=5.0):
                            row_match = True
                            break
                if row_match:
                    break
            if not row_match:
                sector_ok = False
                sector_detail.append(f"{sector}~{sval:.0f}")
        check("Summary: разбивка по секторам со стоимостью (Energy = GAZP+LKOH)",
              sector_ok, f"Не найдены/неверны: {sector_detail}")


def check_word(agent_ws):
    print("\n=== Check 2: Investment_Memo.docx ===")
    path = os.path.join(agent_ws, "Investment_Memo.docx")
    check("File Investment_Memo.docx exists", os.path.isfile(path))
    if not os.path.isfile(path):
        return

    try:
        from docx import Document
        doc = Document(path)
    except Exception as e:
        check("Word doc is readable", False, str(e))
        return

    full_text = " ".join(p.text for p in doc.paragraphs)
    # also pull table cell text (Holdings table)
    for tbl in doc.tables:
        for r in tbl.rows:
            for cell in r.cells:
                full_text += " " + cell.text
    full_text_lower = full_text.lower()

    check("Memo title содержит 'Portfolio' и 'Report'/'Analysis'",
          "portfolio" in full_text_lower and ("report" in full_text_lower or "analysis" in full_text_lower))
    # Headings accept RU or EN
    check("Memo has Holdings section (Holdings Overview / Обзор активов)",
          "holdings" in full_text_lower or "обзор активов" in full_text_lower or "позици" in full_text_lower)
    check("Memo has Performance section (Performance Analysis / Анализ результатов)",
          "performance" in full_text_lower or "анализ результат" in full_text_lower or "доходност" in full_text_lower)
    check("Memo has Recommendations section (Recommendations / Рекомендации)",
          "recommendation" in full_text_lower or "рекомендац" in full_text_lower)

    for sym in SYMBOLS:
        check(f"Memo mentions stock {sym}", sym in full_text)

    # Holdings table contains numeric values (non-critical structural sanity)
    has_numbers = any(
        any(ch.isdigit() for ch in cell.text)
        for tbl in doc.tables for r in tbl.rows for cell in r.cells
    )
    check("Memo Holdings table содержит числовые значения", has_numbers,
          "В таблицах не найдено чисел")


def check_email(prices, returns):
    print("\n=== Check 3: Email ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT subject, to_addr, from_addr, COALESCE(body_text, '') FROM email.messages
        WHERE to_addr::text ILIKE '%investment-committee%'
        LIMIT 10
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()

    check("Email sent to investment-committee@fund.example.com",
          len(rows) > 0, "No matching email found")
    if not rows:
        return

    subjects = [r[0] or "" for r in rows]
    check("Email subject contains 'Portfolio'",
          any("portfolio" in s.lower() for s in subjects), f"Subjects: {subjects}")
    check("Email subject contains 'Analysis' or 'Report'",
          any("analysis" in s.lower() or "report" in s.lower() for s in subjects), f"Subjects: {subjects}")

    # CRITICAL: body content
    bodies = " ".join((r[3] or "") for r in rows)
    bodies_lower = bodies.lower()

    total_value = sum(HOLDINGS[s] * prices.get(s, 0) for s in SYMBOLS)
    # find any numeric token in the body close to the total
    import re
    tokens = re.findall(r"[-+]?\d[\d\s., ]*", bodies)
    total_ok = False
    for t in tokens:
        # strip space/NBSP thousands separators (RU/EU formatting)
        s = t.replace(" ", "").replace(" ", "")
        candidates = set()
        # RU/EU: comma = decimal separator (e.g. "192071,03")
        if re.search(r",\d{1,2}$", s):
            candidates.add(re.sub(r",(\d{1,2})$", r".\1", s).replace(",", ""))
        # US: comma = thousands separator (e.g. "192,071.03")
        candidates.add(s.replace(",", ""))
        for cleaned in candidates:
            try:
                val = float(cleaned)
            except ValueError:
                continue
            if num_close(val, total_value, tol_pct=5.0):
                total_ok = True
                break
        if total_ok:
            break

    if returns:
        best_sym = max(returns, key=returns.get)
        worst_sym = min(returns, key=returns.get)
    else:
        best_sym = worst_sym = None
    perf_ok = best_sym is not None and best_sym in bodies and worst_sym in bodies

    check("Email: тело содержит итоговую стоимость портфеля и лучший/худший тикер",
          total_ok and perf_ok,
          f"total_ok={total_ok} (~{total_value:.0f}), best={best_sym} worst={worst_sym}, тело: {bodies[:200]}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=== Evaluation: moex-portfolio-analysis-excel-word-email ===")
    prices = get_latest_prices()
    returns = get_30d_returns()
    print(f"Latest prices from DB: {prices}")
    print(f"30d returns from DB: {{ {', '.join(f'{k}: {v:.2f}%' for k, v in returns.items())} }}")

    check_excel(args.agent_workspace, prices, returns)
    check_word(args.agent_workspace)
    check_email(prices, returns)

    print(f"\n=== SUMMARY: {PASS_COUNT} passed, {FAIL_COUNT} failed ===")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"pass": PASS_COUNT, "fail": FAIL_COUNT, "critical_fails": CRITICAL_FAILS}, f)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("No checks ran.")
        sys.exit(1)

    if CRITICAL_FAILS:
        print(f"\nCRITICAL FAILURES ({len(CRITICAL_FAILS)}): {CRITICAL_FAILS}")
        print("=> FAIL (критическая проверка провалена)")
        sys.exit(1)

    pct = PASS_COUNT / total * 100
    print(f"Score: {pct:.1f}%")
    sys.exit(0 if pct >= 70 else 1)


if __name__ == "__main__":
    main()
