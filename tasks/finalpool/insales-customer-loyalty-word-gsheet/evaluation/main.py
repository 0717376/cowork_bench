"""Evaluation for insales-customer-loyalty-word-gsheet (InSales / russified).

Blocking checks: Customer_Loyalty.docx structure & analytical content + gsheet deliverable.
CRITICAL_CHECKS: any failure => immediate FAIL (sys.exit(1)) before the accuracy gate.
Otherwise PASS requires accuracy >= 70 and no critical failure.

NOTE: customer data values (names) are russified CENTRALLY by db/zzz_wc_after_init.sql.
The agent doc is grepped INDEPENDENTLY (we never diff the English groundtruth text for names).
Section HEADINGS stay English by design (eval greps them).
"""
import argparse
import os
import re
import sys

try:
    from docx import Document
except ImportError:
    Document = None

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []

# Checks whose failure makes the whole task FAIL regardless of accuracy.
CRITICAL_CHECKS = set()


def record(name, passed, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    tag = "CRITICAL " if critical else ""
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {tag}{name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {tag}{name}{msg}")
        if critical:
            CRITICAL_FAILS.append(name)


def normalize_numbers(text):
    """Strip '$' and digit-grouping separators (comma/space/NBSP/narrow-NBSP)."""
    t = text.replace("$", "")
    t = re.sub(r"(?<=\d)[,\u00a0\u202f\u2009 ](?=\d{3}(?!\d))", "", t)
    return t.replace("\u00a0", " ").replace("\u202f", " ")


def check_word(agent_workspace):
    """Check Word document structure and analytical content."""
    print("\n=== Checking Word Document ===")

    agent_path = os.path.join(agent_workspace, "Customer_Loyalty.docx")

    if not os.path.isfile(agent_path):
        record("Word file exists", False, f"Not found: {agent_path}", critical=True)
        return None
    record("Word file exists", True)

    if Document is None:
        record("python-docx installed", False, "Cannot import docx", critical=True)
        return None

    doc = Document(agent_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Table cells too: answers laid out in a docx table are legitimate.
    full_text += "\n" + "\n".join(
        c.text for t in doc.tables for r in t.rows for c in r.cells)
    low = full_text.lower()
    # Numeric matching runs on normalized text: "$29,769.82" == "29769.82".
    norm = normalize_numbers(low)

    # ---- Structural (non-critical) ----
    record("Has title 'customer loyalty analysis report'",
           "customer loyalty analysis report" in low, "Title not found")
    record("Mentions report date 2026-03-06", "2026-03-06" in low, "Date not found")
    record("Has 'tier summary' section", "tier summary" in low, "'Tier Summary' not found")
    record("Has 'gold tier' section", "gold tier" in low, "'Gold Tier' not found")
    record("Has 'silver tier' section", "silver tier" in low, "'Silver Tier' not found")
    record("Has 'bronze tier' section", "bronze tier" in low, "'Bronze Tier' not found")

    # ---- CRITICAL: top spender pinned precisely ----
    # William Gonzalez -> Василий Гончаров (id=15), $4586.91 total, $1146.73 avg.
    # 'гончаров' alone is ambiguous (Евгения Гончаров also exists) -> anchor on the total.
    record("Top spender 'василий гончаров' present",
           "василий гончаров" in low, "Russified top-spender name not found", critical=True)
    record("Top spender exact total 4586.91 present",
           "4586.91" in norm, "Top-spender total not found", critical=True)

    # ---- CRITICAL: tier counts & total active customers ----
    # Counts pass either as a phrase ("14 customers"/"14 клиентов") or in row
    # context (table row / sentence: "Gold ... 14").
    def tier_count_found(tier_re, count):
        if f"{count} customer" in low or f"{count} клиент" in low:
            return True
        return bool(re.search(
            rf"(?:{tier_re})\D{{0,40}}(?<!\d){count}(?!\d)", norm))

    record("Total Active Customers = 43",
           "43" in norm, "Total active count 43 not found", critical=True)
    record("Gold tier count = 14 customers",
           tier_count_found("gold|золот", 14),
           "Gold count (14 customers) not found", critical=True)
    record("Silver tier count = 13 customers",
           tier_count_found("silver|серебр", 13),
           "Silver count (13 customers) not found", critical=True)
    record("Bronze tier count = 16 customers",
           tier_count_found("bronze|бронз", 16),
           "Bronze count (16 customers) not found", critical=True)

    # ---- CRITICAL: tier summary money totals (2dp) ----
    record("Gold tier total spending 29769.82",
           "29769.82" in norm, "Gold total not found", critical=True)
    record("Silver tier total spending 8001.58",
           "8001.58" in norm, "Silver total not found", critical=True)
    record("Bronze tier total spending 1646.77",
           "1646.77" in norm, "Bronze total not found", critical=True)
    record("Total revenue 39418.17",
           "39418.17" in norm, "Total revenue not found", critical=True)

    # ---- Per-tier ordering sanity (non-critical, spot-check top entries) ----
    # Gold #1 (4586.91) must appear before Gold #2 Scarlett->? (3362.01).
    record("Gold ordering: top total before 2nd",
           ("4586.91" in norm and "3362.01" in norm
            and norm.index("4586.91") < norm.index("3362.01")),
           "Gold tier not sorted desc by total")

    return full_text


def check_gsheet(top_name="Гончаров"):
    """Blocking: verify the 'Customer Loyalty Dashboard' / 'Tier Data' deliverable in gsheet schema."""
    print("\n=== Checking Google Sheet (gsheet schema) ===")
    try:
        import psycopg2
    except ImportError:
        record("psycopg2 available for gsheet check", False, "psycopg2 not installed")
        return
    try:
        conn = psycopg2.connect(
            host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym",
            user="eigent", password="camel"
        )
    except Exception as e:
        record("Connect to gsheet DB", False, str(e))
        return

    cur = conn.cursor()

    # Spreadsheet titled "Customer Loyalty Dashboard"
    cur.execute("SELECT id, title FROM gsheet.spreadsheets")
    sheets = cur.fetchall()
    ss_match = [s for s in sheets if "customer loyalty dashboard" in (s[1] or "").lower()]
    record("Spreadsheet 'Customer Loyalty Dashboard' exists",
           len(ss_match) > 0,
           f"Titles found: {[s[1] for s in sheets]}")

    ss_id = ss_match[0][0] if ss_match else (sheets[0][0] if sheets else None)

    # Sheet named "Tier Data"
    tier_sheet_id = None
    if ss_id is not None:
        cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (ss_id,))
        tsheets = cur.fetchall()
        tier_match = [t for t in tsheets if "tier data" in (t[1] or "").lower()]
        record("Sheet 'Tier Data' exists",
               len(tier_match) > 0,
               f"Sheet titles: {[t[1] for t in tsheets]}")
        if tier_match:
            tier_sheet_id = tier_match[0][0]
        elif tsheets:
            tier_sheet_id = tsheets[0][0]

    # Cells: header columns + ~43 data rows + russified top name
    if tier_sheet_id is not None:
        cur.execute(
            "SELECT value FROM gsheet.cells WHERE spreadsheet_id = %s AND sheet_id = %s",
            (ss_id, tier_sheet_id))
        vals = [r[0] for r in cur.fetchall() if r[0] is not None]
        blob = " | ".join(str(v) for v in vals)
        low_blob = blob.lower()

        expected_cols = ["customer_id", "name", "email", "orders",
                         "total_spent", "avg_order_value", "tier"]
        cols_ok = sum(1 for c in expected_cols if c in low_blob)
        record("Tier Data has the 7 header columns",
               cols_ok == len(expected_cols),
               f"Found {cols_ok}/7 columns")

        # 43 active customers: at least 43 non-empty rows worth of data.
        # Count distinct row indices if available, else fall back on email count.
        cur.execute(
            "SELECT COUNT(DISTINCT row_index) FROM gsheet.cells "
            "WHERE spreadsheet_id = %s AND sheet_id = %s",
            (ss_id, tier_sheet_id))
        try:
            distinct_rows = cur.fetchone()[0] or 0
        except Exception:
            distinct_rows = 0
        # header + 43 data rows => >= 44 distinct rows.
        record("Tier Data has ~43 data rows (>=44 incl. header)",
               distinct_rows >= 44,
               f"distinct rows = {distinct_rows}")

        record("Russified top name in gsheet",
               top_name.lower() in low_blob,
               "Top customer name not present in Tier Data")
    else:
        record("Tier Data sheet present for cell checks", False, "No sheet to inspect")

    cur.close()
    conn.close()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gsheet()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100.0) if total else 0.0

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Accuracy: {accuracy:.1f}%")

    if CRITICAL_FAILS:
        print("\n=== RESULT: FAIL (critical check failed) ===")
        for c in CRITICAL_FAILS:
            print(f"  [CRITICAL FAIL] {c}")
        sys.exit(1)

    if accuracy >= 70:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)
    else:
        print(f"\n=== RESULT: FAIL (accuracy {accuracy:.1f}% < 70%) ===")
        sys.exit(1)


if __name__ == "__main__":
    main()
