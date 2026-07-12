"""
Evaluation for yt-veritasium-arxiv-survey-word-gcal-email task.

Checks:
1. Science_Communication_Survey.docx exists
2. Word doc has >= 6 headings
3. Word doc text contains Veritasium and at least 3 science terms
4. Word doc text contains at least 3 paper titles or author names
5. GCal has 2 new events in April 2026
6. Email sent to seminar@science.edu
7. Email sent to collab@research.org

CRITICAL_CHECKS gate semantic completion: a correct agent passes them all,
a non-doer (no document / no new calendar event / no emails) fails regardless
of accuracy. Any critical failure forces sys.exit(1) before the accuracy gate.
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

# Semantic gates. Any failure here => hard FAIL regardless of accuracy.
# Paper titles/authors and the "Veritasium" service identity stay English,
# so these are stable markers even when the survey prose is Russian.
CRITICAL_CHECKS = {
    "Science_Communication_Survey.docx exists",
    "Word doc references >= 3 paper authors/titles",
    "GCal has >= 2 events added by agent in April 2026",
    "Email sent to seminar@science.edu",
    "Email sent to collab@research.org",
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Word Document Science_Communication_Survey.docx ===")
    docx_path = os.path.join(agent_workspace, "Science_Communication_Survey.docx")
    if not os.path.exists(docx_path):
        record("Science_Communication_Survey.docx exists", False, f"Not found at {docx_path}")
        return
    record("Science_Communication_Survey.docx exists", True)

    try:
        import docx
        doc = docx.Document(docx_path)
    except ImportError:
        # python-docx unavailable: cannot inspect content, but the critical
        # reference check below would be unverifiable. Treat as a hard miss.
        record("Word doc readable (python-docx)", False,
               "python-docx not installed; cannot verify document content")
        record("Word doc references >= 3 paper authors/titles", False,
               "python-docx not installed")
        return
    except Exception as e:
        record("Word doc readable (python-docx)", False, str(e))
        record("Word doc references >= 3 paper authors/titles", False, str(e))
        return

    # Count headings
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    record("Word doc has >= 6 headings", len(headings) >= 6,
           f"Found {len(headings)} headings: {[h.text[:50] for h in headings[:8]]}")

    # Check text content
    all_text = " ".join(p.text for p in doc.paragraphs).lower()

    has_veritasium = "veritasium" in all_text
    record("Word doc mentions Veritasium", has_veritasium, "Veritasium not found in text")

    science_terms = ["quantum", "evolution", "cognitive", "mathematical", "fluid", "fermi", "brain",
                     "paradox", "decoherence", "neuroplasticity", "biomimetic", "game theory"]
    found_terms = [t for t in science_terms if t in all_text]
    record("Word doc contains >= 3 science topic terms", len(found_terms) >= 3,
           f"Found: {found_terms}")

    # Check for paper author/title references. These are injected verbatim
    # (English) into arxiv.papers by preprocess, so the agent can only produce
    # them by actually reading the database and writing them into the survey.
    key_refs = ["sean carroll", "martin nowak", "daniel kahneman", "timothy gowers",
                "john dabiri", "anders sandberg", "michael merzenich",
                "many-worlds", "game theory and evolution", "cognitive biases",
                "paradoxes in mathematics", "fluid dynamics in nature", "fermi paradox",
                "neuroplasticity"]
    found_refs = [r for r in key_refs if r in all_text]
    record("Word doc references >= 3 paper authors/titles", len(found_refs) >= 3,
           f"Found refs: {found_refs}")

    # Executive summary heading. Accept English markers (kept verbatim per task)
    # and Russian equivalents in case the agent localises the heading.
    summary_markers = ["executive summary", "summary", "резюме", "краткое содержание", "обзор"]
    has_exec_summary = any(m in all_text for m in summary_markers)
    record("Word doc contains Executive Summary section", has_exec_summary,
           "No Executive Summary found")

    # References section. Accept English markers / arxiv ids / Russian equivalent.
    ref_markers = ["references", "arxiv:", "arxiv.org", "список литературы", "литература"]
    has_references = any(m in all_text for m in ref_markers)
    record("Word doc contains References section", has_references,
           "No References section found")

    # Check word count is substantial
    word_count = len(all_text.split())
    record("Word doc has substantial content (>= 500 words)", word_count >= 500,
           f"Word count: {word_count}")


def check_gcal():
    print("\n=== Check 2: GCal April 2026 Events ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime FROM gcal.events
        WHERE start_datetime >= '2026-04-01' AND start_datetime < '2026-05-01'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    # The preseeded "Lab Meeting" must NOT be enough to pass. Require the agent
    # to add at least 2 new April events (Science Communication Seminar +
    # Survey Review Session). A non-doer leaves only Lab Meeting -> FAIL.
    new_events = [e for e in events if "lab meeting" not in (e[0] or "").lower()]
    new_summaries = " ".join(e[0] or "" for e in new_events).lower()
    topic_hit = any(k in new_summaries for k in ("seminar", "survey", "review", "communication"))
    record("GCal has >= 2 events added by agent in April 2026",
           len(new_events) >= 2 and topic_hit,
           f"New events: {[e[0] for e in new_events]} (total April events: {len(events)})")


def check_emails_sent():
    print("\n=== Check 3: Emails Sent ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        # Check messages in Sent/SENT folders
        cur.execute("""
            SELECT m.to_addr FROM email.messages m
            JOIN email.folders f ON m.folder_id = f.id
            WHERE UPPER(f.name) = 'SENT'
        """)
        sent_rows = cur.fetchall()
        # Also check via sent_log join
        cur.execute("""
            SELECT m.to_addr FROM email.sent_log sl
            JOIN email.messages m ON sl.message_id = m.id
        """)
        sent_rows += cur.fetchall()
        sent_text = " ".join(str(row[0]) for row in sent_rows).lower()

        record("Email sent to seminar@science.edu",
               "seminar@science.edu" in sent_text,
               f"Sent entries: {len(sent_rows)}")
        record("Email sent to collab@research.org",
               "collab@research.org" in sent_text,
               f"Sent entries: {len(sent_rows)}")
    except Exception as e:
        record("Email sent to seminar@science.edu", False, str(e))
        record("Email sent to collab@research.org", False, str(e))
    finally:
        cur.close()
        conn.close()


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_gcal()
    check_emails_sent()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    # CRITICAL gate: any semantic failure is a hard FAIL before accuracy.
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    if critical_failed:
        print(f"CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")
        print("FAIL")
        sys.exit(1)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
