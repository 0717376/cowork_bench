"""
Evaluation script for gsheet-pdf-financial-report task.

Dynamically queries PostgreSQL (ClickHouse-backed sf_data schema) to compute
expected FY2024 values, then checks agent output files for correctness.

Critical checks (CRITICAL_CHECKS): any failure => overall FAIL regardless of
accuracy. Otherwise PASS requires accuracy >= 70%.
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

import psycopg2
import openpyxl

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Critical checks: any failure => overall FAIL regardless of accuracy.
# These reflect the SUBSTANCE of the deliverable (correct values from the DB),
# not merely structural presence.
CRITICAL_CHECKS = {
    "Excel Quarterly Revenue: 4 rows Q1-Q4 with revenue/count/avg matching DB",
    "Excel Top Products: top-3 names and revenue match DB descending order",
    "Word doc has all 4 section headings",
    "Word doc Executive Summary references total annual revenue",
    "Dashboard gsheet quarterly revenue values match DB",
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")


def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").replace("₽", "").replace("RUB", "").strip())
    except (ValueError, TypeError):
        return default


def get_quarterly_revenue():
    """Query PostgreSQL for FY2024 quarterly revenue."""
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute('''
        SELECT
          CASE
            WHEN EXTRACT(MONTH FROM o."ORDER_DATE"::date) BETWEEN 1 AND 3 THEN 'Q1'
            WHEN EXTRACT(MONTH FROM o."ORDER_DATE"::date) BETWEEN 4 AND 6 THEN 'Q2'
            WHEN EXTRACT(MONTH FROM o."ORDER_DATE"::date) BETWEEN 7 AND 9 THEN 'Q3'
            ELSE 'Q4'
          END as quarter,
          ROUND(SUM(o."TOTAL_AMOUNT"::float)::numeric, 2) as revenue,
          COUNT(*) as order_count
        FROM sf_data."SALES_DW__PUBLIC__ORDERS" o
        WHERE o."ORDER_DATE" >= '2024-01-01' AND o."ORDER_DATE" < '2025-01-01'
        GROUP BY quarter ORDER BY quarter
    ''')
    rows = cur.fetchall()
    result = []
    for q, rev, cnt in rows:
        rev = float(rev)
        cnt = int(cnt)
        avg = round(rev / cnt, 2)
        result.append((q, rev, cnt, avg))
    conn.close()
    return result


def get_top_products():
    """Query PostgreSQL for top 10 FY2024 products by revenue."""
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute('''
        SELECT p."PRODUCT_NAME", p."CATEGORY",
               SUM(o."QUANTITY"::int) as units,
               ROUND(SUM(o."TOTAL_AMOUNT"::float)::numeric, 2) as revenue
        FROM sf_data."SALES_DW__PUBLIC__ORDERS" o
        JOIN sf_data."SALES_DW__PUBLIC__PRODUCTS" p ON o."PRODUCT_ID" = p."PRODUCT_ID"
        WHERE o."ORDER_DATE" >= '2024-01-01' AND o."ORDER_DATE" < '2025-01-01'
        GROUP BY p."PRODUCT_NAME", p."CATEGORY"
        ORDER BY revenue DESC LIMIT 10
    ''')
    rows = [(r[0], r[1], int(r[2]), float(r[3])) for r in cur.fetchall()]
    conn.close()
    return rows


def check_excel(workspace, expected_quarters, expected_products):
    """Check FY2024_Financial_Analysis.xlsx for correctness."""
    print("\n--- Check 1: Excel File ---")
    xlsx_path = Path(workspace) / "FY2024_Financial_Analysis.xlsx"
    if not xlsx_path.exists():
        check("FY2024_Financial_Analysis.xlsx exists", False, f"not found in {workspace}")
        # Cannot proceed; mark dependent critical checks as failed.
        check("Excel Quarterly Revenue: 4 rows Q1-Q4 with revenue/count/avg matching DB", False, "no xlsx")
        check("Excel Top Products: top-3 names and revenue match DB descending order", False, "no xlsx")
        return
    check("FY2024_Financial_Analysis.xlsx exists", True)

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)

    has_qr = "Quarterly Revenue" in wb.sheetnames
    has_tp = "Top Products" in wb.sheetnames
    check("Excel 'Quarterly Revenue' sheet present", has_qr, wb.sheetnames)
    check("Excel 'Top Products' sheet present", has_tp, wb.sheetnames)

    # --- Check Quarterly Revenue ---
    qr_ok = False
    if has_qr:
        ws1 = wb["Quarterly Revenue"]
        rows1 = list(ws1.iter_rows(values_only=True))
        if len(rows1) >= 2:
            header1 = [str(h).strip() if h else "" for h in rows1[0]]
            expected_cols = ["Quarter", "Revenue", "Order_Count", "Avg_Order_Value"]
            cols_ok = all(c in header1 for c in expected_cols)
            check("Quarterly Revenue has required columns", cols_ok, header1)
            if cols_ok:
                idx = {col: header1.index(col) for col in expected_cols}
                data_rows = rows1[1:]
                qr_ok = (len(data_rows) == 4)
                if qr_ok:
                    for i, (exp_q, exp_rev, exp_cnt, exp_avg) in enumerate(expected_quarters):
                        row = data_rows[i]
                        q_val = str(row[idx["Quarter"]]).strip() if row[idx["Quarter"]] else ""
                        rev_val = safe_float(row[idx["Revenue"]])
                        cnt_val = row[idx["Order_Count"]]
                        avg_val = safe_float(row[idx["Avg_Order_Value"]])
                        if q_val != exp_q:
                            qr_ok = False
                            break
                        if rev_val is None or abs(rev_val - exp_rev) > 5.0:
                            qr_ok = False
                            break
                        try:
                            if cnt_val is None or abs(int(cnt_val) - exp_cnt) > 2:
                                qr_ok = False
                                break
                        except (ValueError, TypeError):
                            qr_ok = False
                            break
                        if avg_val is None or abs(avg_val - exp_avg) > 1.0:
                            qr_ok = False
                            break
    check("Excel Quarterly Revenue: 4 rows Q1-Q4 with revenue/count/avg matching DB",
          qr_ok, "values mismatch or wrong shape" if not qr_ok else "")

    # --- Check Top Products ---
    tp_top3_ok = False
    tp_full_ok = False
    if has_tp:
        ws2 = wb["Top Products"]
        rows2 = list(ws2.iter_rows(values_only=True))
        if len(rows2) >= 2:
            header2 = [str(h).strip() if h else "" for h in rows2[0]]
            expected_cols2 = ["Product_Name", "Category", "Units_Sold", "Revenue"]
            cols_ok = all(c in header2 for c in expected_cols2)
            check("Top Products has required columns", cols_ok, header2)
            if cols_ok:
                idx2 = {col: header2.index(col) for col in expected_cols2}
                data_rows2 = rows2[1:]

                def row_match(i):
                    exp_name, exp_cat, exp_units, exp_rev = expected_products[i]
                    row = data_rows2[i]
                    name_val = str(row[idx2["Product_Name"]]).strip() if row[idx2["Product_Name"]] else ""
                    rev_val = safe_float(row[idx2["Revenue"]])
                    units_val = row[idx2["Units_Sold"]]
                    if name_val[:30].lower() != exp_name[:30].lower():
                        return False
                    if rev_val is None or abs(rev_val - exp_rev) > 5.0:
                        return False
                    try:
                        if units_val is None or abs(int(units_val) - exp_units) > 2:
                            return False
                    except (ValueError, TypeError):
                        return False
                    return True

                if len(data_rows2) >= 3:
                    tp_top3_ok = all(row_match(i) for i in range(3))
                if len(data_rows2) == 10:
                    tp_full_ok = all(row_match(i) for i in range(10))
    check("Excel Top Products: top-3 names and revenue match DB descending order",
          tp_top3_ok, "top-3 mismatch")
    check("Top Products: all 10 rows match DB", tp_full_ok, "not all 10 rows correct")

    wb.close()


def check_gsheet(expected_quarters):
    """Check FY2024 Dashboard spreadsheet exists with quarterly revenue matching DB."""
    print("\n--- Check 2: Google Sheet ---")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        cur.execute(
            "SELECT id, title FROM gsheet.spreadsheets WHERE LOWER(title) LIKE '%fy2024%' OR LOWER(title) LIKE '%dashboard%'"
        )
        spreadsheets = cur.fetchall()
        ss_present = bool(spreadsheets)
        check("FY2024 Dashboard spreadsheet exists", ss_present,
              "no spreadsheet with 'FY2024'/'dashboard' in title")
        if not ss_present:
            check("Dashboard gsheet quarterly revenue values match DB", False, "no spreadsheet")
            return

        ss_id, ss_title = spreadsheets[0][0], spreadsheets[0][1]
        print(f"  Found spreadsheet: '{ss_title}' (id={ss_id})")

        cur.execute("SELECT COUNT(*) FROM gsheet.cells WHERE spreadsheet_id = %s", (ss_id,))
        cell_count = cur.fetchone()[0]
        check("Dashboard has >= 16 cells", cell_count >= 16, f"{cell_count} cells")

        cur.execute(
            """SELECT value FROM gsheet.cells
               WHERE spreadsheet_id = %s AND LOWER(value) IN ('q1', 'q2', 'q3', 'q4')""",
            (ss_id,),
        )
        found_quarters = cur.fetchall()
        check("Dashboard has all 4 quarter labels", len(found_quarters) >= 4,
              f"{len(found_quarters)} quarter labels")

        # SEMANTIC: every quarterly revenue value from the DB must appear among
        # the dashboard cell values (numeric, tolerance-matched). This verifies
        # the dashboard mirrors real figures, not just placeholder labels.
        cur.execute("SELECT value FROM gsheet.cells WHERE spreadsheet_id = %s", (ss_id,))
        cell_numbers = []
        for (v,) in cur.fetchall():
            fv = safe_float(v)
            if fv is not None:
                cell_numbers.append(fv)

        all_rev_present = True
        for (_q, exp_rev, _cnt, _avg) in expected_quarters:
            if not any(abs(cn - exp_rev) <= 5.0 for cn in cell_numbers):
                all_rev_present = False
                break
        check("Dashboard gsheet quarterly revenue values match DB", all_rev_present,
              "one or more quarterly revenue values missing from dashboard cells")
    finally:
        cur.close()
        conn.close()


def check_word(workspace, expected_quarters, total_revenue):
    """Check FY2024_Financial_Report.docx for required content."""
    print("\n--- Check 3: Word Document ---")
    docx_path = Path(workspace) / "FY2024_Financial_Report.docx"
    if not docx_path.exists():
        check("FY2024_Financial_Report.docx exists", False, f"not found in {workspace}")
        check("Word doc has all 4 section headings", False, "no docx")
        check("Word doc Executive Summary references total annual revenue", False, "no docx")
        return
    check("FY2024_Financial_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(str(docx_path))
        all_text = " ".join([p.text for p in doc.paragraphs])
        all_text_lower = all_text.lower()

        # Title (English identifier per task.md instructions).
        title_ok = "fy2024 annual financial report" in all_text_lower or any(
            "fy2024" in p.text.lower() and "annual" in p.text.lower() for p in doc.paragraphs
        )
        check("Word doc contains title", title_ok, "title not found")

        # Sections — headings stay English per task.md; accept English markers.
        required_sections = ["executive summary", "quarterly performance",
                             "product analysis", "outlook"]
        sections_ok = all(s in all_text_lower for s in required_sections)
        missing = [s for s in required_sections if s not in all_text_lower]
        check("Word doc has all 4 section headings", sections_ok, f"missing: {missing}")

        check("Word doc references 2024", "2024" in all_text, "no '2024'")

        # SEMANTIC: the document must reference the actual total annual revenue
        # (sum of quarterly revenue) within tolerance, not just any large number.
        # Match against numeric tokens in the prose, tolerant of grouping/decimals.
        numbers = re.findall(r'\d[\d\s .,]*\d|\d', all_text)
        total_int = round(total_revenue)

        def parse_locale_variants(tok):
            """Yield float interpretations of a numeric prose token for both
            US ("1,266,471.66") and RU ("1 266 471,66") formatting."""
            cands = []
            # US: strip spaces/NBSP and commas as thousands separators.
            cands.append(safe_float(tok.replace(" ", " ").replace(" ", "")))
            # RU: spaces/NBSP are thousands separators, last comma is the decimal.
            ru = tok.replace(" ", "").replace(" ", "")
            if "," in ru:
                head, _, dec = ru.rpartition(",")
                ru = head.replace(",", "") + "." + dec
            try:
                cands.append(float(re.sub(r'[^\d.\-]', '', ru)) if ru else None)
            except (ValueError, TypeError):
                cands.append(None)
            return [c for c in cands if c is not None]

        total_ref_ok = False
        for n in numbers:
            for fv in parse_locale_variants(n):
                # Accept full RUB or thousands/millions rounding within 1.5%.
                if abs(fv - total_revenue) <= max(50.0, total_revenue * 0.015):
                    total_ref_ok = True
                    break
                # Accept "in thousands" / "in millions" representations.
                if total_int >= 1000 and abs(fv - total_revenue / 1000.0) <= total_revenue / 1000.0 * 0.015 + 0.5:
                    total_ref_ok = True
                    break
                if total_int >= 1_000_000 and abs(fv - total_revenue / 1_000_000.0) <= total_revenue / 1_000_000.0 * 0.02 + 0.1:
                    total_ref_ok = True
                    break
            if total_ref_ok:
                break
        check("Word doc Executive Summary references total annual revenue", total_ref_ok,
              f"expected ~{total_revenue:,.2f} RUB referenced in text")

        # Revenue/выручка keyword present in ORIGINAL lowercased text (RU+EN).
        kw_ok = ("revenue" in all_text_lower or "выручк" in all_text_lower)
        check("Word doc references revenue/выручка", kw_ok, "no revenue keyword")

        check("Word doc is substantial (>= 200 chars)", len(all_text.strip()) >= 200,
              f"{len(all_text.strip())} chars")

    except ImportError:
        # python-docx missing: cannot verify semantics -> critical checks fail safe.
        check("Word doc has all 4 section headings", False, "python-docx unavailable")
        check("Word doc Executive Summary references total annual revenue", False,
              "python-docx unavailable")


def check_pdf(workspace):
    """Check FY2024_Financial_Report.pdf exists and has reasonable size."""
    print("\n--- Check 4: PDF File ---")
    pdf_path = Path(workspace) / "FY2024_Financial_Report.pdf"
    if not pdf_path.exists():
        check("FY2024_Financial_Report.pdf exists", False, f"not found in {workspace}")
        return
    size = pdf_path.stat().st_size
    check("FY2024_Financial_Report.pdf exists and >= 5KB", size >= 5000, f"{size} bytes")


def run_evaluation(workspace, res_log_file=None):
    global PASS_COUNT, FAIL_COUNT, FAILED_NAMES
    PASS_COUNT = 0
    FAIL_COUNT = 0
    FAILED_NAMES = []

    print("Fetching expected data from database...")
    quarters = get_quarterly_revenue()
    products = get_top_products()
    total_revenue = sum(q[1] for q in quarters)
    print(f"  Quarters: {len(quarters)}, Top Products: {len(products)}")
    print(f"  Total FY2024 Revenue: {total_revenue:,.2f} RUB")

    try:
        check_excel(workspace, quarters, products)
    except Exception as e:
        check("Excel Quarterly Revenue: 4 rows Q1-Q4 with revenue/count/avg matching DB", False, e)
        check("Excel Top Products: top-3 names and revenue match DB descending order", False, e)

    try:
        check_gsheet(quarters)
    except Exception as e:
        check("Dashboard gsheet quarterly revenue values match DB", False, e)

    try:
        check_word(workspace, quarters, total_revenue)
    except Exception as e:
        check("Word doc has all 4 section headings", False, e)
        check("Word doc Executive Summary references total annual revenue", False, e)

    try:
        check_pdf(workspace)
    except Exception as e:
        check("FY2024_Financial_Report.pdf exists and >= 5KB", False, e)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total else 0
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\n=== Results: {PASS_COUNT}/{total} passed ({accuracy:.1f}%) ===")
    if critical_failed:
        print(f"CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")

    if res_log_file:
        try:
            with open(res_log_file, "w") as f:
                json.dump({
                    "total_passed": PASS_COUNT, "total_checks": total,
                    "accuracy": accuracy, "critical_failed": critical_failed,
                }, f, indent=2)
        except Exception:
            pass

    success = (not critical_failed) and accuracy >= 70
    return success, f"Passed {PASS_COUNT}/{total} checks ({accuracy:.1f}%)"


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--launch_time", required=False, help="Launch time")
    args = parser.parse_args()

    try:
        success, message = run_evaluation(args.agent_workspace, args.res_log_file)
    except Exception as e:
        print(f"Error during evaluation: {e}")
        sys.exit(1)

    print(message)
    if success:
        print("\nPass all tests!")
        sys.exit(0)
    else:
        print("\nSome checks failed.")
        sys.exit(1)
