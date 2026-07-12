"""Evaluation для event-catering-logistics-optimization (RU-стек: kulinar/forms;
keep-foreign: excel/word/emails/google_calendar).

Проверяет пять частей задачи:
  - Excel Catering_Plan.xlsx: лист Menu (4 колонки в порядке, 6-8 строк),
    лист Dietary_Summary (счётчики по диетам совпадают с participant_list.xlsx),
    лист Summary (Attendees=500, Budget_Per_Person=1500, Total_Budget=750000).
  - Меню состоит из реальных блюд из базы kulinar.
  - Forms (gform.*): опрос пищевых предпочтений с >=3 вопросами.
  - Word Catering_Proposal.docx: меню + общий бюджет 750000 + раздел резервного плана.
  - Email на events@company.ru + событие(я) календаря с дедлайнами вендоров.

CRITICAL_CHECKS: любой их провал => FAIL (sys.exit(1)) даже при accuracy >= 70%.
Порог: accuracy >= 70% И нет критических провалов.

Счётчики по диетам читаются ЧЕСТНО из засеянного participant_list.xlsx
(initial_workspace), а не хардкодятся — задача остаётся валидной при пересеве данных.
"""
import argparse
import json
import os
import sys
import unicodedata
from collections import Counter

import psycopg2

DB = dict(
    host=os.environ.get("PGHOST", "localhost"),
    port=int(os.environ.get("PGPORT", "5432")),
    dbname=os.environ.get("PGDATABASE", "cowork_gym"),
    user="eigent",
    password="camel",
)

HERE = os.path.dirname(os.path.abspath(__file__))
SEED_PARTICIPANTS = os.path.join(HERE, "..", "initial_workspace", "participant_list.xlsx")

# Канонический набор блюд kulinar (источник: local_servers/kulinar-mcp -> all_recipes.json).
KULINAR_RECIPES = {
    "Салат Оливье", "Винегрет", "Сельдь под шубой", "Салат Мимоза", "Крабовый салат",
    "Греческий салат", "Салат с курицей и грибами", "Холодец", "Икра кабачковая",
    "Грибы маринованные", "Сало солёное", "Селёдка с луком", "Борщ",
    "Щи из квашеной капусты", "Солянка мясная", "Уха", "Окрошка", "Грибной суп",
    "Рассольник", "Куриный бульон с лапшой", "Бефстроганов", "Пельмени домашние",
    "Голубцы", "Котлеты домашние", "Жаркое в горшочках", "Курица в сметане",
    "Рыба запечённая по-русски", "Цыплёнок табака", "Гречка с тушёнкой",
    "Плов узбекский", "Картофельное пюре", "Гречневая каша", "Перловая каша",
    "Картофель отварной с укропом", "Рис отварной", "Пирожки с капустой жареные",
    "Пирожки с мясом печёные", "Блины тонкие", "Кулебяка с капустой и яйцом",
    "Расстегаи с рыбой", "Медовик", "Наполеон", "Сырники", "Пасха творожная",
    "Ватрушки с творогом", "Кисель ягодный", "Морс клюквенный",
    "Компот из сухофруктов", "Сбитень", "Квас домашний",
}

CRITICAL_CHECKS = {
    "Excel Menu: 4 колонки в порядке + 6-8 строк данных",
    "Excel Dietary_Summary: счётчики совпадают с participant_list.xlsx",
    "Excel Summary: Attendees=500, Budget_Per_Person=1500, Total_Budget=750000",
    "Рецепты Menu — реальные блюда из базы kulinar",
    "Forms: опрос пищевых предпочтений с >=3 вопросами",
    "Word: меню + общий бюджет 750000 + раздел резервного плана",
    "Email на events@company.ru + событие календаря с дедлайном вендора",
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []


def record(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILED
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = f": {str(detail)[:250]}" if detail else ""
        print(f"  [FAIL] {name}{d}")
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILED.append(name)


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def normalize(s: str) -> str:
    """Lowercase + схлопывание кириллических/латинских двойников (А/A, С/C...).
    Только для ID-сопоставлений, НЕ для поиска русских ключевых слов."""
    s = (s or "").lower()
    s = unicodedata.normalize("NFKD", s)
    table = str.maketrans({"а": "a", "о": "o", "е": "e", "р": "p",
                           "с": "c", "у": "y", "к": "k", "х": "x"})
    return s.translate(table)


def norm_recipe(s: str) -> str:
    s = (s or "").strip().lower().replace("ё", "е")
    return " ".join(s.split())


CANON_NORM = {norm_recipe(r) for r in KULINAR_RECIPES}


def get_sheet(wb, name):
    for s in wb.sheetnames:
        if s.strip().lower() == name.strip().lower():
            return wb[s]
    return None


# ---------------------------------------------------------------------------
# Эталонные счётчики по диетам — ЧЕСТНО из засеянного participant_list.xlsx
# ---------------------------------------------------------------------------
def expected_dietary_counts():
    try:
        import openpyxl
    except ImportError:
        return None
    if not os.path.isfile(SEED_PARTICIPANTS):
        return None
    try:
        wb = openpyxl.load_workbook(SEED_PARTICIPANTS, data_only=True)
    except Exception:
        return None
    ws = get_sheet(wb, "Participants") or wb.worksheets[0]
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return None
    header = [str(h).strip().lower() if h is not None else "" for h in rows[0]]
    try:
        di = next(i for i, h in enumerate(header) if "dietary" in h)
    except StopIteration:
        di = len(header) - 1
    cnt = Counter()
    for r in rows[1:]:
        if di < len(r) and r[di] is not None and str(r[di]).strip():
            cnt[str(r[di]).strip()] += 1
    return dict(cnt)


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------
def check_excel(agent_workspace):
    print("\n=== Проверка Excel Catering_Plan.xlsx ===")
    menu_cells = []
    try:
        import openpyxl
    except ImportError:
        record("openpyxl доступен", False, "openpyxl не установлен")
        return menu_cells

    agent_file = os.path.join(agent_workspace, "Catering_Plan.xlsx")
    record("Catering_Plan.xlsx существует", os.path.isfile(agent_file), f"Ожидался {agent_file}")
    if not os.path.isfile(agent_file):
        record("Excel Menu: 4 колонки в порядке + 6-8 строк данных", False, "нет файла")
        record("Excel Dietary_Summary: счётчики совпадают с participant_list.xlsx", False, "нет файла")
        record("Excel Summary: Attendees=500, Budget_Per_Person=1500, Total_Budget=750000", False, "нет файла")
        return menu_cells

    try:
        wb = openpyxl.load_workbook(agent_file, data_only=True)
    except Exception as e:
        record("Excel читается", False, str(e))
        record("Excel Menu: 4 колонки в порядке + 6-8 строк данных", False, "файл не читается")
        record("Excel Dietary_Summary: счётчики совпадают с participant_list.xlsx", False, "файл не читается")
        record("Excel Summary: Attendees=500, Budget_Per_Person=1500, Total_Budget=750000", False, "файл не читается")
        return menu_cells

    # --- Лист Menu ---
    print("\n--- Лист Menu ---")
    menu_ws = get_sheet(wb, "Menu")
    record("Лист 'Menu' существует", menu_ws is not None, f"Найдены: {wb.sheetnames}")
    if menu_ws is not None:
        headers = [c.value for c in list(menu_ws.rows)[0]] if menu_ws.max_row > 0 else []
        hl = [str(h).strip().lower() if h is not None else "" for h in headers]
        record("Menu: колонка Recipe_Name", any("recipe" in h for h in hl), f"Headers: {headers}")
        record("Menu: колонка Category", any("category" in h for h in hl), f"Headers: {headers}")
        record("Menu: колонка Servings", any("serving" in h for h in hl), f"Headers: {headers}")
        record("Menu: колонка Cost_Per_Person", any("cost" in h for h in hl), f"Headers: {headers}")

        data_rows = [row for row in menu_ws.iter_rows(min_row=2, values_only=True)
                     if any(v is not None and str(v).strip() for v in row)]
        expected = ["recipe", "category", "serving", "cost"]
        order_ok = all(i < len(hl) and expected[i] in hl[i] for i in range(4))
        rows_ok = 6 <= len(data_rows) <= 8
        record("Excel Menu: 4 колонки в порядке + 6-8 строк данных",
               order_ok and rows_ok,
               f"order_ok={order_ok}, rows={len(data_rows)}, headers={headers}")
        for row in data_rows:
            if row and row[0] is not None and str(row[0]).strip():
                menu_cells.append(str(row[0]).strip())

    # --- Лист Dietary_Summary ---
    print("\n--- Лист Dietary_Summary ---")
    ds_ws = get_sheet(wb, "Dietary_Summary")
    record("Лист 'Dietary_Summary' существует", ds_ws is not None, f"Найдены: {wb.sheetnames}")
    expected_counts = expected_dietary_counts()
    if ds_ws is not None and expected_counts:
        produced = {}
        for row in ds_ws.iter_rows(min_row=2, values_only=True):
            if row and row[0] is not None and str(row[0]).strip():
                produced[str(row[0]).strip().lower()] = row[1]
        # сравнение: каждый эталонный тип присутствует и его число совпадает
        all_ok = True
        details = []
        for diet, exp in expected_counts.items():
            got = produced.get(diet.strip().lower())
            ok = got is not None and num_close(got, exp, 0)
            details.append(f"{diet}: ждали {exp}, получили {got}")
            if not ok:
                all_ok = False
        record("Excel Dietary_Summary: счётчики совпадают с participant_list.xlsx",
               all_ok and len(produced) >= len(expected_counts), "; ".join(details))
    else:
        record("Excel Dietary_Summary: счётчики совпадают с participant_list.xlsx",
               False, f"лист есть={ds_ws is not None}, эталон={expected_counts}")

    # --- Лист Summary ---
    print("\n--- Лист Summary ---")
    sum_ws = get_sheet(wb, "Summary")
    record("Лист 'Summary' существует", sum_ws is not None, f"Найдены: {wb.sheetnames}")
    if sum_ws is not None:
        sdata = {}
        for row in sum_ws.iter_rows(min_row=2, values_only=True):
            if row and row[0]:
                sdata[str(row[0]).strip().lower()] = row[1]
        att_ok = any(("attendee" in k or "участник" in k) and num_close(v, 500, 0)
                     for k, v in sdata.items())
        bpp_ok = any(("budget_per_person" in k or "budget per person" in k or "на челов" in k)
                     and num_close(v, 1500, 0) for k, v in sdata.items())
        tot_ok = any(("total_budget" in k or "total budget" in k or "общий бюджет" in k)
                     and num_close(v, 750000, 1) for k, v in sdata.items())
        record("Excel Summary: Attendees=500, Budget_Per_Person=1500, Total_Budget=750000",
               att_ok and bpp_ok and tot_ok,
               f"att={att_ok}, bpp={bpp_ok}, total={tot_ok}, data={sdata}")
    else:
        record("Excel Summary: Attendees=500, Budget_Per_Person=1500, Total_Budget=750000",
               False, "нет листа Summary")

    return menu_cells


# ---------------------------------------------------------------------------
# Recipes vs kulinar
# ---------------------------------------------------------------------------
def check_recipes(menu_cells):
    print("\n=== Проверка рецептов против базы kulinar ===")
    if not menu_cells:
        record("Рецепты Menu — реальные блюда из базы kulinar", False, "нет названий блюд")
        return
    matched, unknown = 0, []
    for name in menu_cells:
        n = norm_recipe(name)
        if n in CANON_NORM or any(c in n or n in c for c in CANON_NORM):
            matched += 1
        else:
            unknown.append(name)
    record("Рецепты Menu — реальные блюда из базы kulinar",
           matched == len(menu_cells) and matched >= 6,
           f"совпало {matched}/{len(menu_cells)}; неизвестные: {unknown[:5]}")


# ---------------------------------------------------------------------------
# Forms (gform)
# ---------------------------------------------------------------------------
def check_forms():
    print("\n=== Проверка Forms (опрос предпочтений) ===")
    try:
        conn = psycopg2.connect(**DB)
    except Exception as e:
        record("Forms: опрос пищевых предпочтений с >=3 вопросами", False, str(e))
        return
    cur = conn.cursor()
    cur.execute("SELECT id, title FROM gform.forms")
    forms = cur.fetchall()
    record("Создана хотя бы одна форма", len(forms) >= 1, f"Найдено {len(forms)} форм")

    def title_matches(title):
        t = (title or "").lower()
        ru = any(k in t for k in ("пищев", "предпочт", "питани", "аллерг", "диет", "ограничен"))
        en = any(k in t for k in ("dietary", "preference", "menu", "diet"))
        return ru or en

    target = None
    for fid, title in forms:
        if title_matches(title):
            target = fid
            break
    if target is None and forms:
        best_fid, best_q = None, -1
        for fid, _t in forms:
            cur.execute("SELECT COUNT(*) FROM gform.questions WHERE form_id = %s", (fid,))
            qc = cur.fetchone()[0]
            if qc > best_q:
                best_fid, best_q = fid, qc
        target = best_fid

    q_count = 0
    if target is not None:
        cur.execute("SELECT COUNT(*) FROM gform.questions WHERE form_id = %s", (target,))
        q_count = cur.fetchone()[0]

    record("Forms: найден опрос с подходящим заголовком",
           any(title_matches(t) for _f, t in forms), f"Заголовки: {[t for _f, t in forms]}")

    topics_ok = False
    if target is not None:
        cur.execute(
            "SELECT lower(title) || ' ' || coalesce(lower(description),'') "
            "FROM gform.questions WHERE form_id = %s", (target,))
        qtext = " ".join(r[0] for r in cur.fetchall())
        has_diet = any(k in qtext for k in ("ограничен", "предпочт", "вегетариан", "веган",
                                            "халяль", "глютен", "diet", "restrict", "preference"))
        has_allergy = any(k in qtext for k in ("аллерг", "allerg"))
        has_comment = any(k in qtext for k in ("комментар", "пожелан", "дополнит",
                                               "comment", "request", "additional"))
        topics_ok = has_diet and has_allergy and has_comment
        record("Forms: вопросы покрывают ограничения/аллергии/комментарии",
               topics_ok, f"diet={has_diet}, allergy={has_allergy}, comment={has_comment}")
    else:
        record("Forms: вопросы покрывают ограничения/аллергии/комментарии", False, "форма не найдена")

    record("Forms: опрос пищевых предпочтений с >=3 вопросами",
           target is not None and any(title_matches(t) for _f, t in forms) and q_count >= 3,
           f"вопросов={q_count}")
    conn.close()


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------
def check_word(agent_workspace):
    print("\n=== Проверка Word Catering_Proposal.docx ===")
    try:
        from docx import Document
    except ImportError:
        record("Word: меню + общий бюджет 750000 + раздел резервного плана",
               False, "python-docx не установлен")
        return
    path = os.path.join(agent_workspace, "Catering_Proposal.docx")
    record("Catering_Proposal.docx существует", os.path.isfile(path), f"Ожидался {path}")
    if not os.path.isfile(path):
        record("Word: меню + общий бюджет 750000 + раздел резервного плана", False, "нет файла")
        return
    try:
        doc = Document(path)
    except Exception as e:
        record("Word читается", False, str(e))
        record("Word: меню + общий бюджет 750000 + раздел резервного плана", False, "файл не читается")
        return
    text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    tl = text.lower()
    nb = norm_recipe(text)

    dishes_found = sum(1 for c in CANON_NORM if c and c in nb)
    has_menu = dishes_found >= 3
    # бюджет: 750000 / 750 000 / 750к
    budget_variants = ["750000", "750 000", "750 000", "750.000"]
    has_budget = any(v in text for v in budget_variants)
    has_contingency = any(k in tl for k in ("резерв", "запасн", "контингенс", "backup",
                                            "contingency", "на случай"))
    record("Word: упомянуто >=3 блюда из меню", has_menu, f"найдено блюд={dishes_found}")
    record("Word: указан общий бюджет 750000", has_budget, "")
    record("Word: есть раздел резервного плана", has_contingency, "")
    record("Word: меню + общий бюджет 750000 + раздел резервного плана",
           has_menu and has_budget and has_contingency,
           f"menu={has_menu}, budget={has_budget}, contingency={has_contingency}")


# ---------------------------------------------------------------------------
# Email + Calendar
# ---------------------------------------------------------------------------
def check_email_calendar():
    print("\n=== Проверка Email + Calendar ===")
    try:
        conn = psycopg2.connect(**DB)
    except Exception as e:
        record("Email на events@company.ru + событие календаря с дедлайном вендора", False, str(e))
        return
    cur = conn.cursor()

    # Email: отправлено на events@company.ru (не от него же)
    cur.execute(
        """
        SELECT subject, body_text FROM email.messages
         WHERE to_addr::text ILIKE %s AND from_addr NOT ILIKE %s
        """,
        ("%events@company.ru%", "%events@company.ru%"),
    )
    rows = cur.fetchall()
    record("Письмо отправлено на events@company.ru", len(rows) >= 1, f"matched={len(rows)}")
    body = " ".join(((s or "") + " " + (b or "")) for s, b in rows)
    bl = body.lower()
    email_ok = (len(rows) >= 1 and ("500" in body)
                and any(k in bl for k in ("бюджет", "750", "budget"))
                and any(k in bl for k in ("меню", "блюд", "кейтеринг", "menu")))
    record("Email: в теле число участников, бюджет и меню", email_ok,
           f"rows={len(rows)}")

    # Calendar: событие с дедлайном вендора
    cur.execute("SELECT summary, COALESCE(description,'') FROM gcal.events")
    events = cur.fetchall()
    etext = " ".join((str(s) + " " + str(d)) for s, d in events).lower()
    has_vendor = any(k in etext for k in ("вендор", "поставщик", "дедлайн", "vendor", "deadline"))
    record("Calendar: создано хотя бы одно событие", len(events) >= 1, f"events={len(events)}")
    record("Calendar: событие упоминает вендора/дедлайн", has_vendor,
           f"summaries={[e[0] for e in events][:5]}")

    record("Email на events@company.ru + событие календаря с дедлайном вендора",
           len(rows) >= 1 and email_ok and len(events) >= 1 and has_vendor,
           f"email={len(rows)}, events={len(events)}, vendor={has_vendor}")
    conn.close()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("EVENT CATERING LOGISTICS OPTIMIZATION - EVALUATION")
    print("=" * 70)

    menu = check_excel(args.agent_workspace)
    check_recipes(menu or [])
    check_forms()
    check_word(args.agent_workspace)
    check_email_calendar()

    total = PASS_COUNT + FAIL_COUNT
    pct = 100.0 * PASS_COUNT / total if total else 0.0
    print("\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Accuracy: {pct:.1f}%")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": pct}
    if args.res_log_file:
        try:
            with open(args.res_log_file, "w") as f:
                json.dump(result, f, indent=2)
        except Exception:
            pass

    if CRITICAL_FAILED:
        print(f"  CRITICAL FAIL: {CRITICAL_FAILED}")
        print("  Overall: FAIL")
        sys.exit(1)

    overall = pct >= 70.0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
