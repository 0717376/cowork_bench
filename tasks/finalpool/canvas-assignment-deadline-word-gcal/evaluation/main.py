"""
Evaluation script for canvas-assignment-deadline-word-gcal task.

Канва не меняется (read-only live data в Canvas). Проверяется:
1. Excel Assignment_Deadlines_FFF2013J.xlsx — 2 листа с корректными данными
2. Word Assignment_Schedule_FFF2013J.docx — заголовок и таблица заданий
3. Google Calendar — события-напоминания "Assignment Due: ..."
4. Email на fff2013j.students@university.edu

Гейт: любой провал из CRITICAL_CHECKS => общий FAIL независимо от accuracy.
Иначе порог: accuracy >= 70%.
"""

import argparse
import json
import os
import sys
from datetime import datetime, timedelta

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "cowork_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Любой провал этих проверок => общий FAIL независимо от accuracy.
CRITICAL_CHECKS = {
    "Total_Assignments = 13",
    "Total_Points_Possible = 900",
    "Avg_Points_Per_Assignment correct",
    "All Assignments row count matches",
    "Email to fff2013j.students@university.edu found",
    "Assignment reminder events: count/title correct",
}


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_by_name(wb, name):
    for sname in wb.sheetnames:
        if sname.strip().lower() == name.strip().lower():
            return [[cell.value for cell in row] for row in wb[sname].iter_rows()]
    return None


# ============================================================================
# Check 1: Excel file
# ============================================================================

def check_excel(agent_workspace, groundtruth_workspace):
    print("\n=== Checking Assignment_Deadlines_FFF2013J.xlsx ===")

    try:
        import openpyxl
    except ImportError:
        record("openpyxl available", False, "pip install openpyxl")
        return False

    agent_file = os.path.join(agent_workspace, "Assignment_Deadlines_FFF2013J.xlsx")
    gt_file = os.path.join(groundtruth_workspace, "Assignment_Deadlines_FFF2013J.xlsx")

    if not os.path.isfile(agent_file):
        record("Excel file exists", False, f"Not found: {agent_file}")
        return False
    record("Excel file exists", True)

    if not os.path.isfile(gt_file):
        record("Groundtruth file exists", False, f"Not found: {gt_file}")
        return False

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    all_ok = True

    # Check All Assignments sheet
    a_all = load_sheet_by_name(agent_wb, "All Assignments")
    g_all = load_sheet_by_name(gt_wb, "All Assignments")
    record("Sheet 'All Assignments' exists", a_all is not None)
    if a_all is None:
        all_ok = False

    if a_all is not None and g_all is not None:
        a_data = [r for r in a_all[1:] if any(v is not None for v in r)]
        g_data = [r for r in g_all[1:] if any(v is not None for v in r)]
        record("All Assignments row count matches",
               len(a_data) == len(g_data),
               f"Expected {len(g_data)}, got {len(a_data)}")
        if len(a_data) != len(g_data):
            all_ok = False

        # Build lookup by assignment name
        a_lookup = {}
        for row in a_data:
            if row and row[0] is not None:
                a_lookup[str(row[0]).strip().lower()] = row

        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            name = str(g_row[0]).strip()
            key = name.lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                record(f"Assignment row: {name}", False, "Not found")
                all_ok = False
                continue
            record(f"Assignment row: {name}", True)

            # Points_Possible (col 1) — несовпадение должно ронять all_ok
            if len(g_row) > 1 and len(a_row) > 1:
                pts_ok = num_close(a_row[1], g_row[1], 0.01)
                record(f"{name}: Points_Possible correct", pts_ok,
                       f"got {a_row[1]}, expected {g_row[1]}")
                if not pts_ok:
                    all_ok = False

            # Due_Date (col 2) — формат YYYY-MM-DD, значение совпадает с GT
            if len(g_row) > 2 and len(a_row) > 2:
                a_due = str(a_row[2]).strip()[:10] if a_row[2] is not None else ""
                g_due = str(g_row[2]).strip()[:10] if g_row[2] is not None else ""
                due_ok = a_due == g_due
                record(f"{name}: Due_Date correct", due_ok,
                       f"got {a_due!r}, expected {g_due!r}")

    # Check Summary sheet
    a_summ = load_sheet_by_name(agent_wb, "Summary")
    g_summ = load_sheet_by_name(gt_wb, "Summary")
    record("Sheet 'Summary' exists", a_summ is not None)
    if a_summ is None:
        all_ok = False

    if a_summ is not None and g_summ is not None:
        a_data = [r for r in a_summ[1:] if any(v is not None for v in r)]
        g_data = [r for r in g_summ[1:] if any(v is not None for v in r)]

        a_lookup = {}
        for row in a_data:
            if row and row[0] is not None:
                a_lookup[str(row[0]).strip().lower()] = row

        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                record(f"Summary row: {g_row[0]}", False, "Not found")
                all_ok = False
                continue
            record(f"Summary row: {g_row[0]}", True)

            if key == "total_assignments":
                ok = num_close(a_row[1], g_row[1], 0)
                record("Total_Assignments = 13", ok,
                       f"got {a_row[1]}, expected {g_row[1]}")
                if not ok:
                    all_ok = False
            elif key == "total_points_possible":
                ok = num_close(a_row[1], g_row[1], 0)
                record("Total_Points_Possible = 900", ok,
                       f"got {a_row[1]}, expected {g_row[1]}")
                if not ok:
                    all_ok = False
            elif key == "avg_points_per_assignment":
                ok = num_close(a_row[1], g_row[1], 0.01)
                record("Avg_Points_Per_Assignment correct", ok,
                       f"got {a_row[1]}, expected {g_row[1]}")
                if not ok:
                    all_ok = False

    return all_ok


# ============================================================================
# Check 2: Word document
# ============================================================================

def check_word(agent_workspace):
    print("\n=== Checking Assignment_Schedule_FFF2013J.docx ===")

    docx_path = os.path.join(agent_workspace, "Assignment_Schedule_FFF2013J.docx")
    if not os.path.isfile(docx_path):
        record("Word file exists", False, f"Not found: {docx_path}")
        return False
    record("Word file exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        headings_text = " ".join(p.text for p in doc.paragraphs
                                 if p.style.name.startswith("Heading")).lower()

        record("Word doc has content", len(all_text.strip()) >= 100,
               f"Content length: {len(all_text.strip())}")
        record("Word doc heading mentions Finance or FFF",
               any(term in (all_text + headings_text) for term in
                   ["finance", "fff", "foundations", "assignment", "schedule"]),
               "Missing Finance/FFF content in doc")

        tables = doc.tables
        record("Word doc has at least 1 table", len(tables) >= 1,
               f"Found {len(tables)} tables")

        # Check table has assignment data: имена, даты YYYY-MM-DD, баллы
        if tables:
            table_text = " ".join(
                cell.text.lower()
                for row in tables[0].rows
                for cell in row.cells
            )
            record("Table has TMA assignments",
                   "tma" in table_text,
                   f"Table text: {table_text[:200]}")
            record("Table has CMA assignments",
                   "cma" in table_text,
                   f"Table text: {table_text[:200]}")
            import re
            record("Table has YYYY-MM-DD dates",
                   bool(re.search(r"\d{4}-\d{2}-\d{2}", table_text)),
                   "No ISO date found in table")

        return True

    except ImportError:
        size = os.path.getsize(docx_path)
        record("Word file has content (>3KB)", size > 3000, f"Size: {size} bytes")
        return size > 3000
    except Exception as e:
        record("Word file readable", False, str(e))
        return False


# ============================================================================
# Check 3: Google Calendar
# ============================================================================

# (assignment_name, due_date) для заданий с датой сдачи — читается честно
# из тех же GT-данных, что и Excel; используется для проверки заголовка,
# количества и смещения "за 7 дней до, 08:00".
EXPECTED_DUE = {
    "TMA 34873": "2013-10-20",
    "TMA 34874": "2013-11-17",
    "TMA 34875": "2014-01-05",
    "TMA 34876": "2014-02-09",
    "TMA 34877": "2014-03-23",
    "CMA 34878": "2014-05-25",
    "CMA 34879": "2014-05-25",
    "CMA 34880": "2014-05-25",
    "CMA 34881": "2014-05-25",
    "CMA 34882": "2014-05-25",
    "CMA 34883": "2014-05-25",
    "CMA 34884": "2014-05-25",
    "Final Exam 34885": "2014-05-25",
}


def check_gcal():
    print("\n=== Checking Google Calendar ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT summary, start_datetime, end_datetime
        FROM gcal.events
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    print(f"[check_gcal] Found {len(events)} calendar events.")

    record("At least 1 calendar event created",
           len(events) >= 1,
           f"Found {len(events)}")

    # События-напоминания с корректным заголовком "Assignment Due: <name>"
    due_events = [e for e in events
                  if e[0] and e[0].strip().lower().startswith("assignment due:")]
    record("Assignment reminder events: count/title correct",
           len(due_events) == len(EXPECTED_DUE),
           f"Found {len(due_events)} 'Assignment Due:' events, expected {len(EXPECTED_DUE)}")

    # По заданию: проверяем, что есть событие на каждое имя и что старт =
    # due_date - 7 дней в 08:00 (без жёсткой привязки к tz-строке в БД).
    title_map = {}
    for summary, start_dt, end_dt in due_events:
        nm = summary.split(":", 1)[1].strip().lower() if ":" in summary else ""
        title_map[nm] = (start_dt, end_dt)

    offset_ok = 0
    matched = 0
    for name, due in EXPECTED_DUE.items():
        hit = title_map.get(name.lower())
        if not hit:
            continue
        matched += 1
        start_dt, end_dt = hit
        if start_dt is None:
            continue
        try:
            if isinstance(start_dt, str):
                sdt = datetime.fromisoformat(start_dt.replace("Z", "+00:00"))
            else:
                sdt = start_dt
            sdt = sdt.replace(tzinfo=None)
            expected_start = datetime.strptime(due, "%Y-%m-%d") - timedelta(days=7)
            expected_start = expected_start.replace(hour=8, minute=0)
            if sdt.date() == expected_start.date() and sdt.hour == 8:
                offset_ok += 1
        except Exception:
            pass

    record("Reminder events titled per assignment",
           matched >= len(EXPECTED_DUE) - 1,
           f"matched {matched}/{len(EXPECTED_DUE)} by title")
    record("Reminder start = due_date - 7 days at 08:00",
           offset_ok >= max(1, len(EXPECTED_DUE) - 2),
           f"{offset_ok}/{len(EXPECTED_DUE)} events with correct 7-day/08:00 offset")

    return len(due_events) == len(EXPECTED_DUE)


# ============================================================================
# Check 4: Email
# ============================================================================

def check_emails():
    print("\n=== Checking Emails ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    print(f"[check_emails] Found {len(all_emails)} total emails.")
    record("At least 1 email sent", len(all_emails) >= 1, f"Found {len(all_emails)}")

    found_email = False
    for subject, from_addr, to_addr, body_text in all_emails:
        to_str = str(to_addr or "").lower()
        subject_lower = (subject or "").lower()
        # Получатель ОБЯЗАТЕЛЕН: адрес должен присутствовать в to_addr.
        if "fff2013j.students@university.edu" not in to_str:
            continue
        found_email = True
        record("Email to fff2013j.students@university.edu found", True)

        record("Email subject mentions assignment or deadline",
               "assignment" in subject_lower or "deadline" in subject_lower
               or "задани" in subject_lower or "дедлайн" in subject_lower
               or "срок" in subject_lower,
               f"Subject: {subject}")

        body_lower = (body_text or "").lower()
        # Тело письма на русском перечисляет задания — допускаем RU+EN маркеры.
        record("Email body lists assignments",
               any(term in body_lower for term in
                   ["tma", "cma", "assignment", "due date",
                    "задани", "срок", "дедлайн", "балл"]),
               "Body missing assignment list")
        break

    if not found_email:
        record("Email to fff2013j.students@university.edu found", False,
               f"Emails: {[(e[0], str(e[2])[:60]) for e in all_emails[:3]]}")

    return found_email


# ============================================================================
# Main
# ============================================================================

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    check_excel(args.agent_workspace, gt_dir)
    check_word(args.agent_workspace)
    check_gcal()
    check_emails()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\n=== Results: {PASS_COUNT}/{total} passed ({accuracy:.1f}%) ===")

    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    if critical_failed:
        print(f"CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")

    result = {
        "passed": PASS_COUNT,
        "failed": FAIL_COUNT,
        "accuracy": accuracy,
        "critical_failed": critical_failed,
        "success": (not critical_failed) and accuracy >= 70,
    }
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if critical_failed:
        print("FAIL (critical check failed)")
        sys.exit(1)
    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    print("FAIL")
    sys.exit(1)


if __name__ == "__main__":
    main()
