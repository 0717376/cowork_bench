"""Evaluation for insales-shipping-performance-gsheet-word-email.

Checks:
1. Google Sheet "Shipping Performance Dashboard" with zone data (at least 3 rows)
2. Word doc Shipping_Performance_Report.docx with heading, table, and paragraphs
3. Email to logistics@store.example.com from operations@store.example.com
   Subject: "Monthly Shipping Performance Report"
"""
import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="cowork_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILS = []

# Expected shipping data (from actual InSales / wc DB query; data values
# russified centrally by db/zzz_wc_after_init.sql, in sync with groundtruth)
EXPECTED_ZONES = [
    {"name": "стандартная доставка", "orders": 64, "revenue": 25369.91},
    {"name": "бесплатная доставка", "orders": 34, "revenue": 12734.76},
    {"name": "экспресс-доставка", "orders": 30, "revenue": 12673.21},
]


def check(name, condition, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS]{' [CRITICAL]' if critical else ''} {name}")
    else:
        FAIL_COUNT += 1
        if critical:
            CRITICAL_FAILS.append(name)
        detail_str = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL]{' [CRITICAL]' if critical else ''} {name}{detail_str}")


def num_close(a, b, tol=5.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, title FROM gsheet.spreadsheets
        WHERE LOWER(title) LIKE '%shipping%'
    """)
    sheets = cur.fetchall()
    check("Shipping Performance Dashboard spreadsheet exists", len(sheets) >= 1,
          f"Found {len(sheets)} matching spreadsheets")

    if sheets:
        ss_id = sheets[0][0]
        cur.execute("""
            SELECT c.value FROM gsheet.cells c
            JOIN gsheet.sheets s ON c.spreadsheet_id = s.spreadsheet_id AND c.sheet_id = s.id
            WHERE c.spreadsheet_id = %s
        """, (ss_id,))
        cells = cur.fetchall()
        all_values = " ".join(str(c[0]) for c in cells if c[0])
        all_lower = all_values.lower()

        # Collect numeric tokens present in the sheet for per-zone num_close()
        numeric_tokens = []
        for c in cells:
            if c[0] is None:
                continue
            tok = str(c[0]).replace(" ", "").replace(" ", "").replace(",", ".")
            tok = "".join(ch for ch in tok if ch.isdigit() or ch == "." or ch == "-")
            try:
                numeric_tokens.append(float(tok))
            except (TypeError, ValueError):
                pass

        # NON-CRITICAL structure: all three RU method names present
        zone_names_found = sum(
            1 for z in EXPECTED_ZONES
            if any(part in all_lower for part in z["name"].split())
        )
        check("GSheet has data for all 3 shipping methods (names)",
              zone_names_found >= 3,
              f"Found {zone_names_found}/3 shipping method names")
        check("GSheet contains Standard Shipping (стандартная)",
              "стандартная" in all_lower, "Standard not found")
        check("GSheet contains Express Shipping (экспресс)",
              "экспресс" in all_lower, "Express not found")
        check("GSheet contains Free Shipping (бесплатн)",
              "бесплатн" in all_lower, "Free not found")

        # CRITICAL: per-zone Order_Count for all three groups (64/34/30 together)
        counts_ok = all(
            any(num_close(t, z["orders"], tol=0.5) for t in numeric_tokens)
            for z in EXPECTED_ZONES
        )
        check("GSheet Order_Count correct for all 3 groups (64, 34, 30)",
              counts_ok,
              f"Expected 64/34/30; numeric tokens: {sorted(set(numeric_tokens))[:40]}",
              critical=True)

        # CRITICAL: per-zone Total_Revenue for all three groups via num_close()
        revenue_ok = all(
            any(num_close(t, z["revenue"], tol=5.0) for t in numeric_tokens)
            for z in EXPECTED_ZONES
        )
        check("GSheet Total_Revenue correct for all 3 zones (25369.91, 12734.76, 12673.21)",
              revenue_ok,
              f"Expected per-zone revenue; numeric tokens: {sorted(set(numeric_tokens))[:40]}",
              critical=True)

    cur.close()
    conn.close()


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    doc_path = os.path.join(agent_workspace, "Shipping_Performance_Report.docx")
    check("Shipping_Performance_Report.docx exists", os.path.isfile(doc_path),
          f"Expected at {doc_path}", critical=True)
    if not os.path.isfile(doc_path):
        return

    try:
        from docx import Document
        doc = Document(doc_path)
    except Exception as e:
        check("Word doc readable", False, str(e), critical=True)
        return

    full_text = " ".join(p.text.lower() for p in doc.paragraphs)
    headings = [p.text.lower() for p in doc.paragraphs
                if p.style.name.startswith("Heading")]

    check("Word doc has 'Shipping Performance Report' heading",
          any("shipping performance report" in h for h in headings) or
          "shipping performance report" in full_text,
          "Heading not found")

    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    check("Word doc has at least 3 paragraphs", para_count >= 3,
          f"Found {para_count} non-empty paragraphs")

    check("Word doc has a table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")

    if doc.tables:
        table = doc.tables[0]
        table_text = " ".join(
            cell.text.lower()
            for row in table.rows
            for cell in row.cells
        )
        # NON-CRITICAL: expected column headers present
        headers_ok = all(h in table_text for h in
                         ("zone_name", "order_count", "total_revenue", "avg_order_value"))
        check("Table contains expected headers (Zone_Name/Order_Count/Total_Revenue/Avg_Order_Value)",
              headers_ok,
              f"Headers not all found in table")
        # CRITICAL: all three RU method names present in the table
        methods_ok = ("стандартн" in table_text and "экспресс" in table_text
                      and "бесплатн" in table_text)
        check("Table contains all 3 RU shipping method names (стандартная, экспресс, бесплатн)",
              methods_ok,
              "Not all 3 method names in table",
              critical=True)

    check("Word doc mentions Standard Shipping (стандартная)",
          "стандартн" in full_text, "Standard shipping not mentioned")
    check("Word doc mentions Express Shipping (экспресс)",
          "экспресс" in full_text, "Express shipping not mentioned")


def check_emails():
    print("\n=== Checking Emails ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    conn.close()

    def parse_recipients(to_addr):
        if to_addr is None:
            return []
        if isinstance(to_addr, list):
            return [str(r).strip().lower() for r in to_addr]
        to_str = str(to_addr).strip()
        try:
            parsed = json.loads(to_str)
            if isinstance(parsed, list):
                return [str(r).strip().lower() for r in parsed]
            return [to_str.lower()]
        except (json.JSONDecodeError, TypeError):
            return [to_str.lower()]

    target = "logistics@store.example.com"
    found = None
    for subj, from_addr, to_addr, body in all_emails:
        recipients = parse_recipients(to_addr)
        if target in recipients:
            found = (subj, from_addr, to_addr, body)
            break

    check("Email sent to logistics@store.example.com", found is not None,
          critical=True)
    if found:
        subj, from_addr, to_addr, body = found
        check("Email from operations@store.example.com",
              "operations@store.example.com" in (from_addr or "").lower(),
              f"From: {from_addr}",
              critical=True)
        check("Subject contains shipping + performance",
              "shipping" in (subj or "").lower() and "performance" in (subj or "").lower(),
              f"Subject: {subj}",
              critical=True)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("WC SHIPPING PERFORMANCE GSHEET WORD EMAIL - EVALUATION")
    print("=" * 70)

    check_gsheet()
    check_word(args.agent_workspace)
    check_emails()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100.0) if total else 0.0

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Accuracy: {accuracy:.1f}%")

    if CRITICAL_FAILS:
        print(f"  CRITICAL FAILURES ({len(CRITICAL_FAILS)}):")
        for c in CRITICAL_FAILS:
            print(f"    - {c}")
        print(f"  Overall: FAIL (critical check failed)")
        sys.exit(1)

    overall = accuracy >= 70.0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
