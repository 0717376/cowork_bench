"""Evaluation для yf-stock-comparison-word-gcal (RU / moex-finance).

Агент готовит квартальный сравнительный отчёт по пяти акциям MOEX
(SBER.ME, GAZP.ME, LKOH.ME, TCSG.ME, MTSS.ME), данные берёт из MCP
`moex-finance` (схема moex.*). Затем создаёт событие в Google Calendar и
отправляет письмо.

CRITICAL_CHECKS (семантические): любой их провал => общий FAIL независимо от
accuracy. Структурные проверки (заголовки столбцов, наличие разделов) —
не критические.

Замечания:
- recommendationKey у всех тикеров MOEX = 'none', поэтому проверка рекомендаций
  не делается по конкретным значениям.
- currentPrice (снимок) может отличаться от последнего синтетического исторического
  close, поэтому цены проверяем «свободно» по префиксу, а критичность вешаем на
  наличие символов + RUB.
- email subject и адрес — английские (eval их грепает), тело письма — русское.
"""
import argparse
import os
import sys

import psycopg2

# Пять отслеживаемых тикеров MOEX (источник данных moex-finance)
TRACKED = ["SBER", "GAZP", "LKOH", "TCSG", "MTSS"]

# Полные символы в moex.stock_prices (с суффиксом .ME)
TRACKED_FULL = ["SBER.ME", "GAZP.ME", "LKOH.ME", "TCSG.ME", "MTSS.ME"]

# Обязательные английские заголовки столбцов (сохраняются на английском)
REQUIRED_HEADERS = ["Symbol", "Company_Name", "Current_Price",
                    "Analyst_Recommendation", "Sector"]

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = False

# Имена критических проверок (семантика, не структура)
CRITICAL_CHECKS = {
    "Word: документ содержит все пять символов MOEX (SBER, GAZP, LKOH, TCSG, MTSS)",
    "Word: упомянут RUB и >= 3 корректных цен-снимков (свободный префикс)",
    "GCal: событие 2026-04-10 14:00-15:30 с релевантным названием",
    "Email: письмо на analyst@investment.com с темой 'Stock Comparison Report - Q2 2026' и непустым телом",
}


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT, CRITICAL_FAILED
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        tag = " (CRITICAL)" if name in CRITICAL_CHECKS else ""
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL]{tag} {name}{msg}")
        if name in CRITICAL_CHECKS:
            CRITICAL_FAILED = True


def pg():
    return psycopg2.connect(host=os.environ.get("PGHOST", "localhost"), port=5432,
                            dbname="cowork_gym", user="eigent", password="camel")


def latest_close_prefixes():
    """Ожидаемые цены = последний close из moex.stock_prices (как их форсирует
    MCP moex-finance в currentPrice/regularMarketPrice). Возвращаем целочисленные
    префиксы для формат-устойчивого, но значимого сравнения с текстом документа."""
    conn = pg(); cur = conn.cursor()
    prefixes = []
    for sym in TRACKED_FULL:
        cur.execute(
            "SELECT close FROM moex.stock_prices WHERE symbol = %s "
            "ORDER BY date DESC LIMIT 1", (sym,))
        row = cur.fetchone()
        if row and row[0] is not None:
            prefixes.append(str(int(float(row[0]))))
    cur.close(); conn.close()
    return prefixes


def check_word_doc(agent_workspace):
    doc_path = os.path.join(agent_workspace, "Stock_Comparison_Report.docx")
    if not os.path.exists(doc_path):
        record("Word: документ содержит все пять символов MOEX (SBER, GAZP, LKOH, TCSG, MTSS)",
               False, "Stock_Comparison_Report.docx не найден в рабочей директории")
        record("Word: упомянут RUB и >= 3 корректных цен-снимков (свободный префикс)",
               False, "Stock_Comparison_Report.docx не найден")
        record("Word: присутствуют английские заголовки столбцов", False, "файл не найден")
        record("Word: присутствует заголовок 'Stock Comparison' (RU+EN)", False, "файл не найден")
        record("Word: присутствуют разделы Stock Overview / Analysis / Portfolio Strategy",
               False, "файл не найден")
        return
    try:
        from docx import Document
        doc = Document(doc_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text
        low = full_text.lower()

        # CRITICAL: все пять символов
        missing_syms = [s for s in TRACKED if s not in full_text]
        record("Word: документ содержит все пять символов MOEX (SBER, GAZP, LKOH, TCSG, MTSS)",
               len(missing_syms) == 0, f"отсутствуют: {missing_syms}" if missing_syms else "")

        # CRITICAL: RUB + >= 3 корректных цен.
        # Ожидаемые цены берём ЖИВЫМИ из moex.stock_prices (последний close),
        # т.к. MCP форсирует currentPrice = этому значению. Так проверка следит
        # за сидом и не устаревает.
        has_rub = "rub" in low or "руб" in low
        try:
            expected_prefixes = latest_close_prefixes()
        except Exception as e:
            expected_prefixes = []
            print(f"  [warn] не удалось получить цены из БД: {e}")
        price_hits = sum(1 for p in expected_prefixes if p in full_text)
        record("Word: упомянут RUB и >= 3 корректных цен-снимков (свободный префикс)",
               has_rub and price_hits >= 3,
               f"RUB={has_rub}, ожидаемые префиксы={expected_prefixes}, совпавших цен={price_hits}")

        # NON-critical: английские заголовки столбцов
        missing_hdr = [h for h in REQUIRED_HEADERS if h not in full_text]
        record("Word: присутствуют английские заголовки столбцов",
               len(missing_hdr) == 0, f"отсутствуют: {missing_hdr}" if missing_hdr else "")

        # NON-critical: заголовок (RU+EN)
        record("Word: присутствует заголовок 'Stock Comparison' (RU+EN)",
               ("stock comparison" in low or "сравнительн" in low), "")

        # NON-critical: разделы
        has_sections = all(s in low for s in ["stock overview", "analysis", "portfolio strategy"])
        record("Word: присутствуют разделы Stock Overview / Analysis / Portfolio Strategy",
               has_sections, "")
    except Exception as e:
        record("Word: документ содержит все пять символов MOEX (SBER, GAZP, LKOH, TCSG, MTSS)",
               False, f"ошибка чтения docx: {e}")
        record("Word: упомянут RUB и >= 3 корректных цен-снимков (свободный префикс)",
               False, f"ошибка чтения docx: {e}")


def check_gcal_event():
    try:
        conn = pg(); cur = conn.cursor()
        cur.execute("""
            SELECT summary, start_datetime, end_datetime
            FROM gcal.events
            WHERE start_datetime::date = '2026-04-10'
            ORDER BY start_datetime
        """)
        rows = cur.fetchall()
        cur.close(); conn.close()
    except Exception as e:
        record("GCal: событие 2026-04-10 14:00-15:30 с релевантным названием",
               False, f"ошибка БД: {e}")
        return

    if not rows:
        record("GCal: событие 2026-04-10 14:00-15:30 с релевантным названием",
               False, "событие на 2026-04-10 не найдено")
        return

    ok = False
    detail = ""
    for summary, sdt, edt in rows:
        s = (summary or "").lower()
        title_ok = any(k in s for k in ["портфел", "обзор", "portfolio", "review"])
        st = sdt.strftime("%H:%M") if sdt else ""
        et = edt.strftime("%H:%M") if edt else ""
        # task.md не задаёт таймзону. Агент в RU-контексте резонно выбирает
        # Europe/Moscow, и gcal MCP конвертирует наивное 14:00 MSK в 11:00 UTC
        # перед хранением. Поэтому принимаем 14:00-15:30 (наивно/UTC) ЛИБО
        # 11:00-12:30 UTC (тот же интервал, сдвинутый на смещение MSK).
        time_ok = ((st == "14:00" and et == "15:30") or
                   (st == "11:00" and et == "12:30"))
        if title_ok and time_ok:
            ok = True
            break
        detail = f"summary={summary!r}, start={st}, end={et}"
    record("GCal: событие 2026-04-10 14:00-15:30 с релевантным названием", ok, detail)


def check_email():
    try:
        conn = pg(); cur = conn.cursor()
        cur.execute("""
            SELECT subject, to_addr, body_text FROM email.messages
            WHERE to_addr::text ILIKE '%analyst@investment.com%'
            ORDER BY id DESC LIMIT 10
        """)
        rows = cur.fetchall()
        cur.close(); conn.close()
    except Exception as e:
        record("Email: письмо на analyst@investment.com с темой 'Stock Comparison Report - Q2 2026' и непустым телом",
               False, f"ошибка БД: {e}")
        return

    if not rows:
        record("Email: письмо на analyst@investment.com с темой 'Stock Comparison Report - Q2 2026' и непустым телом",
               False, "письмо на analyst@investment.com не найдено")
        return

    ok = False
    detail = ""
    for subject, to_addr, body in rows:
        subj = (subject or "")
        subj_ok = "Stock Comparison Report - Q2 2026" in subj
        body_ok = bool(body and body.strip())
        if subj_ok and body_ok:
            ok = True
            break
        detail = f"subject={subject!r}, body_len={len((body or '').strip())}"
    record("Email: письмо на analyst@investment.com с темой 'Stock Comparison Report - Q2 2026' и непустым телом",
           ok, detail)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()
    agent_ws = args.agent_workspace or os.path.join(os.path.dirname(__file__), "..", "groundtruth_workspace")

    print("  Проверка документа Word...")
    check_word_doc(agent_ws)
    print("  Проверка события Google Calendar...")
    check_gcal_event()
    print("  Проверка письма...")
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    print(f"\n  Пройдено {PASS_COUNT}/{total} ({accuracy:.1f}%)")

    if CRITICAL_FAILED:
        print("\n=== RESULT: FAIL (провалена критическая проверка) ===")
        sys.exit(1)

    if accuracy >= 70:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)
    else:
        print(f"\n=== RESULT: FAIL (accuracy {accuracy:.1f}% < 70%) ===")
        sys.exit(1)


if __name__ == "__main__":
    main()
