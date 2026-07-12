"""
Evaluation script for arxiv-research-report task.
Checks that LLM_Reasoning_Survey.docx and .pdf exist with the expected content.

Pass model:
  - accuracy >= 70% of all checks AND no CRITICAL check failed => PASS
  - any CRITICAL check failure => FAIL regardless of accuracy

CRITICAL checks (semantic — reflect the core deliverable):
  - all 5 target paper titles present AND zero noise topics present
  - at least one correct citation count AND at least one correct venue present
    (proves real scholarly data was retrieved, not hallucinated)
  - all four section labels present as styled-heading paragraphs
  - key author names (Jason Wei, Shunyu Yao, Xuezhi Wang) all present
  - LLM_Reasoning_Survey.pdf exists, > 5KB and starts with %PDF magic bytes

Usage:
  python -m evaluation.main --agent_workspace <path> --groundtruth_workspace <path> --launch_time <time>
"""
import argparse
import json
import os
import re
import sys


PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Names of CRITICAL checks. Any of these failing => overall FAIL.
CRITICAL_CHECKS = {
    "CRITICAL: all 5 target paper titles present and no noise topics",
    "CRITICAL: real citation count and venue present",
    "CRITICAL: four section labels are styled headings",
    "CRITICAL: key author names present",
    "CRITICAL: PDF exists, >5KB and starts with %PDF",
}


def check(name: str, condition: bool, detail: str = ""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        detail_truncated = (detail[:200] + "...") if len(detail) > 200 else detail
        print(f"  [FAIL] {name}: {detail_truncated}")


def normalize(text: str) -> str:
    """Normalize text for comparison: lowercase, collapse whitespace.

    NOTE: this is lowercase + whitespace-collapse only (no transliteration).
    All matched identifiers here (titles, authors, venues, domain terms) stay
    English, so we match against `normalized` for them. RU prose is not graded.
    """
    return re.sub(r'\s+', ' ', text.lower().strip())


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", type=str, required=True)
    parser.add_argument("--groundtruth_workspace", type=str, required=True)
    parser.add_argument("--launch_time", type=str, required=False)
    parser.add_argument("--res_log_file", type=str, required=False)
    args = parser.parse_args()

    docx_path = os.path.join(args.agent_workspace, "LLM_Reasoning_Survey.docx")
    pdf_path = os.path.join(args.agent_workspace, "LLM_Reasoning_Survey.pdf")

    # ── Check 1: Word document exists ────────────────────────────────────────
    check("LLM_Reasoning_Survey.docx exists", os.path.exists(docx_path),
          f"File not found at {docx_path}")

    if not os.path.exists(docx_path):
        _finish(args)
        return

    # Read the Word document
    try:
        from docx import Document
        doc = Document(docx_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        # Table cells too: answers laid out in a docx table are legitimate.
        full_text += "\n" + "\n".join(
            c.text for t in doc.tables for r in t.rows for c in r.cells)
    except Exception as e:
        check("Word document readable", False, str(e))
        _finish(args)
        return

    normalized = normalize(full_text)

    # Collect styled-heading paragraph texts (lowercased).
    headings = []
    for para in doc.paragraphs:
        if para.style and para.style.name and "heading" in para.style.name.lower():
            if para.text.strip():
                headings.append(para.text.lower())
    heading_blob = " ".join(headings)

    # ── Check 2: Minimum content length (NON-critical) ───────────────────────
    check("Document has at least 500 characters",
          len(full_text.strip()) >= 500,
          f"Document has {len(full_text.strip())} characters")

    # ── Check 3: Title mentions survey or reasoning (NON-critical) ───────────
    check("Document mentions survey or reasoning",
          "survey" in normalized or "reasoning" in normalized,
          "Neither 'survey' nor 'reasoning' found")

    # ── Check 4: Required section labels present somewhere (NON-critical) ─────
    has_intro = "introduction" in normalized
    has_lit_review = "literature review" in normalized or "literature" in normalized
    has_methodology = "methodology" in normalized
    has_conclusion = "conclusion" in normalized or "summary" in normalized

    check("Has Introduction section (text)", has_intro, "No 'Introduction' text found")
    check("Has Literature Review section (text)", has_lit_review, "No 'Literature Review' text found")
    check("Has Methodology Comparison section (text)", has_methodology, "No 'Methodology' text found")
    check("Has Conclusion section (text)", has_conclusion, "No 'Conclusion' text found")

    # ── CRITICAL: four section labels are STYLED headings ────────────────────
    # The word MCP emits Heading-styled paragraphs; require real structure,
    # not a degenerate body that merely lists the four words.
    styled_intro = "introduction" in heading_blob
    styled_lit = "literature review" in heading_blob or "literature" in heading_blob
    styled_method = "methodology" in heading_blob
    styled_concl = "conclusion" in heading_blob
    all_styled = styled_intro and styled_lit and styled_method and styled_concl
    check("CRITICAL: four section labels are styled headings",
          all_styled,
          f"Styled headings found: {sorted(headings)}; "
          f"intro={styled_intro} lit={styled_lit} method={styled_method} concl={styled_concl}")

    # ── Check 5: At least 4 styled headings (NON-critical) ───────────────────
    check("Has at least 4 styled headings",
          len(headings) >= 4,
          f"Found only {len(headings)} styled headings")

    # ── Check 6: All 5 target paper titles present (NON-critical, per-title) ──
    paper_titles = [
        "Chain-of-Thought Prompting Elicits Reasoning in Large Language Models",
        "Tree of Thoughts: Deliberate Problem Solving with Large Language Models",
        "Self-Consistency Improves Chain of Thought Reasoning in Language Models",
        "Process Supervision for Mathematical Reasoning",
        "Scaling LLM Reasoning with Reinforcement Learning",
    ]
    titles_found = [t for t in paper_titles if t.lower() in normalized]
    for title in paper_titles:
        check(f"Paper title present: {title[:50]}",
              title.lower() in normalized,
              "Title not found in document text")

    # ── Check 7: Noise topics NOT featured (NON-critical, per-topic) ─────────
    noise_topics = ["image classification", "federated learning", "protein structure"]
    noise_present = [t for t in noise_topics if t.lower() in normalized]
    for topic in noise_topics:
        check(f"Noise topic NOT present: {topic}",
              topic.lower() not in normalized,
              f"Noise topic '{topic}' found -- should not be included")

    # ── CRITICAL: correct paper selection (all 5 titles, zero noise) ─────────
    check("CRITICAL: all 5 target paper titles present and no noise topics",
          len(titles_found) == 5 and not noise_present,
          f"titles_found={len(titles_found)}/5, noise_present={noise_present}")

    # ── Check 8: Key author names present (NON-critical, per-author) ─────────
    key_authors = ["Jason Wei", "Shunyu Yao", "Xuezhi Wang"]
    authors_found = [a for a in key_authors if a.lower() in normalized]
    for author in key_authors:
        check(f"Author present: {author}",
              author.lower() in normalized,
              f"Author '{author}' not found")

    # ── CRITICAL: all key author names present ───────────────────────────────
    check("CRITICAL: key author names present",
          len(authors_found) == len(key_authors),
          f"Only found: {authors_found}")

    # ── Check 9: Key domain terms present (NON-critical, per-term) ───────────
    key_terms = ["chain-of-thought", "tree of thoughts", "self-consistency", "process supervision"]
    for term in key_terms:
        check(f"Key term present: {term}",
              term.lower() in normalized,
              f"Term '{term}' not found")

    # ── CRITICAL: real scholarly data — a source citation count AND a venue ──
    # OR across the 5 source numbers and 4 venues to tolerate phrasing variance.
    citation_counts = ["850", "420", "650", "280", "190"]
    venues = ["neurips", "iclr", "icml", "aaai"]
    citation_hit = next((c for c in citation_counts if c in normalized), None)
    venue_hit = next((v for v in venues if v in normalized), None)
    check("CRITICAL: real citation count and venue present",
          citation_hit is not None and venue_hit is not None,
          f"citation_hit={citation_hit}, venue_hit={venue_hit}")

    # ── Check 10: PDF exists with reasonable size (NON-critical) ─────────────
    check("LLM_Reasoning_Survey.pdf exists", os.path.exists(pdf_path),
          f"PDF not found at {pdf_path}")

    pdf_ok = False
    pdf_detail = "PDF not found"
    if os.path.exists(pdf_path):
        pdf_size = os.path.getsize(pdf_path)
        check("PDF file size > 5KB", pdf_size > 5000, f"PDF is only {pdf_size} bytes")
        magic = b""
        try:
            with open(pdf_path, "rb") as fh:
                magic = fh.read(5)
        except Exception as e:
            pdf_detail = f"read error: {e}"
        pdf_ok = pdf_size > 5000 and magic.startswith(b"%PDF")
        pdf_detail = f"size={pdf_size}, magic={magic!r}"

    # ── CRITICAL: PDF export deliverable ─────────────────────────────────────
    check("CRITICAL: PDF exists, >5KB and starts with %PDF",
          pdf_ok, pdf_detail)

    _finish(args)


def _finish(args):
    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0.0
    print(f"\nResults: {PASS_COUNT}/{total} passed, {FAIL_COUNT} failed ({accuracy:.1f}%)")

    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    if critical_failed:
        print(f"CRITICAL FAILURES: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_failed": critical_failed,
    }
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if critical_failed:
        print("FAIL (critical check failed)")
        sys.exit(1)
    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    print("FAIL")
    sys.exit(1)


if __name__ == "__main__":
    main()
