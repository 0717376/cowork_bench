"""Evaluation для terminal-kulinar-pw-nutrition-gsheet-word.

Проверки:
1. Облачная таблица "Nutrition Dashboard" с листами Recipe Comparison и Daily Plan.
2. Документ Wellness_Diet_Plan.docx с обязательными разделами (RU+EN ключевые слова).
3. Скрипт nutrition_calculator.py существует.

CRITICAL_CHECKS — содержательные проверки: любой их провал => общий FAIL
независимо от accuracy >= 70%. Структурные проверки (лист есть, заголовок есть)
остаются НЕкритическими.
"""
import argparse
import json
import os
import re
import sys
import unicodedata
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent", "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Содержательные (критические) проверки.
CRITICAL_CHECKS = {
    "Стандарты: значения со страницы попали в Daily Plan/docx (>=4 из 6 чисел)",
    "Recipe Comparison: >=9 строк И все обязательные заголовки столбцов",
    "Daily Plan: есть итоговая строка и Pct_Daily_Calories согласуется с total/2000",
    "docx: присутствуют все 4 обязательных раздела (стандарты/рецепты/план/дефицит)",
    "Recipe Comparison: >=9 названий блюд соответствуют реальным рецептам kulinar",
}

# Эталонные значения со справочной страницы (calories/protein/carbs/fat/fiber/sodium).
STANDARD_VALUES = [2000, 50, 300, 65, 25, 2300]

# Резервный список названий рецептов kulinar (на случай, если JSON недоступен в eval-окружении).
FALLBACK_KULINAR_NAMES = [
    'Бефстроганов', 'Блины тонкие', 'Борщ', 'Ватрушки с творогом', 'Винегрет',
    'Голубцы', 'Греческий салат', 'Гречка с тушёнкой', 'Гречневая каша', 'Грибной суп',
    'Грибы маринованные', 'Жаркое в горшочках', 'Икра кабачковая',
    'Картофель отварной с укропом', 'Картофельное пюре', 'Квас домашний',
    'Кисель ягодный', 'Компот из сухофруктов', 'Котлеты домашние', 'Крабовый салат',
    'Кулебяка с капустой и яйцом', 'Куриный бульон с лапшой', 'Курица в сметане',
    'Медовик', 'Морс клюквенный', 'Наполеон', 'Окрошка', 'Пасха творожная',
    'Пельмени домашние', 'Перловая каша', 'Пирожки с капустой жареные',
    'Пирожки с мясом печёные', 'Плов узбекский', 'Рассольник', 'Расстегаи с рыбой',
    'Рис отварной', 'Рыба запечённая по-русски', 'Салат Мимоза', 'Салат Оливье',
    'Салат с курицей и грибами', 'Сало солёное', 'Сбитень', 'Сельдь под шубой',
    'Селёдка с луком', 'Солянка мясная', 'Сырники', 'Уха', 'Холодец',
    'Цыплёнок табака', 'Щи из квашеной капусты',
]


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        print(f"  [FAIL] {name}: {str(detail)[:200]}")


def norm(s):
    """NFKD + lower для устойчивого сравнения названий блюд (ё->е, регистр, пробелы)."""
    s = unicodedata.normalize("NFKD", str(s)).lower()
    s = s.replace("ё", "е")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def safe_float(val):
    try:
        if val is None:
            return None
        v = re.sub(r"[^0-9.\-]", "", str(val).replace(",", "."))
        if v in ("", "-", ".", "-."):
            return None
        return float(v)
    except (ValueError, TypeError):
        return None


def load_kulinar_names():
    """Грузим названия рецептов из all_recipes.json; иначе fallback-список."""
    candidates = []
    here = os.path.dirname(os.path.abspath(__file__))
    # подняться к корню проекта и зайти в local_servers
    root = here
    for _ in range(8):
        candidates.append(os.path.join(root, "local_servers", "kulinar-mcp", "src", "data", "all_recipes.json"))
        root = os.path.dirname(root)
    candidates.append("/opt/local_servers/kulinar-mcp/src/data/all_recipes.json")
    for p in candidates:
        try:
            if os.path.isfile(p):
                with open(p, encoding="utf-8") as f:
                    data = json.load(f)
                names = [r.get("name") for r in data if r.get("name")]
                if names:
                    return [norm(n) for n in names]
        except Exception:
            continue
    return [norm(n) for n in FALLBACK_KULINAR_NAMES]


def get_conn():
    return psycopg2.connect(**DB_CONFIG)


def fetch_sheet_grid(cur, ss_id, sheet_id):
    """Возвращает dict[(row,col)] -> value для листа."""
    cur.execute("""
        SELECT row_index, col_index, value FROM gsheet.cells
        WHERE spreadsheet_id = %s AND sheet_id = %s
    """, (ss_id, sheet_id))
    grid = {}
    for r, c, v in cur.fetchall():
        grid[(r, c)] = v
    return grid


def check_gsheet():
    print("\n=== Проверка 1: Google Sheets Nutrition Dashboard ===")
    conn = get_conn()
    cur = conn.cursor()

    cur.execute("SELECT id, title FROM gsheet.spreadsheets")
    spreadsheets = cur.fetchall()
    dashboard = None
    for ss_id, title in spreadsheets:
        if title and "nutrition" in title.lower():
            dashboard = (ss_id, title)
            break

    check("Таблица Nutrition Dashboard существует", dashboard is not None,
          f"Найдены таблицы: {[s[1] for s in spreadsheets]}")

    if not dashboard:
        cur.close()
        conn.close()
        return

    ss_id = dashboard[0]
    cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (ss_id,))
    sheet_list = cur.fetchall()
    sheet_titles = [s[1].lower() for s in sheet_list]
    check("Минимум 2 листа", len(sheet_list) >= 2,
          f"Найдено {len(sheet_list)}: {[s[1] for s in sheet_list]}")

    has_recipe = any("recipe" in t or "comparison" in t for t in sheet_titles)
    has_daily = any("daily" in t or "plan" in t for t in sheet_titles)
    check("Есть лист Recipe Comparison", has_recipe, f"Листы: {sheet_titles}")
    check("Есть лист Daily Plan", has_daily, f"Листы: {sheet_titles}")

    # ----- Recipe Comparison -----
    recipe_sheet = next((s for s in sheet_list
                         if "recipe" in s[1].lower() or "comparison" in s[1].lower()), None)
    recipe_names_in_sheet = []
    if recipe_sheet:
        rgrid = fetch_sheet_grid(cur, ss_id, recipe_sheet[0])
        data_rows = sorted({r for (r, c) in rgrid if r > 0})
        row_count = len(data_rows)

        # заголовки (row 0)
        header_cells = [rgrid.get((0, c)) for c in sorted({c for (r, c) in rgrid if r == 0})]
        header_text = " ".join(str(h) for h in header_cells if h)
        header_low = header_text.lower()
        required_cols = ["recipe_name", "meal_type", "calories", "protein", "carbs", "fat", "fiber", "sodium"]
        cols_ok = all(rc in header_low for rc in required_cols)
        # CRITICAL: >=9 строк И все обязательные заголовки
        check("Recipe Comparison: >=9 строк И все обязательные заголовки столбцов",
              row_count >= 9 and cols_ok,
              f"строк={row_count}, заголовки='{header_text[:160]}'")

        # определить колонку Recipe_Name
        name_col = None
        for c in sorted({c for (r, c) in rgrid if r == 0}):
            if "recipe_name" in str(rgrid.get((0, c), "")).lower():
                name_col = c
                break
        if name_col is None:
            # фолбэк: первая колонка
            name_col = min((c for (r, c) in rgrid if r == 0), default=0)
        for r in data_rows:
            v = rgrid.get((r, name_col))
            if v:
                recipe_names_in_sheet.append(str(v))
    else:
        check("Recipe Comparison: >=9 строк И все обязательные заголовки столбцов",
              False, "лист Recipe Comparison не найден")

    # CRITICAL: названия блюд соответствуют реальным рецептам kulinar
    kulinar_names = load_kulinar_names()
    matched = 0
    for nm in recipe_names_in_sheet:
        nn = norm(nm)
        if any(kn and (kn in nn or nn in kn) for kn in kulinar_names):
            matched += 1
    check("Recipe Comparison: >=9 названий блюд соответствуют реальным рецептам kulinar",
          matched >= 9, f"совпало {matched} из {len(recipe_names_in_sheet)} названий")

    # ----- Daily Plan -----
    daily_sheet = next((s for s in sheet_list
                        if "daily" in s[1].lower() or "plan" in s[1].lower()), None)
    daily_grid = {}
    if daily_sheet:
        daily_grid = fetch_sheet_grid(cur, ss_id, daily_sheet[0])
        d_rows = sorted({r for (r, c) in daily_grid if r > 0})
        check("Daily Plan: минимум 3 строки приёмов пищи", len(d_rows) >= 3,
              f"найдено {len(d_rows)} строк данных")

        # карта заголовков -> col
        dheaders = {str(daily_grid.get((0, c), "")).lower(): c
                    for c in sorted({c for (r, c) in daily_grid if r == 0})}

        def find_col(*keys):
            for h, c in dheaders.items():
                if any(k in h for k in keys):
                    return c
            return None

        cal_col = find_col("calori", "kcal", "калори")
        pct_col = find_col("pct_daily_calories", "pct_daily", "pct", "процент")

        # найти итоговую строку (totals): Meal-колонка содержит total/итог,
        # либо последняя строка данных.
        meal_col = find_col("meal", "приём", "прием")
        totals_row = None
        for r in d_rows:
            mv = str(daily_grid.get((r, meal_col), "")).lower() if meal_col is not None else ""
            if any(k in mv for k in ("total", "итог", "сумм", "всего")):
                totals_row = r
                break

        # числовые строки приёмов пищи (не итог)
        meal_rows = [r for r in d_rows if r != totals_row]
        sum_calories = 0.0
        for r in meal_rows:
            cv = safe_float(daily_grid.get((r, cal_col))) if cal_col is not None else None
            if cv:
                sum_calories += cv

        pct_consistent = False
        detail = ""
        if totals_row is not None and pct_col is not None:
            # значение калорий в итоговой строке (если есть) предпочтительнее
            tot_cal = safe_float(daily_grid.get((totals_row, cal_col))) if cal_col is not None else None
            if not tot_cal:
                tot_cal = sum_calories
            pct_val = safe_float(daily_grid.get((totals_row, pct_col)))
            if tot_cal and pct_val is not None:
                expected = tot_cal / 2000.0 * 100.0
                # допускаем как долю (0.x), так и проценты
                if pct_val <= 5:
                    pct_val_pct = pct_val * 100.0
                else:
                    pct_val_pct = pct_val
                tol = max(8.0, expected * 0.15)
                pct_consistent = abs(pct_val_pct - expected) <= tol
                detail = f"total_cal={tot_cal}, pct={pct_val} (~{pct_val_pct:.1f}%), ожидалось ~{expected:.1f}%"
            else:
                detail = f"totals_row={totals_row}, cal={tot_cal}, pct={pct_val}"
        else:
            detail = f"totals_row={totals_row}, pct_col={pct_col}"

        # CRITICAL
        check("Daily Plan: есть итоговая строка и Pct_Daily_Calories согласуется с total/2000",
              totals_row is not None and pct_consistent, detail)
    else:
        check("Daily Plan: минимум 3 строки приёмов пищи", False, "лист Daily Plan не найден")
        check("Daily Plan: есть итоговая строка и Pct_Daily_Calories согласуется с total/2000",
              False, "лист Daily Plan не найден")

    # заголовки упоминают калории (по всей таблице, row 0)
    cur.execute("""
        SELECT value FROM gsheet.cells
        WHERE spreadsheet_id = %s AND row_index = 0
        ORDER BY col_index
    """, (ss_id,))
    header_cells = cur.fetchall()
    header_text = " ".join(str(c[0]) for c in header_cells).lower()
    check("Заголовки упоминают калории", "calori" in header_text or "kcal" in header_text,
          f"Заголовки: {header_text[:120]}")

    # ----- CRITICAL: значения стандартов из источника присутствуют в Daily Plan -----
    daily_blob = " ".join(str(v) for v in daily_grid.values())
    daily_numbers = set(re.findall(r"\d+", daily_blob))
    present = sum(1 for v in STANDARD_VALUES if str(v) in daily_numbers)
    # ослабляем: если в Daily Plan не нашлось — допроверим в docx (см. check_word)
    # сохраним результат глобально для совместной критической проверки
    global STANDARDS_IN_SHEET
    STANDARDS_IN_SHEET = present

    cur.close()
    conn.close()


STANDARDS_IN_SHEET = 0


def check_word(workspace):
    print("\n=== Проверка 2: Wellness_Diet_Plan.docx ===")
    path = os.path.join(workspace, "Wellness_Diet_Plan.docx")
    if not os.path.exists(path):
        check("Документ Word существует", False, f"Не найден: {path}")
        # критические doc-проверки тоже падают
        check("docx: присутствуют все 4 обязательных раздела (стандарты/рецепты/план/дефицит)",
              False, "docx отсутствует")
        check("Стандарты: значения со страницы попали в Daily Plan/docx (>=4 из 6 чисел)",
              STANDARDS_IN_SHEET >= 4, f"в таблице найдено {STANDARDS_IN_SHEET}/6, docx отсутствует")
        return
    check("Документ Word существует", True)

    try:
        from docx import Document
        doc = Document(path)
        # включаем текст параграфов И таблиц
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        full_text = " ".join(parts)
        low = full_text.lower()
    except Exception as e:
        check("Документ Word читается", False, str(e))
        check("docx: присутствуют все 4 обязательных раздела (стандарты/рецепты/план/дефицит)",
              False, str(e))
        check("Стандарты: значения со страницы попали в Daily Plan/docx (>=4 из 6 чисел)",
              STANDARDS_IN_SHEET >= 4, f"docx не читается, в таблице {STANDARDS_IN_SHEET}/6")
        return

    def has(*keys):
        return any(k in low for k in keys)

    # некритические RU+EN проверки разделов/ключевых слов
    sec_wellness = has("благополуч", "диет", "wellness", "diet plan")
    sec_standards = has("стандарт", "рекоменд", "нормы", "норма", "standard", "recommended")
    sec_breakfast = has("завтрак", "breakfast")
    sec_lunch = has("обед", "lunch")
    sec_dinner = has("ужин", "dinner")
    sec_calories = has("калори", "calori", "ккал", "kcal")
    sec_gap = has("дефицит", "недостаток", "восполн", "gap", "supplement", "supplemented")

    check("Документ упоминает благополучие/план питания", sec_wellness, full_text[:120])
    check("Документ упоминает нормы/рекомендации", sec_standards, full_text[:120])
    check("Документ упоминает завтрак", sec_breakfast, full_text[:120])
    check("Документ упоминает обед", sec_lunch, full_text[:120])
    check("Документ упоминает ужин", sec_dinner, full_text[:120])
    check("Документ упоминает калории", sec_calories, full_text[:120])
    check("Документ содержит раздел gap-анализа (дефицит/восполнение)", sec_gap, full_text[:200])

    # CRITICAL: все 4 обязательных раздела
    section_standards = sec_standards
    section_recipes = sec_breakfast and sec_lunch and sec_dinner
    section_plan = has("план питания", "суточный план", "meal plan", "daily plan", "рекомендуемый")
    section_gap = sec_gap
    check("docx: присутствуют все 4 обязательных раздела (стандарты/рецепты/план/дефицит)",
          section_standards and section_recipes and section_plan and section_gap,
          f"стандарты={section_standards}, рецепты(зав/об/уж)={section_recipes}, "
          f"план={section_plan}, дефицит={section_gap}")

    # CRITICAL: значения стандартов из источника (>=4 из 6) — в docx ИЛИ в таблице
    doc_numbers = set(re.findall(r"\d+", full_text))
    present_doc = sum(1 for v in STANDARD_VALUES if str(v) in doc_numbers)
    present_total = max(present_doc, STANDARDS_IN_SHEET)
    check("Стандарты: значения со страницы попали в Daily Plan/docx (>=4 из 6 чисел)",
          present_total >= 4,
          f"в docx {present_doc}/6, в таблице {STANDARDS_IN_SHEET}/6 (нужно >=4)")


def check_script(workspace):
    print("\n=== Проверка 3: nutrition_calculator.py ===")
    path = os.path.join(workspace, "nutrition_calculator.py")
    check("nutrition_calculator.py существует", os.path.exists(path))


def check_reverse_validation(workspace):
    print("\n=== Обратная валидация ===")
    path = os.path.join(workspace, "Wellness_Diet_Plan.docx")
    if os.path.isfile(path):
        try:
            from docx import Document
            doc = Document(path)
            full_text = " ".join(p.text for p in doc.paragraphs).lower()
            check("В документе нет шаблонных заглушек",
                  "[insert" not in full_text and "todo" not in full_text
                  and "xxx" not in full_text and "[вставьте" not in full_text,
                  "Найден шаблонный текст-заглушка")
        except Exception:
            pass

    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT title, COUNT(*) FROM gsheet.spreadsheets
            WHERE LOWER(title) LIKE '%%nutrition%%'
            GROUP BY title HAVING COUNT(*) > 1
        """)
        dupes = cur.fetchall()
        check("Нет дублей таблицы Nutrition Dashboard", len(dupes) == 0,
              f"Найдены дубли: {dupes}")
        cur.close()
        conn.close()
    except Exception:
        pass


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_gsheet()
    check_word(args.agent_workspace)
    check_script(args.agent_workspace)
    check_reverse_validation(args.agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: проверки не выполнены.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    print(f"\nИтого: {PASS_COUNT}/{total} проверок пройдено ({accuracy:.1f}%)")
    if critical_failed:
        print(f"КРИТИЧЕСКИЕ ПРОВАЛЫ: {len(critical_failed)}")
        for n in critical_failed:
            print(f"  - {n}")

    result = {
        "total_passed": PASS_COUNT, "total_checks": total,
        "accuracy": accuracy, "critical_failed": critical_failed,
    }
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if critical_failed:
        print("FAIL (критический провал)")
        sys.exit(1)
    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
