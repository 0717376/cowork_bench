"""Проверка (evaluation) для задачи canvas-quiz-analysis-gsheet-email.

Порог прохождения: accuracy >= 70 И ни одна CRITICAL-проверка не провалена.

CRITICAL-проверки (семантические, любой провал => немедленный FAIL):
  - Достаточное число строк из эталона (_quiz_data.txt) совпадает с GSheet
    по (Course_Name, Quiz_Title, Avg_Score, округл. до 1 знака) в пределах
    допуска — проверяет, что средний балл реально вычислен правильно.
  - Pass_Rate в GSheet совпадает с эталонным определением (% попыток с баллом
    >= 60% от Max_Score, округл. до 1 знака) для выборки тестов — проверяет
    основное аналитическое правило.
  - Число различимых строк-тестов в GSheet попадает в ожидаемый диапазон
    (~76, с допуском) — проверяет правило фильтрации из task.md
    (заменяет слабую проверку '>= 10').
  - Письмо отправлено на academic_coordinator@university.edu с темой,
    содержащей 'quiz', и непустым телом (> 30 символов).
  - Документ Quiz_Performance_Summary.docx существует, содержит реальную
    таблицу (doc.tables непусто) и раздел рекомендаций (RU+EN).

Структурные/мягкие проверки (таблица/лист/файл существуют, колонки на месте,
письмо найдено) помечены как НЕ критические. Названия курсов/тестов читаются
'честно' из общей фикстуры Canvas и не хардкодятся по значениям.
"""
import argparse
import ast
import json
import os
import re
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "cowork_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Семантические проверки: любой провал => немедленный FAIL независимо от accuracy.
CRITICAL_CHECKS = {
    "Средние баллы (Avg_Score) в GSheet совпадают с эталоном для выборки тестов",
    "Pass_Rate в GSheet совпадает с эталоном для выборки тестов",
    "Число строк-тестов в GSheet в ожидаемом диапазоне (~76)",
    "Письмо отправлено координатору с темой про quiz и непустым телом",
    "Word-документ содержит реальную таблицу и раздел рекомендаций",
}

# Ожидаемое число строк-тестов после фильтрации (task.md: ~76).
EXPECTED_ROWS = 76
ROW_TOLERANCE = 12  # допускаем 64..88
# Минимум совпадений из эталона, чтобы считать вычисления корректными.
MIN_GT_MATCH = 8
SCORE_TOL = 0.2  # допуск на округление/способ агрегации (1 знак)


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def load_groundtruth(groundtruth_workspace):
    """Загружает эталонные кортежи тестов из _quiz_data.txt.

    Формат строки:
      ('Course_Name', 'Quiz_Title', Total_Attempts, Avg_Score, Max_Score, Pass_Rate)
    Возвращает список dict с float-значениями.
    """
    candidates = []
    if groundtruth_workspace:
        candidates.append(os.path.join(groundtruth_workspace, "_quiz_data.txt"))
    # Запасной вариант: рядом с этим файлом проверки.
    here = os.path.dirname(os.path.abspath(__file__))
    candidates.append(os.path.join(here, "..", "groundtruth_workspace", "_quiz_data.txt"))

    path = next((p for p in candidates if os.path.isfile(p)), None)
    if not path:
        return []

    rows = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or not line.startswith("("):
                continue
            # ast.literal_eval не понимает вызовы Decimal('..') —
            # разворачиваем их в обычные числовые литералы.
            norm = re.sub(r"Decimal\(\s*'([^']*)'\s*\)", r"\1", line)
            try:
                tup = ast.literal_eval(norm)
            except Exception:
                continue
            if not isinstance(tup, tuple) or len(tup) < 6:
                continue
            rows.append({
                "course": str(tup[0]),
                "quiz": str(tup[1]),
                "attempts": int(tup[2]),
                "avg_score": float(tup[3]),
                "max_score": float(tup[4]),
                "pass_rate": float(tup[5]),
            })
    return rows


def _to_float(val):
    try:
        return float(str(val).strip())
    except Exception:
        return None


def check_gsheet(gt_rows):
    print("\n=== Проверка Google Sheet ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        # Находим таблицу 'Quiz Performance Tracker'
        cur.execute("""
            SELECT id, title FROM gsheet.spreadsheets
            WHERE LOWER(title) LIKE '%quiz%' AND LOWER(title) LIKE '%performance%'
        """)
        sheets = cur.fetchall()
        record("Google Sheet 'Quiz Performance Tracker' существует",
               len(sheets) >= 1,
               f"Найдено таблиц: {len(sheets)}")
        if not sheets:
            conn.close()
            return

        ss_id = sheets[0][0]

        # Лист 'Quiz Scores'
        cur.execute("""
            SELECT id, title FROM gsheet.sheets
            WHERE spreadsheet_id = %s AND LOWER(title) LIKE '%%quiz%%score%%'
        """, (ss_id,))
        quiz_sheets = cur.fetchall()
        record("Лист 'Quiz Scores' существует в таблице",
               len(quiz_sheets) >= 1,
               f"Найдено листов: {len(quiz_sheets)}")
        sheet_id = quiz_sheets[0][0] if quiz_sheets else None

        # Читаем ячейки целевого листа (либо всей таблицы как запасной вариант)
        if sheet_id is not None:
            cur.execute("""
                SELECT row_index, col_index, value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s
                ORDER BY row_index, col_index
            """, (ss_id, sheet_id))
        else:
            cur.execute("""
                SELECT row_index, col_index, value FROM gsheet.cells
                WHERE spreadsheet_id = %s
                ORDER BY row_index, col_index
            """, (ss_id,))
        cells = cur.fetchall()

        grid = {}
        for row_idx, col_idx, val in cells:
            grid.setdefault(row_idx, {})[col_idx] = val

        all_values = " ".join(str(v) for r in grid.values() for v in r.values() if v)
        record("GSheet содержит названия курсов",
               "биохими" in all_values.lower() or "биоинформат" in all_values.lower()
               or "креативн" in all_values.lower() or "вычислен" in all_values.lower(),
               f"Пример значений: {all_values[:200]}")
        record("GSheet содержит названия тестов (CMA)",
               "cma" in all_values.lower(),
               f"Пример значений: {all_values[:200]}")

        if not grid:
            record("Число строк-тестов в GSheet в ожидаемом диапазоне (~76)",
                   False, "Ячейки не найдены")
            conn.close()
            return

        min_row = min(grid.keys())
        header_row = grid.get(min_row, {})
        max_h = max(header_row.keys()) if header_row else -1
        header_vals = [str(header_row.get(i, "") or "").strip() for i in range(max_h + 1)]

        def find_col(name):
            for i, h in enumerate(header_vals):
                if h.lower() == name.lower():
                    return i
            return None

        course_col = find_col("Course_Name")
        quiz_col = find_col("Quiz_Title")
        avg_col = find_col("Avg_Score")
        max_col = find_col("Max_Score")
        pass_col = find_col("Pass_Rate")

        record("Колонка Course_Name присутствует", course_col is not None, f"Заголовки: {header_vals}")
        record("Колонка Quiz_Title присутствует", quiz_col is not None, f"Заголовки: {header_vals}")
        record("Колонка Avg_Score присутствует", avg_col is not None, f"Заголовки: {header_vals}")
        record("Колонка Pass_Rate присутствует", pass_col is not None, f"Заголовки: {header_vals}")

        # Строки данных
        data_rows = [grid[r] for r in sorted(grid.keys()) if r > min_row]
        row_count = len(data_rows)

        # CRITICAL: число строк в ожидаемом диапазоне
        record("Число строк-тестов в GSheet в ожидаемом диапазоне (~76)",
               abs(row_count - EXPECTED_ROWS) <= ROW_TOLERANCE,
               f"Найдено строк: {row_count} (ожидалось ~{EXPECTED_ROWS}±{ROW_TOLERANCE})")

        # Индексируем строки агента по (course, quiz) -> {avg, pass}
        agent_idx = {}
        if course_col is not None and quiz_col is not None:
            for row in data_rows:
                c = str(row.get(course_col, "") or "").strip()
                q = str(row.get(quiz_col, "") or "").strip()
                if not c or not q:
                    continue
                agent_idx[(c.lower(), q.lower())] = {
                    "avg": _to_float(row.get(avg_col)) if avg_col is not None else None,
                    "pass": _to_float(row.get(pass_col)) if pass_col is not None else None,
                }

        # CRITICAL: сверка Avg_Score с эталоном
        avg_matches = 0
        avg_checked = 0
        avg_detail = []
        for gt in gt_rows:
            key = (gt["course"].lower(), gt["quiz"].lower())
            if key not in agent_idx:
                continue
            avg_checked += 1
            a = agent_idx[key]["avg"]
            if a is not None and abs(a - gt["avg_score"]) <= SCORE_TOL:
                avg_matches += 1
            elif len(avg_detail) < 5:
                avg_detail.append(f"{gt['quiz']}: agent={a} gt={gt['avg_score']}")
        record("Средние баллы (Avg_Score) в GSheet совпадают с эталоном для выборки тестов",
               avg_matches >= min(MIN_GT_MATCH, len(gt_rows)),
               f"Совпало {avg_matches}/{avg_checked} (нужно >= {min(MIN_GT_MATCH, len(gt_rows))}); расхождения: {avg_detail}")

        # CRITICAL: сверка Pass_Rate с эталоном
        pass_matches = 0
        pass_checked = 0
        pass_detail = []
        for gt in gt_rows:
            key = (gt["course"].lower(), gt["quiz"].lower())
            if key not in agent_idx:
                continue
            pass_checked += 1
            p = agent_idx[key]["pass"]
            if p is not None and abs(p - gt["pass_rate"]) <= SCORE_TOL:
                pass_matches += 1
            elif len(pass_detail) < 5:
                pass_detail.append(f"{gt['quiz']}: agent={p} gt={gt['pass_rate']}")
        record("Pass_Rate в GSheet совпадает с эталоном для выборки тестов",
               pass_matches >= min(MIN_GT_MATCH, len(gt_rows)),
               f"Совпало {pass_matches}/{pass_checked} (нужно >= {min(MIN_GT_MATCH, len(gt_rows))}); расхождения: {pass_detail}")

        cur.close()
        conn.close()
    except Exception as e:
        record("Проверка GSheet", False, str(e))
        # Критические проверки, которых мы не достигли, помечаем как провал
        for name in ("Средние баллы (Avg_Score) в GSheet совпадают с эталоном для выборки тестов",
                     "Pass_Rate в GSheet совпадает с эталоном для выборки тестов",
                     "Число строк-тестов в GSheet в ожидаемом диапазоне (~76)"):
            if name not in FAILED_NAMES:
                record(name, False, "не достигнуто из-за ошибки GSheet")


def check_word(agent_workspace):
    print("\n=== Проверка документа Word ===")
    docx_path = os.path.join(agent_workspace, "Quiz_Performance_Summary.docx")
    if not os.path.isfile(docx_path):
        record("Quiz_Performance_Summary.docx существует", False, f"Не найдено: {docx_path}")
        record("Word-документ содержит реальную таблицу и раздел рекомендаций",
               False, "Файл отсутствует")
        return
    record("Quiz_Performance_Summary.docx существует", True)

    try:
        from docx import Document
        doc = Document(docx_path)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += " " + cell.text.lower()

        record("Word-документ содержит осмысленный контент (>= 100 символов)",
               len(all_text.strip()) >= 100,
               f"Длина контента: {len(all_text)}")
        record("Word-документ содержит контент про тесты",
               "quiz" in all_text or "score" in all_text or "performance" in all_text
               or "тест" in all_text or "балл" in all_text,
               f"Фрагмент: {all_text[:200]}")

        # CRITICAL: реальная таблица + раздел рекомендаций (RU+EN)
        has_table = len(doc.tables) >= 1
        has_recs = any(kw in all_text for kw in
                       ("recommend", "support", "intervention",
                        "рекомендац", "поддержк", "вмешательств"))
        record("Word-документ содержит реальную таблицу и раздел рекомендаций",
               has_table and has_recs,
               f"tables={len(doc.tables)}, рекомендации={has_recs}")
    except ImportError:
        record("Word-документ содержит контент", os.path.getsize(docx_path) > 1000,
               f"Размер: {os.path.getsize(docx_path)}")
        record("Word-документ содержит реальную таблицу и раздел рекомендаций",
               False, "python-docx недоступен — не удалось проверить таблицу")
    except Exception as e:
        record("Word-документ читается", False, str(e))
        record("Word-документ содержит реальную таблицу и раздел рекомендаций",
               False, str(e))


def check_email():
    print("\n=== Проверка письма ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT id, subject, to_addr, body_text
            FROM email.messages
            WHERE to_addr::text ILIKE '%%academic_coordinator@university.edu%%'
        """)
        emails = cur.fetchall()
        # Запасной матч по теме, если адрес записан иначе
        if not emails:
            cur.execute("""
                SELECT id, subject, to_addr, body_text
                FROM email.messages
                WHERE subject ILIKE '%%quiz%%performance%%'
                   OR subject ILIKE '%%quiz%%analysis%%'
            """)
            emails = cur.fetchall()

        record("Письмо отправлено координатору (academic_coordinator@university.edu)",
               len(emails) >= 1, "Подходящее письмо не найдено")

        # CRITICAL: тема про quiz + непустое тело
        ok = False
        detail = "Письмо не найдено"
        if emails:
            email = emails[0]
            subject = str(email[1]).lower() if email[1] else ""
            body = str(email[3]) if email[3] else ""
            ok = ("quiz" in subject) and (len(body) > 30)
            detail = f"Тема='{email[1]}', длина тела={len(body)}"
            record("Тема письма содержит 'quiz'", "quiz" in subject, f"Тема: {email[1]}")
            record("Тело письма непустое (> 30 символов)", len(body) > 30,
                   f"Длина тела: {len(body)}")
        record("Письмо отправлено координатору с темой про quiz и непустым телом",
               ok, detail)

        cur.close()
        conn.close()
    except Exception as e:
        record("Проверка письма", False, str(e))
        record("Письмо отправлено координатору с темой про quiz и непустым телом",
               False, str(e))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    gt_rows = load_groundtruth(args.groundtruth_workspace)
    print(f"[eval] Загружено эталонных строк: {len(gt_rows)}")

    check_gsheet(gt_rows)
    check_word(args.agent_workspace)
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\n=== Итого: {PASS_COUNT}/{total} проверок пройдено ({accuracy:.1f}%) ===")

    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_failed": critical_failed,
    }
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if critical_failed:
        print(f"CRITICAL FAILURES: {critical_failed}")
        print("FAIL (провалена критическая проверка)")
        sys.exit(1)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
