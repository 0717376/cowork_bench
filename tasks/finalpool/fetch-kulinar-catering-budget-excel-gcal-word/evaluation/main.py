"""Evaluation для fetch-kulinar-catering-budget-excel-gcal-word.

Структурные проверки (NON-critical): наличие файлов/листов/колонок, пороги
числа строк, наличие итоговой строки, события доставки в календаре.

CRITICAL-проверки (любой провал => немедленный FAIL до порога точности):
  C1. Цены в "Ingredient Costs" взяты из API поставщика (Unit_Price совпадает
      с ценой ингредиента в ingredients.json), и хотя бы для одной строки
      Line_Total == Quantity_Needed * Unit_Price * (1 - Discount_Pct/100).
  C2. Логика оптовых скидок применена корректно: ингредиент с суммарным
      Quantity_Needed >= 100 кг имеет Discount_Pct == 15; >= 50 (и < 100) => 10;
      иначе 0.
  C3. Масштабирование на 80 человек: хотя бы для одного рецепта
      Scaling_Factor == 80 / Default_Servings (в пределах допуска).
  C4. Сводный бюджет согласован: итоговая сумма (grand total) == сумме Meal_Cost
      по строкам, в пределах допуска.

Порог: точность >= 70% по NON-critical-проверкам И отсутствие critical-провалов.
"""
import json
import os
import sys
import tarfile
from argparse import ArgumentParser

import openpyxl
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
CRITICAL_FAILED = []


def record(name, passed, detail="", critical=False):
    global PASS_COUNT, FAIL_COUNT
    tag = "CRIT " if critical else ""
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {tag}{name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {tag}{name}{msg}")
        if critical:
            CRITICAL_FAILED.append(name)


# ---------------------------------------------------------------------------
# Supplier pricing fixture
# ---------------------------------------------------------------------------

def load_supplier_prices():
    """Канонические цены поставщика из mock_pages.tar.gz (RU-имена, RUB)."""
    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    candidates = [
        os.path.join(task_root, "tmp", "mock_pages", "api", "ingredients.json"),
    ]
    for c in candidates:
        if os.path.exists(c):
            with open(c, encoding="utf-8") as f:
                return json.load(f)
    tar_path = os.path.join(task_root, "files", "mock_pages.tar.gz")
    with tarfile.open(tar_path, "r:gz") as tar:
        member = next(m for m in tar.getmembers()
                      if m.name.endswith("api/ingredients.json"))
        return json.load(tar.extractfile(member))


def norm_name(s):
    return str(s).strip().lower().replace("ё", "е")


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------

def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def check_excel(agent_ws, supplier):
    print("\n=== Excel: Catering_Budget.xlsx ===")
    agent_file = os.path.join(agent_ws, "Catering_Budget.xlsx")
    if not os.path.exists(agent_file):
        record("Catering_Budget.xlsx существует", False, agent_file)
        # Невозможно проверить производные данные -> critical провал.
        record("Цены/скидки/масштаб согласованы с API", False,
               "файл отсутствует", critical=True)
        return
    record("Catering_Budget.xlsx существует", True)

    try:
        wb = openpyxl.load_workbook(agent_file, data_only=True)
    except Exception as e:
        record("Catering_Budget.xlsx читается", False, str(e))
        record("Цены/скидки/масштаб согласованы с API", False,
               "файл не читается", critical=True)
        return
    record("Catering_Budget.xlsx читается", True)

    # ---- Meal Plan ----
    mp = load_sheet_rows(wb, "Meal Plan")
    meal_data = []
    if mp is None:
        record("Лист 'Meal Plan' присутствует", False)
    else:
        record("Лист 'Meal Plan' присутствует", True)
        meal_data = mp[1:] if len(mp) > 1 else []
        record("Meal Plan: >= 9 строк", len(meal_data) >= 9,
               f"строк: {len(meal_data)}")
        days = set()
        meals = set()
        for row in meal_data:
            if row and row[0] is not None:
                try:
                    days.add(int(row[0]))
                except (TypeError, ValueError):
                    pass
            if row and len(row) > 1 and row[1]:
                meals.add(norm_name(row[1]))
        record("Meal Plan: покрыты 3 дня", len(days) >= 3, f"дни: {sorted(days)}")
        # Принимаем RU и EN названия приёмов пищи.
        has_b = any(("breakfast" in m or "завтрак" in m) for m in meals)
        has_l = any(("lunch" in m or "обед" in m) for m in meals)
        has_d = any(("dinner" in m or "ужин" in m) for m in meals)
        record("Meal Plan: есть завтрак/обед/ужин (RU/EN)",
               has_b and has_l and has_d,
               f"meals={sorted(meals)}")

    # ---- Ingredient Costs ----
    ic = load_sheet_rows(wb, "Ingredient Costs")
    cost_data = []
    if ic is None:
        record("Лист 'Ingredient Costs' присутствует", False)
    else:
        record("Лист 'Ingredient Costs' присутствует", True)
        cost_data = ic[1:] if len(ic) > 1 else []
        record("Ingredient Costs: >= 10 строк", len(cost_data) >= 10,
               f"строк: {len(cost_data)}")
        totals = []
        for row in cost_data:
            if row and len(row) > 6 and row[6] is not None:
                try:
                    totals.append(float(row[6]))
                except (TypeError, ValueError):
                    pass
        record("Ingredient Costs: >= 10 значений Line_Total",
               len(totals) >= 10, f"значений: {len(totals)}")
        record("Ingredient Costs: все Line_Total > 0",
               len(totals) > 0 and all(t > 0 for t in totals))

    # ---- Budget Summary ----
    bs = load_sheet_rows(wb, "Budget Summary")
    summary_data = []
    if bs is None:
        record("Лист 'Budget Summary' присутствует", False)
    else:
        record("Лист 'Budget Summary' присутствует", True)
        summary_data = bs[1:] if len(bs) > 1 else []
        record("Budget Summary: >= 9 строк", len(summary_data) >= 9,
               f"строк: {len(summary_data)}")
        has_total = any(
            cell and ("total" in str(cell).lower() or "итог" in str(cell).lower())
            for row in summary_data if row for cell in row
        )
        record("Budget Summary: есть итоговая строка (total/итог)", has_total)

    # ============================ CRITICAL ============================
    # Карта цен поставщика по нормализованным именам.
    price_map = {norm_name(i["name"]): float(i["price_per_kg"])
                 for i in supplier.get("ingredients", [])}

    # C1: цены из API + хотя бы одна строка с корректной формулой Line_Total.
    price_match = 0
    formula_ok = False
    rows_priced = 0
    for row in cost_data:
        if not row or len(row) < 7:
            continue
        # позиции: Recipe_Name(0), Ingredient(1), Quantity_Needed(2), Unit(3),
        #          Unit_Price(4), Discount_Pct(5), Line_Total(6)
        ing, qn, up, dp, lt = row[1], row[2], row[4], row[5], row[6]
        nm = norm_name(ing) if ing is not None else None
        if nm in price_map and up is not None:
            try:
                up_f = float(up)
                if abs(up_f - price_map[nm]) <= max(0.02 * price_map[nm], 0.5):
                    price_match += 1
            except (TypeError, ValueError):
                pass
        # формула Line_Total
        try:
            qn_f, up_f, dp_f, lt_f = float(qn), float(up), float(dp or 0), float(lt)
            rows_priced += 1
            expected = qn_f * up_f * (1 - dp_f / 100.0)
            if abs(lt_f - expected) <= max(0.01 * abs(expected), 1.0):
                formula_ok = True
        except (TypeError, ValueError):
            pass
    record("C1: Unit_Price из API поставщика и Line_Total по формуле",
           price_match >= 3 and formula_ok,
           f"совпадений цен={price_match}, формула_ок={formula_ok}, "
           f"строк_с_числами={rows_priced}",
           critical=True)

    # C2: оптовые скидки по суммарному весу (только строки в кг).
    weight = {}
    disc_by_name = {}
    for row in cost_data:
        if not row or len(row) < 7:
            continue
        ing, qn, unit, dp = row[1], row[2], row[3], row[5]
        if ing is None:
            continue
        nm = norm_name(ing)
        u = norm_name(unit) if unit is not None else ""
        if u in ("kg", "кг"):
            try:
                weight[nm] = weight.get(nm, 0.0) + float(qn)
            except (TypeError, ValueError):
                pass
        if dp is not None:
            try:
                disc_by_name.setdefault(nm, set()).add(round(float(dp)))
            except (TypeError, ValueError):
                pass

    disc_violations = []
    checked = 0
    for nm, w in weight.items():
        if nm not in disc_by_name:
            continue
        expected = 15 if w >= 100 else (10 if w >= 50 else 0)
        for d in disc_by_name[nm]:
            checked += 1
            if d != expected:
                disc_violations.append(f"{nm}: вес~{w:.1f}кг, скидка={d}, ждали {expected}")
    # Требуем, что проверяемые строки в кг есть и нет нарушений.
    record("C2: оптовые скидки соответствуют суммарному весу (50кг/10%, 100кг/15%)",
           checked >= 1 and not disc_violations,
           f"проверено={checked}, нарушения={disc_violations[:5]}",
           critical=True)

    # C3: масштабирование на 80 человек.
    scale_ok = False
    scale_detail = []
    for row in meal_data:
        if not row or len(row) < 5:
            continue
        default_servings, scaling = row[3], row[4]
        try:
            ds = float(default_servings)
            sf = float(scaling)
            if ds > 0:
                expected = 80.0 / ds
                scale_detail.append((ds, sf, round(expected, 3)))
                if abs(sf - expected) <= max(0.02 * expected, 0.05):
                    scale_ok = True
        except (TypeError, ValueError):
            pass
    record("C3: Scaling_Factor == 80 / Default_Servings (>=1 рецепт)",
           scale_ok, f"примеры (ds,sf,ожид)={scale_detail[:6]}",
           critical=True)

    # C4: согласованность сводного бюджета (grand total == сумме Meal_Cost).
    meal_costs = []
    grand = None
    for row in summary_data:
        if not row:
            continue
        # ищем total-строку
        row_is_total = any(
            cell and ("total" in str(cell).lower() or "итог" in str(cell).lower())
            for cell in row
        )
        # числовое значение Meal_Cost (колонка 3) или любое число в total-строке
        numeric_cells = []
        for cell in row:
            try:
                numeric_cells.append(float(cell))
            except (TypeError, ValueError):
                pass
        if row_is_total:
            if numeric_cells:
                grand = max(numeric_cells)  # обычно один большой итог
        else:
            if len(row) > 3 and row[3] is not None:
                try:
                    meal_costs.append(float(row[3]))
                except (TypeError, ValueError):
                    pass
    sum_meals = sum(meal_costs)
    c4_ok = (grand is not None and meal_costs and
             abs(grand - sum_meals) <= max(0.01 * sum_meals, 1.0))
    record("C4: итог Budget Summary == сумме Meal_Cost",
           c4_ok,
           f"grand={grand}, сумма_строк={round(sum_meals, 2)}, строк={len(meal_costs)}",
           critical=True)


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def check_word(agent_ws):
    print("\n=== Word: Catering_Proposal.docx ===")
    doc_path = os.path.join(agent_ws, "Catering_Proposal.docx")
    if not os.path.exists(doc_path):
        record("Catering_Proposal.docx существует", False, doc_path)
        return
    record("Catering_Proposal.docx существует", True)
    try:
        from docx import Document
        doc = Document(doc_path)
        full = "\n".join(p.text for p in doc.paragraphs)
        # Table cells too: answers laid out in a docx table are legitimate.
        full += "\n" + "\n".join(
            c.text for t in doc.tables for r in t.rows for c in r.cells)
        low = full.lower()
    except Exception as e:
        record("Catering_Proposal.docx читается", False, str(e))
        return

    record("Word: упомянуто 80 участников", "80" in full)
    record("Word: упомянут День 1 (RU/EN)",
           "day 1" in low or "day1" in low or "день 1" in low or "день1" in low)
    record("Word: есть итог стоимости (total/итог)",
           "total" in low or "итог" in low)
    record("Word: есть заголовок/обзор (catering/retreat/кейтеринг/выезд)",
           any(k in low for k in ("catering", "retreat", "кейтеринг", "выезд", "ретрит")))


# ---------------------------------------------------------------------------
# Calendar
# ---------------------------------------------------------------------------

def check_gcal():
    print("\n=== Calendar: события доставки ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute(
            """
            SELECT summary, start_datetime::date
              FROM gcal.events
             WHERE summary ILIKE '%доставк%'
                OR summary ILIKE '%кейтеринг%'
                OR summary ILIKE '%catering%delivery%'
                OR summary ILIKE '%day%delivery%'
             ORDER BY start_datetime
            """
        )
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Календарь доступен", False, str(e))
        return

    record("GCal: >= 3 события доставки", len(rows) >= 3,
           f"найдено: {[(r[0], str(r[1])) for r in rows]}")
    dates = [str(r[1]) for r in rows]
    record("GCal: событие на 2026-04-13", "2026-04-13" in dates)
    record("GCal: событие на 2026-04-14", "2026-04-14" in dates)
    record("GCal: событие на 2026-04-15", "2026-04-15" in dates)


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    supplier = load_supplier_prices()
    print(f"Загружено цен поставщика: {len(supplier.get('ingredients', []))} "
          f"(валюта {supplier.get('currency')})")

    check_excel(args.agent_workspace, supplier)
    check_word(args.agent_workspace)
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total else 0
    print(f"\nИтого: {PASS_COUNT}/{total} проверок пройдено ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total,
              "accuracy": accuracy, "critical_failed": CRITICAL_FAILED}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

    if CRITICAL_FAILED:
        print(f"FAIL: провалены критические проверки: {CRITICAL_FAILED}")
        sys.exit(1)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
