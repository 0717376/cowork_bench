"""
Оценка задачи yt-transcript-afrobeat-song-list-excel-email.

Проверки:
1. Afrobeat_Tracklist.xlsx существует в agent_workspace
2. Лист "Tracklist" содержит >= 8 строк данных
3. Лист "Tracklist" содержит столбцы Track_Number, Song_Title, Artist
4. Лист "Artist_Summary" существует и содержит >= 4 строки
5. Curator_Notes.docx существует с >= 3 заголовками
6. Curator_Notes.docx содержит музыкальные ключевые слова
7. Сверка значений с эталоном (groundtruth) для обоих листов
8. Страница в Teamly с "Afrobeat"/"Mix"/"Tracklist" в заголовке и реальными треками
9. Письмо отправлено на music@label.com с содержательной темой/телом

Критические проверки (CRITICAL_CHECKS): любой провал => общий FAIL независимо
от accuracy. Структурные проверки (наличие файла) — не критичны; критично
семантическое наполнение (реальные треки/исполнители из эталона).
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "cowork_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
FAILED_NAMES = []

# Семантические критические проверки. Корректный агент проходит их все;
# «недеятель» (пустой/заглушечный результат) проваливает хотя бы одну.
CRITICAL_CHECKS = {
    "Tracklist: >= 8 строк данных и столбцы Track_Number/Song_Title/Artist",
    "Tracklist: первая строка совпадает с эталоном (реальный трек)",
    "Teamly: страница трек-листа с реальными исполнителями из микса",
    "Email на music@label.com: тема и тело по теме трек-листа",
}


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        FAILED_NAMES.append(name)
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def check_excel(agent_workspace, groundtruth_workspace="."):
    print("\n=== Проверки 1-4, 7: Afrobeat_Tracklist.xlsx ===")
    xlsx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".xlsx") and ("afrobeat" in fname.lower() or "tracklist" in fname.lower()):
            xlsx_path = os.path.join(agent_workspace, fname)
            break

    record("Afrobeat_Tracklist.xlsx существует", xlsx_path is not None,
           f"Подходящий xlsx не найден в {agent_workspace}")

    if not xlsx_path:
        record("Tracklist: >= 8 строк данных и столбцы Track_Number/Song_Title/Artist", False, "xlsx не найден")
        record("Tracklist: первая строка совпадает с эталоном (реальный трек)", False, "xlsx не найден")
        record("Artist_Summary содержит >= 4 строки данных", False, "xlsx не найден")
        return

    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_path)

        # Лист Tracklist
        tracklist_sheet = None
        for name in wb.sheetnames:
            if "tracklist" in name.lower() or "track" in name.lower():
                tracklist_sheet = wb[name]
                break
        if tracklist_sheet is None and wb.sheetnames:
            tracklist_sheet = wb[wb.sheetnames[0]]

        rows_ok = False
        cols_ok = False
        if tracklist_sheet:
            data_rows = [r for r in tracklist_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            rows_ok = len(data_rows) >= 8

            headers = [str(c.value).strip() if c.value else "" for c in next(tracklist_sheet.iter_rows(max_row=1))]
            headers_lower = [h.lower() for h in headers]
            has_track_num = any(("track" in h and "num" in h) or h == "track_number" for h in headers_lower)
            has_song = any("song" in h or "title" in h for h in headers_lower)
            has_artist = any("artist" in h for h in headers_lower)
            cols_ok = has_track_num and has_song and has_artist
            print(f"  (info) строк данных={len(data_rows)}, заголовки={headers}")

        # CRITICAL: объединённая структурная+столбцовая проверка Tracklist
        record("Tracklist: >= 8 строк данных и столбцы Track_Number/Song_Title/Artist",
               rows_ok and cols_ok,
               f"rows_ok={rows_ok}, cols_ok={cols_ok}")

        # Лист Artist_Summary
        summary_sheet = None
        for name in wb.sheetnames:
            if "artist" in name.lower() or "summary" in name.lower():
                summary_sheet = wb[name]
                break

        record("Artist_Summary существует", summary_sheet is not None,
               f"Листы: {wb.sheetnames}")

        if summary_sheet:
            s_rows = [r for r in summary_sheet.iter_rows(min_row=2, values_only=True)
                      if any(c is not None for c in r)]
            record("Artist_Summary содержит >= 4 строки данных", len(s_rows) >= 4,
                   f"Найдено {len(s_rows)} строк")
        else:
            record("Artist_Summary содержит >= 4 строки данных", False, "Лист не найден")

        # --- Сверка значений с эталоном (groundtruth) ---
        first_row_ok = False
        gt_path = os.path.join(groundtruth_workspace, "Afrobeat_Tracklist.xlsx")
        if os.path.isfile(gt_path):
            gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
            for gt_sname in gt_wb.sheetnames:
                gt_ws = gt_wb[gt_sname]
                a_ws = None
                for asn in wb.sheetnames:
                    if asn.strip().lower() == gt_sname.strip().lower():
                        a_ws = wb[asn]
                        break
                if a_ws is None:
                    record(f"GT лист '{gt_sname}' есть в xlsx агента", False, f"Доступно: {wb.sheetnames}")
                    continue
                gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                a_rows = [r for r in a_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                record(f"GT '{gt_sname}' число строк", len(a_rows) == len(gt_rows),
                       f"Ожидалось {len(gt_rows)}, получено {len(a_rows)}")
                for ri in range(min(3, len(gt_rows))):
                    if ri >= len(a_rows):
                        break
                    ok = True
                    for ci in range(min(len(gt_rows[ri]), len(a_rows[ri]))):
                        gv, av = gt_rows[ri][ci], a_rows[ri][ci]
                        if gv is None:
                            continue
                        if isinstance(gv, (int, float)):
                            if not num_close(av, gv, max(abs(gv) * 0.1, 1.0)):
                                ok = False
                                break
                        else:
                            if not str_match(av, gv):
                                ok = False
                                break
                    record(f"GT '{gt_sname}' строка {ri+1} значения", ok,
                           f"gt={gt_rows[ri][:4]}, agent={a_rows[ri][:4] if ri < len(a_rows) else 'отсутствует'}")
                    # первая строка листа Tracklist — критический сигнал реального трека
                    if ri == 0 and gt_sname.strip().lower() == "tracklist":
                        first_row_ok = ok
            gt_wb.close()
        else:
            # без эталона критическую первую строку определяем по непустым song/artist
            first_row_ok = rows_ok and cols_ok

        # CRITICAL: первая строка трек-листа = реальный трек из эталона
        record("Tracklist: первая строка совпадает с эталоном (реальный трек)",
               first_row_ok, f"first_row_ok={first_row_ok}")

    except Exception as e:
        record("Tracklist: >= 8 строк данных и столбцы Track_Number/Song_Title/Artist", False, str(e))
        record("Tracklist: первая строка совпадает с эталоном (реальный трек)", False, str(e))
        record("Artist_Summary существует", False, str(e))
        record("Artist_Summary содержит >= 4 строки данных", False, str(e))


def check_word(agent_workspace):
    print("\n=== Проверки 5-6: Curator_Notes.docx ===")
    docx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".docx") and ("curator" in fname.lower() or "notes" in fname.lower()):
            docx_path = os.path.join(agent_workspace, fname)
            break
    if not docx_path:
        for fname in os.listdir(agent_workspace):
            if fname.lower().endswith(".docx"):
                docx_path = os.path.join(agent_workspace, fname)
                break

    record("Curator_Notes.docx существует", docx_path is not None,
           f"Подходящий docx не найден в {agent_workspace}")

    if not docx_path:
        record("Word: >= 3 заголовков", False, "docx не найден")
        record("Word: содержит музыкальные ключевые слова", False, "docx не найден")
        return

    try:
        from docx import Document
        doc = Document(docx_path)
        headings = [p for p in doc.paragraphs if p.style.name.lower().startswith("heading")]
        record("Word: >= 3 заголовков", len(headings) >= 3,
               f"Найдено {len(headings)} заголовков")

        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        keywords = ["track", "song", "artist", "afrobeat", "mix", "music"]
        found = [k for k in keywords if k in full_text]
        record("Word: содержит музыкальные ключевые слова", len(found) >= 3,
               f"Найденные ключевые слова: {found}")
    except Exception as e:
        record("Word: >= 3 заголовков", False, str(e))
        record("Word: содержит музыкальные ключевые слова", False, str(e))


def check_teamly(groundtruth_workspace="."):
    print("\n=== Проверка 8: Страница трек-листа в Teamly ===")
    # Реальные исполнители из эталонного xlsx — для содержательной проверки тела.
    gt_artists = []
    try:
        import openpyxl
        gt_path = os.path.join(groundtruth_workspace, "Afrobeat_Tracklist.xlsx")
        if os.path.isfile(gt_path):
            gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
            for sname in gt_wb.sheetnames:
                if sname.strip().lower() == "artist_summary":
                    for r in gt_wb[sname].iter_rows(min_row=2, values_only=True):
                        if r and r[0]:
                            gt_artists.append(str(r[0]).strip().lower())
            gt_wb.close()
    except Exception:
        pass
    # Запасной набор реальных исполнителей микса, если эталон недоступен.
    if not gt_artists:
        gt_artists = ["burna boy", "wizkid", "rema", "davido", "ckay", "asake", "ayra starr"]

    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        # Только страницы, созданные агентом (сидовые id <= 63).
        cur.execute("""
            SELECT id, title, body FROM teamly.pages
            WHERE id > 63
            ORDER BY id DESC
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Teamly: страница трек-листа с реальными исполнителями из микса", False, f"DB error: {e}")
        return

    # Заголовок по теме + тело с >= 2 реальными исполнителями.
    matched = None
    for pid, title, body in rows:
        title_l = (title or "").lower()
        if any(k in title_l for k in ["afrobeat", "mix", "tracklist", "трек-лист", "треклист"]):
            body_l = (body or "").lower()
            artist_hits = sum(1 for a in gt_artists if a in body_l)
            if artist_hits >= 2:
                matched = (pid, title, artist_hits)
                break

    record("Teamly: страница трек-листа с реальными исполнителями из микса",
           matched is not None,
           f"страниц агента: {len(rows)}, заголовки: {[r[1] for r in rows][:5]}")


def check_email():
    print("\n=== Проверка 9: Письмо на music@label.com ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages")
    messages = cur.fetchall()
    cur.close()
    conn.close()

    matching = None
    for subject, from_addr, to_addr, body_text in messages:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                to_str = " ".join(str(r).lower() for r in parsed) if isinstance(parsed, list) else to_addr.lower()
            except Exception:
                to_str = str(to_addr).lower()
        # Отправленное агентом письмо: исходящий адрес НЕ music@label.com
        # (это вшитый отправитель запроса), а music@label.com — в получателях.
        if "music@label.com" in to_str and "music@label.com" not in str(from_addr).lower():
            matching = (subject, from_addr, to_addr, body_text)
            break

    has_content = False
    if matching:
        subj = matching[0] or ""
        body = matching[3] or ""
        has_content = any(k in (subj + " " + body).lower()
                          for k in ["afrobeat", "tracklist", "track", "artist", "mix"])

    # CRITICAL: письмо отправлено агентом И по теме трек-листа.
    record("Email на music@label.com: тема и тело по теме трек-листа",
           matching is not None and has_content,
           f"всего сообщений: {len(messages)}, найдено отправленное={matching is not None}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    agent_ws = args.agent_workspace

    check_excel(agent_ws, args.groundtruth_workspace)
    check_word(agent_ws)
    check_teamly(args.groundtruth_workspace)
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: проверки не выполнялись.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nИтого: {PASS_COUNT}/{total} проверок пройдено ({accuracy:.1f}%)")

    critical_failed = [n for n in FAILED_NAMES if n in CRITICAL_CHECKS]
    if critical_failed:
        print(f"  КРИТИЧЕСКИЕ ПРОВАЛЫ: {len(critical_failed)}")
        for n in critical_failed:
            print(f"    - {n}")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "critical_failed": critical_failed,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if critical_failed:
        print("Overall: FAIL (провалена критическая проверка)")
        sys.exit(1)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
